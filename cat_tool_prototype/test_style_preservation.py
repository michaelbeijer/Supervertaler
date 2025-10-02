"""
Test script for style preservation on export
Tests Phase B: Ensure exported documents maintain original styles
"""

import os
from docx import Document
from docx_handler import DOCXHandler

def create_test_document():
    """Create a test document with various styles"""
    doc = Document()
    
    # Title
    para = doc.add_paragraph("Document Title")
    para.style = 'Title'
    
    # Subtitle
    para = doc.add_paragraph("Document Subtitle")
    para.style = 'Subtitle'
    
    # Heading 1
    para = doc.add_paragraph("Chapter 1: Introduction")
    para.style = 'Heading 1'
    
    # Normal paragraph
    para = doc.add_paragraph("This is a normal paragraph with regular text.")
    
    # Heading 2
    para = doc.add_paragraph("Section 1.1: Background")
    para.style = 'Heading 2'
    
    # Normal paragraph
    para = doc.add_paragraph("Another normal paragraph with some content to translate.")
    
    # Heading 3
    para = doc.add_paragraph("Subsection 1.1.1: Details")
    para.style = 'Heading 3'
    
    # Normal paragraph
    para = doc.add_paragraph("More content in a normal paragraph.")
    
    # Add a table with styled content
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Table header (first row)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Column A"
    header_cells[1].text = "Column B"
    
    # Table data (second row)
    data_cells = table.rows[1].cells
    data_cells[0].text = "Data in cell A"
    data_cells[1].text = "Data in cell B"
    
    # Save the document
    output_path = "test_style_preservation_input.docx"
    doc.save(output_path)
    print(f"Created test document: {output_path}")
    return output_path

def test_style_preservation():
    """Test that styles are preserved during export"""
    print("=" * 60)
    print("PHASE B: STYLE PRESERVATION ON EXPORT TEST")
    print("=" * 60)
    
    # Create test document
    input_path = create_test_document()
    
    # Initialize handler and import
    handler = DOCXHandler()
    segments = handler.import_docx(input_path)
    
    print(f"\n✅ Imported {len(segments)} segments")
    
    # Display segments with their styles
    print("\n📋 Segments and their styles:")
    print("-" * 60)
    for idx, text in enumerate(segments):
        # Get style from paragraphs_info
        para_info = handler.paragraphs_info[idx]
        style = para_info.style if para_info.style else 'Normal'
        source = text[:50] + "..." if len(text) > 50 else text
        is_table = " [TABLE]" if para_info.is_table_cell else ""
        print(f"  {idx:2d}. {style:20s} {source}{is_table}")
    
    # Create "translations" (just add prefix for testing)
    translated_segments = []
    for idx, text in enumerate(segments):
        translated_seg = {
            'paragraph_id': idx,  # Use index as paragraph_id
            'source': text,
            'target': f"[TRANSLATED] {text}"  # Simple test translation
        }
        translated_segments.append(translated_seg)
    
    # Export with style preservation
    output_path = "test_style_preservation_output.docx"
    handler.export_docx(translated_segments, output_path, preserve_formatting=True)
    
    print(f"\n✅ Exported to: {output_path}")
    
    # Verify the output by reading it back
    print("\n🔍 Verifying exported document...")
    verify_doc = Document(output_path)
    
    print("\n📄 Exported document structure:")
    print("-" * 60)
    para_count = 0
    for para in verify_doc.paragraphs:
        if para.text.strip():
            style = para.style.name
            text = para.text[:50] + "..." if len(para.text) > 50 else para.text
            print(f"  {para_count:2d}. {style:20s} {text}")
            para_count += 1
    
    # Check tables
    if verify_doc.tables:
        print(f"\n📊 Tables in exported document: {len(verify_doc.tables)}")
        for table_idx, table in enumerate(verify_doc.tables):
            print(f"  Table {table_idx + 1}:")
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    if cell.text.strip():
                        text = cell.text[:40] + "..." if len(cell.text) > 40 else cell.text
                        print(f"    R{row_idx+1}C{cell_idx+1}: {text}")
    
    print("\n" + "=" * 60)
    print("✅ PHASE B TEST COMPLETE")
    print("=" * 60)
    print(f"\nPlease open the following files to verify style preservation:")
    print(f"  Input:  {input_path}")
    print(f"  Output: {output_path}")
    print(f"\nCheck that:")
    print(f"  ✓ Title remains Title style")
    print(f"  ✓ Subtitle remains Subtitle style")
    print(f"  ✓ Heading 1 remains Heading 1 style")
    print(f"  ✓ Heading 2 remains Heading 2 style")
    print(f"  ✓ Heading 3 remains Heading 3 style")
    print(f"  ✓ Normal paragraphs remain Normal style")
    print(f"  ✓ Table content is preserved")

if __name__ == "__main__":
    test_style_preservation()
