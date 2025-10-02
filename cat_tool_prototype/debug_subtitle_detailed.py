"""Debug script with detailed import logging"""
from docx import Document
from docx_handler import DOCXHandler

# First, check what python-docx sees
print("=" * 80)
print("STEP 1: What python-docx sees in document.paragraphs")
print("=" * 80)
doc = Document('test_style_preservation_input.docx')

# Build table paragraph set
table_paragraph_ids = set()
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                table_paragraph_ids.add(id(para._element))

print(f"Table paragraphs to skip: {len(table_paragraph_ids)}")

for idx, para in enumerate(doc.paragraphs):
    is_table = id(para._element) in table_paragraph_ids
    is_empty = not para.text.strip()
    style = para.style.name
    text = para.text[:50]
    skip_reason = ""
    if is_table:
        skip_reason = " [SKIPPED: In table]"
    elif is_empty:
        skip_reason = " [SKIPPED: Empty]"
    else:
        skip_reason = " [INCLUDED]"
    
    print(f"{idx:2d}. {style:20s} {skip_reason:25s} {text}")

print("\n" + "=" * 80)
print("STEP 2: What DOCXHandler imports")
print("=" * 80)

handler = DOCXHandler()
segments = handler.import_docx('test_style_preservation_input.docx')

print(f"\nImported {len(segments)} segments:")
for i, seg in enumerate(segments):
    para_info = handler.paragraphs_info[i]
    print(f"{i:2d}. {para_info.style:20s} {seg[:50]}")
