"""Debug: Check why Subtitle is considered a table paragraph"""
from docx import Document

doc = Document('test_style_preservation_input.docx')

print("Checking table structure:")
print("=" * 80)

# Check each table
for table_idx, table in enumerate(doc.tables):
    print(f"\nTable {table_idx}:")
    print(f"  Rows: {len(table.rows)}")
    print(f"  Columns: {len(table.rows[0].cells) if table.rows else 0}")
    
    for row_idx, row in enumerate(table.rows):
        print(f"\n  Row {row_idx}:")
        for cell_idx, cell in enumerate(row.cells):
            print(f"    Cell {cell_idx}:")
            for para_idx, para in enumerate(cell.paragraphs):
                text = para.text[:40]
                style = para.style.name
                elem_id = id(para._element)
                print(f"      Para {para_idx}: {style:20s} {text} [ID: {elem_id}]")

print("\n" + "=" * 80)
print("Now checking document.paragraphs:")
print("=" * 80)

for idx, para in enumerate(doc.paragraphs):
    text = para.text[:40]
    style = para.style.name
    elem_id = id(para._element)
    print(f"{idx:2d}. {style:20s} {text:40s} [ID: {elem_id}]")
