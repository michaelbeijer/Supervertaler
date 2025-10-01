"""
DOCX Handler
Import and export DOCX files with formatting preservation
"""

import os
from typing import List, Dict, Any
from dataclasses import dataclass

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("ERROR: python-docx not installed. Run: pip install python-docx")

# Import tag manager for inline formatting
try:
    from tag_manager import TagManager
except ImportError:
    print("WARNING: tag_manager not found. Inline formatting will not be preserved.")
    TagManager = None


@dataclass
class ParagraphInfo:
    """Information about a paragraph for reconstruction"""
    text: str
    style: str = None
    alignment: str = None
    paragraph_index: int = 0


class DOCXHandler:
    """Handle DOCX import and export operations"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library is required. Install with: pip install python-docx")
        
        self.original_document = None
        self.original_path = None
        self.paragraphs_info: List[ParagraphInfo] = []
        self.tag_manager = TagManager() if TagManager else None
    
    def import_docx(self, file_path: str, extract_formatting: bool = True) -> List[str]:
        """
        Import DOCX file and extract paragraphs with formatting tags
        
        Args:
            file_path: Path to DOCX file
            extract_formatting: If True, convert formatting to inline tags
        
        Returns: List of paragraph texts (with tags if extract_formatting=True)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        print(f"[DOCX Handler] Importing: {file_path}")
        if extract_formatting and self.tag_manager:
            print("[DOCX Handler] Extracting inline formatting as tags")
        
        # Load document
        self.original_document = Document(file_path)
        self.original_path = file_path
        self.paragraphs_info = []
        
        paragraphs = []
        
        # Extract paragraphs with metadata
        for idx, para in enumerate(self.original_document.paragraphs):
            text = para.text.strip()
            
            if text:  # Only include non-empty paragraphs
                # Extract formatting if requested
                if extract_formatting and self.tag_manager:
                    runs = self.tag_manager.extract_runs(para)
                    text_with_tags = self.tag_manager.runs_to_tagged_text(runs)
                    paragraphs.append(text_with_tags)
                else:
                    paragraphs.append(text)
                
                # Store paragraph info for reconstruction
                para_info = ParagraphInfo(
                    text=text,
                    style=para.style.name if para.style else None,
                    alignment=str(para.alignment) if para.alignment else None,
                    paragraph_index=idx
                )
                self.paragraphs_info.append(para_info)
        
        print(f"[DOCX Handler] Extracted {len(paragraphs)} paragraphs")
        return paragraphs
    
    def export_docx(self, segments: List[Dict[str, Any]], output_path: str, 
                    preserve_formatting: bool = True):
        """
        Export translated segments back to DOCX
        
        Args:
            segments: List of segment dictionaries with 'paragraph_id', 'source', 'target'
            output_path: Path to save the translated document
            preserve_formatting: Whether to preserve original formatting (default True)
        """
        print(f"[DOCX Handler] Exporting to: {output_path}")
        
        if not self.original_document:
            raise ValueError("No original document loaded. Import a DOCX first.")
        
        # Create a new document based on the original
        if preserve_formatting and self.original_path:
            # Copy the original document structure
            doc = Document(self.original_path)
        else:
            # Create new blank document
            doc = Document()
        
        # Group segments by paragraph
        para_segments = {}
        for seg in segments:
            para_id = seg.get('paragraph_id', 0)
            if para_id not in para_segments:
                para_segments[para_id] = []
            para_segments[para_id].append(seg)
        
        # Track which paragraphs we've processed
        processed_paras = set()
        
        # Replace paragraph text with translations
        # We need to match non-empty paragraphs from import with segment paragraph IDs
        non_empty_para_index = 0
        for para in doc.paragraphs:
            # Only process non-empty paragraphs (same logic as import)
            if not para.text.strip():
                continue
                
            if non_empty_para_index in para_segments:
                # Combine all segments from this paragraph
                translations = [s['target'] for s in para_segments[non_empty_para_index] 
                              if s['target'].strip()]
                
                if translations:
                    # Join segments back into paragraph (single space, no extra newlines)
                    new_text = ' '.join(translations)
                    
                    # Replace text while preserving formatting
                    self._replace_paragraph_text(para, new_text)
                    processed_paras.add(non_empty_para_index)
                elif not any(s['target'].strip() for s in para_segments[non_empty_para_index]):
                    # No translations provided - keep original or clear it
                    # For now, let's keep the original text if nothing was translated
                    pass
            
            non_empty_para_index += 1
        
        # Save the document
        doc.save(output_path)
        print(f"[DOCX Handler] Export complete: {output_path}")
        print(f"[DOCX Handler] Translated {len(processed_paras)} paragraphs")
    
    def _replace_paragraph_text(self, paragraph, new_text: str):
        """
        Replace paragraph text while preserving or applying formatting
        
        If new_text contains inline tags (e.g., <b>text</b>), they will be
        converted to proper formatting runs.
        """
        # Check if text contains formatting tags
        if self.tag_manager and ('<b>' in new_text or '<i>' in new_text or '<u>' in new_text or '<bi>' in new_text):
            self._replace_paragraph_with_formatting(paragraph, new_text)
            return
        
        # Simple replacement (no tags) - preserve original formatting
        # Store original formatting from first run (if any)
        original_font_name = None
        original_font_size = None
        original_bold = False
        original_italic = False
        
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if first_run.font:
                original_font_name = first_run.font.name
                original_font_size = first_run.font.size
                original_bold = first_run.font.bold or False
                original_italic = first_run.font.italic or False
        
        # Clear paragraph - delete all runs except first
        while len(paragraph.runs) > 1:
            paragraph._element.remove(paragraph.runs[-1]._element)
        
        # If no runs exist, create one
        if not paragraph.runs:
            run = paragraph.add_run()
        else:
            run = paragraph.runs[0]
        
        # Set the new text (strip any trailing/leading whitespace to avoid extra newlines)
        run.text = new_text.strip()
        
        # Restore formatting
        if original_font_name:
            run.font.name = original_font_name
        if original_font_size:
            run.font.size = original_font_size
        if original_bold:
            run.font.bold = True
        if original_italic:
            run.font.italic = True
    
    def _replace_paragraph_with_formatting(self, paragraph, tagged_text: str):
        """
        Replace paragraph text with formatted runs based on inline tags
        
        Example: "Hello <b>world</b>!" creates runs with proper bold formatting
        """
        if not self.tag_manager:
            # Fallback: strip tags and use simple replacement
            clean_text = tagged_text.replace('<b>', '').replace('</b>', '')
            clean_text = clean_text.replace('<i>', '').replace('</i>', '')
            clean_text = clean_text.replace('<u>', '').replace('</u>', '')
            clean_text = clean_text.replace('<bi>', '').replace('</bi>', '')
            self._replace_paragraph_text(paragraph, clean_text)
            return
        
        # Store original font properties
        original_font_name = None
        original_font_size = None
        
        if paragraph.runs and paragraph.runs[0].font:
            first_run = paragraph.runs[0]
            original_font_name = first_run.font.name
            original_font_size = first_run.font.size
        
        # Clear all runs
        for run in paragraph.runs:
            paragraph._element.remove(run._element)
        
        # Convert tagged text to run specifications
        run_specs = self.tag_manager.tagged_text_to_runs(tagged_text)
        
        # Create runs with proper formatting
        for spec in run_specs:
            run = paragraph.add_run(spec['text'])
            
            # Apply formatting
            if spec.get('bold'):
                run.font.bold = True
            if spec.get('italic'):
                run.font.italic = True
            if spec.get('underline'):
                run.font.underline = True
            
            # Restore original font properties
            if original_font_name:
                run.font.name = original_font_name
            if original_font_size:
                run.font.size = original_font_size
    
    def export_bilingual_docx(self, segments: List[Dict[str, Any]], output_path: str):
        """
        Export as bilingual document (source | target in table)
        Useful for review purposes
        """
        print(f"[DOCX Handler] Exporting bilingual document: {output_path}")
        
        doc = Document()
        doc.add_heading('Bilingual Translation Document', 0)
        
        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = '#'
        header_cells[1].text = 'Source'
        header_cells[2].text = 'Target'
        
        # Add segments
        for seg in segments:
            row_cells = table.add_row().cells
            row_cells[0].text = str(seg.get('id', ''))
            row_cells[1].text = seg.get('source', '')
            row_cells[2].text = seg.get('target', '')
        
        doc.save(output_path)
        print(f"[DOCX Handler] Bilingual export complete")
    
    def get_document_info(self) -> Dict[str, Any]:
        """Get information about the loaded document"""
        if not self.original_document:
            return {}
        
        return {
            'paragraphs': len(self.original_document.paragraphs),
            'sections': len(self.original_document.sections),
            'tables': len(self.original_document.tables),
            'path': self.original_path
        }


# Quick test
if __name__ == "__main__":
    print("DOCX Handler Test")
    print("To test, you need a sample DOCX file.")
    
    if DOCX_AVAILABLE:
        print("✓ python-docx is installed")
    else:
        print("✗ python-docx is NOT installed")
        print("  Run: pip install python-docx")
