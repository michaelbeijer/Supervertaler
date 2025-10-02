"""
Test script for style support in CAT Editor v0.3.1
Creates a document with various heading styles and tests display
"""

from docx import Document
from docx.shared import Pt
import os

def create_test_document_with_styles():
    """Create a test DOCX document with various styles"""
    doc = Document()
    
    # Title style
    doc.add_heading('Software Development Agreement', level=0)  # Title
    
    # Add subtitle (using paragraph with Subtitle style)
    subtitle = doc.add_paragraph('Between Company A and Company B')
    subtitle.style = 'Subtitle'
    
    # Heading 1
    doc.add_heading('1. Introduction', level=1)
    
    # Normal paragraph
    doc.add_paragraph(
        'This Software Development Agreement ("Agreement") is entered into as of '
        'January 1, 2025, by and between Party A and Party B.'
    )
    
    # Another normal paragraph
    doc.add_paragraph(
        'The purpose of this Agreement is to define the terms and conditions under '
        'which the Developer will provide software development services to the Client.'
    )
    
    # Heading 2
    doc.add_heading('1.1 Definitions', level=2)
    
    # Normal paragraph
    doc.add_paragraph(
        'For the purposes of this Agreement, the following terms shall have the meanings set forth below:'
    )
    
    # Heading 3
    doc.add_heading('1.1.1 Software', level=3)
    
    # Normal paragraph
    doc.add_paragraph(
        '"Software" means the computer programs, applications, and related documentation '
        'to be developed under this Agreement.'
    )
    
    # Heading 3
    doc.add_heading('1.1.2 Services', level=3)
    
    # Normal paragraph
    doc.add_paragraph(
        '"Services" means the software development, testing, and implementation services '
        'to be provided by the Developer.'
    )
    
    # Heading 2
    doc.add_heading('1.2 Scope of Work', level=2)
    
    # Normal paragraph
    doc.add_paragraph(
        'The Developer agrees to develop and deliver the Software in accordance with '
        'the specifications outlined in Exhibit A attached hereto.'
    )
    
    # Heading 1
    doc.add_heading('2. Project Details', level=1)
    
    # Add table with project information
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Fill table
    cells = table.rows[0].cells
    cells[0].text = 'Project Name'
    cells[1].text = 'Enterprise Resource Planning System'
    
    cells = table.rows[1].cells
    cells[0].text = 'Duration'
    cells[1].text = '12 months'
    
    cells = table.rows[2].cells
    cells[0].text = 'Budget'
    cells[1].text = '$500,000'
    
    cells = table.rows[3].cells
    cells[0].text = 'Start Date'
    cells[1].text = 'February 1, 2025'
    
    # Heading 1
    doc.add_heading('3. Payment Terms', level=1)
    
    # Normal paragraph
    doc.add_paragraph(
        'The Client agrees to pay the Developer according to the following schedule:'
    )
    
    # Heading 2
    doc.add_heading('3.1 Payment Schedule', level=2)
    
    # Add payment table
    table2 = doc.add_table(rows=4, cols=3)
    table2.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table2.rows[0].cells
    header_cells[0].text = 'Milestone'
    header_cells[1].text = 'Percentage'
    header_cells[2].text = 'Amount'
    
    # Data rows
    data = [
        ('Project Initiation', '25%', '$125,000'),
        ('Development Phase', '50%', '$250,000'),
        ('Final Delivery', '25%', '$125,000')
    ]
    
    for i, (milestone, percentage, amount) in enumerate(data, 1):
        cells = table2.rows[i].cells
        cells[0].text = milestone
        cells[1].text = percentage
        cells[2].text = amount
    
    # Heading 2
    doc.add_heading('3.2 Payment Terms', level=2)
    
    # Normal paragraph
    doc.add_paragraph(
        'Payment shall be made within 30 days of receipt of invoice. '
        'Late payments will incur interest at a rate of 1.5% per month.'
    )
    
    # Heading 1
    doc.add_heading('4. Confidentiality', level=1)
    
    # Normal paragraph with formatting
    para = doc.add_paragraph()
    para.add_run('Both parties agree to maintain the ').font.size = Pt(11)
    conf_run = para.add_run('confidentiality')
    conf_run.bold = True
    conf_run.font.size = Pt(11)
    para.add_run(' of all proprietary information disclosed during the term of this Agreement.').font.size = Pt(11)
    
    # Quote style
    quote = doc.add_paragraph(
        '"Confidential Information" includes all technical and business information, '
        'source code, designs, and any other materials disclosed by either party.'
    )
    quote.style = 'Intense Quote'
    
    # Heading 1
    doc.add_heading('5. Termination', level=1)
    
    # Normal paragraph
    doc.add_paragraph(
        'Either party may terminate this Agreement upon 30 days written notice to the other party.'
    )
    
    # Heading 2
    doc.add_heading('5.1 Effect of Termination', level=2)
    
    # Normal paragraph
    doc.add_paragraph(
        'Upon termination, all work in progress shall be delivered to the Client, '
        'and final payment shall be calculated on a pro-rata basis.'
    )
    
    # Save the document
    output_path = os.path.join(os.path.dirname(__file__), 'test_document_with_styles.docx')
    doc.save(output_path)
    print(f"✓ Created test document: {output_path}")
    return output_path


def test_style_import():
    """Test importing the document with style support"""
    from docx_handler import DOCXHandler
    from simple_segmenter import SimpleSegmenter
    
    # Create test document
    doc_path = create_test_document_with_styles()
    
    # Import it
    handler = DOCXHandler()
    paragraphs = handler.import_docx(doc_path)
    
    print(f"\n=== Import Results ===")
    print(f"Total items extracted: {len(paragraphs)}")
    
    # Show document info
    info = handler.get_document_info()
    print(f"\nDocument Information:")
    print(f"  - Regular paragraphs: {info['regular_paragraphs']}")
    print(f"  - Table cells: {info['table_cells']}")
    print(f"  - Tables: {info['tables']}")
    print(f"  - Total items: {info['total_items']}")
    
    # Show paragraph info with styles
    print(f"\n=== Extracted Items with Styles ===")
    style_counts = {}
    for i, para_info in enumerate(handler.paragraphs_info):
        style = para_info.style or "Normal"
        style_counts[style] = style_counts.get(style, 0) + 1
        
        if para_info.is_table_cell:
            print(f"{i+1:2d}. [TABLE T{para_info.table_index+1}R{para_info.row_index+1}C{para_info.cell_index+1}] [{style:15s}] {para_info.text[:40]}...")
        else:
            print(f"{i+1:2d}. [PARA] [{style:15s}] {para_info.text[:50]}...")
    
    # Show style statistics
    print(f"\n=== Style Statistics ===")
    for style, count in sorted(style_counts.items()):
        print(f"  {style:20s}: {count:2d} segments")
    
    return doc_path, handler


if __name__ == "__main__":
    print("=" * 70)
    print("Style Support Test for CAT Editor v0.3.1")
    print("=" * 70)
    
    try:
        doc_path, handler = test_style_import()
        print("\n✓ Test completed successfully!")
        print(f"\nYou can now open '{os.path.basename(doc_path)}' in the CAT Editor")
        print("to see headings displayed with proper styling:")
        print("  - Heading 1: Bold, dark blue")
        print("  - Heading 2: Bold, medium blue")
        print("  - Heading 3: Bold, light blue")
        print("  - Title: Bold, larger, purple")
        print("  - Normal: Regular text")
        
    except Exception as e:
        print(f"\n✗ Test failed: {str(e)}")
        import traceback
        traceback.print_exc()
