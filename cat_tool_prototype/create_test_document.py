"""
Test script to create a sample DOCX with formatted text
This can be used to test the inline formatting tag feature
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    print("ERROR: python-docx not installed")
    print("Run: pip install python-docx")
    exit(1)

def create_test_document():
    """Create a test document with various formatting"""
    
    doc = Document()
    doc.add_heading('Test Document with Formatting', 0)
    
    # Paragraph 1: Bold text
    p1 = doc.add_paragraph()
    p1.add_run('This document contains ')
    p1.add_run('bold').bold = True
    p1.add_run(' text for testing.')
    
    # Paragraph 2: Italic text
    p2 = doc.add_paragraph()
    p2.add_run('It also has ')
    p2.add_run('italic').italic = True
    p2.add_run(' and ')
    p2.add_run('underlined').underline = True
    p2.add_run(' words.')
    
    # Paragraph 3: Multiple bold words
    p3 = doc.add_paragraph()
    p3.add_run('The ')
    p3.add_run('API key').bold = True
    p3.add_run(' is required for ')
    p3.add_run('authentication').bold = True
    p3.add_run('.')
    
    # Paragraph 4: Bold and italic combined
    p4 = doc.add_paragraph()
    p4.add_run('This is ')
    run = p4.add_run('very important')
    run.bold = True
    run.italic = True
    p4.add_run(' information.')
    
    # Paragraph 5: Complex formatting
    p5 = doc.add_paragraph()
    p5.add_run('The ')
    p5.add_run('User-Agent').bold = True
    p5.add_run(' header should include the ')
    p5.add_run('application name').italic = True
    p5.add_run(' and ')
    p5.add_run('version number').italic = True
    p5.add_run('.')
    
    # Paragraph 6: Technical text
    p6 = doc.add_paragraph()
    p6.add_run('To install the package, run: ')
    p6.add_run('pip install package-name').bold = True
    p6.add_run(' in your terminal.')
    
    # Paragraph 7: Plain text (no formatting)
    p7 = doc.add_paragraph()
    p7.add_run('This paragraph has no special formatting at all.')
    
    # Paragraph 8: Multiple formats in sequence
    p8 = doc.add_paragraph()
    p8.add_run('Format examples: ')
    p8.add_run('bold').bold = True
    p8.add_run(', ')
    p8.add_run('italic').italic = True
    p8.add_run(', ')
    p8.add_run('underline').underline = True
    p8.add_run(', and ')
    run = p8.add_run('all three')
    run.bold = True
    run.italic = True
    run.underline = True
    p8.add_run('.')
    
    # Save
    filename = 'test_document_with_formatting.docx'
    doc.save(filename)
    print(f"✓ Created test document: {filename}")
    print(f"\nThis document contains:")
    print("  - Bold text")
    print("  - Italic text")
    print("  - Underlined text")
    print("  - Bold + Italic combination")
    print("  - Complex mixed formatting")
    print("  - Plain text (no formatting)")
    print(f"\nImport this file into the CAT Editor to see inline tags!")

if __name__ == "__main__":
    create_test_document()
