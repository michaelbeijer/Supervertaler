"""
Test script for table support in CAT Editor
Creates a sample DOCX with tables and tests the import/export
"""

from docx import Document
from docx.shared import Pt, RGBColor
import os

def create_test_document():
    """Create a test DOCX document with tables"""
    doc = Document()
    
    # Add title
    doc.add_heading('Sample Contract Document', 0)
    
    # Add intro paragraph
    para = doc.add_paragraph()
    para.add_run('This is an introduction paragraph with ').font.size = Pt(12)
    bold_run = para.add_run('bold text')
    bold_run.bold = True
    bold_run.font.size = Pt(12)
    para.add_run(' and ').font.size = Pt(12)
    italic_run = para.add_run('italic text')
    italic_run.italic = True
    italic_run.font.size = Pt(12)
    para.add_run('.').font.size = Pt(12)
    
    # Add another paragraph
    doc.add_paragraph('This document contains both regular paragraphs and tables to demonstrate table cell segmentation.')
    
    # Add a heading
    doc.add_heading('Party Information', level=1)
    
    # Add table 1: Party details
    table1 = doc.add_table(rows=3, cols=2)
    table1.style = 'Light Grid Accent 1'
    
    # Fill table 1
    cells = table1.rows[0].cells
    cells[0].text = 'Party A'
    cells[1].text = 'XYZ Corporation, Inc.'
    
    cells = table1.rows[1].cells
    cells[0].text = 'Party B'
    cells[1].text = 'ABC Industries Ltd.'
    
    cells = table1.rows[2].cells
    cells[0].text = 'Effective Date'
    cells[1].text = 'January 1, 2025'
    
    # Add paragraph between tables
    doc.add_paragraph('The terms and conditions of this agreement are outlined below.')
    
    # Add heading
    doc.add_heading('Payment Terms', level=1)
    
    # Add table 2: Payment schedule
    table2 = doc.add_table(rows=4, cols=3)
    table2.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table2.rows[0].cells
    header_cells[0].text = 'Installment'
    header_cells[1].text = 'Amount'
    header_cells[2].text = 'Due Date'
    
    # Data rows
    data = [
        ('First Payment', '$10,000', 'March 1, 2025'),
        ('Second Payment', '$15,000', 'June 1, 2025'),
        ('Final Payment', '$25,000', 'September 1, 2025')
    ]
    
    for i, (installment, amount, due_date) in enumerate(data, 1):
        cells = table2.rows[i].cells
        cells[0].text = installment
        cells[1].text = amount
        cells[2].text = due_date
    
    # Add closing paragraph
    doc.add_paragraph('This agreement constitutes the entire understanding between the parties.')
    
    # Add another paragraph with formatting
    closing = doc.add_paragraph()
    closing.add_run('Signed on this date by ').font.size = Pt(11)
    auth_run = closing.add_run('authorized representatives')
    auth_run.bold = True
    auth_run.underline = True
    auth_run.font.size = Pt(11)
    closing.add_run(' of both parties.').font.size = Pt(11)
    
    # Save the document
    output_path = os.path.join(os.path.dirname(__file__), 'test_document_with_tables.docx')
    doc.save(output_path)
    print(f"✓ Created test document: {output_path}")
    return output_path


def test_import():
    """Test importing the document with table support"""
    from docx_handler import DOCXHandler
    from simple_segmenter import SimpleSegmenter
    
    # Create test document
    doc_path = create_test_document()
    
    # Import it
    handler = DOCXHandler()
    paragraphs = handler.import_docx(doc_path)
    
    print(f"\n=== Import Results ===")
    print(f"Total items extracted: {len(paragraphs)}")
    
    # Show document info
    info = handler.get_document_info()
    print(f"\nDocument Information:")
    print(f"  - Regular paragraphs: {info['regular_paragraphs']}")
    print(f"  - Table cells: {info['table_cells']}")
    print(f"  - Tables: {info['tables']}")
    print(f"  - Total items: {info['total_items']}")
    
    # Show paragraph info breakdown
    print(f"\n=== Extracted Items ===")
    for i, para_info in enumerate(handler.paragraphs_info):
        if para_info.is_table_cell:
            print(f"{i+1}. [TABLE T{para_info.table_index+1}R{para_info.row_index+1}C{para_info.cell_index+1}] {para_info.text[:50]}...")
        else:
            print(f"{i+1}. [PARA] {para_info.text[:50]}...")
    
    # Test segmentation
    print(f"\n=== Segmentation Test ===")
    segmenter = SimpleSegmenter()
    segments = segmenter.segment_paragraphs(paragraphs)
    print(f"Total segments: {len(segments)}")
    
    # Show first few segments
    print("\nFirst 5 segments:")
    for i, (para_id, text) in enumerate(segments[:5], 1):
        para_info = handler._get_para_info(para_id)
        if para_info and para_info.is_table_cell:
            print(f"  {i}. [TABLE] {text[:60]}...")
        else:
            print(f"  {i}. [PARA] {text[:60]}...")
    
    return doc_path, handler, segments


if __name__ == "__main__":
    print("=" * 70)
    print("Table Support Test for CAT Editor")
    print("=" * 70)
    
    try:
        doc_path, handler, segments = test_import()
        print("\n✓ Test completed successfully!")
        print(f"\nYou can now open '{os.path.basename(doc_path)}' in the CAT Editor")
        print("to see table cells displayed with proper segmentation.")
        
    except Exception as e:
        print(f"\n✗ Test failed: {str(e)}")
        import traceback
        traceback.print_exc()
