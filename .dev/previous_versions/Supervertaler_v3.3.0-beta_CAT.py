"""  
Supervertaler v3.3.0-beta (CAT Editor)
AI-Powered Computer-Aided Translation Tool

Features:
- Grid Pagination System (50 segments/page, 10x faster loading) ⚡
- Smart Paragraph Detection for document view 🧠
- Unified Prompt Library (System Prompts + Custom Instructions) 🎯
- LLM Translation (OpenAI GPT-4, Anthropic Claude, Google Gemini)
- Custom Prompts with variable substitution
- Translation Memory with fuzzy matching
- Dual Selection in Grid (memoQ-style)
- Import/Export DOCX, TSV, JSON
- CafeTran & memoQ bilingual DOCX support
- Auto-export options (Session reports, TMX, TSV, XLIFF, Excel)
- Sentence segmentation
- Editable grid interface with inline editing
- Find/Replace functionality
- Status tracking and progress monitoring
- Project save/load with context preservation
- Dev mode with parallel folder structure (user data/ vs user data_private/)

Author: Michael Beijer + AI Assistant
Date: October 12, 2025
"""

# Version constant
APP_VERSION = "3.3.0-beta"# --- Private Features Flag ---
# Check for .supervertaler.local file to enable private features (for developers only)
# Users won't have this file, so they won't see confusing private folder options
import os as _os_temp
ENABLE_PRIVATE_FEATURES = _os_temp.path.exists(
    _os_temp.path.join(_os_temp.path.dirname(_os_temp.path.abspath(__file__)), ".supervertaler.local")
)
if ENABLE_PRIVATE_FEATURES:
    print("[DEV MODE] Private features enabled (.supervertaler.local found)")
del _os_temp

def get_user_data_path(folder_name):
    """
    Get the appropriate user data path based on dev mode.
    
    In dev mode (with .supervertaler.local file):
        Returns: user data_private/<folder_name>
    In user mode:
        Returns: user data/<folder_name>
    
    Examples:
        Dev mode:  get_user_data_path('TMs') -> 'user data_private/TMs'
        User mode: get_user_data_path('TMs') -> 'user data/TMs'
    """
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    if ENABLE_PRIVATE_FEATURES:
        # Dev mode: use parallel private structure
        return os.path.join(script_dir, "user data_private", folder_name)
    else:
        # User mode: use standard structure
        return os.path.join(script_dir, "user data", folder_name)

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import json
import os
import re
import zipfile
import io
import base64
import shutil
import queue
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
from difflib import SequenceMatcher
import xml.etree.ElementTree as ET

# LLM API imports (optional - graceful degradation if not installed)
try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
    ANTHROPIC_VERSION = getattr(anthropic, '__version__', 'unknown')
except ImportError as e:
    ANTHROPIC_AVAILABLE = False
    ANTHROPIC_VERSION = "not installed"
    print(f"Note: anthropic library not available: {e}")

try:
    import openai
    OPENAI_AVAILABLE = True
    OPENAI_VERSION = getattr(openai, '__version__', 'unknown')
except ImportError as e:
    OPENAI_AVAILABLE = False
    OPENAI_VERSION = "not installed"
    print(f"Note: openai library not available: {e}")

try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
    GEMINI_VERSION = "available"
except ImportError as e:
    GEMINI_AVAILABLE = False
    GEMINI_VERSION = "not installed"
    print(f"Note: google-generativeai library not available: {e}")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError as e:
    PIL_AVAILABLE = False
    print(f"Note: PIL/Pillow library not available (image support disabled): {e}")

# Import our custom modules
try:
    from modules.simple_segmenter import SimpleSegmenter
    from modules.docx_handler import DOCXHandler
    from modules.tag_manager import TagManager
except ImportError:
    print("ERROR: Could not import required modules")
    print("Make sure the 'modules' folder exists with simple_segmenter.py, docx_handler.py, and tag_manager.py")
    import sys
    sys.exit(1)


# --- API Key Loading ---
def load_api_keys():
    """Load API keys from api_keys.txt file in the same directory as the script"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    api_keys_file = os.path.join(script_dir, "api_keys.txt")
    
    api_keys = {
        "google": "",
        "claude": "",
        "openai": ""
    }
    
    if os.path.exists(api_keys_file):
        try:
            with open(api_keys_file, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith('#') and '=' in line:
                        key, value = line.split('=', 1)
                        key = key.strip().lower()
                        value = value.strip()
                        if key in ["google", "google_api_key", "gemini"]:
                            api_keys["google"] = value
                        elif key in ["claude", "claude_api_key", "anthropic"]:
                            api_keys["claude"] = value
                        elif key in ["openai", "openai_api_key", "chatgpt"]:
                            api_keys["openai"] = value
        except Exception as e:
            print(f"Error reading api_keys.txt: {e}")
    else:
        # Create template file
        try:
            with open(api_keys_file, 'w', encoding='utf-8') as f:
                f.write("# API Keys Configuration\n")
                f.write("# Format: key_name = your_api_key_here\n")
                f.write("# Remove the # at the beginning of the line to uncomment\n\n")
                f.write("# Google API Key for Gemini models\n")
                f.write("#google = YOUR_GOOGLE_API_KEY_HERE\n\n")
                f.write("# Claude API Key for Anthropic models\n")
                f.write("#claude = YOUR_CLAUDE_API_KEY_HERE\n\n")
                f.write("# OpenAI API Key for ChatGPT models\n")
                f.write("#openai = YOUR_OPENAI_API_KEY_HERE\n")
            print(f"Created template api_keys.txt file at: {api_keys_file}")
        except Exception as e:
            print(f"Could not create api_keys.txt template: {e}")
    
    return api_keys


# Load API keys at module level
API_KEYS = load_api_keys()


# --- DOCX Track Changes Parsing Utilities ---
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def tag(name: str) -> str:
    """Create fully qualified XML tag name"""
    return f"{{{W_NS}}}{name}"

def collect_text(node, mode: str):
    """
    Recursively collect visible text from a node.
    mode='original' -> exclude insertions (<w:ins>), include deletions (<w:del> and <w:delText>)
    mode='final'    -> include insertions, exclude deletions
    """
    parts = []
    t = node.tag

    if t == tag("ins"):
        if mode == "final":
            for child in node:
                parts.extend(collect_text(child, mode))
        return parts  # in 'original' we ignore insertions

    if t == tag("del"):
        if mode == "original":
            for child in node:
                parts.extend(collect_text(child, mode))
        return parts  # in 'final' we ignore deletions

    # Runs and their children
    if t == tag("r"):
        for child in node:
            if child.tag == tag("t"):
                parts.append(child.text or "")
            elif child.tag == tag("delText"):
                if mode == "original":
                    parts.append(child.text or "")
            elif child.tag == tag("tab"):
                parts.append("\t")
            elif child.tag == tag("br"):
                parts.append("\n")
            else:
                parts.extend(collect_text(child, mode))
        return parts

    # Plain text nodes
    if t == tag("t"):
        parts.append(node.text or "")
        return parts
    if t == tag("delText"):
        if mode == "original":
            parts.append(node.text or "")
        return parts
    if t == tag("tab"):
        parts.append("\t")
        return parts
    if t == tag("br"):
        parts.append("\n")
        return parts

    # Generic recursion
    for child in node:
        parts.extend(collect_text(child, mode))
    return parts

def tidy_text(s: str) -> str:
    """Clean up text by collapsing whitespace"""
    s = re.sub(r"[ \t]+\n", "\n", s)
    s = re.sub(r"\n+", "\n", s)
    return s.strip()

def parse_docx_pairs(docx_path):
    """
    Extract (original_text, final_text) pairs from DOCX file with track changes.
    Returns list of tuples where text actually changed.
    """
    try:
        with zipfile.ZipFile(docx_path) as z:
            try:
                xml = z.read("word/document.xml").decode("utf-8")
            except KeyError:
                raise RuntimeError("This file does not contain word/document.xml; is it a valid .docx?")

        root = ET.fromstring(xml)
        ns = {"w": W_NS}

        rows = []
        for p in root.findall(".//w:p", ns):
            original = "".join(collect_text(p, "original"))
            final = "".join(collect_text(p, "final"))
            o_clean = tidy_text(original)
            f_clean = tidy_text(final)
            if o_clean != f_clean:
                rows.append((o_clean, f_clean))
        return rows
    except Exception as e:
        raise RuntimeError(f"Error parsing DOCX file: {e}")

def format_tracked_changes_context(tracked_changes_list, max_length=1000):
    """Format tracked changes for AI context, keeping within token limits"""
    if not tracked_changes_list:
        return ""
    
    context_parts = ["TRACKED CHANGES REFERENCE (Original→Final editing patterns):"]
    current_length = len(context_parts[0])
    
    for i, (original, final) in enumerate(tracked_changes_list):
        change_text = f"• \"{original}\" → \"{final}\""
        if current_length + len(change_text) > max_length:
            if i > 0:  # Only add if we have at least one example
                context_parts.append("(Additional examples truncated to save space)")
            break
        context_parts.append(change_text)
        current_length += len(change_text)
    
    return "\n".join(context_parts) + "\n"

def pil_image_to_base64_png(img):
    """Encode a PIL image to base64 PNG (ascii) for Claude/OpenAI data URLs."""
    try:
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return base64.b64encode(buf.getvalue()).decode("ascii")
    except Exception:
        return None

def normalize_figure_ref(ref_text):
    """Normalize figure references to match drawing filenames"""
    if not ref_text:
        return None
    match = re.search(r"(?:figure|figuur|fig\.?)\s*([\w\d]+(?:[\s\.\-]*[\w\d]+)?)", ref_text, re.IGNORECASE)
    if match:
        identifier = match.group(1)
        return re.sub(r"[\s\.\-]", "", identifier).lower()
    base_name = os.path.splitext(ref_text)[0]
    cleaned_base = re.sub(r"(?:figure|figuur|fig\.?)\s*", "", base_name, flags=re.IGNORECASE)
    normalized = re.sub(r"[\s\.\-]", "", cleaned_base).lower()
    if normalized:
        return normalized
    return None

def get_simple_lang_code(lang_name_or_code_input):
    """Convert language name to simple 2-letter code"""
    if not lang_name_or_code_input:
        return ""
    lang_lower = lang_name_or_code_input.strip().lower()
    lang_map = {
        "english": "en", "dutch": "nl", "german": "de", "french": "fr",
        "spanish": "es", "italian": "it", "japanese": "ja", "chinese": "zh",
        "russian": "ru", "portuguese": "pt",
    }
    if lang_lower in lang_map:
        return lang_map[lang_lower]
    base_code = lang_lower.split('-')[0].split('_')[0]
    if len(base_code) == 2:
        return base_code
    return lang_lower[:2]


# --- Tracked Changes Agent ---
class TrackedChangesAgent:
    """
    Manages tracked changes from DOCX files or TSV files.
    Provides AI with examples of preferred editing patterns to learn translator style.
    """
    def __init__(self, log_callback=None):
        self.change_data = []  # List of (original_text, final_text) tuples
        self.files_loaded = []  # Track which files have been loaded
        self.log_callback = log_callback or print
    
    def log(self, message):
        """Log a message"""
        if callable(self.log_callback):
            self.log_callback(message)
    
    def load_docx_changes(self, docx_path):
        """Load tracked changes from a DOCX file"""
        if not docx_path:
            return False
            
        self.log(f"[Tracked Changes] Loading changes from: {os.path.basename(docx_path)}")
        
        try:
            new_changes = parse_docx_pairs(docx_path)
            
            # Clear existing changes to prevent duplicates
            self.change_data.clear()
            self.files_loaded.clear()
            
            # Add new changes
            self.change_data.extend(new_changes)
            self.files_loaded.append(os.path.basename(docx_path))
            
            self.log(f"[Tracked Changes] Loaded {len(new_changes)} change pairs from {os.path.basename(docx_path)}")
            self.log(f"[Tracked Changes] Total change pairs available: {len(self.change_data)}")
            
            return True
        except Exception as e:
            self.log(f"[Tracked Changes] Error loading {docx_path}: {e}")
            messagebox.showerror("Tracked Changes Error", 
                               f"Failed to load tracked changes from {os.path.basename(docx_path)}:\\n{e}")
            return False
    
    def load_tsv_changes(self, tsv_path):
        """Load tracked changes from a TSV file (original_text<tab>final_text format)"""
        if not tsv_path:
            return False
            
        self.log(f"[Tracked Changes] Loading changes from: {os.path.basename(tsv_path)}")
        
        try:
            new_changes = []
            with open(tsv_path, 'r', encoding='utf-8') as f:
                for line_num, line in enumerate(f, 1):
                    line = line.rstrip('\n\r')
                    if not line.strip():
                        continue
                    
                    # Skip header line if it looks like one
                    if line_num == 1 and ('original' in line.lower() and 'final' in line.lower()):
                        continue
                    
                    parts = line.split('\t')
                    if len(parts) >= 2:
                        original = parts[0].strip()
                        final = parts[1].strip()
                        if original and final and original != final:  # Only add if actually different
                            new_changes.append((original, final))
                    else:
                        self.log(f"[Tracked Changes] Skipping line {line_num}: insufficient columns")
            
            # Add to existing changes
            self.change_data.extend(new_changes)
            self.files_loaded.append(os.path.basename(tsv_path))
            
            self.log(f"[Tracked Changes] Loaded {len(new_changes)} change pairs from {os.path.basename(tsv_path)}")
            self.log(f"[Tracked Changes] Total change pairs available: {len(self.change_data)}")
            
            return True
        except Exception as e:
            self.log(f"[Tracked Changes] Error loading {tsv_path}: {e}")
            messagebox.showerror("Tracked Changes Error", 
                               f"Failed to load tracked changes from {os.path.basename(tsv_path)}:\\n{e}")
            return False
    
    def clear_changes(self):
        """Clear all loaded tracked changes"""
        self.change_data.clear()
        self.files_loaded.clear()
        self.log("[Tracked Changes] All tracked changes cleared")
    
    def search_changes(self, search_text, exact_match=False):
        """Search for changes containing the search text"""
        if not search_text.strip():
            return self.change_data
        
        search_lower = search_text.lower()
        results = []
        
        for original, final in self.change_data:
            if exact_match:
                if search_text == original or search_text == final:
                    results.append((original, final))
            else:
                if (search_lower in original.lower() or 
                    search_lower in final.lower()):
                    results.append((original, final))
        
        return results

    def find_relevant_changes(self, source_segments, max_changes=10):
        """
        Find tracked changes relevant to the current source segments being processed.
        Uses two-pass algorithm: exact matches first, then partial word overlap.
        """
        if not self.change_data or not source_segments:
            return []
        
        relevant_changes = []
        
        # First pass: exact matches
        for segment in source_segments:
            segment_lower = segment.lower().strip()
            for original, final in self.change_data:
                original_lower = original.lower().strip()
                if segment_lower == original_lower and (original, final) not in relevant_changes:
                    relevant_changes.append((original, final))
                    if len(relevant_changes) >= max_changes:
                        return relevant_changes
        
        # Second pass: partial matches (word overlap)
        if len(relevant_changes) < max_changes:
            for segment in source_segments:
                segment_words = set(word.lower() for word in segment.split() if len(word) > 3)
                for original, final in self.change_data:
                    if (original, final) in relevant_changes:
                        continue
                    
                    original_words = set(word.lower() for word in original.split() if len(word) > 3)
                    # Check if there's significant word overlap
                    if segment_words and original_words:
                        overlap = len(segment_words.intersection(original_words))
                        min_overlap = min(2, len(segment_words) // 2)
                        if overlap >= min_overlap:
                            relevant_changes.append((original, final))
                            if len(relevant_changes) >= max_changes:
                                return relevant_changes
        
        return relevant_changes
    
    def get_entry_count(self):
        """Get number of loaded change pairs"""
        return len(self.change_data)


# --- TMX Generator Class ---

class TrackedChangesBrowser:
    def __init__(self, parent, tracked_changes_agent, parent_app=None, log_queue=None):
        self.parent = parent
        self.tracked_changes_agent = tracked_changes_agent
        self.parent_app = parent_app  # Reference to main app for AI settings
        self.log_queue = log_queue if log_queue else queue.Queue()
        self.window = None
    
    def show_browser(self):
        """Show the tracked changes browser window"""
        if not self.tracked_changes_agent.change_data:
            messagebox.showinfo("No Changes", "No tracked changes loaded. Load a DOCX or TSV file with tracked changes first.")
            return
        
        # Create window if it doesn't exist
        if self.window is None or not self.window.winfo_exists():
            self.create_window()
        else:
            self.window.lift()
    
    def create_window(self):
        """Create the browser window"""
        self.window = tk.Toplevel(self.parent)
        self.window.title(f"Tracked Changes Browser ({len(self.tracked_changes_agent.change_data)} pairs)")
        self.window.geometry("900x700")  # Taller to accommodate detail view
        
        # Search frame
        search_frame = tk.Frame(self.window)
        search_frame.pack(fill=tk.X, padx=10, pady=5)
        
        tk.Label(search_frame, text="Search:").pack(side=tk.LEFT)
        self.search_var = tk.StringVar()
        search_entry = tk.Entry(search_frame, textvariable=self.search_var, width=40)
        search_entry.pack(side=tk.LEFT, padx=(5,0))
        search_entry.bind('<KeyRelease>', self.on_search)
        
        self.exact_match_var = tk.BooleanVar()
        tk.Checkbutton(search_frame, text="Exact match", variable=self.exact_match_var, 
                      command=self.on_search).pack(side=tk.LEFT, padx=(10,0))
        
        tk.Button(search_frame, text="Clear", command=self.clear_search).pack(side=tk.LEFT, padx=(10,0))
        
        # Results info
        self.results_label = tk.Label(self.window, text="")
        self.results_label.pack(pady=2)
        
        # Main content frame (results + detail)
        main_frame = tk.Frame(self.window)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Results frame with scrollbar (top half)
        results_frame = tk.Frame(main_frame)
        results_frame.pack(fill=tk.BOTH, expand=True)
        
        # Create Treeview for displaying changes
        columns = ('Original', 'Final')
        self.tree = ttk.Treeview(results_frame, columns=columns, show='headings', height=12)
        
        # Define headings
        self.tree.heading('Original', text='Original Text')
        self.tree.heading('Final', text='Final Text')
        
        # Configure column widths
        self.tree.column('Original', width=400)
        self.tree.column('Final', width=400)
        
        # Add scrollbars
        v_scrollbar = ttk.Scrollbar(results_frame, orient=tk.VERTICAL, command=self.tree.yview)
        h_scrollbar = ttk.Scrollbar(results_frame, orient=tk.HORIZONTAL, command=self.tree.xview)
        self.tree.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
        
        # Pack tree and scrollbars
        self.tree.grid(row=0, column=0, sticky="nsew")
        v_scrollbar.grid(row=0, column=1, sticky="ns")
        h_scrollbar.grid(row=1, column=0, sticky="ew")
        
        results_frame.grid_rowconfigure(0, weight=1)
        results_frame.grid_columnconfigure(0, weight=1)
        
        # Detail view frame (bottom half)
        detail_frame = tk.LabelFrame(main_frame, text="Selected Change Details", padx=5, pady=5)
        detail_frame.pack(fill=tk.BOTH, expand=False, pady=(10,0))
        
        # Original text display
        tk.Label(detail_frame, text="Original Text:", font=("Segoe UI", 10, "bold")).pack(anchor="w")
        self.original_text = tk.Text(detail_frame, height=4, wrap=tk.WORD, state="disabled", 
                                    bg="#f8f8f8", relief="solid", borderwidth=1)
        self.original_text.pack(fill=tk.X, pady=(2,5))
        
        # Final text display
        tk.Label(detail_frame, text="Final Text:", font=("Segoe UI", 10, "bold")).pack(anchor="w")
        self.final_text = tk.Text(detail_frame, height=4, wrap=tk.WORD, state="disabled",
                                 bg="#f0f8ff", relief="solid", borderwidth=1)
        self.final_text.pack(fill=tk.X, pady=(2,0))
        
        # Bind selection event
        self.tree.bind('<<TreeviewSelect>>', self.on_selection_change)
        
        # Context menu for copying
        self.context_menu = tk.Menu(self.window, tearoff=0)
        self.context_menu.add_command(label="Copy Original", command=self.copy_original)
        self.context_menu.add_command(label="Copy Final", command=self.copy_final)
        self.context_menu.add_command(label="Copy Both", command=self.copy_both)
        
        self.tree.bind("<Button-3>", self.show_context_menu)  # Right click
        
        # Export button frame
        export_frame = tk.Frame(self.window)
        export_frame.pack(fill=tk.X, padx=10, pady=5)
        
        tk.Button(export_frame, text="📊 Export Report (MD)", command=self.export_to_md_report,
                 bg="#4CAF50", fg="white", font=("Segoe UI", 10, "bold"),
                 relief="raised", padx=10, pady=5).pack(side=tk.LEFT)
        
        tk.Label(export_frame, text="Export tracked changes report with AI-powered change analysis",
                fg="gray").pack(side=tk.LEFT, padx=(10,0))
        
        # Status bar
        status_frame = tk.Frame(self.window)
        status_frame.pack(fill=tk.X, padx=10, pady=2)
        
        files_text = f"Files loaded: {', '.join(self.tracked_changes_agent.files_loaded)}" if self.tracked_changes_agent.files_loaded else "No files loaded"
        tk.Label(status_frame, text=files_text, anchor=tk.W).pack(fill=tk.X)
        
        # Load all changes initially
        self.load_results(self.tracked_changes_agent.change_data)
    
    def on_selection_change(self, event=None):
        """Handle selection change in the tree"""
        selection = self.tree.selection()
        if not selection:
            # Clear detail view if no selection
            self.original_text.config(state="normal")
            self.original_text.delete(1.0, tk.END)
            self.original_text.config(state="disabled")
            self.final_text.config(state="normal")
            self.final_text.delete(1.0, tk.END)
            self.final_text.config(state="disabled")
            return
        
        # Get the selected change pair
        original, final = self.get_selected_change()
        if original and final:
            # Update original text display
            self.original_text.config(state="normal")
            self.original_text.delete(1.0, tk.END)
            self.original_text.insert(1.0, original)
            self.original_text.config(state="disabled")
            
            # Update final text display
            self.final_text.config(state="normal")
            self.final_text.delete(1.0, tk.END)
            self.final_text.insert(1.0, final)
            self.final_text.config(state="disabled")
    
    def on_search(self, event=None):
        """Handle search input"""
        search_text = self.search_var.get()
        exact_match = self.exact_match_var.get()
        
        results = self.tracked_changes_agent.search_changes(search_text, exact_match)
        self.load_results(results)
    
    def clear_search(self):
        """Clear search and show all results"""
        self.search_var.set("")
        self.load_results(self.tracked_changes_agent.change_data)
    
    def load_results(self, results):
        """Load results into the treeview"""
        # Clear existing items
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        # Add new items
        for i, (original, final) in enumerate(results):
            # Truncate long text for display
            display_original = (original[:100] + "...") if len(original) > 100 else original
            display_final = (final[:100] + "...") if len(final) > 100 else final
            
            self.tree.insert('', 'end', values=(display_original, display_final))
        
        # Update results label
        total_changes = len(self.tracked_changes_agent.change_data)
        showing = len(results)
        if showing == total_changes:
            self.results_label.config(text=f"Showing all {total_changes} change pairs")
        else:
            self.results_label.config(text=f"Showing {showing} of {total_changes} change pairs")
    
    def show_context_menu(self, event):
        """Show context menu for copying"""
        item = self.tree.identify_row(event.y)
        if item:
            self.tree.selection_set(item)
            self.context_menu.post(event.x_root, event.y_root)
    
    def get_selected_change(self):
        """Get the currently selected change pair"""
        selection = self.tree.selection()
        if not selection:
            return None, None
        
        item = selection[0]
        index = self.tree.index(item)
        
        # Get current results (might be filtered)
        search_text = self.search_var.get()
        exact_match = self.exact_match_var.get()
        current_results = self.tracked_changes_agent.search_changes(search_text, exact_match)
        
        if 0 <= index < len(current_results):
            return current_results[index]
        return None, None
    
    def copy_original(self):
        """Copy original text to clipboard"""
        original, _ = self.get_selected_change()
        if original:
            self.window.clipboard_clear()
            self.window.clipboard_append(original)
    
    def copy_final(self):
        """Copy final text to clipboard"""
        _, final = self.get_selected_change()
        if final:
            self.window.clipboard_clear()
            self.window.clipboard_append(final)
    
    def copy_both(self):
        """Copy both texts to clipboard"""
        original, final = self.get_selected_change()
        if original and final:
            both_text = f"Original: {original}\n\nFinal: {final}"
            self.window.clipboard_clear()
            self.window.clipboard_append(both_text)

    
    def export_to_md_report(self):
        """Export tracked changes to a Markdown report with AI-powered change analysis"""
        from tkinter import filedialog, messagebox
        
        if not self.tracked_changes_agent.change_data:
            messagebox.showwarning("No Data", "No tracked changes available to export.")
            return
        
        # Ask user whether to export all or filtered results
        search_text = self.search_var.get()
        if search_text:
            # User has active search filter
            result = messagebox.askyesnocancel(
                "Export Scope",
                f"You have an active search filter showing {len(self.tree.get_children())} of {len(self.tracked_changes_agent.change_data)} changes.\n\n"
                "Yes = Export filtered results only\n"
                "No = Export all tracked changes\n"
                "Cancel = Cancel export"
            )
            if result is None:  # Cancel
                return
            export_filtered = result
        else:
            export_filtered = False
        
        # Get the data to export
        if export_filtered:
            exact_match = self.exact_match_var.get()
            data_to_export = self.tracked_changes_agent.search_changes(search_text, exact_match)
            default_filename = "tracked_changes_filtered_report.md"
        else:
            data_to_export = self.tracked_changes_agent.change_data
            default_filename = "tracked_changes_report.md"
        
        # Ask for save location
        filepath = filedialog.asksaveasfilename(
            title="Export Tracked Changes Report",
            defaultextension=".md",
            filetypes=(("Markdown files", "*.md"), ("All files", "*.*")),
            initialfile=default_filename
        )
        
        if not filepath:
            return
        
        # Ask if user wants AI analysis
        ai_analysis = messagebox.askyesno(
            "AI Analysis",
            f"Generate AI-powered change summaries?\n\n"
            f"This will analyze {len(data_to_export)} changes using the currently selected AI model.\n\n"
            f"Note: This may take a few minutes and will use API credits.\n\n"
            f"Click 'No' to export without AI analysis."
        )
        
        # If AI analysis enabled, let user choose batch size
        batch_size = 25  # Default
        if ai_analysis:
            batch_dialog = tk.Toplevel(self.window)
            batch_dialog.title("Batch Size Configuration")
            batch_dialog.geometry("450x280")
            batch_dialog.transient(self.window)
            batch_dialog.grab_set()
            
            tk.Label(batch_dialog, text="Configure Batch Processing", 
                    font=("Segoe UI", 11, "bold")).pack(pady=10)
            tk.Label(batch_dialog, 
                    text=f"Choose how many segments to process per AI request\n"
                         f"Larger batches = faster but more tokens per request",
                    font=("Segoe UI", 9)).pack(pady=5)
            
            # Slider for batch size
            batch_var = tk.IntVar(value=25)
            
            slider_frame = tk.Frame(batch_dialog)
            slider_frame.pack(pady=10, fill='x', padx=20)
            
            tk.Label(slider_frame, text="Batch Size:", font=("Segoe UI", 9)).pack(side='left')
            batch_label = tk.Label(slider_frame, text="25", font=("Segoe UI", 10, "bold"), fg="blue")
            batch_label.pack(side='right')
            
            def update_label(val):
                batch_label.config(text=str(int(float(val))))
            
            slider = tk.Scale(batch_dialog, from_=1, to=100, orient='horizontal',
                            variable=batch_var, command=update_label, length=350)
            slider.pack(pady=5)
            
            # Info label
            info_label = tk.Label(batch_dialog, 
                                text=f"Total changes: {len(data_to_export)} | "
                                     f"Estimated batches at size 25: {(len(data_to_export) + 24) // 25}",
                                font=("Segoe UI", 8), fg="gray")
            info_label.pack(pady=5)
            
            def update_info(*args):
                size = batch_var.get()
                batches = (len(data_to_export) + size - 1) // size
                info_label.config(text=f"Total changes: {len(data_to_export)} | "
                                      f"Estimated batches at size {size}: {batches}")
            
            batch_var.trace('w', update_info)
            
            # OK button
            def on_ok():
                nonlocal batch_size
                batch_size = batch_var.get()
                batch_dialog.destroy()
            
            tk.Button(batch_dialog, text="OK", command=on_ok, 
                     font=("Segoe UI", 10), width=15).pack(pady=10)
            
            # Wait for dialog to close
            batch_dialog.wait_window()
        
        try:
            # Prepare report content
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            # Build AI prompt info for report header
            ai_prompt_info = ""
            if ai_analysis and hasattr(self, 'parent_app') and self.parent_app:
                provider = self.parent_app.current_llm_provider
                model = self.parent_app.current_llm_model
                # Capitalize provider name for display
                provider_display = provider.capitalize()
                ai_prompt_info = f"""

### AI Analysis Configuration

**Provider:** {provider_display}  
**Model:** {model}

**Prompt Template Used:**
```
You are a precision editor analyzing changes between two versions of text.
Compare the original and revised text and identify EXACTLY what changed.

CRITICAL INSTRUCTIONS:
- Be extremely specific and precise
- PAY SPECIAL ATTENTION to quote marks: " vs " vs " (curly vs straight)
- Check for apostrophe changes: ' vs ' (curly vs straight)  
- Check for dash changes: - vs – vs — (hyphen vs en-dash vs em-dash)
- Quote the exact words/phrases that changed
- Use this format: "X" → "Y"
- For single word changes: quote both words
- For multiple changes: put each on its own line
- For punctuation/formatting: describe precisely
- For additions: "Added: [exact text]"
- For deletions: "Removed: [exact text]"
- DO NOT say "No change" unless texts are 100% identical
- DO NOT use vague terms like "clarified", "improved", "fixed"
- DO quote the actual changed text

Examples of single changes:
✓ "pre-cut" → "incision"
✓ Curly quotes → straight quotes: "word" → "word"
✓ Curly apostrophe → straight: don't → don't
✓ "package" → "packaging"

Examples of multiple changes (one per line):
✓ "split portions" → "divided portions"
  "connected by a" → "connected, via a"
  Curly quotes → straight quotes throughout
✓ Added: "carefully"
✓ "color" → "colour" (US to UK spelling)
✗ Clarified terminology (too vague)
✗ Fixed grammar (not specific)
✗ Improved word choice (not helpful)
```

---

"""
            
            md_content = f"""# Tracked Changes Analysis Report ([Supervertaler](https://github.com/michaelbeijer/Supervertaler) {APP_VERSION})

## What is this report?

This report analyzes the differences between AI-generated translations and your final edited versions exported from your CAT tool (memoQ, CafeTran, etc.). It shows exactly what you changed during post-editing, helping you review your editing decisions and track your translation workflow improvements.

**Use case:** After completing a translation project in your CAT tool with tracked changes enabled, export the bilingual document and load it here to see a detailed breakdown of all modifications made to the AI-generated baseline.

---

**Generated:** {timestamp}  
**Total Changes:** {len(data_to_export)}  
**Filter Applied:** {"Yes - " + search_text if export_filtered else "No"}  
**AI Analysis:** {"Enabled" if ai_analysis else "Disabled"}
{ai_prompt_info}
"""
            
            # Process changes with paragraph format
            if ai_analysis:
                # Show progress window
                self.log_queue.put(f"[Export] Generating AI summaries for {len(data_to_export)} changes in batches...")
                
                progress_window = tk.Toplevel(self.window)
                progress_window.title("Generating AI Analysis...")
                progress_window.geometry("400x150")
                progress_window.transient(self.window)
                progress_window.grab_set()
                
                tk.Label(progress_window, text="Analyzing tracked changes with AI (batched)...", 
                        font=("Segoe UI", 10)).pack(pady=10)
                progress_label = tk.Label(progress_window, text="Processing batch 0 of 0")
                progress_label.pack()
                batch_info_label = tk.Label(progress_window, text="", font=("Segoe UI", 8), fg="gray")
                batch_info_label.pack()
                
                # Process in batches (user-configured)
                # batch_size already set from dialog above
                total_batches = (len(data_to_export) + batch_size - 1) // batch_size
                all_summaries = {}
                
                for batch_num in range(total_batches):
                    start_idx = batch_num * batch_size
                    end_idx = min(start_idx + batch_size, len(data_to_export))
                    batch = data_to_export[start_idx:end_idx]
                    
                    progress_label.config(text=f"Processing batch {batch_num + 1} of {total_batches}")
                    batch_info_label.config(text=f"Segments {start_idx + 1}-{end_idx} of {len(data_to_export)}")
                    progress_window.update()
                    
                    # Generate AI summaries for this batch
                    try:
                        batch_summaries = self.get_ai_change_summaries_batch(batch, start_idx)
                        all_summaries.update(batch_summaries)
                        self.log_queue.put(f"[Export] Batch {batch_num + 1}/{total_batches} complete ({len(batch)} segments)")
                    except Exception as e:
                        self.log_queue.put(f"[Export] Error in batch {batch_num + 1}: {e}")
                        # Fill in error messages for failed batch
                        for i in range(start_idx, end_idx):
                            all_summaries[i] = f"_Error generating summary: {str(e)}_"
                
                progress_window.destroy()
                
                # Now build the markdown content with the summaries
                for i, (original, final) in enumerate(data_to_export):
                    summary = all_summaries.get(i, "_No summary available_")
                    
                    # Add segment in paragraph format
                    md_content += f"""### Segment {i + 1}

**Target (Original):**  
{original}

**Target (Revised):**  
{final}

**Change Summary:**  
{summary}

---

"""
            else:
                # No AI analysis - simpler paragraph format
                for i, (original, final) in enumerate(data_to_export, 1):
                    md_content += f"""### Segment {i}

**Target (Original):**  
{original}

**Target (Revised):**  
{final}

---

"""
            
            md_content += f"""

---

## Summary Statistics

- **Total Segments Analyzed:** {len(data_to_export)}
- **AI Analysis:** {"Enabled" if ai_analysis else "Disabled"}
- **Export Type:** {"Filtered" if export_filtered else "Complete"}

*This report was generated by [Supervertaler](https://github.com/michaelbeijer/Supervertaler) {APP_VERSION}*
"""
            
            # Write to file
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(md_content)
            
            messagebox.showinfo(
                "Export Successful",
                f"Exported {len(data_to_export)} tracked changes to:\n{filepath}\n\n"
                + ("AI change summaries included." if ai_analysis else "Export completed without AI analysis.")
            )
            
            self.log_queue.put(f"[Export] Report saved to: {filepath}")
            
        except Exception as e:
            messagebox.showerror(
                "Export Error",
                f"Failed to export tracked changes report:\n{str(e)}"
            )
            self.log_queue.put(f"[Export] Error: {e}")
    
    def get_ai_change_summaries_batch(self, changes_batch, start_index):
        """Get AI summaries for a batch of changes - much faster than one-by-one"""
        if not hasattr(self, 'parent_app') or not self.parent_app:
            # Fallback for batch
            return {i: "Modified text" for i in range(start_index, start_index + len(changes_batch))}
        
        try:
            provider = self.parent_app.current_llm_provider
            model_name = self.parent_app.current_llm_model
            api_key = ""
            
            # Debug logging
            self.log_queue.put(f"[Export] Using provider: {provider}, model: {model_name}")
            
            if provider == "claude":
                api_key = self.parent_app.api_keys.get("claude", "")
            elif provider == "gemini":
                api_key = self.parent_app.api_keys.get("google", "")
            elif provider == "openai":
                api_key = self.parent_app.api_keys.get("openai", "")
            
            if not api_key:
                self.log_queue.put(f"[Export] ERROR: No API key found for provider: {provider}")
                return {i: "AI unavailable - no API key" for i in range(start_index, start_index + len(changes_batch))}
            
            self.log_queue.put(f"[Export] API key found, calling {provider}...")
            
            # Build batch prompt with all changes
            batch_prompt = """You are a precision editor analyzing changes between multiple text versions.
For each numbered pair below, identify EXACTLY what changed.

CRITICAL INSTRUCTIONS:
- Be extremely specific and precise
- PAY SPECIAL ATTENTION to quote marks: " vs " vs " (curly vs straight)
- Check for apostrophe changes: ' vs ' (curly vs straight)
- Check for dash changes: - vs – vs — (hyphen vs en-dash vs em-dash)
- Quote the exact words/phrases that changed
- Use format: "X" → "Y"
- For multiple changes in one segment: put each on its own line
- For punctuation/formatting: describe precisely (e.g., 'Curly quotes → straight quotes: "word" → "word"')
- DO NOT say "No change" unless texts are 100% identical (byte-for-byte)
- DO NOT use vague terms like "clarified", "improved", "fixed"
- DO quote the actual changed text

IMPORTANT: If only punctuation changed (quotes, apostrophes, dashes), you MUST report it!

"""
            
            # Add all changes to the prompt
            for i, (original, final) in enumerate(changes_batch):
                batch_prompt += f"""
[{i + 1}] ORIGINAL: {original}
    REVISED: {final}

"""
            
            batch_prompt += """
Now provide the change summary for each segment, formatted as:

[1] your precise summary here
[2] your precise summary here
[3] your precise summary here

(etc. for all segments)"""
            
            # Call AI based on provider
            self.log_queue.put(f"[Export] Checking provider condition: {provider} == gemini? {provider == 'gemini'}, GEMINI_AVAILABLE? {GEMINI_AVAILABLE}")
            if provider == "gemini" and GEMINI_AVAILABLE:
                import google.generativeai as genai
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel(model_name)
                
                response = model.generate_content(batch_prompt)
                response_text = response.text.strip()
                
            elif provider == "claude" and ANTHROPIC_AVAILABLE:
                import anthropic
                client = anthropic.Anthropic(api_key=api_key)
                
                message = client.messages.create(
                    model=model_name,
                    max_tokens=2000,  # Larger for batch
                    messages=[{
                        "role": "user",
                        "content": batch_prompt
                    }]
                )
                
                response_text = message.content[0].text.strip()
                
            elif provider == "openai" and OPENAI_AVAILABLE:
                import openai
                client = openai.OpenAI(api_key=api_key)
                
                response = client.chat.completions.create(
                    model=model_name,
                    max_tokens=2000,  # Larger for batch
                    messages=[{
                        "role": "user",
                        "content": batch_prompt
                    }]
                )
                
                response_text = response.choices[0].message.content.strip()
            else:
                self.log_queue.put(f"[Export] ERROR: No matching provider condition for {provider}")
                return {i: "Provider not available" for i in range(start_index, start_index + len(changes_batch))}
            
            # Parse the response to extract individual summaries
            summaries = {}
            current_num = None
            current_summary_lines = []
            
            for line in response_text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                
                # Check if line starts with [N]
                import re
                match = re.match(r'^\[(\d+)\]\s*(.*)$', line)
                if match:
                    # Save previous summary if any
                    if current_num is not None:
                        summary_text = '\n'.join(current_summary_lines).strip()
                        summaries[start_index + current_num - 1] = summary_text
                    
                    # Start new summary
                    current_num = int(match.group(1))
                    summary_start = match.group(2).strip()
                    current_summary_lines = [summary_start] if summary_start else []
                elif current_num is not None:
                    # Continuation of current summary
                    current_summary_lines.append(line)
            
            # Save last summary
            if current_num is not None:
                summary_text = '\n'.join(current_summary_lines).strip()
                summaries[start_index + current_num - 1] = summary_text
            
            # Fill in any missing summaries
            for i in range(len(changes_batch)):
                if (start_index + i) not in summaries:
                    summaries[start_index + i] = "_Summary not parsed correctly_"
            
            return summaries
            
        except Exception as e:
            self.log_queue.put(f"[AI Batch] Error: {e}")
            return {i: f"Analysis failed: {str(e)}" for i in range(start_index, start_index + len(changes_batch))}
    
    def get_ai_change_summary(self, original_text, revised_text):
        """Get AI summary of what changed between original and revised text"""
        # This method needs to access the parent app's AI configuration
        # We'll need to pass the parent app reference when creating the browser
        
        # For now, return a simple diff-based summary as fallback
        # The parent app integration will be added in the next step
        
        if not hasattr(self, 'parent_app'):
            # Fallback to simple analysis
            if original_text == revised_text:
                return "No change"
            elif len(revised_text) > len(original_text):
                return "Expanded/added content"
            elif len(revised_text) < len(original_text):
                return "Shortened/removed content"
            else:
                return "Modified wording"
        
        # If parent_app is available, use its AI agent
        try:
            provider = self.parent_app.current_llm_provider
            model_name = self.parent_app.current_llm_model
            api_key = ""
            
            if provider == "claude":
                api_key = self.parent_app.api_keys.get("claude", "")
            elif provider == "gemini":
                api_key = self.parent_app.api_keys.get("google", "")
            elif provider == "openai":
                api_key = self.parent_app.api_keys.get("openai", "")
            
            if not api_key:
                return "AI unavailable"
            
            # Create a temporary agent for this analysis
            if provider == "gemini" and GEMINI_AVAILABLE:
                import google.generativeai as genai
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel(model_name)
                
                prompt = f"""You are a precision editor analyzing changes between two versions of text.
Compare the original and revised text and identify EXACTLY what changed.

Original: {original_text}
Revised: {revised_text}

CRITICAL INSTRUCTIONS:
- Be extremely specific and precise
- PAY SPECIAL ATTENTION to quote marks: " vs " vs " (curly vs straight)
- Check for apostrophe changes: ' vs ' (curly vs straight)  
- Check for dash changes: - vs – vs — (hyphen vs en-dash vs em-dash)
- Quote the exact words/phrases that changed
- Use this format: "X" → "Y"
- For single word changes: quote both words
- For multiple changes: put each on its own line
- For punctuation/formatting: describe precisely
- For additions: "Added: [exact text]"
- For deletions: "Removed: [exact text]"
- DO NOT say "No change" unless texts are 100% identical
- DO NOT use vague terms like "clarified", "improved", "fixed"
- DO quote the actual changed text

Examples of single changes:
✓ "pre-cut" → "incision"
✓ Curly quotes → straight quotes: "word" → "word"
✓ Curly apostrophe → straight: don't → don't
✓ "package" → "packaging"

Examples of multiple changes (one per line):
✓ "split portions" → "divided portions"
  "connected by a" → "connected, via a"
  Curly quotes → straight quotes throughout
✗ Clarified terminology
✗ Fixed grammar

Your precise change summary:"""
                
                response = model.generate_content(prompt)
                summary = response.text.strip()
                # Clean up the response
                summary = summary.split('.')[0].split('\n')[0]
                words = summary.split()
                if len(words) > 10:
                    summary = ' '.join(words[:10]) + '...'
                return summary
                
            elif provider == "claude" and ANTHROPIC_AVAILABLE:
                import anthropic
                client = anthropic.Anthropic(api_key=api_key)
                
                message = client.messages.create(
                    model=model_name,
                    max_tokens=100,
                    messages=[{
                        "role": "user",
                        "content": f"""You are a precision editor analyzing changes between two versions of text.
Compare the original and revised text and identify EXACTLY what changed.

Original: {original_text}
Revised: {revised_text}

CRITICAL INSTRUCTIONS:
- Be extremely specific and precise
- PAY SPECIAL ATTENTION to quote marks: " vs " vs " (curly vs straight)
- Check for apostrophe changes: ' vs ' (curly vs straight)  
- Check for dash changes: - vs – vs — (hyphen vs en-dash vs em-dash)
- Quote the exact words/phrases that changed
- Use this format: "X" → "Y"
- For single word changes: quote both words
- For multiple changes: put each on its own line
- For punctuation/formatting: describe precisely
- For additions: "Added: [exact text]"
- For deletions: "Removed: [exact text]"
- DO NOT say "No change" unless texts are 100% identical
- DO NOT use vague terms like "clarified", "improved", "fixed"
- DO quote the actual changed text

Examples of single changes:
✓ "pre-cut" → "incision"
✓ Curly quotes → straight quotes: "word" → "word"
✓ Curly apostrophe → straight: don't → don't
✓ "package" → "packaging"

Examples of multiple changes (one per line):
✓ "split portions" → "divided portions"
  "connected by a" → "connected, via a"
  Curly quotes → straight quotes throughout
✗ Clarified terminology
✗ Fixed grammar

Your precise change summary:"""
                    }]
                )
                
                summary = message.content[0].text.strip()
                # Clean up the response - remove extra formatting
                summary = summary.replace('Your precise change summary:', '').strip()
                summary = summary.split('\n')[0]  # First line only
                return summary
                
            elif provider == "openai" and OPENAI_AVAILABLE:
                import openai
                client = openai.OpenAI(api_key=api_key)
                
                response = client.chat.completions.create(
                    model=model_name,
                    max_tokens=100,
                    messages=[{
                        "role": "user",
                        "content": f"""You are a precision editor analyzing changes between two versions of text.
Compare the original and revised text and identify EXACTLY what changed.

Original: {original_text}
Revised: {revised_text}

CRITICAL INSTRUCTIONS:
- Be extremely specific and precise
- PAY SPECIAL ATTENTION to quote marks: " vs " vs " (curly vs straight)
- Check for apostrophe changes: ' vs ' (curly vs straight)  
- Check for dash changes: - vs – vs — (hyphen vs en-dash vs em-dash)
- Quote the exact words/phrases that changed
- Use this format: "X" → "Y"
- For single word changes: quote both words
- For multiple changes: put each on its own line
- For punctuation/formatting: describe precisely
- For additions: "Added: [exact text]"
- For deletions: "Removed: [exact text]"
- DO NOT say "No change" unless texts are 100% identical
- DO NOT use vague terms like "clarified", "improved", "fixed"
- DO quote the actual changed text

Examples of single changes:
✓ "pre-cut" → "incision"
✓ Curly quotes → straight quotes: "word" → "word"
✓ Curly apostrophe → straight: don't → don't
✓ "package" → "packaging"

Examples of multiple changes (one per line):
✓ "split portions" → "divided portions"
  "connected by a" → "connected, via a"
  Curly quotes → straight quotes throughout
✗ Clarified terminology
✗ Fixed grammar

Your precise change summary:"""
                    }]
                )
                
                summary = response.choices[0].message.content.strip()
                # Clean up the response - remove extra formatting
                summary = summary.replace('Your precise change summary:', '').strip()
                summary = summary.split('\n')[0]  # First line only
                return summary
            else:
                return "Simple text change"
                
        except Exception as e:
            self.log_queue.put(f"[AI Summary] Error: {e}")
            return "Analysis failed"

# --- Helper Functions ---
def get_simple_lang_code(lang_name_or_code_input):
    if not lang_name_or_code_input: return ""
    lang_lower = lang_name_or_code_input.strip().lower()
    lang_map = {
        "english": "en", "dutch": "nl", "german": "de", "french": "fr",
        "spanish": "es", "italian": "it", "japanese": "ja", "chinese": "zh",
        "russian": "ru", "portuguese": "pt",
    }
    if lang_lower in lang_map: return lang_map[lang_lower]
    base_code = lang_lower.split('-')[0].split('_')[0]
    if len(base_code) == 2: return base_code
    return lang_lower[:2]

def normalize_figure_ref(ref_text):
    if not ref_text: return None
    match = re.search(r"(?:figure|figuur|fig\.?)\s*([\w\d]+(?:[\s\.\-]*[\w\d]+)?)", ref_text, re.IGNORECASE)
    if match:
        identifier = match.group(1); return re.sub(r"[\s\.\-]", "", identifier).lower()
    base_name = os.path.splitext(ref_text)[0]
    cleaned_base = re.sub(r"(?:figure|figuur|fig\.?)\s*", "", base_name, flags=re.IGNORECASE)
    normalized = re.sub(r"[\s\.\-]", "", cleaned_base).lower()
    if normalized: return normalized
    return None

def format_tracked_changes_context(tracked_changes_list, max_length=1000):
    """Format tracked changes for AI context, keeping within token limits"""
    if not tracked_changes_list:
        return ""
    
    context_parts = ["TRACKED CHANGES REFERENCE (Original→Final editing patterns):"]
    current_length = len(context_parts[0])
    
    for i, (original, final) in enumerate(tracked_changes_list):
        change_text = f"• \"{original}\" → \"{final}\""
        if current_length + len(change_text) > max_length:
            if i > 0:  # Only add if we have at least one example
                context_parts.append("(Additional examples truncated to save space)")
            break
        context_parts.append(change_text)
        current_length += len(change_text)
    
    return "\n".join(context_parts) + "\n"

def pil_image_to_base64_png(img):
    """Encode a PIL image to base64 PNG (ascii) for Claude/OpenAI data URLs."""
    try:
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return base64.b64encode(buf.getvalue()).decode("ascii")
    except Exception:
        return None

# --- TMX Generator Class ---


class TMXGenerator:
    """Helper class for generating TMX (Translation Memory eXchange) files"""
    
    def __init__(self, log_callback=None):
        self.log = log_callback if log_callback else lambda msg: None
    
    def generate_tmx(self, source_segments, target_segments, source_lang, target_lang):
        """Generate TMX content from parallel segments"""
        # Basic TMX structure
        tmx = ET.Element('tmx')
        tmx.set('version', '1.4')
        
        header = ET.SubElement(tmx, 'header')
        header.set('creationdate', datetime.now().strftime('%Y%m%dT%H%M%SZ'))
        header.set('srclang', get_simple_lang_code(source_lang))
        header.set('adminlang', 'en')
        header.set('segtype', 'sentence')
        header.set('creationtool', 'Supervertaler')
        header.set('creationtoolversion', '2.5.0')
        header.set('datatype', 'plaintext')
        
        body = ET.SubElement(tmx, 'body')
        
        # Add translation units
        added_count = 0
        for src, tgt in zip(source_segments, target_segments):
            if not src.strip() or not tgt or '[ERR' in str(tgt) or '[Missing' in str(tgt):
                continue
                
            tu = ET.SubElement(body, 'tu')
            
            # Source segment
            tuv_src = ET.SubElement(tu, 'tuv')
            tuv_src.set('xml:lang', get_simple_lang_code(source_lang))
            seg_src = ET.SubElement(tuv_src, 'seg')
            seg_src.text = src.strip()
            
            # Target segment
            tuv_tgt = ET.SubElement(tu, 'tuv')
            tuv_tgt.set('xml:lang', get_simple_lang_code(target_lang))
            seg_tgt = ET.SubElement(tuv_tgt, 'seg')
            seg_tgt.text = str(tgt).strip()
            
            added_count += 1
        
        self.log(f"[TMX Generator] Created TMX with {added_count} translation units")
        return ET.ElementTree(tmx)
    
    def save_tmx(self, tmx_tree, output_path):
        """Save TMX tree to file with proper XML formatting"""
        try:
            # Pretty print with indentation
            self._indent(tmx_tree.getroot())
            tmx_tree.write(output_path, encoding='utf-8', xml_declaration=True)
            self.log(f"[TMX Generator] Saved TMX file: {output_path}")
            return True
        except Exception as e:
            self.log(f"[TMX Generator] Error saving TMX: {e}")
            return False
    
    def _indent(self, elem, level=0):
        """Add indentation to XML for pretty printing"""
        i = "\n" + level * "  "
        if len(elem):
            if not elem.text or not elem.text.strip():
                elem.text = i + "  "
            if not elem.tail or not elem.tail.strip():
                elem.tail = i
            for child in elem:
                self._indent(child, level + 1)
            if not child.tail or not child.tail.strip():
                child.tail = i
        else:
            if level and (not elem.tail or not elem.tail.strip()):
                elem.tail = i


# --- Prompt Library Manager ---
class PromptLibrary:
    """
    Manages translation prompts with domain-specific expertise.
    Supports two types:
    - System Prompts: Define AI role and expertise
    - Custom Instructions: Additional context and preferences
    
    Loads JSON files from appropriate folders based on dev mode.
    """
    def __init__(self, system_prompts_dir, log_callback=None):
        # Use path resolver for both folder types
        self.system_prompts_dir = get_user_data_path("System_prompts")
        self.custom_instructions_dir = get_user_data_path("Custom_instructions")
        
        self.log = log_callback if log_callback else print
        
        # Create directories if they don't exist
        os.makedirs(self.system_prompts_dir, exist_ok=True)
        os.makedirs(self.custom_instructions_dir, exist_ok=True)
        
        # Available prompts: {filename: prompt_data}
        self.prompts = {}
        self.active_prompt = None  # Currently selected prompt
        self.active_prompt_name = None
        
    def load_all_prompts(self):
        """Load all prompts (system prompts and custom instructions) from appropriate directories"""
        self.prompts = {}
        
        # Load from the appropriate directories based on dev mode
        sys_count = self._load_from_directory(self.system_prompts_dir, prompt_type="system_prompt")
        inst_count = self._load_from_directory(self.custom_instructions_dir, prompt_type="custom_instruction")
        
        total = sys_count + inst_count
        self.log(f"✓ Loaded {total} prompts ({sys_count} system prompts, {inst_count} custom instructions)")
        return total
    
    def _load_from_directory(self, directory, prompt_type="system_prompt"):
        """Load prompts from a specific directory
        
        Args:
            directory: Path to directory
            prompt_type: Either 'system_prompt' or 'custom_instruction'
        """
        count = 0
        
        if not os.path.exists(directory):
            return count
        
        for filename in os.listdir(directory):
            if not filename.endswith('.json'):
                continue
            
            filepath = os.path.join(directory, filename)
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    prompt_data = json.load(f)
                    
                    # Add metadata
                    prompt_data['_filename'] = filename
                    prompt_data['_filepath'] = filepath
                    prompt_data['_type'] = prompt_type  # Add type field
                    
                    # Validate required fields
                    if 'name' not in prompt_data or 'translate_prompt' not in prompt_data:
                        self.log(f"⚠ Skipping {filename}: missing required fields (name, translate_prompt)")
                        continue
                    
                    self.prompts[filename] = prompt_data
                    count += 1
                    
            except Exception as e:
                self.log(f"⚠ Failed to load {filename}: {e}")
                
        return count
    
    def get_prompt_list(self):
        """Get list of available prompts with metadata"""
        prompt_list = []
        for filename, data in sorted(self.prompts.items()):
            prompt_list.append({
                'filename': filename,
                'name': data.get('name', 'Unnamed'),
                'description': data.get('description', ''),
                'domain': data.get('domain', 'General'),
                'version': data.get('version', '1.0'),
                'filepath': data.get('_filepath', ''),
                '_type': data.get('_type', 'system_prompt')  # Include type for filtering
            })
        return prompt_list
    
    def get_prompt(self, filename):
        """Get full prompt data by filename"""
        return self.prompts.get(filename)
    
    def set_active_prompt(self, filename):
        """Set the active custom prompt"""
        if filename not in self.prompts:
            self.log(f"✗ Prompt not found: {filename}")
            return False
        
        self.active_prompt = self.prompts[filename]
        self.active_prompt_name = self.active_prompt.get('name', filename)
        self.log(f"✓ Active prompt: {self.active_prompt_name}")
        return True
    
    def clear_active_prompt(self):
        """Clear active prompt (use default)"""
        self.active_prompt = None
        self.active_prompt_name = None
        self.log("✓ Using default translation prompt")
    
    def get_translate_prompt(self):
        """Get the translate_prompt from active prompt, or None if using default"""
        if self.active_prompt:
            return self.active_prompt.get('translate_prompt')
        return None
    
    def get_proofread_prompt(self):
        """Get the proofread_prompt from active prompt, or None if using default"""
        if self.active_prompt:
            return self.active_prompt.get('proofread_prompt')
        return None
    
    def search_prompts(self, search_text):
        """Search prompts by name, description, or domain"""
        if not search_text:
            return self.get_prompt_list()
        
        search_lower = search_text.lower()
        results = []
        
        for filename, data in sorted(self.prompts.items()):
            name = data.get('name', '').lower()
            desc = data.get('description', '').lower()
            domain = data.get('domain', '').lower()
            
            if search_lower in name or search_lower in desc or search_lower in domain:
                results.append({
                    'filename': filename,
                    'name': data.get('name', 'Unnamed'),
                    'description': data.get('description', ''),
                    'domain': data.get('domain', 'General'),
                    'version': data.get('version', '1.0'),
                    'filepath': data.get('_filepath', '')
                })
        
        return results
    
    def create_new_prompt(self, name, description, domain, translate_prompt, proofread_prompt="", 
                         version="1.0", prompt_type="system_prompt"):
        """Create a new prompt and save to JSON
        
        Args:
            prompt_type: Either 'system_prompt' or 'custom_instruction'
        """
        # Create filename from name
        filename = name.replace(' ', '_').replace('/', '_') + '.json'
        
        # Choose directory based on type using path resolver
        if prompt_type == "custom_instruction":
            directory = self.custom_instructions_dir
        else:  # system_prompt
            directory = self.system_prompts_dir
            
        filepath = os.path.join(directory, filename)
        
        # Create prompt data
        prompt_data = {
            'name': name,
            'description': description,
            'domain': domain,
            'version': version,
            'created': datetime.now().strftime('%Y-%m-%d'),
            'translate_prompt': translate_prompt,
            'proofread_prompt': proofread_prompt
        }
        
        # Save to file
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(prompt_data, f, indent=2, ensure_ascii=False)
            
            # Add to loaded prompts
            prompt_data['_filename'] = filename
            prompt_data['_filepath'] = filepath
            prompt_data['_type'] = prompt_type
            self.prompts[filename] = prompt_data
            
            self.log(f"✓ Created new prompt: {name}")
            return True
            
        except Exception as e:
            self.log(f"✗ Failed to create prompt: {e}")
            messagebox.showerror("Save Error", f"Failed to save prompt:\n{e}")
            return False
    
    def update_prompt(self, filename, name, description, domain, translate_prompt, 
                     proofread_prompt="", version="1.0"):
        """Update an existing prompt"""
        if filename not in self.prompts:
            self.log(f"✗ Prompt not found: {filename}")
            return False
        
        filepath = self.prompts[filename]['_filepath']
        
        # Update prompt data
        prompt_data = {
            'name': name,
            'description': description,
            'domain': domain,
            'version': version,
            'created': self.prompts[filename].get('created', datetime.now().strftime('%Y-%m-%d')),
            'modified': datetime.now().strftime('%Y-%m-%d'),
            'translate_prompt': translate_prompt,
            'proofread_prompt': proofread_prompt
        }
        
        # Save to file
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(prompt_data, f, indent=2, ensure_ascii=False)
            
            # Update loaded prompts
            prompt_data['_filename'] = filename
            prompt_data['_filepath'] = filepath
            prompt_data['_type'] = self.prompts[filename].get('_type', 'system_prompt')
            self.prompts[filename] = prompt_data
            
            self.log(f"✓ Updated prompt: {name}")
            return True
            
        except Exception as e:
            self.log(f"✗ Failed to update prompt: {e}")
            messagebox.showerror("Save Error", f"Failed to update prompt:\n{e}")
            return False
    
    def delete_prompt(self, filename):
        """Delete a custom prompt"""
        if filename not in self.prompts:
            return False
        
        filepath = self.prompts[filename]['_filepath']
        prompt_name = self.prompts[filename].get('name', filename)
        
        try:
            os.remove(filepath)
            del self.prompts[filename]
            
            # Clear active if this was active
            if self.active_prompt and self.active_prompt.get('_filename') == filename:
                self.clear_active_prompt()
            
            self.log(f"✓ Deleted prompt: {prompt_name}")
            return True
            
        except Exception as e:
            self.log(f"✗ Failed to delete prompt: {e}")
            messagebox.showerror("Delete Error", f"Failed to delete prompt:\n{e}")
            return False
    
    def export_prompt(self, filename, export_path):
        """Export a prompt to a specific location"""
        if filename not in self.prompts:
            return False
        
        try:
            source = self.prompts[filename]['_filepath']
            shutil.copy2(source, export_path)
            self.log(f"✓ Exported prompt to: {export_path}")
            return True
        except Exception as e:
            self.log(f"✗ Export failed: {e}")
            return False
    
    def import_prompt(self, import_path, prompt_type="system_prompt"):
        """Import a prompt from an external file
        
        Args:
            import_path: Path to JSON file to import
            prompt_type: Either 'system_prompt' or 'custom_instruction'
        """
        try:
            with open(import_path, 'r', encoding='utf-8') as f:
                prompt_data = json.load(f)
            
            # Validate
            if 'name' not in prompt_data or 'translate_prompt' not in prompt_data:
                messagebox.showerror("Invalid Prompt", "Missing required fields: name, translate_prompt")
                return False
            
            # Copy to appropriate directory based on type
            filename = os.path.basename(import_path)
            if prompt_type == "custom_instruction":
                directory = self.custom_instructions_dir
            else:  # system_prompt
                directory = self.system_prompts_dir
            dest_path = os.path.join(directory, filename)
            
            shutil.copy2(import_path, dest_path)
            
            # Add metadata and load
            prompt_data['_filename'] = filename
            prompt_data['_filepath'] = dest_path
            prompt_data['_type'] = prompt_type
            self.prompts[filename] = prompt_data
            
            self.log(f"✓ Imported prompt: {prompt_data['name']}")
            return True
            
        except Exception as e:
            self.log(f"✗ Import failed: {e}")
            messagebox.showerror("Import Error", f"Failed to import prompt:\n{e}")
            return False


# --- Translation Memory Architecture ---

class TM:
    """Individual Translation Memory with metadata"""
    
    def __init__(self, name: str, tm_id: str, enabled: bool = True, read_only: bool = False):
        self.name = name
        self.tm_id = tm_id
        self.enabled = enabled
        self.read_only = read_only
        self.entries: Dict[str, str] = {}  # source -> target mapping
        self.metadata = {
            'source_lang': None,
            'target_lang': None,
            'file_path': None,
            'created': datetime.now().isoformat(),
            'modified': datetime.now().isoformat()
        }
        self.fuzzy_threshold = 0.75
    
    def add_entry(self, source: str, target: str):
        """Add translation pair to this TM"""
        if not self.read_only and source and target:
            self.entries[source.strip()] = target.strip()
            self.metadata['modified'] = datetime.now().isoformat()
    
    def get_exact_match(self, source: str) -> Optional[str]:
        """Get exact match from this TM"""
        return self.entries.get(source.strip())
    
    def calculate_similarity(self, text1: str, text2: str) -> float:
        """Calculate similarity ratio between two texts"""
        return SequenceMatcher(None, text1.lower(), text2.lower()).ratio()
    
    def get_fuzzy_matches(self, source: str, max_matches: int = 5) -> List[Dict]:
        """Get fuzzy matches from this TM"""
        source = source.strip()
        matches = []
        
        for tm_source, tm_target in self.entries.items():
            similarity = self.calculate_similarity(source, tm_source)
            if similarity >= self.fuzzy_threshold:
                matches.append({
                    'source': tm_source,
                    'target': tm_target,
                    'similarity': similarity,
                    'match_pct': int(similarity * 100),
                    'tm_name': self.name,
                    'tm_id': self.tm_id
                })
        
        matches.sort(key=lambda x: x['similarity'], reverse=True)
        return matches[:max_matches]
    
    def get_entry_count(self) -> int:
        """Get number of entries in this TM"""
        return len(self.entries)
    
    def to_dict(self) -> Dict:
        """Serialize TM to dictionary for JSON storage"""
        return {
            'name': self.name,
            'tm_id': self.tm_id,
            'enabled': self.enabled,
            'read_only': self.read_only,
            'entries': self.entries,
            'metadata': self.metadata,
            'fuzzy_threshold': self.fuzzy_threshold
        }
    
    @staticmethod
    def from_dict(data: Dict) -> 'TM':
        """Deserialize TM from dictionary"""
        tm = TM(
            name=data.get('name', 'Unnamed TM'),
            tm_id=data.get('tm_id', 'unknown'),
            enabled=data.get('enabled', True),
            read_only=data.get('read_only', False)
        )
        tm.entries = data.get('entries', {})
        tm.metadata = data.get('metadata', {})
        tm.fuzzy_threshold = data.get('fuzzy_threshold', 0.75)
        return tm


class TMDatabase:
    """Manages multiple Translation Memories"""
    
    def __init__(self):
        # Core TMs
        self.project_tm = TM(name='Project TM', tm_id='project', enabled=True, read_only=False)
        self.big_mama_tm = TM(name='Big Mama', tm_id='big_mama', enabled=True, read_only=False)
        
        # Custom TMs (user-loaded TMX files)
        self.custom_tms: Dict[str, TM] = {}
        
        # Global fuzzy threshold (can be overridden per TM)
        self.fuzzy_threshold = 0.75
    
    def get_tm(self, tm_id: str) -> Optional[TM]:
        """Get TM by ID"""
        if tm_id == 'project':
            return self.project_tm
        elif tm_id == 'big_mama' or tm_id == 'main':  # Support legacy 'main' ID
            return self.big_mama_tm
        else:
            return self.custom_tms.get(tm_id)
    
    def get_all_tms(self, enabled_only: bool = False) -> List[TM]:
        """Get all TMs (optionally only enabled ones)"""
        tms = [self.project_tm, self.big_mama_tm] + list(self.custom_tms.values())
        if enabled_only:
            tms = [tm for tm in tms if tm.enabled]
        return tms
    
    def add_custom_tm(self, name: str, tm_id: str = None, read_only: bool = False) -> TM:
        """Add a new custom TM"""
        if tm_id is None:
            tm_id = f"custom_{len(self.custom_tms)}"
        tm = TM(name=name, tm_id=tm_id, enabled=True, read_only=read_only)
        self.custom_tms[tm_id] = tm
        return tm
    
    def remove_custom_tm(self, tm_id: str) -> bool:
        """Remove a custom TM"""
        if tm_id in self.custom_tms:
            del self.custom_tms[tm_id]
            return True
        return False
    
    def search_all(self, source: str, tm_ids: List[str] = None, enabled_only: bool = True) -> List[Dict]:
        """
        Search across multiple TMs
        Args:
            source: Source text to search for
            tm_ids: Specific TM IDs to search (None = search all)
            enabled_only: Only search enabled TMs
        Returns:
            List of match dictionaries sorted by similarity
        """
        all_matches = []
        
        # Determine which TMs to search
        if tm_ids:
            tms = [self.get_tm(tm_id) for tm_id in tm_ids if self.get_tm(tm_id)]
        else:
            tms = self.get_all_tms(enabled_only=enabled_only)
        
        # Search each TM
        for tm in tms:
            if tm and (not enabled_only or tm.enabled):
                matches = tm.get_fuzzy_matches(source, max_matches=10)
                all_matches.extend(matches)
        
        # Sort by similarity (highest first)
        all_matches.sort(key=lambda x: x['similarity'], reverse=True)
        return all_matches
    
    def add_to_project_tm(self, source: str, target: str):
        """Add entry to Project TM (convenience method)"""
        self.project_tm.add_entry(source, target)
    
    def get_entry_count(self, enabled_only: bool = False) -> int:
        """Get total entry count across all TMs"""
        tms = self.get_all_tms(enabled_only=enabled_only)
        return sum(tm.get_entry_count() for tm in tms)
    
    def load_tmx_file(self, filepath: str, src_lang: str, tgt_lang: str, 
                      tm_name: str = None, read_only: bool = False) -> tuple[str, int]:
        """
        Load TMX file into a new custom TM
        Returns: (tm_id, entry_count)
        """
        if tm_name is None:
            tm_name = os.path.basename(filepath).replace('.tmx', '')
        
        # Create new custom TM
        tm_id = f"custom_{os.path.basename(filepath).replace('.', '_')}"
        tm = self.add_custom_tm(tm_name, tm_id, read_only=read_only)
        
        # Load TMX content
        loaded_count = self._load_tmx_into_tm(filepath, src_lang, tgt_lang, tm)
        
        # Update metadata
        tm.metadata['file_path'] = filepath
        tm.metadata['source_lang'] = src_lang
        tm.metadata['target_lang'] = tgt_lang
        
        return tm_id, loaded_count
    
    def _load_tmx_into_tm(self, filepath: str, src_lang: str, tgt_lang: str, tm: TM) -> int:
        """Internal: Load TMX content into specific TM"""
        loaded_count = 0
        
        try:
            tree = ET.parse(filepath)
            root = tree.getroot()
            xml_ns = "http://www.w3.org/XML/1998/namespace"
            
            # Normalize language codes
            src_lang = src_lang.split('-')[0].split('_')[0].lower()
            tgt_lang = tgt_lang.split('-')[0].split('_')[0].lower()
            
            for tu in root.findall('.//tu'):
                src_text, tgt_text = None, None
                
                for tuv_node in tu.findall('tuv'):
                    lang_attr = tuv_node.get(f'{{{xml_ns}}}lang')
                    if not lang_attr:
                        continue
                    
                    tmx_lang = lang_attr.split('-')[0].split('_')[0].lower()
                    
                    seg_node = tuv_node.find('seg')
                    if seg_node is not None:
                        try:
                            text = ET.tostring(seg_node, encoding='unicode', method='text').strip()
                        except:
                            text = "".join(seg_node.itertext()).strip()
                        
                        if tmx_lang == src_lang:
                            src_text = text
                        elif tmx_lang == tgt_lang:
                            tgt_text = text
                
                if src_text and tgt_text:
                    tm.add_entry(src_text, tgt_text)
                    loaded_count += 1
            
            return loaded_count
        except Exception as e:
            print(f"Error loading TMX: {e}")
            return 0
    
    def detect_tmx_languages(self, filepath: str) -> List[str]:
        """Detect all language codes present in a TMX file"""
        try:
            tree = ET.parse(filepath)
            root = tree.getroot()
            xml_ns = "http://www.w3.org/XML/1998/namespace"
            
            languages = set()
            for tuv in root.findall('.//tuv'):
                lang_attr = tuv.get(f'{{{xml_ns}}}lang')
                if lang_attr:
                    languages.add(lang_attr)
            
            return sorted(list(languages))
        except:
            return []
    
    def to_dict(self) -> Dict:
        """Serialize entire database to dictionary"""
        return {
            'project_tm': self.project_tm.to_dict(),
            'big_mama_tm': self.big_mama_tm.to_dict(),
            'custom_tms': {tm_id: tm.to_dict() for tm_id, tm in self.custom_tms.items()},
            'fuzzy_threshold': self.fuzzy_threshold
        }
    
    @staticmethod
    def from_dict(data: Dict) -> 'TMDatabase':
        """Deserialize database from dictionary"""
        db = TMDatabase()
        
        if 'project_tm' in data:
            db.project_tm = TM.from_dict(data['project_tm'])
        if 'big_mama_tm' in data:
            db.big_mama_tm = TM.from_dict(data['big_mama_tm'])
        elif 'main_tm' in data:  # Legacy support
            db.big_mama_tm = TM.from_dict(data['main_tm'])
            db.big_mama_tm.name = 'Big Mama'  # Update name
            db.big_mama_tm.tm_id = 'big_mama'
        if 'custom_tms' in data:
            db.custom_tms = {tm_id: TM.from_dict(tm_data) 
                            for tm_id, tm_data in data['custom_tms'].items()}
        db.fuzzy_threshold = data.get('fuzzy_threshold', 0.75)
        
        return db


# Legacy TMAgent for backwards compatibility
class TMAgent:
    """Wrapper for backwards compatibility - delegates to TMDatabase"""
    
    def __init__(self):
        self.tm_database = TMDatabase()
        self.fuzzy_threshold = 0.75
    
    @property
    def tm_data(self):
        """Legacy property - returns Project TM entries"""
        return self.tm_database.project_tm.entries
    
    @tm_data.setter
    def tm_data(self, value: Dict[str, str]):
        """Legacy property setter"""
        self.tm_database.project_tm.entries = value
    
    def add_entry(self, source: str, target: str):
        """Add to Project TM"""
        self.tm_database.add_to_project_tm(source, target)
    
    def get_exact_match(self, source: str) -> Optional[str]:
        """Search all enabled TMs for exact match"""
        matches = self.tm_database.search_all(source, enabled_only=True)
        for match in matches:
            if match['match_pct'] == 100:
                return match['target']
        return None
    
    def get_fuzzy_matches(self, source: str, max_matches: int = 5) -> List[Tuple[str, str, float]]:
        """Legacy format - returns tuples"""
        matches = self.tm_database.search_all(source, enabled_only=True)
        return [(m['source'], m['target'], m['similarity']) for m in matches[:max_matches]]
    
    def get_best_match(self, source: str) -> Optional[Tuple[str, str, float]]:
        """Get best match in legacy format"""
        matches = self.get_fuzzy_matches(source, max_matches=1)
        return matches[0] if matches else None
    
    def load_from_tmx(self, filepath: str, src_lang: str = "en", tgt_lang: str = "nl") -> int:
        """Legacy TMX load - loads into a new custom TM"""
        tm_id, count = self.tm_database.load_tmx_file(filepath, src_lang, tgt_lang)
        return count
    
    def get_entry_count(self) -> int:
        """Get total entry count"""
        return self.tm_database.get_entry_count(enabled_only=False)
    
    def clear(self):
        """Clear Project TM only"""
        self.tm_database.project_tm.entries.clear()


# Model definitions (fallbacks if API fetch fails)
GEMINI_MODELS = [
    "gemini-2.5-pro-preview-05-06",
    "gemini-2.5-flash-preview-05-06",
    "gemini-1.5-pro-latest",
    "gemini-1.5-flash-latest"
]

CLAUDE_MODELS = [
    "claude-3-5-sonnet-20241022",
    "claude-3-5-haiku-20241022",
    "claude-3-opus-20240229"
]

OPENAI_MODELS = [
    "gpt-4o",
    "gpt-4o-mini",
    "gpt-4-turbo",
    "gpt-4"
]


def fetch_available_models(provider: str, api_key: str) -> List[str]:
    """Fetch available models from API provider.
    
    Args:
        provider: 'openai', 'claude', or 'gemini'
        api_key: API key for the provider
        
    Returns:
        List of available model names, or fallback list if fetch fails
    """
    if not api_key or not api_key.strip():
        return []
    
    try:
        if provider == "openai" and OPENAI_AVAILABLE:
            client = openai.OpenAI(api_key=api_key)
            models_response = client.models.list()
            available_models = [model.id for model in models_response.data]
            # Filter for chat models only
            chat_models = [m for m in available_models if 'gpt' in m.lower()]
            # Sort by preference (4o, 4o-mini, 4-turbo, etc.)
            preferred_order = ['gpt-4o', 'gpt-4o-mini', 'gpt-4-turbo', 'gpt-4', 'gpt-3.5-turbo']
            sorted_models = []
            for pref in preferred_order:
                matching = [m for m in chat_models if m.startswith(pref)]
                sorted_models.extend(sorted(matching, reverse=True))
            # Add any remaining models
            remaining = [m for m in chat_models if m not in sorted_models]
            sorted_models.extend(sorted(remaining, reverse=True))
            return sorted_models if sorted_models else OPENAI_MODELS
            
        elif provider == "claude" and ANTHROPIC_AVAILABLE:
            # Anthropic doesn't have a models.list() endpoint
            # Return known models
            return CLAUDE_MODELS
            
        elif provider == "gemini" and GEMINI_AVAILABLE:
            genai.configure(api_key=api_key)
            models_response = genai.list_models()
            available_models = []
            for model in models_response:
                if 'generateContent' in model.supported_generation_methods:
                    model_name = model.name.replace('models/', '')
                    available_models.append(model_name)
            return available_models if available_models else GEMINI_MODELS
            
    except Exception as e:
        print(f"Warning: Could not fetch models for {provider}: {e}")
        # Return fallback models
        if provider == "openai":
            return OPENAI_MODELS
        elif provider == "claude":
            return CLAUDE_MODELS
        elif provider == "gemini":
            return GEMINI_MODELS
    
    return []


class LayoutMode:
    """Layout mode constants"""
    GRID = "grid"      # memoQ-style inline editing
    SPLIT = "split"    # List view with editor panel
    DOCUMENT = "document" # Document flow view with clickable segments


class Segment:
    """Represents a translation segment"""
    
    def __init__(self, seg_id: int, source: str, paragraph_id: int = 0, 
                 is_table_cell: bool = False, table_info: tuple = None,
                 style: str = None, document_position: int = 0):
        self.id = seg_id
        self.source = source
        self.target = ""
        self.status = "untranslated"  # untranslated, translated, approved
        self.locked = False  # Lock status to prevent editing
        self.paragraph_id = paragraph_id
        self.document_position = document_position  # Position in original document
        self.notes = ""
        self.modified = False
        self.created_at = datetime.now().isoformat()
        self.modified_at = datetime.now().isoformat()
        
        # Table information
        self.is_table_cell = is_table_cell
        self.table_info = table_info  # (table_idx, row_idx, cell_idx) if is_table_cell
        
        # Style information (Heading 1, Normal, Title, etc.)
        self.style = style or "Normal"
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for JSON serialization"""
        return {
            'id': self.id,
            'source': self.source,
            'target': self.target,
            'status': self.status,
            'locked': self.locked,
            'paragraph_id': self.paragraph_id,
            'document_position': self.document_position,
            'notes': self.notes,
            'modified': self.modified,
            'created_at': self.created_at,
            'modified_at': self.modified_at,
            'is_table_cell': self.is_table_cell,
            'table_info': self.table_info,
            'style': self.style
        }
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'Segment':
        """Create Segment from dictionary"""
        seg = cls(data['id'], data['source'], data.get('paragraph_id', 0),
                  data.get('is_table_cell', False), data.get('table_info'),
                  data.get('style', 'Normal'), data.get('document_position', 0))
        seg.target = data.get('target', '')
        seg.status = data.get('status', 'untranslated')
        seg.locked = data.get('locked', False)
        seg.notes = data.get('notes', '')
        seg.modified = data.get('modified', False)
        seg.created_at = data.get('created_at', datetime.now().isoformat())
        seg.modified_at = data.get('modified_at', datetime.now().isoformat())
        return seg


class Supervertaler:
    """Main Supervertaler application - AI-Powered CAT Tool"""
    
    def __init__(self, root):
        self.root = root
        self.root.title("Supervertaler v3.0.0-beta - AI-Powered CAT Tool")
        self.root.geometry("1200x800")
        
        # Layout mode
        self.layout_mode = LayoutMode.GRID  # Default to memoQ-style
        self.current_edit_widget = None  # For inline editing in Grid mode
        
        # Data
        self.segments: List[Segment] = []
        self.current_segment: Optional[Segment] = None
        self.project_file: Optional[str] = None
        self.original_docx: Optional[str] = None
        self.modified = False
        
        # Multi-selection state
        self.selected_segments: Set[int] = set()  # Set of selected segment IDs
        self.last_selected_index: Optional[int] = None  # For Shift+Click range selection
        
        # Filter state
        self.filtered_segments = []
        self.filter_active = False
        
        # Components
        self.segmenter = SimpleSegmenter()
        self.docx_handler = DOCXHandler()
        self.tag_manager = TagManager()
        
        # Default language list (user-editable)
        self.available_languages = [
            "Afrikaans", "Albanian", "Arabic", "Armenian", "Basque", "Bengali",
            "Bulgarian", "Catalan", "Chinese (Simplified)", "Chinese (Traditional)",
            "Croatian", "Czech", "Danish", "Dutch", "English", "Estonian",
            "Finnish", "French", "Galician", "Georgian", "German", "Greek",
            "Hebrew", "Hindi", "Hungarian", "Icelandic", "Indonesian", "Irish",
            "Italian", "Japanese", "Korean", "Latvian", "Lithuanian", "Macedonian",
            "Malay", "Norwegian", "Persian", "Polish", "Portuguese", "Romanian",
            "Russian", "Serbian", "Slovak", "Slovenian", "Spanish", "Swahili",
            "Swedish", "Thai", "Turkish", "Ukrainian", "Urdu", "Vietnamese", "Welsh"
        ]
        
        # LLM settings
        self.api_keys = API_KEYS.copy()  # Local copy of API keys
        self.current_llm_provider = "openai"  # Default provider
        self.current_llm_model = "gpt-4o"  # Default model
        self.source_language = "English"
        self.target_language = "Dutch"
        
        # Translation prompts - context-aware for different modes
        # Single segment translation (Ctrl+T)
        self.single_segment_prompt = (
            "You are an expert {{SOURCE_LANGUAGE}} to {{TARGET_LANGUAGE}} translator with deep understanding of context and nuance.\n\n"
            "**CONTEXT**: Full document context is provided for reference below.\n\n"
            "**YOUR TASK**: Translate ONLY the text in the 'TEXT TO TRANSLATE' section.\n\n"
            "**IMPORTANT INSTRUCTIONS**:\n"
            "- Provide ONLY the translated text\n"
            "- Do NOT include numbering, labels, or commentary\n"
            "- Do NOT repeat the source text\n"
            "- Maintain accuracy and natural fluency\n\n"
            "**CRITICAL: CAT TOOL TAG PRESERVATION**:\n"
            "- Source may contain CAT tool formatting tags in various formats:\n"
            "  • memoQ: [1}, {2], [3}, {4] (asymmetric bracket-brace pairs)\n"
            "  • Trados Studio: <410>text</410>, <434>text</434> (XML-style opening/closing tags)\n"
            "  • CafeTran: |formatted text| (pipe symbols mark formatted text - bold, italic, underline, etc.)\n"
            "  • Other CAT tools: various bracketed or special character sequences\n"
            "- These are placeholder tags representing formatting (bold, italic, links, etc.)\n"
            "- PRESERVE ALL tags - if source has N tags, target must have exactly N tags\n"
            "- Keep tags with their content and adjust position for natural target language word order\n"
            "- Never translate, omit, or modify the tags themselves - only reposition them\n"
            "- Examples:\n"
            "  • memoQ: '[1}De uitvoer{2]' → '[1}The exports{2]'\n"
            "  • Trados: '<410>De uitvoer van machines</410>' → '<410>Exports of machinery</410>'\n"
            "  • CafeTran: 'He debuted against |Juventus FC| in 2001' → 'Hij debuteerde tegen |Juventus FC| in 2001'\n"
            "  • Multiple: '[1}De uitvoer{2] [3}stelt niets voor{4]' → '[1}Exports{2] [3}mean nothing{4]'\n\n"
            "**LANGUAGE-SPECIFIC NUMBER FORMATTING**:\n"
            "- If the target language is **Dutch**, **French**, **German**, **Italian**, **Spanish**, or another **continental European language**, use a **comma** as the decimal separator and a **space or non-breaking space** between the number and unit (e.g., 17,1 cm).\n"
            "- If the target language is **English** or **Irish**, use a **full stop (period)** as the decimal separator and **no space** before the unit (e.g., 17.1 cm).\n"
            "- Always follow the **number formatting conventions** of the target language.\n\n"
            "If the text refers to figures (e.g., 'Figure 1A'), relevant images may be provided for visual context.\n\n"
            "{{SOURCE_LANGUAGE}} text:\n{{SOURCE_TEXT}}"
        )
        
        # Batch DOCX translation
        self.batch_docx_prompt = (
            "You are an expert {{SOURCE_LANGUAGE}} to {{TARGET_LANGUAGE}} translator specializing in document translation.\n\n"
            "**YOUR TASK**: Translate ALL segments below while maintaining document structure and formatting.\n\n"
            "**IMPORTANT INSTRUCTIONS**:\n"
            "- Translate each segment completely and accurately\n"
            "- Preserve paragraph breaks and structure\n"
            "- Maintain consistent terminology throughout\n"
            "- Consider document-wide context for accuracy\n"
            "- Output translations in the same order as source segments\n\n"
            "**CRITICAL: CAT TOOL TAG PRESERVATION**:\n"
            "- Segments may contain CAT tool formatting tags in various formats:\n"
            "  • memoQ: [1}, {2] | Trados: <410>text</410> | CafeTran: |formatted text|\n"
            "- These are placeholder tags representing formatting - preserve ALL of them\n"
            "- Keep tags with their content, repositioning as needed for natural target language structure\n"
            "- Never translate, omit, or modify the tags themselves\n"
            "- Example: '<410>De uitvoer</410> <434>stelt niets voor</434>' → '<410>Exports</410> <434>mean nothing</434>'\n"
            "- Example: '[1}Text{2]' → '[1}Translation{2]'\n\n"
            "**LANGUAGE-SPECIFIC NUMBER FORMATTING**:\n"
            "- If the target language is **Dutch**, **French**, **German**, **Italian**, **Spanish**, or another **continental European language**, use a **comma** as the decimal separator and a **space or non-breaking space** between the number and unit (e.g., 17,1 cm).\n"
            "- If the target language is **English** or **Irish**, use a **full stop (period)** as the decimal separator and **no space** before the unit (e.g., 17.1 cm).\n"
            "- Always follow the **number formatting conventions** of the target language.\n\n"
            "{{SOURCE_LANGUAGE}} text:\n{{SOURCE_TEXT}}"
        )
        
        # Batch bilingual (TXT/memoQ export) translation
        self.batch_bilingual_prompt = (
            "You are an expert {{SOURCE_LANGUAGE}} to {{TARGET_LANGUAGE}} translator working with a bilingual translation file.\n\n"
            "**YOUR TASK**: Translate each source segment below.\n\n"
            "**FILE FORMAT**: This is a bilingual export (e.g., from memoQ) where each segment is numbered.\n\n"
            "**IMPORTANT INSTRUCTIONS**:\n"
            "- Translate each numbered segment\n"
            "- Maintain segment numbering in your output\n"
            "- Keep translations aligned with source segment numbers\n"
            "- Ensure consistency across all segments\n\n"
            "**CRITICAL: CAT TOOL TAG PRESERVATION**:\n"
            "- Segments often contain CAT tool formatting tags in various formats:\n"
            "  • memoQ: [1}, {2] | Trados: <410>text</410> | CafeTran: |formatted text|\n"
            "- These are placeholder tags representing formatting (bold, italic, links, etc.)\n"
            "- You MUST preserve ALL tags - if source has 4 tags, target must have 4 tags\n"
            "- Keep tags with the content they wrap, repositioning if sentence structure requires it\n"
            "- Never translate, omit, or modify the tags - only reposition appropriately\n"
            "- Example: '<410>De uitvoer van machines</410> <434>stelt niets voor</434>' → '<410>Exports of machinery</410> <434>mean nothing</434>'\n"
            "- Never translate or omit tags - only move them to appropriate positions for target language\n"
            "- Example: '[1}De uitvoer van de USSR naar de BLEU{2]' → '[1}USSR exports to the BLEU{2]'\n"
            "- Example: '[1}De uitvoer van machines{2] [3}stelt niets voor{4]' → '[1}Exports of machinery{2] [3}mean nothing{4]'\n"
            "- Empty tag pairs like '[3} {4]' must also be preserved\n\n"
            "**LANGUAGE-SPECIFIC NUMBER FORMATTING**:\n"
            "- If the target language is **Dutch**, **French**, **German**, **Italian**, **Spanish**, or another **continental European language**, use a **comma** as the decimal separator and a **space or non-breaking space** between the number and unit (e.g., 17,1 cm).\n"
            "- If the target language is **English** or **Irish**, use a **full stop (period)** as the decimal separator and **no space** before the unit (e.g., 17.1 cm).\n"
            "- Always follow the **number formatting conventions** of the target language.\n\n"
            "{{SOURCE_LANGUAGE}} text:\n{{SOURCE_TEXT}}"
        )
        
        # Default translation prompt (backwards compatibility)
        self.default_translate_prompt = self.single_segment_prompt
        
        self.default_proofread_prompt = (
            "You are an expert proofreader and editor for {{SOURCE_LANGUAGE}} → {{TARGET_LANGUAGE}} translations, skilled in various document types and domains.\n\n"
            "For each segment you receive a SOURCE SEGMENT and EXISTING TRANSLATION. "
            "Your tasks: improve accuracy, ensure terminology consistency, enhance readability, correct grammar, improve fluency, verify completeness, and maintain consistency with visual elements.\n\n"
            "**CRITICAL: CAT TOOL TAG PRESERVATION**:\n"
            "- Source/target may contain CAT tool formatting tags in various formats:\n"
            "  • memoQ: [1}, {2] | Trados: <410>text</410> | CafeTran: |formatted text|\n"
            "- These are placeholder tags representing formatting - preserve ALL of them\n"
            "- Keep tags with their content, repositioning as needed for natural target language structure\n"
            "- Never translate, omit, or modify the tags themselves\n"
            "- Example: '<410>The exports</410>' remains '<410>The exports</410>'\n\n"
            "**LANGUAGE-SPECIFIC NUMBER FORMATTING**:\n"
            "- If the target language is **Dutch**, **French**, **German**, **Italian**, **Spanish**, or another **continental European language**, use a **comma** as the decimal separator and a **space or non-breaking space** between the number and unit (e.g., 17,1 cm).\n"
            "- If the target language is **English** or **Irish**, use a **full stop (period)** as the decimal separator and **no space** before the unit (e.g., 17.1 cm).\n"
            "- Always follow the **number formatting conventions** of the target language.\n\n"
            "OUTPUT FORMAT STRICTLY:\n"
            "1) Numbered list of revised {{TARGET_LANGUAGE}} translations (use same numbering; if no changes needed, reproduce the original).\n"
            "2) Then a section:\n---CHANGES SUMMARY START---\n"
            "Per modified line: '<line>. <brief description of changes>' OR if none changed: 'No changes made to any segment in this batch.'\n"
            "---CHANGES SUMMARY END---"
        )
        
        # Store current prompts (initially same as defaults)
        self.current_translate_prompt = self.default_translate_prompt
        self.current_proofread_prompt = self.default_proofread_prompt
        
        # System prompts directory (uses path resolver for dev mode support)
        self.system_prompts_dir = get_user_data_path("System_prompts")
        os.makedirs(self.system_prompts_dir, exist_ok=True)
        
        # Translation memory - new multi-TM architecture
        self.tm_database = TMDatabase()
        self.tm_agent = TMAgent()  # Legacy wrapper for backward compatibility
        self.tm_agent.tm_database = self.tm_database  # Share same database
        self.translation_memory: List[Dict[str, str]] = []  # Deprecated - keeping for backward compatibility
        self.tm_auto_search_timer = None  # Timer for automatic TM search
        
        # Tracked changes agent (learns from editing patterns)
        self.tracked_changes_agent = TrackedChangesAgent(log_callback=self.log)
        
        # Prompt library (system prompts for domain-specific translation)
        self.prompt_library = PromptLibrary(self.system_prompts_dir, log_callback=self.log)
        
        # TMX Generator for export
        self.tmx_generator = TMXGenerator(log_callback=self.log)
        
        # Drawings/Images support for multimodal translation
        self.drawings_images_map = {}  # Normalized figure name -> PIL Image
        self.drawings_folder = None  # Path to folder containing drawing images
        
        # Setup UI
        self.setup_ui()
        
        # Load prompts after UI is ready (so logging works)
        self.prompt_library.load_all_prompts()
        
        # Status
        self.log("Supervertaler v3.0.0-beta ready. Import a DOCX file to begin.")
        self.log(f"✨ LLM APIs: OpenAI={OPENAI_AVAILABLE}, Claude={ANTHROPIC_AVAILABLE}, Gemini={GEMINI_AVAILABLE}")
        self.log("✨ Layout modes available: Grid (memoQ-style), List, Document")
    
    def get_context_aware_prompt(self, mode: str = "single") -> str:
        """Get the appropriate translation prompt based on context.
        
        Args:
            mode: Translation mode - 'single', 'batch_docx', or 'batch_bilingual'
            
        Returns:
            The appropriate prompt template for the given context
        """
        # If user has selected a custom prompt, use that
        if hasattr(self, 'current_translate_prompt') and self.current_translate_prompt != self.single_segment_prompt:
            return self.current_translate_prompt
        
        # Otherwise, select based on mode
        if mode == "single":
            return self.single_segment_prompt
        elif mode == "batch_docx":
            return self.batch_docx_prompt
        elif mode == "batch_bilingual":
            return self.batch_bilingual_prompt
        else:
            return self.single_segment_prompt  # Default fallback
    
    def setup_ui(self):
        """Create the user interface"""
        
        # Menu bar
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # Project menu - Project management (moved to first position)
        project_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Project", menu=project_menu)
        project_menu.add_command(label="Save Project", command=self.save_project, accelerator="Ctrl+S")
        project_menu.add_command(label="Save Project As...", command=self.save_project_as)
        project_menu.add_command(label="Open Project...", command=self.load_project, accelerator="Ctrl+L")
        project_menu.add_command(label="Close Project", command=self.close_project)
        project_menu.add_separator()
        project_menu.add_command(label="API Settings...", command=self.show_api_settings)
        project_menu.add_command(label="Language Settings...", command=self.show_language_settings)
        
        # File menu - Import/Export focus
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        
        # Import submenu
        import_submenu = tk.Menu(file_menu, tearoff=0)
        file_menu.add_cascade(label="Import", menu=import_submenu)
        import_submenu.add_command(label="DOCX (Monolingual)...", command=self.import_docx, accelerator="Ctrl+O")
        import_submenu.add_command(label="TXT (Mono/Bilingual)...", command=self.import_txt_bilingual)
        import_submenu.add_command(label="memoQ DOCX (Bilingual)...", command=self.import_memoq_bilingual)
        import_submenu.add_command(label="CafeTran DOCX (Bilingual)...", command=self.import_cafetran_bilingual)
        import_submenu.add_command(label="Trados Studio DOCX (Bilingual)...", command=self.import_trados_bilingual)
        
        # Export submenu
        export_submenu = tk.Menu(file_menu, tearoff=0)
        file_menu.add_cascade(label="Export", menu=export_submenu)
        export_submenu.add_command(label="DOCX (Monolingual)...", command=self.export_docx)
        export_submenu.add_command(label="DOCX (Bilingual)...", command=self.export_bilingual_docx)
        export_submenu.add_command(label="TMX...", command=self.export_tmx)
        export_submenu.add_command(label="TSV...", command=self.export_tsv)
        export_submenu.add_separator()
        export_submenu.add_command(label="memoQ DOCX...", command=self.export_memoq_bilingual)
        export_submenu.add_command(label="CafeTran DOCX...", command=self.export_cafetran_bilingual)
        export_submenu.add_command(label="Trados Studio DOCX...", command=self.export_trados_bilingual)
        export_submenu.add_command(label="TXT (Bilingual)...", command=self.export_txt_bilingual)
        export_submenu.add_separator()
        export_submenu.add_command(label="Session Report...", command=self.generate_session_report)
        
        file_menu.add_separator()
        
        # Recent projects submenu (placeholder)
        recent_menu = tk.Menu(file_menu, tearoff=0)
        file_menu.add_cascade(label="Recent Projects", menu=recent_menu, state='disabled')
        
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.on_closing)
        
        # Edit menu
        edit_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Edit", menu=edit_menu)
        edit_menu.add_command(label="Find/Replace...", command=self.show_find_replace, accelerator="Ctrl+F")
        edit_menu.add_separator()
        edit_menu.add_command(label="Copy Source to Target", command=self.copy_source_to_target, accelerator="Ctrl+D")
        edit_menu.add_command(label="Copy Source to Target (All Segments)", command=self.copy_source_to_target_all)
        edit_menu.add_command(label="Clear Target", command=self.clear_target)
        edit_menu.add_separator()
        
        # Bulk operations submenu
        bulk_menu = tk.Menu(edit_menu, tearoff=0)
        edit_menu.add_cascade(label="Bulk Operations", menu=bulk_menu)
        bulk_menu.add_command(label="Select All Segments", command=self.select_all_segments, accelerator="Ctrl+A")
        bulk_menu.add_command(label="Clear All Targets...", command=self.clear_all_targets)
        bulk_menu.add_separator()
        bulk_menu.add_command(label="Change Status (All)...", command=self.change_status_all)
        bulk_menu.add_command(label="Change Status (Filtered)...", command=self.change_status_filtered)
        bulk_menu.add_separator()
        bulk_menu.add_command(label="Lock All Segments", command=self.lock_all_segments)
        bulk_menu.add_command(label="Unlock All Segments", command=self.unlock_all_segments)
        bulk_menu.add_command(label="Lock Filtered Segments", command=self.lock_filtered_segments)
        bulk_menu.add_command(label="Unlock Filtered Segments", command=self.unlock_filtered_segments)
        edit_menu.add_separator()
        
        # Segment operations submenu
        segment_menu = tk.Menu(edit_menu, tearoff=0)
        edit_menu.add_cascade(label="Segment", menu=segment_menu)
        segment_menu.add_command(label="Translate Current", command=self.translate_current_segment, accelerator="Ctrl+T")
        segment_menu.add_command(label="Translate All Untranslated", command=self.translate_all_untranslated)
        segment_menu.add_separator()
        segment_menu.add_command(label="Lock Current Segment", command=self.lock_current_segment)
        segment_menu.add_command(label="Unlock Current Segment", command=self.unlock_current_segment)
        
        # View menu - Layout & Assistant panel
        view_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="View", menu=view_menu)
        view_menu.add_command(label="Grid View", command=lambda: self.switch_layout(LayoutMode.GRID), accelerator="Ctrl+1")
        view_menu.add_command(label="List View", command=lambda: self.switch_layout(LayoutMode.SPLIT), accelerator="Ctrl+2")
        view_menu.add_command(label="Document View", command=lambda: self.switch_layout(LayoutMode.DOCUMENT), accelerator="Ctrl+3")
        view_menu.add_separator()
        view_menu.add_command(label="Grid Columns...", command=self.show_column_visibility_dialog)
        view_menu.add_command(label="Toggle Style Colors", command=self.toggle_grid_style_colors)
        view_menu.add_separator()
        
        # Assistant panel layout submenu
        resources_layout_menu = tk.Menu(view_menu, tearoff=0)
        view_menu.add_cascade(label="Assistant panel", menu=resources_layout_menu)
        resources_layout_menu.add_command(label="Tabbed Layout", command=lambda: self.switch_assistance_layout('tabbed'))
        resources_layout_menu.add_command(label="Stacked Layout", command=lambda: self.switch_assistance_layout('stacked'))
        resources_layout_menu.add_command(label="List Layout", command=lambda: self.switch_assistance_layout('list'))
        
        view_menu.add_separator()
        view_menu.add_command(label="Toggle Filter Mode", command=self.toggle_filter_mode, accelerator="Ctrl+M")
        view_menu.add_command(label="Apply Filters", command=self.apply_filters, accelerator="Ctrl+Shift+A")
        view_menu.add_command(label="Clear Filters", command=self.clear_filters)
        view_menu.add_command(label="Focus Filter", command=self.focus_filter_source, accelerator="Ctrl+Shift+F")
        
        # Resources menu - TM, Glossaries, etc.
        resources_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Resources", menu=resources_menu)
        resources_menu.add_command(label="Translation Memory...", command=self.show_tm_manager)
        resources_menu.add_command(label="Load TM File...", command=self.load_tm_file)
        resources_menu.add_separator()
        resources_menu.add_command(label="📝 Load Tracked Changes (DOCX)...", command=self.load_tracked_changes_docx)
        resources_menu.add_command(label="🗑️ Clear Tracked Changes", command=self.clear_tracked_changes)
        resources_menu.add_separator()
        resources_menu.add_command(label="🎨 Load Drawing Images...", command=self.load_drawings)
        resources_menu.add_command(label="🗑️ Clear Drawings", command=self.clear_drawings)
        resources_menu.add_separator()
        resources_menu.add_command(label="📚 Prompt Library", command=self.show_custom_prompts, accelerator="Ctrl+P")
        resources_menu.add_command(label="🎭 System Prompts", command=self.show_system_prompts)
        resources_menu.add_command(label="📝 Custom Instructions", command=self.show_custom_instructions)
        
        # Help menu
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Help", menu=help_menu)
        help_menu.add_command(label="User Guide", command=self.show_user_guide)
        help_menu.add_command(label="Changelog", command=self.show_changelog)
        help_menu.add_separator()
        help_menu.add_command(label="About", command=self.show_about)
        
        # Keyboard shortcuts (translate)
        self.root.bind('<Control-t>', lambda e: self.translate_current_segment())
        self.root.bind('<Control-p>', lambda e: self.show_custom_prompts())  # Prompt Library shortcut
        self.root.bind('<Control-a>', lambda e: self.select_all_segments())  # Select All Segments
        
        # Keyboard shortcuts
        self.root.bind('<Control-o>', lambda e: self.import_docx())
        self.root.bind('<Control-s>', lambda e: self.save_project())
        self.root.bind('<Control-l>', lambda e: self.load_project())
        self.root.bind('<Control-f>', lambda e: self.show_find_replace())
        self.root.bind('<Control-d>', lambda e: self.copy_source_to_target())
        
        # Layout switching shortcuts
        self.root.bind('<Control-Key-1>', lambda e: self.switch_layout(LayoutMode.GRID))
        self.root.bind('<Control-Key-2>', lambda e: self.switch_layout(LayoutMode.SPLIT))
        self.root.bind('<Control-Key-3>', lambda e: self.switch_layout(LayoutMode.DOCUMENT))
        
        # Filter shortcuts
        self.root.bind('<Control-m>', lambda e: self.toggle_filter_mode())  # Toggle filter mode
        self.root.bind('<Control-Shift-F>', lambda e: self.focus_filter_source())  # Focus source filter
        self.root.bind('<Control-Shift-A>', lambda e: self.apply_filters())  # Apply filters
        
        # Navigation shortcuts
        self.root.bind('<Control-Down>', lambda e: self.navigate_segment('next'))
        self.root.bind('<Control-Up>', lambda e: self.navigate_segment('prev'))
        
        # Grid View display options
        self.grid_style_colors_enabled = True  # Toggle for style-based font colors
        
        # Dual text selection state (for Grid View)
        self.dual_selection_row = None  # Currently active row index
        self.dual_selection_source = None  # Source Text widget with selection
        self.dual_selection_target = None  # Target Text widget with selection
        self.dual_selection_focused_widget = None  # 'source' or 'target' for keyboard selection
        self.dual_selection_source_cursor = None  # Remember cursor position in source
        self.dual_selection_target_cursor = None  # Remember cursor position in target
        
        # F2 to enter edit mode (works globally when segment is selected)
        self.root.bind('<F2>', lambda e: self.enter_edit_mode_global())
        
        # Toolbar
        self.toolbar = tk.Frame(self.root, bg='#f0f0f0', height=40)
        self.toolbar.pack(side='top', fill='x', padx=5, pady=5)
        
        # Import dropdown button
        import_btn = tk.Menubutton(self.toolbar, text="Import ▼", 
                                   bg='#4CAF50', fg='white', padx=15, relief='raised',
                                   activebackground='#45a049')
        import_btn.pack(side='left', padx=2)
        import_menu = tk.Menu(import_btn, tearoff=0)
        import_btn.config(menu=import_menu)
        import_menu.add_command(label="DOCX (Monolingual)", command=self.import_docx)
        import_menu.add_command(label="TXT (Mono/Bilingual)", command=self.import_txt_bilingual)
        import_menu.add_command(label="memoQ DOCX (Bilingual)", command=self.import_memoq_bilingual)
        import_menu.add_command(label="CafeTran DOCX (Bilingual)", command=self.import_cafetran_bilingual)
        import_menu.add_command(label="Trados Studio DOCX (Bilingual)", command=self.import_trados_bilingual)
        
        # Translate dropdown button (workflow: import → translate → export → save)
        translate_btn = tk.Menubutton(self.toolbar, text="Translate ▼", 
                                      bg='#9C27B0', fg='white', padx=15, relief='raised',
                                      activebackground='#7B1FA2')
        translate_btn.pack(side='left', padx=2)
        translate_menu = tk.Menu(translate_btn, tearoff=0)
        translate_btn.config(menu=translate_menu)
        translate_menu.add_command(label="⚡ Translate Current Segment (Ctrl+T)", 
                                  command=self.translate_current_segment)
        translate_menu.add_command(label="🚀 Translate All Untranslated", 
                                  command=self.translate_all_untranslated)
        
        # Export dropdown button
        export_btn = tk.Menubutton(self.toolbar, text="Export ▼", 
                                   bg='#FF9800', fg='white', padx=15, relief='raised',
                                   activebackground='#f57c00')
        export_btn.pack(side='left', padx=2)
        export_menu = tk.Menu(export_btn, tearoff=0)
        export_btn.config(menu=export_menu)
        export_menu.add_command(label="DOCX (Monolingual)", command=self.export_docx)
        export_menu.add_command(label="DOCX (Bilingual)", command=self.export_bilingual_docx)
        export_menu.add_command(label="TMX", command=self.export_tmx)
        export_menu.add_command(label="TSV", command=self.export_tsv)
        export_menu.add_command(label="memoQ DOCX", command=self.export_memoq_bilingual)
        export_menu.add_command(label="CafeTran DOCX", command=self.export_cafetran_bilingual)
        export_menu.add_command(label="Trados Studio DOCX", command=self.export_trados_bilingual)
        export_menu.add_separator()
        export_menu.add_command(label="Session Report", command=self.generate_session_report)
        
        # Save Project button (single, frequently used)
        tk.Button(self.toolbar, text="Save", command=self.save_project,
                 bg='#2196F3', fg='white', padx=15).pack(side='left', padx=2)
        
        ttk.Separator(self.toolbar, orient='vertical').pack(side='left', fill='y', padx=10)
        
        # Layout mode buttons (neutral colors, toggle style)
        self.layout_btn_grid = tk.Button(self.toolbar, text="Grid", 
                                         command=lambda: self.switch_layout(LayoutMode.GRID),
                                         bg='#e0e0e0', padx=12, relief='sunken')
        self.layout_btn_grid.pack(side='left', padx=1)
        
        self.layout_btn_split = tk.Button(self.toolbar, text="List",
                                          command=lambda: self.switch_layout(LayoutMode.SPLIT),
                                          bg='#f5f5f5', padx=12, relief='raised')
        self.layout_btn_split.pack(side='left', padx=1)
        
        self.layout_btn_document = tk.Button(self.toolbar, text="Document",
                                             command=lambda: self.switch_layout(LayoutMode.DOCUMENT),
                                             bg='#f5f5f5', padx=12, relief='raised')
        self.layout_btn_document.pack(side='left', padx=1)
        
        ttk.Separator(self.toolbar, orient='vertical').pack(side='left', fill='y', padx=10)
        
        # Tools dropdown button
        tools_btn = tk.Menubutton(self.toolbar, text="Tools ▼", 
                                  padx=12, relief='raised')
        tools_btn.pack(side='left', padx=2)
        tools_menu = tk.Menu(tools_btn, tearoff=0)
        tools_btn.config(menu=tools_menu)
        tools_menu.add_command(label="Find/Replace", command=self.show_find_replace)
        tools_menu.add_command(label="Prompt Library", command=self.show_custom_prompts)
        tools_menu.add_separator()
        tools_menu.add_command(label="API Settings", command=self.show_api_settings)
        tools_menu.add_command(label="Language Settings", command=self.show_language_settings)
        
        # Progress info
        self.progress_label = tk.Label(self.toolbar, text="No document loaded", bg='#f0f0f0')
        self.progress_label.pack(side='right', padx=10)
        
        # Add dev mode indicator if enabled (subtle, right-aligned in toolbar)
        if ENABLE_PRIVATE_FEATURES:
            dev_mode_label = tk.Label(self.toolbar, text="DEV MODE", 
                                     font=("Segoe UI", 8, "bold"), 
                                     fg="white", bg="#dc3545", 
                                     padx=8, pady=2, relief='raised', bd=1)
            dev_mode_label.pack(side='right', padx=5)
        
        # Use PanedWindow to make log resizable
        self.main_paned = ttk.PanedWindow(self.root, orient='vertical')
        self.main_paned.pack(fill='both', expand=True, padx=5, pady=5)
        
        # Main content area (top pane)
        self.content_frame = tk.Frame(self.main_paned)
        self.main_paned.add(self.content_frame, weight=1)
        
        # Create the appropriate layout
        self.create_layout_ui()
        
        # Log/Status area (bottom pane - resizable by dragging sash upward)
        log_frame = tk.LabelFrame(self.main_paned, text="Log (↕ drag to resize)", padx=5, pady=5)
        self.main_paned.add(log_frame, weight=0)
        
        self.log_text = scrolledtext.ScrolledText(log_frame, height=4, wrap='word',
                                                  font=('Consolas', 9), state='disabled')
        self.log_text.pack(fill='both', expand=True)
        
        # Add visual grip to the sash for better discoverability
        # The sash is positioned between the panes, slightly above the log frame border
        self.root.after(100, self.add_sash_grip)
        
        # Handle window close
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
    
    def create_layout_ui(self):
        """Create UI based on current layout mode"""
        # Show Start Screen if no document is loaded
        if not self.segments:
            self.create_start_screen()
        elif self.layout_mode == LayoutMode.GRID:
            self.create_grid_layout()
        elif self.layout_mode == LayoutMode.SPLIT:
            self.create_split_layout()
        elif self.layout_mode == LayoutMode.DOCUMENT:
            self.create_document_layout()
    
    def create_start_screen(self):
        """Create Start Screen (Splash Screen) shown when no document is loaded"""
        
        # Create main horizontal paned window (splash screen on left, assistance panel on right)
        self.start_paned = ttk.PanedWindow(self.content_frame, orient='horizontal')
        self.start_paned.pack(fill='both', expand=True)
        
        # Left side: Splash Screen
        splash_container = tk.Frame(self.start_paned, bg='#f5f5f5')
        self.start_paned.add(splash_container, weight=3)
        
        # Center content vertically
        center_frame = tk.Frame(splash_container, bg='#f5f5f5')
        center_frame.place(relx=0.5, rely=0.5, anchor='center')
        
        # Store original logo for resizing
        self.splash_logo_original = None
        self.splash_logo_label = None
        self.splash_center_frame = center_frame
        self.splash_container = splash_container
        
        # Load and display logo
        try:
            from PIL import Image, ImageTk
            logo_path = os.path.join(os.path.dirname(__file__), 'assets', 'supervertaler_icon_colours.png')
            if os.path.exists(logo_path):
                self.splash_logo_original = Image.open(logo_path)
                # Initial size based on window (will be updated dynamically)
                initial_size = min(300, int(splash_container.winfo_width() * 0.4)) if splash_container.winfo_width() > 1 else 250
                logo_image = self.splash_logo_original.resize((initial_size, initial_size), Image.Resampling.LANCZOS)
                logo_photo = ImageTk.PhotoImage(logo_image)
                self.splash_logo_label = tk.Label(center_frame, image=logo_photo, bg='#f5f5f5')
                self.splash_logo_label.image = logo_photo  # Keep reference
                self.splash_logo_label.pack(pady=(0, 30))
                
                # Bind resize event to update logo size
                splash_container.bind('<Configure>', self.resize_splash_logo)
        except Exception as e:
            self.log(f"⚠ Could not load logo: {e}")
        
        # Program title
        tk.Label(center_frame, text="Supervertaler", 
                font=('Segoe UI', 32, 'bold'), bg='#f5f5f5', fg='#2196F3').pack(pady=(0, 5))
        
        tk.Label(center_frame, text="v3.3.0-beta", 
                font=('Segoe UI', 14), bg='#f5f5f5', fg='#666').pack(pady=(0, 20))
        
        # Subtitle
        tk.Label(center_frame, text="AI-Powered Computer-Assisted Translation Tool", 
                font=('Segoe UI', 12), bg='#f5f5f5', fg='#333').pack(pady=(0, 30))
        
        # Quick start buttons
        button_frame = tk.Frame(center_frame, bg='#f5f5f5')
        button_frame.pack(pady=20)
        
        btn_style = {
            'font': ('Segoe UI', 9),
            'width': 18,
            'relief': 'flat',
            'cursor': 'hand2',
            'padx': 12,
            'pady': 10
        }
        
        tk.Button(button_frame, text="📄 Open DOCX File", bg='#4CAF50', fg='white',
                 command=self.import_docx, **btn_style).pack(side='left', padx=5)
        
        tk.Button(button_frame, text="📂 Open Project", bg='#2196F3', fg='white',
                 command=self.load_project, **btn_style).pack(side='left', padx=5)
        
        tk.Button(button_frame, text="📥 Import Bilingual", bg='#FF9800', fg='white',
                 command=self.show_import_bilingual_menu, **btn_style).pack(side='left', padx=5)
        
        # Footer with info
        footer_frame = tk.Frame(splash_container, bg='#f5f5f5')
        footer_frame.pack(side='bottom', fill='x', pady=20)
        
        tk.Label(footer_frame, text="Use the Assistant panel on the right to configure settings →", 
                font=('Segoe UI', 10), bg='#f5f5f5', fg='#999').pack()
        
        # Right side: Assistance panel (with settings access)
        self.create_assistance_panel(parent_paned=self.start_paned)
    
    def show_import_bilingual_menu(self):
        """Show popup menu for bilingual import options"""
        menu = tk.Menu(self.root, tearoff=0, font=('Segoe UI', 10))
        menu.add_command(label="📋 memoQ DOCX", command=self.import_memoq_bilingual)
        menu.add_command(label="📋 CafeTran DOCX", command=self.import_cafetran_bilingual)
        menu.add_command(label="📋 Trados Studio DOCX", command=self.import_trados_bilingual)
        
        # Position menu at mouse cursor
        try:
            menu.post(self.root.winfo_pointerx(), self.root.winfo_pointery())
        except:
            menu.post(self.root.winfo_x() + 200, self.root.winfo_y() + 200)
    
    def resize_splash_logo(self, event=None):
        """Dynamically resize splash screen logo based on window size"""
        if not hasattr(self, 'splash_logo_original') or self.splash_logo_original is None:
            return
        
        if not hasattr(self, 'splash_logo_label') or self.splash_logo_label is None:
            return
        
        try:
            from PIL import ImageTk
            
            # Get container width
            container_width = self.splash_container.winfo_width()
            container_height = self.splash_container.winfo_height()
            
            # Calculate logo size (30-40% of container width, but between 200-400px)
            if container_width > 1 and container_height > 1:
                logo_size = max(200, min(400, int(container_width * 0.35)))
                
                # Resize logo
                resized_logo = self.splash_logo_original.resize((logo_size, logo_size), Image.Resampling.LANCZOS)
                logo_photo = ImageTk.PhotoImage(resized_logo)
                
                # Update label
                self.splash_logo_label.configure(image=logo_photo)
                self.splash_logo_label.image = logo_photo  # Keep reference
        except Exception as e:
            pass  # Silently ignore resize errors
    
    def switch_from_start_to_grid(self):
        """Switch from Start Screen to Grid View when document is loaded"""
        # Clear content frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()
        
        # Create grid layout (the import method will call load_segments_to_grid() after this)
        self.create_grid_layout()
    
    def create_grid_layout(self):
        """Create Grid View layout (memoQ-style with inline editing and dynamic row heights)"""
        
        # Create main horizontal paned window (grid on left, assistance panel on right)
        self.grid_paned = ttk.PanedWindow(self.content_frame, orient='horizontal')
        self.grid_paned.pack(fill='both', expand=True)
        
        # Left side: Grid and editor
        left_container = tk.Frame(self.grid_paned)
        self.grid_paned.add(left_container, weight=3)
        
        # Grid frame (top part - expandable)
        grid_frame = tk.LabelFrame(left_container, text="Translation Grid - Grid View (Click target to edit)", padx=5, pady=5)
        grid_frame.pack(side='top', fill='both', expand=True)
        
        # Filter panel (above the grid header)
        filter_frame = tk.Frame(grid_frame, bg='#f0f0f0', relief='ridge', borderwidth=1)
        filter_frame.pack(side='top', fill='x', pady=(0, 5))
        
        # Filter label
        tk.Label(filter_frame, text="🔍 Filter:", bg='#f0f0f0', 
                font=('Segoe UI', 9, 'bold')).pack(side='left', padx=5, pady=5)
        
        # Filter mode selection (radio-style buttons)
        self.filter_mode = 'filter'  # 'filter' or 'highlight'
        
        mode_frame = tk.Frame(filter_frame, bg='#f0f0f0', relief='solid', bd=1)
        mode_frame.pack(side='left', padx=5)
        
        self.filter_mode_btn = tk.Button(mode_frame, text="\U0001F50D Filter",
                                        command=lambda: self.set_filter_mode('filter'),
                                        bg='#4CAF50', fg='white', font=('Segoe UI', 9, 'bold'),
                                        relief='sunken', bd=2, cursor='hand2', width=10)
        self.filter_mode_btn.pack(side='left', padx=1, pady=1)
        
        self.highlight_mode_btn = tk.Button(mode_frame, text="💡 Highlight",
                                            command=lambda: self.set_filter_mode('highlight'),
                                            bg='#e0e0e0', fg='#666', font=('Segoe UI', 9),
                                            relief='raised', bd=2, cursor='hand2', width=10)
        self.highlight_mode_btn.pack(side='left', padx=1, pady=1)
        
        tk.Label(filter_frame, text="│", bg='#f0f0f0', fg='#ccc',
                font=('Segoe UI', 10)).pack(side='left', padx=5)
        
        # Source filter
        tk.Label(filter_frame, text="Source:", bg='#f0f0f0', 
                font=('Segoe UI', 9)).pack(side='left', padx=(10, 2))
        self.filter_source_var = tk.StringVar()
        source_filter_entry = tk.Entry(filter_frame, textvariable=self.filter_source_var,
                                      font=('Segoe UI', 9), width=25, relief='solid', borderwidth=1)
        source_filter_entry.pack(side='left', padx=2, pady=3)
        # Bind Enter key to apply filters
        source_filter_entry.bind('<Return>', lambda e: self.apply_filters())
        
        # Target filter
        tk.Label(filter_frame, text="Target:", bg='#f0f0f0', 
                font=('Segoe UI', 9)).pack(side='left', padx=(10, 2))
        self.filter_target_var = tk.StringVar()
        target_filter_entry = tk.Entry(filter_frame, textvariable=self.filter_target_var,
                                      font=('Segoe UI', 9), width=25, relief='solid', borderwidth=1)
        target_filter_entry.pack(side='left', padx=2, pady=3)
        # Bind Enter key to apply filters
        target_filter_entry.bind('<Return>', lambda e: self.apply_filters())
        
        # Status filter
        tk.Label(filter_frame, text="Status:", bg='#f0f0f0', 
                font=('Segoe UI', 9)).pack(side='left', padx=(10, 2))
        self.filter_status_var = tk.StringVar(value="All")
        status_filter_combo = ttk.Combobox(filter_frame, textvariable=self.filter_status_var,
                                          values=["All", "untranslated", "draft", "translated", "approved"],
                                          state='readonly', width=12, font=('Segoe UI', 9))
        status_filter_combo.pack(side='left', padx=2, pady=3)
        # Combobox selection triggers apply automatically (user expectation for dropdowns)
        status_filter_combo.bind('<<ComboboxSelected>>', lambda e: self.apply_filters())
        
        # Apply filters button
        tk.Button(filter_frame, text="🔍 Apply", command=self.apply_filters,
                 bg='#2196F3', fg='white', font=('Segoe UI', 9, 'bold'),
                 relief='raised', padx=10, pady=2, cursor='hand2').pack(side='left', padx=5)
        
        # Clear filters button
        tk.Button(filter_frame, text="✕ Clear", command=self.clear_filters,
                 bg='#ff9800', fg='white', font=('Segoe UI', 8),
                 relief='flat', padx=8, pady=2).pack(side='left', padx=5)
        
        # Filter results label
        self.filter_results_label = tk.Label(filter_frame, text="", bg='#f0f0f0',
                                            font=('Segoe UI', 9), fg='#666')
        self.filter_results_label.pack(side='left', padx=10)
        
        # Column configuration with adjustable source/target widths
        self.fixed_columns_width = 200  # ID + Type + Status combined
        self.source_width = 450  # Initial source width
        self.target_width = 450  # Initial target width
        self.dragging_splitter = False
        
        self.grid_columns = {
            'id': {'title': '#', 'width': 35, 'anchor': 'center', 'visible': True},
            'type': {'title': 'Type', 'width': 50, 'anchor': 'center', 'visible': True},
            'style': {'title': 'Style', 'width': 70, 'anchor': 'center', 'visible': True},
            'source': {'title': '📄 Source', 'width': self.source_width, 'anchor': 'w', 'visible': True},
            'target': {'title': '🎯 Target', 'width': self.target_width, 'anchor': 'w', 'visible': True},
            'status': {'title': '✓', 'width': 30, 'anchor': 'center', 'visible': True}  # Icon column, after target
        }
        
        # Initialize filter state
        self.filtered_segments = []
        self.filter_active = False
        
        # Initialize pagination
        self.grid_page_size = 50  # Segments per page
        self.grid_current_page = 0  # 0-indexed
        
        # Create STICKY header row (outside canvas, fixed at top)
        self.header_container = tk.Frame(grid_frame, bg='white')
        self.header_container.pack(side='top', fill='x')
        self.create_grid_header()
        
        # Create pagination controls
        self.pagination_frame = tk.Frame(grid_frame, bg='#f8f8f8', height=35)
        self.pagination_frame.pack(side='top', fill='x', pady=(0, 2))
        self.pagination_frame.pack_propagate(False)
        
        # Pagination label (left side)
        self.pagination_label = tk.Label(self.pagination_frame, 
                                         text="Segments 1-50 of 355",
                                         font=('Segoe UI', 9),
                                         bg='#f8f8f8', fg='#555')
        self.pagination_label.pack(side='left', padx=10)
        
        # Selection counter (center-left) - shows how many segments are selected
        self.selection_counter_label = tk.Label(self.pagination_frame,
                                                text="",
                                                font=('Segoe UI', 9, 'bold'),
                                                bg='#f8f8f8', fg='#2196F3')  # Blue color
        self.selection_counter_label.pack(side='left', padx=15)
        
        # Pagination controls (right side)
        pagination_controls = tk.Frame(self.pagination_frame, bg='#f8f8f8')
        pagination_controls.pack(side='right', padx=10)
        
        # First page button
        self.first_page_btn = tk.Button(pagination_controls, text='⏮ First', 
                                        command=self.go_to_first_page,
                                        font=('Segoe UI', 8),
                                        relief='flat', bd=1,
                                        bg='white', fg='#333',
                                        padx=8, pady=2)
        self.first_page_btn.pack(side='left', padx=2)
        
        # Previous page button
        self.prev_page_btn = tk.Button(pagination_controls, text='◀ Prev', 
                                       command=self.go_to_prev_page,
                                       font=('Segoe UI', 8),
                                       relief='flat', bd=1,
                                       bg='white', fg='#333',
                                       padx=8, pady=2)
        self.prev_page_btn.pack(side='left', padx=2)
        
        # Page number input
        page_input_frame = tk.Frame(pagination_controls, bg='#f8f8f8')
        page_input_frame.pack(side='left', padx=5)
        
        tk.Label(page_input_frame, text='Page:', font=('Segoe UI', 8),
                bg='#f8f8f8', fg='#555').pack(side='left', padx=(0, 3))
        
        self.page_number_var = tk.StringVar(value='1')
        self.page_number_entry = tk.Entry(page_input_frame, 
                                          textvariable=self.page_number_var,
                                          width=5, font=('Segoe UI', 9),
                                          justify='center')
        self.page_number_entry.pack(side='left')
        self.page_number_entry.bind('<Return>', lambda e: self.go_to_page())
        
        self.total_pages_label = tk.Label(page_input_frame, text='of 8',
                                          font=('Segoe UI', 8),
                                          bg='#f8f8f8', fg='#555')
        self.total_pages_label.pack(side='left', padx=(3, 0))
        
        # Next page button
        self.next_page_btn = tk.Button(pagination_controls, text='Next ▶', 
                                       command=self.go_to_next_page,
                                       font=('Segoe UI', 8),
                                       relief='flat', bd=1,
                                       bg='white', fg='#333',
                                       padx=8, pady=2)
        self.next_page_btn.pack(side='left', padx=2)
        
        # Last page button
        self.last_page_btn = tk.Button(pagination_controls, text='Last ⏭', 
                                       command=self.go_to_last_page,
                                       font=('Segoe UI', 8),
                                       relief='flat', bd=1,
                                       bg='white', fg='#333',
                                       padx=8, pady=2)
        self.last_page_btn.pack(side='left', padx=2)
        
        # Page size selector (far right)
        pagesize_frame = tk.Frame(self.pagination_frame, bg='#f8f8f8')
        pagesize_frame.pack(side='right', padx=(0, 10))
        
        tk.Label(pagesize_frame, text='Per page:', font=('Segoe UI', 8),
                bg='#f8f8f8', fg='#555').pack(side='left', padx=(0, 3))
        
        self.page_size_var = tk.StringVar(value='50')
        page_size_combo = ttk.Combobox(pagesize_frame, 
                                      textvariable=self.page_size_var,
                                      values=['25', '50', '100', '200', 'All'],
                                      width=6, state='readonly',
                                      font=('Segoe UI', 8))
        page_size_combo.pack(side='left')
        page_size_combo.bind('<<ComboboxSelected>>', self.on_page_size_changed)
        
        # Create scrollable content area
        content_area = tk.Frame(grid_frame, bg='white')
        content_area.pack(side='top', fill='both', expand=True)
        
        # Create custom grid canvas for rows
        self.grid_canvas = tk.Canvas(content_area, bg='white', highlightthickness=0)
        self.grid_canvas.pack(side='left', fill='both', expand=True)
        
        # Scrollbars
        v_scroll = ttk.Scrollbar(content_area, orient='vertical', command=self.grid_canvas.yview)
        h_scroll = ttk.Scrollbar(content_area, orient='horizontal', command=self.grid_canvas.xview)
        self.grid_canvas.configure(yscrollcommand=v_scroll.set, xscrollcommand=h_scroll.set)
        
        v_scroll.pack(side='right', fill='y')
        h_scroll.pack(side='bottom', fill='x')
        
        # Create a frame inside canvas to hold all rows
        self.grid_inner_frame = tk.Frame(self.grid_canvas, bg='white')
        self.canvas_window = self.grid_canvas.create_window((0, 0), window=self.grid_inner_frame, anchor='nw')
        
        # Storage for row widgets
        self.grid_rows = []  # List of row data: {'segment': seg, 'widgets': {...}, 'row_frame': frame}
        self.current_row_index = -1
        
        # Bind canvas resize to update scroll region
        self.grid_inner_frame.bind('<Configure>', self.update_grid_scroll_region)
        self.grid_canvas.bind('<Configure>', self.on_grid_canvas_resize)
        
        # Keyboard bindings on canvas
        self.grid_canvas.bind('<Double-1>', self.on_grid_double_click)
        self.grid_canvas.bind('<Button-1>', self.on_grid_click)
        self.grid_canvas.bind('<F2>', self.enter_edit_mode)
        self.grid_canvas.bind('<Return>', lambda e: self.enter_edit_mode())
        self.grid_canvas.bind('<Control-d>', lambda e: self.copy_source_to_target())
        
        # Mouse wheel scrolling
        self.grid_canvas.bind('<MouseWheel>', self.on_grid_mousewheel)
        
        # Context menu
        self.create_context_menu()
        self.grid_canvas.bind('<Button-3>', self.show_grid_context_menu)
        
        # Focus the canvas so keyboard events work
        self.grid_canvas.focus_set()
        
        # Store reference to container for editor panel
        self.grid_left_container = left_container
        
        # Editor panel at bottom (hideable)
        self.grid_editor_visible = True
        self.create_grid_editor_panel()
        
        # Right side: Assistance panel (MT, TM, Glossary, etc.)
        self.create_assistance_panel()
    
    def create_grid_editor_panel(self):
        """Create the editor panel for Grid View"""
        # Toggle button above the editor panel
        toggle_btn_frame = tk.Frame(self.grid_left_container, bg='#f0f0f0')
        toggle_btn_frame.pack(side='bottom', fill='x', pady=(2, 0))
        
        self.grid_editor_toggle_btn = tk.Button(toggle_btn_frame, 
                                               text="🔽 Hide Editor Panel", 
                                               command=self.toggle_grid_editor,
                                               font=('Segoe UI', 9),
                                               relief='raised',
                                               bg='#4CAF50',
                                               fg='white',
                                               padx=10,
                                               pady=2)
        self.grid_editor_toggle_btn.pack(side='left', padx=5, pady=2)
        
        # Editor frame (bottom part) - compact padding for better fit
        self.grid_editor_frame = tk.LabelFrame(self.grid_left_container, text="Segment Editor", padx=5, pady=5)
        self.grid_editor_frame.pack(side='bottom', fill='x', pady=(0, 0))
        
        # Segment info and status on same row
        info_frame = tk.Frame(self.grid_editor_frame)
        info_frame.pack(fill='x', pady=(0, 5))
        
        self.grid_seg_info_label = tk.Label(info_frame, text="No segment selected",
                                           font=('Segoe UI', 9, 'bold'))
        self.grid_seg_info_label.pack(side='left')
        
        # Status selector
        status_frame = tk.Frame(info_frame)
        status_frame.pack(side='right')
        
        tk.Label(status_frame, text="Status:").pack(side='left', padx=(0, 5))
        self.grid_status_var = tk.StringVar(value="untranslated")
        grid_status_combo = ttk.Combobox(status_frame, textvariable=self.grid_status_var,
                                        values=["untranslated", "translated", "approved"],
                                        state='readonly', width=12)
        grid_status_combo.pack(side='left')
        grid_status_combo.bind('<<ComboboxSelected>>', self.on_grid_editor_status_change)
        
        # Source (read-only) - compact for better fit
        tk.Label(self.grid_editor_frame, text="Source:", font=('Segoe UI', 9, 'bold')).pack(anchor='w')
        self.grid_source_text = tk.Text(self.grid_editor_frame, height=1, wrap='word', bg='#f5f5f5',
                                       state='disabled', font=('Segoe UI', 9))
        self.grid_source_text.pack(fill='x', pady=(2, 5))
        
        # Target (editable) with validation label
        target_header_frame = tk.Frame(self.grid_editor_frame)
        target_header_frame.pack(fill='x', anchor='w')
        tk.Label(target_header_frame, text="Target:", font=('Segoe UI', 9, 'bold')).pack(side='left')
        self.grid_tag_validation_label = tk.Label(target_header_frame, text="", font=('Segoe UI', 8))
        self.grid_tag_validation_label.pack(side='left', padx=(10, 0))
        
        self.grid_target_text = tk.Text(self.grid_editor_frame, height=1, wrap='word', font=('Segoe UI', 9))
        self.grid_target_text.pack(fill='x', pady=(2, 3))
        self.grid_target_text.bind('<KeyRelease>', self.on_grid_target_change)
        self.grid_target_text.bind('<Control-Return>', lambda e: self.save_grid_editor_and_next())
        self.grid_target_text.bind('<Control-b>', lambda e: self.insert_tag_grid('b'))
        self.grid_target_text.bind('<Control-i>', lambda e: self.insert_tag_grid('i'))
        self.grid_target_text.bind('<Control-u>', lambda e: self.insert_tag_grid('u'))
        
        # Tag buttons - wrapping layout for narrow screens
        tag_button_frame = tk.Frame(self.grid_editor_frame)
        tag_button_frame.pack(fill='x', pady=(0, 3))
        
        # Use a Label frame for better organization and wrapping
        tag_container = tk.Frame(tag_button_frame, relief='flat')
        tag_container.pack(fill='x')
        
        # Row 1: Label and formatting buttons
        tag_row1 = tk.Frame(tag_container)
        tag_row1.pack(fill='x', pady=2)
        
        tk.Label(tag_row1, text="Insert:").pack(side='left', padx=(0, 5))
        tk.Button(tag_row1, text="<b>Bold</b>", command=lambda: self.insert_tag_grid('b'),
                 relief='flat', bg='#ffcccc', font=('Segoe UI', 8)).pack(side='left', padx=2)
        tk.Button(tag_row1, text="<i>Italic</i>", command=lambda: self.insert_tag_grid('i'),
                 relief='flat', bg='#ccccff', font=('Segoe UI', 8, 'italic')).pack(side='left', padx=2)
        tk.Button(tag_row1, text="<u>Underline</u>", command=lambda: self.insert_tag_grid('u'),
                 relief='flat', bg='#ccffcc', font=('Segoe UI', 8, 'underline')).pack(side='left', padx=2)
        tk.Button(tag_row1, text="Strip Tags", command=self.strip_tags_from_grid_target,
                 relief='flat', bg='#eeeeee', font=('Segoe UI', 8)).pack(side='left', padx=10)
        tk.Button(tag_row1, text="Copy Source Tags", command=self.copy_source_tags_grid,
                 relief='flat', bg='#e6f3ff', font=('Segoe UI', 8)).pack(side='left', padx=2)
        
        # Action buttons - wrapping layout for narrow screens
        button_frame = tk.Frame(self.grid_editor_frame)
        button_frame.pack(fill='x', pady=(0, 0))
        
        # Left side buttons
        left_buttons = tk.Frame(button_frame)
        left_buttons.pack(side='left', fill='x', expand=True)
        
        tk.Button(left_buttons, text="Copy Source → Target", command=self.copy_source_to_target_grid_editor,
                 font=('Segoe UI', 8)).pack(side='left', padx=(0, 5))
        tk.Button(left_buttons, text="Clear Target", command=self.clear_grid_target,
                 font=('Segoe UI', 8)).pack(side='left', padx=(0, 5))
        
        # Right side button (most important action)
        tk.Button(button_frame, text="Save & Next (Ctrl+Enter)", command=self.save_grid_editor_and_next,
                 bg='#4CAF50', fg='white', font=('Segoe UI', 8, 'bold')).pack(side='right')

    
    def toggle_grid_editor(self):
        """Toggle visibility of Grid View editor panel"""
        if self.grid_editor_visible:
            # Hide the editor
            self.grid_editor_frame.pack_forget()
            self.grid_editor_visible = False
            # Update toolbar button
            self.grid_editor_toggle_btn.config(text="🔼 Show Editor Panel", relief='raised', bg='#E0E0E0', fg='black')
        else:
            # Show the editor
            self.grid_editor_frame.pack(side='bottom', fill='x', pady=(5, 0))
            self.grid_editor_visible = True
            # Update toolbar button
            self.grid_editor_toggle_btn.config(text="🔽 Hide Editor Panel", relief='sunken', bg='#4CAF50', fg='white')
    
    def on_grid_editor_status_change(self, event):
        """Handle status change from grid editor panel"""
        if hasattr(self, 'current_segment') and self.current_segment:
            new_status = self.grid_status_var.get()
            self.current_segment.status = new_status
            self.current_segment.modified = True
            self.modified = True
            self.update_progress()
            # Update the grid row if visible
            if self.current_row_index >= 0:
                self.update_grid_row(self.current_row_index)
    
    def copy_source_to_target_grid_editor(self):
        """Copy source to target in grid editor panel"""
        if hasattr(self, 'current_segment') and self.current_segment:
            self.grid_target_text.delete('1.0', 'end')
            self.grid_target_text.insert('1.0', self.current_segment.source)
    
    def save_grid_editor_segment(self):
        """Save the current segment from grid editor panel"""
        if hasattr(self, 'current_segment') and self.current_segment:
            target = self.grid_target_text.get('1.0', 'end-1c').strip()
            status = self.grid_status_var.get()
            
            self.current_segment.target = target
            self.current_segment.status = status
            self.current_segment.modified = True
            self.modified = True
            
            # Add to Project TM if translated or approved and has content
            if status in ['translated', 'approved'] and target:
                self.tm_database.add_to_project_tm(self.current_segment.source, target)
                self.log(f"✓ Added to Project TM: {self.current_segment.source[:50]}...")
            
            self.update_progress()
            # Update the grid row
            if self.current_row_index >= 0:
                self.update_grid_row(self.current_row_index)
            self.log(f"✓ Segment #{self.current_segment.id} saved")
    
    def save_grid_editor_and_next(self):
        """Save current segment with 'translated' status and move to next untranslated segment"""
        # Set status to translated before saving
        self.grid_status_var.set('translated')
        self.save_grid_editor_segment()
        
        # Find next untranslated segment
        if not self.current_segment:
            return
        
        current_id = self.current_segment.id
        next_untranslated = None
        
        for seg in self.segments:
            if seg.id > current_id and seg.status == 'untranslated':
                next_untranslated = seg
                break
        
        # If found, navigate to it
        if next_untranslated:
            # Find the tree item with this segment ID
            for item in self.tree.get_children():
                values = self.tree.item(item, 'values')
                if int(values[0]) == next_untranslated.id:
                    self.tree.selection_set(item)
                    self.tree.see(item)
                    self.select_grid_row(self.tree.index(item))
                    break
        else:
            self.log("No more untranslated segments")
    
    def load_segment_to_grid_editor(self, segment):
        """Load a segment into the grid editor panel"""
        if not hasattr(self, 'grid_editor_frame'):
            return
        
        self.current_segment = segment
        self.grid_seg_info_label.config(text=f"Segment #{segment.id} | Paragraph {segment.paragraph_id}")
        
        # Source
        self.grid_source_text.config(state='normal')
        self.grid_source_text.delete('1.0', 'end')
        self.grid_source_text.insert('1.0', segment.source)
        self.grid_source_text.config(state='disabled')
        
        # Target
        self.grid_target_text.delete('1.0', 'end')
        if segment.target:
            self.grid_target_text.insert('1.0', segment.target)
        
        # Status
        self.grid_status_var.set(segment.status)
    
    def on_grid_target_change(self, event):
        """Handle target text changes in grid editor to validate tags"""
        if hasattr(self, 'current_segment') and self.current_segment:
            self.validate_tags_grid()
    
    def validate_tags_grid(self):
        """Validate HTML tags in grid editor target text"""
        if not hasattr(self, 'current_segment') or not self.current_segment:
            return
        
        source = self.current_segment.source
        target = self.grid_target_text.get('1.0', 'end-1c')
        
        # Extract tags from source and target
        source_tags = re.findall(r'<[^>]+>', source)
        target_tags = re.findall(r'<[^>]+>', target)
        
        # Compare
        if source_tags == target_tags:
            self.grid_tag_validation_label.config(text="✓ Tags match", fg='#4CAF50')
        elif not source_tags and not target_tags:
            self.grid_tag_validation_label.config(text="", fg='#666')
        else:
            self.grid_tag_validation_label.config(text="⚠ Tags differ from source", fg='#FF9800')
    
    def clear_grid_target(self):
        """Clear target text in grid editor"""
        self.grid_target_text.delete('1.0', 'end')
        self.grid_tag_validation_label.config(text="", fg='#666')
    
    def strip_tags_from_grid_target(self):
        """Remove all HTML tags from grid editor target text"""
        target = self.grid_target_text.get('1.0', 'end-1c')
        cleaned = re.sub(r'<[^>]+>', '', target)
        self.grid_target_text.delete('1.0', 'end')
        self.grid_target_text.insert('1.0', cleaned)
        self.validate_tags_grid()
    
    def copy_source_tags_grid(self):
        """Copy HTML tags from source to target in grid editor, preserving target text"""
        if not hasattr(self, 'current_segment') or not self.current_segment:
            return
        
        source = self.current_segment.source
        target = self.grid_target_text.get('1.0', 'end-1c')
        
        # Extract tags from source
        source_tags = re.findall(r'<[^>]+>', source)
        
        if not source_tags:
            self.log("No tags found in source")
            return
        
        # Remove existing tags from target
        clean_target = re.sub(r'<[^>]+>', '', target)
        
        # Try to intelligently place tags
        # For now, just wrap the target with the same tags
        result = clean_target
        for tag in source_tags:
            if not tag.startswith('</'):
                # Opening tag - add at start
                result = tag + result
            else:
                # Closing tag - add at end
                result = result + tag
        
        self.grid_target_text.delete('1.0', 'end')
        self.grid_target_text.insert('1.0', result)
        self.validate_tags_grid()
        self.log("✓ Tags copied from source")
    
    def create_assistance_panel(self, parent_paned=None):
        """Create the right-side assistance panel with dockable/stackable panes"""
        # Use provided parent or fall back to grid_paned for backward compatibility
        if parent_paned is None:
            parent_paned = self.grid_paned
        
        # Right panel container
        right_container = tk.Frame(parent_paned, bg='#f9f9f9')
        parent_paned.add(right_container, weight=1)
        
        # Header with controls
        header_frame = tk.Frame(right_container, bg='#e0e0e0', height=35)
        header_frame.pack(side='top', fill='x')
        header_frame.pack_propagate(False)
        
        tk.Label(header_frame, text="Assistant panel", font=('Segoe UI', 10, 'bold'),
                bg='#e0e0e0').pack(side='left', padx=10, pady=5)
        
        # Toolbar buttons (middle section)
        toolbar_frame = tk.Frame(header_frame, bg='#e0e0e0')
        toolbar_frame.pack(side='left', padx=20, pady=2)
        
        tk.Button(toolbar_frame, text="🧪 Preview Prompt", 
                 command=self.preview_combined_prompt,
                 font=('Segoe UI', 8), bg='#2196F3', fg='white',
                 relief='raised', padx=8, pady=2,
                 cursor='hand2').pack(side='left', padx=2)
        
        # Refresh tabs button (right side, before layout toggle)
        tk.Button(header_frame, text="🔄 Refresh Tabs",
                 command=self.refresh_assist_tabs,
                 font=('Segoe UI', 8), bg='#9C27B0', fg='white',
                 relief='raised', padx=8, pady=2,
                 cursor='hand2').pack(side='right', padx=2)
        
        # Layout mode button (right side)
        self.assist_layout_btn = tk.Button(header_frame, text="⊞ Stacked View",
                                          command=self.toggle_assistance_layout,
                                          font=('Segoe UI', 8), bg='#4CAF50', fg='white',
                                          relief='raised', padx=8, pady=2)
        self.assist_layout_btn.pack(side='right', padx=5)
        
        # Container for dynamic content (will switch between tabbed and stacked)
        self.assist_content_frame = tk.Frame(right_container, bg='#f9f9f9')
        self.assist_content_frame.pack(fill='both', expand=True)
        
        # Track which panels are visible
        self.assist_visible_panels = {
            'projects': True,          # Project Library
            'system_prompts': True,    # System Prompt Library
            'custom_instructions': True, # Custom translation instructions
            'mt': True,                # Machine Translation suggestions
            'llm': True,               # LLM Translation
            'tm': True,                # Translation Memory (matches + management)
            'glossary': True,          # Glossary
            'reference_images': True,  # Reference images for context
            'nontrans': True,          # Non-translatables
            'settings': True,          # Translation Settings
            'tracked_changes': True,   # Post-Translation Analysis
            'log': True                # Session Log (synchronized with main log)
        }
        self.assist_layout_mode = 'tabbed'  # 'tabbed' or 'stacked'
        
        # Create initial tabbed layout
        self.create_tabbed_assistance()
        
        # Create detail view (always at bottom)
        self.create_detail_panel(right_container)
    
    def toggle_assistance_layout(self):
        """Toggle between tabbed and stacked layout for assistance panels"""
        if self.assist_layout_mode == 'tabbed':
            # Switch to stacked
            self.assist_layout_mode = 'stacked'
            self.assist_layout_btn.config(text="⊟ Tabbed View", bg='#FF9800')
            self.rebuild_assistance_layout()
        else:
            # Switch to tabbed
            self.assist_layout_mode = 'tabbed'
            self.assist_layout_btn.config(text="⊞ Stacked View", bg='#4CAF50')
            self.rebuild_assistance_layout()
    
    def refresh_assist_tabs(self):
        """Refresh tab layout by toggling to stacked and back to tabbed view"""
        if self.assist_layout_mode == 'tabbed':
            # Quick toggle to stacked and back to force tab reflow
            self.assist_layout_mode = 'stacked'
            self.rebuild_assistance_layout()
            self.assist_layout_mode = 'tabbed'
            self.rebuild_assistance_layout()
            self.log("✅ Tab layout refreshed")
    
    def rebuild_assistance_layout(self):
        """Rebuild the assistance panel layout"""
        # Clear current content
        for widget in self.assist_content_frame.winfo_children():
            widget.destroy()
        
        if self.assist_layout_mode == 'tabbed':
            self.create_tabbed_assistance()
        else:
            self.create_stacked_assistance()
    
    def create_tabbed_assistance(self):
        """Create tabbed layout with custom tab bar and overflow menu"""
        # Container for custom tab bar + content
        tab_container = tk.Frame(self.assist_content_frame, bg='#f0f0f0')
        tab_container.pack(fill='both', expand=True)
        
        # Custom tab bar at top
        self.assist_tab_bar = tk.Frame(tab_container, bg='#e0e0e0', height=32)
        self.assist_tab_bar.pack(side='top', fill='x')
        self.assist_tab_bar.pack_propagate(False)
        
        # Container for visible tab buttons
        self.assist_tab_buttons_frame = tk.Frame(self.assist_tab_bar, bg='#e0e0e0')
        self.assist_tab_buttons_frame.pack(side='left', fill='both', expand=True)
        
        # Overflow menu button (shown when tabs don't fit)
        self.assist_overflow_btn = tk.Menubutton(self.assist_tab_bar, 
                                                 text='More ▼',
                                                 font=('Segoe UI', 8),
                                                 bg='#d0d0d0',
                                                 fg='#333',
                                                 relief='raised',
                                                 bd=1,
                                                 padx=5,
                                                 pady=2,
                                                 cursor='hand2')
        self.assist_overflow_menu = tk.Menu(self.assist_overflow_btn, tearoff=0)
        self.assist_overflow_btn.config(menu=self.assist_overflow_menu)
        # Don't pack overflow button yet - only show when needed
        
        # Content area for tab content
        self.assist_content_area = tk.Frame(tab_container, bg='white')
        self.assist_content_area.pack(fill='both', expand=True)
        
        # Define all tabs with their full names and short names
        self.assist_tabs = []
        
        if self.assist_visible_panels.get('projects', True):
            self.assist_tabs.append({
                'key': 'projects',
                'name': '📁 Projects',
                'short': 'Projects',
                'frame': None,
                'button': None,
                'create_func': self.create_projects_tab
            })
        
        if self.assist_visible_panels.get('system_prompts', True):
            self.assist_tabs.append({
                'key': 'system_prompts',
                'name': '📝 System Prompts',
                'short': 'Prompts',
                'frame': None,
                'button': None,
                'create_func': self.create_prompts_tab
            })
        
        if self.assist_visible_panels.get('custom_instructions', True):
            self.assist_tabs.append({
                'key': 'custom_instructions',
                'name': '📋 Custom Instructions',
                'short': 'Custom',
                'frame': None,
                'button': None,
                'create_func': self.create_custom_instructions_tab
            })
        
        if self.assist_visible_panels.get('mt', True):
            self.assist_tabs.append({
                'key': 'mt',
                'name': '🤖 MT',
                'short': 'MT',
                'frame': None,
                'button': None,
                'create_func': self.create_mt_tab
            })
        
        if self.assist_visible_panels.get('llm', True):
            self.assist_tabs.append({
                'key': 'llm',
                'name': '✨ LLM',
                'short': 'LLM',
                'frame': None,
                'button': None,
                'create_func': self.create_llm_tab
            })
        
        if self.assist_visible_panels.get('tm', True):
            self.assist_tabs.append({
                'key': 'tm',
                'name': '💾 TM',
                'short': 'TM',
                'frame': None,
                'button': None,
                'create_func': self.create_tm_tab
            })
        
        if self.assist_visible_panels.get('glossary', True):
            self.assist_tabs.append({
                'key': 'glossary',
                'name': '📚 Glossary',
                'short': 'Glossary',
                'frame': None,
                'button': None,
                'create_func': self.create_glossary_tab
            })
        
        if self.assist_visible_panels.get('reference_images', True):
            self.assist_tabs.append({
                'key': 'reference_images',
                'name': '🖼 Images',
                'short': 'Images',
                'frame': None,
                'button': None,
                'create_func': self.create_reference_images_tab
            })
        
        if self.assist_visible_panels.get('nontrans', True):
            self.assist_tabs.append({
                'key': 'nontrans',
                'name': '🔒 Non-trans',
                'short': 'Non-trans',
                'frame': None,
                'button': None,
                'create_func': self.create_nontrans_tab
            })
        
        if self.assist_visible_panels.get('tracked_changes', True):
            self.assist_tabs.append({
                'key': 'tracked_changes',
                'name': '📊 Changes',
                'short': 'Changes',
                'frame': None,
                'button': None,
                'create_func': self.create_tracked_changes_tab
            })
        
        if self.assist_visible_panels.get('settings', True):
            self.assist_tabs.append({
                'key': 'settings',
                'name': '⚙ Settings',
                'short': 'Settings',
                'frame': None,
                'button': None,
                'create_func': self.create_settings_tab
            })
        
        if self.assist_visible_panels.get('log', True):
            self.assist_tabs.append({
                'key': 'log',
                'name': '📋 Log',
                'short': 'Log',
                'frame': None,
                'button': None,
                'create_func': self.create_log_tab
            })
        
        # Create all tab frames (hidden initially)
        for tab in self.assist_tabs:
            frame = tk.Frame(self.assist_content_area, bg='white')
            tab['frame'] = frame
            tab['create_func'](frame)
        
        # Build tab buttons and determine which ones fit
        self.assist_current_tab = 0
        self._build_assist_tab_buttons()
        
        # Show first tab
        self._show_assist_tab(0)
    
    def _build_assist_tab_buttons(self):
        """Build tab buttons and overflow menu based on available space"""
        # Clear existing buttons
        for widget in self.assist_tab_buttons_frame.winfo_children():
            widget.destroy()
        
        # Clear overflow menu
        self.assist_overflow_menu.delete(0, 'end')
        self.assist_overflow_btn.pack_forget()
        
        # Create buttons for all tabs
        for i, tab in enumerate(self.assist_tabs):
            btn = tk.Button(self.assist_tab_buttons_frame,
                          text=tab['name'],
                          font=('Segoe UI', 8),
                          bg='#e0e0e0' if i != self.assist_current_tab else 'white',
                          fg='#333',
                          relief='raised' if i != self.assist_current_tab else 'sunken',
                          bd=1,
                          padx=8,
                          pady=4,
                          cursor='hand2',
                          command=lambda idx=i: self._show_assist_tab(idx))
            btn.pack(side='left', padx=1, pady=2)
            tab['button'] = btn
        
        # Schedule overflow check after widgets are rendered
        self.assist_tab_buttons_frame.after(100, self._check_tab_overflow)
    
    def _check_tab_overflow(self):
        """Check if tabs overflow and move excess to dropdown"""
        try:
            # Get available width
            available_width = self.assist_tab_buttons_frame.winfo_width()
            if available_width <= 1:
                return  # Not rendered yet
            
            # Reserve space for overflow button
            overflow_btn_width = 70
            usable_width = available_width - overflow_btn_width
            
            # Calculate cumulative width of buttons
            total_width = 0
            overflow_index = len(self.assist_tabs)  # All fit by default
            
            for i, tab in enumerate(self.assist_tabs):
                if tab['button'] and tab['button'].winfo_exists():
                    btn_width = tab['button'].winfo_width() + 2  # +2 for padding
                    total_width += btn_width
                    
                    if total_width > usable_width and overflow_index == len(self.assist_tabs):
                        # This button doesn't fit
                        overflow_index = i
            
            # If some tabs don't fit, move them to overflow menu
            if overflow_index < len(self.assist_tabs):
                # Hide overflow buttons
                for i in range(overflow_index, len(self.assist_tabs)):
                    if self.assist_tabs[i]['button']:
                        self.assist_tabs[i]['button'].pack_forget()
                
                # Rebuild overflow menu
                self.assist_overflow_menu.delete(0, 'end')
                for i in range(overflow_index, len(self.assist_tabs)):
                    tab = self.assist_tabs[i]
                    self.assist_overflow_menu.add_command(
                        label=tab['name'],
                        command=lambda idx=i: self._show_assist_tab(idx)
                    )
                
                # Show overflow button
                self.assist_overflow_btn.pack(side='right', padx=2, pady=2)
            else:
                # All tabs fit, hide overflow button
                self.assist_overflow_btn.pack_forget()
                
        except tk.TclError:
            pass  # Widget was destroyed
    
    def _show_assist_tab(self, index):
        """Show the specified tab and hide others"""
        if index < 0 or index >= len(self.assist_tabs):
            return
        
        # Hide all tab frames
        for tab in self.assist_tabs:
            if tab['frame']:
                tab['frame'].pack_forget()
        
        # Update button styles
        for i, tab in enumerate(self.assist_tabs):
            if tab['button'] and tab['button'].winfo_exists():
                if i == index:
                    tab['button'].config(bg='white', relief='sunken', font=('Segoe UI', 8, 'bold'))
                else:
                    tab['button'].config(bg='#e0e0e0', relief='raised', font=('Segoe UI', 8))
        
        # Show selected tab frame
        self.assist_tabs[index]['frame'].pack(fill='both', expand=True)
        self.assist_current_tab = index
    
    def create_stacked_assistance(self):
        """Create stacked collapsible panels layout with resizable panes"""
        # Create a vertical PanedWindow for resizable panels
        self.assist_stacked_paned = ttk.PanedWindow(self.assist_content_frame, orient='vertical')
        self.assist_stacked_paned.pack(fill='both', expand=True, padx=2, pady=2)
        
        # Track panel weights for initial sizing
        panel_weights = []
        
        # Create collapsible panels in the PanedWindow
        if self.assist_visible_panels.get('mt', True):
            mt_container = self.create_collapsible_panel_resizable('🤖 Machine Translation', 'mt',
                                                                   self.create_mt_tab, expanded=True)
            self.assist_stacked_paned.add(mt_container, weight=1)
            panel_weights.append(1)
        
        if self.assist_visible_panels.get('llm', True):
            llm_container = self.create_collapsible_panel_resizable('✨ LLM Translation', 'llm',
                                                                    self.create_llm_tab, expanded=True)
            self.assist_stacked_paned.add(llm_container, weight=1)
            panel_weights.append(1)
        
        if self.assist_visible_panels.get('tm', True):
            tm_container = self.create_collapsible_panel_resizable('� TM Matches', 'tm_matches',
                                                                   self.create_tm_tab, expanded=False)
            self.assist_stacked_paned.add(tm_container, weight=1)
            panel_weights.append(1)
        
        if self.assist_visible_panels.get('glossary', True):
            glossary_container = self.create_collapsible_panel_resizable('📚 Glossary', 'glossary',
                                                                         self.create_glossary_tab, expanded=False)
            self.assist_stacked_paned.add(glossary_container, weight=1)
            panel_weights.append(1)
        
        if self.assist_visible_panels.get('nontrans', True):
            nontrans_container = self.create_collapsible_panel_resizable('🔒 Non-translatables', 'nontrans',
                                                                         self.create_nontrans_tab, expanded=False)
            self.assist_stacked_paned.add(nontrans_container, weight=1)
            panel_weights.append(1)

    
    
    def create_collapsible_panel_resizable(self, title, panel_id, content_creator, expanded=True):
        """Create a resizable collapsible panel with header and content for use in PanedWindow"""
        # Panel container - this will be added to the PanedWindow
        panel_container = tk.Frame(self.assist_stacked_paned, bg='#ffffff', relief='solid', borderwidth=1)
        
        # Header with collapse/expand button and visibility toggle
        header = tk.Frame(panel_container, bg='#e3f2fd', cursor='hand2')
        header.pack(fill='x')
        
        # Collapse/Expand button
        collapse_btn = tk.Label(header, text='▼' if expanded else '▶',
                               bg='#e3f2fd', font=('Segoe UI', 10),
                               cursor='hand2', width=2)
        collapse_btn.pack(side='left', padx=5, pady=3)
        
        # Title
        title_label = tk.Label(header, text=title, bg='#e3f2fd',
                              font=('Segoe UI', 9, 'bold'), cursor='hand2')
        title_label.pack(side='left', pady=3)
        
        # Resize indicator
        resize_label = tk.Label(header, text='⇕', bg='#e3f2fd',
                               font=('Segoe UI', 9), fg='#666')
        resize_label.pack(side='left', padx=10)
        
        # Undock button (removes from stack, returns to tabs)
        undock_btn = tk.Button(header, text='⊟', command=lambda: self.undock_panel(panel_id),
                              bg='#ff9800', fg='white', font=('Segoe UI', 8),
                              relief='flat', padx=4, pady=0, cursor='hand2')
        undock_btn.pack(side='right', padx=5, pady=3)
        
        # Content frame with minimum height
        content_frame = tk.Frame(panel_container, bg='white')
        if expanded:
            content_frame.pack(fill='both', expand=True, padx=2, pady=2)
        
        # Set minimum height for the panel
        panel_container.update_idletasks()
        
        # Create content using the provided creator function
        content_creator(content_frame)
        
        # Toggle function
        def toggle_panel():
            if content_frame.winfo_ismapped():
                content_frame.pack_forget()
                collapse_btn.config(text='▶')
                panel_container.config(height=30)  # Collapsed height (just header)
            else:
                content_frame.pack(fill='both', expand=True, padx=2, pady=2)
                collapse_btn.config(text='▼')
                panel_container.config(height=200)  # Expanded default height
        
        # Bind click events to header elements (but not the undock button)
        collapse_btn.bind('<Button-1>', lambda e: toggle_panel())
        title_label.bind('<Button-1>', lambda e: toggle_panel())
        
        # Return the container to be added to PanedWindow
        return panel_container
    
    def create_collapsible_panel(self, parent, title, panel_id, content_creator, expanded=True):
        """Create a collapsible panel with header and content"""
        # Panel container
        panel_container = tk.Frame(parent, bg='#ffffff', relief='solid', borderwidth=1)
        panel_container.pack(fill='both', expand=False, padx=3, pady=2)
        
        # Header with collapse/expand button and visibility toggle
        header = tk.Frame(panel_container, bg='#e3f2fd', cursor='hand2')
        header.pack(fill='x')
        
        # Collapse/Expand button
        collapse_btn = tk.Label(header, text='▼' if expanded else '▶',
                               bg='#e3f2fd', font=('Segoe UI', 10),
                               cursor='hand2', width=2)
        collapse_btn.pack(side='left', padx=5, pady=3)
        
        # Title
        title_label = tk.Label(header, text=title, bg='#e3f2fd',
                              font=('Segoe UI', 9, 'bold'), cursor='hand2')
        title_label.pack(side='left', pady=3)
        
        # Undock button (removes from stack, returns to tabs)
        undock_btn = tk.Button(header, text='⊟', command=lambda: self.undock_panel(panel_id),
                              bg='#ff9800', fg='white', font=('Segoe UI', 8),
                              relief='flat', padx=4, pady=0, cursor='hand2')
        undock_btn.pack(side='right', padx=5, pady=3)
        
        # Content frame
        content_frame = tk.Frame(panel_container, bg='white')
        if expanded:
            content_frame.pack(fill='both', expand=True, padx=2, pady=2)
        
        # Create content using the provided creator function
        content_creator(content_frame)
        
        # Toggle function
        def toggle_panel():
            if content_frame.winfo_ismapped():
                content_frame.pack_forget()
                collapse_btn.config(text='▶')
            else:
                content_frame.pack(fill='both', expand=True, padx=2, pady=2)
                collapse_btn.config(text='▼')
        
        # Bind click events to header elements
        collapse_btn.bind('<Button-1>', lambda e: toggle_panel())
        title_label.bind('<Button-1>', lambda e: toggle_panel())
        header.bind('<Button-1>', lambda e: toggle_panel())
    
    def undock_panel(self, panel_id):
        """Remove a panel from stacked view (will reappear when switching to tabbed)"""
        self.assist_visible_panels[panel_id] = False
        self.rebuild_assistance_layout()
        self.log(f"ℹ Panel '{panel_id}' hidden. Switch to Tabbed View to restore.")
    
    def create_detail_panel(self, parent):
        """Create the suggestion detail panel (shows source/target of selected item)"""
        detail_frame = tk.LabelFrame(parent, text="Suggestion Detail", padx=5, pady=5)
        detail_frame.pack(side='bottom', fill='x', padx=2, pady=2)
        
        # Source of selected suggestion
        tk.Label(detail_frame, text="Source:", font=('Segoe UI', 9, 'bold')).pack(anchor='w')
        self.assist_detail_source = tk.Text(detail_frame, height=2, wrap='word', 
                                            bg='#f5f5f5', state='disabled',
                                            font=('Segoe UI', 9))
        self.assist_detail_source.pack(fill='x', pady=(2, 8))
        
        # Target of selected suggestion
        tk.Label(detail_frame, text="Target:", font=('Segoe UI', 9, 'bold')).pack(anchor='w')
        self.assist_detail_target = tk.Text(detail_frame, height=2, wrap='word',
                                            bg='white', state='disabled',
                                            font=('Segoe UI', 9))
        self.assist_detail_target.pack(fill='x', pady=(2, 5))
        
        # Action buttons
        detail_btn_frame = tk.Frame(detail_frame)
        detail_btn_frame.pack(fill='x', pady=(5, 0))
        
        tk.Button(detail_btn_frame, text="📋 Copy to Target", 
                 command=self.copy_suggestion_to_target,
                 bg='#4CAF50', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=2)
        tk.Button(detail_btn_frame, text="➕ Insert at Cursor",
                 command=self.insert_suggestion_at_cursor,
                 bg='#2196F3', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=2)
    
    # === NEW TAB CREATORS ===
    
    def create_prompts_tab(self, parent):
        """Create System Prompts tab - global AI behavior settings"""
        # Info section
        info_frame = tk.Frame(parent, bg='#e3f2fd', relief='solid', borderwidth=1)
        info_frame.pack(fill='x', padx=5, pady=5)
        
        tk.Label(info_frame, text="📝 System Prompts (Global AI Behavior)", font=('Segoe UI', 10, 'bold'),
                bg='#e3f2fd').pack(anchor='w', padx=10, pady=5)
        tk.Label(info_frame, text="Configure how the AI handles translation and proofreading tasks",
                font=('Segoe UI', 9), bg='#e3f2fd', fg='#666').pack(anchor='w', padx=10, pady=(0, 5))
        
        # Create notebook for translate/proofread tabs
        prompts_notebook = ttk.Notebook(parent)
        prompts_notebook.pack(fill='both', expand=True, padx=5, pady=5)
        
        # === TRANSLATION PROMPT TAB ===
        translate_frame = tk.Frame(prompts_notebook, bg='white')
        prompts_notebook.add(translate_frame, text='🌐 Translation')
        
        # Translation prompt text
        translate_scroll = tk.Scrollbar(translate_frame)
        translate_scroll.pack(side='right', fill='y')
        
        self.translate_prompt_text = tk.Text(translate_frame, wrap='word', height=15,
                                            yscrollcommand=translate_scroll.set,
                                            font=('Consolas', 9))
        self.translate_prompt_text.pack(fill='both', expand=True, padx=5, pady=5)
        translate_scroll.config(command=self.translate_prompt_text.yview)
        
        # Insert current translate prompt
        self.translate_prompt_text.insert("1.0", self.default_translate_prompt)
        
        # Translation prompt buttons
        translate_btn_frame = tk.Frame(translate_frame)
        translate_btn_frame.pack(fill='x', padx=5, pady=5)
        
        tk.Button(translate_btn_frame, text="💾 Save Prompt", 
                 command=self.save_translate_prompt,
                 bg='#4CAF50', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=2)
        tk.Button(translate_btn_frame, text="🔄 Reset to Default", 
                 command=self.reset_translate_prompt,
                 bg='#9E9E9E', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=2)
        tk.Button(translate_btn_frame, text="👁️ Preview", 
                 command=self.preview_translate_prompt,
                 bg='#2196F3', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=2)
        tk.Button(translate_btn_frame, text="📚 Prompt Library", 
                 command=self.show_custom_prompts,
                 bg='#FF9800', fg='white', font=('Segoe UI', 9, 'bold')).pack(side='right', padx=2)
        
        # === PROOFREADING PROMPT TAB ===
        proofread_frame = tk.Frame(prompts_notebook, bg='white')
        prompts_notebook.add(proofread_frame, text='✏️ Proofreading')
        
        # Proofreading prompt text
        proofread_scroll = tk.Scrollbar(proofread_frame)
        proofread_scroll.pack(side='right', fill='y')
        
        self.proofread_prompt_text = tk.Text(proofread_frame, wrap='word', height=15,
                                            yscrollcommand=proofread_scroll.set,
                                            font=('Consolas', 9))
        self.proofread_prompt_text.pack(fill='both', expand=True, padx=5, pady=5)
        proofread_scroll.config(command=self.proofread_prompt_text.yview)
        
        # Insert current proofread prompt
        self.proofread_prompt_text.insert("1.0", self.default_proofread_prompt)
        
        # Proofreading prompt buttons
        proofread_btn_frame = tk.Frame(proofread_frame)
        proofread_btn_frame.pack(fill='x', padx=5, pady=5)
        
        tk.Button(proofread_btn_frame, text="💾 Save Prompt", 
                 command=self.save_proofread_prompt,
                 bg='#4CAF50', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=2)
        tk.Button(proofread_btn_frame, text="🔄 Reset to Default", 
                 command=self.reset_proofread_prompt,
                 bg='#9E9E9E', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=2)
        tk.Button(proofread_btn_frame, text="👁️ Preview", 
                 command=self.preview_proofread_prompt,
                 bg='#2196F3', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=2)
        tk.Button(proofread_btn_frame, text="📚 Prompt Library", 
                 command=self.show_custom_prompts,
                 bg='#FF9800', fg='white', font=('Segoe UI', 9, 'bold')).pack(side='right', padx=2)
        
        # Help text
        help_text = tk.Label(parent, 
                           text="💡 Variables: {{SOURCE_LANGUAGE}}, {{TARGET_LANGUAGE}}, {{SOURCE_TEXT}}",
                           font=('Segoe UI', 8), fg='#666', bg='#f5f5f5')
        help_text.pack(fill='x', padx=5, pady=2)
    
    def create_projects_tab(self, parent):
        """Create Project Library tab - quick access to recent projects"""
        # Info section
        info_frame = tk.Frame(parent, bg='#e8f5e9', relief='solid', borderwidth=1)
        info_frame.pack(fill='x', padx=5, pady=5)
        
        tk.Label(info_frame, text="📁 Project Library", font=('Segoe UI', 10, 'bold'),
                bg='#e8f5e9').pack(anchor='w', padx=10, pady=5)
        tk.Label(info_frame, text="Quick access to your translation projects",
                font=('Segoe UI', 9), bg='#e8f5e9', fg='#666').pack(anchor='w', padx=10, pady=(0, 5))
        
        # Current project info
        current_frame = tk.LabelFrame(parent, text="Current Project", padx=5, pady=5)
        current_frame.pack(fill='x', padx=5, pady=5)
        
        project_name = os.path.basename(self.project_file) if self.project_file else "No project loaded"
        self.project_current_label = tk.Label(current_frame, text=f"📄 {project_name}",
                                             font=('Segoe UI', 9, 'bold'), fg='#4CAF50')
        self.project_current_label.pack(anchor='w')
        
        stats_text = f"Segments: {len(self.segments)} | Translated: {sum(1 for s in self.segments if s.target)}"
        tk.Label(current_frame, text=stats_text,
                font=('Segoe UI', 8), fg='#666').pack(anchor='w')
        
        # Recent projects list
        list_frame = tk.LabelFrame(parent, text="Recent Projects", padx=5, pady=5)
        list_frame.pack(fill='both', expand=True, padx=5, pady=5)
        
        # Treeview for recent projects
        tree_scroll = ttk.Scrollbar(list_frame, orient='vertical')
        tree_scroll.pack(side='right', fill='y')
        
        self.project_tree = ttk.Treeview(list_frame, columns=('modified', 'segments'),
                                        show='tree headings', yscrollcommand=tree_scroll.set)
        tree_scroll.config(command=self.project_tree.yview)
        
        self.project_tree.heading('#0', text='Project Name')
        self.project_tree.heading('modified', text='Modified')
        self.project_tree.heading('segments', text='Segments')
        
        self.project_tree.column('#0', width=200)
        self.project_tree.column('modified', width=100)
        self.project_tree.column('segments', width=80)
        
        self.project_tree.pack(fill='both', expand=True)
        
        # Add current project if loaded
        if self.project_file:
            self.project_tree.insert('', 'end', 
                                    text=os.path.basename(self.project_file),
                                    values=('Today', len(self.segments)))
        
        # Action buttons
        button_frame = tk.Frame(parent)
        button_frame.pack(fill='x', padx=5, pady=5)
        
        tk.Button(button_frame, text="📂 Browse Projects", command=self.load_project,
                 bg='#4CAF50', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=2)
    
    def create_tm_manager_tab(self, parent):
        """Create TM Manager tab - Load, view, and configure Translation Memory"""
        tm_info = tk.Frame(parent, bg='#e3f2fd', relief='solid', borderwidth=1)
        tm_info.pack(fill='x', padx=5, pady=5)
        tk.Label(tm_info, text="Translation Memory Manager", font=('Segoe UI', 10, 'bold'),
                bg='#e3f2fd').pack(anchor='w', padx=10, pady=5)
        tk.Label(tm_info, text="Load, view, and manage your translation memory",
                font=('Segoe UI', 9), bg='#e3f2fd', fg='#666').pack(anchor='w', padx=10, pady=(0, 5))
        
        # Status display
        status_frame = tk.LabelFrame(parent, text="Current Status", padx=10, pady=10)
        status_frame.pack(fill='x', padx=5, pady=5)
        
        self.tm_context_status_label = tk.Label(status_frame, 
                                        text=f"📊 {self.tm_agent.get_entry_count()} entries | Threshold: {int(self.tm_agent.fuzzy_threshold*100)}%",
                                        font=('Segoe UI', 9))
        self.tm_context_status_label.pack(anchor='w', pady=5)
        
        # Action buttons
        tm_buttons = tk.Frame(parent)
        tm_buttons.pack(fill='x', padx=5, pady=5)
        
        tk.Button(tm_buttons, text="📂 Load TM File", command=self.load_tm_file,
                 bg='#4CAF50', fg='white', font=('Segoe UI', 9), width=15).pack(side='left', padx=2)
        tk.Button(tm_buttons, text="⚙️ TM Manager", command=self.show_tm_manager,
                 bg='#2196F3', fg='white', font=('Segoe UI', 9), width=15).pack(side='left', padx=2)
        
        # Quick info
        info_text = tk.Text(parent, height=8, wrap='word', font=('Segoe UI', 9),
                           bg='#f5f5f5', relief='flat', padx=10, pady=10)
        info_text.pack(fill='both', expand=True, padx=5, pady=5)
        info_text.insert('1.0', 
            "💡 Translation Memory helps you:\n\n"
            "• Reuse previous translations automatically\n"
            "• Find similar segments (fuzzy matching)\n"
            "• Maintain consistency across projects\n"
            "• Save money on API costs\n\n"
            "Load a TMX or TXT file to get started, or your TM will build "
            "automatically as you translate segments.")
        info_text.config(state='disabled')
    
    def create_reference_images_tab(self, parent):
        """Create Reference Images tab - Visual context for multimodal AI"""
        img_info = tk.Frame(parent, bg='#fff3e0', relief='solid', borderwidth=1)
        img_info.pack(fill='x', padx=5, pady=5)
        tk.Label(img_info, text="Reference Images", font=('Segoe UI', 10, 'bold'),
                bg='#fff3e0').pack(anchor='w', padx=10, pady=5)
        tk.Label(img_info, text="Provide visual context for multimodal AI models (GPT-4 Vision, Claude Vision)",
                font=('Segoe UI', 9), bg='#fff3e0', fg='#666').pack(anchor='w', padx=10, pady=(0, 5))
        
        img_folder_frame = tk.LabelFrame(parent, text="Image Folder", padx=10, pady=10)
        img_folder_frame.pack(fill='x', padx=5, pady=5)
        
        self.image_folder_var = tk.StringVar(value="Not yet configured")
        tk.Label(img_folder_frame, textvariable=self.image_folder_var,
                font=('Segoe UI', 9), fg='#999').pack(anchor='w', pady=5)
        
        tk.Button(img_folder_frame, text="📁 Select Image Folder",
                 bg='#FF9800', fg='white', font=('Segoe UI', 9),
                 state='disabled').pack(anchor='w', pady=5)
        
        # Feature info
        info_text = tk.Text(parent, height=10, wrap='word', font=('Segoe UI', 9),
                           bg='#f5f5f5', relief='flat', padx=10, pady=10)
        info_text.pack(fill='both', expand=True, padx=5, pady=5)
        info_text.insert('1.0', 
            "🚀 Coming Soon!\n\n"
            "This feature will allow you to:\n\n"
            "• Attach screenshots or diagrams to segments\n"
            "• Provide visual context to AI models\n"
            "• Improve translation quality for UI elements\n"
            "• Help AI understand layout and formatting\n\n"
            "Supported models:\n"
            "• GPT-4 Vision (OpenAI)\n"
            "• Claude 3 Vision (Anthropic)\n"
            "• Gemini Pro Vision (Google)")
        info_text.config(state='disabled')
    
    def create_custom_instructions_tab(self, parent):
        """Create Custom Instructions tab - project-specific translation guidance"""
        # Info section with clear distinction from system prompts
        info_frame = tk.Frame(parent, bg='#fff3e0', relief='solid', borderwidth=1)
        info_frame.pack(fill='x', padx=5, pady=5)
        
        tk.Label(info_frame, text="📋 Custom Instructions (Project-Specific)", font=('Segoe UI', 10, 'bold'),
                bg='#fff3e0').pack(anchor='w', padx=10, pady=5)
        tk.Label(info_frame, text="Add specific guidance for THIS project (appended to system prompts)",
                font=('Segoe UI', 9), bg='#fff3e0', fg='#666').pack(anchor='w', padx=10, pady=(0, 5))
        
        # Explanation box
        explain_frame = tk.Frame(parent, bg='#e8f5e9', relief='solid', borderwidth=1)
        explain_frame.pack(fill='x', padx=5, pady=5)
        
        tk.Label(explain_frame, text="💡 How it works:",
                font=('Segoe UI', 9, 'bold'), bg='#e8f5e9').pack(anchor='w', padx=10, pady=(5, 2))
        tk.Label(explain_frame, 
                text="• System Prompts (global) define HOW to translate\n"
                     "• Custom Instructions (per-project) define WHAT to focus on\n"
                     "• Both are combined when translating segments",
                font=('Segoe UI', 8), bg='#e8f5e9', justify='left').pack(anchor='w', padx=20, pady=(0, 5))
        
        # Document Context Info
        context_info_frame = tk.Frame(parent, bg='#e3f2fd', relief='solid', borderwidth=1)
        context_info_frame.pack(fill='x', padx=5, pady=5)
        
        tk.Label(context_info_frame, text="📄 Document Context Status",
                font=('Segoe UI', 9, 'bold'), bg='#e3f2fd').pack(anchor='w', padx=10, pady=(5, 2))
        
        self.context_status_label = tk.Label(context_info_frame, 
                text="Context: Enabled | 0 segments | 0 characters",
                font=('Segoe UI', 8), bg='#e3f2fd', fg='#1976D2')
        self.context_status_label.pack(anchor='w', padx=20, pady=(0, 5))
        
        tk.Label(context_info_frame, 
                text="Full document context is sent to AI for better terminology consistency.\n"
                     "Toggle in Settings tab if needed.",
                font=('Segoe UI', 8), bg='#e3f2fd', fg='#666', justify='left').pack(anchor='w', padx=20, pady=(0, 5))
        
        # Instructions text area - starts empty
        self.custom_instructions_text = scrolledtext.ScrolledText(parent, wrap='word', height=10,
                                                                 font=('Segoe UI', 9))
        self.custom_instructions_text.pack(fill='both', expand=True, padx=5, pady=5)
        
        # Save button
        button_frame = tk.Frame(parent)
        button_frame.pack(fill='x', padx=5, pady=5)
        
        tk.Button(button_frame, text="💾 Save to Project", 
                 command=self.save_custom_instructions,
                 bg='#4CAF50', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=2)
        tk.Button(button_frame, text="\U0001F4CB Load Example Template",
                 command=self.reload_default_instructions,
                 bg='#2196F3', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=2)
        
        # Hint about global preview button
        tk.Label(button_frame, text="💡 Use '🧪 Preview Prompt' button in workspace header to test",
                font=('Segoe UI', 8), fg='#666').pack(side='right', padx=10)
    
    def save_custom_instructions(self):
        """Save custom instructions to project"""
        self.log("✓ Custom instructions saved")
        # TODO: Save to project file when project is saved
    
    def on_custom_instructions_focus(self, event):
        """Clear placeholder when user focuses on the text field"""
        # Check if it's still the placeholder (gray text)
        current_text = self.custom_instructions_text.get('1.0', 'end-1c')
        if "💡 This is just an example" in current_text or current_text.startswith("# Example: Custom Translation"):
            # Clear placeholder
            self.custom_instructions_text.delete('1.0', tk.END)
            self.custom_instructions_text.tag_remove("placeholder", "1.0", "end")
            # Set to normal black color
            self.custom_instructions_text.insert('1.0', "# Custom Translation Instructions for This Project\n\n")
    
    def on_custom_instructions_key(self, event):
        """Handle key press to clear placeholder styling"""
        # Remove gray placeholder tag when user starts typing
        if self.custom_instructions_text.tag_ranges("placeholder"):
            self.custom_instructions_text.tag_remove("placeholder", "1.0", "end")
    
    def is_custom_instructions_placeholder(self, text):
        """Check if custom instructions text is just a placeholder (should not be sent to AI)"""
        if not text:
            return True
        
        # List of placeholder patterns that should be ignored
        placeholders = [
            "# Custom Translation Instructions for This Project",
            "# Example: Custom Translation",
            "💡 This is just an example"
        ]
        
        # Check if text starts with any placeholder or contains the example notice
        for placeholder in placeholders:
            if text.startswith(placeholder) and len(text) < 200:  # Short text is likely just placeholder
                return True
            if "💡 This is just an example" in text:
                return True
        
        return False
    
    def reload_default_instructions(self):
        """Load example template as editable text (not placeholder)"""
        self.custom_instructions_text.delete('1.0', tk.END)
        self.custom_instructions_text.tag_remove("placeholder", "1.0", "end")  # Remove gray styling
        self.custom_instructions_text.insert('1.0', 
            "# Custom Translation Instructions for This Project\n\n"
            "## Style Guidelines:\n"
            "- Use formal tone (use 'u' instead of 'je' in Dutch)\n"
            "- Maintain professional business language\n\n"
            "## Terminology Preferences:\n"
            "- Keep technical terms in English unless standard translation exists\n"
            "- 'user interface' → 'gebruikersinterface' (not 'interface')\n"
            "- 'dashboard' → 'dashboard' (keep in English)\n\n"
            "## Formatting Rules:\n"
            "- Preserve all HTML/XML tags exactly as they appear\n"
            "- Maintain capitalization for proper nouns\n\n"
            "## Context:\n"
            "This is a software localization project for a business analytics platform.\n"
            "Target audience: Business professionals and data analysts.\n")
        self.log("✓ Example template loaded - you can now edit it")
    
    def test_custom_instructions(self):
        """Test custom instructions with current segment"""
        if not self.current_segment:
            messagebox.showinfo("No Segment", "Please select a segment first")
            return
        
        instructions = self.custom_instructions_text.get('1.0', tk.END).strip()
        combined_prompt = self.current_translate_prompt + "\n\n**SPECIAL INSTRUCTIONS:**\n" + instructions
        
        # Show preview dialog
        dialog = tk.Toplevel(self.root)
        dialog.title("Combined Prompt Preview")
        dialog.geometry("700x500")
        
        tk.Label(dialog, text="This is how your prompts will be combined:",
                font=('Segoe UI', 10, 'bold')).pack(padx=10, pady=5)
        
        text = scrolledtext.ScrolledText(dialog, wrap='word', font=('Consolas', 9))
        text.pack(fill='both', expand=True, padx=10, pady=5)
        text.insert('1.0', combined_prompt)
        text.config(state='disabled')
        
        tk.Button(dialog, text="Close", command=dialog.destroy).pack(pady=5)
    
    # === SYSTEM PROMPT MANAGEMENT ===
    
    def save_translate_prompt(self):
        """Save translation prompt"""
        self.current_translate_prompt = self.translate_prompt_text.get('1.0', tk.END).strip()
        self.log("✓ Translation prompt saved")
        messagebox.showinfo("Saved", "Translation prompt updated successfully")
    
    def reset_translate_prompt(self):
        """Reset translation prompt to default"""
        if messagebox.askyesno("Reset Prompt", "Reset translation prompt to default?"):
            self.translate_prompt_text.delete('1.0', tk.END)
            self.translate_prompt_text.insert('1.0', self.default_translate_prompt)
            self.current_translate_prompt = self.default_translate_prompt
            self.log("✓ Translation prompt reset to default")
    
    def preview_translate_prompt(self):
        """Preview translation prompt with current variables"""
        prompt = self.translate_prompt_text.get('1.0', tk.END).strip()
        
        # Replace variables
        preview = prompt.replace('{{SOURCE_LANGUAGE}}', self.source_language)
        preview = preview.replace('{{TARGET_LANGUAGE}}', self.target_language)
        preview = preview.replace('{{SOURCE_TEXT}}', '[Current segment text will appear here]')
        
        # Show preview dialog
        dialog = tk.Toplevel(self.root)
        dialog.title("Translation Prompt Preview")
        dialog.geometry("700x500")
        
        tk.Label(dialog, text=f"Preview: {self.source_language} → {self.target_language}",
                font=('Segoe UI', 10, 'bold')).pack(padx=10, pady=5)
        
        text = scrolledtext.ScrolledText(dialog, wrap='word', font=('Consolas', 9))
        text.pack(fill='both', expand=True, padx=10, pady=5)
        text.insert('1.0', preview)
        text.config(state='disabled')
        
        tk.Button(dialog, text="Close", command=dialog.destroy).pack(pady=5)
    
    def save_proofread_prompt(self):
        """Save proofreading prompt"""
        self.current_proofread_prompt = self.proofread_prompt_text.get('1.0', tk.END).strip()
        self.log("✓ Proofreading prompt saved")
        messagebox.showinfo("Saved", "Proofreading prompt updated successfully")
    
    def reset_proofread_prompt(self):
        """Reset proofreading prompt to default"""
        if messagebox.askyesno("Reset Prompt", "Reset proofreading prompt to default?"):
            self.proofread_prompt_text.delete('1.0', tk.END)
            self.proofread_prompt_text.insert('1.0', self.default_proofread_prompt)
            self.current_proofread_prompt = self.default_proofread_prompt
            self.log("✓ Proofreading prompt reset to default")
    
    def preview_proofread_prompt(self):
        """Preview proofreading prompt with current variables"""
        prompt = self.proofread_prompt_text.get('1.0', tk.END).strip()
        
        # Replace variables
        preview = prompt.replace('{{SOURCE_LANGUAGE}}', self.source_language)
        preview = preview.replace('{{TARGET_LANGUAGE}}', self.target_language)
        
        # Show preview dialog
        dialog = tk.Toplevel(self.root)
        dialog.title("Proofreading Prompt Preview")
        dialog.geometry("700x500")
        
        tk.Label(dialog, text=f"Preview: {self.source_language} → {self.target_language}",
                font=('Segoe UI', 10, 'bold')).pack(padx=10, pady=5)
        
        text = scrolledtext.ScrolledText(dialog, wrap='word', font=('Consolas', 9))
        text.pack(fill='both', expand=True, padx=10, pady=5)
        text.insert('1.0', preview)
        text.config(state='disabled')
        
        tk.Button(dialog, text="Close", command=dialog.destroy).pack(pady=5)
    
    def preview_combined_prompt(self):
        """Preview the complete combined prompt (system + custom instructions) with current segment"""
        # Check if there's a current segment
        if not self.current_segment:
            messagebox.showinfo("No Segment Selected", 
                              "Please select a segment to preview the prompt.\n\n"
                              "The preview will show how the system prompt and custom instructions "
                              "are combined for the current segment.")
            return
        
        # Build the complete prompt
        prompt = self.current_translate_prompt
        prompt = prompt.replace("{{SOURCE_LANGUAGE}}", self.source_language)
        prompt = prompt.replace("{{TARGET_LANGUAGE}}", self.target_language)
        prompt = prompt.replace("{{SOURCE_TEXT}}", self.current_segment.source)
        
        # Add custom instructions if provided (skip if just the placeholder)
        custom_instructions = self.custom_instructions_text.get('1.0', tk.END).strip()
        if custom_instructions and not self.is_custom_instructions_placeholder(custom_instructions):
            prompt += "\n\n**SPECIAL INSTRUCTIONS FOR THIS PROJECT:**\n" + custom_instructions
        
        # Show preview dialog
        dialog = tk.Toplevel(self.root)
        dialog.title("Complete Prompt Preview - What AI Will Receive")
        dialog.geometry("800x600")
        dialog.transient(self.root)
        
        # Header
        header_frame = tk.Frame(dialog, bg='#2196F3', height=60)
        header_frame.pack(fill='x')
        header_frame.pack_propagate(False)
        
        tk.Label(header_frame, 
                text="🧪 Complete Prompt Preview",
                font=('Segoe UI', 14, 'bold'), bg='#2196F3', fg='white').pack(pady=5)
        tk.Label(header_frame,
                text=f"{self.source_language} → {self.target_language} | Segment #{self.current_segment.id}",
                font=('Segoe UI', 9), bg='#2196F3', fg='white').pack()
        
        # Info panel
        info_frame = tk.Frame(dialog, bg='#e3f2fd')
        info_frame.pack(fill='x', padx=10, pady=5)
        
        tk.Label(info_frame, text="💡 This is the EXACT prompt that will be sent to the AI",
                font=('Segoe UI', 9, 'bold'), bg='#e3f2fd').pack(anchor='w', padx=10, pady=5)
        
        # Composition breakdown
        composition_text = "📋 Composition:\n"
        composition_text += f"  • System Prompt (Translation): {len(self.current_translate_prompt)} characters\n"
        if custom_instructions and not self.is_custom_instructions_placeholder(custom_instructions):
            composition_text += f"  • Custom Instructions: {len(custom_instructions)} characters\n"
        composition_text += f"  • Total prompt length: {len(prompt)} characters"
        
        tk.Label(info_frame, text=composition_text,
                font=('Segoe UI', 8), bg='#e3f2fd', fg='#666', justify='left').pack(anchor='w', padx=20, pady=(0, 5))
        
        # Prompt text area
        text_frame = tk.Frame(dialog)
        text_frame.pack(fill='both', expand=True, padx=10, pady=5)
        
        scrollbar = tk.Scrollbar(text_frame)
        scrollbar.pack(side='right', fill='y')
        
        text = tk.Text(text_frame, wrap='word', font=('Consolas', 9),
                      yscrollcommand=scrollbar.set, bg='#f5f5f5')
        text.pack(fill='both', expand=True)
        scrollbar.config(command=text.yview)
        
        # Insert prompt with highlighting
        text.insert('1.0', prompt)
        
        # Highlight the source text section
        start_idx = prompt.find(self.current_segment.source)
        if start_idx != -1:
            start_line = prompt[:start_idx].count('\n') + 1
            start_col = start_idx - prompt[:start_idx].rfind('\n') - 1
            end_idx = start_idx + len(self.current_segment.source)
            end_line = prompt[:end_idx].count('\n') + 1
            end_col = end_idx - prompt[:end_idx].rfind('\n') - 1
            
            text.tag_add("source", f"{start_line}.{start_col}", f"{end_line}.{end_col}")
            text.tag_config("source", background="#fff9c4", font=('Consolas', 9, 'bold'))
        
        text.config(state='disabled')
        
        # Button frame
        btn_frame = tk.Frame(dialog)
        btn_frame.pack(fill='x', padx=10, pady=10)
        
        tk.Button(btn_frame, text="📋 Copy to Clipboard", 
                 command=lambda: self.copy_to_clipboard(prompt),
                 bg='#4CAF50', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=5)
        tk.Button(btn_frame, text="Close", command=dialog.destroy,
                 bg='#757575', fg='white', font=('Segoe UI', 9)).pack(side='right', padx=5)
    
    def copy_to_clipboard(self, text):
        """Copy text to clipboard"""
        self.root.clipboard_clear()
        self.root.clipboard_append(text)
        self.log("✓ Copied to clipboard")
        messagebox.showinfo("Copied", "Prompt copied to clipboard!")
    
    def create_context_tab(self, parent):
        """DEPRECATED: Old Context tab - kept for backward compatibility but not used"""
        # Create sub-notebook for different context types
        context_notebook = ttk.Notebook(parent)
        context_notebook.pack(fill='both', expand=True, padx=2, pady=2)
        
        # === TM Context Sub-tab ===
        tm_context_frame = tk.Frame(context_notebook, bg='white')
        context_notebook.add(tm_context_frame, text='💾 TM')
        
        tm_info = tk.Frame(tm_context_frame, bg='#e3f2fd', relief='solid', borderwidth=1)
        tm_info.pack(fill='x', padx=5, pady=5)
        tk.Label(tm_info, text="Translation Memory Status", font=('Segoe UI', 10, 'bold'),
                bg='#e3f2fd').pack(anchor='w', padx=10, pady=5)
        
        self.tm_context_status_label = tk.Label(tm_info, 
                                        text=f"📊 {self.tm_agent.get_entry_count()} entries | Threshold: {int(self.tm_agent.fuzzy_threshold*100)}%",
                                        font=('Segoe UI', 9), bg='#e3f2fd')
        self.tm_context_status_label.pack(anchor='w', padx=10, pady=(0, 5))
        
        tm_buttons = tk.Frame(tm_context_frame)
        tm_buttons.pack(fill='x', padx=5, pady=5)
        
        tk.Button(tm_buttons, text="📂 Load TM File", command=self.load_tm_file,
                 bg='#4CAF50', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=2)
        tk.Button(tm_buttons, text="⚙️ TM Manager", command=self.show_tm_manager,
                 bg='#2196F3', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=2)
        
        # === Reference Images Sub-tab ===
        images_context_frame = tk.Frame(context_notebook, bg='white')
        context_notebook.add(images_context_frame, text='🖼️ Images')
        
        img_info = tk.Frame(images_context_frame, bg='#fff3e0', relief='solid', borderwidth=1)
        img_info.pack(fill='x', padx=5, pady=5)
        tk.Label(img_info, text="Reference Images", font=('Segoe UI', 10, 'bold'),
                bg='#fff3e0').pack(anchor='w', padx=10, pady=5)
        tk.Label(img_info, text="Provide visual context for multimodal AI (future feature)",
                font=('Segoe UI', 9), bg='#fff3e0', fg='#666').pack(anchor='w', padx=10, pady=(0, 5))
        
        img_folder_frame = tk.LabelFrame(images_context_frame, text="Image Folder", padx=5, pady=5)
        img_folder_frame.pack(fill='x', padx=5, pady=5)
        
        self.image_folder_var = tk.StringVar(value="Not yet implemented")
        tk.Label(img_folder_frame, textvariable=self.image_folder_var,
                font=('Segoe UI', 9), fg='#999').pack(anchor='w')
        
        # === Custom Instructions Sub-tab ===
        instructions_frame = tk.Frame(context_notebook, bg='white')
        context_notebook.add(instructions_frame, text='📋 Instructions')
        
        inst_info = tk.Frame(instructions_frame, bg='#f3e5f5', relief='solid', borderwidth=1)
        inst_info.pack(fill='x', padx=5, pady=5)
        tk.Label(inst_info, text="Custom Instructions", font=('Segoe UI', 10, 'bold'),
                bg='#f3e5f5').pack(anchor='w', padx=10, pady=5)
        tk.Label(inst_info, text="Additional context for AI (terminology, style preferences)",
                font=('Segoe UI', 9), bg='#f3e5f5', fg='#666').pack(anchor='w', padx=10, pady=(0, 5))
        
        # DEPRECATED: Use separate variable name to avoid overwriting the main custom_instructions_text
        self.deprecated_custom_instructions_text = scrolledtext.ScrolledText(instructions_frame, 
                                                                 height=10, wrap='word',
                                                                 font=('Segoe UI', 9))
        self.deprecated_custom_instructions_text.pack(fill='both', expand=True, padx=5, pady=5)
        self.deprecated_custom_instructions_text.insert('1.0', 
            "Example:\n"
            "- Use formal tone\n"
            "- Keep technical terms in English\n"
            "- Maintain consistent terminology for 'user interface' → 'gebruikersinterface'\n")
        
        tk.Button(instructions_frame, text="💾 Save Instructions", 
                 command=lambda: self.log("✓ Custom instructions saved"),
                 bg='#4CAF50', fg='white', font=('Segoe UI', 9)).pack(padx=5, pady=5)
    
    def create_settings_tab(self, parent):
        """Create Settings tab - LLM provider, languages, preferences"""
        # LLM Provider Settings
        provider_frame = tk.LabelFrame(parent, text="LLM Provider", padx=10, pady=10)
        provider_frame.pack(fill='x', padx=5, pady=5)
        
        tk.Label(provider_frame, text="Current:", font=('Segoe UI', 9, 'bold')).grid(row=0, column=0, sticky='w', pady=5)
        
        # Use StringVar for dynamic updating
        self.settings_llm_display = tk.StringVar(value=f"{self.current_llm_provider.upper()} / {self.current_llm_model}")
        tk.Label(provider_frame, textvariable=self.settings_llm_display,
                font=('Segoe UI', 9), fg='#4CAF50').grid(row=0, column=1, sticky='w', pady=5)
        
        tk.Button(provider_frame, text="⚙️ Configure API Settings", command=self.show_api_settings,
                 bg='#2196F3', fg='white', font=('Segoe UI', 9)).grid(row=1, column=0, columnspan=2, pady=5, sticky='w')
        
        # Language Settings
        lang_frame = tk.LabelFrame(parent, text="Language Pair", padx=10, pady=10)
        lang_frame.pack(fill='x', padx=5, pady=5)
        
        tk.Label(lang_frame, text="Source:", font=('Segoe UI', 9)).grid(row=0, column=0, sticky='w', pady=5)
        
        # Use StringVar for dynamic updating
        self.settings_source_lang_display = tk.StringVar(value=self.source_language)
        tk.Label(lang_frame, textvariable=self.settings_source_lang_display, font=('Segoe UI', 9, 'bold'),
                fg='#666').grid(row=0, column=1, sticky='w', pady=5)
        
        tk.Label(lang_frame, text="Target:", font=('Segoe UI', 9)).grid(row=1, column=0, sticky='w', pady=5)
        
        # Use StringVar for dynamic updating
        self.settings_target_lang_display = tk.StringVar(value=self.target_language)
        tk.Label(lang_frame, textvariable=self.settings_target_lang_display, font=('Segoe UI', 9, 'bold'),
                fg='#666').grid(row=1, column=1, sticky='w', pady=5)
        
        # Buttons frame for language actions
        lang_buttons = tk.Frame(lang_frame)
        lang_buttons.grid(row=2, column=0, columnspan=2, pady=5, sticky='w')
        
        tk.Button(lang_buttons, text="🌍 Change Languages", command=self.show_language_settings,
                 bg='#FF9800', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=(0, 5))
        tk.Button(lang_buttons, text="🔄 Swap", command=self.swap_languages,
                 bg='#4CAF50', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=(0, 5))
        tk.Button(lang_buttons, text="✏️ Edit Language List", command=self.edit_language_list,
                 bg='#2196F3', fg='white', font=('Segoe UI', 9)).pack(side='left')
        
        # Translation Preferences
        pref_frame = tk.LabelFrame(parent, text="Translation Preferences", padx=10, pady=10)
        pref_frame.pack(fill='x', padx=5, pady=5)
        
        # Batch size setting (default: 100 segments per API call)
        self.chunk_size_var = tk.StringVar(value="100")
        chunk_frame = tk.Frame(pref_frame)
        chunk_frame.pack(anchor='w', pady=2)
        tk.Label(chunk_frame, text="Batch Size (segments per API call):",
                font=('Segoe UI', 9)).pack(side='left')
        tk.Spinbox(chunk_frame, from_=1, to=500, textvariable=self.chunk_size_var,
                  width=8, font=('Segoe UI', 9)).pack(side='left', padx=5)
        tk.Label(pref_frame, text="  ⓘ Larger batches = faster but higher API cost per call. Default: 100",
                font=('Segoe UI', 8), fg='gray').pack(anchor='w', padx=20)
        
        # Surrounding segments context for single-segment translation
        context_row = tk.Frame(pref_frame)
        context_row.pack(anchor='w', pady=5, fill='x')
        tk.Label(context_row, text="Surrounding segments (single-segment translation):",
                font=('Segoe UI', 9)).pack(side='left')
        self.surrounding_segments_var = tk.StringVar(value="5")
        tk.Spinbox(context_row, from_=0, to=20, width=5, textvariable=self.surrounding_segments_var,
                  font=('Segoe UI', 9)).pack(side='left', padx=10)
        tk.Label(context_row, text="segments before/after",
                font=('Segoe UI', 9)).pack(side='left')
        tk.Label(pref_frame, text="  ⓘ Send nearby segments for context without full document. 0 = no context. Default: 5",
                font=('Segoe UI', 8), fg='gray').pack(anchor='w', padx=20)
        
        self.use_context_var = tk.BooleanVar(value=False)  # Changed default to False to save API costs
        tk.Checkbutton(pref_frame, text="Include full document context in batch translation (increases API usage)",
                      variable=self.use_context_var, font=('Segoe UI', 9)).pack(anchor='w', pady=2)
        tk.Label(pref_frame, text="  ⓘ Context helps with consistency but sends more data - use for technical docs",
                font=('Segoe UI', 8), fg='gray').pack(anchor='w', padx=20)
        
        self.check_tm_var = tk.BooleanVar(value=True)
        tk.Checkbutton(pref_frame, text="Check TM before API call",
                      variable=self.check_tm_var, font=('Segoe UI', 9)).pack(anchor='w', pady=2)
        
        self.auto_propagate_var = tk.BooleanVar(value=False)
        tk.Checkbutton(pref_frame, text="Auto-propagate 100% TM matches",
                      variable=self.auto_propagate_var, font=('Segoe UI', 9)).pack(anchor='w', pady=2)
        
        # Auto-Export Options
        export_frame = tk.LabelFrame(parent, text="Auto-Export Options", padx=10, pady=10)
        export_frame.pack(fill='x', padx=5, pady=5)
        
        tk.Label(export_frame, text="Automatically export these formats alongside your primary export:",
                font=('Segoe UI', 9), fg='#666').pack(anchor='w', pady=(0, 10))
        
        # Session reports
        reports_subframe = tk.Frame(export_frame)
        reports_subframe.pack(anchor='w', fill='x', pady=2)
        
        self.auto_export_session_md_var = tk.BooleanVar(value=False)
        tk.Checkbutton(reports_subframe, text="Session Report (.md)",
                      variable=self.auto_export_session_md_var, font=('Segoe UI', 9)).pack(side='left', padx=(0, 20))
        
        self.auto_export_session_html_var = tk.BooleanVar(value=False)
        tk.Checkbutton(reports_subframe, text="Session Report (.html)",
                      variable=self.auto_export_session_html_var, font=('Segoe UI', 9)).pack(side='left')
        
        # Translation memory and exchange formats
        tm_subframe = tk.Frame(export_frame)
        tm_subframe.pack(anchor='w', fill='x', pady=2)
        
        self.auto_export_tmx_var = tk.BooleanVar(value=False)
        tk.Checkbutton(tm_subframe, text="Translation Memory (.tmx)",
                      variable=self.auto_export_tmx_var, font=('Segoe UI', 9)).pack(side='left', padx=(0, 20))
        
        self.auto_export_tsv_var = tk.BooleanVar(value=False)
        tk.Checkbutton(tm_subframe, text="Tab-separated (.tsv)",
                      variable=self.auto_export_tsv_var, font=('Segoe UI', 9)).pack(side='left')
        
        # Bilingual formats
        bilingual_subframe = tk.Frame(export_frame)
        bilingual_subframe.pack(anchor='w', fill='x', pady=2)
        
        self.auto_export_bilingual_txt_var = tk.BooleanVar(value=False)
        tk.Checkbutton(bilingual_subframe, text="Bilingual Text (.txt)",
                      variable=self.auto_export_bilingual_txt_var, font=('Segoe UI', 9)).pack(side='left', padx=(0, 20))
        
        self.auto_export_xliff_var = tk.BooleanVar(value=False)
        tk.Checkbutton(bilingual_subframe, text="XLIFF (.xliff)",
                      variable=self.auto_export_xliff_var, font=('Segoe UI', 9)).pack(side='left')
        
        # Excel format
        excel_subframe = tk.Frame(export_frame)
        excel_subframe.pack(anchor='w', fill='x', pady=2)
        
        self.auto_export_excel_var = tk.BooleanVar(value=False)
        tk.Checkbutton(excel_subframe, text="Excel Bilingual (.xlsx)",
                      variable=self.auto_export_excel_var, font=('Segoe UI', 9)).pack(side='left')
        
        tk.Label(export_frame, text="  ⓘ Exports will be saved in the same directory as your primary export",
                font=('Segoe UI', 8), fg='gray').pack(anchor='w', pady=(5, 0))
        
        # Info section
        info_frame = tk.Frame(parent, bg='#f0f0f0', relief='solid', borderwidth=1)
        info_frame.pack(fill='x', padx=5, pady=10)
        
        tk.Label(info_frame, text="ℹ️ Settings are automatically saved with your project",
                font=('Segoe UI', 9), bg='#f0f0f0', fg='#666').pack(padx=10, pady=10)
    
    def create_tracked_changes_tab(self, parent):
        """Create the Tracked Changes Analysis tab"""
        parent.configure(bg='white')
        
        # Header
        header_frame = tk.Frame(parent, bg='#2c3e50', pady=8)
        header_frame.pack(fill='x')
        
        tk.Label(header_frame, text="📊 Post-Translation Analysis", 
                font=('Segoe UI', 12, 'bold'), bg='#2c3e50', fg='white').pack()
        tk.Label(header_frame, 
                text="Review differences between AI baseline and your final translations",
                font=('Segoe UI', 8), bg='#2c3e50', fg='#ecf0f1').pack()
        
        # Main content
        content_frame = tk.Frame(parent, bg='white', padx=10, pady=10)
        content_frame.pack(fill='both', expand=True)
        
        # Status section
        status_frame = tk.LabelFrame(content_frame, text="Current Status", 
                                    font=('Segoe UI', 10, 'bold'), bg='white', padx=10, pady=10)
        status_frame.pack(fill='x', pady=(0,10))
        
        self.tracked_changes_status_label = tk.Label(status_frame, 
                                                     text="No tracked changes loaded",
                                                     font=('Segoe UI', 10), fg='gray', bg='white')
        self.tracked_changes_status_label.pack(anchor='w', pady=5)
        
        # Update status if changes are loaded
        if hasattr(self, 'tracked_changes_agent') and self.tracked_changes_agent.change_data:
            count = len(self.tracked_changes_agent.change_data)
            files = len(self.tracked_changes_agent.files_loaded)
            self.tracked_changes_status_label.config(
                text=f"✅ {count} changes loaded from {files} file(s)",
                fg='green'
            )
        
        # Actions section
        actions_frame = tk.LabelFrame(content_frame, text="Actions", 
                                     font=('Segoe UI', 10, 'bold'), bg='white', padx=10, pady=10)
        actions_frame.pack(fill='x', pady=(0,10))
        
        # Load button
        load_btn = tk.Button(actions_frame, text="📂 Load Tracked Changes (DOCX)", 
                           command=self.load_tracked_changes_docx,
                           font=('Segoe UI', 10), bg='#3498db', fg='white',
                           cursor='hand2', relief='raised', bd=2, padx=15, pady=8)
        load_btn.pack(fill='x', pady=(0,5))
        
        # Browse/Export button
        browse_btn = tk.Button(actions_frame, text="📊 Browse & Export Analysis Report", 
                             command=self.browse_tracked_changes,
                             font=('Segoe UI', 10), bg='#27ae60', fg='white',
                             cursor='hand2', relief='raised', bd=2, padx=15, pady=8)
        browse_btn.pack(fill='x', pady=(0,5))
        
        # Clear button
        clear_btn = tk.Button(actions_frame, text="🗑 Clear All Changes", 
                            command=self.clear_tracked_changes,
                            font=('Segoe UI', 10), bg='#e74c3c', fg='white',
                            cursor='hand2', relief='raised', bd=2, padx=15, pady=8)
        clear_btn.pack(fill='x')
        
        # Info section
        info_frame = tk.LabelFrame(content_frame, text="How It Works", 
                                  font=('Segoe UI', 10, 'bold'), bg='white', padx=10, pady=10)
        info_frame.pack(fill='both', expand=True)
        
        info_text = """1. Complete your translation project in a CAT tool (memoQ, CafeTran, etc.)
   with tracked changes enabled

2. Export the bilingual document containing tracked changes

3. Click 'Load Tracked Changes' above to import the file

4. Click 'Browse & Export' to:
   • View all changes in a searchable browser
   • Export AI-powered analysis report (Markdown)
   • Configure batch processing (1-100 segments)
   • Get precise change summaries

Use this feature AFTER translation to:
✓ Review your editing decisions
✓ Track workflow improvements
✓ Document translation rationale
✓ Quality assurance checks"""
        
        info_label = tk.Label(info_frame, text=info_text, 
                            font=('Segoe UI', 9), bg='white', fg='#34495e',
                            justify='left', anchor='nw')
        info_label.pack(fill='both', expand=True, padx=5, pady=5)
    
    def create_log_tab(self, parent):
        """Create Log tab - synchronized with main log window"""
        # Header with description
        header_frame = tk.Frame(parent, bg='#f0f0f0', relief='solid', borderwidth=1)
        header_frame.pack(fill='x', padx=5, pady=5)
        
        tk.Label(header_frame, text="📋 Session Log",
                font=('Segoe UI', 10, 'bold'), bg='#f0f0f0').pack(anchor='w', padx=10, pady=(10, 2))
        tk.Label(header_frame, text="All system messages, API calls, and operations are logged here in real-time.",
                font=('Segoe UI', 9), bg='#f0f0f0', fg='#666').pack(anchor='w', padx=10, pady=(0, 10))
        
        # Log display area (synchronized with main log window)
        log_display_frame = tk.Frame(parent)
        log_display_frame.pack(fill='both', expand=True, padx=5, pady=5)
        
        # Create a synchronized log text widget
        self.log_tab_text = scrolledtext.ScrolledText(log_display_frame, wrap='word',
                                                       font=('Consolas', 9), state='disabled',
                                                       bg='white', fg='black')
        self.log_tab_text.pack(fill='both', expand=True)
        
        # Configure text tags for different log levels (same as main log)
        self.log_tab_text.tag_config('info', foreground='black')
        self.log_tab_text.tag_config('success', foreground='green')
        self.log_tab_text.tag_config('warning', foreground='orange')
        self.log_tab_text.tag_config('error', foreground='red')
        
        # SYNC: Copy existing log content from bottom log to this tab
        if hasattr(self, 'log_text'):
            try:
                # Get current content from bottom log
                bottom_log_content = self.log_text.get('1.0', 'end-1c')
                
                # Copy to tab log
                self.log_tab_text.config(state='normal')
                self.log_tab_text.insert('1.0', bottom_log_content)
                self.log_tab_text.see('end')  # Scroll to bottom
                self.log_tab_text.config(state='disabled')
            except (tk.TclError, AttributeError):
                pass  # Log widget not yet created or already destroyed
        
        # Toolbar with clear button
        toolbar = tk.Frame(parent, bg='#f0f0f0')
        toolbar.pack(fill='x', padx=5, pady=(0, 5))
        
        tk.Button(toolbar, text="🗑️ Clear Log", command=self.clear_log,
                 bg='#757575', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=5, pady=5)
        
        tk.Label(toolbar, text="ℹ️ This log is synchronized with the main log window at the bottom",
                font=('Segoe UI', 8), bg='#f0f0f0', fg='#666').pack(side='left', padx=10)
    
    # === END OF NEW TAB CREATORS ===
    
    def create_mt_tab(self, parent):
        """Create Machine Translation tab content"""
        # Toolbar with provider selection and translate button
        toolbar = tk.Frame(parent, bg='#f0f0f0')
        toolbar.pack(side='top', fill='x', padx=5, pady=5)
        
        tk.Label(toolbar, text="Provider:", bg='#f0f0f0', font=('Segoe UI', 9)).pack(side='left', padx=2)
        self.mt_provider_var = tk.StringVar(value="Google Translate")
        mt_provider_combo = ttk.Combobox(toolbar, textvariable=self.mt_provider_var,
                                        values=["Google Translate", "DeepL", "Microsoft Translator", 
                                               "Amazon Translate", "Custom MT"],
                                        state='readonly', width=18, font=('Segoe UI', 9))
        mt_provider_combo.pack(side='left', padx=5)
        
        tk.Button(toolbar, text="🔄 Translate", command=self.get_mt_translation,
                 bg='#4CAF50', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=5)
        
        # Listbox for MT results
        list_frame = tk.Frame(parent)
        list_frame.pack(fill='both', expand=True, padx=5, pady=5)
        
        scrollbar = ttk.Scrollbar(list_frame, orient='vertical')
        self.mt_listbox = tk.Listbox(list_frame, yscrollcommand=scrollbar.set,
                                     font=('Segoe UI', 9), selectmode='single')
        scrollbar.config(command=self.mt_listbox.yview)
        
        self.mt_listbox.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')
        
        self.mt_listbox.bind('<<ListboxSelect>>', self.on_mt_select)
        self.mt_listbox.bind('<Double-Button-1>', lambda e: self.copy_suggestion_to_target())
        
        # Store MT results
        self.mt_results = []
    
    def create_llm_tab(self, parent):
        """Create LLM Translation tab content"""
        # Toolbar with model selection and translate button
        toolbar = tk.Frame(parent, bg='#f0f0f0')
        toolbar.pack(side='top', fill='x', padx=5, pady=5)
        
        tk.Label(toolbar, text="Model:", bg='#f0f0f0', font=('Segoe UI', 9)).pack(side='left', padx=2)
        self.llm_model_var = tk.StringVar(value="Claude 3.5 Sonnet")
        llm_model_combo = ttk.Combobox(toolbar, textvariable=self.llm_model_var,
                                      values=["Claude 3.5 Sonnet", "GPT-4", "Gemini Pro",
                                             "Claude 3 Opus", "GPT-4 Turbo", "Supervertaler Custom"],
                                      state='readonly', width=18, font=('Segoe UI', 9))
        llm_model_combo.pack(side='left', padx=5)
        
        tk.Button(toolbar, text="✨ Generate", command=self.get_llm_translation,
                 bg='#9C27B0', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=5)
        
        # Prompt template selector
        prompt_frame = tk.Frame(parent, bg='#f9f9f9')
        prompt_frame.pack(side='top', fill='x', padx=5, pady=5)
        
        tk.Label(prompt_frame, text="Prompt:", bg='#f9f9f9', font=('Segoe UI', 9)).pack(side='left', padx=2)
        self.llm_prompt_var = tk.StringVar(value="Standard Translation")
        llm_prompt_combo = ttk.Combobox(prompt_frame, textvariable=self.llm_prompt_var,
                                       values=["Standard Translation", "Legal Translation",
                                              "Technical Translation", "Marketing Translation",
                                              "Medical Translation", "Custom Supervertaler Prompt"],
                                       state='readonly', width=25, font=('Segoe UI', 9))
        llm_prompt_combo.pack(side='left', padx=5)
        
        # Listbox for LLM results (can have multiple alternatives)
        list_frame = tk.Frame(parent)
        list_frame.pack(fill='both', expand=True, padx=5, pady=5)
        
        scrollbar = ttk.Scrollbar(list_frame, orient='vertical')
        self.llm_listbox = tk.Listbox(list_frame, yscrollcommand=scrollbar.set,
                                      font=('Segoe UI', 9), selectmode='single')
        scrollbar.config(command=self.llm_listbox.yview)
        
        self.llm_listbox.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')
        
        self.llm_listbox.bind('<<ListboxSelect>>', self.on_llm_select)
        self.llm_listbox.bind('<Double-Button-1>', lambda e: self.copy_suggestion_to_target())
        
        # Store LLM results
        self.llm_results = []
    
    def create_tm_tab(self, parent):
        """Create Translation Memory tab - Matches (top) + Management (bottom)"""
        # === TM MATCHES SECTION (TOP) ===
        matches_frame = tk.LabelFrame(parent, text="TM Matches for Current Segment", 
                                      font=('Segoe UI', 10, 'bold'), padx=5, pady=5)
        matches_frame.pack(fill='both', expand=True, padx=5, pady=(5, 2))
        
        # Toolbar with TM source selection and search button
        toolbar = tk.Frame(matches_frame, bg='#f0f0f0')
        toolbar.pack(side='top', fill='x', padx=2, pady=2)
        
        tk.Label(toolbar, text="TM:", bg='#f0f0f0', font=('Segoe UI', 9)).pack(side='left', padx=2)
        self.tm_source_var = tk.StringVar(value="All Active TMs")
        self.tm_source_combo = ttk.Combobox(toolbar, textvariable=self.tm_source_var,
                                      values=["All Active TMs"],
                                      state='readonly', width=15, font=('Segoe UI', 9))
        self.tm_source_combo.pack(side='left', padx=5)
        
        tk.Button(toolbar, text="🔍 Search", command=self.search_tm,
                 bg='#2196F3', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=5)
        
        # Settings frame
        settings_frame = tk.Frame(matches_frame, bg='#f9f9f9')
        settings_frame.pack(side='top', fill='x', padx=2, pady=2)
        
        tk.Label(settings_frame, text="Min. Match:", bg='#f9f9f9', 
                font=('Segoe UI', 8)).pack(side='left', padx=2)
        self.tm_threshold_var = tk.StringVar(value="75")
        tm_threshold_spin = ttk.Spinbox(settings_frame, from_=50, to=100, increment=5,
                                       textvariable=self.tm_threshold_var, width=5,
                                       font=('Segoe UI', 8))
        tm_threshold_spin.pack(side='left', padx=2)
        tk.Label(settings_frame, text="%", bg='#f9f9f9', font=('Segoe UI', 8)).pack(side='left')
        
        # Treeview for TM matches (shows match %, source, target)
        tree_frame = tk.Frame(matches_frame)
        tree_frame.pack(fill='both', expand=True, padx=2, pady=2)
        
        scrollbar = ttk.Scrollbar(tree_frame, orient='vertical')
        self.tm_tree = ttk.Treeview(tree_frame, columns=('match', 'text', 'source_tm'),
                                   show='headings', yscrollcommand=scrollbar.set,
                                   selectmode='browse', height=8)
        scrollbar.config(command=self.tm_tree.yview)
        
        self.tm_tree.heading('match', text='Match %', anchor='w')
        self.tm_tree.heading('text', text='Translation', anchor='w')
        self.tm_tree.heading('source_tm', text='TM Source', anchor='w')
        
        self.tm_tree.column('match', width=70, minwidth=70, stretch=False)
        self.tm_tree.column('text', width=250, minwidth=150, stretch=True)
        self.tm_tree.column('source_tm', width=120, minwidth=100, stretch=False)
        
        self.tm_tree.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')
        
        # Tag configuration for match percentage colors
        self.tm_tree.tag_configure('exact', background='#c8e6c9')  # Light green
        self.tm_tree.tag_configure('high', background='#fff9c4')   # Light yellow
        self.tm_tree.tag_configure('medium', background='#ffecb3') # Light orange
        
        self.tm_tree.bind('<<TreeviewSelect>>', self.on_tm_select)
        self.tm_tree.bind('<Double-Button-1>', lambda e: self.copy_suggestion_to_target())
        
        # Store TM results
        self.tm_results = []
        
        # === TM MANAGEMENT SECTION (BOTTOM) ===
        management_frame = tk.LabelFrame(parent, text="TM Settings & Management",
                                        font=('Segoe UI', 10, 'bold'), padx=5, pady=5)
        management_frame.pack(fill='x', padx=5, pady=(2, 5))
        
        # Status display
        status_container = tk.Frame(management_frame, bg='#e3f2fd', relief='solid', borderwidth=1)
        status_container.pack(fill='x', padx=2, pady=2)
        
        self.tm_context_status_label = tk.Label(status_container, 
                                        text=f"📊 {self.tm_agent.get_entry_count()} entries | Threshold: {int(self.tm_agent.fuzzy_threshold*100)}%",
                                        font=('Segoe UI', 9), bg='#e3f2fd')
        self.tm_context_status_label.pack(pady=5)
        
        # Action buttons in a single row
        tm_buttons = tk.Frame(management_frame)
        tm_buttons.pack(fill='x', padx=2, pady=5)
        
        tk.Button(tm_buttons, text="📂 Load TM File", command=self.load_tm_file,
                 bg='#4CAF50', fg='white', font=('Segoe UI', 9), width=14).pack(side='left', padx=2)
        tk.Button(tm_buttons, text="⚙️ TM Manager", command=self.show_tm_manager,
                 bg='#2196F3', fg='white', font=('Segoe UI', 9), width=14).pack(side='left', padx=2)
        tk.Button(tm_buttons, text="💾 Save TM", 
                 command=lambda: self.log("✓ TM saved (auto-saves with project)"),
                 bg='#FF9800', fg='white', font=('Segoe UI', 9), width=14).pack(side='left', padx=2)
    
    def create_glossary_tab(self, parent):
        """Create Glossary/Termbase tab content"""
        # Toolbar
        toolbar = tk.Frame(parent, bg='#f0f0f0')
        toolbar.pack(side='top', fill='x', padx=5, pady=5)
        
        tk.Label(toolbar, text="Glossary:", bg='#f0f0f0', font=('Segoe UI', 9)).pack(side='left', padx=2)
        self.glossary_source_var = tk.StringVar(value="Project Glossary")
        glossary_combo = ttk.Combobox(toolbar, textvariable=self.glossary_source_var,
                                     values=["Project Glossary", "Main Termbase",
                                            "Domain Glossary", "All Glossaries"],
                                     state='readonly', width=18, font=('Segoe UI', 9))
        glossary_combo.pack(side='left', padx=5)
        
        tk.Button(toolbar, text="🔍 Search", command=self.search_glossary,
                 bg='#FF9800', fg='white', font=('Segoe UI', 9)).pack(side='left', padx=5)
        
        # Treeview for glossary matches
        tree_frame = tk.Frame(parent)
        tree_frame.pack(fill='both', expand=True, padx=5, pady=5)
        
        scrollbar = ttk.Scrollbar(tree_frame, orient='vertical')
        self.glossary_tree = ttk.Treeview(tree_frame, columns=('source_term', 'target_term', 'domain'),
                                         show='headings', yscrollcommand=scrollbar.set,
                                         selectmode='browse', height=10)
        scrollbar.config(command=self.glossary_tree.yview)
        
        self.glossary_tree.heading('source_term', text='Source Term')
        self.glossary_tree.heading('target_term', text='Target Term')
        self.glossary_tree.heading('domain', text='Domain')
        
        self.glossary_tree.column('source_term', width=120, minwidth=80, stretch=True)
        self.glossary_tree.column('target_term', width=120, minwidth=80, stretch=True)
        self.glossary_tree.column('domain', width=80, minwidth=60, stretch=False)
        
        self.glossary_tree.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')
        
        self.glossary_tree.bind('<<TreeviewSelect>>', self.on_glossary_select)
        self.glossary_tree.bind('<Double-Button-1>', lambda e: self.insert_suggestion_at_cursor())
        
        # Store glossary results
        self.glossary_results = []
    
    def create_nontrans_tab(self, parent):
        """Create Non-translatables tab content"""
        # Info label
        info_frame = tk.Frame(parent, bg='#e3f2fd')
        info_frame.pack(side='top', fill='x', padx=5, pady=5)
        
        tk.Label(info_frame, text="💡 Non-translatable elements detected in source",
                bg='#e3f2fd', font=('Segoe UI', 9)).pack(padx=10, pady=5)
        
        # Listbox for non-translatables
        list_frame = tk.Frame(parent)
        list_frame.pack(fill='both', expand=True, padx=5, pady=5)
        
        scrollbar = ttk.Scrollbar(list_frame, orient='vertical')
        self.nontrans_listbox = tk.Listbox(list_frame, yscrollcommand=scrollbar.set,
                                          font=('Segoe UI', 9), selectmode='single')
        scrollbar.config(command=self.nontrans_listbox.yview)
        
        self.nontrans_listbox.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')
        
        self.nontrans_listbox.bind('<<ListboxSelect>>', self.on_nontrans_select)
        self.nontrans_listbox.bind('<Double-Button-1>', lambda e: self.insert_suggestion_at_cursor())
        
        # Store non-translatables
        self.nontrans_results = []
        
    # Assistance panel action methods (placeholder implementations for now)
    
    def get_mt_translation(self):
        """Get machine translation for current segment"""
        if not hasattr(self, 'current_segment') or not self.current_segment:
            self.log("⚠ No segment selected")
            return
        
        provider = self.mt_provider_var.get()
        source_text = self.current_segment.source
        
        self.log(f"🤖 Requesting MT from {provider}...")
        
        # Placeholder: In real implementation, call MT API
        # For now, show example results
        self.mt_results = [
            {"provider": provider, "text": f"[MT] {source_text} (translated)", "confidence": 0.95},
            {"provider": provider, "text": f"[MT Alt] {source_text} (alternative)", "confidence": 0.85}
        ]
        
        self.mt_listbox.delete(0, 'end')
        for i, result in enumerate(self.mt_results):
            conf_pct = int(result['confidence'] * 100)
            self.mt_listbox.insert('end', f"{conf_pct}% - {result['text']}")
        
        self.log(f"✓ Received {len(self.mt_results)} MT suggestions")
    
    def get_llm_translation(self):
        """Get LLM translation for current segment"""
        if not hasattr(self, 'current_segment') or not self.current_segment:
            self.log("⚠ No segment selected")
            return
        
        model = self.llm_model_var.get()
        prompt_type = self.llm_prompt_var.get()
        source_text = self.current_segment.source
        
        self.log(f"✨ Requesting LLM translation from {model} with {prompt_type}...")
        
        # Placeholder: In real implementation, call LLM API
        # This is where Supervertaler integration will happen!
        self.llm_results = [
            {"model": model, "prompt": prompt_type, "text": f"[LLM] {source_text} (professional translation)", "quality": 0.98},
            {"model": model, "prompt": prompt_type, "text": f"[LLM Alt] {source_text} (alternative phrasing)", "quality": 0.92}
        ]
        
        self.llm_listbox.delete(0, 'end')
        for i, result in enumerate(self.llm_results):
            quality_pct = int(result['quality'] * 100)
            self.llm_listbox.insert('end', f"{quality_pct}% - {result['text']}")
        
        self.log(f"✓ Received {len(self.llm_results)} LLM suggestions")
        self.log(f"💡 Future: This will integrate with Supervertaler prompts and models")
    
    def schedule_auto_tm_search(self):
        """Schedule automatic TM search after a delay (2 seconds)"""
        # Cancel any pending search
        if self.tm_auto_search_timer is not None:
            self.root.after_cancel(self.tm_auto_search_timer)
            self.tm_auto_search_timer = None
        
        # Schedule new search after 2 seconds (2000 ms)
        # This prevents searches when user is quickly navigating between segments
        self.tm_auto_search_timer = self.root.after(2000, self.auto_search_tm)
        self.log("⏱ Auto-search scheduled (2 seconds)...")
    
    def auto_search_tm(self):
        """Automatically search TM (called after delay)"""
        self.tm_auto_search_timer = None  # Clear timer reference
        
        if not hasattr(self, 'current_segment') or not self.current_segment:
            self.log("⚠ Auto-search cancelled: No segment selected")
            return
        
        # Trigger the search (auto_triggered flag prevents logging)
        self.log("🔍 Auto-search executing...")
        self.search_tm(auto_triggered=True)
    
    def search_tm(self, auto_triggered=False):
        """Search translation memory for matches across enabled TMs"""
        if not hasattr(self, 'current_segment') or not self.current_segment:
            if not auto_triggered:  # Only log if manually triggered
                self.log("⚠ No segment selected")
            return
        
        # Check if TM tree exists
        if not hasattr(self, 'tm_tree'):
            self.log("⚠ Auto-search cancelled: TM tree not initialized")
            return
        
        tm_source = self.tm_source_var.get()
        threshold = int(self.tm_threshold_var.get())
        source_text = self.current_segment.source
        
        if not auto_triggered:  # Only log if manually triggered
            self.log(f"🔍 Searching {tm_source} (min {threshold}% match)...")
        
        # Determine which TMs to search
        if tm_source == "All Active TMs":
            # Search all enabled TMs
            matches = self.tm_database.search_all(source_text, enabled_only=True)
        else:
            # Parse selected TM name (format: "✓ TM Name" or "✗ TM Name")
            tm_name = tm_source.split(' ', 1)[1] if ' ' in tm_source else tm_source
            
            # Find matching TM by name
            tm_ids = []
            for tm in self.tm_database.get_all_tms(enabled_only=False):
                if tm.name == tm_name:
                    tm_ids.append(tm.tm_id)
                    break
            
            matches = self.tm_database.search_all(source_text, tm_ids=tm_ids, enabled_only=False)
        
        if auto_triggered:
            self.log(f"🔍 Auto-search found {len(matches)} raw matches (before threshold filter)")
        
        # Filter by threshold and format for display
        self.tm_results = []
        for match in matches:
            if match['match_pct'] >= threshold:
                self.tm_results.append({
                    "match": match['match_pct'],
                    "source": match['source'],
                    "target": match['target'],
                    "tm": match['tm_name']
                })
        
        if auto_triggered:
            self.log(f"✓ Auto-search: {len(self.tm_results)} matches above {threshold}% threshold")
        
        # Clear and populate tree
        for item in self.tm_tree.get_children():
            self.tm_tree.delete(item)
        
        for result in self.tm_results:
            match_pct = result['match']
            if match_pct == 100:
                tag = 'exact'
            elif match_pct >= 90:
                tag = 'high'
            else:
                tag = 'medium'
            
            # Use separate column for TM source
            self.tm_tree.insert('', 'end', values=(f"{match_pct}%", result['target'], result['tm']), tags=(tag,))
        
        if auto_triggered:
            self.log(f"✓ Auto-search: Populated TM tree with {len(self.tm_results)} matches")
        
        if not auto_triggered:  # Only log if manually triggered
            self.log(f"✓ Found {len(self.tm_results)} TM matches")
    
    def search_glossary(self):
        """Search glossary/termbase for terms"""
        if not hasattr(self, 'current_segment') or not self.current_segment:
            self.log("⚠ No segment selected")
            return
        
        glossary = self.glossary_source_var.get()
        source_text = self.current_segment.source
        
        self.log(f"📚 Searching {glossary}...")
        
        # Placeholder: In real implementation, search glossary database
        # Extract potential terms from source
        self.glossary_results = [
            {"source_term": "pivot point", "target_term": "draaipunt", "domain": "Technical"},
            {"source_term": "coupling bar", "target_term": "koppelstang", "domain": "Technical"},
            {"source_term": "frame", "target_term": "frame", "domain": "General"}
        ]
        
        # Clear and populate tree
        for item in self.glossary_tree.get_children():
            self.glossary_tree.delete(item)
        
        for result in self.glossary_results:
            self.glossary_tree.insert('', 'end', values=(
                result['source_term'],
                result['target_term'],
                result['domain']
            ))
        
        self.log(f"✓ Found {len(self.glossary_results)} glossary terms")
    
    def on_mt_select(self, event):
        """Handle MT suggestion selection"""
        selection = self.mt_listbox.curselection()
        if selection and self.mt_results:
            idx = selection[0]
            result = self.mt_results[idx]
            self.update_suggestion_detail(self.current_segment.source if hasattr(self, 'current_segment') else "",
                                         result['text'])
    
    def on_llm_select(self, event):
        """Handle LLM suggestion selection"""
        selection = self.llm_listbox.curselection()
        if selection and self.llm_results:
            idx = selection[0]
            result = self.llm_results[idx]
            self.update_suggestion_detail(self.current_segment.source if hasattr(self, 'current_segment') else "",
                                         result['text'])
    
    def on_tm_select(self, event):
        """Handle TM match selection"""
        selection = self.tm_tree.selection()
        if selection and self.tm_results:
            idx = self.tm_tree.index(selection[0])
            result = self.tm_results[idx]
            self.update_suggestion_detail(result['source'], result['target'])
    
    def on_glossary_select(self, event):
        """Handle glossary term selection"""
        selection = self.glossary_tree.selection()
        if selection and self.glossary_results:
            idx = self.glossary_tree.index(selection[0])
            result = self.glossary_results[idx]
            self.update_suggestion_detail(result['source_term'], result['target_term'])
    
    def on_nontrans_select(self, event):
        """Handle non-translatable selection"""
        selection = self.nontrans_listbox.curselection()
        if selection and self.nontrans_results:
            idx = selection[0]
            result = self.nontrans_results[idx]
            self.update_suggestion_detail("Non-translatable", result)
    
    def update_suggestion_detail(self, source, target):
        """Update the suggestion detail panel with source/target"""
        # Update source
        self.assist_detail_source.config(state='normal')
        self.assist_detail_source.delete('1.0', 'end')
        self.assist_detail_source.insert('1.0', source)
        self.assist_detail_source.config(state='disabled')
        
        # Update target
        self.assist_detail_target.config(state='normal')
        self.assist_detail_target.delete('1.0', 'end')
        self.assist_detail_target.insert('1.0', target)
        self.assist_detail_target.config(state='disabled')
    
    def copy_suggestion_to_target(self):
        """Copy the selected suggestion to the target field"""
        target_text = self.assist_detail_target.get('1.0', 'end-1c')
        if not target_text.strip():
            self.log("⚠ No suggestion selected")
            return
        
        # Insert into grid editor if visible
        if hasattr(self, 'grid_target_text'):
            self.grid_target_text.delete('1.0', 'end')
            self.grid_target_text.insert('1.0', target_text)
            self.log("✓ Suggestion copied to target")
    
    def insert_suggestion_at_cursor(self):
        """Insert the selected suggestion at cursor position in target field"""
        target_text = self.assist_detail_target.get('1.0', 'end-1c')
        if not target_text.strip():
            self.log("⚠ No suggestion selected")
            return
        
        # Insert at cursor position in grid editor
        if hasattr(self, 'grid_target_text'):
            self.grid_target_text.insert('insert', target_text)
            self.log("✓ Suggestion inserted at cursor")
    
    def create_split_layout(self):
        """Create List View layout (list with editor panel)"""
        
        # Create main horizontal paned window (list/editor on left, assistance panel on right)
        self.main_paned = ttk.PanedWindow(self.content_frame, orient='horizontal')
        self.main_paned.pack(fill='both', expand=True)
        
        # Left side: Grid and editor
        left_container = tk.Frame(self.main_paned)
        self.main_paned.add(left_container, weight=3)
        
        # Grid frame (top part)
        grid_frame = tk.LabelFrame(left_container, text="Translation Grid", padx=5, pady=5)
        grid_frame.pack(side='top', fill='both', expand=True)
        
        # Filter panel (above the treeview) - using same filter variables as Grid View
        filter_frame = tk.Frame(grid_frame, bg='#f0f0f0', relief='ridge', borderwidth=1)
        filter_frame.pack(side='top', fill='x', pady=(0, 5))
        
        # Filter label
        tk.Label(filter_frame, text="🔍 Filter:", bg='#f0f0f0', 
                font=('Segoe UI', 9, 'bold')).pack(side='left', padx=5, pady=5)
        
        # Filter mode selection - always create new buttons for this view
        mode_frame = tk.Frame(filter_frame, bg='#f0f0f0', relief='solid', bd=1)
        mode_frame.pack(side='left', padx=5)
        
        # Create buttons with current mode state
        filter_btn = tk.Button(mode_frame, text="🔍 Filter",
                              command=lambda: self.set_filter_mode('filter'),
                              bg='#4CAF50' if self.filter_mode == 'filter' else '#e0e0e0',
                              fg='white' if self.filter_mode == 'filter' else '#666',
                              font=('Segoe UI', 9, 'bold') if self.filter_mode == 'filter' else ('Segoe UI', 9),
                              relief='sunken' if self.filter_mode == 'filter' else 'raised',
                              bd=2, cursor='hand2', width=10)
        filter_btn.pack(side='left', padx=1, pady=1)
        
        highlight_btn = tk.Button(mode_frame, text="💡 Highlight",
                                 command=lambda: self.set_filter_mode('highlight'),
                                 bg='#FFA500' if self.filter_mode == 'highlight' else '#e0e0e0',
                                 fg='white' if self.filter_mode == 'highlight' else '#666',
                                 font=('Segoe UI', 9, 'bold') if self.filter_mode == 'highlight' else ('Segoe UI', 9),
                                 relief='sunken' if self.filter_mode == 'highlight' else 'raised',
                                 bd=2, cursor='hand2', width=10)
        highlight_btn.pack(side='left', padx=1, pady=1)
        
        # Store references for this view
        self.list_filter_mode_btn = filter_btn
        self.list_highlight_mode_btn = highlight_btn
        # Also update main references to point to these buttons
        self.filter_mode_btn = filter_btn
        self.highlight_mode_btn = highlight_btn
        
        tk.Label(filter_frame, text="│", bg='#f0f0f0', fg='#ccc',
                font=('Segoe UI', 10)).pack(side='left', padx=5)
        
        # Source filter - using shared variables
        if not hasattr(self, 'filter_source_var'):
            self.filter_source_var = tk.StringVar()
        
        tk.Label(filter_frame, text="Source:", bg='#f0f0f0',
                font=('Segoe UI', 9)).pack(side='left', padx=(0, 2))
        source_entry = tk.Entry(filter_frame, textvariable=self.filter_source_var, width=20,
                               font=('Segoe UI', 9))
        source_entry.pack(side='left', padx=(0, 10))
        # Bind Enter key to apply filters
        source_entry.bind('<Return>', lambda e: self.apply_filters())
        
        # Target filter
        if not hasattr(self, 'filter_target_var'):
            self.filter_target_var = tk.StringVar()
        
        tk.Label(filter_frame, text="Target:", bg='#f0f0f0',
                font=('Segoe UI', 9)).pack(side='left', padx=(0, 2))
        target_entry = tk.Entry(filter_frame, textvariable=self.filter_target_var, width=20,
                               font=('Segoe UI', 9))
        target_entry.pack(side='left', padx=(0, 10))
        # Bind Enter key to apply filters
        target_entry.bind('<Return>', lambda e: self.apply_filters())
        
        # Status filter
        if not hasattr(self, 'filter_status_var'):
            self.filter_status_var = tk.StringVar(value="All")
        
        tk.Label(filter_frame, text="Status:", bg='#f0f0f0',
                font=('Segoe UI', 9)).pack(side='left', padx=(0, 2))
        status_combo = ttk.Combobox(filter_frame, textvariable=self.filter_status_var,
                                   values=["All", "untranslated", "translated", "approved"],
                                   state='readonly', width=12, font=('Segoe UI', 9))
        status_combo.pack(side='left', padx=(0, 10))
        # Combobox selection triggers apply automatically (user expectation for dropdowns)
        status_combo.bind('<<ComboboxSelected>>', lambda e: self.apply_filters())
        
        # Apply button
        tk.Button(filter_frame, text="🔍 Apply", command=self.apply_filters,
                 bg='#2196F3', fg='white', font=('Segoe UI', 9, 'bold'),
                 relief='raised', cursor='hand2', padx=8).pack(side='left', padx=5)
        
        # Clear button
        tk.Button(filter_frame, text="✕ Clear", command=self.clear_filters,
                 bg='#ff5252', fg='white', font=('Segoe UI', 8),
                 relief='raised', cursor='hand2').pack(side='left', padx=5)
        
        # Results label - always create new for this view
        self.filter_results_label = tk.Label(filter_frame, text="", bg='#f0f0f0',
                                            font=('Segoe UI', 9, 'italic'), fg='#666')
        self.filter_results_label.pack(side='left', padx=10)
        
        # Create a frame to hold treeview and scrollbars with grid layout
        tree_container = tk.Frame(grid_frame)
        tree_container.pack(side='top', fill='both', expand=True)
        
        # Create treeview for segments (parent is tree_container now)
        self.tree = ttk.Treeview(tree_container,
                                columns=('id', 'type', 'style', 'status', 'source', 'target'),
                                show='headings',
                                selectmode='browse')
        
        # Define columns
        self.tree.heading('id', text='#')
        self.tree.heading('type', text='Type')
        self.tree.heading('style', text='Style')
        self.tree.heading('status', text='Status')
        self.tree.heading('source', text='Source')
        self.tree.heading('target', text='Target')
        
        self.tree.column('id', width=40, minwidth=40, anchor='center', stretch=False)
        self.tree.column('type', width=65, minwidth=60, anchor='center', stretch=False)
        self.tree.column('style', width=80, minwidth=70, anchor='w', stretch=False)
        self.tree.column('status', width=95, minwidth=90, anchor='center', stretch=False)
        self.tree.column('source', width=400, minwidth=200, anchor='w', stretch=True)
        self.tree.column('target', width=400, minwidth=200, anchor='w', stretch=True)
        
        # Scrollbars
        v_scroll = ttk.Scrollbar(tree_container, orient='vertical', command=self.tree.yview)
        h_scroll = ttk.Scrollbar(tree_container, orient='horizontal', command=self.tree.xview)
        self.tree.configure(yscrollcommand=v_scroll.set, xscrollcommand=h_scroll.set)
        
        self.tree.grid(row=0, column=0, sticky='nsew')
        v_scroll.grid(row=0, column=1, sticky='ns')
        h_scroll.grid(row=1, column=0, sticky='ew')
        
        tree_container.grid_rowconfigure(0, weight=1)
        tree_container.grid_columnconfigure(0, weight=1)
        
        # Configure row colors
        self.tree.tag_configure('untranslated', background='#ffe6e6')
        self.tree.tag_configure('translated', background='#e6ffe6')
        self.tree.tag_configure('approved', background='#e6f3ff')
        self.tree.tag_configure('table_cell', foreground='#0066cc', font=('TkDefaultFont', 9, 'italic'))
        
        # Configure style-specific formatting
        self.tree.tag_configure('heading1', font=('TkDefaultFont', 10, 'bold'), foreground='#003366')
        self.tree.tag_configure('heading2', font=('TkDefaultFont', 9, 'bold'), foreground='#0066cc')
        self.tree.tag_configure('heading3', font=('TkDefaultFont', 9, 'bold'), foreground='#3399ff')
        self.tree.tag_configure('title', font=('TkDefaultFont', 11, 'bold'), foreground='#663399')
        self.tree.tag_configure('subtitle', font=('TkDefaultFont', 9, 'italic'), foreground='#663399')
        
        # Bind events
        self.tree.bind('<<TreeviewSelect>>', self.on_segment_select)
        self.tree.bind('<Return>', lambda e: self.focus_target_editor())
        self.tree.bind('<Double-1>', lambda e: self.focus_target_editor())
        
        # Editor frame (bottom part)
        editor_frame = tk.LabelFrame(left_container, text="Segment Editor", padx=10, pady=10)
        editor_frame.pack(side='bottom', fill='x', pady=(5, 0))
        
        # Segment info
        info_frame = tk.Frame(editor_frame)
        info_frame.pack(fill='x', pady=(0, 5))
        
        self.seg_info_label = tk.Label(info_frame, text="No segment selected",
                                       font=('Segoe UI', 9, 'bold'))
        self.seg_info_label.pack(side='left')
        
        # Status selector
        status_frame = tk.Frame(info_frame)
        status_frame.pack(side='right')
        
        tk.Label(status_frame, text="Status:").pack(side='left', padx=(0, 5))
        self.status_var = tk.StringVar(value="untranslated")
        self.status_combo = ttk.Combobox(status_frame, textvariable=self.status_var,
                                        values=["untranslated", "translated", "approved"],
                                        state='readonly', width=12)
        self.status_combo.pack(side='left')
        self.status_combo.bind('<<ComboboxSelected>>', self.on_status_change)
        
        # Source (read-only)
        tk.Label(editor_frame, text="Source:", font=('Segoe UI', 9, 'bold')).pack(anchor='w')
        self.source_text = tk.Text(editor_frame, height=2, wrap='word', bg='#f5f5f5',
                                   state='disabled', font=('Segoe UI', 10))
        self.source_text.pack(fill='x', pady=(2, 10))
        
        # Target (editable)
        target_header_frame = tk.Frame(editor_frame)
        target_header_frame.pack(fill='x', anchor='w')
        tk.Label(target_header_frame, text="Target:", font=('Segoe UI', 9, 'bold')).pack(side='left')
        self.tag_validation_label = tk.Label(target_header_frame, text="", font=('Segoe UI', 8))
        self.tag_validation_label.pack(side='left', padx=(10, 0))
        
        self.target_text = tk.Text(editor_frame, height=2, wrap='word', font=('Segoe UI', 10))
        self.target_text.pack(fill='x', pady=(2, 5))
        self.target_text.bind('<KeyRelease>', self.on_target_change)
        self.target_text.bind('<Control-Return>', lambda e: self.save_segment_and_next())
        self.target_text.bind('<Control-b>', lambda e: self.insert_tag('b'))
        self.target_text.bind('<Control-i>', lambda e: self.insert_tag('i'))
        self.target_text.bind('<Control-u>', lambda e: self.insert_tag('u'))
        
        # Tag buttons - wrapping layout for narrow screens
        tag_button_frame = tk.Frame(editor_frame)
        tag_button_frame.pack(fill='x', pady=(0, 10))
        
        # Container for better organization
        tag_container = tk.Frame(tag_button_frame, relief='flat')
        tag_container.pack(fill='x')
        
        # Row 1: Label and formatting buttons
        tag_row1 = tk.Frame(tag_container)
        tag_row1.pack(fill='x', pady=2)
        
        tk.Label(tag_row1, text="Insert:").pack(side='left', padx=(0, 5))
        tk.Button(tag_row1, text="<b>Bold</b>", command=lambda: self.insert_tag('b'),
                 relief='flat', bg='#ffcccc', font=('Segoe UI', 8)).pack(side='left', padx=2)
        tk.Button(tag_row1, text="<i>Italic</i>", command=lambda: self.insert_tag('i'),
                 relief='flat', bg='#ccccff', font=('Segoe UI', 8, 'italic')).pack(side='left', padx=2)
        tk.Button(tag_row1, text="<u>Underline</u>", command=lambda: self.insert_tag('u'),
                 relief='flat', bg='#ccffcc', font=('Segoe UI', 8, 'underline')).pack(side='left', padx=2)
        tk.Button(tag_row1, text="Strip Tags", command=self.strip_tags_from_target,
                 relief='flat', bg='#eeeeee', font=('Segoe UI', 8)).pack(side='left', padx=10)
        tk.Button(tag_row1, text="Copy Source Tags", command=self.copy_source_tags,
                 relief='flat', bg='#e6f3ff', font=('Segoe UI', 8)).pack(side='left', padx=2)
        
        # Action buttons - wrapping layout for narrow screens
        button_frame = tk.Frame(editor_frame)
        button_frame.pack(fill='x')
        
        # Left side buttons
        left_buttons = tk.Frame(button_frame)
        left_buttons.pack(side='left', fill='x', expand=True)
        
        tk.Button(left_buttons, text="Copy Source → Target", command=self.copy_source_to_target,
                 font=('Segoe UI', 8)).pack(side='left', padx=(0, 5))
        tk.Button(left_buttons, text="Clear Target", command=self.clear_target,
                 font=('Segoe UI', 8)).pack(side='left', padx=(0, 5))
        
        # Right side button (most important action)
        tk.Button(button_frame, text="Save & Next (Ctrl+Enter)", command=self.save_segment_and_next,
                 bg='#4CAF50', fg='white', font=('Segoe UI', 8, 'bold')).pack(side='right')
        
        # Right side: Assistance panel
        self.create_assistance_panel()
    
    def create_document_layout(self):
        """Create Document View layout - shows text in document flow with clickable segments"""
        
        # Create main horizontal paned window (document/editor on left, assistance panel on right)
        self.main_paned = ttk.PanedWindow(self.content_frame, orient='horizontal')
        self.main_paned.pack(fill='both', expand=True)
        
        # Left side: Document and editor
        left_container = tk.Frame(self.main_paned)
        self.main_paned.add(left_container, weight=3)
        
        # Filter panel (at the top) - using same filter variables as Grid View
        filter_frame = tk.Frame(left_container, bg='#f0f0f0', relief='ridge', borderwidth=1)
        filter_frame.pack(side='top', fill='x', pady=(0, 5))
        
        # Filter label
        tk.Label(filter_frame, text="🔍 Filter:", bg='#f0f0f0', 
                font=('Segoe UI', 9, 'bold')).pack(side='left', padx=5, pady=5)
        
        # Filter mode selection - always create new buttons for this view
        mode_frame = tk.Frame(filter_frame, bg='#f0f0f0', relief='solid', bd=1)
        mode_frame.pack(side='left', padx=5)
        
        # Create buttons with current mode state
        filter_btn = tk.Button(mode_frame, text="🔍 Filter",
                              command=lambda: self.set_filter_mode('filter'),
                              bg='#4CAF50' if self.filter_mode == 'filter' else '#e0e0e0',
                              fg='white' if self.filter_mode == 'filter' else '#666',
                              font=('Segoe UI', 9, 'bold') if self.filter_mode == 'filter' else ('Segoe UI', 9),
                              relief='sunken' if self.filter_mode == 'filter' else 'raised',
                              bd=2, cursor='hand2', width=10)
        filter_btn.pack(side='left', padx=1, pady=1)
        
        highlight_btn = tk.Button(mode_frame, text="💡 Highlight",
                                 command=lambda: self.set_filter_mode('highlight'),
                                 bg='#FFA500' if self.filter_mode == 'highlight' else '#e0e0e0',
                                 fg='white' if self.filter_mode == 'highlight' else '#666',
                                 font=('Segoe UI', 9, 'bold') if self.filter_mode == 'highlight' else ('Segoe UI', 9),
                                 relief='sunken' if self.filter_mode == 'highlight' else 'raised',
                                 bd=2, cursor='hand2', width=10)
        highlight_btn.pack(side='left', padx=1, pady=1)
        
        # Store references for this view
        self.doc_filter_mode_btn = filter_btn
        self.doc_highlight_mode_btn = highlight_btn
        # Also update main references to point to these buttons
        self.filter_mode_btn = filter_btn
        self.highlight_mode_btn = highlight_btn
        
        tk.Label(filter_frame, text="│", bg='#f0f0f0', fg='#ccc',
                font=('Segoe UI', 10)).pack(side='left', padx=5)
        
        # Source filter - using shared variables
        if not hasattr(self, 'filter_source_var'):
            self.filter_source_var = tk.StringVar()
        
        tk.Label(filter_frame, text="Source:", bg='#f0f0f0',
                font=('Segoe UI', 9)).pack(side='left', padx=(0, 2))
        source_entry = tk.Entry(filter_frame, textvariable=self.filter_source_var, width=20,
                               font=('Segoe UI', 9))
        source_entry.pack(side='left', padx=(0, 10))
        source_entry.bind('<Return>', lambda e: self.apply_filters())
        
        # Target filter
        if not hasattr(self, 'filter_target_var'):
            self.filter_target_var = tk.StringVar()
        
        tk.Label(filter_frame, text="Target:", bg='#f0f0f0',
                font=('Segoe UI', 9)).pack(side='left', padx=(0, 2))
        target_entry = tk.Entry(filter_frame, textvariable=self.filter_target_var, width=20,
                               font=('Segoe UI', 9))
        target_entry.pack(side='left', padx=(0, 10))
        target_entry.bind('<Return>', lambda e: self.apply_filters())
        
        # Status filter
        if not hasattr(self, 'filter_status_var'):
            self.filter_status_var = tk.StringVar(value="All")
        
        tk.Label(filter_frame, text="Status:", bg='#f0f0f0',
                font=('Segoe UI', 9)).pack(side='left', padx=(0, 2))
        status_combo = ttk.Combobox(filter_frame, textvariable=self.filter_status_var,
                                   values=["All", "untranslated", "translated", "approved"],
                                   state='readonly', width=12, font=('Segoe UI', 9))
        status_combo.pack(side='left', padx=(0, 10))
        status_combo.bind('<<ComboboxSelected>>', lambda e: self.apply_filters())
        
        # Apply button
        tk.Button(filter_frame, text="🔍 Apply", command=self.apply_filters,
                 bg='#2196F3', fg='white', font=('Segoe UI', 9, 'bold'),
                 relief='raised', cursor='hand2', padx=8).pack(side='left', padx=5)
        
        # Clear button
        tk.Button(filter_frame, text="✕ Clear", command=self.clear_filters,
                 bg='#ff5252', fg='white', font=('Segoe UI', 8),
                 relief='raised', cursor='hand2').pack(side='left', padx=5)
        
        # Results label - always create new for this view
        self.filter_results_label = tk.Label(filter_frame, text="", bg='#f0f0f0',
                                            font=('Segoe UI', 9, 'italic'), fg='#666')
        self.filter_results_label.pack(side='left', padx=10)
        
        # Split into document area (top) and editor panel (bottom)
        # Document area
        doc_container = tk.Frame(left_container)
        doc_container.pack(side='top', fill='both', expand=True)
        
        # Canvas for scrolling
        self.doc_canvas = tk.Canvas(doc_container, bg='white', highlightthickness=0)
        scrollbar = ttk.Scrollbar(doc_container, orient='vertical', command=self.doc_canvas.yview)
        
        self.doc_canvas.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side='right', fill='y')
        self.doc_canvas.pack(side='left', fill='both', expand=True)
        
        # Inner frame to hold content
        self.doc_inner_frame = tk.Frame(self.doc_canvas, bg='white')
        self.doc_canvas_window = self.doc_canvas.create_window((0, 0), window=self.doc_inner_frame, anchor='nw')
        
        # Bind canvas resizing to update inner frame width
        def on_canvas_resize(event):
            # Set the inner frame width to match canvas width (minus scrollbar and padding)
            canvas_width = event.width - 80  # Account for padding (40px each side)
            self.doc_canvas.itemconfig(self.doc_canvas_window, width=canvas_width + 80)
            self.doc_inner_frame.config(width=canvas_width)
        
        self.doc_canvas.bind('<Configure>', on_canvas_resize)
        
        # Update scroll region when inner frame changes
        self.doc_inner_frame.bind('<Configure>', 
                                  lambda e: self.doc_canvas.configure(scrollregion=self.doc_canvas.bbox('all')))
        
        # Bind mouse wheel
        self.doc_canvas.bind('<MouseWheel>', lambda e: self.doc_canvas.yview_scroll(int(-1*(e.delta/120)), "units"))
        self.doc_inner_frame.bind('<MouseWheel>', lambda e: self.doc_canvas.yview_scroll(int(-1*(e.delta/120)), "units"))
        
        # Editor panel (bottom) - similar to Split View
        editor_frame = tk.LabelFrame(left_container, text="Segment Editor", padx=10, pady=10)
        editor_frame.pack(side='bottom', fill='x', pady=(5, 0))
        
        # Segment info
        info_frame = tk.Frame(editor_frame)
        info_frame.pack(fill='x', pady=(0, 5))
        
        self.doc_seg_info_label = tk.Label(info_frame, text="Click on any segment in the document to edit",
                                       font=('Segoe UI', 9, 'bold'))
        self.doc_seg_info_label.pack(side='left')
        
        # Status selector
        status_frame = tk.Frame(info_frame)
        status_frame.pack(side='right')
        
        tk.Label(status_frame, text="Status:").pack(side='left', padx=(0, 5))
        self.doc_status_var = tk.StringVar(value="untranslated")
        self.doc_status_combo = ttk.Combobox(status_frame, textvariable=self.doc_status_var,
                                        values=["untranslated", "translated", "approved"],
                                        state='readonly', width=12)
        self.doc_status_combo.pack(side='left')
        self.doc_status_combo.bind('<<ComboboxSelected>>', self.on_doc_status_change)
        
        # Source (read-only)
        tk.Label(editor_frame, text="Source:", font=('Segoe UI', 9, 'bold')).pack(anchor='w')
        self.doc_source_text = tk.Text(editor_frame, height=2, wrap='word', bg='#f5f5f5',
                                   state='disabled', font=('Segoe UI', 10))
        self.doc_source_text.pack(fill='x', pady=(2, 10))
        
        # Target (editable)
        tk.Label(editor_frame, text="Target:", font=('Segoe UI', 9, 'bold')).pack(anchor='w')
        self.doc_target_text = tk.Text(editor_frame, height=2, wrap='word', font=('Segoe UI', 10))
        self.doc_target_text.pack(fill='x', pady=(2, 5))
        self.doc_target_text.bind('<KeyRelease>', self.on_doc_target_change)
        self.doc_target_text.bind('<Control-Return>', lambda e: self.save_doc_segment_and_next())
        self.doc_target_text.bind('<Control-d>', lambda e: self.copy_source_to_target_doc())
        
        # Action buttons - wrapping layout for narrow screens
        button_frame = tk.Frame(editor_frame)
        button_frame.pack(fill='x')
        
        # Left side buttons
        left_buttons = tk.Frame(button_frame)
        left_buttons.pack(side='left', fill='x', expand=True)
        
        tk.Button(left_buttons, text="Copy Source → Target", command=self.copy_source_to_target_doc,
                 font=('Segoe UI', 8)).pack(side='left', padx=(0, 5))
        tk.Button(left_buttons, text="Clear Target", command=self.clear_doc_target,
                 font=('Segoe UI', 8)).pack(side='left', padx=(0, 5))
        
        # Right side button (most important action)
        tk.Button(button_frame, text="Save & Next (Ctrl+Enter)", command=self.save_doc_segment_and_next,
                 bg='#4CAF50', fg='white', font=('Segoe UI', 8, 'bold')).pack(side='right')
        
        # Store segment widgets for later reference
        self.doc_segment_widgets = {}
        self.doc_current_segment = None
        
        # Right side: Assistance panel (pass main_paned as parent)
        self.create_assistance_panel(parent_paned=self.main_paned)
    
    def load_segments_to_document(self):
        """Load segments into document view (FAST - single Text widget approach)"""
        # Clear existing content
        for widget in self.doc_inner_frame.winfo_children():
            widget.destroy()
        
        self.doc_segment_widgets = {}
        
        # Use filtered segments if filter is active, otherwise all segments
        segments_to_show = self.filtered_segments if self.filter_active else self.segments
        
        # Debug: Check if we have segments to show
        if not segments_to_show:
            # No segments to display
            no_content_label = tk.Label(self.doc_inner_frame, 
                                       text="No segments to display.\n\nPlease import a document first.",
                                       bg='white', fg='#666', font=('Segoe UI', 12),
                                       pady=50)
            no_content_label.pack(fill='both', expand=True)
            self.log(f"⚠ Document view: No segments to display")
            return
        
        self.log(f"📄 Loading {len(segments_to_show)} segments to document view...")
        
        # OPTIMIZATION: Use single scrollable Text widget instead of 355 separate widgets
        # This is MUCH faster - one widget instead of hundreds
        
        # Create a single Text widget for the entire document
        doc_text = tk.Text(self.doc_inner_frame, wrap='word', bg='white', relief='flat',
                          font=('Segoe UI', 11), fg='#000000', 
                          highlightthickness=0, borderwidth=0,
                          padx=40, pady=30, 
                          spacing1=0,  # Space above each line
                          spacing2=0,  # Space between wrapped lines  
                          spacing3=12,  # Space below each paragraph (segment)
                          cursor='arrow', state='normal')
        doc_text.pack(fill='both', expand=True)
        
        # Bind mouse wheel
        doc_text.bind('<MouseWheel>', lambda e: doc_text.yview_scroll(int(-1*(e.delta/120)), "units"))
        
        # Insert all segments into the single Text widget
        prev_paragraph_id = None
        
        for i, seg in enumerate(segments_to_show):
            # Determine paragraph break vs. sentence flow
            if i > 0:
                # Check if this is a new paragraph (different paragraph_id)
                if seg.paragraph_id != prev_paragraph_id:
                    # New paragraph - add double newline
                    doc_text.insert('end', '\n\n')
                else:
                    # Same paragraph - just add a space
                    doc_text.insert('end', ' ')
            
            # Track current paragraph
            prev_paragraph_id = seg.paragraph_id
            
            # Create unique tag for this segment
            tag_name = f"seg_{seg.id}"
            
            # Determine what to display
            if seg.target and seg.target.strip():
                display_text = seg.target
            elif seg.modified and seg.target == '':
                display_text = f"[Segment {seg.id} - Empty - Click to edit]"
            elif seg.source:
                display_text = seg.source
            else:
                display_text = f"[Segment {seg.id} - Empty segment]"
            
            # Insert segment text with tag
            start_pos = doc_text.index('insert')
            doc_text.insert('end', display_text, tag_name)
            end_pos = doc_text.index('insert')
            
            # Configure tag appearance based on status
            if seg.status == 'untranslated':
                bg_color = '#ffe6e6'
                hover_bg = '#ffcccc'
            elif seg.status == 'draft':
                bg_color = '#fff9e6'
                hover_bg = '#ffe6b3'
            elif seg.status == 'translated':
                bg_color = '#e6ffe6'
                hover_bg = '#ccffcc'
            elif seg.status == 'approved':
                bg_color = '#e6f3ff'
                hover_bg = '#cce6ff'
            else:
                bg_color = 'white'
                hover_bg = '#f0f0f0'
            
            # Apply tag styling with proper spacing and borders
            doc_text.tag_config(tag_name, 
                              background=bg_color, 
                              relief='flat',
                              borderwidth=0,
                              lmargin1=0,
                              lmargin2=0,
                              rmargin=0)
            
            # Bind click event
            doc_text.tag_bind(tag_name, '<Button-1>', 
                             lambda e, s=seg, dt=doc_text, tn=tag_name: self.on_doc_segment_click(s, dt, tn))
            
            # Bind hover effects
            doc_text.tag_bind(tag_name, '<Enter>', 
                             lambda e, dt=doc_text, tn=tag_name, hbg=hover_bg: dt.tag_config(tn, background=hbg, relief='raised'))
            doc_text.tag_bind(tag_name, '<Leave>', 
                             lambda e, dt=doc_text, tn=tag_name, bg=bg_color: dt.tag_config(tn, background=bg, relief='flat'))
            
            # Store widget reference
            self.doc_segment_widgets[seg.id] = {
                'text_widget': doc_text,
                'tag_name': tag_name,
                'segment': seg,
                'start': start_pos,
                'end': end_pos
            }
        
        # Make document read-only
        doc_text.config(state='disabled')
        
        self.log(f"✓ Document view loaded ({len(segments_to_show)} segments)")
    
    def render_paragraph(self, para_segments):
        """Render a paragraph with its segments"""
        # Check if this is a heading or special style
        first_seg = para_segments[0]
        style = first_seg.style
        
        # Create paragraph frame with padding
        para_frame = tk.Frame(self.doc_inner_frame, bg='white')
        para_frame.pack(fill='x', pady=(0, 15), padx=40, anchor='w')
        
        # Determine text style based on paragraph style
        font_size = 11
        font_weight = 'normal'
        font_style = 'roman'
        text_color = '#000000'
        
        if 'Heading 1' in style or 'Title' in style:
            font_size = 16
            font_weight = 'bold'
            text_color = '#003366'
        elif 'Heading 2' in style:
            font_size = 14
            font_weight = 'bold'
            text_color = '#0066cc'
        elif 'Heading 3' in style:
            font_size = 12
            font_weight = 'bold'
            text_color = '#3399ff'
        elif 'Subtitle' in style:
            font_size = 12
            font_style = 'italic'
            text_color = '#663399'
        
        # Create a Text widget for the paragraph (allows inline segments)
        # Set height to a large value initially, will be adjusted after content is added
        para_text = tk.Text(para_frame, wrap='word', bg='white', relief='flat',
                          font=('Segoe UI', font_size, font_weight),
                          fg=text_color, highlightthickness=0, borderwidth=0,
                          cursor='arrow', state='normal', height=50)  # Start with large height
        para_text.pack(fill='both', expand=True)
        
        # Bind mouse wheel to this text widget too
        para_text.bind('<MouseWheel>', lambda e: self.doc_canvas.yview_scroll(int(-1*(e.delta/120)), "units"))
        
        # Insert each segment as a clickable span
        for i, seg in enumerate(para_segments):
            # Add space between sentences (except for first)
            if i > 0:
                para_text.insert('end', ' ')
            
            # Create tag for this segment
            tag_name = f"seg_{seg.id}"
            
            # Determine what to display:
            # - If target has content: show target
            # - If segment was modified but target is empty: show placeholder (user cleared it)
            # - If segment not modified yet: show source (initial state)
            if seg.target and seg.target.strip():
                # Target has content
                display_text = seg.target
            elif seg.modified and seg.target == '':
                # Segment was edited but target was cleared
                display_text = f"[Segment {seg.id} - Empty - Click to edit]"
            elif seg.source:
                # Not yet translated or target is empty, show source
                display_text = seg.source
            else:
                # Both source and target empty
                display_text = f"[Segment {seg.id} - Empty segment]"
            
            start_pos = para_text.index('insert')
            para_text.insert('end', display_text, tag_name)
            end_pos = para_text.index('insert')
            
            # Configure tag appearance based on status
            if seg.status == 'untranslated':
                bg_color = '#ffe6e6'
                hover_bg = '#ffcccc'
            elif seg.status == 'draft':
                bg_color = '#fff9e6'
                hover_bg = '#ffe6b3'
            elif seg.status == 'translated':
                bg_color = '#e6ffe6'
                hover_bg = '#ccffcc'
            elif seg.status == 'approved':
                bg_color = '#e6f3ff'
                hover_bg = '#cce6ff'
            else:
                bg_color = 'white'
                hover_bg = '#f0f0f0'
            
            # Apply tag styling
            para_text.tag_config(tag_name, background=bg_color, relief='flat')
            
            # Apply filter highlighting if active and segment matches
            if self.filter_active and self.should_highlight_segment(seg):
                # Highlight only the specific search terms within the segment
                self.highlight_search_terms_in_segment(para_text, start_pos, end_pos, seg, tag_name)
            
            # Bind click event to enter edit mode
            para_text.tag_bind(tag_name, '<Button-1>', 
                             lambda e, s=seg, pt=para_text, tn=tag_name: self.on_doc_segment_click(s, pt, tn))
            
            # Bind hover effects
            para_text.tag_bind(tag_name, '<Enter>', 
                             lambda e, pt=para_text, tn=tag_name, hbg=hover_bg: pt.tag_config(tn, background=hbg, relief='raised'))
            para_text.tag_bind(tag_name, '<Leave>', 
                             lambda e, pt=para_text, tn=tag_name, bg=bg_color: pt.tag_config(tn, background=bg, relief='flat'))
            
            # Store widget reference
            self.doc_segment_widgets[seg.id] = {
                'text_widget': para_text,
                'tag_name': tag_name,
                'segment': seg,
                'start': start_pos,
                'end': end_pos
            }
        
        # Make paragraph read-only
        para_text.config(state='disabled')
        
        # OPTIMIZATION: Return text widget for batch height calculation
        # Height will be calculated later in batch for better performance
        return para_text
    
    def render_table(self, table_id, table_cells):
        """Render a table with its cells"""
        # Find table dimensions
        max_row = max(pos[0] for pos in table_cells.keys())
        max_col = max(pos[1] for pos in table_cells.keys())
        
        # Create table frame with padding
        table_frame = tk.Frame(self.doc_inner_frame, bg='white')
        table_frame.pack(fill='x', pady=(0, 15), padx=40, anchor='w')
        
        # Create table grid
        for row_idx in range(max_row + 1):
            for col_idx in range(max_col + 1):
                seg = table_cells.get((row_idx, col_idx))
                
                if seg:
                    # Determine what to display
                    if seg.target and seg.target.strip():
                        display_text = seg.target
                    elif seg.modified and seg.target == '':
                        display_text = f"[Segment {seg.id} - Empty - Click to edit]"
                    elif seg.source:
                        display_text = seg.source
                    else:
                        display_text = f"[Segment {seg.id} - Empty segment]"
                    
                    # Configure tag appearance based on status
                    if seg.status == 'untranslated':
                        bg_color = '#ffe6e6'
                        hover_bg = '#ffcccc'
                    elif seg.status == 'draft':
                        bg_color = '#fff9e6'
                        hover_bg = '#ffe6b3'
                    elif seg.status == 'translated':
                        bg_color = '#e6ffe6'
                        hover_bg = '#ccffcc'
                    elif seg.status == 'approved':
                        bg_color = '#e6f3ff'
                        hover_bg = '#cce6ff'
                    else:
                        bg_color = 'white'
                        hover_bg = '#f0f0f0'
                    
                    # Create cell frame
                    cell_frame = tk.Frame(table_frame, relief='solid', bd=1, bg=bg_color)
                    cell_frame.grid(row=row_idx, column=col_idx, sticky='nsew', padx=1, pady=1)
                    
                    # Create text widget for cell
                    cell_text = tk.Text(cell_frame, wrap='word', font=('Segoe UI', 9),
                                       bg=bg_color, relief='flat', bd=0,
                                       highlightthickness=0, cursor='arrow',
                                       state='normal', height=1, width=20)
                    cell_text.pack(fill='both', expand=True, padx=3, pady=3)
                    
                    # Create tag for this cell
                    tag_name = f"seg_{seg.id}"
                    cell_text.insert('1.0', display_text, tag_name)
                    
                    # Apply tag styling
                    cell_text.tag_config(tag_name, background=bg_color, relief='flat')
                    
                    # Apply filter highlighting if active and segment matches
                    if self.filter_active and self.should_highlight_segment(seg):
                        # Highlight only the specific search terms within the cell
                        self.highlight_search_terms_in_segment(cell_text, '1.0', 'end-1c', seg, tag_name)
                    
                    # Bind click event
                    cell_text.tag_bind(tag_name, '<Button-1>',
                                     lambda e, s=seg, ct=cell_text, tn=tag_name: self.on_doc_segment_click(s, ct, tn))
                    
                    # Bind hover effects
                    cell_text.tag_bind(tag_name, '<Enter>',
                                     lambda e, ct=cell_text, tn=tag_name, hbg=hover_bg: ct.tag_config(tn, background=hbg, relief='raised'))
                    cell_text.tag_bind(tag_name, '<Leave>',
                                     lambda e, ct=cell_text, tn=tag_name, bg=bg_color: ct.tag_config(tn, background=bg, relief='flat'))
                    
                    # Bind mouse wheel
                    cell_text.bind('<MouseWheel>', lambda e: self.doc_canvas.yview_scroll(int(-1*(e.delta/120)), "units"))
                    
                    # Make read-only
                    cell_text.config(state='disabled')
                    
                    # Calculate height
                    cell_text.update_idletasks()
                    actual_lines = int(cell_text.index('end-1c').split('.')[0])
                    cell_text.config(height=max(1, actual_lines))
                    
                    # Store widget reference
                    self.doc_segment_widgets[seg.id] = {
                        'text_widget': cell_text,
                        'tag_name': tag_name,
                        'segment': seg,
                        'start': '1.0',
                        'end': 'end'
                    }
        
        # Make columns expand equally
        for col_idx in range(max_col + 1):
            table_frame.grid_columnconfigure(col_idx, weight=1, uniform="table_col")
    
    def on_doc_segment_click(self, segment, para_text, tag_name):
        """Handle click on a segment in document view - load into editor panel"""
        # Save any current edit first
        if self.doc_current_segment and self.doc_current_segment != segment:
            self.save_doc_segment()
        
        # Highlight the selected segment in the document
        if self.doc_current_segment:
            # Remove highlight from previous segment
            old_widget_info = self.doc_segment_widgets.get(self.doc_current_segment.id)
            if old_widget_info:
                old_tag = old_widget_info['tag_name']
                old_text = old_widget_info['text_widget']
                # Restore original status color
                if self.doc_current_segment.status == 'untranslated':
                    bg_color = '#ffe6e6'
                elif self.doc_current_segment.status == 'draft':
                    bg_color = '#fff9e6'
                elif self.doc_current_segment.status == 'translated':
                    bg_color = '#e6ffe6'
                elif self.doc_current_segment.status == 'approved':
                    bg_color = '#e6f3ff'
                else:
                    bg_color = 'white'
                old_text.tag_config(old_tag, background=bg_color, relief='flat', borderwidth=0)
        
        # Set current segment
        self.doc_current_segment = segment
        self.current_segment = segment
        
        # Highlight selected segment with a border
        para_text.tag_config(tag_name, relief='solid', borderwidth=2)
        
        # Scroll to make the segment visible
        # Get the text widget's position and scroll to it
        try:
            para_text.see(tag_name + '.first')
            # Also scroll the canvas to show this text widget
            self.doc_canvas.update_idletasks()
            # Get the bbox of the paragraph's parent frame
            para_frame = para_text.master
            para_frame.update_idletasks()
            # Scroll canvas to show the frame
            self.doc_canvas.yview_moveto(0)  # Reset
            bbox = self.doc_canvas.bbox('all')
            if bbox:
                # Calculate relative position
                para_y = para_frame.winfo_y()
                total_height = bbox[3] - bbox[1]
                if total_height > 0:
                    scroll_pos = para_y / total_height
                    self.doc_canvas.yview_moveto(max(0, scroll_pos - 0.1))  # Scroll with 10% margin
        except:
            pass
        
        # Load segment into editor panel
        self.doc_seg_info_label.config(text=f"Segment #{segment.id} - {segment.style}")
        
        # Update source
        self.doc_source_text.config(state='normal')
        self.doc_source_text.delete('1.0', 'end')
        self.doc_source_text.insert('1.0', segment.source)
        self.doc_source_text.config(state='disabled')
        
        # Update target
        self.doc_target_text.delete('1.0', 'end')
        if segment.target:
            self.doc_target_text.insert('1.0', segment.target)
        
        # Update status
        self.doc_status_var.set(segment.status)
        
        # Focus target text
        self.doc_target_text.focus_set()
        
        self.log(f"Editing Segment #{segment.id}")
    
    def on_doc_target_change(self, event=None):
        """Handle changes to target text in document view"""
        # No need to do anything here - just allow typing
        pass
    
    def on_doc_status_change(self, event=None):
        """Handle status change in document view"""
        if not self.doc_current_segment:
            return
        
        new_status = self.doc_status_var.get()
        self.doc_current_segment.status = new_status
        self.doc_current_segment.modified = True
        self.modified = True
        self.update_progress()
        
        # Update the segment's background color in the document
        widget_info = self.doc_segment_widgets.get(self.doc_current_segment.id)
        if widget_info:
            tag_name = widget_info['tag_name']
            para_text = widget_info['text_widget']
            
            if new_status == 'untranslated':
                bg_color = '#ffe6e6'
            elif new_status == 'draft':
                bg_color = '#fff9e6'
            elif new_status == 'translated':
                bg_color = '#e6ffe6'
            elif new_status == 'approved':
                bg_color = '#e6f3ff'
            else:
                bg_color = 'white'
            
            para_text.tag_config(tag_name, background=bg_color)
        
        self.log(f"✓ Status changed to {new_status}")
    
    def save_doc_segment(self):
        """Save current segment in document view"""
        if not self.doc_current_segment:
            return
        
        # Get new target text
        new_target = self.doc_target_text.get('1.0', 'end-1c')
        
        # Update segment
        self.doc_current_segment.target = new_target
        self.doc_current_segment.modified = True
        
        # Update the text in the document view
        widget_info = self.doc_segment_widgets.get(self.doc_current_segment.id)
        if widget_info:
            tag_name = widget_info['tag_name']
            para_text = widget_info['text_widget']
            
            # Get the range of the tag
            ranges = para_text.tag_ranges(tag_name)
            if len(ranges) >= 2:
                start, end = ranges[0], ranges[1]
                
                # Determine what to display based on new target
                if new_target.strip():
                    # Target has content
                    display_text = new_target
                elif new_target == '' and self.doc_current_segment.modified:
                    # Target explicitly cleared by user (and segment was modified)
                    display_text = f"[Segment {self.doc_current_segment.id} - Empty - Click to edit]"
                elif self.doc_current_segment.source:
                    # Fallback to source
                    display_text = self.doc_current_segment.source
                else:
                    # Both empty
                    display_text = f"[Segment {self.doc_current_segment.id} - Empty segment]"
                
                # Replace the text
                para_text.config(state='normal')
                para_text.delete(start, end)
                para_text.insert(start, display_text, tag_name)
                para_text.config(state='disabled')
                
                # Update background color based on status
                if self.doc_current_segment.status == 'draft':
                    bg_color = '#fff9e6'
                elif self.doc_current_segment.status == 'translated':
                    bg_color = '#e6ffe6'
                elif self.doc_current_segment.status == 'approved':
                    bg_color = '#e6f3ff'
                else:
                    bg_color = '#ffe6e6'
                
                para_text.tag_config(tag_name, background=bg_color)
        
        self.modified = True
        self.update_progress()
        self.log(f"✓ Segment #{self.doc_current_segment.id} saved")
    
    def save_doc_segment_and_next(self):
        """Save current segment with 'translated' status and move to next untranslated segment"""
        if not self.doc_current_segment:
            return 'break'
        
        # Set status to translated before saving
        self.doc_status_var.set('translated')
        self.save_doc_segment()
        
        # Find next untranslated segment
        current_id = self.doc_current_segment.id
        next_untranslated = None
        
        for seg in self.segments:
            if seg.id > current_id and seg.status == 'untranslated':
                if seg.id in self.doc_segment_widgets:
                    next_untranslated = seg
                    break
        
        # If found, navigate to it
        if next_untranslated:
            widget_info = self.doc_segment_widgets[next_untranslated.id]
            self.on_doc_segment_click(next_untranslated, widget_info['text_widget'], widget_info['tag_name'])
        else:
            self.log("No more untranslated segments")
        
        return 'break'
    
    def copy_source_to_target_doc(self):
        """Copy source to target in document view"""
        if not self.doc_current_segment:
            return 'break'
        
        # Copy source to target text widget
        self.doc_target_text.delete('1.0', 'end')
        self.doc_target_text.insert('1.0', self.doc_current_segment.source)
        
        self.log(f"✓ Copied source to target (Segment #{self.doc_current_segment.id})")
        return 'break'
    
    def clear_doc_target(self):
        """Clear target text in document view"""
        self.doc_target_text.delete('1.0', 'end')
    
    def navigate_document_segment(self, direction='next'):
        """Navigate to next/previous segment in document view"""
        if not self.doc_current_segment:
            # No current segment, select first
            if self.segments and self.segments[0].id in self.doc_segment_widgets:
                widget_info = self.doc_segment_widgets[self.segments[0].id]
                self.on_doc_segment_click(self.segments[0], widget_info['text_widget'], widget_info['tag_name'])
            return
        
        # Save current segment first
        self.save_doc_segment()
        
        current_id = self.doc_current_segment.id
        
        if direction == 'next':
            # Find next segment
            for seg in self.segments:
                if seg.id > current_id:
                    if seg.id in self.doc_segment_widgets:
                        widget_info = self.doc_segment_widgets[seg.id]
                        self.on_doc_segment_click(seg, widget_info['text_widget'], widget_info['tag_name'])
                    break
        else:
            # Find previous segment
            for seg in reversed(self.segments):
                if seg.id < current_id:
                    if seg.id in self.doc_segment_widgets:
                        widget_info = self.doc_segment_widgets[seg.id]
                        self.on_doc_segment_click(seg, widget_info['text_widget'], widget_info['tag_name'])
                    break
    
    def log(self, message: str):
        """Add message to log (both main window and workspace tab)"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        formatted_message = f"[{timestamp}] {message}\n"
        
        # Update main log window
        self.log_text.config(state='normal')
        self.log_text.insert('end', formatted_message)
        self.log_text.see('end')
        self.log_text.config(state='disabled')
        
        # Also update log tab if it exists (Translation Workspace)
        if hasattr(self, 'log_tab_text'):
            try:
                self.log_tab_text.config(state='normal')
                self.log_tab_text.insert('end', formatted_message)
                self.log_tab_text.see('end')
                self.log_tab_text.config(state='disabled')
            except tk.TclError:
                pass  # Widget was destroyed during layout switch
    
    def clear_log(self):
        """Clear both main log window and workspace log tab"""
        # Clear main log window
        self.log_text.config(state='normal')
        self.log_text.delete('1.0', 'end')
        self.log_text.config(state='disabled')
        
        # Clear log tab if it exists
        if hasattr(self, 'log_tab_text'):
            try:
                self.log_tab_text.config(state='normal')
                self.log_tab_text.delete('1.0', 'end')
                self.log_tab_text.config(state='disabled')
            except tk.TclError:
                pass  # Widget was destroyed during layout switch
        
        self.log("Log cleared")
    
    def switch_layout(self, new_mode: str):
        """Switch between layout modes"""
        if new_mode == self.layout_mode:
            return  # Already in this mode
        
        self.log(f"Switching to {new_mode} layout...")
        
        # Save current segment if editing
        if hasattr(self, 'current_edit_widget') and self.current_edit_widget:
            try:
                # Check if widget still exists before trying to save
                if self.current_edit_widget.winfo_exists():
                    self.save_grid_edit(go_next=False)
            except:
                pass  # Widget already destroyed
            finally:
                self.current_edit_widget = None
        elif self.layout_mode == LayoutMode.DOCUMENT and hasattr(self, 'doc_current_segment') and self.doc_current_segment:
            try:
                self.save_doc_segment()
            except:
                pass  # Widget already destroyed
        
        # Remember current selection from any view mode
        current_seg_id = None
        if self.layout_mode == LayoutMode.GRID and self.current_segment:
            current_seg_id = self.current_segment.id
        elif self.layout_mode == LayoutMode.DOCUMENT and hasattr(self, 'doc_current_segment') and self.doc_current_segment:
            current_seg_id = self.doc_current_segment.id
        elif self.layout_mode == LayoutMode.SPLIT and hasattr(self, 'current_segment') and self.current_segment:
            current_seg_id = self.current_segment.id
        
        # Update layout mode
        self.layout_mode = new_mode
        
        # Update button states
        self.update_layout_buttons()
        
        # Clear content frame
        for widget in self.content_frame.winfo_children():
            widget.destroy()
        
        # Rebuild UI based on layout mode
        if new_mode == LayoutMode.GRID:
            self.create_grid_layout()
            self.load_segments_to_grid()
            # Restore selection
            if current_seg_id:
                for i, row_data in enumerate(self.grid_rows):
                    if row_data['segment'].id == current_seg_id:
                        self.select_grid_row(i)
                        break
        elif new_mode == LayoutMode.SPLIT:
            self.create_split_layout()
            self.load_segments_to_tree()
            # Restore selection
            if current_seg_id:
                for item in self.tree.get_children():
                    if int(self.tree.item(item)['values'][0]) == current_seg_id:
                        self.tree.selection_set(item)
                        self.tree.see(item)
                        # Trigger selection event to load segment in editor
                        self.on_segment_select(None)
                        break
        elif new_mode == LayoutMode.DOCUMENT:
            self.create_document_layout()
            self.load_segments_to_document()
            # Restore selection by clicking first segment
            if current_seg_id and current_seg_id in self.doc_segment_widgets:
                widget_info = self.doc_segment_widgets[current_seg_id]
                self.on_doc_segment_click(widget_info['segment'], 
                                        widget_info['text_widget'], 
                                        widget_info['tag_name'])
        
        mode_names = {
            LayoutMode.GRID: "Grid View (memoQ-style)",
            LayoutMode.SPLIT: "List View",
            LayoutMode.DOCUMENT: "Document View (Flow)"
        }
        self.log(f"✓ Switched to {mode_names.get(new_mode, new_mode)}")
    
    def update_layout_buttons(self):
        """Update layout button visual states"""
        # Reset all buttons to neutral gray
        self.layout_btn_grid.config(relief='raised', bg='#f5f5f5', fg='black')
        self.layout_btn_split.config(relief='raised', bg='#f5f5f5', fg='black')
        self.layout_btn_document.config(relief='raised', bg='#f5f5f5', fg='black')
        
        # Highlight active button with darker gray and sunken effect
        if self.layout_mode == LayoutMode.GRID:
            self.layout_btn_grid.config(relief='sunken', bg='#e0e0e0', fg='black')
        elif self.layout_mode == LayoutMode.SPLIT:
            self.layout_btn_split.config(relief='sunken', bg='#e0e0e0', fg='black')
        elif self.layout_mode == LayoutMode.DOCUMENT:
            self.layout_btn_document.config(relief='sunken', bg='#e0e0e0', fg='black')
    
    # Custom Grid View methods
    
    def create_grid_header(self):
        """Create sticky header row for custom grid with resizable columns"""
        header_frame = tk.Frame(self.header_container, bg='#e0e0e0', relief='raised', bd=1)
        header_frame.pack(fill='x', side='top', pady=0)
        
        # Store header labels for updating
        self.header_labels = {}
        
        # Fixed columns with resize handles - ID, Type, Style (status moved to end)
        for col_name in ['id', 'type', 'style']:
            col_info = self.grid_columns[col_name]
            if not col_info.get('visible', True):
                continue  # Skip hidden columns
            
            # Container for label + resize handle
            col_container = tk.Frame(header_frame, bg='#e0e0e0')
            col_container.pack(side='left', fill='y')
            
            header_label = tk.Label(col_container, 
                                   text=col_info['title'],
                                   font=('Segoe UI', 9, 'bold'),
                                   bg='#e0e0e0',
                                   fg='black',
                                   width=col_info['width'] // 8,
                                   anchor=col_info['anchor'],
                                   relief='raised',
                                   bd=1)
            header_label.pack(side='left', fill='both', expand=True, padx=1, pady=1)
            self.header_labels[col_name] = header_label
            
            # Resize handle (thin draggable border)
            resize_handle = tk.Frame(col_container, bg='#999', width=2, cursor='sb_h_double_arrow')
            resize_handle.pack(side='left', fill='y')
            resize_handle.bind('<Button-1>', lambda e, c=col_name: self.start_column_resize(e, c))
            resize_handle.bind('<B1-Motion>', lambda e, c=col_name: self.on_column_resize(e, c))
            resize_handle.bind('<ButtonRelease-1>', self.end_column_resize)
        
        # Status column header (packed on right side first, before content)
        if self.grid_columns['status'].get('visible', True):
            col_info = self.grid_columns['status']
            status_header = tk.Label(header_frame, 
                                    text=col_info['title'],  # '\u2713' icon
                                    font=('Segoe UI', 9, 'bold'),
                                    bg='#e0e0e0',
                                    fg='black',
                                    width=col_info['width'] // 8,
                                    anchor=col_info['anchor'],
                                    relief='raised',
                                    bd=1)
            status_header.pack(side='right', fill='y', padx=1, pady=1)
            self.header_labels['status'] = status_header
        
        # Create container frame to match row structure
        header_content = tk.Frame(header_frame, bg='#e0e0e0')
        header_content.pack(side='left', fill='both', expand=True)
        
        # Source header - fixed width frame to match row source_frame
        source_header_frame = tk.Frame(header_content, bg='#e0e0e0', width=self.source_width)
        source_header_frame.pack(side='left', fill='both', expand=False)
        source_header_frame.pack_propagate(False)
        
        self.source_header = tk.Label(source_header_frame, 
                               text='📄 Source',
                               font=('Segoe UI', 9, 'bold'),
                               bg='#e0e0e0',
                               fg='black',
                               anchor='w',
                               relief='raised',
                               bd=1)
        self.source_header.pack(fill='both', expand=True, padx=1, pady=1)
        
        # Draggable splitter - exactly 4px to match row splitter_space
        self.splitter = tk.Frame(header_content, bg='#666', width=4, cursor='sb_h_double_arrow', relief='raised')
        self.splitter.pack(side='left', fill='y', pady=0)
        
        # Bind splitter dragging events
        self.splitter.bind('<Button-1>', self.start_splitter_drag)
        self.splitter.bind('<B1-Motion>', self.on_splitter_drag)
        self.splitter.bind('<ButtonRelease-1>', self.end_splitter_drag)
        
        # Target header - expanding frame to match row target_frame
        target_header_frame = tk.Frame(header_content, bg='#e0e0e0')
        target_header_frame.pack(side='left', fill='both', expand=True)
        
        self.target_header = tk.Label(target_header_frame, 
                               text='🎯 Target',
                               font=('Segoe UI', 9, 'bold'),
                               bg='#e0e0e0',
                               fg='black',
                               anchor='w',
                               relief='raised',
                               bd=1)
        self.target_header.pack(fill='both', expand=True, padx=1, pady=1)
        
        # Store header frames for width updates
        self.source_header_frame = source_header_frame
        self.target_header_frame = target_header_frame
    
    def update_header_widths(self):
        """Update header frame widths based on current source/target widths"""
        self.source_header_frame.config(width=self.source_width)
    
    def start_splitter_drag(self, event):
        """Start dragging the column splitter"""
        self.dragging_splitter = True
        self.drag_start_x = event.x_root
        self.drag_start_source_width = self.source_width
        self.drag_start_target_width = self.target_width
    
    def on_splitter_drag(self, event):
        """Handle splitter dragging"""
        if not self.dragging_splitter:
            return
        
        # Calculate delta
        delta = event.x_root - self.drag_start_x
        
        # Update widths (keep minimum of 200px each)
        new_source_width = max(200, self.drag_start_source_width + delta)
        new_target_width = max(200, self.drag_start_target_width - delta)
        
        self.source_width = new_source_width
        self.target_width = new_target_width
        
        # Update column config
        self.grid_columns['source']['width'] = self.source_width
        self.grid_columns['target']['width'] = self.target_width
        
        # Update headers
        self.update_header_widths()
        
        # Update all rows
        self.update_all_row_widths()
    
    def end_splitter_drag(self, event):
        """End dragging the column splitter"""
        self.dragging_splitter = False
    
    def start_column_resize(self, event, col_name):
        """Start resizing a fixed column"""
        self.resizing_column = col_name
        self.resize_start_x = event.x_root
        self.resize_start_width = self.grid_columns[col_name]['width']
    
    def on_column_resize(self, event, col_name):
        """Handle fixed column resizing"""
        if not hasattr(self, 'resizing_column') or self.resizing_column != col_name:
            return
        
        # Calculate delta
        delta = event.x_root - self.resize_start_x
        
        # Update width (keep minimum of 25px)
        new_width = max(25, self.resize_start_width + delta)
        self.grid_columns[col_name]['width'] = new_width
        
        # Update header label
        if col_name in self.header_labels:
            self.header_labels[col_name].config(width=new_width // 8)
        
        # Reload grid to apply new widths
        # Note: For better performance, we could update existing rows instead
        # but full reload ensures consistency
        if hasattr(self, 'grid_rows') and len(self.grid_rows) > 0:
            # Store current scroll position and selection
            current_page = self.grid_current_page if hasattr(self, 'grid_current_page') else 0
            # Reload will restore to same page
            pass  # Width changes are applied when cells are created
    
    def end_column_resize(self, event):
        """End resizing fixed column and reload grid"""
        if hasattr(self, 'resizing_column'):
            # Reload grid to apply new column widths
            current_selection = self.current_row_index if hasattr(self, 'current_row_index') else 0
            self.load_segments_to_grid()
            # Try to restore selection
            if current_selection >= 0 and current_selection < len(self.grid_rows):
                self.select_grid_row(current_selection)
            del self.resizing_column
    
    def update_all_row_widths(self):
        """Update all row column widths after splitter drag"""
        for row_data in self.grid_rows:
            widgets = row_data['widgets']
            
            # Update source frame width (fixed)
            if 'source_frame' in widgets:
                widgets['source_frame'].config(width=self.source_width)
            
            # Target frame expands automatically, no need to set width
        
        # Update canvas
        self.grid_canvas.update_idletasks()
        self.update_grid_scroll_region()
    
    def show_column_visibility_dialog(self):
        """Show dialog to toggle column visibility"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Grid View - Show/Hide Columns")
        dialog.geometry("300x250")
        dialog.transient(self.root)
        dialog.grab_set()
        
        tk.Label(dialog, text="Select columns to display:", 
                font=('Segoe UI', 10, 'bold')).pack(pady=10)
        
        # Create checkboxes for each column (except source/target which are always visible)
        # Order matches the actual column order: Type, Style, Status
        checkbox_vars = {}
        for col_name in ['id', 'type', 'style', 'status']:
            var = tk.BooleanVar(value=self.grid_columns[col_name].get('visible', True))
            checkbox_vars[col_name] = var
            
            cb = tk.Checkbutton(dialog, 
                              text=self.grid_columns[col_name]['title'],
                              variable=var,
                              font=('Segoe UI', 10))
            cb.pack(anchor='w', padx=20, pady=5)
        
        # Note about source/target
        tk.Label(dialog, text="Note: Source and Target are always visible",
                fg='gray', font=('Segoe UI', 8)).pack(pady=5)
        
        # Buttons
        button_frame = tk.Frame(dialog)
        button_frame.pack(pady=10)
        
        def apply_changes():
            # Update column visibility
            for col_name, var in checkbox_vars.items():
                self.grid_columns[col_name]['visible'] = var.get()
            
            # Rebuild grid
            self.rebuild_grid()
            dialog.destroy()
            self.log("Column visibility updated")
        
        def cancel():
            dialog.destroy()
        
        tk.Button(button_frame, text="Apply", command=apply_changes, width=10).pack(side='left', padx=5)
        tk.Button(button_frame, text="Cancel", command=cancel, width=10).pack(side='left', padx=5)
    
    def rebuild_grid(self):
        """Rebuild the entire grid (e.g., after column visibility changes)"""
        if self.layout_mode != LayoutMode.GRID:
            return
        
        # Store current selection
        current_selection = self.current_row_index
        
        # Clear existing grid
        for widget in self.grid_inner_frame.winfo_children():
            widget.destroy()
        self.grid_rows.clear()
        
        # Clear header
        for widget in self.header_container.winfo_children():
            widget.destroy()
        
        # Recreate header
        self.create_grid_header()
        
        # Recreate all rows
        segments = self.segments if not self.filter_active else self.filtered_segments
        for segment in segments:
            self.add_grid_row(segment)
        
        # Restore selection
        if 0 <= current_selection < len(self.grid_rows):
            self.select_grid_row(current_selection)
        
        # Update scroll region
        self.grid_canvas.update_idletasks()
        self.update_grid_scroll_region()
    
    def toggle_grid_style_colors(self):
        """Toggle style-based font colors in Grid View"""
        self.grid_style_colors_enabled = not self.grid_style_colors_enabled
        status = "enabled" if self.grid_style_colors_enabled else "disabled"
        self.log(f"Style colors {status}")
        
        # Rebuild grid to apply changes
        if self.layout_mode == LayoutMode.GRID:
            self.rebuild_grid()
    
    def apply_formatting_to_text_widget(self, text_widget, text_content, segment=None):
        """Apply visual formatting (bold, italic, underline) and style colors to text in a Text widget"""
        import re
        
        # Configure text tags for formatting
        text_widget.tag_configure('bold', font=('Segoe UI', 9, 'bold'))
        text_widget.tag_configure('italic', font=('Segoe UI', 9, 'italic'))
        text_widget.tag_configure('underline', font=('Segoe UI', 9, 'underline'))
        
        # Apply style-based colors if enabled (matching List View)
        if self.grid_style_colors_enabled and segment:
            if segment.is_table_cell:
                text_widget.tag_configure('style_color', foreground='#0066cc', font=('Segoe UI', 9, 'italic'))
                text_widget.tag_add('style_color', '1.0', 'end')
            elif segment.style:
                style_lower = segment.style.lower()
                if 'heading 1' in style_lower:
                    text_widget.tag_configure('style_color', foreground='#003366', font=('Segoe UI', 10, 'bold'))
                    text_widget.tag_add('style_color', '1.0', 'end')
                elif 'heading 2' in style_lower:
                    text_widget.tag_configure('style_color', foreground='#0066cc', font=('Segoe UI', 9, 'bold'))
                    text_widget.tag_add('style_color', '1.0', 'end')
                elif 'heading 3' in style_lower:
                    text_widget.tag_configure('style_color', foreground='#3399ff', font=('Segoe UI', 9, 'bold'))
                    text_widget.tag_add('style_color', '1.0', 'end')
                elif 'title' in style_lower and 'subtitle' not in style_lower:
                    text_widget.tag_configure('style_color', foreground='#663399', font=('Segoe UI', 11, 'bold'))
                    text_widget.tag_add('style_color', '1.0', 'end')
                elif 'subtitle' in style_lower:
                    text_widget.tag_configure('style_color', foreground='#663399', font=('Segoe UI', 9, 'italic'))
                    text_widget.tag_add('style_color', '1.0', 'end')
        
        # Find and apply bold tags <b>...</b>
        for match in re.finditer(r'<b>(.*?)</b>', text_content, re.DOTALL):
            start_pos = f"1.0 + {match.start()} chars"
            end_pos = f"1.0 + {match.end()} chars"
            # Get the actual position of the content (without tags)
            content_start = f"1.0 + {match.start(1)} chars"
            content_end = f"1.0 + {match.end(1)} chars"
            text_widget.tag_add('bold', content_start, content_end)
        
        # Find and apply italic tags <i>...</i>
        for match in re.finditer(r'<i>(.*?)</i>', text_content, re.DOTALL):
            content_start = f"1.0 + {match.start(1)} chars"
            content_end = f"1.0 + {match.end(1)} chars"
            text_widget.tag_add('italic', content_start, content_end)
        
        # Find and apply underline tags <u>...</u>
        for match in re.finditer(r'<u>(.*?)</u>', text_content, re.DOTALL):
            content_start = f"1.0 + {match.start(1)} chars"
            content_end = f"1.0 + {match.end(1)} chars"
            text_widget.tag_add('underline', content_start, content_end)
    
    def get_status_icon(self, status, is_locked=False):
        """Get memoQ-style status icon"""
        if is_locked:
            return '🔒'  # Blue padlock for locked
        elif status == 'approved':
            return '✓✓'  # Double checkmark for approved
        elif status == 'translated':
            return '✓'   # Green checkmark for translated
        elif status == 'draft':
            return '~'   # Tilde for draft
        elif status == 'untranslated':
            return '✗'   # Red X for untranslated
        else:
            return '?'   # Unknown status
    
    def get_status_icon_color(self, status, is_locked=False):
        """Get color for status icon (foreground color)"""
        if is_locked:
            return '#2196F3'  # Blue for locked
        elif status == 'approved':
            return '#1976D2'  # Dark blue for approved
        elif status == 'translated':
            return '#4CAF50'  # Green for translated
        elif status == 'draft':
            return '#FF9800'  # Orange for draft
        elif status == 'untranslated':
            return '#F44336'  # Red for untranslated
        else:
            return '#666'     # Gray for unknown
    
    def calculate_row_height(self, segment):
        """Calculate appropriate row height based on content with word wrapping"""
        import textwrap
        
        # Calculate character width based on actual column widths
        # Approximate: 8px per character
        source_width_chars = max(30, self.source_width // 8)
        target_width_chars = max(30, self.target_width // 8)
        
        # Calculate wrapped lines for source
        source_lines = 0
        if segment.source:
            # Split by existing newlines first
            source_paragraphs = segment.source.split('\n')
            for para in source_paragraphs:
                if para:
                    # Calculate how many lines this paragraph will wrap to
                    wrapped = textwrap.fill(para, width=source_width_chars)
                    source_lines += wrapped.count('\n') + 1
                else:
                    source_lines += 1
        
        # Calculate wrapped lines for target
        target_lines = 0
        if segment.target:
            target_paragraphs = segment.target.split('\n')
            for para in target_paragraphs:
                if para:
                    wrapped = textwrap.fill(para, width=target_width_chars)
                    target_lines += wrapped.count('\n') + 1
                else:
                    target_lines += 1
        else:
            target_lines = 1
        
        # Use the maximum of source and target
        max_lines = max(source_lines, target_lines)
        
        # Calculate height: 18px per line + 15px padding
        # Minimum 35px, maximum 250px
        height = max(35, min(250, max_lines * 18 + 15))
        
        return height
    
    def update_grid_scroll_region(self, event=None):
        """Update the scroll region of the canvas"""
        self.grid_canvas.configure(scrollregion=self.grid_canvas.bbox('all'))
    
    def on_grid_canvas_resize(self, event):
        """Handle canvas resize to adjust inner frame width"""
        canvas_width = event.width
        self.grid_canvas.itemconfig(self.canvas_window, width=canvas_width)
    
    def on_grid_mousewheel(self, event):
        """Handle mouse wheel scrolling"""
        self.grid_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
    
    def on_grid_click(self, event):
        """Handle click on grid to select row (with Ctrl/Shift multi-selection support)"""
        # Find which row was clicked
        y_pos = self.grid_canvas.canvasy(event.y)
        
        # Check for Ctrl and Shift modifiers (Trados/memoQ style)
        ctrl_pressed = (event.state & 0x4) != 0  # Ctrl key
        shift_pressed = (event.state & 0x1) != 0  # Shift key
        
        for i, row_data in enumerate(self.grid_rows):
            row_frame = row_data['row_frame']
            if row_frame.winfo_y() <= y_pos <= row_frame.winfo_y() + row_frame.winfo_height():
                self.select_grid_row(i, ctrl_pressed, shift_pressed)
                self.grid_canvas.focus_set()
                break
    
    def on_grid_double_click(self, event):
        """Handle double-click to enter edit mode or show source"""
        # Find which row and column was clicked
        y_pos = self.grid_canvas.canvasy(event.y)
        x_pos = self.grid_canvas.canvasx(event.x)
        
        # Double-click always selects just this row (ignores Ctrl/Shift)
        for i, row_data in enumerate(self.grid_rows):
            row_frame = row_data['row_frame']
            if row_frame.winfo_y() <= y_pos <= row_frame.winfo_y() + row_frame.winfo_height():
                self.select_grid_row(i, ctrl_pressed=False, shift_pressed=False)
                
                # Determine which column was clicked
                col_x = 0
                for col_name, col_info in self.grid_columns.items():
                    if col_x <= x_pos <= col_x + col_info['width']:
                        if col_name == 'target':
                            self.enter_edit_mode()
                        elif col_name == 'source':
                            self.show_source_popup(event)
                        break
                    col_x += col_info['width']
                break
    
    def show_grid_context_menu(self, event):
        """Show context menu for grid"""
        # Select row first
        self.on_grid_click(event)
        self.context_menu.post(event.x_root, event.y_root)
    
    def on_text_widget_click(self, event, row_index):
        """Handle click on source text widget - select row with multi-selection support"""
        # Check for Ctrl and Shift modifiers
        ctrl_pressed = (event.state & 0x4) != 0
        shift_pressed = (event.state & 0x1) != 0
        
        self.select_grid_row(row_index, ctrl_pressed, shift_pressed)
        return 'break'  # Prevent default Text widget behavior
    
    def on_target_click(self, event, row_index):
        """Handle single click on target - select row and enter edit mode"""
        # Select the row first
        self.select_grid_row(row_index)
        
        # Enter edit mode immediately
        self.enter_edit_mode()
        
        return 'break'  # Prevent default Text widget behavior
    
    def on_target_double_click(self, event, row_index):
        """Handle double-click on target - enter edit mode directly"""
        # Select the row first
        self.select_grid_row(row_index)
        
        # Enter edit mode for this row
        self.enter_edit_mode()
        
        return 'break'
    
    def on_source_right_click(self, event, row_index):
        """Handle right-click on source - select row and show source popup"""
        self.select_grid_row(row_index)
        self.show_source_popup(event)
        return 'break'
    
    def on_target_right_click(self, event, row_index):
        """Handle right-click on target - select row and show context menu"""
        self.select_grid_row(row_index)
        self.context_menu.post(event.x_root, event.y_root)
        return 'break'
    
    def select_grid_row(self, row_index, ctrl_pressed=False, shift_pressed=False):
        """Select a row in the custom grid with multi-selection support (Trados/memoQ style)"""
        if row_index < 0 or row_index >= len(self.grid_rows):
            return
        
        # Safety check: ensure we're in grid mode
        if self.layout_mode != LayoutMode.GRID:
            return
        
        # Safety check: ensure grid_canvas exists
        if not hasattr(self, 'grid_canvas') or not self.grid_canvas.winfo_exists():
            return
        
        row_data = self.grid_rows[row_index]
        segment = row_data['segment']
        
        # Clear dual selection if changing rows
        if self.dual_selection_row is not None and self.dual_selection_row != row_index:
            self.clear_dual_selection()
        
        # Multi-selection handling (Trados/memoQ style)
        if shift_pressed and self.last_selected_index is not None:
            # Range selection: select all between last and current
            start = min(self.last_selected_index, row_index)
            end = max(self.last_selected_index, row_index)
            for i in range(start, end + 1):
                if i < len(self.grid_rows):
                    seg = self.grid_rows[i]['segment']
                    self.selected_segments.add(seg.id)
                    self._apply_selection_highlight(i, True)
        elif ctrl_pressed:
            # Toggle selection on this row
            if segment.id in self.selected_segments:
                self.selected_segments.remove(segment.id)
                self._apply_selection_highlight(row_index, False)
            else:
                self.selected_segments.add(segment.id)
                self._apply_selection_highlight(row_index, True)
        else:
            # Normal click - clear other selections and select only this row
            self._clear_all_selection_highlights()
            self.selected_segments.clear()
            self.selected_segments.add(segment.id)
            self._apply_selection_highlight(row_index, True)
        
        # Update last selected index for range selection
        self.last_selected_index = row_index
        
        # Update current row/segment (for editing)
        old_index = self.current_row_index
        self.current_row_index = row_index
        self.current_segment = segment
        
        # Update border on current row (different from selection highlight)
        if old_index >= 0 and old_index < len(self.grid_rows):
            try:
                old_row = self.grid_rows[old_index]['row_frame']
                if old_row.winfo_exists() and old_row['relief'] == 'solid':
                    # Keep selection color but remove focus border if not selected
                    if self.grid_rows[old_index]['segment'].id not in self.selected_segments:
                        old_row.config(relief='flat', bd=1)
            except (tk.TclError, KeyError):
                pass
        
        # Add focus border to new current row
        row_frame = row_data['row_frame']
        try:
            if segment.id in self.selected_segments:
                row_frame.config(relief='solid', bd=2)  # Thicker border for focus
        except tk.TclError:
            return  # Widget destroyed
        
        # Update selection counter
        self.update_selection_counter()
        
        # Ensure row is visible - scroll to it smoothly
        try:
            self.grid_canvas.update_idletasks()
            
            # Get canvas viewport height and row position
            canvas_height = self.grid_canvas.winfo_height()
            row_y = row_frame.winfo_y()
            row_height = row_frame.winfo_height()
            
            # Get current scroll position
            scroll_region = self.grid_canvas.cget('scrollregion').split()
            if len(scroll_region) == 4:
                total_height = float(scroll_region[3])
                
                # Check if row is visible
                current_view = self.grid_canvas.yview()
                view_top = current_view[0] * total_height
                view_bottom = current_view[1] * total_height
                
                # Only scroll if row is not fully visible
                if row_y < view_top:
                    # Row is above viewport - scroll to show it at top
                    self.grid_canvas.yview_moveto(row_y / total_height)
                elif row_y + row_height > view_bottom:
                    # Row is below viewport - scroll to show it at bottom
                    self.grid_canvas.yview_moveto((row_y + row_height - canvas_height) / total_height)
        except (tk.TclError, ValueError, ZeroDivisionError):
            pass  # Geometry not ready or widget destroyed
        
        # Load segment into editor panel
        self.load_segment_to_grid_editor(self.current_segment)
        
        # Trigger automatic TM search after delay
        self.schedule_auto_tm_search()
    
    def _apply_selection_highlight(self, row_index, selected):
        """Apply or remove blue selection highlight to a row (Trados/memoQ style)"""
        if row_index < 0 or row_index >= len(self.grid_rows):
            return
        
        row_data = self.grid_rows[row_index]
        row_frame = row_data['row_frame']
        
        # Blue highlight color like Trados/memoQ
        selection_color = '#CCE5FF' if selected else None
        
        try:
            if selected:
                # Store original background for restoration
                if 'original_bg' not in row_data:
                    row_data['original_bg'] = row_frame.cget('bg')
                row_frame.config(bg=selection_color)
                # Update all child widgets
                for widget in row_frame.winfo_children():
                    self._update_widget_bg(widget, selection_color)
            else:
                # Restore original background
                original_bg = row_data.get('original_bg', 'white')
                row_frame.config(bg=original_bg)
                for widget in row_frame.winfo_children():
                    self._update_widget_bg(widget, original_bg)
        except tk.TclError:
            pass
    
    def _update_widget_bg(self, widget, bg_color):
        """Recursively update background color of widget and children"""
        try:
            widget_type = widget.winfo_class()
            if widget_type in ('Label', 'Frame', 'Text'):
                widget.config(bg=bg_color)
            # Recurse for containers
            for child in widget.winfo_children():
                self._update_widget_bg(child, bg_color)
        except tk.TclError:
            pass
    
    def _clear_all_selection_highlights(self):
        """Remove selection highlight from all rows"""
        for i in range(len(self.grid_rows)):
            self._apply_selection_highlight(i, False)
    
    def update_selection_counter(self):
        """Update the selection counter in status bar and log"""
        count = len(self.selected_segments)
        
        # Update the status bar label (if in Grid layout)
        if hasattr(self, 'selection_counter_label') and self.selection_counter_label.winfo_exists():
            if count > 0:
                self.selection_counter_label.config(text=f"📌 {count} segment{'s' if count != 1 else ''} selected")
            else:
                self.selection_counter_label.config(text="")
        
        # Also log for visibility
        if count > 0:
            self.log(f"📌 {count} segment{'s' if count != 1 else ''} selected")
    
    def should_highlight_segment(self, segment):
        """
        Check if a segment matches the current filter criteria.
        Used in Highlight Mode to determine which segments to highlight.
        In Filter Mode, all visible segments already match, so highlighting is always applied.
        In Highlight Mode, all segments are visible, so we only highlight those that match.
        """
        # In filter mode, all visible segments match (already filtered)
        if self.filter_mode == 'filter':
            return True
        
        # In highlight mode, check if this segment matches the filters
        source_filter = self.filter_source_var.get().strip().lower()
        target_filter = self.filter_target_var.get().strip().lower()
        status_filter = self.filter_status_var.get()
        
        # Check source filter
        if source_filter and source_filter not in segment.source.lower():
            return False
        
        # Check target filter
        if target_filter and target_filter not in segment.target.lower():
            return False
        
        # Check status filter
        if status_filter != "All" and segment.status != status_filter:
            return False
        
        return True
    
    def highlight_search_terms_in_segment(self, text_widget, start_index, end_index, segment, tag_name):
        """
        Highlight only the specific search terms within a segment in Document View.
        This provides more precise visual feedback than highlighting the entire segment.
        """
        source_filter = self.filter_source_var.get().strip().lower()
        target_filter = self.filter_target_var.get().strip().lower()
        
        # Get the text content from the widget
        segment_text = text_widget.get(start_index, end_index)
        
        # Determine which text to search in based on what's displayed
        # If target has content, we're showing target; otherwise showing source
        if segment.target and segment.target.strip():
            display_text = segment.target
            search_term = target_filter
        else:
            display_text = segment.source
            search_term = source_filter
        
        # If we have a search term, highlight all occurrences
        if search_term:
            # Search for all occurrences (case-insensitive)
            search_lower = display_text.lower()
            search_term_lower = search_term.lower()
            
            start_pos = 0
            occurrence_count = 0
            
            while True:
                # Find next occurrence
                pos = search_lower.find(search_term_lower, start_pos)
                if pos == -1:
                    break
                
                # Calculate the actual position in the Text widget
                # We need to account for the start_index offset
                match_start = f"{start_index} + {pos} chars"
                match_end = f"{start_index} + {pos + len(search_term)} chars"
                
                # Create a unique tag for this highlight
                highlight_tag = f"search_highlight_{segment.id}_{occurrence_count}"
                
                # Add the tag to this occurrence
                text_widget.tag_add(highlight_tag, match_start, match_end)
                
                # Configure the tag with bright yellow background and bold
                text_widget.tag_config(highlight_tag, 
                                     background='#FFFF00',  # Bright yellow
                                     foreground='#000000',  # Black text
                                     font=('Segoe UI', 11, 'bold'))
                
                # Raise the tag so it appears above the segment background
                text_widget.tag_raise(highlight_tag)
                
                occurrence_count += 1
                start_pos = pos + 1  # Move past this occurrence to find next one
    
    def add_grid_row(self, segment):
        """Add a row to the custom grid with dynamic height"""
        # Calculate row height based on content
        row_height = self.calculate_row_height(segment)
        
        # Determine type label
        if segment.is_table_cell and segment.table_info:
            type_text = f"T{segment.table_info[0]+1}R{segment.table_info[1]+1}C{segment.table_info[2]+1}"
        else:
            type_text = "Para"
        
        # Determine background color based on status
        status_colors = {
            'untranslated': '#ffe6e6',
            'draft': '#fff9e6',
            'translated': '#e6ffe6',
            'approved': '#e6f3ff'
        }
        bg_color = status_colors.get(segment.status, 'white')
        
        # Get status icon (memoQ style)
        status_icon = self.get_status_icon(segment.status, segment.locked if hasattr(segment, 'locked') else False)
        
        # Create row frame
        row_frame = tk.Frame(self.grid_inner_frame, bg=bg_color, relief='flat', bd=1, height=row_height)
        row_frame.pack(fill='x', side='top', pady=0)
        row_frame.pack_propagate(False)  # Prevent frame from shrinking
        
        # Create widgets for each column
        widgets = {}
        
        # ID column
        if self.grid_columns['id'].get('visible', True):
            id_label = tk.Label(row_frame, text=str(segment.id), 
                               bg=bg_color, font=('Segoe UI', 9),
                               width=self.grid_columns['id']['width'] // 8,
                               anchor='center', cursor='hand2')  # Hand cursor to indicate clickable
            id_label.pack(side='left', padx=1)
            
            # Make ID label clickable for row selection (Trados/memoQ style)
            row_index = len(self.grid_rows)
            id_label.bind('<Button-1>', lambda e, idx=row_index: self.on_text_widget_click(e, idx))
            id_label.bind('<MouseWheel>', self.on_grid_mousewheel)
            
            widgets['id'] = id_label
        
        # Type column
        if self.grid_columns['type'].get('visible', True):
            type_label = tk.Label(row_frame, text=type_text, 
                                 bg=bg_color, font=('Segoe UI', 9),
                                 width=self.grid_columns['type']['width'] // 8,
                                 anchor='center')
            type_label.pack(side='left', padx=1)
            widgets['type'] = type_label
        
        # Style column
        if self.grid_columns['style'].get('visible', True):
            style_text = segment.style if segment.style else "Normal"
            style_label = tk.Label(row_frame, text=style_text, 
                                  bg=bg_color, font=('Segoe UI', 9),
                                  width=self.grid_columns['style']['width'] // 8,
                                  anchor='center')
            style_label.pack(side='left', padx=1)
            widgets['style'] = style_label
        
        # Status column (packed first on right side, before content takes remaining space)
        if self.grid_columns['status'].get('visible', True):
            status_color = self.get_status_icon_color(segment.status, segment.locked if hasattr(segment, 'locked') else False)
            # Create fixed-width container for status icon
            status_container = tk.Frame(row_frame, bg=bg_color, width=30)
            status_container.pack(side='right', fill='y')
            status_container.pack_propagate(False)
            
            status_label = tk.Label(status_container, text=status_icon, 
                                   bg=bg_color, font=('Segoe UI', 14, 'bold'),
                                   fg=status_color,
                                   anchor='center')
            status_label.pack(fill='both', expand=True)
            status_label.bind('<MouseWheel>', self.on_grid_mousewheel)
            widgets['status'] = status_label
            widgets['status_container'] = status_container
        
        # Create a container frame for source + splitter + target to match header layout
        content_container = tk.Frame(row_frame, bg=bg_color)
        content_container.pack(side='left', fill='both', expand=True)
        
        # Source column (Text widget for multi-line, read-only) - fixed width
        source_frame = tk.Frame(content_container, bg=bg_color, width=self.source_width)
        source_frame.pack(side='left', fill='both', expand=False)
        source_frame.pack_propagate(False)
        
        source_text = tk.Text(source_frame, wrap='word', font=('Segoe UI', 9),
                             bg=bg_color, relief='solid', bd=1,
                             state='normal',
                             highlightthickness=0,
                             cursor='arrow',
                             padx=2, pady=2)
        source_text.insert('1.0', segment.source)
        
        # Apply formatting tags to source text (with optional style colors)
        self.apply_formatting_to_text_widget(source_text, segment.source, segment)
        
        # Highlight filter matches in source (only if segment matches filters in highlight mode)
        if self.filter_active and hasattr(self, 'filter_source_var'):
            source_filter = self.filter_source_var.get().strip()
            if source_filter and self.should_highlight_segment(segment):
                self.highlight_text_in_widget(source_text, source_filter, 'yellow')
        
        source_text.config(state='disabled')
        source_text.pack(fill='both', expand=True, padx=1, pady=1)
        
        # Make source clickable to select row and support dual selection
        row_index = len(self.grid_rows)  # Current index before appending
        source_text.bind('<Button-1>', lambda e, idx=row_index: self.on_source_text_click(e, idx))
        source_text.bind('<ButtonRelease-1>', lambda e, idx=row_index: self.on_source_selection_made(e, idx))
        source_text.bind('<Button-3>', lambda e, idx=row_index: self.on_source_right_click(e, idx))
        source_text.bind('<MouseWheel>', self.on_grid_mousewheel)
        
        # Keyboard selection bindings (Ctrl+Shift+Arrow keys - memoQ style)
        source_text.bind('<Control-Shift-Left>', 
                        lambda e, w=source_text: self.extend_selection_keyboard(w, 'left', 'char', 'source'))
        source_text.bind('<Control-Shift-Right>', 
                        lambda e, w=source_text: self.extend_selection_keyboard(w, 'right', 'char', 'source'))
        source_text.bind('<Control-Shift-Control-Left>', 
                        lambda e, w=source_text: self.extend_selection_keyboard(w, 'left', 'word', 'source'))
        source_text.bind('<Control-Shift-Control-Right>', 
                        lambda e, w=source_text: self.extend_selection_keyboard(w, 'right', 'word', 'source'))
        source_text.bind('<Tab>', lambda e: self.switch_dual_selection_focus())
        source_text.bind('<Escape>', lambda e: (self.clear_dual_selection(), None)[1] or 'break')
        
        widgets['source'] = source_text
        widgets['source_frame'] = source_frame
        
        # Splitter placeholder (matches header splitter width)
        splitter_space = tk.Frame(content_container, bg=bg_color, width=4)
        splitter_space.pack(side='left', fill='y')
        
        # Target column (Text widget for multi-line, will be editable) - expands
        target_frame = tk.Frame(content_container, bg=bg_color)
        target_frame.pack(side='left', fill='both', expand=True)
        
        target_text = tk.Text(target_frame, wrap='word', font=('Segoe UI', 9),
                             bg=bg_color, relief='solid', bd=1,
                             state='normal',
                             highlightthickness=0,
                             cursor='xterm',
                             padx=2, pady=2)
        if segment.target:
            target_text.insert('1.0', segment.target)
            # Apply formatting tags to target text (with optional style colors)
            self.apply_formatting_to_text_widget(target_text, segment.target, segment)
        
        # Highlight filter matches in target (only if segment matches filters in highlight mode)
        if self.filter_active and hasattr(self, 'filter_target_var'):
            target_filter = self.filter_target_var.get().strip()
            if target_filter and self.should_highlight_segment(segment):
                self.highlight_text_in_widget(target_text, target_filter, 'lightgreen')
        
        target_text.config(state='disabled')  # Initially disabled, enabled in edit mode
        target_text.pack(fill='both', expand=True, padx=1, pady=1)
        
        # Make target clickable - support dual selection and edit mode
        target_text.bind('<Button-1>', lambda e, idx=row_index: self.on_target_text_click(e, idx))
        target_text.bind('<ButtonRelease-1>', lambda e, idx=row_index: self.on_target_selection_made(e, idx))
        target_text.bind('<Double-Button-1>', lambda e, idx=row_index: self.on_target_double_click(e, idx))
        target_text.bind('<Button-3>', lambda e, idx=row_index: self.on_target_right_click(e, idx))
        target_text.bind('<MouseWheel>', self.on_grid_mousewheel)
        
        # Keyboard selection bindings (Ctrl+Shift+Arrow keys - memoQ style)
        target_text.bind('<Control-Shift-Left>', 
                        lambda e, w=target_text: self.extend_selection_keyboard(w, 'left', 'char', 'target'))
        target_text.bind('<Control-Shift-Right>', 
                        lambda e, w=target_text: self.extend_selection_keyboard(w, 'right', 'char', 'target'))
        target_text.bind('<Control-Shift-Control-Left>', 
                        lambda e, w=target_text: self.extend_selection_keyboard(w, 'left', 'word', 'target'))
        target_text.bind('<Control-Shift-Control-Right>', 
                        lambda e, w=target_text: self.extend_selection_keyboard(w, 'right', 'word', 'target'))
        target_text.bind('<Tab>', lambda e: self.switch_dual_selection_focus())
        target_text.bind('<Escape>', lambda e: (self.clear_dual_selection(), None)[1] or 'break')
        
        widgets['target'] = target_text
        widgets['target_frame'] = target_frame
        
        # Also bind mousewheel to the frames and labels
        content_container.bind('<MouseWheel>', self.on_grid_mousewheel)
        row_frame.bind('<MouseWheel>', self.on_grid_mousewheel)
        type_label.bind('<MouseWheel>', self.on_grid_mousewheel)
        status_label.bind('<MouseWheel>', self.on_grid_mousewheel)
        
        # Store row data
        row_data = {
            'segment': segment,
            'widgets': widgets,
            'row_frame': row_frame
        }
        self.grid_rows.append(row_data)
        
        # Update scroll region
        self.grid_canvas.update_idletasks()
        self.update_grid_scroll_region()
    
    # Grid View inline editing methods
    
    def create_context_menu(self):
        """Create context menu for Grid View"""
        self.context_menu = tk.Menu(self.root, tearoff=0)
        self.context_menu.add_command(label="📋 Copy Source → Target (Ctrl+D)", 
                                     command=self.copy_source_to_target)
        self.context_menu.add_command(label="📄 View Source Text", 
                                     command=lambda: self.show_source_popup_from_menu())
        self.context_menu.add_separator()
        self.context_menu.add_command(label="Insert <b>Bold</b> Tag (Ctrl+B)", 
                                     command=lambda: self.insert_tag_inline('b'))
        self.context_menu.add_command(label="Insert <i>Italic</i> Tag (Ctrl+I)", 
                                     command=lambda: self.insert_tag_inline('i'))
        self.context_menu.add_command(label="Insert <u>Underline</u> Tag (Ctrl+U)", 
                                     command=lambda: self.insert_tag_inline('u'))
        self.context_menu.add_separator()
        self.context_menu.add_command(label="🗑️ Clear Target", 
                                     command=self.clear_target_inline)
        self.context_menu.add_separator()
        self.context_menu.add_command(label="✅ Mark as Translated", 
                                     command=lambda: self.set_status_inline('translated'))
        self.context_menu.add_command(label="⭐ Mark as Approved", 
                                     command=lambda: self.set_status_inline('approved'))
        self.context_menu.add_command(label="📝 Mark as Draft", 
                                     command=lambda: self.set_status_inline('draft'))
        self.context_menu.add_separator()
        self.context_menu.add_command(label="⬇️ Next Segment (Ctrl+Down)", 
                                     command=lambda: self.navigate_segment('next'))
        self.context_menu.add_command(label="⬆️ Previous Segment (Ctrl+Up)", 
                                     command=lambda: self.navigate_segment('prev'))
    
    def show_context_menu(self, event):
        """Show context menu on right-click"""
        # Select the row under cursor
        item = self.tree.identify_row(event.y)
        if item:
            self.tree.selection_set(item)
            self.context_menu.post(event.x_root, event.y_root)
    
    def show_source_popup_from_menu(self):
        """Show source popup from context menu"""
        # Create a fake event at center of screen
        class FakeEvent:
            pass
        event = FakeEvent()
        event.x_root = self.root.winfo_x() + self.root.winfo_width() // 2
        event.y_root = self.root.winfo_y() + self.root.winfo_height() // 2
        self.show_source_popup(event)
    
    def on_segment_select_grid(self, event):
        """Handle segment selection in Grid View (minimal action)"""
        selection = self.tree.selection()
        if not selection:
            return
        
        item = selection[0]
        seg_id = int(self.tree.item(item)['values'][0])
        
        # Find and set current segment
        for seg in self.segments:
            if seg.id == seg_id:
                self.current_segment = seg
                break
    
    def on_cell_double_click(self, event):
        """Handle double-click on cell - check if it's the target column"""
        # Identify what was clicked
        region = self.tree.identify_region(event.x, event.y)
        if region != "cell":
            return
        
        column = self.tree.identify_column(event.x)
        # Column #4 is source, #5 is target (0-indexed: #0 is tree column, #1-5 are our columns)
        if column == '#4':  # Source column
            self.show_source_popup(event)
        elif column == '#5':  # Target column
            self.enter_edit_mode()
    
    def enter_edit_mode_global(self):
        """Enter edit mode from anywhere - works in Grid or Split mode"""
        if not self.current_segment:
            return
        
        if self.layout_mode == LayoutMode.GRID:
            # In Grid mode, enter inline editing
            self.enter_edit_mode()
        else:
            # In Split mode, focus the target text widget
            if hasattr(self, 'target_text'):
                self.target_text.focus()
    
    def enter_edit_mode(self, event=None):
        """Enter inline edit mode for target cell"""
        # Clear dual selection when entering edit mode
        self.clear_dual_selection()
        
        if self.layout_mode == LayoutMode.GRID:
            # Custom grid edit mode
            if self.current_row_index < 0 or self.current_row_index >= len(self.grid_rows):
                return
            
            row_data = self.grid_rows[self.current_row_index]
            target_widget = row_data['widgets']['target']
            
            # Enable the target Text widget
            target_widget.config(state='normal', bg='#ffffcc', relief='solid', bd=2)
            target_widget.focus_set()
            
            # Bind save/cancel keys
            target_widget.bind('<Control-Return>', lambda e: self.save_grid_edit(go_next=True))
            target_widget.bind('<Escape>', lambda e: self.cancel_grid_edit())
            
            # Bind tag shortcuts
            target_widget.bind('<Control-b>', lambda e: self.insert_tag_grid('b'))
            target_widget.bind('<Control-i>', lambda e: self.insert_tag_grid('i'))
            target_widget.bind('<Control-u>', lambda e: self.insert_tag_grid('u'))
            
            # Real-time tag validation
            target_widget.bind('<KeyRelease>', self.validate_tags_grid)
            
            # Store current edit widget
            self.current_edit_widget = target_widget
            
            self.log("Edit mode: Ctrl+Enter to save & next, Escape to cancel")
            
        else:
            # Treeview edit mode (Split/Compact)
            selection = self.tree.selection()
            if not selection:
                return
            
            item = selection[0]
            
            # Get cell bounds for target column
            bbox = self.tree.bbox(item, column='#5')  # Target column
            if not bbox:
                return
            
            # Create Entry widget overlay
            self.current_edit_widget = tk.Entry(
                self.tree,
                font=('Segoe UI', 10),
                relief='solid',
                borderwidth=2,
                bg='#ffffcc'  # Yellow background to indicate edit mode
            )
            
            # Position over cell
            x, y, width, height = bbox
            self.current_edit_widget.place(
                x=x, y=y, width=width, height=height
            )
            
            # Get current target text
            segment = self.current_segment
            if segment:
                self.current_edit_widget.insert(0, segment.target)
                self.current_edit_widget.select_range(0, tk.END)
                self.current_edit_widget.focus()
            
            # Bind keys
            self.current_edit_widget.bind('<Return>', 
                lambda e: self.save_inline_edit(go_next=False))
            self.current_edit_widget.bind('<Control-Return>', 
                lambda e: self.save_inline_edit(go_next=True))
            self.current_edit_widget.bind('<Escape>', 
                lambda e: self.cancel_inline_edit())
            self.current_edit_widget.bind('<FocusOut>', 
                lambda e: self.save_inline_edit(go_next=False))
            
            # Real-time tag validation
            self.current_edit_widget.bind('<KeyRelease>', 
                self.validate_tags_inline)
            
            self.log("Edit mode: Press Enter to save, Ctrl+Enter for save & next, Escape to cancel")
    
    def save_inline_edit(self, go_next=False):
        """Save inline edit and optionally move to next"""
        if not self.current_edit_widget:
            return
        
        # Get edited text
        new_text = self.current_edit_widget.get()
        
        # Validate tags
        if self.tag_manager and new_text:
            is_valid, error = self.tag_manager.validate_tags(new_text)
            if not is_valid:
                messagebox.showerror("Tag Error", error)
                self.current_edit_widget.focus()
                return
        
        # Save to segment
        if self.current_segment:
            self.current_segment.target = new_text
            if new_text and self.current_segment.status == 'untranslated':
                self.current_segment.status = 'translated'
            self.current_segment.modified = True
            self.current_segment.modified_at = datetime.now().isoformat()
            
            # Update grid
            self.update_segment_in_grid(self.current_segment)
            
            self.modified = True
            self.update_progress()
            self.log(f"Segment #{self.current_segment.id} saved")
        
        # Destroy edit widget
        self.current_edit_widget.destroy()
        self.current_edit_widget = None
        
        # Clear validation label
        self.tag_validation_label.config(text="")
        
        # Move to next if requested
        if go_next:
            self.next_segment()
    
    def cancel_inline_edit(self):
        """Cancel inline editing without saving"""
        if self.current_edit_widget:
            self.current_edit_widget.destroy()
            self.current_edit_widget = None
            self.tag_validation_label.config(text="")
            self.log("Edit cancelled")
    
    def save_grid_edit(self, go_next=False):
        """Save grid edit for custom grid (Text widget)"""
        if not self.current_edit_widget or self.current_row_index < 0:
            return
        
        # Get edited text
        new_text = self.current_edit_widget.get('1.0', 'end-1c')
        
        # Validate tags
        if self.tag_manager and new_text:
            is_valid, error = self.tag_manager.validate_tags(new_text)
            if not is_valid:
                messagebox.showerror("Tag Error", error)
                self.current_edit_widget.focus()
                return
        
        # Save to segment
        row_data = self.grid_rows[self.current_row_index]
        segment = row_data['segment']
        segment.target = new_text
        
        # Set status based on go_next flag
        if go_next:
            # Ctrl+Enter was pressed - always set to translated
            segment.status = 'translated'
        else:
            # Regular save - only auto-set if untranslated
            if new_text and segment.status == 'untranslated':
                segment.status = 'translated'
        
        segment.modified = True
        segment.modified_at = datetime.now().isoformat()
        
        # Add to Project TM if translated or approved
        if segment.status in ['translated', 'approved'] and new_text:
            self.tm_database.add_to_project_tm(segment.source, new_text)
            self.log(f"✓ Added to Project TM: {segment.source[:50]}...")
        
        # Update the display
        self.update_grid_row(self.current_row_index)
        
        # Disable editing - keep border
        target_widget = row_data['widgets']['target']
        target_widget.config(state='disabled', bg=self.get_status_color(segment.status), relief='solid', bd=1)
        
        self.current_edit_widget = None
        self.modified = True
        self.update_progress()
        self.log(f"Segment #{segment.id} saved")
        
        # Move to next if requested
        if go_next:
            # Find next untranslated segment
            current_id = segment.id
            next_untranslated_index = None
            
            for i in range(self.current_row_index + 1, len(self.grid_rows)):
                if self.grid_rows[i]['segment'].status == 'untranslated':
                    next_untranslated_index = i
                    break
            
            if next_untranslated_index is not None:
                # Navigate to next untranslated segment
                self.select_grid_row(next_untranslated_index)
                # Automatically enter edit mode on the next segment
                self.root.after(50, self.enter_edit_mode)  # Small delay to ensure row is selected
            else:
                self.log("No more untranslated segments")
                self.grid_canvas.focus_set()
        else:
            # Return focus to canvas
            self.grid_canvas.focus_set()
    
    def cancel_grid_edit(self):
        """Cancel grid editing without saving"""
        if self.current_row_index >= 0 and self.current_row_index < len(self.grid_rows):
            row_data = self.grid_rows[self.current_row_index]
            target_widget = row_data['widgets']['target']
            segment = row_data['segment']
            
            # Restore original text
            target_widget.delete('1.0', 'end')
            if segment.target:
                target_widget.insert('1.0', segment.target)
            
            # Disable editing - keep border
            target_widget.config(state='disabled', bg=self.get_status_color(segment.status), relief='solid', bd=1)
        
        self.current_edit_widget = None
        self.tag_validation_label.config(text="")
        self.grid_canvas.focus_set()
        self.log("Edit cancelled")
    
    def get_status_color(self, status):
        """Get background color for status"""
        status_colors = {
            'untranslated': '#ffe6e6',
            'draft': '#fff9e6',
            'translated': '#e6ffe6',
            'approved': '#e6f3ff'
        }
        return status_colors.get(status, 'white')
    
    def update_grid_row(self, row_index):
        """Update a grid row after editing"""
        if row_index < 0 or row_index >= len(self.grid_rows):
            return
        
        row_data = self.grid_rows[row_index]
        segment = row_data['segment']
        widgets = row_data['widgets']
        row_frame = row_data['row_frame']
        
        # Update background color
        bg_color = self.get_status_color(segment.status)
        row_frame.config(bg=bg_color)
        for widget_name, widget in widgets.items():
            if widget_name == 'status_container':
                widget.config(bg=bg_color)
            elif widget_name != 'status':
                widget.config(bg=bg_color)
            # Keep borders on source and target
            if widget_name in ['source', 'target']:
                widget.config(relief='solid', bd=1)
        
        # Update status icon and color
        status_icon = self.get_status_icon(segment.status, segment.locked if hasattr(segment, 'locked') else False)
        status_color = self.get_status_icon_color(segment.status, segment.locked if hasattr(segment, 'locked') else False)
        widgets['status'].config(text=status_icon, fg=status_color, bg=bg_color)
        
        # Update target text
        target_widget = widgets['target']
        target_widget.config(state='normal')
        target_widget.delete('1.0', 'end')
        if segment.target:
            target_widget.insert('1.0', segment.target)
        target_widget.config(state='disabled')
        
        # Recalculate and update row height
        new_height = self.calculate_row_height(segment)
        row_frame.config(height=new_height)
        
        # Update scroll region
        self.grid_canvas.update_idletasks()
        self.update_grid_scroll_region()
    
    def validate_tags_grid(self, event=None):
        """Validate tags and resize row dynamically during editing"""
        if not self.current_edit_widget or not self.tag_manager:
            return
        
        text = self.current_edit_widget.get('1.0', 'end-1c')
        
        # Validate tags
        if not text:
            self.grid_tag_validation_label.config(text="", fg='#666')
        else:
            is_valid, error = self.tag_manager.validate_tags(text)
            if is_valid:
                self.grid_tag_validation_label.config(text="✓ Tags valid", fg='green')
            else:
                self.grid_tag_validation_label.config(text=f"✗ {error}", fg='red')
        
        # Dynamic row resizing while typing
        if self.current_row_index >= 0 and self.current_row_index < len(self.grid_rows):
            self.resize_row_during_edit()
    
    def resize_row_during_edit(self):
        """Dynamically resize the current row based on content while editing"""
        if self.current_row_index < 0 or self.current_row_index >= len(self.grid_rows):
            return
        
        import textwrap
        
        row_data = self.grid_rows[self.current_row_index]
        segment = row_data['segment']
        row_frame = row_data['row_frame']
        target_widget = row_data['widgets']['target']
        source_widget = row_data['widgets']['source']
        
        # Get current text from target widget (being edited)
        current_target_text = target_widget.get('1.0', 'end-1c')
        
        # Use dynamic column widths
        source_width_chars = max(30, self.source_width // 8)
        target_width_chars = max(30, self.target_width // 8)
        
        # Calculate wrapped lines for source (unchanged)
        source_lines = 0
        if segment.source:
            source_paragraphs = segment.source.split('\n')
            for para in source_paragraphs:
                if para:
                    wrapped = textwrap.fill(para, width=source_width_chars)
                    source_lines += wrapped.count('\n') + 1
                else:
                    source_lines += 1
        
        # Calculate wrapped lines for current target text
        target_lines = 0
        if current_target_text:
            target_paragraphs = current_target_text.split('\n')
            for para in target_paragraphs:
                if para:
                    wrapped = textwrap.fill(para, width=target_width_chars)
                    target_lines += wrapped.count('\n') + 1
                else:
                    target_lines += 1
        else:
            target_lines = 1
        
        # Use the maximum of source and target
        max_lines = max(source_lines, target_lines)
        
        # Calculate new height
        new_height = max(35, min(250, max_lines * 18 + 15))
        
        # Only update if height changed significantly (avoid constant tiny updates)
        current_height = row_frame.winfo_height()
        if abs(new_height - current_height) > 5:
            row_frame.config(height=new_height)
            
            # Update scroll region
            self.grid_canvas.update_idletasks()
            self.update_grid_scroll_region()
    
    def insert_tag_grid(self, tag_type):
        """Insert tag at cursor position in grid editor"""
        if self.current_edit_widget:
            opening = f"<{tag_type}>"
            closing = f"</{tag_type}>"
            self.current_edit_widget.insert(tk.INSERT, opening + closing)
            # Move cursor between tags
            current_pos = self.current_edit_widget.index(tk.INSERT)
            line, col = current_pos.split('.')
            new_pos = f"{line}.{int(col) - len(closing)}"
            self.current_edit_widget.mark_set(tk.INSERT, new_pos)
            self.current_edit_widget.focus()
            return 'break'  # Prevent default behavior
    
    def validate_tags_inline(self, event=None):
        """Validate tags in real-time during inline editing"""
        if not self.current_edit_widget or not self.tag_manager:
            return
        
        text = self.current_edit_widget.get()
        if not text:
            self.grid_tag_validation_label.config(text="", fg='#666')
            return
        
        is_valid, error = self.tag_manager.validate_tags(text)
        if is_valid:
            self.grid_tag_validation_label.config(text="✓ Tags valid", fg='green')
        else:
            self.grid_tag_validation_label.config(text=f"✗ {error}", fg='red')
    
    # Dual text selection methods (Grid View)
    
    def on_source_text_click(self, event, row_index):
        """Handle click in source text - potential start of dual selection (with multi-selection support)"""
        # Check for Ctrl and Shift modifiers
        ctrl_pressed = (event.state & 0x4) != 0
        shift_pressed = (event.state & 0x1) != 0
        
        # First, handle row selection with multi-selection support
        self.select_grid_row(row_index, ctrl_pressed, shift_pressed)
        
        # Clear any existing dual selection from other rows
        if self.dual_selection_row is not None and self.dual_selection_row != row_index:
            self.clear_dual_selection()
        
        self.dual_selection_row = row_index
        row_data = self.grid_rows[row_index]
        self.dual_selection_source = row_data['widgets']['source']
    
    def on_source_selection_made(self, event, row_index):
        """Handle selection made in source text"""
        if row_index != self.dual_selection_row:
            return
        
        source_widget = self.grid_rows[row_index]['widgets']['source']
        
        # Check if there's a selection
        try:
            selection_start = source_widget.index(tk.SEL_FIRST)
            selection_end = source_widget.index(tk.SEL_LAST)
            
            # Store the selection and highlight it
            if selection_start != selection_end:
                # Remove existing source selection tag
                source_widget.tag_remove('dual_sel_source', '1.0', tk.END)
                
                # Add colored tag for source selection
                source_widget.tag_add('dual_sel_source', selection_start, selection_end)
                source_widget.tag_config('dual_sel_source',
                                        background='#B3E5FC',  # Light blue
                                        foreground='#01579B')  # Dark blue
                
                # Raise the tag above other tags
                source_widget.tag_raise('dual_sel_source')
                
                selected_text = source_widget.get(selection_start, selection_end)
                self.log(f"Source selection: '{selected_text}' ({len(selected_text)} chars)")
        except tk.TclError:
            # No selection made
            pass
    
    def on_target_text_click(self, event, row_index):
        """Handle click in target text - potential start of dual selection (with multi-selection support)"""
        # Check for Ctrl and Shift modifiers
        ctrl_pressed = (event.state & 0x4) != 0
        shift_pressed = (event.state & 0x1) != 0
        
        # First, handle row selection with multi-selection support
        self.select_grid_row(row_index, ctrl_pressed, shift_pressed)
        
        # Check if this is the same row
        if self.dual_selection_row is not None and self.dual_selection_row != row_index:
            self.clear_dual_selection()
        
        self.dual_selection_row = row_index
        row_data = self.grid_rows[row_index]
        self.dual_selection_target = row_data['widgets']['target']
        
        # Don't enter edit mode on single click - only on double-click
        # This allows for selection without editing
    
    def on_target_selection_made(self, event, row_index):
        """Handle selection made in target text"""
        if row_index != self.dual_selection_row:
            return
        
        target_widget = self.grid_rows[row_index]['widgets']['target']
        
        # Check if there's a selection
        try:
            selection_start = target_widget.index(tk.SEL_FIRST)
            selection_end = target_widget.index(tk.SEL_LAST)
            
            # Store the selection and highlight it
            if selection_start != selection_end:
                # Remove existing target selection tag
                target_widget.tag_remove('dual_sel_target', '1.0', tk.END)
                
                # Add colored tag for target selection
                target_widget.tag_add('dual_sel_target', selection_start, selection_end)
                target_widget.tag_config('dual_sel_target',
                                        background='#C8E6C9',  # Light green
                                        foreground='#1B5E20')  # Dark green
                
                # Raise the tag above other tags
                target_widget.tag_raise('dual_sel_target')
                
                selected_text = target_widget.get(selection_start, selection_end)
                self.log(f"Target selection: '{selected_text}' ({len(selected_text)} chars)")
        except tk.TclError:
            # No selection made
            pass
    
    def clear_dual_selection(self):
        """Clear dual text selection highlights"""
        if self.dual_selection_row is not None and self.dual_selection_row < len(self.grid_rows):
            row_data = self.grid_rows[self.dual_selection_row]
            
            # Clear source selection tag
            if 'source' in row_data['widgets']:
                source_widget = row_data['widgets']['source']
                source_widget.tag_remove('dual_sel_source', '1.0', tk.END)
            
            # Clear target selection tag
            if 'target' in row_data['widgets']:
                target_widget = row_data['widgets']['target']
                target_widget.tag_remove('dual_sel_target', '1.0', tk.END)
        
        self.dual_selection_row = None
        self.dual_selection_source = None
        self.dual_selection_target = None
        self.dual_selection_focused_widget = None
        # Reset cursor memory when clearing selection
        self.dual_selection_source_cursor = None
        self.dual_selection_target_cursor = None
    
    def switch_dual_selection_focus(self):
        """Switch focus between source and target for keyboard selection (Tab key)"""
        # Don't interfere with edit mode - Tab should save & next in edit mode
        if self.current_edit_widget:
            return  # Allow default Tab behavior (save & next)
        
        if self.current_row_index < 0 or self.current_row_index >= len(self.grid_rows):
            return 'break'
        
        row_data = self.grid_rows[self.current_row_index]
        
        if self.dual_selection_focused_widget == 'source':
            # Switch to target
            self.focus_target_for_selection(self.current_row_index)
        elif self.dual_selection_focused_widget == 'target':
            # Switch to source
            self.focus_source_for_selection(self.current_row_index)
        else:
            # Start with source
            self.focus_source_for_selection(self.current_row_index)
        
        return 'break'  # Prevent default Tab behavior
    
    def focus_source_for_selection(self, row_index):
        """Focus source widget for keyboard-based dual selection"""
        if row_index < 0 or row_index >= len(self.grid_rows):
            return
        
        row_data = self.grid_rows[row_index]
        source_widget = row_data['widgets']['source']
        
        # Enable source widget temporarily for keyboard input (usually disabled/readonly)
        source_widget.config(state='normal')
        source_widget.focus_set()
        
        # Position cursor - use remembered position or start at 1.0
        if self.dual_selection_source_cursor is not None:
            # Return to where we left off
            cursor_pos = self.dual_selection_source_cursor
        else:
            # First time - start at beginning
            cursor_pos = '1.0'
        
        source_widget.mark_set(tk.INSERT, cursor_pos)
        source_widget.see(cursor_pos)
        
        # Visual indicator - subtle border
        source_widget.config(relief='solid', bd=2, highlightthickness=1,
                           highlightbackground='#2196F3', highlightcolor='#2196F3')
        
        # Save target's cursor position before switching away
        if 'target' in row_data['widgets']:
            target_widget = row_data['widgets']['target']
            try:
                self.dual_selection_target_cursor = target_widget.index(tk.INSERT)
            except:
                pass
            target_widget.config(state='disabled', relief='solid', bd=1, highlightthickness=0)
        
        self.dual_selection_focused_widget = 'source'
        self.dual_selection_row = row_index
        self.log("Source focused - use Ctrl+Shift+Arrows to select, Tab to switch")
    
    def focus_target_for_selection(self, row_index):
        """Focus target widget for keyboard-based dual selection"""
        if row_index < 0 or row_index >= len(self.grid_rows):
            return
        
        row_data = self.grid_rows[row_index]
        target_widget = row_data['widgets']['target']
        
        # Enable target widget temporarily for keyboard input (may be disabled)
        target_widget.config(state='normal')
        target_widget.focus_set()
        
        # Position cursor - use remembered position or start at 1.0
        if self.dual_selection_target_cursor is not None:
            # Return to where we left off
            cursor_pos = self.dual_selection_target_cursor
        else:
            # First time - start at beginning
            cursor_pos = '1.0'
        
        target_widget.mark_set(tk.INSERT, cursor_pos)
        target_widget.see(cursor_pos)
        
        # Visual indicator - subtle border
        target_widget.config(relief='solid', bd=2, highlightthickness=1,
                           highlightbackground='#4CAF50', highlightcolor='#4CAF50')
        
        # Save source's cursor position before switching away
        if 'source' in row_data['widgets']:
            source_widget = row_data['widgets']['source']
            try:
                self.dual_selection_source_cursor = source_widget.index(tk.INSERT)
            except:
                pass
            source_widget.config(state='disabled', relief='solid', bd=1, highlightthickness=0)
        
        self.dual_selection_focused_widget = 'target'
        self.dual_selection_row = row_index
        self.log("Target focused - use Ctrl+Shift+Arrows to select, Tab to switch")
    
    def extend_selection_keyboard(self, widget, direction, unit, widget_type):
        """Extend selection using keyboard (Ctrl+Shift+Arrow keys)
        
        Args:
            widget: The Text widget (source or target)
            direction: 'left' or 'right'
            unit: 'char' or 'word'
            widget_type: 'source' or 'target'
        """
        try:
            # Get current insert position (cursor)
            current_pos = widget.index(tk.INSERT)
            
            # Check if we already have a selection
            try:
                sel_start = widget.index(tk.SEL_FIRST)
                sel_end = widget.index(tk.SEL_LAST)
                has_selection = True
                # Determine if cursor is at start or end to know direction of extension
                cursor_at_start = widget.compare(current_pos, "==", sel_start)
            except tk.TclError:
                # No existing selection, start from cursor
                sel_start = current_pos
                sel_end = current_pos
                has_selection = False
                cursor_at_start = False
            
            # Calculate new position based on direction and unit
            if direction == 'right':
                if unit == 'char':
                    new_pos = widget.index(f"{current_pos} + 1 char")
                else:  # word
                    new_pos = widget.index(f"{current_pos} wordend")
                    # If not at word boundary, move to next word end
                    if widget.compare(new_pos, "==", current_pos):
                        new_pos = widget.index(f"{new_pos} + 1 char wordend")
                
                # Extend selection
                if has_selection and cursor_at_start:
                    # Cursor at start, moving right shrinks selection
                    new_start = new_pos
                    new_end = sel_end
                else:
                    # Cursor at end or no selection, extend right
                    new_start = sel_start if has_selection else current_pos
                    new_end = new_pos
                
                cursor_pos = new_end
                
            else:  # left
                if unit == 'char':
                    new_pos = widget.index(f"{current_pos} - 1 char")
                else:  # word
                    new_pos = widget.index(f"{current_pos} wordstart")
                    # If not at word boundary, move to previous word start
                    if widget.compare(new_pos, "==", current_pos):
                        new_pos = widget.index(f"{new_pos} - 1 char wordstart")
                
                # Extend selection
                if has_selection and not cursor_at_start:
                    # Cursor at end, moving left shrinks selection
                    new_start = sel_start
                    new_end = new_pos
                else:
                    # Cursor at start or no selection, extend left
                    new_start = new_pos
                    new_end = sel_end if has_selection else current_pos
                
                cursor_pos = new_start
            
            # Don't go beyond text boundaries
            text_start = widget.index("1.0")
            text_end = widget.index("end-1c")
            
            # Clamp to boundaries
            if widget.compare(new_start, "<", text_start):
                new_start = text_start
            if widget.compare(new_end, ">", text_end):
                new_end = text_end
            
            # Ensure start is before end
            if widget.compare(new_start, ">", new_end):
                new_start, new_end = new_end, new_start
            
            # Apply the selection
            widget.tag_remove('sel', '1.0', tk.END)  # Clear default selection
            
            # Remove existing colored tag
            tag_name = f'dual_sel_{widget_type}'
            widget.tag_remove(tag_name, '1.0', tk.END)
            
            # Add new colored selection if there's a range
            if widget.compare(new_start, "!=", new_end):
                # Add Tkinter selection for cursor positioning
                widget.tag_add('sel', new_start, new_end)
                
                # Add colored tag
                widget.tag_add(tag_name, new_start, new_end)
                
                # Configure tag colors
                if widget_type == 'source':
                    widget.tag_config(tag_name,
                                    background='#B3E5FC',  # Light blue
                                    foreground='#01579B')  # Dark blue
                else:  # target
                    widget.tag_config(tag_name,
                                    background='#C8E6C9',  # Light green
                                    foreground='#1B5E20')  # Dark green
                
                widget.tag_raise(tag_name)
                
                # Update cursor position
                widget.mark_set(tk.INSERT, cursor_pos)
                widget.see(cursor_pos)  # Scroll to cursor
                
                # Show selection in status bar
                selected_text = widget.get(new_start, new_end)
                self.log(f"{widget_type.capitalize()} selection: '{selected_text}' ({len(selected_text)} chars)")
        
        except tk.TclError as e:
            # Handle edge cases silently
            pass
        
        return 'break'  # Prevent default behavior
    
    def insert_tag_inline(self, tag_type):
        """Insert tag at cursor position in inline editor"""
        if self.current_edit_widget:
            # Insert in the edit widget
            pos = self.current_edit_widget.index(tk.INSERT)
            opening = f"<{tag_type}>"
            closing = f"</{tag_type}>"
            self.current_edit_widget.insert(pos, opening + closing)
            # Move cursor between tags
            self.current_edit_widget.icursor(pos + len(opening))
            self.current_edit_widget.focus()
    
    def clear_target_inline(self):
        """Clear target of currently selected segment"""
        if self.current_segment:
            self.current_segment.target = ""
            self.current_segment.status = "untranslated"
            
            if self.layout_mode == LayoutMode.GRID:
                # Update custom grid
                if self.current_row_index >= 0:
                    self.update_grid_row(self.current_row_index)
            else:
                # Update treeview
                self.update_segment_in_grid(self.current_segment)
            
            self.modified = True
            self.update_progress()
            self.log(f"Segment #{self.current_segment.id} target cleared")
    
    def set_status_inline(self, status):
        """Set status of currently selected segment"""
        if self.current_segment:
            self.current_segment.status = status
            
            if self.layout_mode == LayoutMode.GRID:
                # Update custom grid
                if self.current_row_index >= 0:
                    self.update_grid_row(self.current_row_index)
            else:
                # Update treeview
                self.update_segment_in_grid(self.current_segment)
            
            self.modified = True
            self.log(f"Segment #{self.current_segment.id} marked as {status}")
    
    def next_segment(self):
        """Move to next segment"""
        selection = self.tree.selection()
        if not selection:
            return
        
        # Get next item
        item = selection[0]
        next_item = self.tree.next(item)
        
        if next_item:
            self.tree.selection_set(next_item)
            self.tree.see(next_item)
            self.tree.focus(next_item)
            
            # If in Grid mode, could auto-enter edit mode
            if self.layout_mode == LayoutMode.GRID:
                # Small delay to allow selection to update
                self.root.after(100, self.enter_edit_mode)
    
    def navigate_segment(self, direction='next'):
        """Navigate to next or previous segment"""
        if self.layout_mode == LayoutMode.GRID:
            # Custom grid navigation
            if not self.grid_rows:
                return
            
            if self.current_row_index < 0:
                # No selection, select first
                self.select_grid_row(0)
                segment = self.grid_rows[0]['segment']
                seg_type = self.get_segment_type_label(segment)
                self.log(f"Segment #{segment.id} ({seg_type})")
                return
            
            if direction == 'next':
                if self.current_row_index < len(self.grid_rows) - 1:
                    self.select_grid_row(self.current_row_index + 1)
                    segment = self.grid_rows[self.current_row_index]['segment']
                    seg_type = self.get_segment_type_label(segment)
                    self.log(f"Segment #{segment.id} ({seg_type})")
                else:
                    self.log("Already at last segment")
            elif direction == 'prev':
                if self.current_row_index > 0:
                    self.select_grid_row(self.current_row_index - 1)
                    segment = self.grid_rows[self.current_row_index]['segment']
                    seg_type = self.get_segment_type_label(segment)
                    self.log(f"Segment #{segment.id} ({seg_type})")
                else:
                    self.log("Already at first segment")
        elif self.layout_mode == LayoutMode.DOCUMENT:
            # Document view navigation
            self.navigate_document_segment(direction)
            if self.current_segment:
                seg_type = self.get_segment_type_label(self.current_segment)
                self.log(f"Segment #{self.current_segment.id} ({seg_type})")
        else:
            # Treeview navigation
            selection = self.tree.selection()
            if not selection:
                # No selection, select first item
                children = self.tree.get_children()
                if children:
                    self.tree.selection_set(children[0])
                    self.tree.see(children[0])
                    self.tree.focus(children[0])
                    # Get segment info
                    if self.current_segment:
                        seg_type = self.get_segment_type_label(self.current_segment)
                        self.log(f"Segment #{self.current_segment.id} ({seg_type})")
                return
            
            item = selection[0]
            
            if direction == 'next':
                next_item = self.tree.next(item)
                if next_item:
                    self.tree.selection_set(next_item)
                    self.tree.see(next_item)
                    self.tree.focus(next_item)
                    if self.current_segment:
                        seg_type = self.get_segment_type_label(self.current_segment)
                        self.log(f"Segment #{self.current_segment.id} ({seg_type})")
                else:
                    self.log("Already at last segment")
            elif direction == 'prev':
                prev_item = self.tree.prev(item)
                if prev_item:
                    self.tree.selection_set(prev_item)
                    self.tree.see(prev_item)
                    self.tree.focus(prev_item)
                    if self.current_segment:
                        seg_type = self.get_segment_type_label(self.current_segment)
                        self.log(f"Segment #{self.current_segment.id} ({seg_type})")
                else:
                    self.log("Already at first segment")
    
    def get_segment_type_label(self, segment):
        """Get a user-friendly label for the segment type"""
        if segment.is_table_cell:
            return "Table"
        elif segment.style and "Heading" in segment.style:
            return segment.style  # e.g., "Heading 1", "Heading 2"
        elif segment.style and segment.style != "Normal":
            return segment.style  # e.g., "Title", "Subtitle"
        else:
            return "Para"  # Regular paragraph
    
    def show_source_popup(self, event):
        """Show source and target in a memoQ-style popup with clear divider"""
        if not self.current_segment:
            return
        
        # Calculate content-aware size
        source_lines = self.current_segment.source.count('\n') + 1
        target_lines = self.current_segment.target.count('\n') + 1 if self.current_segment.target else 1
        
        # Estimate height needed (30px per line + padding)
        source_height = max(60, min(200, source_lines * 30 + 40))
        target_height = max(60, min(200, target_lines * 30 + 40))
        total_height = source_height + target_height + 120  # Extra for labels and buttons
        
        # Width based on content length
        max_line_length = max(
            max(len(line) for line in self.current_segment.source.split('\n') or ['']) if self.current_segment.source else 0,
            max(len(line) for line in self.current_segment.target.split('\n') or ['']) if self.current_segment.target else 0
        )
        popup_width = max(500, min(800, max_line_length * 8 + 100))
        
        # Create popup window
        popup = tk.Toplevel(self.root)
        popup.title(f"Segment #{self.current_segment.id} - Source & Target")
        popup.geometry(f"{popup_width}x{total_height}")
        popup.transient(self.root)
        
        # Position near cursor
        popup.geometry(f"+{event.x_root + 10}+{event.y_root + 10}")
        
        # Main content frame
        main_frame = tk.Frame(popup, padx=15, pady=10)
        main_frame.pack(fill='both', expand=True)
        
        # Source section
        source_label_frame = tk.Frame(main_frame)
        source_label_frame.pack(fill='x', pady=(0, 5))
        tk.Label(source_label_frame, text="📄 Source", 
                font=('Segoe UI', 10, 'bold'), fg='#0066cc').pack(side='left')
        tk.Label(source_label_frame, text="(Read-only, select to copy)", 
                font=('Segoe UI', 8), fg='#666').pack(side='left', padx=(10, 0))
        
        source_widget = scrolledtext.ScrolledText(main_frame, wrap='word', 
                                                  font=('Segoe UI', 10),
                                                  bg='#f0f8ff',
                                                  relief='solid',
                                                  borderwidth=1,
                                                  height=max(2, source_lines))
        source_widget.pack(fill='both', expand=True, pady=(0, 10))
        source_widget.insert('1.0', self.current_segment.source)
        source_widget.config(state='normal')  # Allow selection for copying
        
        # Clear divider line (memoQ-style)
        divider = tk.Frame(main_frame, height=2, bg='#0066cc', relief='solid')
        divider.pack(fill='x', pady=10)
        
        # Target section
        target_label_frame = tk.Frame(main_frame)
        target_label_frame.pack(fill='x', pady=(0, 5))
        tk.Label(target_label_frame, text="🎯 Target", 
                font=('Segoe UI', 10, 'bold'), fg='#009900').pack(side='left')
        tk.Label(target_label_frame, text="(Read-only)", 
                font=('Segoe UI', 8), fg='#666').pack(side='left', padx=(10, 0))
        
        target_widget = scrolledtext.ScrolledText(main_frame, wrap='word', 
                                                  font=('Segoe UI', 10),
                                                  bg='#f0fff0',
                                                  relief='solid',
                                                  borderwidth=1,
                                                  height=max(2, target_lines))
        target_widget.pack(fill='both', expand=True, pady=(0, 10))
        if self.current_segment.target:
            target_widget.insert('1.0', self.current_segment.target)
        else:
            target_widget.insert('1.0', '[No translation yet]')
            target_widget.config(fg='#999')
        target_widget.config(state='normal')  # Allow selection for copying
        
        # Button frame
        btn_frame = tk.Frame(popup, padx=15, pady=10, bg='#f0f0f0')
        btn_frame.pack(fill='x', side='bottom')
        
        tk.Button(btn_frame, text="📋 Copy Source to Clipboard", 
                 command=lambda: self.copy_to_clipboard(self.current_segment.source, popup),
                 bg='#2196F3', fg='white', padx=10).pack(side='left', padx=5)
        tk.Button(btn_frame, text="✨ Copy Source → Target", 
                 command=lambda: [self.copy_source_to_target(), popup.destroy()],
                 bg='#4CAF50', fg='white', padx=10).pack(side='left', padx=5)
        tk.Button(btn_frame, text="✖ Close", 
                 command=popup.destroy,
                 padx=10).pack(side='right', padx=5)
        
        # Auto-select source text for easy copying
        source_widget.tag_add('sel', '1.0', 'end')
        source_widget.focus()
        
        # Close on Escape
        popup.bind('<Escape>', lambda e: popup.destroy())
    
    def copy_to_clipboard(self, text, popup=None):
        """Copy text to clipboard"""
        self.root.clipboard_clear()
        self.root.clipboard_append(text)
        self.log("✓ Copied to clipboard")
        if popup:
            popup.destroy()
    
    def import_txt_bilingual(self):
        """Import bilingual TXT file (memoQ/CAT tool format) or source-only TXT"""
        file_path = filedialog.askopenfilename(
            title="Select Bilingual TXT file",
            filetypes=[("Text Files", "*.txt"), ("TSV Files", "*.tsv"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return
        
        self.import_txt_from_path(file_path)
    
    def import_txt_from_path(self, file_path):
        """Import TXT file from a given path"""
        try:
            self.log(f"Importing text file: {os.path.basename(file_path)}")
            
            # Clear existing segments
            self.segments = []
            
            # Detect delimiter (tab or comma)
            with open(file_path, 'r', encoding='utf-8') as f:
                first_line = f.readline()
                
                # Smart delimiter detection:
                # - If file has tabs, use tab delimiter
                # - If file has NO tabs but has commas in EVERY line, assume CSV
                # - Otherwise, treat as single-column (no delimiter splitting)
                if '\t' in first_line:
                    delimiter = '\t'
                else:
                    # Check if this looks like a CSV (commas in multiple lines)
                    f.seek(0)
                    sample_lines = [f.readline() for _ in range(min(5, sum(1 for _ in f)))]
                    f.seek(0)
                    
                    # Count lines with commas
                    comma_count = sum(1 for line in sample_lines if ',' in line)
                    
                    # If most lines have commas, assume CSV
                    if comma_count >= len(sample_lines) * 0.8:
                        delimiter = ','
                    else:
                        # Single column file, use a delimiter that won't match
                        delimiter = '\t'  # Use tab but file won't have any
                
                f.seek(0)  # Reset to beginning
                
                # Check if first line is header
                has_header = any(header in first_line.lower() for header in ['id', 'source', 'target', 'segment'])
                
                if has_header:
                    next(f)  # Skip header row
                
                # Parse segments
                for line_num, line in enumerate(f, 1):
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Split by delimiter
                    parts = line.split(delimiter)
                    
                    if len(parts) < 1:
                        self.log(f"Warning: Line {line_num} is empty, skipping")
                        continue
                    
                    # Determine format based on number of columns
                    if len(parts) == 1:
                        # Single column: Source only (no pre-translation)
                        seg_id = line_num
                        source = parts[0].strip()
                        target = ""
                        
                    elif len(parts) == 2:
                        # Two columns: Check if first is numeric ID or source text
                        if parts[0].strip().isdigit():
                            # Format: ID, Source (no target)
                            seg_id = int(parts[0].strip())
                            source = parts[1].strip()
                            target = ""
                        else:
                            # Format: Source, Target (bilingual)
                            seg_id = line_num
                            source = parts[0].strip()
                            target = parts[1].strip()
                    
                    else:  # 3 or more columns
                        # Format: ID, Source, Target (full bilingual)
                        if parts[0].strip().isdigit():
                            seg_id = int(parts[0].strip())
                            source = parts[1].strip()
                            target = parts[2].strip() if len(parts) > 2 else ""
                        else:
                            # No numeric ID, use line number
                            seg_id = line_num
                            source = parts[0].strip()
                            target = parts[1].strip() if len(parts) > 1 else ""
                    
                    # Create segment
                    segment = Segment(seg_id, source)
                    segment.target = target
                    segment.status = "translated" if target.strip() else "untranslated"
                    segment.modified = False
                    self.segments.append(segment)
            
            # Load into grid
            self.load_segments_to_grid()
            
            # Update status with format detection
            translated_count = sum(1 for seg in self.segments if seg.target.strip())
            
            # Determine what format was detected
            if translated_count == 0:
                format_msg = "source-only format"
            elif translated_count == len(self.segments):
                format_msg = "fully translated bilingual format"
            else:
                format_msg = "partially translated bilingual format"
            
            # Switch from Start Screen to Grid View if needed
            if hasattr(self, 'start_paned'):
                self.switch_from_start_to_grid()
            
            # Load into grid
            self.load_segments_to_grid()
            
            self.log(f"✓ Loaded {len(self.segments)} segments ({format_msg})")
            self.log(f"  Pre-translated: {translated_count}, Untranslated: {len(self.segments) - translated_count}")
            self.update_progress()
            self.modified = False
            
            # Store original file for reference
            self.original_txt = file_path
            
        except Exception as e:
            messagebox.showerror("Import Error", f"Failed to import text file:\n{str(e)}")
            self.log(f"✗ Import failed: {str(e)}")
    
    def import_docx(self):
        """Import a DOCX file"""
        file_path = filedialog.askopenfilename(
            title="Select DOCX file",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return
        
        self.import_docx_from_path(file_path)
    
    def import_docx_from_path(self, file_path):
        """Import a DOCX file from a given path"""
        try:
            self.log(f"Importing: {os.path.basename(file_path)}")
            
            # Import DOCX
            paragraphs = self.docx_handler.import_docx(file_path)
            self.original_docx = file_path
            
            # Segment paragraphs
            self.log("Segmenting text...")
            segmented = self.segmenter.segment_paragraphs(paragraphs)
            
            # Create segments with table and style information
            self.segments = []
            for seg_id, (para_id, text) in enumerate(segmented, 1):
                # Get paragraph info to check if it's a table cell and get style
                para_info = self.docx_handler._get_para_info(para_id)
                
                is_table = False
                table_info = None
                style = "Normal"  # Default style
                doc_position = para_id  # Use para_id as fallback
                
                if para_info:
                    # Get style information
                    style = para_info.style or "Normal"
                    
                    # Get document position for proper ordering
                    doc_position = para_info.document_position
                    
                    # Get table information
                    if para_info.is_table_cell:
                        is_table = True
                        table_info = (para_info.table_index, para_info.row_index, para_info.cell_index)
                
                segment = Segment(seg_id, text, para_id, is_table, table_info, style, doc_position)
                self.segments.append(segment)
            
            # Switch from Start Screen to Grid View if needed
            if hasattr(self, 'start_paned'):
                self.switch_from_start_to_grid()
            
            # Load into grid
            self.load_segments_to_grid()
            
            # Update status
            self.log(f"✓ Loaded {len(self.segments)} segments from {len(paragraphs)} paragraphs")
            self.update_progress()
            self.modified = False
            
        except Exception as e:
            messagebox.showerror("Import Error", f"Failed to import DOCX:\n{str(e)}")
            self.log(f"✗ Import failed: {str(e)}")
    
    def load_drawings(self):
        """Load drawing images from a folder for multimodal translation context"""
        if not PIL_AVAILABLE:
            messagebox.showwarning("PIL Not Available", 
                                 "PIL/Pillow library is not installed. Image support is disabled.\n\n"
                                 "To enable image support, install Pillow:\npip install Pillow")
            return
        
        folder_path = filedialog.askdirectory(title="Select folder containing drawing images")
        if not folder_path:
            return
        
        try:
            self.log(f"[Drawings] Loading images from: {folder_path}")
            self.drawings_folder = folder_path
            self.drawings_images_map = {}
            
            # Supported image formats
            image_extensions = ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff')
            
            # Load all images from folder
            loaded_count = 0
            for filename in os.listdir(folder_path):
                if filename.lower().endswith(image_extensions):
                    img_path = os.path.join(folder_path, filename)
                    try:
                        img = Image.open(img_path)
                        # Normalize the filename to match figure references
                        normalized_name = normalize_figure_ref(filename)
                        if normalized_name:
                            self.drawings_images_map[normalized_name] = img
                            loaded_count += 1
                            self.log(f"[Drawings] Loaded: {filename} as '{normalized_name}'")
                    except Exception as e:
                        self.log(f"[Drawings] Failed to load {filename}: {e}")
            
            if loaded_count > 0:
                self.log(f"[Drawings] ✓ Successfully loaded {loaded_count} drawing images")
                messagebox.showinfo("Drawings Loaded", 
                                  f"Loaded {loaded_count} drawing images.\n\n"
                                  "These will be used as context for figure references during translation.")
            else:
                self.log(f"[Drawings] ⚠ No valid images found in folder")
                messagebox.showwarning("No Images", "No valid image files found in selected folder.")
                
        except Exception as e:
            self.log(f"[Drawings] Error loading images: {e}")
            messagebox.showerror("Load Error", f"Failed to load drawings:\n{str(e)}")
    
    def clear_drawings(self):
        """Clear all loaded drawings"""
        self.drawings_images_map.clear()
        self.drawings_folder = None
        self.log("[Drawings] All drawings cleared")
        messagebox.showinfo("Drawings Cleared", "All loaded drawings have been cleared.")
    
    # ===== GRID PAGINATION METHODS =====
    
    def update_pagination_controls(self, start_idx, end_idx, total, total_pages):
        """Update pagination UI labels and button states"""
        # Update label
        self.pagination_label.config(text=f"Segments {start_idx + 1}-{end_idx} of {total}")
        
        # Update page number
        self.page_number_var.set(str(self.grid_current_page + 1))
        self.total_pages_label.config(text=f"of {total_pages}")
        
        # Enable/disable buttons
        is_first_page = self.grid_current_page == 0
        is_last_page = self.grid_current_page >= total_pages - 1
        
        self.first_page_btn.config(state='disabled' if is_first_page else 'normal')
        self.prev_page_btn.config(state='disabled' if is_first_page else 'normal')
        self.next_page_btn.config(state='disabled' if is_last_page else 'normal')
        self.last_page_btn.config(state='disabled' if is_last_page else 'normal')
    
    def go_to_first_page(self):
        """Navigate to first page"""
        self.grid_current_page = 0
        self.load_segments_to_grid()
    
    def go_to_prev_page(self):
        """Navigate to previous page"""
        if self.grid_current_page > 0:
            self.grid_current_page -= 1
            self.load_segments_to_grid()
    
    def go_to_next_page(self):
        """Navigate to next page"""
        all_segments = self.filtered_segments if self.filter_active else self.segments
        total_segments = len(all_segments)
        
        if self.page_size_var.get() == 'All':
            return  # No next page in "All" mode
        
        page_size = int(self.page_size_var.get())
        total_pages = (total_segments + page_size - 1) // page_size
        
        if self.grid_current_page < total_pages - 1:
            self.grid_current_page += 1
            self.load_segments_to_grid()
    
    def go_to_last_page(self):
        """Navigate to last page"""
        all_segments = self.filtered_segments if self.filter_active else self.segments
        total_segments = len(all_segments)
        
        if self.page_size_var.get() == 'All':
            return  # Already on "last" page in "All" mode
        
        page_size = int(self.page_size_var.get())
        total_pages = max(1, (total_segments + page_size - 1) // page_size)
        
        self.grid_current_page = total_pages - 1
        self.load_segments_to_grid()
    
    def go_to_page(self):
        """Navigate to specific page number"""
        try:
            page_num = int(self.page_number_var.get())
            
            all_segments = self.filtered_segments if self.filter_active else self.segments
            total_segments = len(all_segments)
            
            if self.page_size_var.get() == 'All':
                return  # No page navigation in "All" mode
            
            page_size = int(self.page_size_var.get())
            total_pages = max(1, (total_segments + page_size - 1) // page_size)
            
            # Validate page number (1-indexed for user, 0-indexed internally)
            if 1 <= page_num <= total_pages:
                self.grid_current_page = page_num - 1
                self.load_segments_to_grid()
            else:
                messagebox.showwarning("Invalid Page", 
                                      f"Please enter a page number between 1 and {total_pages}")
                self.page_number_var.set(str(self.grid_current_page + 1))
        except ValueError:
            messagebox.showwarning("Invalid Input", "Please enter a valid page number")
            self.page_number_var.set(str(self.grid_current_page + 1))
    
    def on_page_size_changed(self, event=None):
        """Handle page size change"""
        # Reset to first page when changing page size
        self.grid_current_page = 0
        self.load_segments_to_grid()
    
    # ===== END GRID PAGINATION METHODS =====
    
    def load_segments_to_tree(self):
        """Load segments into Treeview (for List View)"""
        self.log(f"load_segments_to_tree() called. Has tree: {hasattr(self, 'tree')}")
        
        # Check if tree exists
        if not hasattr(self, 'tree'):
            self.log("⚠ Tree view not created yet, skipping load")
            return
        
        self.log(f"Tree exists, attempting to load segments...")
        
        # Clear existing
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        # Use filtered segments if filter is active, otherwise all segments
        segments_to_show = self.filtered_segments if self.filter_active else self.segments
        
        self.log(f"Segments info: total={len(self.segments)}, filter_active={self.filter_active}, filtered={len(self.filtered_segments)}, to_show={len(segments_to_show)}")
        
        # Debug: Check if we have segments
        if not segments_to_show:
            self.log(f"⚠ No segments to display in List View. filter_active={self.filter_active}, segments count={len(self.segments)}, filtered count={len(self.filtered_segments)}")
            return
        
        self.log(f"Loading {len(segments_to_show)} segments to tree...")
        
        # Add segments
        for seg in segments_to_show:
            # Determine type label
            if seg.is_table_cell and seg.table_info:
                type_label = f"T{seg.table_info[0]+1}R{seg.table_info[1]+1}C{seg.table_info[2]+1}"
            else:
                type_label = "Para"
            
            # Format style name for display
            style_display = self._format_style_name(seg.style)
            
            # Set tags for styling
            tags = [seg.status]
            if seg.is_table_cell:
                tags.append('table_cell')
            
            # Add style-specific tag for visual formatting
            style_tag = self._get_style_tag(seg.style)
            if style_tag:
                tags.append(style_tag)
            
            # Prepare source and target text with potential highlighting
            source_text = self._truncate(seg.source, 75)
            target_text = self._truncate(seg.target, 75)
            
            # Note: Treeview doesn't support text highlighting like Text widget
            # In highlight mode, matches are shown but not visually highlighted in the tree
            # The highlighting will be visible in the editor panel below
            
            # Insert into tree
            item_id = self.tree.insert('', 'end',
                           values=(seg.id, type_label, style_display, seg.status.capitalize(),
                                  source_text, target_text),
                           tags=tuple(tags))
            if len(tags) == 0:  # Log first item only
                self.log(f"  Inserted segment {seg.id}: {source_text[:30]}... -> item_id={item_id}")
        
        # Select first segment if available
        children = self.tree.get_children()
        self.log(f"After loading, tree has {len(children)} children")
        if children:
            self.tree.selection_set(children[0])
            self.tree.focus(children[0])
            self.on_segment_select(None)
    
    
    
    def set_filter_mode(self, mode):
        """Set the filter mode and update button states"""
        if self.filter_mode == mode:
            return  # Already in this mode
        
        self.filter_mode = mode
        
        if mode == 'filter':
            # Filter mode active - show only matches
            self.filter_mode_btn.config(bg='#4CAF50', fg='white', 
                                       font=('Segoe UI', 9, 'bold'), relief='sunken')
            self.highlight_mode_btn.config(bg='#e0e0e0', fg='#666',
                                          font=('Segoe UI', 9), relief='raised')
            self.log("� Filter Mode: Only matching segments shown")
        else:
            # Highlight mode active - show all with highlights
            self.filter_mode_btn.config(bg='#e0e0e0', fg='#666',
                                       font=('Segoe UI', 9), relief='raised')
            self.highlight_mode_btn.config(bg='#FFA500', fg='white',
                                          font=('Segoe UI', 9, 'bold'), relief='sunken')
            self.log("� Highlight Mode: All segments shown with matches highlighted")
        
        # Reapply filters with new mode
        self.apply_filters()
    
    def apply_filters(self):
        """Apply filters to the segment list"""
        if not hasattr(self, 'segments') or not self.segments:
            return
        
        source_filter = self.filter_source_var.get().strip().lower()
        target_filter = self.filter_target_var.get().strip().lower()
        status_filter = self.filter_status_var.get()
        
        # Check if any filter is active
        self.filter_active = bool(source_filter or target_filter or (status_filter != "All"))
        
        if not self.filter_active:
            # No filters - show all segments
            self.filtered_segments = self.segments.copy()
            if hasattr(self, 'filter_results_label') and self.filter_results_label.winfo_exists():
                self.filter_results_label.config(text="")
        elif self.filter_mode == 'highlight':
            # Highlight mode: show ALL segments but highlight matches
            self.filtered_segments = self.segments.copy()
            
            # Count how many segments have matches (for info label)
            match_count = 0
            for seg in self.segments:
                has_match = True
                if source_filter and source_filter not in seg.source.lower():
                    has_match = False
                if target_filter and target_filter not in seg.target.lower():
                    has_match = False
                if status_filter != "All" and seg.status != status_filter:
                    has_match = False
                if has_match:
                    match_count += 1
            
            # Update results label
            total = len(self.segments)
            if hasattr(self, 'filter_results_label') and self.filter_results_label.winfo_exists():
                self.filter_results_label.config(
                    text=f"💡 Highlighting {match_count} of {total} segments",
                    fg='#4CAF50' if match_count > 0 else '#FF5722'
                )
        else:
            # Filter mode: show ONLY matching segments
            self.filtered_segments = []
            for seg in self.segments:
                # Source filter
                if source_filter and source_filter not in seg.source.lower():
                    continue
                
                # Target filter
                if target_filter and target_filter not in seg.target.lower():
                    continue
                
                # Status filter
                if status_filter != "All" and seg.status != status_filter:
                    continue
                
                # Segment passed all filters
                self.filtered_segments.append(seg)
            
            # Update results label
            total = len(self.segments)
            filtered = len(self.filtered_segments)
            if hasattr(self, 'filter_results_label') and self.filter_results_label.winfo_exists():
                self.filter_results_label.config(
                    text=f"🔍 Showing {filtered} of {total} segments",
                    fg='#4CAF50' if filtered > 0 else '#FF5722'
                )
        
        # Reload the grid with filtered segments
        self.reload_grid_with_filters()
    
    def clear_filters(self):
        """Clear all filters"""
        self.filter_source_var.set("")
        self.filter_target_var.set("")
        self.filter_status_var.set("All")
        self.filter_active = False
        if hasattr(self, 'filter_results_label') and self.filter_results_label.winfo_exists():
            self.filter_results_label.config(text="")
        self.log("✓ Filters cleared")
    
    def toggle_filter_mode(self):
        """Toggle between filter and highlight modes (Ctrl+M)"""
        new_mode = 'highlight' if self.filter_mode == 'filter' else 'filter'
        self.set_filter_mode(new_mode)
    
    def focus_filter_source(self):
        """Focus the source filter entry field (Ctrl+Shift+F)"""
        # Find the source filter entry in the current view
        # This is a bit tricky since we need to find the widget
        # We'll use a simple approach: find all Entry widgets and check their textvariable
        try:
            for widget in self.root.winfo_children():
                self._focus_filter_recursive(widget)
        except:
            pass
    
    def _focus_filter_recursive(self, widget):
        """Recursively search for source filter entry and focus it"""
        try:
            if isinstance(widget, tk.Entry):
                # Check if this entry uses the source filter variable
                if hasattr(widget, 'cget') and widget.cget('textvariable') == str(self.filter_source_var):
                    widget.focus_set()
                    return True
            # Recurse into children
            for child in widget.winfo_children():
                if self._focus_filter_recursive(child):
                    return True
        except:
            pass
        return False
    
    def highlight_text_in_widget(self, text_widget, search_text, color='yellow'):
        """Highlight all occurrences of search_text in a Text widget"""
        if not search_text:
            return
        
        # Configure tag for highlighting
        tag_name = f"highlight_{color}"
        text_widget.tag_configure(tag_name, background=color, foreground='black')
        
        # Get the content
        content = text_widget.get('1.0', 'end-1c')
        search_lower = search_text.lower()
        content_lower = content.lower()
        
        # Find all occurrences (case-insensitive)
        start_pos = 0
        while True:
            start_pos = content_lower.find(search_lower, start_pos)
            if start_pos == -1:
                break
            
            # Calculate Text widget position
            lines_before = content[:start_pos].count('\n')
            chars_in_line = start_pos - content.rfind('\n', 0, start_pos) - 1
            
            start_index = f"{lines_before + 1}.{chars_in_line}"
            end_index = f"{lines_before + 1}.{chars_in_line + len(search_text)}"
            
            # Apply tag
            text_widget.tag_add(tag_name, start_index, end_index)
            
            start_pos += len(search_text)
    
    def reload_grid_with_filters(self):
        """Reload the current view with filtered segments"""
        # Clear selection when filter changes (selected segments may not be visible)
        self.selected_segments.clear()
        self.last_selected_index = None
        self.update_selection_counter()
        
        if self.layout_mode == LayoutMode.GRID:
            # Grid View: Clear and rebuild grid
            self.grid_rows = []
            for widget in self.grid_inner_frame.winfo_children():
                widget.destroy()
            
            # Use filtered segments if filter is active, otherwise all segments
            segments_to_show = self.filtered_segments if self.filter_active else self.segments
            
            # Add filtered segments to grid
            for seg in segments_to_show:
                self.add_grid_row(seg)
            
            # Select first row if available
            if self.grid_rows:
                self.select_grid_row(0)
        
        elif self.layout_mode == LayoutMode.SPLIT:
            # List View: Reload tree
            self.load_segments_to_tree()
        
        elif self.layout_mode == LayoutMode.DOCUMENT:
            # Document View: Reload document
            self.load_segments_to_document()
        # Log the filter result
        if self.filter_active:
            self.log(f"🔍 Filter applied: {len(segments_to_show)} segments match")
    
    def load_segments_to_grid(self):
        """Load segments into the grid with pagination (optimized for large documents)"""
        # SAFETY: Disable ALL interaction during loading to prevent freezing/resize issues
        self.root.config(cursor='wait')
        
        # Block all mouse/keyboard events on content frame
        loading_blocker = None
        if hasattr(self, 'content_frame'):
            # Create a transparent overlay to block all interactions
            loading_blocker = tk.Frame(self.content_frame, bg='white', cursor='wait')
            loading_blocker.place(relx=0, rely=0, relwidth=1, relheight=1)
            loading_blocker.lift()  # Bring to front
            
            # Show loading message on blocker
            loading_msg = tk.Label(loading_blocker, 
                                  text="Loading page...\nPlease wait.",
                                  font=('Segoe UI', 14, 'bold'),
                                  bg='white', fg='gray')
            loading_msg.place(relx=0.5, rely=0.5, anchor='center')
            self.root.update_idletasks()
        
        try:
            # Disable updates during bulk loading to prevent freezing
            if self.layout_mode == LayoutMode.GRID:
                # Clear existing custom grid rows
                self.grid_rows = []
                for widget in self.grid_inner_frame.winfo_children():
                    widget.destroy()
                
                # Determine which segments to show (filtered or all)
                all_segments = self.filtered_segments if self.filter_active else self.segments
                total_segments = len(all_segments)
                
                # Calculate pagination
                if self.page_size_var.get() == 'All':
                    page_size = total_segments
                else:
                    page_size = int(self.page_size_var.get())
                
                self.grid_page_size = page_size
                total_pages = max(1, (total_segments + page_size - 1) // page_size)  # Ceiling division
                
                # Ensure current page is valid
                if self.grid_current_page >= total_pages:
                    self.grid_current_page = max(0, total_pages - 1)
                
                # Calculate slice for current page
                start_idx = self.grid_current_page * page_size
                end_idx = min(start_idx + page_size, total_segments)
                segments_to_show = all_segments[start_idx:end_idx]
                
                # Update pagination UI
                self.update_pagination_controls(start_idx, end_idx, total_segments, total_pages)
                
                # Add segments to custom grid (only current page)
                for seg in segments_to_show:
                    self.add_grid_row(seg)
                
                # Select first row if available
                if self.grid_rows:
                    self.select_grid_row(0)
                    
            else:
                # Use Treeview for Split/Compact modes
                # Temporarily disable redrawing for performance
                self.tree.configure(height=0)  # Collapse to prevent redraw on each insert
                
                # Clear existing
                for item in self.tree.get_children():
                    self.tree.delete(item)
                
                # Batch insert segments (more efficient)
                items_to_insert = []
                for seg in self.segments:
                    # Determine type label
                    if seg.is_table_cell and seg.table_info:
                        type_label = f"T{seg.table_info[0]+1}R{seg.table_info[1]+1}C{seg.table_info[2]+1}"
                    else:
                        type_label = "Para"
                    
                    # Format style name for display
                    style_display = self._format_style_name(seg.style)
                    
                    # Set tags for styling
                    tags = [seg.status]
                    if seg.is_table_cell:
                        tags.append('table_cell')
                    
                    # Add style-specific tag for visual formatting
                    style_tag = self._get_style_tag(seg.style)
                    if style_tag:
                        tags.append(style_tag)
                    
                    items_to_insert.append((seg, type_label, style_display, tags))
                
                # Insert all at once (prevents multiple redraws)
                for seg, type_label, style_display, tags in items_to_insert:
                    self.tree.insert('', 'end',
                                   values=(seg.id, type_label, style_display, seg.status.capitalize(),
                                              self._truncate(seg.source, 75),
                                              self._truncate(seg.target, 75)),
                                       tags=tuple(tags))
                
                # Re-enable treeview with proper height
                self.tree.configure(height=20)
        
        finally:
            # SAFETY: Re-enable interaction after loading is complete
            if loading_blocker:
                loading_blocker.destroy()
            self.root.config(cursor='')
            if hasattr(self, 'grid_canvas'):
                self.grid_canvas.config(state='normal')
    
    def _truncate(self, text: str, length: int) -> str:
        """Truncate text for display"""
        if len(text) <= length:
            return text
        return text[:length-3] + "..."
    
    def _truncate_multiline(self, text: str, length: int) -> str:
        """Truncate text but preserve newlines for better multi-line display in grid"""
        if not text:
            return ""
        
        # If text is short enough, return as-is
        if len(text) <= length:
            return text
        
        # For longer text, show first 2 lines or up to length chars, whichever is less
        lines = text.split('\n')
        if len(lines) > 2:
            # Show first 2 lines + indicator
            result = '\n'.join(lines[:2])
            if len(result) > length:
                result = result[:length-3] + "..."
            else:
                result += "\n..."
            return result
        else:
            # Single line or 2 lines - just truncate if too long
            if len(text) > length:
                return text[:length-3] + "..."
            return text
    
    def _format_style_name(self, style: str) -> str:
        """Format style name for display in grid"""
        if not style or style == "Normal":
            return "Normal"
        # Shorten common styles for better display
        style = style.replace("Heading", "H").replace("heading", "H")
        # Limit length
        if len(style) > 12:
            return style[:12]
        return style
    
    def _get_style_tag(self, style: str) -> str:
        """Get treeview tag for style-specific formatting"""
        if not style:
            return None
        style_lower = style.lower()
        if 'heading 1' in style_lower or style_lower == 'heading1':
            return 'heading1'
        elif 'heading 2' in style_lower or style_lower == 'heading2':
            return 'heading2'
        elif 'heading 3' in style_lower or style_lower == 'heading3':
            return 'heading3'
        elif 'title' in style_lower and 'sub' not in style_lower:
            return 'title'
        elif 'subtitle' in style_lower:
            return 'subtitle'
        return None
    
    def on_segment_select(self, event):
        """Handle segment selection in grid"""
        self.log("📍 on_segment_select() called")  # DEBUG
        selection = self.tree.selection()
        if not selection:
            self.log("⚠ No selection - returning")  # DEBUG
            return
        
        # Save current segment first
        if self.current_segment:
            self.save_current_segment()
        
        # Get selected segment
        item = selection[0]
        values = self.tree.item(item, 'values')
        seg_id = int(values[0])
        
        # Find segment
        self.current_segment = next((s for s in self.segments if s.id == seg_id), None)
        
        if self.current_segment:
            self.log(f"✓ Segment #{self.current_segment.id} selected")  # DEBUG
            self.load_segment_to_editor(self.current_segment)
            
            # Trigger automatic TM search after delay
            self.log("🎯 About to call schedule_auto_tm_search()")  # DEBUG
            self.schedule_auto_tm_search()
            self.log("✓ schedule_auto_tm_search() called")  # DEBUG
        else:
            self.log("⚠ current_segment is None")  # DEBUG
    
    def load_segment_to_editor(self, segment: Segment):
        """Load segment into editor panel"""
        self.seg_info_label.config(text=f"Segment #{segment.id} | Paragraph {segment.paragraph_id}")
        
        # Source
        self.source_text.config(state='normal')
        self.source_text.delete('1.0', 'end')
        self.source_text.insert('1.0', segment.source)
        self.source_text.config(state='disabled')
        
        # Target
        self.target_text.delete('1.0', 'end')
        if segment.target:
            self.target_text.insert('1.0', segment.target)
        
        # Status
        self.status_var.set(segment.status)
        
        # Focus target
        self.target_text.focus_set()
    
    def save_current_segment(self):
        """Save current segment from editor"""
        if not self.current_segment:
            return
        
        # Check layout mode and get values accordingly
        if self.layout_mode == LayoutMode.GRID:
            # In Grid mode, save from grid editor if active
            if hasattr(self, 'save_grid_editor_segment'):
                self.save_grid_editor_segment()
            return
        
        # List/Document mode - use target_text widget
        if not hasattr(self, 'target_text'):
            return
        
        # Get values
        target = self.target_text.get('1.0', 'end-1c').strip()
        status = self.status_var.get()
        
        # Update segment
        if target != self.current_segment.target or status != self.current_segment.status:
            self.current_segment.target = target
            self.current_segment.status = status
            self.current_segment.modified = True
            self.current_segment.modified_at = datetime.now().isoformat()
            self.modified = True
            
            # Add to Project TM if translated or approved and has content
            if status in ['translated', 'approved'] and target:
                self.tm_database.add_to_project_tm(self.current_segment.source, target)
                self.log(f"✓ Added to Project TM: {self.current_segment.source[:50]}...")
            
            # Update grid
            self.update_segment_in_grid(self.current_segment)
            self.update_progress()
    
    def update_segment_in_grid(self, segment: Segment):
        """Update segment display in grid"""
        for item in self.tree.get_children():
            values = self.tree.item(item, 'values')
            if int(values[0]) == segment.id:
                # Determine type label
                if segment.is_table_cell and segment.table_info:
                    type_label = f"T{segment.table_info[0]+1}R{segment.table_info[1]+1}C{segment.table_info[2]+1}"
                else:
                    type_label = "Para"
                
                # Format style name
                style_display = self._format_style_name(segment.style)
                
                # Set tags for styling
                tags = [segment.status]
                if segment.is_table_cell:
                    tags.append('table_cell')
                style_tag = self._get_style_tag(segment.style)
                if style_tag:
                    tags.append(style_tag)
                
                # Different column layouts for different modes
                if self.layout_mode == LayoutMode.GRID:
                    # Grid: No style column (5 columns)
                    # Use multiline truncate for better display
                    self.tree.item(item,
                                 values=(segment.id, type_label, segment.status.capitalize(),
                                        self._truncate_multiline(segment.source, 150),
                                        self._truncate_multiline(segment.target, 150)),
                                 tags=tuple(tags))
                else:
                    # Split: Include style column (6 columns)
                    self.tree.item(item,
                                 values=(segment.id, type_label, style_display, segment.status.capitalize(),
                                        self._truncate(segment.source, 75),
                                        self._truncate(segment.target, 75)),
                                 tags=tuple(tags))
                break
    
    def on_status_change(self, event):
        """Handle status change"""
        self.save_current_segment()
    
    def on_target_change(self, event):
        """Handle target text change - validate tags"""
        target = self.target_text.get('1.0', 'end-1c')
        
        # Validate tags
        is_valid, error_msg = self.tag_manager.validate_tags(target)
        
        if not is_valid:
            self.tag_validation_label.config(text=f"⚠️ {error_msg}", fg='red')
        else:
            # Count tags
            tag_counts = self.tag_manager.count_tags(target)
            if tag_counts:
                tag_text = ', '.join([f"{count} {tag}" for tag, count in tag_counts.items()])
                self.tag_validation_label.config(text=f"✓ Tags: {tag_text}", fg='green')
            else:
                self.tag_validation_label.config(text="", fg='black')
    
    def insert_tag(self, tag_name: str):
        """Insert formatting tag at cursor position"""
        try:
            # Get current selection or cursor position
            try:
                start = self.target_text.index('sel.first')
                end = self.target_text.index('sel.last')
                selected_text = self.target_text.get(start, end)
                
                # Wrap selection in tags
                tagged_text = f"<{tag_name}>{selected_text}</{tag_name}>"
                self.target_text.delete(start, end)
                self.target_text.insert(start, tagged_text)
            except tk.TclError:
                # No selection, insert empty tags at cursor
                cursor_pos = self.target_text.index('insert')
                self.target_text.insert(cursor_pos, f"<{tag_name}></{tag_name}>")
                # Move cursor between tags
                self.target_text.mark_set('insert', f"{cursor_pos}+{len(tag_name)+2}c")
            
            self.target_text.focus_set()
            # Trigger validation
            self.on_target_change(None)
        except Exception as e:
            self.log(f"Error inserting tag: {str(e)}")
    
    def strip_tags_from_target(self):
        """Remove all formatting tags from target"""
        if not self.current_segment:
            return
        
        target = self.target_text.get('1.0', 'end-1c')
        clean_text = self.tag_manager.strip_tags(target)
        
        self.target_text.delete('1.0', 'end')
        self.target_text.insert('1.0', clean_text)
        self.on_target_change(None)
    
    def copy_source_tags(self):
        """Copy tag structure from source to target"""
        if not self.current_segment:
            return
        
        source = self.current_segment.source
        target = self.target_text.get('1.0', 'end-1c')
        
        # Extract tag positions from source
        source_tags = self.tag_manager.count_tags(source)
        
        if not source_tags:
            messagebox.showinfo("No Tags", "Source segment has no formatting tags.")
            return
        
        # If target is empty, copy source structure
        if not target.strip():
            self.target_text.delete('1.0', 'end')
            self.target_text.insert('1.0', source)
            messagebox.showinfo("Tags Copied", "Source formatting structure copied to target. Now translate the text.")
        else:
            messagebox.showinfo("Tag Structure", 
                              f"Source has: {', '.join([f'{c} {t}' for t, c in source_tags.items()])}\n\n"
                              "Use the Insert buttons to add matching tags to your translation.")
        
        self.on_target_change(None)
    
    def copy_source_to_target(self):
        """Copy source to target - works in both Grid and Split modes"""
        if not self.current_segment:
            self.log("⚠️ No segment selected")
            return
        
        # Update the segment
        self.current_segment.target = self.current_segment.source
        if self.current_segment.status == 'untranslated':
            self.current_segment.status = 'draft'
        self.current_segment.modified = True
        
        # Update UI based on mode
        if self.layout_mode == LayoutMode.GRID:
            # In Grid mode with custom grid
            if self.current_row_index >= 0 and self.current_row_index < len(self.grid_rows):
                self.update_grid_row(self.current_row_index)
                
                # If currently editing, update the edit widget
                if self.current_edit_widget:
                    self.current_edit_widget.delete('1.0', 'end')
                    self.current_edit_widget.insert('1.0', self.current_segment.source)
        else:
            # In Split mode, update the target text widget  
            if hasattr(self, 'target_text'):
                self.target_text.delete('1.0', 'end')
                self.target_text.insert('1.0', self.current_segment.source)
        
        self.modified = True
        self.update_progress()
        self.log(f"✓ Copied source to target (Segment #{self.current_segment.id})")
    
    def copy_source_to_target_all(self):
        """Copy source to target for ALL segments (useful for CAT tool re-import workflow)"""
        if not self.segments:
            messagebox.showwarning("No Segments", "No segments loaded.")
            return
        
        # Confirm action
        if not messagebox.askyesno(
            "Copy Source to Target (All)",
            f"Copy source text to target for ALL {len(self.segments)} segments?\n\n"
            "This is useful for preparing files for Trados/memoQ/CafeTran re-import.\n\n"
            "💡 TIP: After copying, translate with AI, then export.\n"
            "This avoids paying for MT pre-translation in your CAT tool!"
        ):
            return
        
        # Copy source to target for all segments
        copied_count = 0
        for segment in self.segments:
            segment.target = segment.source
            if segment.status == 'untranslated':
                segment.status = 'draft'
            segment.modified = True
            copied_count += 1
        
        # Update UI
        self.modified = True
        if self.layout_mode == LayoutMode.GRID:
            self.load_segments_to_grid()
        elif self.layout_mode == LayoutMode.SPLIT:
            self.load_segments_to_list()
            # Reload current segment if one is selected
            if self.current_segment and hasattr(self, 'target_text'):
                self.target_text.delete('1.0', 'end')
                self.target_text.insert('1.0', self.current_segment.target)
        elif self.layout_mode == LayoutMode.DOCUMENT:
            self.load_segments_to_document_view()
        
        self.update_progress()
        self.log(f"✓ Copied source to target for {copied_count} segments")
        messagebox.showinfo(
            "Complete",
            f"✓ Copied source to target for {copied_count} segments.\n\n"
            "Next steps:\n"
            "1. Translate with AI (Translate → Translate All)\n"
            "2. Export bilingual file\n"
            "3. Re-import into Trados/memoQ/CafeTran"
        )
    
    def clear_target(self):
        """Clear target text"""
        if hasattr(self, 'target_text'):
            self.target_text.delete('1.0', 'end')
    
    # ============================================================================
    # BULK OPERATIONS
    # ============================================================================
    
    def select_all_segments(self):
        """Select all visible segments in the current view (Ctrl+A) - Trados/memoQ style"""
        if not self.segments:
            self.log("⚠️ No segments to select")
            return
        
        # Get the segments to select (filtered or all)
        segments_to_select = self.filtered_segments if self.filter_active else self.segments
        
        if not segments_to_select:
            self.log("⚠️ No segments match current filter")
            return
        
        # Clear current selection
        self.selected_segments.clear()
        if self.layout_mode == LayoutMode.GRID:
            self._clear_all_selection_highlights()
        
        # Add all visible segments to selection
        for segment in segments_to_select:
            self.selected_segments.add(segment.id)
        
        # Apply visual highlight to visible rows in grid
        if self.layout_mode == LayoutMode.GRID:
            for i, row_data in enumerate(self.grid_rows):
                if row_data['segment'].id in self.selected_segments:
                    self._apply_selection_highlight(i, True)
        
        # Update selection counter
        self.update_selection_counter()
        
        filter_note = " (filtered)" if self.filter_active else ""
        messagebox.showinfo(
            "Select All",
            f"✓ Selected {len(segments_to_select)} segments{filter_note}.\n\n"
            f"Selected segments are highlighted in blue.\n\n"
            f"You can now apply bulk operations:\n"
            f"• Change status\n"
            f"• Lock/unlock\n"
            f"• Clear targets"
        )
    
    def clear_all_targets(self):
        """Clear target text for all segments"""
        if not self.segments:
            messagebox.showwarning("No Segments", "No segments loaded.")
            return
        
        # Count segments with targets
        segments_with_targets = sum(1 for seg in self.segments if seg.target.strip())
        
        if segments_with_targets == 0:
            messagebox.showinfo("Nothing to Clear", "All targets are already empty.")
            return
        
        # Confirm action
        if not messagebox.askyesno(
            "Clear All Targets",
            f"Clear target text for ALL {len(self.segments)} segments?\n\n"
            f"{segments_with_targets} segments have targets that will be deleted.\n\n"
            f"⚠️ This action cannot be undone!"
        ):
            return
        
        # Clear all targets
        cleared_count = 0
        for segment in self.segments:
            if segment.target.strip():
                segment.target = ""
                segment.status = "untranslated"
                segment.modified = True
                cleared_count += 1
        
        # Update UI
        self.modified = True
        if self.layout_mode == LayoutMode.GRID:
            self.load_segments_to_grid()
        elif self.layout_mode == LayoutMode.SPLIT:
            self.load_segments_to_list()
            if self.current_segment and hasattr(self, 'target_text'):
                self.target_text.delete('1.0', 'end')
        elif self.layout_mode == LayoutMode.DOCUMENT:
            self.load_segments_to_document_view()
        
        self.update_progress()
        self.log(f"✓ Cleared targets for {cleared_count} segments")
        messagebox.showinfo("Complete", f"✓ Cleared {cleared_count} targets")
    
    def change_status_all(self):
        """Change status for all segments"""
        if not self.segments:
            messagebox.showwarning("No Segments", "No segments loaded.")
            return
        
        # Create dialog for status selection
        dialog = tk.Toplevel(self.root)
        dialog.title("Change Status (All Segments)")
        dialog.geometry("400x200")
        dialog.transient(self.root)
        dialog.grab_set()
        
        tk.Label(dialog, text=f"Change status for all {len(self.segments)} segments:",
                font=('Arial', 10, 'bold')).pack(pady=10)
        
        status_var = tk.StringVar(value="translated")
        
        tk.Radiobutton(dialog, text="Untranslated", variable=status_var,
                      value="untranslated").pack(anchor='w', padx=20)
        tk.Radiobutton(dialog, text="Translated", variable=status_var,
                      value="translated").pack(anchor='w', padx=20)
        tk.Radiobutton(dialog, text="Approved", variable=status_var,
                      value="approved").pack(anchor='w', padx=20)
        tk.Radiobutton(dialog, text="Draft", variable=status_var,
                      value="draft").pack(anchor='w', padx=20)
        
        def apply_status():
            new_status = status_var.get()
            changed_count = 0
            
            for segment in self.segments:
                if segment.status != new_status:
                    segment.status = new_status
                    segment.modified = True
                    changed_count += 1
            
            if changed_count > 0:
                self.modified = True
                self.refresh_current_view()
                self.update_progress()
                self.log(f"✓ Changed status to '{new_status}' for {changed_count} segments")
            
            dialog.destroy()
            messagebox.showinfo("Complete", f"✓ Changed status for {changed_count} segments")
        
        button_frame = tk.Frame(dialog)
        button_frame.pack(pady=20)
        tk.Button(button_frame, text="Apply", command=apply_status,
                 bg='#4CAF50', fg='white', width=10).pack(side='left', padx=5)
        tk.Button(button_frame, text="Cancel", command=dialog.destroy,
                 width=10).pack(side='left', padx=5)
    
    def change_status_filtered(self):
        """Change status for filtered/visible segments only"""
        if not self.segments:
            messagebox.showwarning("No Segments", "No segments loaded.")
            return
        
        # Get the segments to modify
        segments_to_modify = self.filtered_segments if self.filter_active else self.segments
        
        if not segments_to_modify:
            messagebox.showwarning("No Segments", "No segments match current filter.")
            return
        
        # Create dialog for status selection
        dialog = tk.Toplevel(self.root)
        dialog.title("Change Status (Filtered Segments)")
        dialog.geometry("400x220")
        dialog.transient(self.root)
        dialog.grab_set()
        
        filter_info = f"{len(segments_to_modify)} filtered" if self.filter_active else f"all {len(segments_to_modify)}"
        tk.Label(dialog, text=f"Change status for {filter_info} segments:",
                font=('Arial', 10, 'bold')).pack(pady=10)
        
        if self.filter_active:
            tk.Label(dialog, text="(Only visible/filtered segments will be changed)",
                    fg='#666', font=('Arial', 8)).pack()
        
        status_var = tk.StringVar(value="translated")
        
        tk.Radiobutton(dialog, text="Untranslated", variable=status_var,
                      value="untranslated").pack(anchor='w', padx=20)
        tk.Radiobutton(dialog, text="Translated", variable=status_var,
                      value="translated").pack(anchor='w', padx=20)
        tk.Radiobutton(dialog, text="Approved", variable=status_var,
                      value="approved").pack(anchor='w', padx=20)
        tk.Radiobutton(dialog, text="Draft", variable=status_var,
                      value="draft").pack(anchor='w', padx=20)
        
        def apply_status():
            new_status = status_var.get()
            changed_count = 0
            
            for segment in segments_to_modify:
                if segment.status != new_status:
                    segment.status = new_status
                    segment.modified = True
                    changed_count += 1
            
            if changed_count > 0:
                self.modified = True
                self.refresh_current_view()
                self.update_progress()
                self.log(f"✓ Changed status to '{new_status}' for {changed_count} segments")
            
            dialog.destroy()
            messagebox.showinfo("Complete", f"✓ Changed status for {changed_count} segments")
        
        button_frame = tk.Frame(dialog)
        button_frame.pack(pady=20)
        tk.Button(button_frame, text="Apply", command=apply_status,
                 bg='#4CAF50', fg='white', width=10).pack(side='left', padx=5)
        tk.Button(button_frame, text="Cancel", command=dialog.destroy,
                 width=10).pack(side='left', padx=5)
    
    def lock_all_segments(self):
        """Lock all segments to prevent editing"""
        if not self.segments:
            messagebox.showwarning("No Segments", "No segments loaded.")
            return
        
        unlocked_count = sum(1 for seg in self.segments if not seg.locked)
        
        if unlocked_count == 0:
            messagebox.showinfo("Already Locked", "All segments are already locked.")
            return
        
        if not messagebox.askyesno(
            "Lock All Segments",
            f"Lock all {len(self.segments)} segments?\n\n"
            f"{unlocked_count} segments will be locked and cannot be edited.\n\n"
            f"You can unlock them later via Edit → Bulk Operations → Unlock All."
        ):
            return
        
        for segment in self.segments:
            segment.locked = True
            segment.modified = True
        
        self.modified = True
        self.refresh_current_view()
        self.log(f"🔒 Locked {unlocked_count} segments")
        messagebox.showinfo("Complete", f"🔒 Locked {unlocked_count} segments")
    
    def unlock_all_segments(self):
        """Unlock all segments"""
        if not self.segments:
            messagebox.showwarning("No Segments", "No segments loaded.")
            return
        
        locked_count = sum(1 for seg in self.segments if seg.locked)
        
        if locked_count == 0:
            messagebox.showinfo("Already Unlocked", "All segments are already unlocked.")
            return
        
        for segment in self.segments:
            segment.locked = False
            segment.modified = True
        
        self.modified = True
        self.refresh_current_view()
        self.log(f"🔓 Unlocked {locked_count} segments")
        messagebox.showinfo("Complete", f"🔓 Unlocked {locked_count} segments")
    
    def lock_filtered_segments(self):
        """Lock only filtered/visible segments"""
        if not self.segments:
            messagebox.showwarning("No Segments", "No segments loaded.")
            return
        
        segments_to_lock = self.filtered_segments if self.filter_active else self.segments
        
        if not segments_to_lock:
            messagebox.showwarning("No Segments", "No segments match current filter.")
            return
        
        unlocked_count = sum(1 for seg in segments_to_lock if not seg.locked)
        
        if unlocked_count == 0:
            messagebox.showinfo("Already Locked", "All filtered segments are already locked.")
            return
        
        filter_info = "filtered" if self.filter_active else "all"
        if not messagebox.askyesno(
            "Lock Filtered Segments",
            f"Lock {len(segments_to_lock)} {filter_info} segments?\n\n"
            f"{unlocked_count} segments will be locked."
        ):
            return
        
        for segment in segments_to_lock:
            segment.locked = True
            segment.modified = True
        
        self.modified = True
        self.refresh_current_view()
        self.log(f"🔒 Locked {unlocked_count} filtered segments")
        messagebox.showinfo("Complete", f"🔒 Locked {unlocked_count} segments")
    
    def unlock_filtered_segments(self):
        """Unlock only filtered/visible segments"""
        if not self.segments:
            messagebox.showwarning("No Segments", "No segments loaded.")
            return
        
        segments_to_unlock = self.filtered_segments if self.filter_active else self.segments
        
        if not segments_to_unlock:
            messagebox.showwarning("No Segments", "No segments match current filter.")
            return
        
        locked_count = sum(1 for seg in segments_to_unlock if seg.locked)
        
        if locked_count == 0:
            messagebox.showinfo("Already Unlocked", "All filtered segments are already unlocked.")
            return
        
        for segment in segments_to_unlock:
            segment.locked = False
            segment.modified = True
        
        self.modified = True
        self.refresh_current_view()
        self.log(f"🔓 Unlocked {locked_count} filtered segments")
        messagebox.showinfo("Complete", f"🔓 Unlocked {locked_count} segments")
    
    def lock_current_segment(self):
        """Lock the currently selected segment"""
        if not self.current_segment:
            self.log("⚠️ No segment selected")
            return
        
        if self.current_segment.locked:
            messagebox.showinfo("Already Locked", "This segment is already locked.")
            return
        
        self.current_segment.locked = True
        self.current_segment.modified = True
        self.modified = True
        self.refresh_current_view()
        self.log(f"🔒 Locked segment #{self.current_segment.id}")
    
    def unlock_current_segment(self):
        """Unlock the currently selected segment"""
        if not self.current_segment:
            self.log("⚠️ No segment selected")
            return
        
        if not self.current_segment.locked:
            messagebox.showinfo("Already Unlocked", "This segment is already unlocked.")
            return
        
        self.current_segment.locked = False
        self.current_segment.modified = True
        self.modified = True
        self.refresh_current_view()
        self.log(f"🔓 Unlocked segment #{self.current_segment.id}")
    
    def refresh_current_view(self):
        """Refresh the current view after bulk operations"""
        if self.layout_mode == LayoutMode.GRID:
            self.load_segments_to_grid()
        elif self.layout_mode == LayoutMode.SPLIT:
            self.load_segments_to_list()
        elif self.layout_mode == LayoutMode.DOCUMENT:
            self.load_segments_to_document_view()
    
    # ============================================================================
    # END BULK OPERATIONS
    # ============================================================================
    
    def save_segment_and_next(self):
        """Save current segment with 'translated' status and move to next untranslated segment"""
        if not self.current_segment:
            return
        
        # Set status to translated before saving
        self.status_var.set('translated')
        self.save_current_segment()
        
        # Find next untranslated segment
        current_id = self.current_segment.id
        next_untranslated = None
        
        for seg in self.segments:
            if seg.id > current_id and seg.status == 'untranslated':
                next_untranslated = seg
                break
        
        # If found, select it in the tree
        if next_untranslated:
            for item in self.tree.get_children():
                values = self.tree.item(item, 'values')
                if int(values[0]) == next_untranslated.id:
                    self.tree.selection_set(item)
                    self.tree.see(item)
                    break
        else:
            self.log("No more untranslated segments")
    
    def focus_target_editor(self):
        """Focus the target text editor"""
        self.target_text.focus_set()
    
    def update_progress(self):
        """Update progress label"""
        if not self.segments:
            self.progress_label.config(text="No document loaded")
            return
        
        total = len(self.segments)
        translated = sum(1 for s in self.segments if s.status in ['translated', 'approved'])
        percentage = (translated / total * 100) if total > 0 else 0
        
        self.progress_label.config(
            text=f"Progress: {translated}/{total} ({percentage:.1f}%) | "
                 f"{'Modified' if self.modified else 'Saved'}"
        )
        
        # Update context status
        self.update_context_status()
    
    def update_context_status(self):
        """Update document context status display"""
        if not hasattr(self, 'context_status_label'):
            return
        
        context_enabled = self.use_context_var.get()
        segment_count = len(self.segments)
        
        if segment_count > 0:
            context_text = self.get_full_document_context(include_translations=False)
            char_count = len(context_text)
            
            status = "Enabled ✓" if context_enabled else "Disabled ✗"
            color = '#1976D2' if context_enabled else '#F44336'
            
            self.context_status_label.config(
                text=f"Context: {status} | {segment_count} segments | {char_count:,} characters",
                fg=color
            )
        else:
            self.context_status_label.config(
                text="Context: No document loaded",
                fg='#666'
            )
    
    def save_project(self):
        """Save project to JSON"""
        if not self.project_file:
            self.save_project_as()
            return
        
        self.save_current_segment()
        
        try:
            data = {
                'version': '0.3.2',
                'created_at': datetime.now().isoformat(),
                'original_docx': self.original_docx,
                'segments': [s.to_dict() for s in self.segments],
                # Save filter preferences
                'filter_preferences': {
                    'mode': self.filter_mode,
                    'source_filter': self.filter_source_var.get() if hasattr(self, 'filter_source_var') else '',
                    'target_filter': self.filter_target_var.get() if hasattr(self, 'filter_target_var') else '',
                    'status_filter': self.filter_status_var.get() if hasattr(self, 'filter_status_var') else 'All',
                    'active': self.filter_active
                },
                # Save TM database (new multi-TM format)
                'tm_database': self.tm_database.to_dict(),
                # Legacy format for backwards compatibility
                'translation_memory': {
                    'entries': self.tm_database.project_tm.entries,
                    'fuzzy_threshold': self.tm_database.fuzzy_threshold
                },
                # Save LLM settings
                'llm_settings': {
                    'provider': self.current_llm_provider,
                    'model': self.current_llm_model,
                    'source_language': self.source_language,
                    'target_language': self.target_language,
                    'custom_prompt': self.current_translate_prompt
                },
                # Save preferences
                'preferences': {
                    'chunk_size': self.chunk_size_var.get(),
                    'use_context': self.use_context_var.get(),
                    'check_tm': self.check_tm_var.get(),
                    'surrounding_segments': self.surrounding_segments_var.get()
                }
            }
            
            with open(self.project_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            
            self.modified = False
            self.log(f"✓ Project saved: {os.path.basename(self.project_file)}")
            self.update_progress()
            
        except Exception as e:
            messagebox.showerror("Save Error", f"Failed to save project:\n{str(e)}")
    
    def save_project_as(self):
        """Save project as new file"""
        file_path = filedialog.asksaveasfilename(
            title="Save Project As",
            defaultextension=".json",
            filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")]
        )
        
        if file_path:
            self.project_file = file_path
            self.save_project()
    
    def load_project(self):
        """Load project from JSON"""
        if self.modified:
            response = messagebox.askyesnocancel(
                "Unsaved Changes",
                "Save current project before loading?"
            )
            if response is None:
                return
            elif response is True:
                self.save_project()
        
        file_path = filedialog.askopenfilename(
            title="Open Project",
            filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return
        
        self.load_project_from_path(file_path)
    
    def load_project_from_path(self, file_path):
        """Load project from a given file path"""
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Load segments
            self.segments = [Segment.from_dict(s) for s in data['segments']]
            self.original_docx = data.get('original_docx')
            self.project_file = file_path
            
            # Load TM database (try new format first, fall back to legacy)
            if 'tm_database' in data:
                # New multi-TM format
                self.tm_database = TMDatabase.from_dict(data['tm_database'])
                self.tm_agent.tm_database = self.tm_database  # Update legacy wrapper
                enabled_count = self.tm_database.get_entry_count(enabled_only=True)
                total_count = self.tm_database.get_entry_count(enabled_only=False)
                self.log(f"✓ Loaded TM database: {total_count} total entries ({enabled_count} in active TMs)")
            elif 'translation_memory' in data:
                # Legacy single-TM format - migrate to new format
                tm_data = data['translation_memory']
                self.tm_database = TMDatabase()
                self.tm_database.project_tm.entries = tm_data.get('entries', {})
                self.tm_database.fuzzy_threshold = tm_data.get('fuzzy_threshold', 0.75)
                self.tm_agent.tm_database = self.tm_database
                self.log(f"✓ Migrated legacy TM: {len(self.tm_database.project_tm.entries)} entries to Project TM")
            else:
                # No TM data
                self.tm_database = TMDatabase()
                self.tm_agent.tm_database = self.tm_database
                self.log("ℹ No TM data in project")
            
            # Load LLM settings if present
            if 'llm_settings' in data:
                llm_settings = data['llm_settings']
                self.current_llm_provider = llm_settings.get('provider', 'openai')
                self.current_llm_model = llm_settings.get('model', 'gpt-4o')
                self.source_language = llm_settings.get('source_language', 'English')
                self.target_language = llm_settings.get('target_language', 'Dutch')
                self.current_translate_prompt = llm_settings.get('custom_prompt', self.default_translate_prompt)
                self.log(f"✓ Loaded LLM settings: {self.current_llm_provider}/{self.current_llm_model}")
            
            # Load preferences if present
            if 'preferences' in data:
                prefs = data['preferences']
                self.chunk_size_var.set(prefs.get('chunk_size', '100'))
                self.use_context_var.set(prefs.get('use_context', False))
                self.check_tm_var.set(prefs.get('check_tm', True))
                self.surrounding_segments_var.set(prefs.get('surrounding_segments', '5'))
                self.log(f"✓ Loaded preferences: batch_size={self.chunk_size_var.get()}, context={self.surrounding_segments_var.get()}")
            
            # Load filter preferences if they exist
            if 'filter_preferences' in data:
                prefs = data['filter_preferences']
                if hasattr(self, 'filter_source_var'):
                    self.filter_source_var.set(prefs.get('source_filter', ''))
                if hasattr(self, 'filter_target_var'):
                    self.filter_target_var.set(prefs.get('target_filter', ''))
                if hasattr(self, 'filter_status_var'):
                    self.filter_status_var.set(prefs.get('status_filter', 'All'))
                self.filter_mode = prefs.get('mode', 'filter')
                self.filter_active = prefs.get('active', False)
                
                # Update button states to match loaded mode
                if hasattr(self, 'filter_mode_btn') and hasattr(self, 'highlight_mode_btn'):
                    if self.filter_mode == 'filter':
                        self.filter_mode_btn.config(bg='#4CAF50', fg='white', 
                                                   font=('Segoe UI', 9, 'bold'), relief='sunken')
                        self.highlight_mode_btn.config(bg='#e0e0e0', fg='#666',
                                                      font=('Segoe UI', 9), relief='raised')
                    else:
                        self.filter_mode_btn.config(bg='#e0e0e0', fg='#666',
                                                   font=('Segoe UI', 9), relief='raised')
                        self.highlight_mode_btn.config(bg='#FFA500', fg='white',
                                                      font=('Segoe UI', 9, 'bold'), relief='sunken')
                
                # Apply the filters if they were active
                if self.filter_active:
                    self.apply_filters()
            
            # If original DOCX exists, load it for export
            if self.original_docx and os.path.exists(self.original_docx):
                self.docx_handler.import_docx(self.original_docx)
            
            # Load to grid
            self.load_segments_to_grid()
            
            # Switch from Start Screen to Grid View if needed
            if hasattr(self, 'start_paned'):
                self.switch_from_start_to_grid()
            else:
                # Refresh display if already in grid view
                self.refresh_display()
            
            self.modified = False
            self.log(f"✓ Project loaded: {os.path.basename(file_path)}")
            self.update_progress()
            
        except Exception as e:
            messagebox.showerror("Open Error", f"Failed to open project:\n{str(e)}")
    
    def close_project(self):
        """Close current project"""
        if not self.segments:
            messagebox.showinfo("No Project", "No project is currently open")
            return
        
        if self.modified:
            response = messagebox.askyesnocancel(
                "Unsaved Changes",
                "Save current project before closing?"
            )
            if response is None:
                return
            elif response is True:
                self.save_project()
        
        # Clear all data
        self.segments = []
        self.current_segment_index = None
        self.project_file = None
        self.original_docx = None
        self.modified = False
        
        # Clear the current view
        if self.current_layout_mode == LayoutMode.GRID:
            for item in self.tree.get_children():
                self.tree.delete(item)
        elif self.current_layout_mode == LayoutMode.SPLIT:
            for item in self.tree.get_children():
                self.tree.delete(item)
            self.source_text.config(state='normal')
            self.source_text.delete('1.0', 'end')
            self.source_text.config(state='disabled')
            self.target_text.delete('1.0', 'end')
        elif self.current_layout_mode == LayoutMode.DOCUMENT:
            # Clear document canvas
            for widget in self.doc_inner_frame.winfo_children():
                widget.destroy()
            self.doc_segment_widgets.clear()
            self.doc_current_segment = None
        
        self.log("✓ Project closed")
        self.update_progress()
    
    def export_docx(self):
        """Export to translated DOCX"""
        if not self.segments:
            messagebox.showwarning("No Data", "No segments to export")
            return
        
        if not self.original_docx or not os.path.exists(self.original_docx):
            messagebox.showerror("No Original",
                               "Original DOCX file not found. Cannot export.")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="Export to DOCX",
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return
        
        try:
            self.save_current_segment()
            
            # Prepare segments for export
            seg_dicts = [s.to_dict() for s in self.segments]
            
            # Export
            self.docx_handler.export_docx(seg_dicts, file_path)
            
            self.log(f"✓ Exported to DOCX: {os.path.basename(file_path)}")
            
            # Perform auto-exports if enabled
            self.perform_auto_exports(file_path)
            
            messagebox.showinfo("Export Complete",
                              f"Document exported successfully to:\n{file_path}")
            
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export DOCX:\n{str(e)}")
            self.log(f"✗ Export failed: {str(e)}")
    
    def export_bilingual_docx(self):
        """Export to bilingual DOCX (table format)"""
        if not self.segments:
            messagebox.showwarning("No Data", "No segments to export")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="Export to Bilingual DOCX",
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return
        
        try:
            self.save_current_segment()
            seg_dicts = [s.to_dict() for s in self.segments]
            self.docx_handler.export_bilingual_docx(seg_dicts, file_path)
            
            self.log(f"✓ Exported bilingual DOCX: {os.path.basename(file_path)}")
            
            # Perform auto-exports if enabled
            self.perform_auto_exports(file_path)
            
            messagebox.showinfo("Export Complete", "Bilingual document exported successfully")
            
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export:\n{str(e)}")
    
    def export_tsv(self):
        """Export to TSV (tab-separated values)"""
        if not self.segments:
            messagebox.showwarning("No Data", "No segments to export")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="Export to TSV",
            defaultextension=".tsv",
            filetypes=[("TSV Files", "*.tsv"), ("Text Files", "*.txt"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return
        
        try:
            self.save_current_segment()
            
            with open(file_path, 'w', encoding='utf-8') as f:
                # Header
                f.write("ID\tStatus\tSource\tTarget\tParagraph\tNotes\n")
                
                # Data
                for seg in self.segments:
                    f.write(f"{seg.id}\t{seg.status}\t{seg.source}\t{seg.target}\t"
                           f"{seg.paragraph_id}\t{seg.notes}\n")
            
            self.log(f"✓ Exported to TSV: {os.path.basename(file_path)}")
            messagebox.showinfo("Export Complete", "TSV file exported successfully")
            
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export TSV:\n{str(e)}")
    
    def export_txt_bilingual(self):
        """Export to bilingual TXT (memoQ/CAT tool format)"""
        if not self.segments:
            messagebox.showwarning("No Data", "No segments to export")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="Export to Bilingual TXT",
            defaultextension=".txt",
            filetypes=[("Text Files", "*.txt"), ("TSV Files", "*.tsv"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return
        
        try:
            self.save_current_segment()
            
            with open(file_path, 'w', encoding='utf-8') as f:
                # Write tab-delimited format: ID, Source, Target
                for seg in self.segments:
                    f.write(f"{seg.id}\t{seg.source}\t{seg.target}\n")
            
            self.log(f"✓ Exported to bilingual TXT: {os.path.basename(file_path)}")
            messagebox.showinfo("Export Complete", 
                              f"Bilingual TXT file exported successfully!\n\n"
                              f"Format: Tab-delimited (ID, Source, Target)\n"
                              f"Segments: {len(self.segments)}\n"
                              f"File: {os.path.basename(file_path)}")
            
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export bilingual TXT:\n{str(e)}")
    
    def export_tmx(self):
        """Export translation memory to TMX format"""
        if not self.segments:
            messagebox.showwarning("No Data", "No segments to export")
            return
        
        # Count translated segments
        translated_segments = [seg for seg in self.segments if seg.target.strip() and seg.status != "untranslated"]
        
        if not translated_segments:
            messagebox.showwarning("No Translations", "No translated segments to export to TMX")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="Export to TMX",
            defaultextension=".tmx",
            filetypes=[("TMX Files", "*.tmx"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return
        
        try:
            self.save_current_segment()
            
            # Prepare source and target lists
            source_segments = [seg.source for seg in translated_segments]
            target_segments = [seg.target for seg in translated_segments]
            
            # Generate TMX
            tmx_tree = self.tmx_generator.generate_tmx(
                source_segments, 
                target_segments,
                self.source_language,
                self.target_language
            )
            
            # Save TMX
            if self.tmx_generator.save_tmx(tmx_tree, file_path):
                self.log(f"✓ Exported {len(translated_segments)} translation units to TMX: {os.path.basename(file_path)}")
                messagebox.showinfo("Export Complete", 
                                  f"TMX file exported successfully.\n\n"
                                  f"{len(translated_segments)} translation units saved.")
            else:
                messagebox.showerror("Export Error", "Failed to save TMX file")
            
        except Exception as e:
            self.log(f"✗ TMX export failed: {str(e)}")
            messagebox.showerror("Export Error", f"Failed to export TMX:\n{str(e)}")
    
    def perform_auto_exports(self, primary_export_path):
        """
        Automatically export additional formats based on user settings.
        Called after any primary export operation.
        
        Args:
            primary_export_path: Path to the primary export file (used to determine directory)
        """
        if not self.segments:
            return
        
        # Get directory from primary export
        export_dir = os.path.dirname(primary_export_path)
        base_name = os.path.splitext(os.path.basename(primary_export_path))[0]
        
        auto_exports_performed = []
        
        try:
            # Session Report (Markdown)
            if self.auto_export_session_md_var.get():
                md_path = os.path.join(export_dir, f"{base_name}_SessionReport.md")
                self._auto_export_session_report_md(md_path)
                auto_exports_performed.append(f"📄 {os.path.basename(md_path)}")
            
            # Session Report (HTML)
            if self.auto_export_session_html_var.get():
                html_path = os.path.join(export_dir, f"{base_name}_SessionReport.html")
                self._auto_export_session_report_html(html_path)
                auto_exports_performed.append(f"🌐 {os.path.basename(html_path)}")
            
            # TMX (Translation Memory)
            if self.auto_export_tmx_var.get():
                tmx_path = os.path.join(export_dir, f"{base_name}.tmx")
                self._auto_export_tmx(tmx_path)
                auto_exports_performed.append(f"💾 {os.path.basename(tmx_path)}")
            
            # TSV (Tab-separated)
            if self.auto_export_tsv_var.get():
                tsv_path = os.path.join(export_dir, f"{base_name}.tsv")
                self._auto_export_tsv(tsv_path)
                auto_exports_performed.append(f"📊 {os.path.basename(tsv_path)}")
            
            # Bilingual TXT
            if self.auto_export_bilingual_txt_var.get():
                txt_path = os.path.join(export_dir, f"{base_name}_bilingual.txt")
                self._auto_export_bilingual_txt(txt_path)
                auto_exports_performed.append(f"📝 {os.path.basename(txt_path)}")
            
            # XLIFF
            if self.auto_export_xliff_var.get():
                xliff_path = os.path.join(export_dir, f"{base_name}.xliff")
                self._auto_export_xliff(xliff_path)
                auto_exports_performed.append(f"🔄 {os.path.basename(xliff_path)}")
            
            # Excel Bilingual
            if self.auto_export_excel_var.get():
                excel_path = os.path.join(export_dir, f"{base_name}_bilingual.xlsx")
                self._auto_export_excel(excel_path)
                auto_exports_performed.append(f"📗 {os.path.basename(excel_path)}")
            
            # Log auto-exports
            if auto_exports_performed:
                self.log(f"✓ Auto-exported {len(auto_exports_performed)} additional format(s):")
                for export in auto_exports_performed:
                    self.log(f"  • {export}")
        
        except Exception as e:
            self.log(f"⚠️ Some auto-exports failed: {str(e)}")
    
    def _auto_export_session_report_md(self, file_path):
        """Generate markdown session report (auto-export version)"""
        import datetime
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Calculate statistics
        total_segments = len(self.segments)
        translated = sum(1 for seg in self.segments if seg.target.strip() and seg.status != "untranslated")
        untranslated = total_segments - translated
        draft = sum(1 for seg in self.segments if seg.status == "draft")
        approved = sum(1 for seg in self.segments if seg.status == "approved")
        
        current_mode = "batch_bilingual" if hasattr(self, 'original_txt') else "batch_docx"
        current_prompt = self.get_context_aware_prompt(current_mode)
        
        report = f"""# Supervertaler CAT Editor Session Report

## Session Information
- **Date & Time**: {timestamp}
- **Supervertaler Version**: {APP_VERSION}
- **AI Provider**: {self.current_llm_provider}
- **AI Model**: {self.current_llm_model}

## Project Statistics
- **Total Segments**: {total_segments}
- **Translated**: {translated} ({translated/total_segments*100:.1f}%)
- **Untranslated**: {untranslated}
- **Draft**: {draft}
- **Approved**: {approved}

## Language Settings
- **Source**: {self.source_language}
- **Target**: {self.target_language}

## Active System Prompt
```
{current_prompt}
```

---
*Auto-generated by Supervertaler v{APP_VERSION}*
"""
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(report)
    
    def _auto_export_session_report_html(self, file_path):
        """Generate HTML session report (auto-export version)"""
        import datetime
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # First generate markdown
        md_path = file_path.replace('.html', '_temp.md')
        self._auto_export_session_report_md(md_path)
        
        # Read markdown and convert to HTML
        with open(md_path, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        html_content = self.markdown_to_html(markdown_content, f"Session Report - {timestamp}")
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        # Clean up temp markdown
        if os.path.exists(md_path):
            os.remove(md_path)
    
    def _auto_export_tmx(self, file_path):
        """Export TMX silently (auto-export version)"""
        translated_segments = [seg for seg in self.segments if seg.target.strip() and seg.status != "untranslated"]
        
        if not translated_segments:
            return
        
        source_segments = [seg.source for seg in translated_segments]
        target_segments = [seg.target for seg in translated_segments]
        
        tmx_tree = self.tmx_generator.generate_tmx(
            source_segments, 
            target_segments,
            self.source_language,
            self.target_language
        )
        
        self.tmx_generator.save_tmx(tmx_tree, file_path)
    
    def _auto_export_tsv(self, file_path):
        """Export TSV silently (auto-export version)"""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write("ID\tStatus\tSource\tTarget\tParagraph\tNotes\n")
            for seg in self.segments:
                f.write(f"{seg.id}\t{seg.status}\t{seg.source}\t{seg.target}\t"
                       f"{seg.paragraph_id}\t{seg.notes}\n")
    
    def _auto_export_bilingual_txt(self, file_path):
        """Export bilingual TXT silently (auto-export version)"""
        with open(file_path, 'w', encoding='utf-8') as f:
            for seg in self.segments:
                f.write(f"{seg.source}\n")
                f.write(f"{seg.target}\n")
                f.write("\n")
    
    def _auto_export_xliff(self, file_path):
        """Export XLIFF silently (auto-export version)"""
        import xml.etree.ElementTree as ET
        
        # Create XLIFF structure
        xliff = ET.Element('xliff', version='1.2', xmlns='urn:oasis:names:tc:xliff:document:1.2')
        file_elem = ET.SubElement(xliff, 'file', {
            'source-language': self.source_language.lower().replace(' ', '-'),
            'target-language': self.target_language.lower().replace(' ', '-'),
            'datatype': 'plaintext',
            'original': getattr(self, 'original_docx', 'document')
        })
        body = ET.SubElement(file_elem, 'body')
        
        for seg in self.segments:
            trans_unit = ET.SubElement(body, 'trans-unit', id=str(seg.id))
            source_elem = ET.SubElement(trans_unit, 'source')
            source_elem.text = seg.source
            target_elem = ET.SubElement(trans_unit, 'target')
            target_elem.text = seg.target
        
        tree = ET.ElementTree(xliff)
        tree.write(file_path, encoding='utf-8', xml_declaration=True)
    
    def _auto_export_excel(self, file_path):
        """Export Excel bilingual (auto-export version)"""
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment
            
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Translation"
            
            # Header
            headers = ['ID', 'Status', 'Source', 'Target', 'Notes']
            ws.append(headers)
            
            # Style header
            header_fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
            header_font = Font(bold=True, color='FFFFFF')
            for cell in ws[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center')
            
            # Add data
            for seg in self.segments:
                ws.append([seg.id, seg.status, seg.source, seg.target, seg.notes])
            
            # Adjust column widths
            ws.column_dimensions['A'].width = 8
            ws.column_dimensions['B'].width = 12
            ws.column_dimensions['C'].width = 50
            ws.column_dimensions['D'].width = 50
            ws.column_dimensions['E'].width = 30
            
            wb.save(file_path)
        
        except ImportError:
            # If openpyxl not available, skip Excel export
            self.log("⚠️ Excel export skipped (openpyxl not installed)")
    
    def import_cafetran_bilingual(self):
        """
        Import CafeTran bilingual DOCX file.
        Extracts source segments from the CafeTran bilingual table (with pipe symbols)
        and loads them into the CAT editor.
        
        CafeTran Bilingual Structure:
        - Row 0: Header (ID, filename, filename, Notes, *)
        - Row 1+: Segment data with pipe symbols (|) marking formatted text
        """
        file_path = filedialog.askopenfilename(
            title="Select CafeTran Bilingual DOCX File",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return  # User cancelled
        
        self.import_cafetran_bilingual_from_path(file_path)
    
    def import_cafetran_bilingual_from_path(self, file_path):
        """Import CafeTran bilingual DOCX from a given path"""
        try:
            # Import the CafeTran handler
            import sys
            # Add modules directory to path
            modules_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'modules')
            if modules_dir not in sys.path:
                sys.path.insert(0, modules_dir)
            
            from cafetran_docx_handler import CafeTranDOCXHandler
            
            # Check if this is actually a CafeTran bilingual DOCX
            if not CafeTranDOCXHandler.is_cafetran_bilingual_docx(file_path):
                messagebox.showerror(
                    "Invalid Format",
                    "This does not appear to be a CafeTran bilingual DOCX file.\n\n"
                    "Expected format: Table with columns 'ID', source, target, etc."
                )
                return
            
            # Load the CafeTran bilingual DOCX
            handler = CafeTranDOCXHandler()
            if not handler.load(file_path):
                messagebox.showerror("Error", "Failed to load CafeTran bilingual DOCX file.")
                return
            
            # Extract source segments
            segments = handler.extract_source_segments()
            
            if not segments:
                messagebox.showwarning("Warning", "No translatable segments found in the CafeTran bilingual file.")
                return
            
            # Clear current segments and load new ones
            self.segments.clear()
            
            # Smart paragraph detection
            current_paragraph_id = 0
            
            for i, cafetran_seg in enumerate(segments):
                # Detect if this should be a new paragraph
                should_start_new_paragraph = False
                
                if i == 0:
                    # First segment always starts a paragraph
                    should_start_new_paragraph = True
                else:
                    prev_source = segments[i-1].source_with_pipes.strip('|').strip()
                    curr_source = cafetran_seg.source_with_pipes.strip('|').strip()
                    
                    # Heuristics for new paragraph:
                    # 1. Current segment looks like a heading (short, all caps or title case, no period)
                    is_heading = (len(curr_source) < 50 and 
                                 curr_source.isupper() and 
                                 not curr_source.endswith('.'))
                    
                    # 2. Previous segment was very short (likely a heading)
                    prev_was_heading = len(prev_source) < 50 and not prev_source.endswith('.')
                    
                    # 3. Previous ended with period AND current starts with capital (but not mid-paragraph continuation)
                    natural_break = (prev_source.endswith('.') and 
                                   curr_source[0].isupper() and 
                                   (prev_was_heading or is_heading))
                    
                    should_start_new_paragraph = is_heading or prev_was_heading or natural_break
                
                # Increment paragraph ID if starting new paragraph
                if should_start_new_paragraph:
                    current_paragraph_id += 1
                
                # Create Segment object with source text (keeping pipe symbols)
                # Segment(seg_id, source, paragraph_id, is_table_cell, table_info, style, document_position)
                seg = Segment(
                    seg_id=i + 1,  # Sequential ID
                    source=cafetran_seg.source_with_pipes,  # Keep pipes for AI to see
                    paragraph_id=current_paragraph_id,  # Smart paragraph grouping
                    is_table_cell=False,
                    table_info=None,
                    style="Normal",
                    document_position=i
                )
                # Set target and status if translation exists
                if cafetran_seg.target_with_pipes:
                    seg.target = cafetran_seg.target_with_pipes
                    seg.status = "confirmed"
                
                self.segments.append(seg)
            
            # Store the original CafeTran file info
            self.cafetran_source_file = file_path
            self.cafetran_segments = segments
            
            # Update UI - load into grid view
            self.current_segment_index = 0
            self.load_segments_to_grid()
            self.update_progress()
            self.modified = False
            
            # Log success
            self.log(f"✓ Imported {len(segments)} segments from CafeTran bilingual DOCX")
            self.log(f"📌 Pipe symbols (|) mark formatted text - AI will preserve them during translation")
            
            messagebox.showinfo(
                "Success", 
                f"✓ Imported {len(segments)} segment(s) from CafeTran bilingual DOCX!\n\n"
                f"Source file: {os.path.basename(file_path)}\n\n"
                f"📌 Pipe symbols (|) mark formatted text (bold/italic/underline).\n"
                f"   The AI will preserve these symbols in translations.\n\n"
                f"Use 'Export to CafeTran DOCX' from the File menu when complete."
            )
            
        except ImportError as e:
            messagebox.showerror(
                "Module Error",
                f"Could not load CafeTran handler module.\n\n{str(e)}"
            )
        except Exception as e:
            messagebox.showerror("Import Error", f"Failed to import CafeTran bilingual DOCX:\n\n{str(e)}")
            self.log(f"✗ CafeTran import failed: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def export_cafetran_bilingual(self):
        """
        Export to CafeTran bilingual DOCX format.
        Creates a new bilingual DOCX file with translations and pipe symbols formatted as bold+red.
        """
        # Check if we have segments
        if not self.segments:
            messagebox.showwarning("No Data", "No segments to export")
            return
        
        # Check if a CafeTran source file was imported
        if not hasattr(self, 'cafetran_source_file'):
            messagebox.showwarning(
                "No CafeTran Source",
                "No CafeTran bilingual file was imported.\n\n"
                "This feature is only available after importing a CafeTran bilingual DOCX file.\n\n"
                "Use 'Export to Bilingual DOCX' for standard workflows."
            )
            return
        
        try:
            # Import the CafeTran handler
            import sys
            # Add modules directory to path
            modules_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'modules')
            if modules_dir not in sys.path:
                sys.path.insert(0, modules_dir)
            
            from cafetran_docx_handler import CafeTranDOCXHandler
            
            # Save current segment
            self.save_current_segment()
            
            # Collect translations (target texts which should already have pipes from AI)
            translations = [seg.target for seg in self.segments]
            
            if not translations or all(not t.strip() for t in translations):
                messagebox.showwarning("Warning", "No translations found to export.")
                return
            
            # Load the original CafeTran bilingual DOCX
            handler = CafeTranDOCXHandler()
            if not handler.load(self.cafetran_source_file):
                messagebox.showerror("Error", "Failed to load the original CafeTran bilingual file.")
                return
            
            # Extract source segments to ensure we have the structure
            segments = handler.extract_source_segments()
            
            # Update target segments with translations
            # The translations should already have pipes placed by the AI
            if not handler.update_target_segments(translations):
                messagebox.showerror("Error", "Failed to update target segments with translations.")
                return
            
            # Prompt user to save the updated CafeTran bilingual file
            save_path = filedialog.asksaveasfilename(
                title="Save CafeTran Bilingual DOCX",
                defaultextension=".docx",
                initialfile=os.path.basename(self.cafetran_source_file).replace('.docx', '_translated.docx'),
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
            )
            
            if save_path:
                if handler.save(save_path):
                    success_msg = (
                        f"✓ Successfully exported {len(translations)} translation(s) to CafeTran bilingual DOCX!\n\n"
                        f"File saved: {os.path.basename(save_path)}\n\n"
                        f"✓ Translations inserted in target column\n"
                        f"✓ Pipe symbols (|) formatted as BOLD + RED for easy visibility\n"
                        f"✓ Formatting markers preserved by AI at corresponding locations\n"
                        f"✓ Segment IDs maintained\n"
                        f"✓ Table structure preserved\n\n"
                        f"You can now import this file back into CafeTran.\n"
                        f"The red pipe symbols mark formatted text locations."
                    )
                    
                    self.log(f"✓ Exported {len(translations)} translations to CafeTran bilingual DOCX: {os.path.basename(save_path)}")
                    
                    # Perform auto-exports if enabled
                    self.perform_auto_exports(save_path)
                    
                    messagebox.showinfo("Success", success_msg)
                else:
                    messagebox.showerror("Save Error", "Failed to save the CafeTran bilingual DOCX file.")
            
        except ImportError as e:
            messagebox.showerror(
                "Module Error",
                f"Could not load CafeTran handler module.\n\n{str(e)}"
            )
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export CafeTran bilingual DOCX:\n\n{str(e)}")
            self.log(f"✗ CafeTran export failed: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def apply_formatting_to_cell(self, cell, text, formatting_info):
        """
        Apply formatting (bold, italic, underline) to a DOCX cell based on formatting_info.
        Uses PROGRAMMATIC approach to preserve memoQ formatting.
        
        Args:
            cell: DOCX table cell object
            text: Plain text to write
            formatting_info: List of dicts with 'text', 'bold', 'italic', 'underline' keys
        
        Strategy:
        - If the entire source segment (or most of it) had one consistent formatting,
          apply that formatting to the entire target
        - Otherwise, try to match CAT tags and preserve formatting around them
        """
        # Clear existing content
        cell.text = ''
        
        # If no formatting info, just write plain text
        if not formatting_info:
            cell.text = text
            return
        
        # Calculate total length and formatted lengths
        total_chars = sum(len(run_info['text']) for run_info in formatting_info)
        bold_chars = sum(len(run_info['text']) for run_info in formatting_info if run_info['bold'])
        italic_chars = sum(len(run_info['text']) for run_info in formatting_info if run_info['italic'])
        underline_chars = sum(len(run_info['text']) for run_info in formatting_info if run_info['underline'])
        
        # Determine if the whole segment should have formatting
        # Lower threshold to 60% to catch more cases
        threshold = 0.6
        apply_bold = (bold_chars / total_chars) > threshold if total_chars > 0 else False
        apply_italic = (italic_chars / total_chars) > threshold if total_chars > 0 else False
        apply_underline = (underline_chars / total_chars) > threshold if total_chars > 0 else False
        
        # Also check if formatting is at the beginning (first 50 chars or 30% of text)
        # This handles cases like "Biagio Pagano (born...)" where just the name is bold
        beginning_threshold = min(50, int(total_chars * 0.3))
        beginning_bold = any(run_info['bold'] for run_info in formatting_info[:2])  # First 2 runs
        beginning_italic = any(run_info['italic'] for run_info in formatting_info[:2])
        beginning_underline = any(run_info['underline'] for run_info in formatting_info[:2])
        
        # Write text with formatting
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        
        if apply_bold or apply_italic or apply_underline:
            # Apply formatting to entire target text (majority formatted)
            run = paragraph.add_run(text)
            run.bold = apply_bold
            run.italic = apply_italic
            run.underline = apply_underline
        elif beginning_bold or beginning_italic or beginning_underline:
            # Partial formatting at beginning - try to apply to first word(s)
            # Extract first 1-2 words and format them
            import re
            words = text.split(None, 2)  # Split into max 3 parts
            if len(words) >= 2:
                # Format first 1-2 words
                first_part = ' '.join(words[:2])
                rest = ' ' + words[2] if len(words) > 2 else ''
                
                run = paragraph.add_run(first_part)
                run.bold = beginning_bold
                run.italic = beginning_italic
                run.underline = beginning_underline
                
                if rest:
                    paragraph.add_run(rest)
            else:
                # Short text, format all
                run = paragraph.add_run(text)
                run.bold = beginning_bold
                run.italic = beginning_italic
                run.underline = beginning_underline
        else:
            # Partial formatting - try to preserve it intelligently
            # This handles cases like CAT tags within formatted text
            
            # Build a simple word-based mapping
            import re
            formatted_words = {}
            for run_info in formatting_info:
                run_text = run_info['text'].strip()
                if run_text and (run_info['bold'] or run_info['italic'] or run_info['underline']):
                    # Store formatting for exact text matches (useful for CAT tags, URLs, etc.)
                    formatted_words[run_text] = run_info
            
            # Try to apply formatting to matching parts
            if formatted_words:
                # Split text but try to preserve CAT tags and special patterns
                words = re.findall(r'\[[0-9]+\}|\{[0-9]+\]|\S+|\s+', text)
                
                for word in words:
                    if word in formatted_words:
                        # Exact match (CAT tag, URL, etc.)
                        run = paragraph.add_run(word)
                        fmt = formatted_words[word]
                        run.bold = fmt['bold']
                        run.italic = fmt['italic']
                        run.underline = fmt['underline']
                    else:
                        # No formatting
                        paragraph.add_run(word)
            else:
                # No specific word matches - just write plain text
                paragraph.add_run(text)
    
    def import_trados_bilingual(self):
        """Import Trados Studio bilingual DOCX file."""
        filepath = filedialog.askopenfilename(
            title="Select Trados Studio Bilingual DOCX",
            filetypes=[
                ("Word Documents", "*.docx"),
                ("All files", "*.*")
            ]
        )
        
        if filepath:
            self.import_trados_bilingual_from_path(filepath)
    
    def import_trados_bilingual_from_path(self, filepath):
        """
        Import Trados Studio bilingual DOCX from file path.
        Used by both menu action and file explorer.
        """
        try:
            from modules.trados_docx_handler import TradosDOCXHandler
            
            # Verify this is a Trados bilingual file
            if not TradosDOCXHandler.is_trados_bilingual_docx(filepath):
                messagebox.showerror(
                    "Invalid File",
                    "The selected file does not appear to be a Trados Studio bilingual DOCX.\n\n"
                    "Expected format:\n"
                    "- Single table with 4 columns\n"
                    "- Header row with: Segment ID, Segment status, Source segment, Target segment"
                )
                return
            
            # Extract segments
            self.log(f"📥 Importing Trados Studio bilingual DOCX: {filepath}")
            extracted_segments = TradosDOCXHandler.extract_segments(filepath)
            
            if not extracted_segments:
                messagebox.showwarning("No Segments", "No segments found in the Trados file.")
                return
            
            self.log(f"📊 Extracted {len(extracted_segments)} segments from Trados file")
            
            # Clear existing segments
            self.segments = []
            self.segment_counter = 0
            
            # Convert extracted segments to internal format (Segment objects)
            for i, seg_data in enumerate(extracted_segments):
                try:
                    # Map Trados status to Supervertaler status
                    trados_status = seg_data['status']
                    supervertaler_status = TradosDOCXHandler.map_trados_status_to_supervertaler(trados_status)
                    
                    # Use formatted versions if available, otherwise plain text
                    source_text = seg_data.get('source_formatted', seg_data['source'])
                    target_text = seg_data.get('target_formatted', seg_data['target'])
                    
                    # Create Segment object (like memoQ/CafeTran import)
                    seg = Segment(
                        seg_id=seg_data['id'],
                        source=source_text,
                        paragraph_id=i,  # Each segment in own paragraph for now
                        is_table_cell=False,
                        table_info=None,
                        style="Normal",
                        document_position=i
                    )
                    
                    # Set target and status
                    seg.target = target_text
                    seg.status = supervertaler_status
                    
                    # Store Trados-specific metadata for export
                    seg.trados_id = seg_data.get('trados_id', '')
                    seg.trados_original_status = trados_status
                    
                    self.segments.append(seg)
                    self.segment_counter = max(self.segment_counter, seg_data['id'])
                    
                except KeyError as e:
                    self.log(f"⚠ Segment {i+1} missing key: {e}. Segment data: {seg_data}")
                    raise
            
            # Store the original Trados file info for export
            self.trados_source_file = filepath
            
            # Switch from Start Screen to Grid View if needed
            if hasattr(self, 'start_paned') and self.start_paned:
                self.switch_from_start_to_grid()
            
            # Update UI - load into grid view
            self.current_segment_index = 0
            self.load_segments_to_grid()
            self.update_progress()
            self.modified = False
            
            self.log(f"✅ Successfully imported {len(self.segments)} segments from Trados Studio file")
            messagebox.showinfo(
                "Import Successful",
                f"Imported {len(self.segments)} segments from Trados Studio bilingual DOCX.\n\n"
                f"File: {os.path.basename(filepath)}"
            )
            
        except ImportError:
            messagebox.showerror(
                "Module Error",
                "Could not import Trados DOCX handler module.\n\n"
                "Please ensure modules/trados_docx_handler.py exists."
            )
            self.log("❌ Failed to import TradosDOCXHandler module")
        except Exception as e:
            messagebox.showerror(
                "Import Error",
                f"Failed to import Trados Studio bilingual DOCX:\n\n{str(e)}"
            )
            self.log(f"❌ Trados import error: {str(e)}")
    
    def export_trados_bilingual(self):
        """Export segments to Trados Studio bilingual DOCX format."""
        if not self.segments:
            messagebox.showwarning("No Segments", "No segments to export.")
            return
        
        try:
            from modules.trados_docx_handler import TradosDOCXHandler
            
            # Check if we have an original Trados file to preserve format
            if not hasattr(self, 'trados_source_file') or not self.trados_source_file:
                messagebox.showerror(
                    "No Trados Source File",
                    "Cannot export to Trados format: No original Trados bilingual file loaded.\n\n"
                    "To export to Trados format, you must first import a Trados bilingual DOCX file."
                )
                self.log("❌ Trados export failed: No original Trados source file")
                return
            
            # File dialog
            filepath = filedialog.asksaveasfilename(
                title="Export to Trados Studio Bilingual DOCX",
                defaultextension=".docx",
                filetypes=[
                    ("Word Documents", "*.docx"),
                    ("All files", "*.*")
                ],
                initialfile="trados_bilingual_export.docx"
            )
            
            if not filepath:
                return
            
            # Prepare segments for export
            export_segments = []
            for seg in self.segments:
                # Use original Trados status if available, otherwise map from current status
                if hasattr(seg, 'trados_original_status') and seg.trados_original_status:
                    trados_status = seg.trados_original_status
                else:
                    # Map Supervertaler status back to Trados format
                    status_map = {
                        'untranslated': 'Not Translated (0%)',
                        'translated': 'Translated (100%)',
                        'approved': 'Approved Sign-off'
                    }
                    trados_status = status_map.get(seg.status, 'Not Translated (0%)')
                
                # Use original Trados UUID if available
                if hasattr(seg, 'trados_id') and seg.trados_id:
                    segment_id = seg.trados_id
                else:
                    segment_id = f"seg-{seg.id}"
                
                export_seg = {
                    'id': segment_id,
                    'status': trados_status,
                    'source': seg.source,
                    'target': seg.target
                }
                export_segments.append(export_seg)
            
            # Update the original Trados file (preserves exact format)
            TradosDOCXHandler.update_bilingual_docx(
                self.trados_source_file, 
                export_segments, 
                filepath
            )
            
            self.log(f"💾 Exported {len(export_segments)} segments to Trados Studio format: {filepath}")
            messagebox.showinfo(
                "Export Successful",
                f"Exported {len(export_segments)} segments to Trados Studio bilingual DOCX.\n\n"
                f"File: {os.path.basename(filepath)}\n\n"
                f"✓ Original format preserved for Trados re-import"
            )
            
        except ImportError:
            messagebox.showerror(
                "Module Error",
                "Could not import Trados DOCX handler module.\n\n"
                "Please ensure modules/trados_docx_handler.py exists."
            )
            self.log("❌ Failed to import TradosDOCXHandler module")
        except Exception as e:
            messagebox.showerror(
                "Export Error",
                f"Failed to export to Trados Studio bilingual DOCX:\n\n{str(e)}"
            )
            self.log(f"❌ Trados export error: {str(e)}")
    
    def import_memoq_bilingual(self):
        """
        Import memoQ bilingual DOCX file using PROGRAMMATIC formatting approach.
        Extracts source segments from the bilingual table and loads them into the CAT editor.
        
        memoQ Bilingual Structure:
        - Row 0: Header with project metadata
        - Row 1: Column headers (ID, Source Lang, Target Lang, Comment, Status)
        - Row 2+: Segment data
        """
        file_path = filedialog.askopenfilename(
            title="Select memoQ Bilingual DOCX File",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return  # User cancelled
        
        self.import_memoq_bilingual_from_path(file_path)
    
    def import_memoq_bilingual_from_path(self, file_path):
        """Import memoQ bilingual DOCX from a given path"""
        try:
            from docx import Document
            
            # Load the bilingual DOCX
            doc = Document(file_path)
            
            if not doc.tables:
                messagebox.showerror("Error", "No table found in the DOCX file.\n\nExpected memoQ bilingual format with a table.")
                return
            
            table = doc.tables[0]
            
            # Validate table structure (should have at least 3 rows: header, column names, data)
            if len(table.rows) < 3:
                messagebox.showerror("Error", f"Invalid table structure.\n\nExpected at least 3 rows, found {len(table.rows)}.")
                return
            
            # Extract source segments from column 1 (skipping header rows 0 and 1)
            # Also extract formatting information (bold, italic, underline)
            segments_data = []
            formatting_map = {}  # segment_index -> list of (text, bold, italic, underline)
            
            for row_idx in range(2, len(table.rows)):
                row = table.rows[row_idx]
                
                # Column 1 contains the source text
                if len(row.cells) >= 2:
                    cell = row.cells[1]
                    source_text = cell.text.strip()
                    
                    if source_text:  # Only add non-empty segments
                        segment_idx = len(segments_data)
                        segments_data.append(source_text)
                        
                        # Extract formatting from runs in this cell
                        formatting_info = []
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run_text = run.text
                                if run_text:  # Only store runs with actual text
                                    formatting_info.append({
                                        'text': run_text,
                                        'bold': run.bold == True,  # Explicit True check
                                        'italic': run.italic == True,  # Explicit True check
                                        'underline': run.underline == True  # Explicit True check (handles WD_UNDERLINE enum)
                                    })
                        
                        # Store formatting for this segment
                        if formatting_info:
                            formatting_map[segment_idx] = formatting_info
            
            if not segments_data:
                messagebox.showwarning("Warning", "No source segments found in the bilingual file.")
                return
            
            # Clear current segments and load new ones
            self.segments.clear()
            
            # Smart paragraph detection
            current_paragraph_id = 0
            
            for i, source_text in enumerate(segments_data):
                # Detect if this should be a new paragraph
                should_start_new_paragraph = False
                
                if i == 0:
                    # First segment always starts a paragraph
                    should_start_new_paragraph = True
                else:
                    prev_source = segments_data[i-1].strip()
                    curr_source = source_text.strip()
                    
                    # Heuristics for new paragraph:
                    # 1. Current segment looks like a heading (short, all caps or title case, no period)
                    is_heading = (len(curr_source) < 50 and 
                                 curr_source.isupper() and 
                                 not curr_source.endswith('.'))
                    
                    # 2. Previous segment was very short (likely a heading)
                    prev_was_heading = len(prev_source) < 50 and not prev_source.endswith('.')
                    
                    # 3. Previous ended with period AND current starts with capital (but not mid-paragraph continuation)
                    natural_break = (prev_source.endswith('.') and 
                                   curr_source[0].isupper() and 
                                   (prev_was_heading or is_heading))
                    
                    should_start_new_paragraph = is_heading or prev_was_heading or natural_break
                
                # Increment paragraph ID if starting new paragraph
                if should_start_new_paragraph:
                    current_paragraph_id += 1
                
                # Create Segment object
                # Each segment gets smart paragraph grouping for proper document view rendering
                seg = Segment(
                    seg_id=i + 1,
                    source=source_text,
                    paragraph_id=current_paragraph_id,  # Smart paragraph grouping
                    is_table_cell=False,
                    table_info=None,
                    style="Normal",
                    document_position=i
                )
                self.segments.append(seg)
            
            # Store the original memoQ file info for export
            self.memoq_source_file = file_path
            self.memoq_formatting_map = formatting_map
            
            # Update UI - load into grid view
            self.current_segment_index = 0
            self.load_segments_to_grid()
            self.update_progress()
            self.modified = False
            
            # Log formatting information
            formatted_segments = sum(1 for fmt_list in formatting_map.values() 
                                   if any(f['bold'] or f['italic'] or f['underline'] for f in fmt_list))
            
            self.log(f"✓ Imported {len(segments_data)} segments from memoQ bilingual DOCX")
            self.log(f"📌 Formatting detected in {formatted_segments} segment(s) - will be preserved programmatically")
            
            messagebox.showinfo(
                "Success", 
                f"✓ Imported {len(segments_data)} segment(s) from memoQ bilingual DOCX!\n\n"
                f"Source file: {os.path.basename(file_path)}\n\n"
                f"📌 Formatting preservation: PROGRAMMATIC approach\n"
                f"   Detected formatting in {formatted_segments} segment(s)\n"
                f"   (bold, italic, underline)\n\n"
                f"Use 'Export to memoQ DOCX' from the File menu when complete."
            )
            
        except ImportError:
            messagebox.showerror(
                "Missing Dependency",
                "The 'python-docx' library is required for bilingual DOCX import.\n\n"
                "Install it with: pip install python-docx"
            )
        except Exception as e:
            messagebox.showerror("Import Error", f"Failed to import memoQ bilingual DOCX:\n\n{str(e)}")
            self.log(f"✗ memoQ import failed: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def export_memoq_bilingual(self):
        """
        Export to memoQ bilingual DOCX format using PROGRAMMATIC formatting approach.
        Writes translations to the target column and preserves source formatting programmatically.
        """
        # Check if we have segments
        if not self.segments:
            messagebox.showwarning("No Data", "No segments to export")
            return
        
        # Check if a memoQ source file was imported
        if not hasattr(self, 'memoq_source_file'):
            messagebox.showwarning(
                "No memoQ Source",
                "No memoQ bilingual file was imported.\n\n"
                "This feature is only available after importing a memoQ bilingual DOCX file.\n\n"
                "Use 'Export to Bilingual DOCX' for standard workflows."
            )
            return
        
        try:
            from docx import Document
            import re
            
            # Save current segment
            self.save_current_segment()
            
            # Collect translations
            translations = [seg.target for seg in self.segments]
            
            if not translations or all(not t.strip() for t in translations):
                messagebox.showwarning("Warning", "No translations found to export.")
                return
            
            # Load the original bilingual DOCX
            doc = Document(self.memoq_source_file)
            table = doc.tables[0]
            
            # Write translations to target column (column 2) and update status
            segments_updated = 0
            segments_with_formatting = 0
            
            for i, translation in enumerate(translations):
                row_idx = i + 2  # Skip header rows (0 and 1)
                
                if row_idx < len(table.rows):
                    row = table.rows[row_idx]
                    
                    # Write translation to column 2 (target) with PROGRAMMATIC formatting
                    if len(row.cells) >= 3:
                        target_cell = row.cells[2]
                        
                        # Get formatting info for this segment (if available)
                        formatting_info = None
                        if hasattr(self, 'memoq_formatting_map') and i in self.memoq_formatting_map:
                            formatting_info = self.memoq_formatting_map[i]
                            if any(f['bold'] or f['italic'] or f['underline'] for f in formatting_info):
                                segments_with_formatting += 1
                        
                        # Apply formatting programmatically to the target cell
                        self.apply_formatting_to_cell(target_cell, translation, formatting_info)
                        segments_updated += 1
                    
                    # Update status to 'Confirmed' in column 4
                    if len(row.cells) >= 5:
                        row.cells[4].text = 'Confirmed'
            
            # Prompt user to save the updated bilingual file
            save_path = filedialog.asksaveasfilename(
                title="Save memoQ Bilingual DOCX",
                defaultextension=".docx",
                initialfile=os.path.basename(self.memoq_source_file).replace('.docx', '_translated.docx'),
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
            )
            
            if save_path:
                doc.save(save_path)
                
                # Build success message with formatting info
                success_msg = (
                    f"✓ Successfully exported {segments_updated} translation(s) to memoQ bilingual DOCX!\n\n"
                    f"File saved: {os.path.basename(save_path)}\n\n"
                )
                
                if segments_with_formatting > 0:
                    success_msg += (
                        f"✓ Formatting preserved in {segments_with_formatting} segment(s)\n"
                        f"  (bold, italic, underline - PROGRAMMATIC approach)\n\n"
                    )
                
                success_msg += (
                    f"✓ Status updated to 'Confirmed'\n"
                    f"✓ Table structure preserved\n\n"
                    f"You can now import this file back into memoQ."
                )
                
                self.log(f"✓ Exported {segments_updated} translations ({segments_with_formatting} with formatting) to memoQ bilingual DOCX: {os.path.basename(save_path)}")
                
                # Perform auto-exports if enabled
                self.perform_auto_exports(save_path)
                
                messagebox.showinfo("Success", success_msg)
            
        except ImportError:
            messagebox.showerror(
                "Missing Dependency",
                "The 'python-docx' library is required for bilingual DOCX export.\n\n"
                "Install it with: pip install python-docx"
            )
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export memoQ bilingual DOCX:\n\n{str(e)}")
            self.log(f"✗ memoQ export failed: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def markdown_to_html(self, markdown_text, title="Session Report"):
        """Convert markdown to styled HTML"""
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 40px auto;
            padding: 0 20px;
            color: #333;
            background-color: #f5f5f5;
        }}
        .container {{
            background-color: white;
            padding: 40px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
            margin-bottom: 30px;
        }}
        h2 {{
            color: #34495e;
            margin-top: 30px;
            margin-bottom: 15px;
            border-bottom: 2px solid #ecf0f1;
            padding-bottom: 8px;
        }}
        h3 {{
            color: #7f8c8d;
            margin-top: 20px;
            margin-bottom: 10px;
        }}
        h4 {{
            color: #95a5a6;
            margin-top: 15px;
            margin-bottom: 8px;
        }}
        ul, ol {{
            margin: 10px 0;
            padding-left: 30px;
        }}
        li {{
            margin: 5px 0;
        }}
        pre {{
            background-color: #f8f9fa;
            border: 1px solid #dee2e6;
            border-radius: 4px;
            padding: 15px;
            overflow-x: auto;
            font-family: 'Courier New', Courier, monospace;
            font-size: 14px;
        }}
        code {{
            background-color: #f8f9fa;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Courier New', Courier, monospace;
            font-size: 14px;
        }}
        pre code {{
            background-color: transparent;
            padding: 0;
        }}
        strong {{
            color: #2c3e50;
        }}
        .emoji {{
            font-size: 1.2em;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }}
        th, td {{
            border: 1px solid #dee2e6;
            padding: 10px;
            text-align: left;
        }}
        th {{
            background-color: #f8f9fa;
            font-weight: bold;
        }}
        hr {{
            border: none;
            border-top: 1px solid #dee2e6;
            margin: 30px 0;
        }}
        .footer {{
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #dee2e6;
            color: #7f8c8d;
            font-size: 14px;
            font-style: italic;
        }}
    </style>
</head>
<body>
    <div class="container">
"""
        
        # Simple markdown to HTML conversion
        lines = markdown_text.split('\n')
        in_code_block = False
        in_list = False
        html_content = []
        
        for line in lines:
            # Code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    html_content.append('</code></pre>')
                    in_code_block = False
                else:
                    html_content.append('<pre><code>')
                    in_code_block = True
                continue
            
            if in_code_block:
                html_content.append(line.replace('<', '&lt;').replace('>', '&gt;'))
                continue
            
            # Headers
            if line.startswith('# '):
                html_content.append(f'<h1>{line[2:]}</h1>')
            elif line.startswith('## '):
                html_content.append(f'<h2>{line[3:]}</h2>')
            elif line.startswith('### '):
                html_content.append(f'<h3>{line[4:]}</h3>')
            elif line.startswith('#### '):
                html_content.append(f'<h4>{line[5:]}</h4>')
            # Horizontal rule
            elif line.strip() == '---':
                html_content.append('<hr>')
            # Lists
            elif line.strip().startswith('- '):
                if not in_list:
                    html_content.append('<ul>')
                    in_list = True
                content = line.strip()[2:]
                # Handle bold
                content = content.replace('**', '<strong>', 1).replace('**', '</strong>', 1)
                html_content.append(f'<li>{content}</li>')
            else:
                if in_list and not line.strip().startswith('- '):
                    html_content.append('</ul>')
                    in_list = False
                
                if line.strip():
                    # Handle bold
                    line = line.replace('**', '<strong>', 1).replace('**', '</strong>', 1)
                    # Handle inline code
                    import re
                    line = re.sub(r'`([^`]+)`', r'<code>\1</code>', line)
                    html_content.append(f'<p>{line}</p>')
                else:
                    html_content.append('')
        
        if in_list:
            html_content.append('</ul>')
        if in_code_block:
            html_content.append('</code></pre>')
        
        html += '\n'.join(html_content)
        html += """
    </div>
</body>
</html>"""
        
        return html
    
    def generate_session_report(self):
        """Generate comprehensive markdown and HTML reports of current session"""
        if not self.segments:
            messagebox.showwarning("No Data", "No segments to generate report from")
            return
        
        import datetime
        
        file_path = filedialog.asksaveasfilename(
            title="Save Session Report",
            defaultextension=".md",
            filetypes=[("Markdown Files", "*.md"), ("Text Files", "*.txt"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return
        
        try:
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            # Calculate statistics
            total_segments = len(self.segments)
            translated = sum(1 for seg in self.segments if seg.target.strip() and seg.status != "untranslated")
            untranslated = total_segments - translated
            draft = sum(1 for seg in self.segments if seg.status == "draft")
            approved = sum(1 for seg in self.segments if seg.status == "approved")
            
            # Get current prompt (context-aware)
            current_mode = "batch_bilingual" if hasattr(self, 'original_txt') else "batch_docx"
            current_prompt = self.get_context_aware_prompt(current_mode)
            
            # Get custom instructions separately if they exist
            custom_instructions = ""
            if hasattr(self, 'custom_instructions_text'):
                try:
                    custom_instructions = self.custom_instructions_text.get('1.0', 'end-1c').strip()
                    # Remove placeholder text if present
                    if custom_instructions.startswith("(Enter project-specific"):
                        custom_instructions = ""
                except:
                    pass
            
            # Determine prompt source
            is_custom_prompt = (hasattr(self, 'current_translate_prompt') and 
                              self.current_translate_prompt != self.single_segment_prompt)
            custom_prompt_source = "Custom loaded prompt" if is_custom_prompt else "Default system prompt"
            
            # Build comprehensive markdown report
            report = f"""# Supervertaler CAT Editor Session Report

## Session Information
- **Date & Time**: {timestamp}
- **Supervertaler Version**: {APP_VERSION}
- **Mode**: CAT Editor with AI-Assisted Translation
- **AI Provider**: {self.current_llm_provider}
- **AI Model**: {self.current_llm_model}

## Project Statistics
- **Total Segments**: {total_segments}
- **Translated**: {translated} ({translated/total_segments*100:.1f}%)
- **Untranslated**: {untranslated} ({untranslated/total_segments*100:.1f}%)
- **Draft (AI-generated)**: {draft}
- **Approved**: {approved}

## Language Settings
- **Source Language**: {self.source_language}
- **Target Language**: {self.target_language}

## Source File
- **Original File**: {getattr(self, 'original_docx', getattr(self, 'original_txt', 'Not saved'))}

## AI Translation Settings

### Active Provider & Model
- **Provider**: {self.current_llm_provider.upper()}
- **Model**: {self.current_llm_model}

### Prompt Configuration
- **Prompt Source**: {custom_prompt_source}
- **Context Mode**: {current_mode}

### System Prompt (Translation Instructions)
```
{current_prompt}
```
"""
            
            # Add custom instructions section only if they exist and are non-empty
            if custom_instructions:
                report += f"""
### Custom Instructions (Project-Specific)
```
{custom_instructions}
```
"""
            
            report += """
## Translation Features Used

### TM (Translation Memory)
- **Status**: {"✅ Enabled" if self.check_tm_var.get() else "❌ Disabled"}
- **TM Entries**: {len(self.tm_agent.memory) if hasattr(self, 'tm_agent') and self.tm_agent else 0}

### Context Awareness
- **Full Document Context**: {"✅ Enabled" if self.use_context_var.get() else "❌ Disabled"}
- **Description**: Provides surrounding segments to AI for better translation quality

## Library Availability
- **Google AI (Gemini)**: {"✅ Available" if GOOGLE_AI_AVAILABLE else "❌ Not Available"}
- **Anthropic (Claude)**: {"✅ Available" if CLAUDE_AVAILABLE else "❌ Not Available"}
- **OpenAI**: {"✅ Available" if OPENAI_AVAILABLE else "❌ Not Available"}
- **PIL (Image Processing)**: {"✅ Available" if PIL_AVAILABLE else "❌ Not Available"}

## API Key Status
- **Google/Gemini**: {"✅ Configured" if self.api_keys.get("google") else "❌ Not Configured"}
- **Claude**: {"✅ Configured" if self.api_keys.get("claude") else "❌ Not Configured"}
- **OpenAI**: {"✅ Configured" if self.api_keys.get("openai") else "❌ Not Configured"}

## Segment Details

### Segments by Status
"""
            
            # Group segments by status
            status_groups = {}
            for seg in self.segments:
                status = seg.status or "untranslated"
                if status not in status_groups:
                    status_groups[status] = []
                status_groups[status].append(seg)
            
            for status, segs in sorted(status_groups.items()):
                report += f"\n#### {status.upper()} ({len(segs)} segments)\n"
                if len(segs) <= 10:
                    for seg in segs:
                        report += f"- **Segment {seg.id}**: {seg.source[:50]}{'...' if len(seg.source) > 50 else ''}\n"
                else:
                    report += f"- {len(segs)} segments (too many to list)\n"
            
            report += f"""

## Workflow Summary

This session used Supervertaler's CAT Editor mode with the following workflow:
1. **Import**: Document imported and segmented
2. **AI Pre-Translation**: {'Segments translated using ' + self.current_llm_provider if translated > 0 else 'No AI translation performed yet'}
3. **Manual Review**: {'Human review and editing of AI translations' if approved > 0 else 'No segments approved yet'}
4. **Quality Control**: TM matching {'enabled' if self.check_tm_var.get() else 'disabled'}

## Technical Information
- **Processing Method**: Segment-by-segment with optional batch translation
- **Context Mode**: {'Full document context provided to AI' if self.use_context_var.get() else 'Single segment translation only'}
- **Output Formats Available**: DOCX, Bilingual DOCX, TSV, TXT, TMX
- **Report Generated**: {timestamp}

---
*This report was automatically generated by Supervertaler v{APP_VERSION} CAT Editor*
"""
            
            # Write markdown report to file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(report)
            
            # Also generate HTML version
            html_path = file_path.rsplit('.', 1)[0] + '.html'
            html_content = self.markdown_to_html(report, f"Supervertaler Session Report - {timestamp}")
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            self.log(f"✓ Session reports saved: {os.path.basename(file_path)} and {os.path.basename(html_path)}")
            messagebox.showinfo("Reports Generated", 
                              f"Session reports saved successfully!\n\n"
                              f"Markdown: {os.path.basename(file_path)}\n"
                              f"HTML: {os.path.basename(html_path)}\n\n"
                              f"The reports include:\n"
                              f"• Project statistics\n"
                              f"• AI configuration\n"
                              f"• Translation settings\n"
                              f"• Segment details\n\n"
                              f"💡 Tip: Double-click the HTML file to open it in your browser!")
            
        except Exception as e:
            self.log(f"✗ Report generation failed: {str(e)}")
            messagebox.showerror("Report Error", f"Failed to generate report:\n{str(e)}")
    
    def show_find_replace(self):
        """Show find/replace dialog"""
        FindReplaceDialog(self.root, self)
    
    # ===== LLM TRANSLATION METHODS (New in v2.5.0) =====
    
    def show_api_settings(self):
        """Show API settings dialog"""
        dialog = tk.Toplevel(self.root)
        dialog.title("API Settings")
        dialog.geometry("600x500")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Main container
        main_frame = ttk.Frame(dialog, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Title
        title_label = tk.Label(main_frame, text="🤖 LLM API Configuration", 
                              font=('Segoe UI', 12, 'bold'))
        title_label.pack(pady=(0, 10))
        
        # API Keys section
        keys_frame = ttk.LabelFrame(main_frame, text="API Keys", padding="10")
        keys_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # OpenAI
        openai_frame = ttk.Frame(keys_frame)
        openai_frame.pack(fill=tk.X, pady=5)
        ttk.Label(openai_frame, text="OpenAI API Key:", width=20).pack(side=tk.LEFT)
        openai_entry = ttk.Entry(openai_frame, width=50, show="*")
        openai_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        openai_entry.insert(0, self.api_keys.get("openai", ""))
        status_openai = ttk.Label(openai_frame, text="✓" if OPENAI_AVAILABLE else "✗ Not installed",
                                  foreground="green" if OPENAI_AVAILABLE else "red")
        status_openai.pack(side=tk.LEFT, padx=5)
        
        # Claude
        claude_frame = ttk.Frame(keys_frame)
        claude_frame.pack(fill=tk.X, pady=5)
        ttk.Label(claude_frame, text="Claude API Key:", width=20).pack(side=tk.LEFT)
        claude_entry = ttk.Entry(claude_frame, width=50, show="*")
        claude_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        claude_entry.insert(0, self.api_keys.get("claude", ""))
        status_claude = ttk.Label(claude_frame, text="✓" if ANTHROPIC_AVAILABLE else "✗ Not installed",
                                 foreground="green" if ANTHROPIC_AVAILABLE else "red")
        status_claude.pack(side=tk.LEFT, padx=5)
        
        # Gemini
        gemini_frame = ttk.Frame(keys_frame)
        gemini_frame.pack(fill=tk.X, pady=5)
        ttk.Label(gemini_frame, text="Google Gemini Key:", width=20).pack(side=tk.LEFT)
        gemini_entry = ttk.Entry(gemini_frame, width=50, show="*")
        gemini_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        gemini_entry.insert(0, self.api_keys.get("google", ""))
        status_gemini = ttk.Label(gemini_frame, text="✓" if GEMINI_AVAILABLE else "✗ Not installed",
                                 foreground="green" if GEMINI_AVAILABLE else "red")
        status_gemini.pack(side=tk.LEFT, padx=5)
        
        # Provider selection
        provider_frame = ttk.LabelFrame(main_frame, text="Default Provider", padding="10")
        provider_frame.pack(fill=tk.X, pady=5)
        
        provider_var = tk.StringVar(value=self.current_llm_provider)
        ttk.Radiobutton(provider_frame, text="OpenAI (GPT-4)", variable=provider_var, 
                       value="openai").pack(anchor=tk.W)
        ttk.Radiobutton(provider_frame, text="Anthropic (Claude)", variable=provider_var,
                       value="claude").pack(anchor=tk.W)
        ttk.Radiobutton(provider_frame, text="Google (Gemini)", variable=provider_var,
                       value="gemini").pack(anchor=tk.W)
        
        # Model selection
        model_frame = ttk.LabelFrame(main_frame, text="Default Model", padding="10")
        model_frame.pack(fill=tk.X, pady=5)
        
        model_var = tk.StringVar(value=self.current_llm_model)
        model_combo = ttk.Combobox(model_frame, textvariable=model_var, state="readonly", width=40)
        
        # Store dynamically fetched models separately
        fetched_models = {}  # provider -> list of models
        
        # Refresh models button
        refresh_frame = ttk.Frame(model_frame)
        refresh_frame.pack(fill=tk.X, pady=5)
        
        def refresh_models():
            """Fetch available models from API"""
            provider = provider_var.get()
            
            # Get API key
            api_key = None
            if provider == "openai":
                api_key = openai_entry.get()
            elif provider == "claude":
                api_key = claude_entry.get()
            elif provider == "gemini":
                api_key = gemini_entry.get()
            
            if not api_key or not api_key.strip():
                messagebox.showwarning("No API Key", f"Please enter a {provider.upper()} API key first.")
                return
            
            # Show loading message
            status_label = ttk.Label(refresh_frame, text="Fetching models...")
            status_label.pack(side=tk.LEFT, padx=5)
            dialog.update()
            
            try:
                available_models = fetch_available_models(provider, api_key)
                if available_models:
                    # Store the fetched models
                    fetched_models[provider] = available_models
                    model_combo['values'] = available_models
                    if model_var.get() not in available_models:
                        model_var.set(available_models[0])
                    status_label.config(text=f"✓ Found {len(available_models)} models", foreground="green")
                    self.log(f"✓ Loaded {len(available_models)} {provider} models")
                else:
                    status_label.config(text="✗ No models found", foreground="red")
            except Exception as e:
                status_label.config(text=f"✗ Error: {str(e)[:30]}", foreground="red")
            
            # Clear status after 3 seconds
            dialog.after(3000, lambda: status_label.destroy())
        
        def update_models(*args):
            provider = provider_var.get()
            
            # Check if we have fetched models for this provider
            if provider in fetched_models:
                model_combo['values'] = fetched_models[provider]
                # Don't reset the model if it's in the fetched list
                if model_var.get() not in fetched_models[provider]:
                    model_var.set(fetched_models[provider][0])
            else:
                # Use cached/fallback models
                if provider == "openai":
                    model_combo['values'] = OPENAI_MODELS
                    if model_var.get() not in OPENAI_MODELS:
                        model_var.set(OPENAI_MODELS[0])
                elif provider == "claude":
                    model_combo['values'] = CLAUDE_MODELS
                    if model_var.get() not in CLAUDE_MODELS:
                        model_var.set(CLAUDE_MODELS[0])
                elif provider == "gemini":
                    model_combo['values'] = GEMINI_MODELS
                    if model_var.get() not in GEMINI_MODELS:
                        model_var.set(GEMINI_MODELS[0])
        
        provider_var.trace('w', update_models)
        update_models()
        model_combo.pack(fill=tk.X, pady=(0, 5))
        
        ttk.Button(refresh_frame, text="🔄 Refresh Available Models", command=refresh_models).pack(side=tk.LEFT)
        
        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=10)
        
        def save_settings():
            # Save API keys
            self.api_keys["openai"] = openai_entry.get()
            self.api_keys["claude"] = claude_entry.get()
            self.api_keys["google"] = gemini_entry.get()
            
            # Save to file
            script_dir = os.path.dirname(os.path.abspath(__file__))
            api_keys_file = os.path.join(script_dir, "api_keys.txt")
            try:
                with open(api_keys_file, 'w', encoding='utf-8') as f:
                    f.write("# API Keys Configuration\n")
                    f.write("# Format: key_name = your_api_key_here\n\n")
                    if self.api_keys["google"]:
                        f.write(f"google = {self.api_keys['google']}\n")
                    if self.api_keys["claude"]:
                        f.write(f"claude = {self.api_keys['claude']}\n")
                    if self.api_keys["openai"]:
                        f.write(f"openai = {self.api_keys['openai']}\n")
                self.log("✓ API keys saved successfully")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save API keys: {e}")
                return
            
            # Save provider and model
            self.current_llm_provider = provider_var.get()
            self.current_llm_model = model_var.get()
            
            # Update the Settings tab display
            if hasattr(self, 'settings_llm_display'):
                self.settings_llm_display.set(f"{self.current_llm_provider.upper()} / {self.current_llm_model}")
            
            self.log(f"✓ Provider set to: {self.current_llm_provider}")
            self.log(f"✓ Model set to: {self.current_llm_model}")
            
            messagebox.showinfo("Success", f"API settings saved successfully!\n\nProvider: {self.current_llm_provider}\nModel: {self.current_llm_model}")
            dialog.destroy()
        
        ttk.Button(button_frame, text="Save", command=save_settings).pack(side=tk.RIGHT, padx=5)
        ttk.Button(button_frame, text="Cancel", command=dialog.destroy).pack(side=tk.RIGHT)
    
    def swap_languages(self):
        """Swap source and target languages"""
        self.source_language, self.target_language = self.target_language, self.source_language
        
        # Update Settings pane displays
        if hasattr(self, 'settings_source_lang_display'):
            self.settings_source_lang_display.set(self.source_language)
        if hasattr(self, 'settings_target_lang_display'):
            self.settings_target_lang_display.set(self.target_language)
        
        self.log(f"🔄 Languages swapped: {self.source_language} → {self.target_language}")
    
    def edit_language_list(self):
        """Edit the custom language list"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Edit Language List")
        dialog.geometry("500x600")
        dialog.transient(self.root)
        dialog.grab_set()
        
        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Available Languages", 
                 font=('Segoe UI', 11, 'bold')).pack(anchor=tk.W, pady=(0, 10))
        
        ttk.Label(main_frame, text="Edit the list below (one language per line):",
                 font=('Segoe UI', 9)).pack(anchor=tk.W, pady=(0, 5))
        
        # Text widget with scrollbar
        text_frame = ttk.Frame(main_frame)
        text_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        scrollbar = ttk.Scrollbar(text_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        lang_text = tk.Text(text_frame, width=50, height=25, yscrollcommand=scrollbar.set,
                           font=('Segoe UI', 10))
        lang_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=lang_text.yview)
        
        # Insert current languages
        lang_text.insert('1.0', '\n'.join(self.available_languages))
        
        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X)
        
        def save_languages():
            # Get text and split into lines
            text_content = lang_text.get('1.0', 'end-1c')
            new_languages = [line.strip() for line in text_content.split('\n') if line.strip()]
            
            if not new_languages:
                messagebox.showwarning("Empty List", "Language list cannot be empty!")
                return
            
            self.available_languages = sorted(new_languages)
            self.log(f"✓ Language list updated: {len(self.available_languages)} languages")
            messagebox.showinfo("Success", f"Language list saved!\n\n{len(self.available_languages)} languages available.")
            dialog.destroy()
        
        def reset_defaults():
            default_languages = [
                "Afrikaans", "Albanian", "Arabic", "Armenian", "Basque", "Bengali",
                "Bulgarian", "Catalan", "Chinese (Simplified)", "Chinese (Traditional)",
                "Croatian", "Czech", "Danish", "Dutch", "English", "Estonian",
                "Finnish", "French", "Galician", "Georgian", "German", "Greek",
                "Hebrew", "Hindi", "Hungarian", "Icelandic", "Indonesian", "Irish",
                "Italian", "Japanese", "Korean", "Latvian", "Lithuanian", "Macedonian",
                "Malay", "Norwegian", "Persian", "Polish", "Portuguese", "Romanian",
                "Russian", "Serbian", "Slovak", "Slovenian", "Spanish", "Swahili",
                "Swedish", "Thai", "Turkish", "Ukrainian", "Urdu", "Vietnamese", "Welsh"
            ]
            lang_text.delete('1.0', tk.END)
            lang_text.insert('1.0', '\n'.join(default_languages))
        
        ttk.Button(button_frame, text="Reset to Defaults", command=reset_defaults).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Save", command=save_languages).pack(side=tk.RIGHT, padx=5)
        ttk.Button(button_frame, text="Cancel", command=dialog.destroy).pack(side=tk.RIGHT)
    
    def show_language_settings(self):
        """Show language settings dialog with dropdowns"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Language Settings")
        dialog.geometry("450x250")
        dialog.transient(self.root)
        dialog.grab_set()
        
        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Header
        ttk.Label(main_frame, text="Translation Language Pair", 
                 font=('Segoe UI', 11, 'bold')).grid(row=0, column=0, columnspan=3, pady=(0, 15))
        
        # Source language
        ttk.Label(main_frame, text="Source Language:", font=('Segoe UI', 9)).grid(row=1, column=0, sticky=tk.W, pady=5)
        source_var = tk.StringVar(value=self.source_language)
        source_combo = ttk.Combobox(main_frame, textvariable=source_var, values=self.available_languages,
                                   state="readonly", width=25, font=('Segoe UI', 9))
        source_combo.grid(row=1, column=1, pady=5, padx=(10, 0), sticky=tk.W)
        
        # Swap button
        def swap():
            source_val = source_var.get()
            target_val = target_var.get()
            source_var.set(target_val)
            target_var.set(source_val)
        
        ttk.Button(main_frame, text="🔄", command=swap, width=3).grid(row=1, column=2, rowspan=2, padx=10)
        
        # Target language
        ttk.Label(main_frame, text="Target Language:", font=('Segoe UI', 9)).grid(row=2, column=0, sticky=tk.W, pady=5)
        target_var = tk.StringVar(value=self.target_language)
        target_combo = ttk.Combobox(main_frame, textvariable=target_var, values=self.available_languages,
                                   state="readonly", width=25, font=('Segoe UI', 9))
        target_combo.grid(row=2, column=1, pady=5, padx=(10, 0), sticky=tk.W)
        
        # Info label
        ttk.Label(main_frame, text="💡 Tip: Use 'Edit Language List' in Settings to customize available languages",
                 font=('Segoe UI', 8), foreground='#666').grid(row=3, column=0, columnspan=3, pady=(10, 0), sticky=tk.W)
        
        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=4, column=0, columnspan=3, pady=20)
        
        def save():
            self.source_language = source_var.get()
            self.target_language = target_var.get()
            
            # Update Settings pane displays
            if hasattr(self, 'settings_source_lang_display'):
                self.settings_source_lang_display.set(self.source_language)
            if hasattr(self, 'settings_target_lang_display'):
                self.settings_target_lang_display.set(self.target_language)
            
            self.log(f"✓ Language pair: {self.source_language} → {self.target_language}")
            dialog.destroy()
        
        ttk.Button(button_frame, text="Save", command=save).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancel", command=dialog.destroy).pack(side=tk.LEFT)
    
    def show_custom_prompts(self, initial_filter="all"):
        """Show comprehensive prompt library browser
        
        Args:
            initial_filter: Initial type filter - "all", "system_prompt", or "custom_instruction"
        """
        dialog = tk.Toplevel(self.root)
        dialog.title("🎯 Prompt Library - System Prompts & Custom Instructions")
        dialog.geometry("1000x700")
        dialog.transient(self.root)
        
        main_frame = ttk.Frame(dialog, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Header with active prompt info
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(header_frame, text="📚 Prompt Library", 
                 font=('Segoe UI', 12, 'bold')).pack(side=tk.LEFT)
        
        active_label = ttk.Label(header_frame, text="", font=('Segoe UI', 9))
        active_label.pack(side=tk.RIGHT)
        
        def update_active_label():
            if self.prompt_library.active_prompt_name:
                active_label.config(text=f"✓ Active: {self.prompt_library.active_prompt_name}", 
                                   foreground='green')
            else:
                active_label.config(text="Using default prompt", foreground='gray')
        
        update_active_label()
        
        # Search frame
        search_frame = ttk.Frame(main_frame)
        search_frame.pack(fill=tk.X, pady=(0, 5))
        
        ttk.Label(search_frame, text="🔍 Search:").pack(side=tk.LEFT, padx=(0, 5))
        search_var = tk.StringVar()
        search_entry = ttk.Entry(search_frame, textvariable=search_var, width=40)
        search_entry.pack(side=tk.LEFT, padx=(0, 10))
        
        # Type filter frame
        filter_frame = ttk.Frame(main_frame)
        filter_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(filter_frame, text="📋 Type:", font=('Segoe UI', 9, 'bold')).pack(side=tk.LEFT, padx=(0, 10))
        
        type_filter_var = tk.StringVar(value=initial_filter)
        ttk.Radiobutton(filter_frame, text="All", variable=type_filter_var, value="all").pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(filter_frame, text="🎭 System prompts", variable=type_filter_var, value="system_prompt").pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(filter_frame, text="📝 Custom instructions", variable=type_filter_var, value="custom_instruction").pack(side=tk.LEFT, padx=5)
        
        # Trigger tree refresh when filter changes
        type_filter_var.trace('w', lambda *args: load_prompts_to_tree())
        
        # Main content - 3 panes
        paned = ttk.PanedWindow(main_frame, orient=tk.HORIZONTAL)
        paned.pack(fill=tk.BOTH, expand=True)
        
        # Left pane - Prompt list
        left_frame = ttk.Frame(paned)
        paned.add(left_frame, weight=1)
        
        ttk.Label(left_frame, text="Available Prompts", font=('Segoe UI', 10, 'bold')).pack(anchor=tk.W)
        
        # Type info
        location_label = ttk.Label(left_frame, 
                                  text="🎭 System prompts  📝 Custom instructions",
                                  font=('Segoe UI', 8), foreground='#666')
        location_label.pack(anchor=tk.W, pady=(0, 5))
        
        # Treeview for prompts
        tree_frame = ttk.Frame(left_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True, pady=(5, 0))
        
        prompt_tree = ttk.Treeview(tree_frame, columns=('Type', 'Domain', 'Location'), show='tree headings', height=15)
        prompt_tree.heading('#0', text='Name ▼')
        prompt_tree.heading('Type', text='Type')
        prompt_tree.heading('Domain', text='Domain')
        prompt_tree.heading('Location', text='Location')
        prompt_tree.column('#0', width=180)
        prompt_tree.column('Type', width=120)
        prompt_tree.column('Domain', width=100)
        prompt_tree.column('Location', width=80)
        
        tree_scroll = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=prompt_tree.yview)
        prompt_tree.configure(yscrollcommand=tree_scroll.set)
        prompt_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Sorting state
        sort_state = {'column': 'name', 'reverse': False}
        
        # Right pane - Preview and details
        right_frame = ttk.Frame(paned)
        paned.add(right_frame, weight=2)
        
        # Details section
        details_frame = ttk.LabelFrame(right_frame, text="Prompt Details", padding=10)
        details_frame.pack(fill=tk.BOTH, expand=True)
        
        # Metadata
        meta_frame = ttk.Frame(details_frame)
        meta_frame.pack(fill=tk.X, pady=(0, 10))
        
        name_label = ttk.Label(meta_frame, text="", font=('Segoe UI', 11, 'bold'))
        name_label.pack(anchor=tk.W)
        
        desc_label = ttk.Label(meta_frame, text="", wraplength=500, justify=tk.LEFT, foreground='#666')
        desc_label.pack(anchor=tk.W, pady=(5, 0))
        
        info_label = ttk.Label(meta_frame, text="", font=('Segoe UI', 8), foreground='#999')
        info_label.pack(anchor=tk.W, pady=(5, 0))
        
        # File path label (clickable)
        filepath_label = ttk.Label(meta_frame, text="", font=('Consolas', 7), foreground='#0066cc', cursor='hand2')
        filepath_label.pack(anchor=tk.W, pady=(3, 0))
        
        # Tabs for translate and proofread prompts
        notebook = ttk.Notebook(details_frame)
        notebook.pack(fill=tk.BOTH, expand=True, pady=(10, 0))
        
        # Translate prompt tab
        translate_frame = ttk.Frame(notebook)
        notebook.add(translate_frame, text="Translation Prompt")
        
        translate_text = tk.Text(translate_frame, wrap=tk.WORD, height=12, font=('Consolas', 9))
        translate_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        translate_scroll = ttk.Scrollbar(translate_frame, command=translate_text.yview)
        translate_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        translate_text.config(yscrollcommand=translate_scroll.set, state='disabled')
        
        # Proofread prompt tab
        proofread_frame = ttk.Frame(notebook)
        notebook.add(proofread_frame, text="Proofreading Prompt")
        
        proofread_text = tk.Text(proofread_frame, wrap=tk.WORD, height=12, font=('Consolas', 9))
        proofread_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        proofread_scroll = ttk.Scrollbar(proofread_frame, command=proofread_text.yview)
        proofread_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        proofread_text.config(yscrollcommand=proofread_scroll.set, state='disabled')
        
        # Selected prompt data
        selected_prompt = {'filename': None, 'data': None}
        
        def sort_prompts(prompts_list, column, reverse=False):
            """Sort prompts by specified column"""
            if column == 'name':
                return sorted(prompts_list, key=lambda x: x['name'].lower(), reverse=reverse)
            elif column == 'domain':
                return sorted(prompts_list, key=lambda x: x['domain'].lower(), reverse=reverse)
            elif column == 'location':
                # Location sorting removed - no longer relevant with auto-routing
                return prompts_list
            elif column == 'type':
                return sorted(prompts_list, key=lambda x: x.get('_type', 'system_prompt'), reverse=reverse)
            return prompts_list
        
        def update_sort_indicators():
            """Update column headers with sort indicators"""
            arrow = ' ▼' if not sort_state['reverse'] else ' ▲'
            
            # Reset all headers
            prompt_tree.heading('#0', text='Name')
            prompt_tree.heading('Type', text='Type')
            prompt_tree.heading('Domain', text='Domain')
            prompt_tree.heading('Location', text='Location')
            
            # Add arrow to sorted column
            if sort_state['column'] == 'name':
                prompt_tree.heading('#0', text='Name' + arrow)
            elif sort_state['column'] == 'type':
                prompt_tree.heading('Type', text='Type' + arrow)
            elif sort_state['column'] == 'domain':
                prompt_tree.heading('Domain', text='Domain' + arrow)
            elif sort_state['column'] == 'location':
                prompt_tree.heading('Location', text='Location' + arrow)
        
        def on_column_click(column):
            """Handle column header click for sorting"""
            # Toggle reverse if clicking same column
            if sort_state['column'] == column:
                sort_state['reverse'] = not sort_state['reverse']
            else:
                sort_state['column'] = column
                sort_state['reverse'] = False
            
            update_sort_indicators()
            load_prompts_to_tree()
        
        # Bind column header clicks
        prompt_tree.heading('#0', command=lambda: on_column_click('name'))
        prompt_tree.heading('Type', command=lambda: on_column_click('type'))
        prompt_tree.heading('Domain', command=lambda: on_column_click('domain'))
        prompt_tree.heading('Location', command=lambda: on_column_click('location'))
        
        def load_prompts_to_tree(prompts_list=None):
            """Load prompts into treeview with type filtering"""
            prompt_tree.delete(*prompt_tree.get_children())
            
            if prompts_list is None:
                prompts_list = self.prompt_library.get_prompt_list()
            
            # Apply type filter
            type_filter = type_filter_var.get()
            if type_filter != "all":
                prompts_list = [p for p in prompts_list if p.get('_type', 'system_prompt') == type_filter]
            
            # Apply sorting
            prompts_list = sort_prompts(prompts_list, sort_state['column'], sort_state['reverse'])
            
            for prompt_info in prompts_list:
                # Determine type label
                prompt_type = prompt_info.get('_type', 'system_prompt')
                type_label = "🎭 System prompts" if prompt_type == 'system_prompt' else "📝 Custom instructions"
                
                # Simple icon based on type
                icon = "\U0001F3AD" if prompt_type == 'system_prompt' else "\U0001F4DD"
                location = "System prompts" if prompt_type == 'system_prompt' else "Custom instructions"
                
                prompt_tree.insert('', 'end', 
                                  text=f"{icon} {prompt_info['name']}", 
                                  values=(type_label, prompt_info['domain'], location),
                                  tags=(prompt_info['filename'],))
        
        def on_search(*args):
            """Filter prompts by search text"""
            search_text = search_var.get()
            if search_text:
                results = self.prompt_library.search_prompts(search_text)
                load_prompts_to_tree(results)
            else:
                load_prompts_to_tree()
        
        search_var.trace('w', on_search)
        
        def on_select(event):
            """Show prompt details when selected"""
            selection = prompt_tree.selection()
            if not selection:
                return
            
            # Get filename from tags
            item = selection[0]
            tags = prompt_tree.item(item, 'tags')
            if not tags:
                return
            
            filename = tags[0]
            prompt_data = self.prompt_library.get_prompt(filename)
            
            if not prompt_data:
                return
            
            selected_prompt['filename'] = filename
            selected_prompt['data'] = prompt_data
            
            # Update metadata
            name_label.config(text=prompt_data.get('name', 'Unnamed'))
            desc_label.config(text=prompt_data.get('description', 'No description'))
            
            domain = prompt_data.get('domain', 'General')
            version = prompt_data.get('version', '1.0')
            created = prompt_data.get('created', 'Unknown')
            
            # Determine location based on type
            prompt_type = prompt_data.get('_type', 'system_prompt')
            
            if prompt_type == 'custom_instruction':
                type_display = "� Custom Instruction"
            else:  # system_prompt
                type_display = "🎭 System Prompt"
            
            info_label.config(text=f"Domain: {domain} | Version: {version} | Created: {created} | Type: {type_display}")
            
            # Update file path
            filepath = prompt_data.get('_filepath', '')
            if filepath:
                filepath_label.config(text=f"📂 {filepath}")
                # Make clickable to open folder
                def open_folder(e):
                    import subprocess
                    folder_path = os.path.dirname(filepath)
                    if os.path.exists(folder_path):
                        subprocess.Popen(f'explorer /select,"{filepath}"')
                filepath_label.bind('<Button-1>', open_folder)
            else:
                filepath_label.config(text="")
            
            # Update translate prompt
            translate_text.config(state='normal')
            translate_text.delete('1.0', 'end')
            translate_text.insert('1.0', prompt_data.get('translate_prompt', 'N/A'))
            translate_text.config(state='disabled')
            
            # Update proofread prompt
            proofread_text.config(state='normal')
            proofread_text.delete('1.0', 'end')
            proofread_text.insert('1.0', prompt_data.get('proofread_prompt', 'N/A'))
            proofread_text.config(state='disabled')
        
        prompt_tree.bind('<<TreeviewSelect>>', on_select)
        
        # Bottom button frame
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=(10, 0))
        
        def apply_prompt():
            """Apply selected prompt as active"""
            if not selected_prompt['filename']:
                messagebox.showwarning("No Selection", "Please select a prompt to apply")
                return
            
            if self.prompt_library.set_active_prompt(selected_prompt['filename']):
                # Update current prompts
                self.current_translate_prompt = self.prompt_library.get_translate_prompt()
                proofread = self.prompt_library.get_proofread_prompt()
                if proofread:
                    self.current_proofread_prompt = proofread
                
                update_active_label()
                messagebox.showinfo("Prompt Applied", 
                                   f"Now using: {selected_prompt['data']['name']}\n\n"
                                   "This custom prompt will be used for all translations.")
        
        def use_default():
            """Clear active prompt and use default"""
            self.prompt_library.clear_active_prompt()
            self.current_translate_prompt = self.default_translate_prompt
            self.current_proofread_prompt = self.default_proofread_prompt
            update_active_label()
            messagebox.showinfo("Default Prompt", "Now using default translation prompt")
        
        def create_new():
            """Create new custom prompt"""
            self.create_prompt_editor(dialog, on_save=lambda: load_prompts_to_tree())
        
        def edit_selected():
            """Edit selected prompt"""
            if not selected_prompt['filename']:
                messagebox.showwarning("No Selection", "Please select a prompt to edit")
                return
            
            self.create_prompt_editor(dialog, 
                                     edit_prompt=selected_prompt['data'],
                                     on_save=lambda: load_prompts_to_tree())
        
        def delete_selected():
            """Delete selected prompt"""
            if not selected_prompt['filename']:
                messagebox.showwarning("No Selection", "Please select a prompt to delete")
                return
            
            prompt_name = selected_prompt['data'].get('name', 'this prompt')
            if messagebox.askyesno("Confirm Delete", 
                                  f"Delete '{prompt_name}'?\n\nThis cannot be undone."):
                if self.prompt_library.delete_prompt(selected_prompt['filename']):
                    load_prompts_to_tree()
                    # Clear selection
                    selected_prompt['filename'] = None
                    selected_prompt['data'] = None
                    update_active_label()
        
        def import_prompt():
            """Import prompt from file"""
            filepath = filedialog.askopenfilename(
                title="Import Custom Prompt",
                filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")]
            )
            
            if not filepath:
                return
            
            if self.prompt_library.import_prompt(filepath):
                load_prompts_to_tree()
                messagebox.showinfo("Import Success", "Prompt imported successfully!")
        
        def export_selected():
            """Export selected prompt"""
            if not selected_prompt['filename']:
                messagebox.showwarning("No Selection", "Please select a prompt to export")
                return
            
            filepath = filedialog.asksaveasfilename(
                title="Export Custom Prompt",
                defaultextension=".json",
                initialfile=selected_prompt['filename'],
                filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")]
            )
            
            if filepath:
                if self.prompt_library.export_prompt(selected_prompt['filename'], filepath):
                    messagebox.showinfo("Export Success", "Prompt exported successfully!")
        
        def refresh_list():
            """Reload all prompts"""
            self.prompt_library.load_all_prompts()
            load_prompts_to_tree()
            update_active_label()
        
        # Left buttons
        ttk.Button(button_frame, text="✓ Apply Selected", command=apply_prompt).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="↺ Use Default", command=use_default).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="➕ New", command=create_new).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="✏️ Edit", command=edit_selected).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="🗑️ Delete", command=delete_selected).pack(side=tk.LEFT, padx=2)
        
        # Separator
        ttk.Separator(button_frame, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=10)
        
        # Right buttons
        ttk.Button(button_frame, text="📥 Import", command=import_prompt).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="📤 Export", command=export_selected).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="🔄 Refresh", command=refresh_list).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="Close", command=dialog.destroy).pack(side=tk.RIGHT, padx=2)
        
        # Initial load
        load_prompts_to_tree()
    
    def show_system_prompts(self):
        """Show Prompt Library filtered to System Prompts only"""
        self.show_custom_prompts(initial_filter="system_prompt")
    
    def show_custom_instructions(self):
        """Show Prompt Library filtered to Custom Instructions only"""
        self.show_custom_prompts(initial_filter="custom_instruction")
    
    def create_prompt_editor(self, parent, edit_prompt=None, on_save=None):
        """Show prompt creation/editing dialog"""
        editor = tk.Toplevel(parent)
        editor.title("Edit Prompt" if edit_prompt else "Create New Prompt")
        editor.geometry("900x700")
        editor.transient(parent)
        editor.grab_set()
        
        main_frame = ttk.Frame(editor, padding=15)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Metadata section
        meta_frame = ttk.LabelFrame(main_frame, text="Prompt Metadata", padding=10)
        meta_frame.pack(fill=tk.X, pady=(0, 10))
        
        # Name
        ttk.Label(meta_frame, text="Name:").grid(row=0, column=0, sticky=tk.W, pady=5)
        name_var = tk.StringVar(value=edit_prompt.get('name', '') if edit_prompt else '')
        name_entry = ttk.Entry(meta_frame, textvariable=name_var, width=50)
        name_entry.grid(row=0, column=1, sticky=tk.EW, pady=5, padx=(5, 0))
        
        # Description
        ttk.Label(meta_frame, text="Description:").grid(row=1, column=0, sticky=tk.W, pady=5)
        desc_var = tk.StringVar(value=edit_prompt.get('description', '') if edit_prompt else '')
        desc_entry = ttk.Entry(meta_frame, textvariable=desc_var, width=50)
        desc_entry.grid(row=1, column=1, sticky=tk.EW, pady=5, padx=(5, 0))
        
        # Domain
        ttk.Label(meta_frame, text="Domain:").grid(row=2, column=0, sticky=tk.W, pady=5)
        domain_var = tk.StringVar(value=edit_prompt.get('domain', 'General') if edit_prompt else 'General')
        domain_entry = ttk.Entry(meta_frame, textvariable=domain_var, width=50)
        domain_entry.grid(row=2, column=1, sticky=tk.EW, pady=5, padx=(5, 0))
        
        # Version
        ttk.Label(meta_frame, text="Version:").grid(row=3, column=0, sticky=tk.W, pady=5)
        version_var = tk.StringVar(value=edit_prompt.get('version', '1.0') if edit_prompt else '1.0')
        version_entry = ttk.Entry(meta_frame, textvariable=version_var, width=50)
        version_entry.grid(row=3, column=1, sticky=tk.EW, pady=5, padx=(5, 0))
        
        # Type dropdown
        ttk.Label(meta_frame, text="Type:").grid(row=4, column=0, sticky=tk.W, pady=5)
        type_var = tk.StringVar(value=edit_prompt.get('_type', 'system_prompt') if edit_prompt else 'system_prompt')
        type_combo = ttk.Combobox(meta_frame, textvariable=type_var, width=48, state='readonly')
        type_combo['values'] = ('🎭 System prompts', '📝 Custom instructions')
        # Map display values to internal values
        type_display_map = {'🎭 System prompts': 'system_prompt', '📝 Custom instructions': 'custom_instruction'}
        type_reverse_map = {'system_prompt': '🎭 System prompts', 'custom_instruction': '📝 Custom instructions'}
        # Set initial display value
        current_type = edit_prompt.get('_type', 'system_prompt') if edit_prompt else 'system_prompt'
        type_var.set(type_reverse_map.get(current_type, '🎭 System prompts'))
        type_combo.grid(row=4, column=1, sticky=tk.EW, pady=5, padx=(5, 0))
        
        # Helper text for type
        type_help = ttk.Label(meta_frame, 
                             text="🎭 System prompts: Defines AI role/expertise  |  📝 Custom instructions: User preferences/context",
                             font=('Segoe UI', 7), foreground='#666')
        type_help.grid(row=5, column=1, sticky=tk.W, pady=(0, 5), padx=(5, 0))
        
        # Note about auto-routing in dev mode
        if ENABLE_PRIVATE_FEATURES:
            dev_note = ttk.Label(meta_frame, 
                               text="🔒 DEV MODE: All prompts auto-save to private folders (user data_private/)",
                               font=('Segoe UI', 7), foreground='#cc6600')
            dev_note.grid(row=6, column=1, sticky=tk.W, pady=5, padx=(5, 0))
        
        meta_frame.columnconfigure(1, weight=1)
        
        # Prompts section
        prompts_frame = ttk.Frame(main_frame)
        prompts_frame.pack(fill=tk.BOTH, expand=True)
        
        notebook = ttk.Notebook(prompts_frame)
        notebook.pack(fill=tk.BOTH, expand=True)
        
        # Translation prompt
        translate_frame = ttk.Frame(notebook)
        notebook.add(translate_frame, text="Translation Prompt")
        
        ttk.Label(translate_frame, text="Variables: {source_lang}, {target_lang}", 
                 foreground='#666').pack(anchor=tk.W, pady=(5, 2))
        
        translate_text = tk.Text(translate_frame, wrap=tk.WORD, height=20, font=('Consolas', 9))
        translate_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        translate_scroll = ttk.Scrollbar(translate_frame, command=translate_text.yview)
        translate_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        translate_text.config(yscrollcommand=translate_scroll.set)
        
        if edit_prompt:
            translate_text.insert('1.0', edit_prompt.get('translate_prompt', ''))
        
        # Proofread prompt
        proofread_frame = ttk.Frame(notebook)
        notebook.add(proofread_frame, text="Proofreading Prompt")
        
        ttk.Label(proofread_frame, text="Variables: {source_lang}, {target_lang}", 
                 foreground='#666').pack(anchor=tk.W, pady=(5, 2))
        
        proofread_text = tk.Text(proofread_frame, wrap=tk.WORD, height=20, font=('Consolas', 9))
        proofread_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        proofread_scroll = ttk.Scrollbar(proofread_frame, command=proofread_text.yview)
        proofread_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        proofread_text.config(yscrollcommand=proofread_scroll.set)
        
        if edit_prompt:
            proofread_text.insert('1.0', edit_prompt.get('proofread_prompt', ''))
        
        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=(10, 0))
        
        def save_prompt():
            """Save the prompt"""
            name = name_var.get().strip()
            if not name:
                messagebox.showwarning("Missing Name", "Please enter a prompt name")
                return
            
            translate_prompt = translate_text.get('1.0', 'end-1c').strip()
            if not translate_prompt:
                messagebox.showwarning("Missing Prompt", "Please enter a translation prompt")
                return
            
            description = desc_var.get().strip()
            domain = domain_var.get().strip()
            version = version_var.get().strip()
            proofread_prompt = proofread_text.get('1.0', 'end-1c').strip()
            # Convert display value to internal value
            prompt_type_display = type_var.get()
            prompt_type = type_display_map.get(prompt_type_display, 'system_prompt')
            
            if edit_prompt:
                # Update existing
                filename = edit_prompt['_filename']
                success = self.prompt_library.update_prompt(
                    filename, name, description, domain, translate_prompt, 
                    proofread_prompt, version
                )
            else:
                # Create new
                success = self.prompt_library.create_new_prompt(
                    name, description, domain, translate_prompt, 
                    proofread_prompt, version, prompt_type
                )
            
            if success:
                if on_save:
                    on_save()
                editor.destroy()
        
        ttk.Button(button_frame, text="💾 Save", command=save_prompt).pack(side=tk.RIGHT, padx=2)
        ttk.Button(button_frame, text="Cancel", command=editor.destroy).pack(side=tk.RIGHT, padx=2)
    
    def load_tm_file(self):
        """Load a TM file (TMX or TXT) with language code selection for TMX"""
        filepath = filedialog.askopenfilename(
            title="Load Translation Memory",
            filetypes=[
                ("TMX Files", "*.tmx"),
                ("Text Files", "*.txt"),
                ("All Files", "*.*")
            ]
        )
        
        if not filepath:
            return
        
        _, ext = os.path.splitext(filepath)
        
        # For TMX files, show language code selector
        if ext.lower() == ".tmx":
            # Detect available languages in TMX
            detected_langs = self.tm_database.detect_tmx_languages(filepath)
            
            if not detected_langs:
                messagebox.showerror("TMX Error", "Could not detect languages in TMX file.")
                return
            
            # Show language selector dialog
            src_lang, tgt_lang, read_only = self.show_tmx_language_selector(filepath, detected_langs)
            
            if not src_lang or not tgt_lang:
                return  # User cancelled
        else:
            src_lang, tgt_lang, read_only = None, None, False
        
        # Create progress dialog
        progress_dialog = tk.Toplevel(self.root)
        progress_dialog.title("Loading Translation Memory")
        progress_dialog.geometry("400x150")
        progress_dialog.transient(self.root)
        progress_dialog.grab_set()
        
        frame = ttk.Frame(progress_dialog, padding="20")
        frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(frame, text=f"Loading {os.path.basename(filepath)}...", 
                 font=('Segoe UI', 10)).pack(pady=(0, 10))
        
        progress_label = ttk.Label(frame, text="Initializing...")
        progress_label.pack(pady=5)
        
        progress_bar = ttk.Progressbar(frame, mode='indeterminate')
        progress_bar.pack(fill=tk.X, pady=10)
        progress_bar.start(10)
        
        result_tm_id = [None]
        result_count = [0]
        error_msg = [None]
        
        def load_in_thread():
            try:
                if ext.lower() == ".tmx":
                    # Load into new custom TM
                    tm_name = os.path.basename(filepath).replace('.tmx', '')
                    result_tm_id[0], result_count[0] = self.tm_database.load_tmx_file(
                        filepath, src_lang, tgt_lang, tm_name=tm_name, read_only=read_only
                    )
                    
                elif ext.lower() == ".txt":
                    # Load into Big Mama (legacy behavior)
                    result_count[0] = self.tm_agent.load_from_txt(filepath)
                else:
                    error_msg[0] = f"Unsupported file type: {ext}"
            
            except Exception as e:
                error_msg[0] = str(e)
            
            # Update UI from main thread
            self.root.after(0, finish_loading)
        
        def finish_loading():
            progress_bar.stop()
            progress_dialog.destroy()
            
            if error_msg[0]:
                messagebox.showerror("Load Error", f"Failed to load TM file:\n{error_msg[0]}")
                self.log(f"✗ TM load failed: {error_msg[0]}")
            elif result_count[0] == 0:
                messagebox.showwarning("No Entries", 
                    f"Loaded 0 translation pairs from {ext.upper()} file.\n\n"
                    f"Possible reasons:\n"
                    f"• Selected language codes not found in TMX\n"
                    f"• File format is incorrect")
                self.log(f"⚠ Loaded 0 entries from {os.path.basename(filepath)}")
            else:
                tm_info = ""
                if result_tm_id[0]:
                    tm = self.tm_database.get_tm(result_tm_id[0])
                    tm_info = f"\n\nTM Name: {tm.name}\nLanguages: {src_lang} → {tgt_lang}"
                
                messagebox.showinfo("TM Loaded", 
                    f"Successfully loaded {result_count[0]} translation pairs{tm_info}")
                self.log(f"✓ Loaded {result_count[0]} entries into TM: {result_tm_id[0] or 'Big Mama'}")
                
                # Refresh TM dropdown and manager
                self.update_tm_dropdown()
                self.refresh_tm_manager()
        
        # Start loading in background thread
        import threading
        thread = threading.Thread(target=load_in_thread, daemon=True)
        thread.start()
    
    def show_tmx_language_selector(self, filepath: str, detected_langs: List[str]) -> tuple:
        """
        Show dialog for selecting source/target languages from TMX
        Returns: (source_lang_code, target_lang_code, read_only) or (None, None, False) if cancelled
        """
        dialog = tk.Toplevel(self.root)
        dialog.title("Select TMX Languages")
        dialog.geometry("500x300")
        dialog.transient(self.root)
        dialog.grab_set()
        
        result = {'src': None, 'tgt': None, 'read_only': False}
        
        main_frame = ttk.Frame(dialog, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Info
        ttk.Label(main_frame, text=f"TMX File: {os.path.basename(filepath)}", 
                 font=('Segoe UI', 10, 'bold')).pack(anchor=tk.W, pady=(0, 10))
        
        ttk.Label(main_frame, text=f"Detected {len(detected_langs)} language(s) in file:", 
                 font=('Segoe UI', 9)).pack(anchor=tk.W, pady=(0, 5))
        
        lang_display = ", ".join(detected_langs[:10])  # Show first 10
        if len(detected_langs) > 10:
            lang_display += f" ... and {len(detected_langs) - 10} more"
        
        ttk.Label(main_frame, text=lang_display, 
                 font=('Segoe UI', 8), foreground='gray').pack(anchor=tk.W, pady=(0, 15))
        
        # Language selection
        select_frame = ttk.Frame(main_frame)
        select_frame.pack(fill=tk.X, pady=10)
        
        # Source language
        ttk.Label(select_frame, text="Source Language:").grid(row=0, column=0, sticky=tk.W, pady=5)
        src_var = tk.StringVar()
        src_combo = ttk.Combobox(select_frame, textvariable=src_var, values=detected_langs, 
                                state='readonly', width=25)
        src_combo.grid(row=0, column=1, sticky=tk.EW, padx=(10, 0), pady=5)
        
        # Target language
        ttk.Label(select_frame, text="Target Language:").grid(row=1, column=0, sticky=tk.W, pady=5)
        tgt_var = tk.StringVar()
        tgt_combo = ttk.Combobox(select_frame, textvariable=tgt_var, values=detected_langs, 
                                state='readonly', width=25)
        tgt_combo.grid(row=1, column=1, sticky=tk.EW, padx=(10, 0), pady=5)
        
        select_frame.columnconfigure(1, weight=1)
        
        # Auto-detect based on current GUI languages
        def auto_detect():
            """Try to match GUI languages to TMX codes"""
            gui_src = self.source_language[:2].lower()
            gui_tgt = self.target_language[:2].lower()
            
            # Try to find matching codes
            for lang in detected_langs:
                lang_code = lang.split('-')[0].split('_')[0].lower()
                if lang_code == gui_src:
                    src_var.set(lang)
                if lang_code == gui_tgt:
                    tgt_var.set(lang)
        
        # Auto-detect button
        auto_btn = ttk.Button(main_frame, text="🔍 Auto-Detect from GUI Languages", 
                             command=auto_detect)
        auto_btn.pack(pady=10)
        
        # Try auto-detect on startup
        auto_detect()
        
        # Options
        options_frame = ttk.LabelFrame(main_frame, text="Options", padding=10)
        options_frame.pack(fill=tk.X, pady=10)
        
        read_only_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(options_frame, text="Load as read-only TM (prevent modifications)", 
                       variable=read_only_var).pack(anchor=tk.W)
        
        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=(10, 0))
        
        def on_ok():
            if not src_var.get() or not tgt_var.get():
                messagebox.showwarning("Selection Required", 
                    "Please select both source and target languages")
                return
            
            if src_var.get() == tgt_var.get():
                messagebox.showwarning("Invalid Selection", 
                    "Source and target languages must be different")
                return
            
            result['src'] = src_var.get()
            result['tgt'] = tgt_var.get()
            result['read_only'] = read_only_var.get()
            dialog.destroy()
        
        def on_cancel():
            dialog.destroy()
        
        ttk.Button(button_frame, text="✓ Load TM", command=on_ok).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="✗ Cancel", command=on_cancel).pack(side=tk.LEFT)
        
        dialog.wait_window()
        return result['src'], result['tgt'], result['read_only']
    
    def refresh_tm_manager(self):
        """Refresh TM Manager if it's open and update TM dropdown"""
        self.update_tm_dropdown()
    
    def update_tm_dropdown(self):
        """Update TM source dropdown with all available TMs"""
        if not hasattr(self, 'tm_source_combo'):
            return
        
        # Build list of TM options
        tm_options = ["All Active TMs"]
        
        for tm in self.tm_database.get_all_tms(enabled_only=False):
            status = "✓" if tm.enabled else "✗"
            tm_options.append(f"{status} {tm.name}")
        
        # Update combobox
        self.tm_source_combo['values'] = tm_options
        
        # Keep current selection if still valid
        current = self.tm_source_var.get()
        if current not in tm_options:
            self.tm_source_var.set("All Active TMs")
    
    def show_tm_manager(self):
        """Show multi-TM management dialog with enable/disable controls"""
        # Update dropdown first
        self.update_tm_dropdown()
        
        dialog = tk.Toplevel(self.root)
        dialog.title("Translation Memory Manager")
        dialog.geometry("900x600")
        dialog.transient(self.root)
        
        main_frame = ttk.Frame(dialog, padding=10)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Header
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 10))
        
        total_entries = self.tm_database.get_entry_count(enabled_only=False)
        enabled_entries = self.tm_database.get_entry_count(enabled_only=True)
        
        ttk.Label(header_frame, text="📚 Translation Memory Database", 
                 font=('Segoe UI', 12, 'bold')).pack(anchor=tk.W)
        ttk.Label(header_frame, text=f"Total: {total_entries} entries ({enabled_entries} in active TMs)", 
                 font=('Segoe UI', 9)).pack(anchor=tk.W)
        
        # TM List
        list_frame = ttk.LabelFrame(main_frame, text="Translation Memories", padding=10)
        list_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # Treeview for TMs
        columns = ("enabled", "name", "entries", "languages", "type")
        tree = ttk.Treeview(list_frame, columns=columns, show='headings', height=12)
        
        tree.heading("enabled", text="✓")
        tree.heading("name", text="TM Name")
        tree.heading("entries", text="Entries")
        tree.heading("languages", text="Languages")
        tree.heading("type", text="Type")
        
        tree.column("enabled", width=40, anchor=tk.CENTER)
        tree.column("name", width=250, anchor=tk.W)
        tree.column("entries", width=100, anchor=tk.CENTER)
        tree.column("languages", width=150, anchor=tk.CENTER)
        tree.column("type", width=120, anchor=tk.CENTER)
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)
        
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        def populate_tree():
            """Populate tree with all TMs"""
            tree.delete(*tree.get_children())
            
            for tm in self.tm_database.get_all_tms(enabled_only=False):
                enabled_icon = "✓" if tm.enabled else "✗"
                entry_count = tm.get_entry_count()
                
                src_lang = tm.metadata.get('source_lang', 'N/A')
                tgt_lang = tm.metadata.get('target_lang', 'N/A')
                lang_pair = f"{src_lang} → {tgt_lang}" if src_lang != 'N/A' else "Not set"
                
                # Determine type
                if tm.tm_id == 'project':
                    tm_type = "Project TM"
                elif tm.tm_id == 'big_mama' or tm.tm_id == 'main':
                    tm_type = "Big Mama"
                else:
                    tm_type = "Custom TM" + (" [RO]" if tm.read_only else "")
                
                tree.insert("", tk.END, iid=tm.tm_id, 
                           values=(enabled_icon, tm.name, entry_count, lang_pair, tm_type))
        
        populate_tree()
        
        # Selection state
        selected_tm_id = [None]
        
        def on_select(event):
            """Update selected TM"""
            selection = tree.selection()
            selected_tm_id[0] = selection[0] if selection else None
        
        tree.bind('<<TreeviewSelect>>', on_select)
        
        # Toggle enabled/disabled on double-click
        def on_double_click(event):
            """Toggle TM enabled state"""
            item = tree.identify_row(event.y)
            if item:
                tm = self.tm_database.get_tm(item)
                if tm:
                    tm.enabled = not tm.enabled
                    populate_tree()
                    self.update_tm_dropdown()  # Update dropdown
                    self.log(f"{'Enabled' if tm.enabled else 'Disabled'} TM: {tm.name}")
        
        tree.bind('<Double-Button-1>', on_double_click)
        
        # Button frame
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X)
        
        def toggle_enabled():
            """Toggle selected TM enabled state"""
            if not selected_tm_id[0]:
                messagebox.showinfo("No Selection", "Please select a TM to enable/disable")
                return
            
            tm = self.tm_database.get_tm(selected_tm_id[0])
            if tm:
                tm.enabled = not tm.enabled
                populate_tree()
                tree.selection_set(selected_tm_id[0])
                self.update_tm_dropdown()  # Update dropdown
        
        def view_tm():
            """View entries in selected TM"""
            if not selected_tm_id[0]:
                messagebox.showinfo("No Selection", "Please select a TM to view")
                return
            
            tm = self.tm_database.get_tm(selected_tm_id[0])
            if tm:
                self.show_tm_entries(tm, dialog)
        
        def remove_tm():
            """Remove selected custom TM"""
            if not selected_tm_id[0]:
                messagebox.showinfo("No Selection", "Please select a TM to remove")
                return
            
            if selected_tm_id[0] in ['project', 'big_mama', 'main']:
                messagebox.showwarning("Cannot Remove", 
                    "Project TM and Big Mama cannot be removed.\n\n"
                    "You can clear their contents from the View dialog.")
                return
            
            tm = self.tm_database.get_tm(selected_tm_id[0])
            if tm:
                if messagebox.askyesno("Confirm Remove", 
                    f"Remove TM '{tm.name}'?\n\n"
                    f"This will delete {tm.get_entry_count()} entries.\n"
                    f"This cannot be undone."):
                    
                    self.tm_database.remove_custom_tm(selected_tm_id[0])
                    selected_tm_id[0] = None
                    populate_tree()
                    self.update_tm_dropdown()  # Update dropdown
                    self.log(f"Removed TM: {tm.name}")
        
        def import_tm():
            """Import new TM"""
            self.load_tm_file()
            populate_tree()
        
        def export_tm():
            """Export selected TM to TMX"""
            if not selected_tm_id[0]:
                messagebox.showinfo("No Selection", "Please select a TM to export")
                return
            
            tm = self.tm_database.get_tm(selected_tm_id[0])
            if tm and tm.get_entry_count() > 0:
                self.export_tm_to_tmx(tm)
            else:
                messagebox.showinfo("Empty TM", "Selected TM has no entries to export")
        
        # Left side buttons
        ttk.Button(button_frame, text="⚡ Enable/Disable", 
                  command=toggle_enabled).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="👁 View Entries", 
                  command=view_tm).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="🗑 Remove", 
                  command=remove_tm).pack(side=tk.LEFT, padx=2)
        
        ttk.Separator(button_frame, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=10)
        
        ttk.Button(button_frame, text="📥 Import TM", 
                  command=import_tm).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="📤 Export TM", 
                  command=export_tm).pack(side=tk.LEFT, padx=2)
        
        ttk.Button(button_frame, text="Close", 
                  command=dialog.destroy).pack(side=tk.RIGHT, padx=2)
        
        # Help text
        help_frame = ttk.Frame(main_frame)
        help_frame.pack(fill=tk.X, pady=(5, 0))
        
        help_text = ("💡 Tip: Double-click a TM to toggle enabled/disabled. "
                    "Only enabled TMs are searched for matches.")
        ttk.Label(help_frame, text=help_text, font=('Segoe UI', 8), 
                 foreground='gray').pack(anchor=tk.W)
    
    def show_tm_entries(self, tm: TM, parent):
        """Show entries in a specific TM"""
        dialog = tk.Toplevel(parent)
        dialog.title(f"TM Entries: {tm.name}")
        dialog.geometry("800x500")
        dialog.transient(parent)
        
        main_frame = ttk.Frame(dialog, padding=10)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Header
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(header_frame, text=f"📖 {tm.name}", 
                 font=('Segoe UI', 11, 'bold')).pack(anchor=tk.W)
        
        info_text = f"Entries: {tm.get_entry_count()} | "
        info_text += f"Status: {'Enabled' if tm.enabled else 'Disabled'} | "
        info_text += f"Mode: {'Read-only' if tm.read_only else 'Editable'}"
        ttk.Label(header_frame, text=info_text).pack(anchor=tk.W)
        
        # Entries list
        list_frame = ttk.Frame(main_frame)
        list_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        columns = ("#", "source", "target")
        tree = ttk.Treeview(list_frame, columns=columns, show='headings')
        
        tree.heading("#", text="#")
        tree.heading("source", text="Source")
        tree.heading("target", text="Target")
        
        tree.column("#", width=50, anchor=tk.CENTER)
        tree.column("source", width=350, anchor=tk.W)
        tree.column("target", width=350, anchor=tk.W)
        
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)
        
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Populate
        for idx, (source, target) in enumerate(tm.entries.items(), 1):
            tree.insert("", tk.END, values=(idx, source, target))
        
        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X)
        
        def clear_tm():
            """Clear all entries"""
            if tm.read_only:
                messagebox.showwarning("Read-Only", "This TM is read-only and cannot be modified")
                return
            
            if messagebox.askyesno("Confirm Clear", 
                f"Clear all {tm.get_entry_count()} entries from '{tm.name}'?\n\n"
                "This cannot be undone."):
                tm.entries.clear()
                tree.delete(*tree.get_children())
                self.log(f"Cleared TM: {tm.name}")
        
        if not tm.read_only:
            ttk.Button(button_frame, text="🗑 Clear All Entries", 
                      command=clear_tm).pack(side=tk.LEFT, padx=2)
        
        ttk.Button(button_frame, text="Close", 
                  command=dialog.destroy).pack(side=tk.RIGHT)
    
    def export_tm_to_tmx(self, tm: TM):
        """Export TM to TMX file"""
        filepath = filedialog.asksaveasfilename(
            title=f"Export {tm.name}",
            defaultextension=".tmx",
            initialfile=f"{tm.name}.tmx",
            filetypes=[("TMX Files", "*.tmx"), ("All Files", "*.*")]
        )
        
        if not filepath:
            return
        
        try:
            # Create TMX structure
            tmx = ET.Element('tmx', version="1.4")
            header = ET.SubElement(tmx, 'header', 
                                  creationtool="Supervertaler",
                                  creationtoolversion="3.0.0",
                                  datatype="PlainText",
                                  segtype="sentence",
                                  adminlang="en-us",
                                  srclang=tm.metadata.get('source_lang', 'en'),
                                  o_tmf="Supervertaler")
            
            body = ET.SubElement(tmx, 'body')
            
            src_lang = tm.metadata.get('source_lang', 'en')
            tgt_lang = tm.metadata.get('target_lang', 'nl')
            
            for source, target in tm.entries.items():
                tu = ET.SubElement(body, 'tu')
                
                # Source TUV
                tuv_src = ET.SubElement(tu, 'tuv')
                tuv_src.set('{http://www.w3.org/XML/1998/namespace}lang', src_lang)
                seg_src = ET.SubElement(tuv_src, 'seg')
                seg_src.text = source
                
                # Target TUV
                tuv_tgt = ET.SubElement(tu, 'tuv')
                tuv_tgt.set('{http://www.w3.org/XML/1998/namespace}lang', tgt_lang)
                seg_tgt = ET.SubElement(tuv_tgt, 'seg')
                seg_tgt.text = target
            
            # Write to file
            tree = ET.ElementTree(tmx)
            ET.indent(tree, space='  ')
            tree.write(filepath, encoding='utf-8', xml_declaration=True)
            
            messagebox.showinfo("Export Success", 
                f"Exported {tm.get_entry_count()} entries to TMX file")
            self.log(f"✓ Exported TM '{tm.name}' to {os.path.basename(filepath)}")
            
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export TM:\n{str(e)}")
            self.log(f"✗ TM export failed: {str(e)}")
    
    # === Tracked Changes Management ===
    
    def load_tracked_changes_docx(self):
        """Load tracked changes from a DOCX file"""
        filepath = filedialog.askopenfilename(
            title="Select DOCX File with Tracked Changes",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if filepath:
            if self.tracked_changes_agent.load_docx_changes(filepath):
                count = self.tracked_changes_agent.get_entry_count()
                # Update status label if it exists
                if hasattr(self, 'tracked_changes_status_label'):
                    files = len(self.tracked_changes_agent.files_loaded)
                    self.tracked_changes_status_label.config(
                        text=f"✅ {count} changes loaded from {files} file(s)",
                        fg='green'
                    )
                messagebox.showinfo("Tracked Changes Loaded", 
                                  f"Successfully loaded tracked changes!\n\n"
                                  f"Change pairs: {count}\n"
                                  f"Files: {len(self.tracked_changes_agent.files_loaded)}\n\n"
                                  f"Ready for analysis in the 📊 Changes tab!")
                
                # Automatically open the browse window
                self.browse_tracked_changes()
    
    def load_tracked_changes_tsv(self):
        """Load tracked changes from a TSV file"""
        filepath = filedialog.askopenfilename(
            title="Select TSV File with Tracked Changes",
            filetypes=[("TSV Files", "*.tsv"), ("Text Files", "*.txt"), ("All Files", "*.*")]
        )
        
        if filepath:
            if self.tracked_changes_agent.load_tsv_changes(filepath):
                count = self.tracked_changes_agent.get_entry_count()
                messagebox.showinfo("Tracked Changes Loaded", 
                                  f"Successfully loaded tracked changes!\\n\\n"
                                  f"Change pairs: {count}\\n"
                                  f"Files: {len(self.tracked_changes_agent.files_loaded)}\\n\\n"
                                  f"These editing patterns will be provided to AI as examples.")
    
    def clear_tracked_changes(self):
        """Clear all loaded tracked changes"""
        if not self.tracked_changes_agent.change_data:
            messagebox.showinfo("No Changes", "No tracked changes are currently loaded.")
            return
        
        count = self.tracked_changes_agent.get_entry_count()
        if messagebox.askyesno("Clear Tracked Changes", 
                              f"Clear all {count} tracked change pairs?\\n\\n"
                              f"Files loaded: {len(self.tracked_changes_agent.files_loaded)}"):
            self.tracked_changes_agent.clear_changes()
            # Update status label if it exists
            if hasattr(self, 'tracked_changes_status_label'):
                self.tracked_changes_status_label.config(
                    text="No tracked changes loaded",
                    fg='gray'
                )
            messagebox.showinfo("Cleared", "All tracked changes have been cleared.")
    
    def browse_tracked_changes(self):
        """Browse and search tracked changes with AI-powered export"""
        if not self.tracked_changes_agent.change_data:
            messagebox.showinfo("No Changes", 
                              "No tracked changes loaded yet.\\n\\n"
                              "Load a DOCX file with tracked changes or TSV file first:\\n"
                              "Translate → Load Tracked Changes...")
            return
        
        # Use the TrackedChangesBrowser class with parent_app reference
        if not hasattr(self, 'tracked_changes_browser') or self.tracked_changes_browser is None:
            # Create a simple log queue adapter for v3 (uses direct log method)
            class LogQueueAdapter:
                def __init__(self, log_func):
                    self.log_func = log_func
                
                def put(self, message):
                    if self.log_func:
                        self.log_func(message)
            
            log_adapter = LogQueueAdapter(self.log)
            
            self.tracked_changes_browser = TrackedChangesBrowser(
                self.root,
                self.tracked_changes_agent,
                parent_app=self,
                log_queue=log_adapter
            )
        
        self.tracked_changes_browser.show_browser()
    
    # === Context Building ===
    
    def get_full_document_context(self, include_translations=False):
        """
        Build full document context string with all segments.
        This provides the AI with complete document understanding for better translation quality.
        
        Args:
            include_translations: If True, include existing translations alongside source text
        
        Returns:
            str: Formatted document context with numbered segments
        """
        if not self.segments:
            return ""
        
        context_lines = []
        
        for segment in self.segments:
            # Format: "[Segment ID]. Source text"
            context_lines.append(f"{segment.id}. {segment.source}")
            
            # Optionally include translation if available (helps with consistency)
            if include_translations and segment.target:
                context_lines.append(f"   → {segment.target}")
        
        return "\n".join(context_lines)
    
    def get_cat_tool_tag_instructions(self):
        """
        Get CAT-tool-specific tag handling instructions based on the source file type.
        Returns detailed instructions if a bilingual CAT file is loaded, empty string otherwise.
        """
        # Check which CAT tool format is loaded
        if hasattr(self, 'trados_source_file') and self.trados_source_file:
            return """
**CRITICAL: TRADOS STUDIO TAG PRESERVATION**

The source text contains Trados Studio formatting tags. These MUST be preserved exactly.

Tag Format:
• Paired tags: <NUMBER>content</NUMBER> (e.g., <13>Figure 1</13>, <226>base beam</226>)
• Self-closing tags: <NUMBER/> (e.g., <231/>, <234/>)
• Nested tags: <233>coupling (3) <231/></233>

MANDATORY RULES:
1. Preserve tag numbers exactly - <13> stays <13>, never becomes <14>
2. Maintain tag pairing - every <N> must have closing </N>
3. Keep self-closing syntax - <231/> must NOT become <231></231>
4. Place tags semantically - wrap the SAME meaning in target as in source
5. Never omit tags - if source has 3 tags, target MUST have exactly 3 tags
6. Never create new tags - only use tags present in source
7. Tags move with their content when word order changes

Examples:
• Simple: "<13>Figuur 1</13> toont..." → "<13>Figure 1</13> shows..."
• With ref: "een <226>basisligger (1)</226> geschikt" → "a <226>base beam</226> (1) suitable"
• Nested: "<233>koppeling (3) <231/></233>met" → "<233>coupling (3) <231/></233>with"
• Multiple: "<240>eerste koppeling</240> aan <243>matrasbeugel</243>" → "<240>first coupling</240> to <243>mattress bracket</243>"

VALIDATION: Count tags in source and target - they must match exactly!
"""
        
        elif hasattr(self, 'memoq_source_file') and self.memoq_source_file:
            return """
**CRITICAL: MEMOQ TAG PRESERVATION**

The source text contains memoQ formatting tags. These MUST be preserved exactly.

Tag Format:
• Opening tags: [N} (e.g., [1}, [3}, [5})
• Closing tags: {N] (e.g., {2], {4], {6])
• Asymmetric brackets/braces - opening uses [}, closing uses {]

MANDATORY RULES:
1. Preserve tag numbers exactly - [1} stays [1}, {2] stays {2]
2. Maintain tag pairing - every [N} must have matching {N]
3. Keep asymmetric syntax - NEVER convert [1} to {1} or {2] to [2]
4. Place tags semantically - wrap the SAME meaning in target as in source
5. Never omit tags - if source has 4 tags, target MUST have exactly 4 tags
6. Never create new tags - only use tags present in source
7. Tags move with their content when word order changes

Examples:
• Simple: "[1}De uitvoer{2]" → "[1}The exports{2]"
• Multiple: "[1}De uitvoer{2] [3}stelt niets voor{4]" → "[1}Exports{2] [3}mean nothing{4]"
• Nested: "[1}De [3}belangrijke{4] uitvoer{2]" → "[1}The [3}important{4] exports{2]"

VALIDATION: Count opening [N} and closing {N] tags - they must match source exactly!
"""
        
        elif hasattr(self, 'cafetran_source_file') and self.cafetran_source_file:
            return """
**CRITICAL: CAFETRAN TAG PRESERVATION**

The source text contains CafeTran pipe-delimited formatting tags. These MUST be preserved exactly.

Tag Format:
• Pipe symbols mark formatted text: |formatted content|
• Content between pipes is bold, italic, underlined, or specially formatted
• Pipes work in pairs - opening | and closing |

MANDATORY RULES:
1. Preserve ALL pipe symbols exactly - every opening | must have closing |
2. Content between pipes represents formatting (not literal pipes)
3. Place pipes semantically - wrap the SAME meaning in target as in source
4. Never omit pipes - if source has 4 pipes (2 pairs), target MUST have exactly 4 pipes
5. Never add pipes - only use pipes present in source
6. Pipes move with their content when word order changes
7. Proper names, technical terms, or formatting often appear between pipes

Examples:
• Simple: "He debuted against |Juventus FC| in 2001" → "Hij debuteerde tegen |Juventus FC| in 2001"
• Multiple: "The |bold text| and |italic text| here" → "De |vetgedrukte tekst| en |cursieve tekst| hier"
• Technical: "Using |ISO 9001| standard" → "Gebruik van |ISO 9001| norm"

VALIDATION: Count pipe symbols in source and target - they must match exactly (always even number)!
"""
        
        else:
            # No bilingual CAT file loaded, return empty string
            return ""
    
    def translate_current_segment(self):
        """Translate the currently selected segment using LLM (with TM lookup first)"""
        if not self.current_segment:
            messagebox.showwarning("No Selection", "Please select a segment to translate")
            return
        
        segment = self.current_segment
        
        # First, check for exact TM match
        exact_match = self.tm_agent.get_exact_match(segment.source)
        if exact_match:
            if messagebox.askyesno("TM Match Found", 
                                  f"Found exact translation in TM:\n\n{exact_match}\n\n"
                                  f"Use this translation?"):
                segment.target = exact_match
                segment.status = "translated"
                segment.modified = True
                self.modified = True
                self.load_segments_to_grid()
                self.update_progress()
                self.log(f"✓ Segment #{segment.id} translated from TM (100% match)")
                return
        
        # Check for fuzzy matches
        fuzzy_matches = self.tm_agent.get_fuzzy_matches(segment.source, max_matches=3)
        if fuzzy_matches:
            match_info = "\n\n".join([
                f"Match {i+1} ({int(similarity*100)}%):\n  Source: {src}\n  Target: {tgt}"
                for i, (src, tgt, similarity) in enumerate(fuzzy_matches)
            ])
            
            if messagebox.askyesno("Fuzzy TM Matches", 
                                  f"Found similar translations in TM:\n\n{match_info}\n\n"
                                  f"Continue with AI translation?"):
                pass  # User chose to proceed with AI
            else:
                return  # User cancelled
        
        # Check if API key is available
        # Normalize provider name (gemini -> google for API keys)
        api_key_name = "google" if self.current_llm_provider == "gemini" else self.current_llm_provider
        if not self.api_keys.get(api_key_name):
            messagebox.showerror("API Key Missing", 
                               f"Please configure your {self.current_llm_provider.upper()} API key in Translate → API Settings")
            return
        
        # Check if library is available
        if self.current_llm_provider == "openai" and not OPENAI_AVAILABLE:
            messagebox.showerror("Library Missing", "OpenAI library not installed. Run: pip install openai")
            return
        elif self.current_llm_provider == "claude" and not ANTHROPIC_AVAILABLE:
            messagebox.showerror("Library Missing", "Anthropic library not installed. Run: pip install anthropic")
            return
        elif self.current_llm_provider == "gemini" and not GEMINI_AVAILABLE:
            messagebox.showerror("Library Missing", "Google GenAI library not installed. Run: pip install google-generativeai")
            return
        
        # Build prompt with full document context (if enabled)
        prompt_parts = []
        
        # 1. System prompt (with language variables replaced) - USE CONTEXT-AWARE PROMPT
        system_prompt = self.get_context_aware_prompt(mode="single")
        system_prompt = system_prompt.replace("{{SOURCE_LANGUAGE}}", self.source_language)
        system_prompt = system_prompt.replace("{{TARGET_LANGUAGE}}", self.target_language)
        system_prompt = system_prompt.replace("{{SOURCE_TEXT}}", segment.source)
        prompt_parts.append(system_prompt)
        
        # 2. Add custom instructions if provided (skip if just placeholder)
        custom_instructions = self.custom_instructions_text.get('1.0', tk.END).strip()
        if custom_instructions and not self.is_custom_instructions_placeholder(custom_instructions):
            prompt_parts.append("\n**SPECIAL INSTRUCTIONS FOR THIS PROJECT:**")
            prompt_parts.append(custom_instructions)
        
        # 2b. Add CAT-tool-specific tag handling instructions (automatic based on source file)
        cat_tag_instructions = self.get_cat_tool_tag_instructions()
        if cat_tag_instructions:
            prompt_parts.append(cat_tag_instructions)
            self.log(f"  Including CAT tool tag preservation instructions (auto-detected)")
        
        # 3. Add tracked changes context (if available)
        if self.tracked_changes_agent.change_data:
            relevant_changes = self.tracked_changes_agent.find_relevant_changes([segment.source], max_changes=10)
            if relevant_changes:
                tracked_context = format_tracked_changes_context(relevant_changes, max_length=1000)
                prompt_parts.append("\n" + tracked_context)
                self.log(f"  Including {len(relevant_changes)} relevant tracked changes as examples")
        
        # 4. Add surrounding segments context (user-configurable)
        try:
            num_surrounding = int(self.surrounding_segments_var.get())
            if num_surrounding > 0:
                current_idx = self.segments.index(segment)
                
                # Get surrounding segments
                start_idx = max(0, current_idx - num_surrounding)
                end_idx = min(len(self.segments), current_idx + num_surrounding + 1)
                surrounding = self.segments[start_idx:end_idx]
                
                # Build context string
                context_parts = []
                for i, seg in enumerate(surrounding):
                    actual_idx = start_idx + i
                    if actual_idx == current_idx:
                        context_parts.append(f">>> {seg.id}. {seg.source} <<<  [TRANSLATE THIS]")
                    else:
                        context_parts.append(f"{seg.id}. {seg.source}")
                        if seg.target:
                            context_parts.append(f"    → {seg.target}")
                
                prompt_parts.append("\n**SURROUNDING SEGMENTS FOR CONTEXT:**")
                prompt_parts.append("(The segment marked with >>> <<< is the one to translate)\n")
                prompt_parts.append("\n".join(context_parts))
                
                self.log(f"  Including {len(surrounding)} surrounding segments ({num_surrounding} before/after)")
        except (ValueError, AttributeError):
            pass  # Skip if setting is invalid
        
        # 5. Full document context removed from single-segment translation to reduce API costs
        # (Context is only used in batch translation mode when explicitly enabled)
        
        # 6. Specify which segment to translate
        prompt_parts.append(f"\n**TEXT TO TRANSLATE:**")
        prompt_parts.append(segment.source)
        prompt_parts.append("\n**YOUR TRANSLATION (provide ONLY the translated text, no numbering or labels):**")
        
        # Combine all parts
        prompt = "\n".join(prompt_parts)
        
        self.log(f"🤖 Translating segment #{segment.id} using {self.current_llm_provider}/{self.current_llm_model}...")
        
        try:
            # Call appropriate API
            if self.current_llm_provider == "openai":
                translation = self.call_openai_api(prompt)
            elif self.current_llm_provider == "claude":
                translation = self.call_claude_api(prompt)
            elif self.current_llm_provider == "gemini":
                translation = self.call_gemini_api(prompt)
            else:
                raise ValueError(f"Unknown provider: {self.current_llm_provider}")
            
            # Update segment
            segment.target = translation.strip()
            segment.status = "translated"
            segment.modified = True
            
            # Add to TM agent (new and improved)
            self.tm_agent.add_entry(segment.source, segment.target)
            
            # Also add to legacy translation memory for backward compatibility
            self.translation_memory.append({
                "source": segment.source,
                "target": segment.target
            })
            
            # Update UI
            self.modified = True
            self.load_segments_to_grid()
            self.update_progress()
            self.log(f"✓ Segment #{segment.id} translated successfully (added to TM)")
            
        except Exception as e:
            messagebox.showerror("Translation Error", f"Failed to translate: {e}")
            self.log(f"✗ Translation failed: {e}")
    
    def translate_all_untranslated(self):
        """Translate all untranslated segments using CHUNKED batch processing (like v2.4.1)"""
        import math
        
        untranslated = [seg for seg in self.segments if not seg.target or seg.status == "untranslated"]
        
        if not untranslated:
            messagebox.showinfo("Complete", "All segments are already translated!")
            return
        
        # Check API key
        api_key_name = "google" if self.current_llm_provider == "gemini" else self.current_llm_provider
        if not self.api_keys.get(api_key_name):
            messagebox.showerror("API Key Missing", 
                               f"Please configure your {self.current_llm_provider.upper()} API key in Translate → API Settings")
            return
        
        # Get chunk size
        try:
            chunk_size = int(self.chunk_size_var.get())
            if chunk_size < 1:
                raise ValueError()
        except:
            messagebox.showerror("Invalid Setting", "Batch size must be a positive number!")
            return
        
        num_chunks = math.ceil(len(untranslated) / chunk_size)
        
        if not messagebox.askyesno("Confirm Batch Translation", 
                                   f"Translate {len(untranslated)} untranslated segments?\n\n"
                                   f"Provider: {self.current_llm_provider}/{self.current_llm_model}\n"
                                   f"Batch Size: {chunk_size} segments per API call\n"
                                   f"API Calls: ~{num_chunks} chunks\n"
                                   f"Context: {'Enabled' if self.use_context_var.get() else 'Disabled'}\n\n"
                                   f"This may take several minutes."):
            return
        
        # Create progress dialog
        progress_dialog = tk.Toplevel(self.root)
        progress_dialog.title("Batch Translation (Chunked)")
        progress_dialog.geometry("500x220")
        progress_dialog.transient(self.root)
        progress_dialog.grab_set()
        
        tk.Label(progress_dialog, text=f"Translating in batches of {chunk_size} segments...", 
                font=('Segoe UI', 12, 'bold')).pack(pady=10)
        
        progress_var = tk.DoubleVar()
        progress_bar = ttk.Progressbar(progress_dialog, variable=progress_var, 
                                      maximum=num_chunks, length=400)
        progress_bar.pack(pady=10)
        
        status_label = tk.Label(progress_dialog, text="Starting...", 
                               font=('Segoe UI', 10))
        status_label.pack(pady=5)
        
        cancel_var = tk.BooleanVar(value=False)
        
        def cancel_translation():
            cancel_var.set(True)
        
        tk.Button(progress_dialog, text="Cancel", command=cancel_translation,
                 bg='#F44336', fg='white').pack(pady=10)
        
        # Process in chunks
        successful = 0
        failed = 0
        
        # Get full document context ONCE if enabled
        full_context = None
        if self.use_context_var.get():
            all_sources = [seg.source for seg in self.segments]
            full_context = "\n".join([f"{i+1}. {src}" for i, src in enumerate(all_sources)])
            self.log(f"✓ Full document context prepared ({len(self.segments)} segments)")
        
        for chunk_idx in range(num_chunks):
            if cancel_var.get():
                self.log(f"⚠ Batch translation cancelled by user ({successful}/{len(untranslated)} completed)")
                break
            
            # Get chunk of segments
            start_idx = chunk_idx * chunk_size
            end_idx = min((chunk_idx + 1) * chunk_size, len(untranslated))
            chunk_segments = untranslated[start_idx:end_idx]
            
            status_label.config(text=f"Processing chunk {chunk_idx+1}/{num_chunks} ({len(chunk_segments)} segments)...")
            progress_var.set(chunk_idx)
            progress_dialog.update()
            
            # Check TM for exact matches first
            segments_needing_llm = []
            for segment in chunk_segments:
                if self.check_tm_var.get():
                    exact_match = self.tm_agent.get_exact_match(segment.source)
                    if exact_match:
                        segment.target = exact_match
                        segment.status = "translated"
                        segment.modified = True
                        self.log(f"✓ Segment #{segment.id} from TM (100% match)")
                        successful += 1
                        continue
                segments_needing_llm.append(segment)
            
            if not segments_needing_llm:
                continue  # All from TM
            
            try:
                # Build BATCH prompt with multiple segments
                prompt_parts = []
                
                # System prompt
                system_prompt = self.get_context_aware_prompt(mode="batch_docx")
                system_prompt = system_prompt.replace("{{SOURCE_LANGUAGE}}", self.source_language)
                system_prompt = system_prompt.replace("{{TARGET_LANGUAGE}}", self.target_language)
                prompt_parts.append(system_prompt)
                
                # Custom instructions
                custom_instructions = self.custom_instructions_text.get('1.0', tk.END).strip()
                if custom_instructions and not self.is_custom_instructions_placeholder(custom_instructions):
                    prompt_parts.append("\n**SPECIAL INSTRUCTIONS FOR THIS PROJECT:**")
                    prompt_parts.append(custom_instructions)
                
                # CAT-tool-specific tag handling (automatic based on source file)
                cat_tag_instructions = self.get_cat_tool_tag_instructions()
                if cat_tag_instructions:
                    prompt_parts.append(cat_tag_instructions)
                
                # Full document context (if enabled)
                if full_context:
                    prompt_parts.append("\n**FULL DOCUMENT CONTEXT FOR REFERENCE:**")
                    prompt_parts.append("(All segments in document - use for terminology consistency)\n")
                    prompt_parts.append(full_context)
                
                # Add the segments to translate in this chunk
                prompt_parts.append(f"\n**SEGMENTS TO TRANSLATE ({len(segments_needing_llm)} in this batch):**")
                prompt_parts.append("Provide ONLY the translations, one per line, in the same order. NO explanations, NO segment numbers, NO labels.\n")
                
                for seg in segments_needing_llm:
                    prompt_parts.append(f"{seg.id}. {seg.source}")
                
                prompt_parts.append("\n**YOUR TRANSLATIONS (one per line, in order):**")
                
                prompt = "\n".join(prompt_parts)
                
                # Call API
                self.log(f"🤖 Translating chunk {chunk_idx+1}/{num_chunks} ({len(segments_needing_llm)} segments)...")
                
                if self.current_llm_provider == "openai":
                    response = self.call_openai_api(prompt)
                elif self.current_llm_provider == "claude":
                    response = self.call_claude_api(prompt)
                elif self.current_llm_provider == "gemini":
                    response = self.call_gemini_api(prompt)
                else:
                    raise ValueError(f"Unknown provider: {self.current_llm_provider}")
                
                # Parse response - should be one translation per line
                translations = [line.strip() for line in response.strip().split('\n') if line.strip()]
                
                # Match translations to segments
                if len(translations) != len(segments_needing_llm):
                    self.log(f"⚠ Warning: Expected {len(segments_needing_llm)} translations, got {len(translations)}")
                
                for i, segment in enumerate(segments_needing_llm):
                    if i < len(translations):
                        translation = translations[i]
                        # Clean up common formatting (remove numbers, dots at start)
                        translation = re.sub(r'^\d+[\.\)]\s*', '', translation)
                        
                        segment.target = translation
                        segment.status = "translated"
                        segment.modified = True
                        
                        # Add to TM
                        self.tm_agent.add_entry(segment.source, segment.target)
                        self.translation_memory.append({"source": segment.source, "target": segment.target})
                        
                        successful += 1
                        self.log(f"✓ Segment #{segment.id} translated")
                    else:
                        failed += 1
                        self.log(f"✗ No translation received for segment #{segment.id}")
                
                self.log(f"✓ Chunk {chunk_idx+1}/{num_chunks} complete")
                
            except Exception as e:
                self.log(f"✗ Failed to translate chunk {chunk_idx+1}: {e}")
                failed += len(segments_needing_llm)
        
        # Update UI
        self.modified = True
        self.load_segments_to_grid()
        self.update_progress()
        
        progress_dialog.destroy()
        
        # Show summary
        messagebox.showinfo("Batch Translation Complete", 
                          f"Successfully translated: {successful}\n"
                          f"Failed: {failed}\n"
                          f"Total processed: {len(untranslated)}\n\n"
                          f"Used {chunk_idx+1} API calls (batch size: {chunk_size})")
        
        self.log(f"✓ Batch translation complete: {successful} successful, {failed} failed, {chunk_idx+1} API calls")

    
    def call_openai_api(self, prompt: str) -> str:
        """Call OpenAI API"""
        from openai import OpenAI
        client = OpenAI(api_key=self.api_keys["openai"])
        
        response = client.chat.completions.create(
            model=self.current_llm_model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3
        )
        
        return response.choices[0].message.content
    
    def call_claude_api(self, prompt: str) -> str:
        """Call Anthropic Claude API"""
        client = anthropic.Anthropic(api_key=self.api_keys["claude"])
        
        response = client.messages.create(
            model=self.current_llm_model,
            max_tokens=8192,  # Increased for potentially longer translations
            messages=[{"role": "user", "content": prompt}]
        )
        
        return response.content[0].text
    
    def call_gemini_api(self, prompt: str) -> str:
        """Call Google Gemini API"""
        genai.configure(api_key=self.api_keys["google"])
        model = genai.GenerativeModel(self.current_llm_model)
        
        response = model.generate_content(prompt)
        return response.text
    
    def add_sash_grip(self):
        """Add visual grip indicator to the sash between content and log"""
        try:
            # Get sash coordinates
            sash_coord = self.main_paned.sash_coord(0)
            if sash_coord:
                x, y = sash_coord
                # Create a small frame with grip pattern on the sash
                grip = tk.Frame(self.main_paned, bg='#999999', height=6, cursor='sb_v_double_arrow')
                grip.place(x=0, y=y-3, relwidth=1, height=6)
                
                # Add dotted pattern for visual feedback
                center_y = 2
                for i in range(0, 60, 4):  # Create dots across the grip
                    dot = tk.Frame(grip, bg='#666666', width=2, height=2)
                    dot.place(x=i, y=center_y)
                
                # Bind drag events to the grip frame
                grip.bind('<Button-1>', lambda e: self.main_paned.event_generate('<Button-1>', x=e.x, y=y))
                grip.bind('<B1-Motion>', lambda e: self.main_paned.sash_place(0, e.x_root, e.y_root))
        except:
            pass  # Silently fail if sash not ready yet
    
    def show_user_guide(self):
        """Open user guide in browser or show file"""
        import webbrowser
        import os
        
        # Try to open USER_GUIDE.md
        guide_path = os.path.join(os.path.dirname(__file__), "USER_GUIDE.md")
        if os.path.exists(guide_path):
            webbrowser.open(f"file://{guide_path}")
        else:
            messagebox.showinfo("User Guide", 
                              "User guide not found.\nPlease check the documentation folder.")
    
    def show_changelog(self):
        """Open changelog in browser or show file"""
        import webbrowser
        import os
        
        # Try to open CHANGELOG-CAT.md for v3
        changelog_path = os.path.join(os.path.dirname(__file__), "CHANGELOG-CAT.md")
        if os.path.exists(changelog_path):
            webbrowser.open(f"file://{changelog_path}")
        else:
            messagebox.showinfo("Changelog", 
                              "Changelog not found.\nPlease check the documentation folder.")
    
    def show_about(self):
        """Show about dialog"""
        about_text = f"""Supervertaler {APP_VERSION}
AI-Powered Computer-Aided Translation Tool

Features:
• Multi-LLM support (OpenAI, Claude, Gemini)
• Translation Memory & Glossaries
• Grid/List/Document views
• CafeTran & memoQ compatibility
• Tracked Changes Analysis
• Custom prompts & instructions

Author: Michael Beijer + AI Assistant
GitHub: github.com/michaelbeijer/Supervertaler

© 2025 - Licensed under MIT"""
        
        messagebox.showinfo("About Supervertaler", about_text)
    
    def on_closing(self):
        """Handle window closing"""
        if self.modified:
            response = messagebox.askyesnocancel(
                "Unsaved Changes",
                "Save project before closing?"
            )
            if response is True:
                self.save_project()
                self.root.destroy()
            elif response is False:
                self.root.destroy()
        else:
            self.root.destroy()


class FindReplaceDialog:
    """Find and replace dialog"""
    
    def __init__(self, parent, app):
        self.app = app
        self.window = tk.Toplevel(parent)
        self.window.title("Find and Replace")
        self.window.geometry("500x250")
        self.window.transient(parent)
        
        # Find
        tk.Label(self.window, text="Find:", font=('Segoe UI', 10, 'bold')).pack(anchor='w', padx=10, pady=(10,2))
        self.find_var = tk.StringVar()
        tk.Entry(self.window, textvariable=self.find_var, width=60).pack(padx=10, pady=2)
        
        # Replace
        tk.Label(self.window, text="Replace with:", font=('Segoe UI', 10, 'bold')).pack(anchor='w', padx=10, pady=(10,2))
        self.replace_var = tk.StringVar()
        tk.Entry(self.window, textvariable=self.replace_var, width=60).pack(padx=10, pady=2)
        
        # Options
        options_frame = tk.Frame(self.window)
        options_frame.pack(pady=10)
        
        self.match_case_var = tk.BooleanVar()
        tk.Checkbutton(options_frame, text="Match case", variable=self.match_case_var).pack(side='left', padx=5)
        
        self.search_in_var = tk.StringVar(value="target")
        tk.Label(options_frame, text="Search in:").pack(side='left', padx=(20,5))
        ttk.Combobox(options_frame, textvariable=self.search_in_var,
                    values=["source", "target", "both"], state='readonly', width=10).pack(side='left')
        
        # Buttons
        button_frame = tk.Frame(self.window)
        button_frame.pack(pady=10)
        
        tk.Button(button_frame, text="Find Next", command=self.find_next, width=12).pack(side='left', padx=5)
        tk.Button(button_frame, text="Replace", command=self.replace_current, width=12).pack(side='left', padx=5)
        tk.Button(button_frame, text="Replace All", command=self.replace_all, width=12,
                 bg='#FF9800', fg='white').pack(side='left', padx=5)
        
        # Results
        self.result_label = tk.Label(self.window, text="", fg='blue')
        self.result_label.pack(pady=5)
        
        self.current_match_index = -1
        self.matches = []
    
    def find_next(self):
        """Find next occurrence"""
        query = self.find_var.get()
        if not query:
            self.result_label.config(text="Please enter search text")
            return
        
        search_in = self.search_in_var.get()
        case_sensitive = self.match_case_var.get()
        
        # Search segments
        self.matches = []
        for seg in self.app.segments:
            source = seg.source if case_sensitive else seg.source.lower()
            target = seg.target if case_sensitive else seg.target.lower()
            query_cmp = query if case_sensitive else query.lower()
            
            found = False
            if search_in in ['source', 'both'] and query_cmp in source:
                found = True
            if search_in in ['target', 'both'] and query_cmp in target:
                found = True
            
            if found:
                self.matches.append(seg.id)
        
        if not self.matches:
            self.result_label.config(text="No matches found")
            return
        
        # Move to next match
        self.current_match_index = (self.current_match_index + 1) % len(self.matches)
        match_id = self.matches[self.current_match_index]
        
        # Select in grid
        for item in self.app.tree.get_children():
            values = self.app.tree.item(item, 'values')
            if int(values[0]) == match_id:
                self.app.tree.selection_set(item)
                self.app.tree.see(item)
                break
        
        self.result_label.config(
            text=f"Match {self.current_match_index + 1} of {len(self.matches)}"
        )
    
    def replace_current(self):
        """Replace current match"""
        if not self.app.current_segment:
            self.result_label.config(text="No segment selected")
            return
        
        find_text = self.find_var.get()
        replace_text = self.replace_var.get()
        
        if not find_text:
            return
        
        # Replace in target
        if self.match_case_var.get():
            new_target = self.app.current_segment.target.replace(find_text, replace_text)
        else:
            import re
            new_target = re.sub(re.escape(find_text), replace_text,
                              self.app.current_segment.target, flags=re.IGNORECASE)
        
        # Update
        self.app.target_text.delete('1.0', 'end')
        self.app.target_text.insert('1.0', new_target)
        self.app.save_current_segment()
        
        self.result_label.config(text="Replaced")
    
    def replace_all(self):
        """Replace all occurrences"""
        find_text = self.find_var.get()
        replace_text = self.replace_var.get()
        
        if not find_text:
            return
        
        search_in = self.search_in_var.get()
        count = 0
        
        for seg in self.app.segments:
            if search_in in ['target', 'both']:
                if self.match_case_var.get():
                    if find_text in seg.target:
                        seg.target = seg.target.replace(find_text, replace_text)
                        count += 1
                else:
                    import re
                    if re.search(re.escape(find_text), seg.target, re.IGNORECASE):
                        seg.target = re.sub(re.escape(find_text), replace_text,
                                          seg.target, flags=re.IGNORECASE)
                        count += 1
        
        # Refresh grid
        self.app.load_segments_to_grid()
        self.app.modified = True
        self.app.update_progress()
        
        self.result_label.config(text=f"Replaced {count} occurrences")


# Main
if __name__ == "__main__":
    root = tk.Tk()
    
    # Set window icon
    try:
        icon_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'assets', 'MB.ico')
        if os.path.exists(icon_path):
            root.iconbitmap(icon_path)
    except Exception as e:
        print(f"Could not load icon: {e}")
    
    app = Supervertaler(root)
    root.mainloop()
