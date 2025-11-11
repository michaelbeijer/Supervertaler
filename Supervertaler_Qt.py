"""
Supervertaler Qt Edition
========================
The ultimate companion tool for translators and writers.
Modern PyQt6 interface with specialised modules to handle any problem.
Version: 1.3.4 (AI Assistant Enhanced Prompt Generation)
Release Date: November 10, 2025
Framework: PyQt6

This is the modern edition of Supervertaler using PyQt6 framework.
For the classic tkinter edition, see Supervertaler_tkinter.py

Key Features:
- Complete Translation Matching: Termbase + TM + MT + Multi-LLM
- Google Cloud Translation API integration
- Multi-LLM Support: OpenAI GPT, Claude, Google Gemini
- 2-Layer Prompt Architecture (System + Custom Prompts) with AI Assistant
- AI Assistant with conversational interface for document analysis
- Universal Lookup with global hotkey (Ctrl+Alt+L)
- Modern theme system (6 themes + custom editor)
- AutoFingers automation for memoQ with TagCleaner module
- memoQ bilingual DOCX import/export
- SQLite-based translation memory with FTS5 search
- Professional TMX editor

Author: Michael Beijer
License: MIT
"""

# Version Information
__version__ = "1.3.4"
__phase__ = "6.4"
__release_date__ = "2025-11-10"
__edition__ = "Qt"

import sys
import json
import os
import subprocess
import atexit
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple
from dataclasses import dataclass, asdict
from datetime import datetime
import threading
import time  # For delays in Universal Lookup
import re

# Fix encoding for Windows console (UTF-8 support)
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

# External dependencies
import pyperclip  # For clipboard operations in Universal Lookup
from modules.universal_lookup import UniversalLookupEngine  # Universal Lookup engine
from modules.voice_dictation_lite import QuickDictationThread  # Voice dictation
from modules.statuses import (
    STATUSES,
    DEFAULT_STATUS,
    StatusDefinition,
    get_status,
    match_memoq_status,
    compose_memoq_status,
)


STATUS_ORDER = [
    "not_started",
    "pretranslated",
    "translated",
    "confirmed",
    "tr_confirmed",
    "proofread",
    "approved",
    "rejected",
]

STATUS_CYCLE = STATUS_ORDER

TRANSLATABLE_STATUSES = {"not_started", "pretranslated", "translated"}

# Check for PyQt6 and offer to install if missing
try:
    from PyQt6.QtWidgets import (
        QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
        QTableWidget, QTableWidgetItem, QHeaderView, QMenuBar, QMenu,
        QFileDialog, QMessageBox, QToolBar, QLabel, QComboBox,
        QPushButton, QSpinBox, QSplitter, QTextEdit, QStatusBar,
        QStyledItemDelegate, QInputDialog, QDialog, QLineEdit, QRadioButton,
        QButtonGroup, QDialogButtonBox, QTabWidget, QGroupBox, QGridLayout, QCheckBox,
        QProgressBar, QFormLayout, QTabBar, QPlainTextEdit, QAbstractItemDelegate,
        QFrame, QListWidget, QListWidgetItem, QStackedWidget, QTreeWidget, QTreeWidgetItem,
        QScrollArea, QSizePolicy
    )
    from PyQt6.QtCore import Qt, QSize, QTimer, pyqtSignal, QObject, QUrl
    from PyQt6.QtGui import QFont, QAction, QKeySequence, QIcon, QTextOption, QColor, QDesktopServices, QTextCharFormat, QTextCursor
except ImportError:
    print("PyQt6 not found. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "PyQt6"])
    print("✓ PyQt6 installed. Please restart the application.")
    sys.exit(0)


# ============================================================================
# PRIVATE DATA PROTECTION SYSTEM
# ============================================================================

# Check for .supervertaler.local file to enable private features (for developers only)
# This ensures that private data (API keys, personal projects, etc.) stays in user_data_private/
# which is .gitignored, preventing accidental upload to GitHub
ENABLE_PRIVATE_FEATURES = os.path.exists(
    os.path.join(os.path.dirname(os.path.abspath(__file__)), ".supervertaler.local")
)
if ENABLE_PRIVATE_FEATURES:
    print("[DEV MODE] Private features enabled (.supervertaler.local found)")
    print("[DEV MODE] Using 'user_data_private/' folder (git-ignored)")
else:
    print("[USER MODE] Using 'user_data/' folder")


# ============================================================================
# GLOBAL AHK PROCESS CLEANUP
# ============================================================================

# Global variable to track AHK process
_ahk_process = None

def cleanup_ahk_process():
    """Cleanup function to kill AHK process on exit"""
    global _ahk_process
    if _ahk_process:
        try:
            _ahk_process.terminate()
            _ahk_process.wait(timeout=1)
            print("[Universal Lookup] AHK process terminated on exit")
        except:
            try:
                _ahk_process.kill()
                print("[Universal Lookup] AHK process killed on exit")
            except:
                pass

# Register cleanup function to run on Python exit
atexit.register(cleanup_ahk_process)


# ============================================================================
# ENUMS
# ============================================================================

class LayoutMode:
    """Layout/view modes for the Project Editor"""
    GRID = "grid"       # Spreadsheet-like table (default)
    LIST = "list"       # List view with editor panel
    DOCUMENT = "document"  # Document flow view with clickable segments


# ============================================================================
# DATA MODELS
# ============================================================================

@dataclass
@dataclass
class Segment:
    """Translation segment (matches tkinter version format)"""
    id: int
    source: str
    target: str = ""
    status: str = DEFAULT_STATUS.key
    type: str = "para"  # para, heading, list_item, table_cell
    notes: str = ""  # Stored as “Comments” in UI
    match_percent: Optional[int] = None  # memoQ match score if provided
    memoQ_status: str = ""  # Raw memoQ status text
    locked: bool = False  # For compatibility with tkinter version
    paragraph_id: int = 0  # Group segments by paragraph for document flow
    style: str = "Normal"  # Heading 1, Heading 2, Title, Subtitle, Normal, etc.
    document_position: int = 0  # Position in original document
    is_table_cell: bool = False  # Whether this segment is in a table
    table_info: Optional[tuple] = None  # (table_idx, row_idx, cell_idx) if is_table_cell
    modified: bool = False  # Track if segment has been edited
    created_at: str = ""  # Creation timestamp
    modified_at: str = ""  # Last modification timestamp
    
    def __post_init__(self):
        """Initialize timestamps if not provided"""
        if not self.created_at:
            self.created_at = datetime.now().isoformat()
        if not self.modified_at:
            self.modified_at = datetime.now().isoformat()
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for JSON serialization"""
        return asdict(self)
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'Segment':
        """Create Segment from dictionary, ignoring unknown fields"""
        # Only use fields that the dataclass knows about
        valid_fields = {f.name for f in cls.__dataclass_fields__.values()}
        filtered_data = {k: v for k, v in data.items() if k in valid_fields}
        return cls(**filtered_data)


@dataclass
class Project:
    """Translation project"""
    name: str
    source_lang: str = "en"
    target_lang: str = "nl"
    segments: List[Segment] = None
    created: str = ""
    modified: str = ""
    prompt_settings: Dict[str, Any] = None  # Store active prompt settings
    
    def __post_init__(self):
        if self.segments is None:
            self.segments = []
        if not self.created:
            self.created = datetime.now().isoformat()
        if not self.modified:
            self.modified = datetime.now().isoformat()
        if self.prompt_settings is None:
            self.prompt_settings = {}
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for JSON serialization"""
        result = {
            'name': self.name,
            'source_lang': self.source_lang,
            'target_lang': self.target_lang,
            'segments': [seg.to_dict() for seg in self.segments],
            'created': self.created,
            'modified': self.modified
        }
        # Add prompt settings if they exist
        if hasattr(self, 'prompt_settings'):
            result['prompt_settings'] = self.prompt_settings
        return result
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'Project':
        """Create Project from dictionary"""
        segments = [Segment.from_dict(seg) for seg in data.get('segments', [])]
        
        # Handle missing name field (use filename or default)
        name = data.get('name', 'Untitled Project')
        
        project = cls(
            name=name,
            source_lang=data.get('source_lang', 'en'),
            target_lang=data.get('target_lang', 'nl'),
            segments=segments,
            created=data.get('created', ''),
            modified=data.get('modified', '')
        )
        # Store prompt settings if they exist
        if 'prompt_settings' in data:
            project.prompt_settings = data['prompt_settings']
        return project


# ============================================================================
# CUSTOM DELEGATES AND EDITORS
# ============================================================================

class GridTableEventFilter:
    """Mixin to pass keyboard shortcuts from editor to table"""
    pass


class GridTextEditor(QTextEdit):
    """Custom QTextEdit for grid cells that passes special shortcuts to parent"""
    
    table_widget = None  # Will be set by delegate
    assistance_panel = None  # Will be set by delegate
    
    def keyPressEvent(self, event):
        """Override to handle Ctrl+1-9, Ctrl+Up/Down shortcuts"""
        if event.modifiers() & Qt.KeyboardModifier.ControlModifier:
            # Ctrl+1 through Ctrl+9: Insert match by number
            # TODO: Update to work with new results_panels system
            if event.key() >= Qt.Key.Key_1 and event.key() <= Qt.Key.Key_9:
                # Disabled for now - needs update to work with results_panels
                pass
            # Ctrl+Up/Down: Send to grid for grid navigation
            elif event.key() in (Qt.Key.Key_Up, Qt.Key.Key_Down):
                if self.table_widget:
                    self.table_widget.keyPressEvent(event)
                    event.accept()
                    return
            # Ctrl+Enter: Insert actual line break
            elif event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
                super().keyPressEvent(event)
                event.accept()
                return
        
        # Shift+Enter: Insert line break (for multi-line content)
        if event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
            if event.modifiers() & Qt.KeyboardModifier.ShiftModifier:
                super().keyPressEvent(event)
                event.accept()
                return
        
        # Enter/Return: Don't insert, let delegate handle it (will close editor)
        if event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
            # Just ignore it - the table's default behavior will close the editor
            event.ignore()
            return
        
        # All other keys: Handle normally (including ESC for closing editor)
        super().keyPressEvent(event)


class ReadOnlyGridTextEditor(QTextEdit):
    """Read-only QTextEdit for source cells - allows easy text selection"""
    
    def __init__(self, text: str = "", parent=None):
        super().__init__(parent)
        self.setReadOnly(True)
        self.setPlainText(text)
        self.setWordWrapMode(QTextOption.WrapMode.WordWrap)
        self.setAcceptRichText(False)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        
        # Style to look like a normal cell with subtle selection
        # IMPORTANT: Use light gray background instead of transparent to avoid text visibility issues
        self.setStyleSheet("""
            QTextEdit {
                border: none;
                background-color: #f5f5f5;
                padding: 0px;
                color: black;
            }
            QTextEdit:focus {
                border: 1px solid #2196F3;
                background-color: white;
                color: black;
            }
            QTextEdit::selection {
                background-color: #D0E7FF;
                color: black;
            }
        """)
        
        # Set document margins to 0 for compact display
        doc = self.document()
        doc.setDocumentMargin(0)
        
        # Configure text option for minimal line spacing
        text_option = QTextOption()
        text_option.setWrapMode(QTextOption.WrapMode.WordWrap)
        doc.setDefaultTextOption(text_option)
        
        # Set minimum height to 0 - let content determine size
        self.setMinimumHeight(0)
        self.setMaximumHeight(16777215)  # Qt's max int
        
        # Store termbase matches for this cell
        self.termbase_matches = {}
    
    def sizeHint(self):
        """Return compact size based on content"""
        doc = self.document()
        ideal_width = self.width() if self.width() > 0 else 200
        doc.setTextWidth(ideal_width)
        height = int(doc.size().height())
        # Add minimal padding (2px total)
        height = max(height + 2, 1)
        current_width = self.width() if self.width() > 0 else 200
        return QSize(current_width, height)
    
    def mousePressEvent(self, event):
        """Allow text selection on click and trigger row selection"""
        super().mousePressEvent(event)
        
        # Find the table and row number by checking which row this widget belongs to
        if self.parent():
            try:
                table = self.parent()
                # Find which row this widget is in
                for row in range(table.rowCount()):
                    if table.cellWidget(row, 2) == self:  # Source is column 2
                        table.selectRow(row)
                        table.setCurrentCell(row, 2)
                        
                        # CRITICAL: Manually trigger on_cell_selected since signals aren't firing
                        # Find the main window and call the method directly
                        main_window = table.parent()
                        while main_window and not hasattr(main_window, 'on_cell_selected'):
                            main_window = main_window.parent()
                        if main_window and hasattr(main_window, 'on_cell_selected'):
                            main_window.on_cell_selected(row, 2, -1, -1)
                        break
            except Exception as e:
                print(f"Error triggering manual cell selection: {e}")
    
    def focusInEvent(self, event):
        """Select text when focused for easy copying and trigger row selection"""
        super().focusInEvent(event)
        # Don't auto-select - let user select manually
        
        # Find the table and row number by checking which row this widget belongs to
        if self.parent():
            try:
                table = self.parent()
                # Find which row this widget is in
                for row in range(table.rowCount()):
                    if table.cellWidget(row, 2) == self:  # Source is column 2
                        table.selectRow(row)
                        table.setCurrentCell(row, 2)
                        
                        # CRITICAL: Manually trigger on_cell_selected since signals aren't firing
                        # Find the main window and call the method directly
                        main_window = table.parent()
                        while main_window and not hasattr(main_window, 'on_cell_selected'):
                            main_window = main_window.parent()
                        if main_window and hasattr(main_window, 'on_cell_selected'):
                            main_window.on_cell_selected(row, 2, -1, -1)
                        break
            except Exception as e:
                print(f"Error triggering manual cell selection: {e}")


class EditableGridTextEditor(QTextEdit):
    """Editable QTextEdit for target cells - allows text selection and editing"""
    
    def __init__(self, text: str = "", parent=None, row: int = -1, table=None):
        super().__init__(parent)
        self.row = row  # Store row number for auto-selection
        self.table = table  # Store table reference
        self.setReadOnly(False)
        self.setPlainText(text)
        self.setWordWrapMode(QTextOption.WrapMode.WordWrap)
        self.setAcceptRichText(False)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        
        # Style to look like a normal cell with subtle selection
        # IMPORTANT: Use white background instead of transparent to avoid text visibility issues
        self.setStyleSheet("""
            QTextEdit {
                border: none;
                background-color: white;
                padding: 0px;
                color: black;
            }
            QTextEdit:focus {
                border: 1px solid #2196F3;
                background-color: white;
                color: black;
            }
            QTextEdit::selection {
                background-color: #D0E7FF;
                color: black;
            }
        """)
        
        # Set document margins to 0 for compact display
        doc = self.document()
        doc.setDocumentMargin(0)
        
        # Configure text option for minimal line spacing
        text_option = QTextOption()
        text_option.setWrapMode(QTextOption.WrapMode.WordWrap)
        doc.setDefaultTextOption(text_option)
        
        # Set minimum height to 0 - let content determine size
        self.setMinimumHeight(0)
        self.setMaximumHeight(16777215)  # Qt's max int
    
    def sizeHint(self):
        """Return compact size based on content"""
        doc = self.document()
        ideal_width = self.width() if self.width() > 0 else 200
        doc.setTextWidth(ideal_width)
        height = int(doc.size().height())
        # Add minimal padding (2px total)
        height = max(height + 2, 1)
        current_width = self.width() if self.width() > 0 else 200
        return QSize(current_width, height)
    
    def mousePressEvent(self, event):
        """Allow text selection on click and auto-select row"""
        super().mousePressEvent(event)
        # Auto-select the row when clicking in the target cell
        if self.table and self.row >= 0:
            self.table.selectRow(self.row)
            self.table.setCurrentCell(self.row, 3)  # Column 3 is Target
            
            # CRITICAL: Manually trigger on_cell_selected since signals aren't firing
            # Find the main window and call the method directly
            try:
                main_window = self.table.parent()
                while main_window and not hasattr(main_window, 'on_cell_selected'):
                    main_window = main_window.parent()
                if main_window and hasattr(main_window, 'on_cell_selected'):
                    main_window.on_cell_selected(self.row, 3, -1, -1)
            except Exception as e:
                print(f"Error triggering manual cell selection: {e}")
    
    def focusInEvent(self, event):
        """Ensure text remains visible when focused and auto-select row"""
        super().focusInEvent(event)
        # Ensure the widget is properly visible
        self.setVisible(True)
        self.show()
        # Auto-select the row when focusing the target cell
        if self.table and self.row >= 0:
            self.table.selectRow(self.row)
            self.table.setCurrentCell(self.row, 3)  # Column 3 is Target
            
            # CRITICAL: Manually trigger on_cell_selected since signals aren't firing
            # Find the main window and call the method directly
            try:
                main_window = self.table.parent()
                while main_window and not hasattr(main_window, 'on_cell_selected'):
                    main_window = main_window.parent()
                if main_window and hasattr(main_window, 'on_cell_selected'):
                    main_window.on_cell_selected(self.row, 3, -1, -1)
            except Exception as e:
                print(f"Error triggering manual cell selection: {e}")


class TermbaseHighlightWidget(QLabel):
    """Custom label widget that displays source text with termbase highlights and tooltips"""
    
    # Signal for term double-click
    term_double_clicked = None
    
    def __init__(self, source_text: str, termbase_matches: Optional[Dict[str, str]] = None, parent=None):
        """
        Args:
            source_text: The full source text to display
            termbase_matches: Dict of {term: translation} for highlighted terms
            parent: Parent widget
        """
        super().__init__(parent)
        self.source_text = source_text
        self.termbase_matches = termbase_matches if termbase_matches is not None else {}
        self.term_ranges = {}  # Store {term: (start_pos, end_pos)}
        self.setWordWrap(True)
        self.setup_display()
        self.setMouseTracking(True)
    
    def setup_display(self):
        """Create rich text display with highlighted terms"""
        html = self._create_highlighted_html()
        self.setText(html)
        self.setTextFormat(Qt.TextFormat.RichText)
    
    def _create_highlighted_html(self) -> str:
        """Create HTML with highlighted termbase matches"""
        text = self.source_text
        
        # Sort terms by length (longest first) to avoid overlaps
        sorted_terms = sorted(self.termbase_matches.keys(), key=len, reverse=True)
        
        # Build HTML with highlights
        html = text
        offset = 0
        
        for term in sorted_terms:
            # Case-insensitive search
            search_term = term.lower()
            search_html = html.lower()
            
            # Find all occurrences
            start = 0
            while True:
                idx = search_html.find(search_term, start)
                if idx == -1:
                    break
                
                # Extract original casing
                original_term = html[idx:idx + len(term)]
                
                # Replace with highlighted version
                highlighted = f'<span style="color: blue; font-weight: bold; text-decoration: underline; cursor: pointer;" class="termbase-match" data-term="{term}">{original_term}</span>'
                
                html = html[:idx] + highlighted + html[idx + len(term):]
                search_html = html.lower()
                
                start = idx + len(highlighted)
        
        return html
    
    def mouseMoveEvent(self, ev):
        """Show tooltip when hovering over highlighted terms"""
        # Find if cursor is over a highlighted term using text format
        # This is a simplified approach - shows generic tooltip
        super().mouseMoveEvent(ev)
    
    def mouseDoubleClickEvent(self, ev):
        """Handle double-click on highlighted terms"""
        # Find which term was clicked by getting the word at the position
        # Since QLabel doesn't provide cursor position easily, we look at the click
        # and try to extract the clicked word from the text
        
        # Get position in text (approximate)
        pos = ev.pos()
        
        # Try to find which term is near the click point
        # This is a simplified approach - in production would need better handling
        for term, translation in self.termbase_matches.items():
            # Check if this term contains the click area (very approximate)
            if self.term_double_clicked is not None and callable(self.term_double_clicked):
                # For now, if we get a double-click and there are matches,
                # this is a signal to insert the first/only match
                # Better implementation would need proper hit detection
                pass
        
        super().mouseDoubleClickEvent(ev)


class WordWrapDelegate(QStyledItemDelegate):
    """Custom delegate to enable word wrap when editing cells"""
    
    def __init__(self, assistance_panel=None, table_widget=None):
        super().__init__()
        self.assistance_panel = assistance_panel
        self.table_widget = table_widget
    
    def createEditor(self, parent, option, index):
        """Create a QTextEdit for multi-line editing with word wrap"""
        # Only use QTextEdit for Target column (column 3)
        if index.column() == 3:
            editor = GridTextEditor(parent)
            editor.assistance_panel = self.assistance_panel
            editor.table_widget = self.table_widget
            editor.setWordWrapMode(QTextOption.WrapMode.WordWrap)
            editor.setAcceptRichText(False)  # Plain text only
            editor.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
            editor.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            
            # Ensure the row is tall enough for editing
            table = parent.parent()
            if hasattr(table, 'resizeRowToContents'):
                # Schedule row resize after editor is shown
                from PyQt6.QtCore import QTimer
                QTimer.singleShot(0, lambda: table.resizeRowToContents(index.row()))
            
            return editor
        else:
            return super().createEditor(parent, option, index)
    
    def setEditorData(self, editor, index):
        """Load data into the editor"""
        if isinstance(editor, QTextEdit):
            text = index.model().data(index, Qt.ItemDataRole.EditRole)
            editor.setPlainText(text or "")
            # Ensure cursor is at start
            cursor = editor.textCursor()
            cursor.movePosition(cursor.MoveOperation.Start)
            editor.setTextCursor(cursor)
        else:
            super().setEditorData(editor, index)
    
    def setModelData(self, editor, model, index):
        """Save data from editor back to model"""
        if isinstance(editor, QTextEdit):
            text = editor.toPlainText()
            model.setData(index, text, Qt.ItemDataRole.EditRole)
        else:
            super().setModelData(editor, model, index)
    
    def updateEditorGeometry(self, editor, option, index):
        """Set the editor geometry to match the cell size"""
        if isinstance(editor, QTextEdit):
            # Make the editor fill the cell properly with some padding
            rect = option.rect
            editor.setGeometry(rect)
        else:
            super().updateEditorGeometry(editor, option, index)


# ============================================================================
# THEME EDITOR DIALOG
# ============================================================================

class ThemeEditorDialog(QDialog):
    """Dialog for editing and managing themes"""
    
    def __init__(self, parent, theme_manager):
        super().__init__(parent)
        self.theme_manager = theme_manager
        self.setWindowTitle("Theme Editor")
        self.setModal(True)
        self.resize(700, 600)
        
        self.setup_ui()
        self.load_themes()
    
    def setup_ui(self):
        """Create the UI"""
        layout = QVBoxLayout(self)
        
        # Theme selection
        theme_group = QGroupBox("Select Theme")
        theme_layout = QHBoxLayout()
        
        self.theme_combo = QComboBox()
        self.theme_combo.currentTextChanged.connect(self.on_theme_selected)
        theme_layout.addWidget(QLabel("Theme:"))
        theme_layout.addWidget(self.theme_combo, 1)
        
        self.apply_btn = QPushButton("✓ Apply")
        self.apply_btn.clicked.connect(self.apply_theme)
        theme_layout.addWidget(self.apply_btn)
        
        theme_group.setLayout(theme_layout)
        layout.addWidget(theme_group)
        
        # Color customization
        colors_group = QGroupBox("Customize Colors")
        colors_layout = QGridLayout()
        
        # Create color pickers for main colors
        self.color_buttons = {}
        color_configs = [
            ("window_bg", "Window Background", 0, 0),
            ("base", "Input Fields", 0, 1),
            ("button", "Buttons", 0, 2),
            ("text", "Text", 1, 0),
            ("highlight", "Highlight", 1, 1),
            ("border", "Borders", 1, 2),
            ("grid_header", "Table Headers", 2, 0),
            ("alternate_bg", "Alternate Rows", 2, 1),
            ("tm_exact", "100% TM Match", 3, 0),
            ("tm_high", "95-99% TM Match", 3, 1),
        ]
        
        for attr, label, row, col in color_configs:
            lbl = QLabel(label + ":")
            btn = QPushButton()
            btn.setMinimumHeight(30)
            btn.setProperty("color_attr", attr)
            btn.clicked.connect(lambda checked, a=attr: self.pick_color(a))
            self.color_buttons[attr] = btn
            
            colors_layout.addWidget(lbl, row * 2, col)
            colors_layout.addWidget(btn, row * 2 + 1, col)
        
        colors_group.setLayout(colors_layout)
        layout.addWidget(colors_group)
        
        # Custom theme actions
        custom_group = QGroupBox("Custom Themes")
        custom_layout = QHBoxLayout()
        
        save_btn = QPushButton("💾 Save as Custom Theme")
        save_btn.clicked.connect(self.save_custom_theme)
        custom_layout.addWidget(save_btn)
        
        delete_btn = QPushButton("🗑️ Delete Custom Theme")
        delete_btn.clicked.connect(self.delete_custom_theme)
        custom_layout.addWidget(delete_btn)
        
        custom_group.setLayout(custom_layout)
        layout.addWidget(custom_group)
        
        # Preview area
        preview_group = QGroupBox("Preview")
        preview_layout = QVBoxLayout()
        
        preview_text = QLabel("This is how text will look in the selected theme.")
        preview_layout.addWidget(preview_text)
        
        preview_table = QTableWidget(3, 2)
        preview_table.setHorizontalHeaderLabels(["Source", "Target"])
        preview_table.setItem(0, 0, QTableWidgetItem("Sample text"))
        preview_table.setItem(0, 1, QTableWidgetItem("Voorbeeldtekst"))
        preview_layout.addWidget(preview_table)
        
        preview_group.setLayout(preview_layout)
        layout.addWidget(preview_group)
        
        # Dialog buttons
        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | 
            QDialogButtonBox.StandardButton.Cancel
        )
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
    
    def load_themes(self):
        """Load available themes into combo box"""
        self.theme_combo.clear()
        themes = self.theme_manager.get_all_themes()
        for theme_name in themes.keys():
            self.theme_combo.addItem(theme_name)
        
        # Select current theme
        current_idx = self.theme_combo.findText(self.theme_manager.current_theme.name)
        if current_idx >= 0:
            self.theme_combo.setCurrentIndex(current_idx)
    
    def on_theme_selected(self, theme_name):
        """When theme is selected from combo box"""
        if not theme_name:
            return
        
        theme = self.theme_manager.get_theme(theme_name)
        if theme:
            self.update_color_buttons(theme)
    
    def update_color_buttons(self, theme):
        """Update color button backgrounds"""
        for attr, btn in self.color_buttons.items():
            color = getattr(theme, attr, "#FFFFFF")
            btn.setStyleSheet(f"background-color: {color}; border: 1px solid #999;")
            btn.setText(color)
    
    def pick_color(self, attr):
        """Open color picker for an attribute"""
        from PyQt6.QtWidgets import QColorDialog
        
        theme_name = self.theme_combo.currentText()
        theme = self.theme_manager.get_theme(theme_name)
        
        if not theme:
            return
        
        current_color = QColor(getattr(theme, attr))
        color = QColorDialog.getColor(current_color, self, f"Pick {attr}")
        
        if color.isValid():
            # Update button
            hex_color = color.name()
            self.color_buttons[attr].setStyleSheet(
                f"background-color: {hex_color}; border: 1px solid #999;"
            )
            self.color_buttons[attr].setText(hex_color)
            
            # Update theme
            setattr(theme, attr, hex_color)
    
    def apply_theme(self):
        """Apply the selected theme"""
        theme_name = self.theme_combo.currentText()
        if self.theme_manager.set_theme(theme_name):
            self.theme_manager.apply_theme(QApplication.instance())
            QMessageBox.information(self, "Theme Applied", 
                                   f"Theme '{theme_name}' has been applied.")
    
    def save_custom_theme(self):
        """Save current settings as a custom theme"""
        from PyQt6.QtWidgets import QInputDialog
        
        name, ok = QInputDialog.getText(self, "Save Theme", "Enter theme name:")
        if ok and name:
            # Create new theme from current settings
            theme_name = self.theme_combo.currentText()
            base_theme = self.theme_manager.get_theme(theme_name)
            
            if base_theme:
                # Create copy with new name
                from modules.theme_manager import Theme
                new_theme = Theme(
                    name=name,
                    **{k: v for k, v in base_theme.to_dict().items() if k != 'name'}
                )
                
                # Update with any modified colors
                for attr, btn in self.color_buttons.items():
                    color = btn.text()
                    if color.startswith('#'):
                        setattr(new_theme, attr, color)
                
                self.theme_manager.save_custom_theme(new_theme)
                self.load_themes()
                
                # Select the new theme
                idx = self.theme_combo.findText(name)
                if idx >= 0:
                    self.theme_combo.setCurrentIndex(idx)
                
                QMessageBox.information(self, "Theme Saved", 
                                       f"Custom theme '{name}' has been saved.")
    
    def delete_custom_theme(self):
        """Delete the selected custom theme"""
        theme_name = self.theme_combo.currentText()
        
        # Can't delete predefined themes
        from modules.theme_manager import ThemeManager
        if theme_name in ThemeManager.PREDEFINED_THEMES:
            QMessageBox.warning(self, "Cannot Delete", 
                               "Cannot delete predefined themes.")
            return
        
        reply = QMessageBox.question(self, "Delete Theme",
                                    f"Delete custom theme '{theme_name}'?",
                                    QMessageBox.StandardButton.Yes | 
                                    QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.Yes:
            if self.theme_manager.delete_custom_theme(theme_name):
                self.load_themes()
                QMessageBox.information(self, "Theme Deleted", 
                                       f"Theme '{theme_name}' has been deleted.")


# ============================================================================
# MAIN WINDOW
# ============================================================================

class SupervertalerQt(QMainWindow):
    """Main application window"""
    
    MAX_RECENT_PROJECTS = 10  # Maximum number of recent projects to track
    
    def __init__(self):
        super().__init__()
        
        # Application state
        self.current_project: Optional[Project] = None
        self.project_file_path: Optional[str] = None
        self.project_modified = False
        
        # memoQ bilingual DOCX import tracking
        self.memoq_source_file = None
        
        # UI Configuration
        self.default_font_family = "Calibri"
        self.default_font_size = 11
        
        # Application settings
        self.allow_replace_in_source = False  # Safety: don't allow replace in source by default
        self.auto_propagate_exact_matches = True  # Auto-fill 100% TM matches for empty segments
        
        # TM and Termbase matching toggle (default: enabled)
        self.enable_tm_matching = True
        self.enable_termbase_matching = True
        self.enable_mt_matching = True  # Machine Translation enabled
        self.enable_llm_matching = True  # LLM Translation enabled
        
        # Translation service availability flags (would be set from config/API keys)
        self.google_translate_enabled = True  # For demo purposes
        self.deepl_enabled = False  # Not implemented yet
        self.openai_enabled = True  # For demo purposes
        self.claude_enabled = True  # For demo purposes
        
        # Timer for delayed lookup (cancel if user moves to another segment)
        self.lookup_timer = None
        self.current_lookup_segment_id = None
        
        # Termbase cache for performance optimization
        # Maps segment ID → {term: translation} dictionary
        self.termbase_cache = {}
        self.termbase_cache_lock = threading.Lock()  # Thread-safe cache access
        self.termbase_batch_worker_thread = None  # Background worker thread
        self.termbase_batch_stop_event = threading.Event()  # Signal to stop background worker
        
        # Global language settings (defaults)
        self.source_language = "English"
        self.target_language = "Dutch"
        
        # View mode tracking
        self.current_view_mode = LayoutMode.GRID  # Default to Grid view
        
        # Document view state (initialized early to prevent AttributeError during project loading)
        self.doc_segment_widgets = {}
        self.doc_current_segment_id = None
        
        # List view state (initialized early to prevent AttributeError during project loading)
        self.list_current_segment_id = None
        
        # Universal Lookup detached window
        self.lookup_detached_window = None
        
        # Database Manager for Termbases
        from modules.database_manager import DatabaseManager
        self.user_data_path = Path("user_data_private" if ENABLE_PRIVATE_FEATURES else "user_data")
        self.db_manager = DatabaseManager(
            db_path=str(self.user_data_path / "Translation_Resources" / "supervertaler.db"),
            log_callback=self.log
        )
        self.db_manager.connect()
        
        # Theme Manager
        from modules.theme_manager import ThemeManager
        self.theme_manager = None  # Will be initialized after UI setup
        
        # User data path - uses safety system to prevent private data leaks
        # If .supervertaler.local exists: uses "user_data_private" (git-ignored)
        # Otherwise: uses "user_data" (safe to commit)
        base_folder = "user_data_private" if ENABLE_PRIVATE_FEATURES else "user_data"
        self.recent_projects_file = self.user_data_path / "recent_projects.json"
        
        # Initialize UI
        self.init_ui()
        
        # Initialize theme manager and apply theme
        self.theme_manager = ThemeManager(self.user_data_path)
        self.theme_manager.apply_theme(QApplication.instance())
        
        # Create example API keys file on first launch (after UI is ready)
        self.ensure_example_api_keys()
        
        self.log("Welcome to Supervertaler Qt v1.3.2")
        self.log("Professional Translation Memory & CAT Tool")
        
        # Load general settings (including auto-propagation)
        self.load_general_settings()
        
        # Load language settings
        self.load_language_settings()
        
        # Restore last project if enabled in settings
        self.restore_last_project_if_enabled()
        
        # Load font sizes from preferences (after UI is fully initialized)
        QApplication.instance().processEvents()  # Allow UI to finish initializing
        self.load_font_sizes_from_preferences()
    
    def init_ui(self):
        """Initialize the user interface"""
        # Build window title with dev mode indicator
        title = "Supervertaler Qt v1.3.2"
        if ENABLE_PRIVATE_FEATURES:
            title += " [🛠️ DEV MODE]"
        self.setWindowTitle(title)
        self.setGeometry(100, 100, 1400, 800)
        
        # Ensure window can be resized (no minimum size constraint)
        self.setMinimumSize(400, 300)  # Very small minimum to allow resizing
        
        # Create menu bar (ribbon removed - using traditional menus)
        self.create_menus()
        
        # Ribbon removed - all functionality moved to menu bar
        # self.create_ribbon()
        
        # Create main layout
        self.create_main_layout()
        
        # Create status bar
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        self.status_bar.showMessage("Ready")

        # Setup global shortcuts
        self.setup_global_shortcuts()

    def setup_global_shortcuts(self):
        """Setup application-wide keyboard shortcuts"""
        from PyQt6.QtGui import QShortcut

        # F9 - Voice dictation
        self.shortcut_dictate = QShortcut(QKeySequence("F9"), self)
        self.shortcut_dictate.activated.connect(self.start_voice_dictation)

    def create_menus(self):
        """Create application menus"""
        menubar = self.menuBar()
        
        # File Menu
        file_menu = menubar.addMenu("&File")
        
        new_action = QAction("&New Project", self)
        new_action.setShortcut(QKeySequence.StandardKey.New)
        new_action.triggered.connect(self.new_project)
        file_menu.addAction(new_action)
        
        open_action = QAction("&Open Project...", self)
        open_action.setShortcut(QKeySequence.StandardKey.Open)
        open_action.triggered.connect(self.open_project)
        file_menu.addAction(open_action)
        
        # Recent projects submenu
        self.recent_menu = file_menu.addMenu("Open &Recent")
        self.update_recent_menu()
        
        file_menu.addSeparator()
        
        save_action = QAction("&Save", self)
        save_action.setShortcut(QKeySequence.StandardKey.Save)
        save_action.triggered.connect(self.save_project)
        file_menu.addAction(save_action)
        
        save_as_action = QAction("Save &As...", self)
        save_as_action.setShortcut(QKeySequence.StandardKey.SaveAs)
        save_as_action.triggered.connect(self.save_project_as)
        file_menu.addAction(save_as_action)
        
        file_menu.addSeparator()
        
        close_action = QAction("&Close Project", self)
        close_action.triggered.connect(self.close_project)
        file_menu.addAction(close_action)
        
        file_menu.addSeparator()
        
        # Import/Export submenu
        import_menu = file_menu.addMenu("&Import")
        
        import_docx_action = QAction("&Monolingual Document (DOCX)...", self)
        import_docx_action.triggered.connect(self.import_docx)
        import_docx_action.setShortcut("Ctrl+O")
        import_menu.addAction(import_docx_action)
        
        import_memoq_action = QAction("memoQ &Bilingual Table (DOCX)...", self)
        import_memoq_action.triggered.connect(self.import_memoq_bilingual)
        import_menu.addAction(import_memoq_action)
        
        export_menu = file_menu.addMenu("&Export")
        
        export_memoq_action = QAction("memoQ &Bilingual Table - Translated (DOCX)...", self)
        export_memoq_action.triggered.connect(self.export_memoq_bilingual)
        export_menu.addAction(export_memoq_action)
        
        export_menu.addSeparator()
        
        export_grid_action = QAction("TMX from &Grid (all segments)...", self)
        export_grid_action.triggered.connect(self.export_tmx_from_grid)
        export_menu.addAction(export_grid_action)
        
        export_selected_action = QAction("TMX from &Selected Segments...", self)
        export_selected_action.triggered.connect(self.export_tmx_from_selected)
        export_menu.addAction(export_selected_action)
        
        export_tm_action = QAction("TMX from &TM(s) for Current Project...", self)
        export_tm_action.triggered.connect(self.export_tmx_from_tm_database)
        export_menu.addAction(export_tm_action)
        
        file_menu.addSeparator()
        
        exit_action = QAction("E&xit", self)
        exit_action.setShortcut(QKeySequence.StandardKey.Quit)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        # Edit Menu
        edit_menu = menubar.addMenu("&Edit")
        
        undo_action = QAction("&Undo", self)
        undo_action.setShortcut(QKeySequence.StandardKey.Undo)
        edit_menu.addAction(undo_action)
        
        redo_action = QAction("&Redo", self)
        redo_action.setShortcut(QKeySequence.StandardKey.Redo)
        edit_menu.addAction(redo_action)
        
        edit_menu.addSeparator()
        
        find_action = QAction("&Find...", self)
        find_action.setShortcut(QKeySequence.StandardKey.Find)
        find_action.triggered.connect(self.show_find_replace_dialog)
        edit_menu.addAction(find_action)
        
        replace_action = QAction("&Replace...", self)
        replace_action.setShortcut(QKeySequence.StandardKey.Replace)
        replace_action.triggered.connect(self.show_find_replace_dialog)
        edit_menu.addAction(replace_action)
        
        edit_menu.addSeparator()
        
        goto_action = QAction("&Go to Segment...", self)
        goto_action.setShortcut("Ctrl+G")
        goto_action.triggered.connect(self.show_goto_dialog)
        edit_menu.addAction(goto_action)
        
        edit_menu.addSeparator()
        
        translate_action = QAction("&Translate Segment", self)
        translate_action.setShortcut("Ctrl+T")
        translate_action.triggered.connect(self.translate_current_segment)
        edit_menu.addAction(translate_action)

        translate_menu = edit_menu.addMenu("Translate &Multiple Segments")

        translate_selected_not_started_action = QAction("Translate selected not-started segments", self)
        translate_selected_not_started_action.triggered.connect(
            lambda checked=False: self.translate_multiple_segments("selected_not_started")
        )
        translate_menu.addAction(translate_selected_not_started_action)

        translate_all_not_started_action = QAction("Translate all not-started segments", self)
        translate_all_not_started_action.triggered.connect(
            lambda checked=False: self.translate_multiple_segments("all_not_started")
        )
        translate_menu.addAction(translate_all_not_started_action)

        translate_all_pretranslated_action = QAction("Translate all pre-translated segments", self)
        translate_all_pretranslated_action.triggered.connect(
            lambda checked=False: self.translate_multiple_segments("all_pretranslated")
        )
        translate_menu.addAction(translate_all_pretranslated_action)

        translate_pending_action = QAction("Translate all not-started & pre-translated", self)
        translate_pending_action.setShortcut("Ctrl+Shift+T")
        translate_pending_action.triggered.connect(
            lambda checked=False: self.translate_multiple_segments("all_not_started_pretranslated")
        )
        translate_menu.addAction(translate_pending_action)

        translate_translatable_action = QAction("Translate all translatable segments", self)
        translate_translatable_action.triggered.connect(
            lambda checked=False: self.translate_multiple_segments("all_translatable")
        )
        translate_menu.addAction(translate_translatable_action)

        translate_all_segments_action = QAction("Translate all segments (all statuses)", self)
        translate_all_segments_action.triggered.connect(
            lambda checked=False: self.translate_multiple_segments("all_segments")
        )
        translate_menu.addAction(translate_all_segments_action)
        
        edit_menu.addSeparator()
        
        # Bulk Operations submenu
        bulk_menu = edit_menu.addMenu("Bulk &Operations")
        
        clear_translations_action = QAction("🗑️ &Clear Translations", self)
        clear_translations_action.setToolTip("Clear translations for selected segments")
        clear_translations_action.triggered.connect(self.clear_selected_translations_from_menu)
        bulk_menu.addAction(clear_translations_action)
        
        edit_menu.addSeparator()
        
        # Universal Lookup
        universal_lookup_action = QAction("🔍 &Universal Lookup...", self)
        universal_lookup_action.setShortcut("Ctrl+Alt+L")
        # Tab indices: Home=0, Resources=1, Modules=2, Settings=3 (Prompt Manager and Editor removed)
        universal_lookup_action.triggered.connect(lambda: self.right_tabs.setCurrentIndex(2) if hasattr(self, 'right_tabs') else None)  # Modules tab
        edit_menu.addAction(universal_lookup_action)
        
        # View Menu
        view_menu = menubar.addMenu("&View")
        
        # Navigation submenu
        nav_menu = view_menu.addMenu("📑 &Navigate To")
        
        go_home_action = QAction("🏠 &Home", self)
        go_home_action.triggered.connect(lambda: self.right_tabs.setCurrentIndex(0) if hasattr(self, 'right_tabs') else None)  # Prompt Manager tab
        nav_menu.addAction(go_home_action)
        
        # Prompt Manager tab removed - it's now integrated in the Home tab
        # go_prompt_action = QAction("💡 &Prompt Manager", self)
        # go_prompt_action.triggered.connect(lambda: self.main_tabs.setCurrentIndex(1) if hasattr(self, 'main_tabs') else None)
        # nav_menu.addAction(go_prompt_action)
        
        # Editor tab removed - functionality moved to Home tab
        # go_editor_action = QAction("📝 &Editor", self)
        # go_editor_action.triggered.connect(lambda: self.main_tabs.setCurrentIndex(2) if hasattr(self, 'main_tabs') else None)
        # nav_menu.addAction(go_editor_action)
        
        go_settings_action = QAction("⚙️ &Settings", self)
        # Tab indices shifted: Resources=1, Modules=2, Settings=3 (Home=0, Prompt Manager and Editor removed)
        go_settings_action.triggered.connect(lambda: self.right_tabs.setCurrentIndex(3) if hasattr(self, 'right_tabs') else None)  # Settings tab
        nav_menu.addAction(go_settings_action)
        
        view_menu.addSeparator()
        
        # View mode switcher
        grid_view_action = QAction("📊 &Grid View", self)
        grid_view_action.triggered.connect(lambda: self.switch_view_mode(LayoutMode.GRID))
        view_menu.addAction(grid_view_action)
        
        list_view_action = QAction("📋 &List View", self)
        list_view_action.triggered.connect(lambda: self.switch_view_mode(LayoutMode.LIST))
        view_menu.addAction(list_view_action)
        
        document_view_action = QAction("📄 &Document View", self)
        document_view_action.triggered.connect(lambda: self.switch_view_mode(LayoutMode.DOCUMENT))
        view_menu.addAction(document_view_action)
        
        view_menu.addSeparator()
        
        # Grid Text section
        grid_zoom_menu = view_menu.addMenu("📊 &Grid Text Zoom")
        
        grid_zoom_in = QAction("Grid Zoom &In", self)
        grid_zoom_in.setShortcut(QKeySequence.StandardKey.ZoomIn)
        grid_zoom_in.triggered.connect(self.zoom_in)
        grid_zoom_menu.addAction(grid_zoom_in)
        
        grid_zoom_out = QAction("Grid Zoom &Out", self)
        grid_zoom_out.setShortcut(QKeySequence.StandardKey.ZoomOut)
        grid_zoom_out.triggered.connect(self.zoom_out)
        grid_zoom_menu.addAction(grid_zoom_out)
        
        grid_zoom_menu.addSeparator()
        
        grid_increase_font = QAction("Grid &Increase Font Size", self)
        grid_increase_font.setShortcut("Ctrl++")
        grid_increase_font.triggered.connect(self.increase_font_size)
        grid_zoom_menu.addAction(grid_increase_font)
        
        grid_decrease_font = QAction("Grid &Decrease Font Size", self)
        grid_decrease_font.setShortcut("Ctrl+-")
        grid_decrease_font.triggered.connect(self.decrease_font_size)
        grid_zoom_menu.addAction(grid_decrease_font)
        
        grid_zoom_menu.addSeparator()
        
        grid_font_family_menu = grid_zoom_menu.addMenu("Grid Font &Family")
        font_families = ["Calibri", "Segoe UI", "Arial", "Consolas", "Verdana", 
                        "Times New Roman", "Georgia", "Courier New"]
        for font_name in font_families:
            font_action = QAction(font_name, self)
            font_action.triggered.connect(lambda checked, f=font_name: self.set_font_family(f))
            grid_font_family_menu.addAction(font_action)
        
        view_menu.addSeparator()
        
        # Translation Results Pane section
        results_zoom_menu = view_menu.addMenu("📋 Translation &Results Pane")
        
        results_zoom_in_action = QAction("Results Zoom &In", self)
        results_zoom_in_action.setShortcut("Ctrl+Shift+=")
        results_zoom_in_action.triggered.connect(self.results_pane_zoom_in)
        results_zoom_menu.addAction(results_zoom_in_action)
        
        results_zoom_out_action = QAction("Results Zoom &Out", self)
        results_zoom_out_action.setShortcut("Ctrl+Shift+-")
        results_zoom_out_action.triggered.connect(self.results_pane_zoom_out)
        results_zoom_menu.addAction(results_zoom_out_action)
        
        results_zoom_reset_action = QAction("Results Zoom &Reset", self)
        results_zoom_reset_action.triggered.connect(self.results_pane_zoom_reset)
        results_zoom_menu.addAction(results_zoom_reset_action)
        
        results_zoom_menu.addSeparator()
        
        results_note = QAction("(Includes match list + compare boxes)", self)
        results_note.setEnabled(False)
        results_zoom_menu.addAction(results_note)
        
        view_menu.addSeparator()
        
        auto_resize_action = QAction("📐 &Auto-Resize Rows", self)
        auto_resize_action.triggered.connect(self.auto_resize_rows)
        auto_resize_action.setToolTip("Automatically resize all rows to fit content")
        view_menu.addAction(auto_resize_action)
        
        # Tools Menu
        tools_menu = menubar.addMenu("&Tools")
        
        tm_manager_action = QAction("&Translation Memory Manager...", self)
        tm_manager_action.setShortcut("Ctrl+M")
        tm_manager_action.triggered.connect(self.show_tm_manager)
        tools_menu.addAction(tm_manager_action)
        
        autofingers_action = QAction("✋ &AutoFingers - CAT Tool Automation...", self)
        autofingers_action.setShortcut("Ctrl+Shift+A")
        autofingers_action.triggered.connect(self.show_autofingers)
        tools_menu.addAction(autofingers_action)
        
        tools_menu.addSeparator()
        
        theme_action = QAction("🎨 &Theme Editor...", self)
        theme_action.triggered.connect(self.show_theme_editor)
        tools_menu.addAction(theme_action)
        
        settings_action = QAction("&Settings...", self)
        settings_action.triggered.connect(lambda: self._go_to_settings_tab())
        tools_menu.addAction(settings_action)
        
        # Help Menu
        help_menu = menubar.addMenu("&Help")
        
        about_action = QAction("&About", self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)
    
    def create_quick_access_toolbar(self):
        """Create Quick Access Toolbar above ribbon"""
        from PyQt6.QtWidgets import QToolBar, QWidget, QHBoxLayout
        from PyQt6.QtCore import QSize
        
        qat = QToolBar("Quick Access")
        qat.setMovable(False)
        qat.setFloatable(False)
        qat.setIconSize(QSize(20, 20))
        qat.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonIconOnly)
        qat.setMaximumHeight(28)
        qat.setMinimumWidth(170)  # Fixed width to align ribbon tabs
        qat.setMaximumWidth(170)
        
        # Styling for compact appearance
        qat.setStyleSheet("""
            QToolBar {
                background: transparent;
                border: none;
                spacing: 2px;
                padding: 2px;
            }
            QToolButton {
                background: transparent;
                border: 1px solid transparent;
                border-radius: 3px;
                padding: 2px;
                margin: 0px;
                font-size: 14pt;
            }
            QToolButton:hover {
                background: rgba(255, 255, 255, 0.1);
                border: 1px solid rgba(255, 255, 255, 0.2);
            }
            QToolButton:pressed {
                background: rgba(0, 0, 0, 0.1);
            }
        """)
        
        # Empty toolbar - all buttons removed, minimize moved to ribbon
        # Add empty separator just to maintain spacing
        qat.addSeparator()
        
        return qat
    
    def create_ribbon(self):
        """
        Create modern ribbon interface - DISABLED
        Ribbon functionality has been moved to traditional menu bar.
        This code is kept for potential future use.
        """
        # Ribbon removed - all actions are now in the menu bar
        # Keeping this method for potential future re-enablement
        return
        
        # from modules.ribbon_widget import RibbonWidget, RibbonTab, RibbonGroup, RibbonButton
        # 
        # # Create ribbon widget
        # self.ribbon = RibbonWidget(self)
        # 
        # # Connect ribbon actions to methods
        # self.ribbon.action_triggered.connect(self.handle_ribbon_action)
        
        # Ribbon code commented out - all functionality moved to menu bar
        # HOME TAB - Project management and navigation
        # home_tab = RibbonTab()
        # project_group = RibbonGroup("Project")
        # project_group.add_button(self.ribbon.create_button("New", "📄", "new", "New project (Ctrl+N)"))
        # ... (all ribbon tab creation code commented out)
        # All actions are now available in the traditional menu bar
    
    def create_ribbon_toolbar(self):
        """Create a toolbar to hold the ribbon - DISABLED"""
        # Ribbon removed - using traditional menu bar instead
        # This method kept for potential future use
        return None
    
    def toggle_ribbon_minimized(self, minimized: bool):
        """Toggle ribbon between full and minimized (tabs-only) mode - DISABLED"""
        # Ribbon removed - no action needed
        self.ribbon_minimized = minimized  # Store state for backwards compatibility
        pass
    
    def show_ribbon_temporarily(self, index: int):
        """Show ribbon temporarily when tab clicked in minimized mode - DISABLED"""
        # Ribbon removed - no action needed
        pass
    
    def on_main_tab_changed(self, index: int):
        """Handle main tab change"""
        # Ribbon removed - menu bar remains constant across all tabs
        # No context switching needed
        pass
    
    def toggle_sidebar(self, visible: bool):
        """Toggle Quick Access Sidebar visibility - DEPRECATED (Quick Access removed)"""
        # The Quick Access sidebar has been removed in favor of Project Home collapsible panel
        pass
    
    def update_sidebar_recent_files(self):
        """Update sidebar recent files list - DEPRECATED (Quick Access removed)"""
        # The Quick Access sidebar has been removed in favor of Project Home collapsible panel
        pass
    
    def handle_ribbon_action(self, action_name: str):
        """
        Handle ribbon button clicks - DEPRECATED
        Ribbon has been removed in favor of traditional menu bar.
        All actions are now accessible through menus.
        This method is kept for backwards compatibility.
        """
        # All ribbon actions have been moved to menu bar
        # This method is kept in case any code still references it
        action_map = {
            # File actions
            "new": self.new_project,
            "open": self.open_project,
            "save": self.save_project,
            "close_project": self.close_project,
            
            # Navigation actions
            "go_to_home": lambda: self.right_tabs.setCurrentIndex(0) if hasattr(self, 'right_tabs') else None,  # Prompt Manager tab
            "go_to_settings": lambda: self.right_tabs.setCurrentIndex(3) if hasattr(self, 'right_tabs') else None,  # Settings tab
            
            # Translation actions
            "translate": self.translate_current_segment,
            "batch_translate": self.translate_batch,
            "tm_manager": self.show_tm_manager,
            "universal_lookup": lambda: self._go_to_universal_lookup() if hasattr(self, 'right_tabs') else None,
            
            # View actions
            "zoom_in": self.zoom_in,
            "zoom_out": self.zoom_out,
            "auto_resize": self.auto_resize_rows,
            "themes": self.show_theme_editor,
            
            # Tools actions
            "autofingers": self.show_autofingers,
            "options": self._go_to_settings_tab,
        }
        
        # Execute action if found (for backwards compatibility)
        action = action_map.get(action_name)
        if action:
            action()
    
    
    def create_toolbar(self):
        """Create main toolbar - REMOVED: Replaced by ribbon interface"""
        # Toolbar removed - replaced by modern ribbon interface
        # All functionality accessible via ribbon tabs and menus
        pass
    
    def create_main_layout(self):
        """Create main application layout with tabs"""
        # Central widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # Main layout with splitter for sidebar and main content
        from PyQt6.QtWidgets import QSplitter, QHBoxLayout
        main_layout = QHBoxLayout(central_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        # Main content area
        content_widget = QWidget()
        from PyQt6.QtWidgets import QSizePolicy
        content_widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        content_layout = QVBoxLayout(content_widget)
        content_layout.setContentsMargins(5, 5, 5, 5)
        
        # ===== NEW RESTRUCTURED UI =====
        # Left: Tab widget with Prompt Manager, Resources, Modules, Settings
        # Right: Editor (always visible, no tabs) 
        from modules.unified_prompt_manager_qt import UnifiedPromptManagerQt
        
        # Main horizontal splitter: Left-side tabs (left) and Editor (right)
        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # ===== LEFT SIDE: Tab widget with Prompt Manager, Resources, Modules, Settings =====
        self.right_tabs = QTabWidget()
        self.right_tabs.setStyleSheet("""
            QTabBar::tab { padding: 8px 15px; }
        """)
        # Ensure left tabs can be resized
        self.right_tabs.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        
        # Connect tab change for navigation tracking
        self.right_tabs.currentChanged.connect(self.on_right_tab_changed)
        
        # 1. UNIFIED PROMPT LIBRARY (first tab)
        prompt_widget = QWidget()
        self.prompt_manager_qt = UnifiedPromptManagerQt(self, standalone=False)
        self.prompt_manager_qt.create_tab(prompt_widget)
        self.right_tabs.addTab(prompt_widget, "🤖 Prompt Manager")
        
        # 2. TRANSLATION RESOURCES
        resources_tab = self.create_resources_tab()
        self.right_tabs.addTab(resources_tab, "📚 Translation Resources")
        
        # 3. MODULES
        modules_tab = self.create_specialised_modules_tab()
        self.right_tabs.addTab(modules_tab, "🧩 Modules")

        # 4. SETTINGS
        settings_tab = self.create_settings_tab()
        self.right_tabs.addTab(settings_tab, "⚙️ Settings")
        
        self.main_splitter.addWidget(self.right_tabs)
        
        # ===== RIGHT SIDE: Editor (no tabs, always visible) =====
        editor_widget = self.create_editor_widget()
        # Ensure editor widget can be resized (splitter handles this automatically, but explicit policy helps)
        from PyQt6.QtWidgets import QSizePolicy
        editor_widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.main_splitter.addWidget(editor_widget)
        
        # Ensure splitter is resizable
        self.main_splitter.setChildrenCollapsible(False)  # Prevent collapsing completely
        self.main_splitter.setHandleWidth(8)  # Make the splitter handle more visible
        self.main_splitter.setOpaqueResize(True)  # Enable real-time resizing
        
        # Set minimum sizes for widgets to allow resizing below default sizes
        self.right_tabs.setMinimumWidth(200)  # Allow left tabs to shrink to 200px
        editor_widget.setMinimumWidth(200)  # Allow editor to shrink to 200px
        
        # Set initial splitter sizes - give more space to translation grid (right panel)
        # Better default: ~40% left (Prompt Manager), ~60% right (grid)
        # For a typical 1920px wide window: ~770px left, ~1150px right
        self.main_splitter.setSizes([770, 1150])  # Initial sizes: reasonable space for Prompt Manager, more space for grid
        
        # Set stretch factors for proportional resizing (right gets more space)
        self.main_splitter.setStretchFactor(0, 2)  # Left tabs: stretch factor 2
        self.main_splitter.setStretchFactor(1, 3)  # Editor widget: stretch factor 3
        
        # Add splitter to content layout with stretch to fill available space
        content_layout.addWidget(self.main_splitter, 1)  # Stretch factor 1
        
        # Add content directly to main layout with stretch
        main_layout.addWidget(content_widget, 1)  # Stretch factor 1
    
    def _create_placeholder_tab(self, title: str, description: str) -> QWidget:
        """Create a simple placeholder tab"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(10, 10, 10, 10)
        
        header = QLabel(title)
        header.setStyleSheet("font-size: 14px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(header)
        
        placeholder = QLabel(description)
        placeholder.setAlignment(Qt.AlignmentFlag.AlignCenter)
        placeholder.setStyleSheet("color: #888; font-size: 12px;")
        layout.addWidget(placeholder, stretch=1)
        
        return tab
    
    def create_non_translatables_tab(self) -> QWidget:
        """Create the Non-Translatables tab - NT management"""
        return self._create_placeholder_tab(
            "🚫 Non-Translatables",
            "Non-Translatables Manager - Coming Soon\n\nFeatures:\n• Manage non-translatable content\n• Define NT patterns\n• Exclude from translation"
        )
    
    def create_prompt_manager_tab(self) -> QWidget:
        """Create the Unified Prompt Library tab - Simplified 2-Layer Architecture"""
        from modules.unified_prompt_manager_qt import UnifiedPromptManagerQt
        
        # Create Unified Prompt Manager widget (embedded mode, not standalone)
        prompt_widget = QWidget()
        self.prompt_manager_qt = UnifiedPromptManagerQt(self, standalone=False)
        self.prompt_manager_qt.create_tab(prompt_widget)
        
        return prompt_widget
    
    def create_tmx_editor_tab(self) -> QWidget:
        """Create the TMX Editor tab - Edit TMs"""
        from modules.tmx_editor_qt import TmxEditorUIQt
        
        # Create container widget with Universal Lookup style header
        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(5)  # Reduced from 10 to 5 for tighter spacing
        
        # Header (matches Universal Lookup / AutoFingers / PDF Rescue style)
        header = QLabel("📝 TMX Editor")
        header.setStyleSheet("font-size: 16pt; font-weight: bold; color: #1976D2;")
        layout.addWidget(header, 0)  # 0 = no stretch, stays compact
        
        # Description box (matches Universal Lookup / AutoFingers / PDF Rescue style)
        description = QLabel(
            "Edit translation memory files directly - inspired by Heartsome TMX Editor.\n"
            "Open, edit, filter, and manage your TMX translation memories."
        )
        description.setWordWrap(True)
        description.setStyleSheet("color: #666; padding: 5px; background-color: #E3F2FD; border-radius: 3px;")
        layout.addWidget(description, 0)  # 0 = no stretch, stays compact
        
        # Create TMX Editor widget (embedded mode) - pass database manager
        tmx_editor = TmxEditorUIQt(parent=None, standalone=False, db_manager=self.db_manager)
        layout.addWidget(tmx_editor, 1)  # 1 = stretch factor, expands to fill space
        
        # Store reference for potential future use
        self.tmx_editor_embedded = tmx_editor
        
        return container
    
    def create_reference_images_tab(self) -> QWidget:
        """Create the Reference Images tab - Visual context"""
        return self._create_placeholder_tab(
            "🖼️ Reference Images",
            "Reference Images - Coming Soon\n\nFeatures:\n• Upload reference images\n• Visual context for translation\n• Screenshot annotation"
        )
    
    def create_pdf_rescue_tab(self) -> QWidget:
        """Create the PDF Rescue tab - AI OCR"""
        from modules.pdf_rescue_Qt import PDFRescueQt
        
        # Create PDF Rescue widget (embedded mode, not standalone)
        pdf_rescue_widget = QWidget()
        self.pdf_rescue_qt = PDFRescueQt(self, standalone=False)
        self.pdf_rescue_qt.create_tab(pdf_rescue_widget)
        
        return pdf_rescue_widget
    
    def create_encoding_repair_tab(self) -> QWidget:
        """Create the Encoding Repair tab - Text Encoding Tool"""
        from modules.encoding_repair_Qt import EncodingRepairQt
        
        # Create Encoding Repair widget (embedded mode, not standalone)
        encoding_repair_widget = QWidget()
        self.encoding_repair_qt = EncodingRepairQt(self, standalone=False)
        self.encoding_repair_qt.create_tab(encoding_repair_widget)
        
        return encoding_repair_widget
    
    def create_tracked_changes_tab(self) -> QWidget:
        """Create the Tracked Changes tab - Post-Translation Analysis"""
        return self._create_placeholder_tab(
            "📊 Tracked Changes",
            "Tracked Changes - Coming Soon\n\nFeatures:\n• Track translation changes\n• Version history\n• Comparison reports"
        )

    def create_llm_leaderboard_tab(self) -> QWidget:
        """Create the LLM Leaderboard tab - Benchmark LLM translation quality"""
        from modules.llm_leaderboard_ui import LLMLeaderboardUI

        # Create LLM client factory that uses existing API keys
        def llm_client_factory(provider: str, model_id: str):
            from modules.llm_clients import LLMClient
            api_keys = self._get_api_keys()

            # Map provider names to API key names
            # (gemini uses "google" as the key name in api_keys.txt)
            provider_to_key = {
                "openai": "openai",
                "claude": "claude",
                "gemini": "google"  # Gemini uses Google API key
            }

            key_name = provider_to_key.get(provider, provider)
            api_key = api_keys.get(key_name, "")

            if not api_key:
                raise ValueError(f"No API key configured for {provider}. Please add it in Settings.")

            return LLMClient(api_key=api_key, provider=provider, model=model_id)

        # Create and return the leaderboard UI widget
        leaderboard_widget = LLMLeaderboardUI(
            parent=self,
            llm_client_factory=llm_client_factory
        )

        return leaderboard_widget

    def _get_api_keys(self) -> dict:
        """Get API keys from settings"""
        from modules.llm_clients import load_api_keys
        return load_api_keys()

    def create_log_tab(self) -> QWidget:
        """Create the Log tab - Session Log"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # Log display area
        self.session_log = QPlainTextEdit()
        self.session_log.setReadOnly(True)
        self.session_log.setStyleSheet("""
            QPlainTextEdit {
                background-color: #ffffff;
                color: #000000;
                font-family: 'Consolas', 'Courier New', monospace;
                font-size: 9px;
                border: 1px solid #ccc;
            }
        """)
        layout.addWidget(self.session_log)
        
        # Initialize with welcome message
        self.session_log.setPlainText("Session Log - Ready\n" + "="*50 + "\n")
        
        return tab
    
    def create_editor_widget(self):
        """Create the Editor widget (left side, always visible, no tabs)"""
        from PyQt6.QtWidgets import QPushButton, QLabel, QStackedWidget, QSizePolicy
        
        widget = QWidget()
        widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        layout = QVBoxLayout(widget)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        # View switcher toolbar
        view_toolbar = QWidget()
        view_toolbar.setStyleSheet("background-color: #f8f9fa; padding: 5px;")
        view_toolbar_layout = QHBoxLayout(view_toolbar)
        view_toolbar_layout.setContentsMargins(10, 5, 10, 5)
        view_toolbar_layout.setSpacing(5)
        
        view_label = QLabel("View:")
        view_label.setStyleSheet("font-weight: bold;")
        view_toolbar_layout.addWidget(view_label)
        
        self.home_grid_view_btn = QPushButton("📊 Grid")
        self.home_grid_view_btn.setCheckable(True)
        self.home_grid_view_btn.setChecked(True)
        self.home_grid_view_btn.clicked.connect(lambda: self.switch_home_view_mode("grid"))
        view_toolbar_layout.addWidget(self.home_grid_view_btn)
        
        self.home_list_view_btn = QPushButton("📋 List")
        self.home_list_view_btn.setCheckable(True)
        self.home_list_view_btn.clicked.connect(lambda: self.switch_home_view_mode("list"))
        view_toolbar_layout.addWidget(self.home_list_view_btn)
        
        self.home_document_view_btn = QPushButton("📄 Document")
        self.home_document_view_btn.setCheckable(True)
        self.home_document_view_btn.clicked.connect(lambda: self.switch_home_view_mode("document"))
        view_toolbar_layout.addWidget(self.home_document_view_btn)
        
        view_toolbar_layout.addStretch()
        
        layout.addWidget(view_toolbar)
        
        # Create assistance panel FIRST (create_grid_view_widget_for_home needs it)
        if not hasattr(self, 'assistance_widget') or self.assistance_widget is None:
            self.create_assistance_panel()
        
        # Stacked widget for different views
        self.home_view_stack = QStackedWidget()
        
        # Create grid view widget with assistance panel at bottom (vertical layout)
        grid_view_widget = self.create_grid_view_widget_for_home()
        self.home_view_stack.addWidget(grid_view_widget)
        
        # List and Document views - create adapted versions for home tab
        list_view_widget = self.create_list_view_widget_for_home()
        self.home_view_stack.addWidget(list_view_widget)
        
        doc_view_widget = self.create_document_view_widget_for_home()
        self.home_view_stack.addWidget(doc_view_widget)
        
        layout.addWidget(self.home_view_stack)
        
        return widget
    
    def create_home_tab(self):
        """DEPRECATED: Replaced by create_editor_widget() - kept for backwards compatibility"""
        return self.create_editor_widget()
    
    def on_right_tab_changed(self, index: int):
        """Handle right-side tab changes"""
        # Refresh AI Assistant LLM client when switching to Prompt Manager tab
        if index == 0 and hasattr(self, 'prompt_manager_qt'):
            # Reload LLM settings in case they changed
            llm_settings = self.load_llm_settings()
            self.current_provider = llm_settings.get('provider', 'openai')
            provider_key = f"{self.current_provider}_model"
            self.current_model = llm_settings.get(provider_key)
            # Reinitialize AI Assistant's LLM client
            self.prompt_manager_qt._init_llm_client()
    
    def detach_universal_lookup(self):
        """Detach Universal Lookup into a separate window for second screen use"""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QPushButton
        
        try:
            # If already detached, just show the window
            if hasattr(self, 'lookup_detached_window') and self.lookup_detached_window and self.lookup_detached_window.isVisible():
                self.lookup_detached_window.raise_()
                self.lookup_detached_window.activateWindow()
                return
            
            # Create detached window
            self.lookup_detached_window = QDialog(self)
            self.lookup_detached_window.setWindowTitle("🔍 Universal Lookup - Supervertaler")
            self.lookup_detached_window.setMinimumSize(600, 700)
            self.lookup_detached_window.resize(700, 800)
            
            # Set window flags to ensure it appears as a proper window
            self.lookup_detached_window.setWindowFlags(
                Qt.WindowType.Window | 
                Qt.WindowType.WindowCloseButtonHint | 
                Qt.WindowType.WindowMinimizeButtonHint |
                Qt.WindowType.WindowMaximizeButtonHint
            )
            
            # Position window to the right of main window, or center on same screen if no space
            main_geometry = self.geometry()
            
            # Get the screen that contains the main window (handles multi-monitor setups)
            app = QApplication.instance()
            screen = app.screenAt(main_geometry.center())
            if not screen:
                # Fallback to primary screen
                screen = app.primaryScreen()
            
            screen_geometry = screen.geometry()
            
            # Try to place to the right of main window
            new_x = main_geometry.right() + 20
            new_y = main_geometry.top()
            
            # If window would go off-screen to the right, center it on the main window's screen
            if new_x + 700 > screen_geometry.right():
                # Center horizontally on screen, place above main window
                new_x = screen_geometry.left() + (screen_geometry.width() - 700) // 2
                new_y = max(screen_geometry.top() + 50, main_geometry.top() - 100)
                
            # Ensure window stays on screen
            new_x = max(screen_geometry.left() + 10, min(new_x, screen_geometry.right() - 710))
            new_y = max(screen_geometry.top() + 10, min(new_y, screen_geometry.bottom() - 810))
            
            self.lookup_detached_window.move(int(new_x), int(new_y))
            
            # Make sure window is raised and activated - use QTimer to ensure it's after window is shown
            QTimer.singleShot(100, lambda: (
                self.lookup_detached_window.raise_(),
                self.lookup_detached_window.activateWindow(),
                self.lookup_detached_window.setFocus()
            ))
            
            layout = QVBoxLayout(self.lookup_detached_window)
            layout.setContentsMargins(10, 10, 10, 10)
            layout.setSpacing(5)
            
            # Header with reattach button
            header_layout = QVBoxLayout()
            
            header_title = QLabel("🔍 Universal Lookup")
            header_title.setStyleSheet("font-size: 16px; font-weight: bold; color: #333;")
            header_layout.addWidget(header_title)
            
            button_layout = QVBoxLayout()
            reattach_btn = QPushButton("📥 Attach to Main Window")
            reattach_btn.setToolTip("Re-attach Universal Lookup to the Home tab")
            reattach_btn.setStyleSheet("font-size: 9pt; padding: 4px 12px; max-width: 200px;")
            reattach_btn.clicked.connect(self.reattach_universal_lookup)
            button_layout.addWidget(reattach_btn, alignment=Qt.AlignmentFlag.AlignRight)
            header_layout.addLayout(button_layout)
            
            layout.addLayout(header_layout)
            
            # Create new Universal Lookup instance for detached window
            # Or move the existing one - better to create new to avoid widget parenting issues
            detached_lookup = UniversalLookupTab(self.lookup_detached_window)
            
            # Copy TM database reference if available
            if hasattr(self, 'tm_database') and self.tm_database:
                detached_lookup.tm_database = self.tm_database
                if detached_lookup.engine:
                    detached_lookup.engine.set_tm_database(self.tm_database)
            
            # Copy home lookup state if it exists
            if hasattr(self, 'home_lookup_widget') and self.home_lookup_widget:
                # Copy source text
                source_text = self.home_lookup_widget.source_text.toPlainText()
                detached_lookup.source_text.setPlainText(source_text)
                
                # Copy mode
                mode = self.home_lookup_widget.mode_combo.currentText()
                detached_lookup.mode_combo.setCurrentText(mode)
                
                # Copy TM database reference
                if hasattr(self.home_lookup_widget, 'tm_database'):
                    detached_lookup.tm_database = self.home_lookup_widget.tm_database
            
            layout.addWidget(detached_lookup, stretch=1)
            
            # Store reference for cleanup
            self.lookup_detached_widget = detached_lookup
            
            # Handle window close
            def on_close():
                self.lookup_detached_window = None
                self.lookup_detached_widget = None
            
            self.lookup_detached_window.finished.connect(on_close)
            
            # Show window (non-modal)
            self.lookup_detached_window.setWindowModality(Qt.WindowModality.NonModal)
            self.lookup_detached_window.show()
            self.log("Universal Lookup detached to separate window")
        except Exception as e:
            import traceback
            error_msg = f"Error detaching Universal Lookup: {str(e)}\n{traceback.format_exc()}"
            self.log(error_msg)
            QMessageBox.warning(self, "Error", f"Could not detach Universal Lookup:\n{str(e)}")
    
    def reattach_universal_lookup(self):
        """Re-attach Universal Lookup to the Home tab"""
        if not self.lookup_detached_window:
            return
        
        # Copy state back to home widget if it exists
        if (hasattr(self, 'home_lookup_widget') and self.home_lookup_widget and 
            hasattr(self, 'lookup_detached_widget') and self.lookup_detached_widget):
            # Copy source text
            source_text = self.lookup_detached_widget.source_text.toPlainText()
            self.home_lookup_widget.source_text.setPlainText(source_text)
            
            # Copy mode
            mode = self.lookup_detached_widget.mode_combo.currentText()
            self.home_lookup_widget.mode_combo.setCurrentText(mode)
        
        # Close detached window
        self.lookup_detached_window.close()
        self.lookup_detached_window = None
        if hasattr(self, 'lookup_detached_widget'):
            self.lookup_detached_widget = None
        
        self.log("Universal Lookup re-attached to Home tab")
    
    def create_projects_manager_tab(self):
        """Create the Projects Manager tab - manage projects, attach TMs and glossaries (legacy - content moved to Home)"""
        # This method is kept for backwards compatibility but content is now in Home tab
        return self.create_home_tab()
    
    def create_resources_tab(self):
        """Create the Resources tab with nested sub-tabs"""
        from PyQt6.QtWidgets import QTabWidget
        
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        # Create nested tab widget
        resources_tabs = QTabWidget()
        self.resources_tabs = resources_tabs  # Store for navigation
        
        # Add nested tabs
        tm_tab = self.create_translation_memories_tab()
        resources_tabs.addTab(tm_tab, "💾 Translation Memories")
        
        termbase_tab = self.create_termbases_tab()
        resources_tabs.addTab(termbase_tab, "🏷️ Termbases")
        
        nt_tab = self.create_non_translatables_tab()
        resources_tabs.addTab(nt_tab, "🚫 Non-Translatables")
        
        # Segmentation Rules
        seg_tab = self.create_segmentation_rules_tab()
        resources_tabs.addTab(seg_tab, "📏 Segmentation Rules")
        
        ref_tab = self.create_reference_images_tab()
        resources_tabs.addTab(ref_tab, "🖼️ Reference Images")
        
        layout.addWidget(resources_tabs)
        
        return tab
    
    def create_specialised_modules_tab(self):
        """Create the Specialised Modules tab with nested sub-tabs"""
        from PyQt6.QtWidgets import QTabWidget
        
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        # Create nested tab widget
        modules_tabs = QTabWidget()
        self.modules_tabs = modules_tabs  # Store for navigation
        
        # Add nested tabs
        tmx_tab = self.create_tmx_editor_tab()
        modules_tabs.addTab(tmx_tab, "✏️ TMX Editor")
        
        pdf_tab = self.create_pdf_rescue_tab()
        modules_tabs.addTab(pdf_tab, "📄 PDF Rescue")
        
        autofingers_tab = AutoFingersWidget(self)
        modules_tabs.addTab(autofingers_tab, "✋ AutoFingers")
        
        encoding_tab = self.create_encoding_repair_tab()
        modules_tabs.addTab(encoding_tab, "🔧 Text Encoding Repair")
        
        tracked_tab = self.create_tracked_changes_tab()
        modules_tabs.addTab(tracked_tab, "🔄 Tracked Changes")
        
        lookup_tab = UniversalLookupTab(self)
        modules_tabs.addTab(lookup_tab, "🔍 Universal Lookup")

        # LLM Leaderboard
        leaderboard_tab = self.create_llm_leaderboard_tab()
        modules_tabs.addTab(leaderboard_tab, "🏆 LLM Leaderboard")

        layout.addWidget(modules_tabs)

        return tab
    
    def create_translation_memories_tab(self):
        """Create the Translation Memories tab - manage all TMs (global)"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(10, 10, 10, 10)
        
        # Header
        header = QLabel("💾 Translation Memories")
        header.setStyleSheet("font-size: 14px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(header)
        
        # Description
        desc = QLabel("Manage translation memories for project matching. View statistics and import/export TMs.")
        desc.setStyleSheet("color: #666; font-size: 11px; margin-bottom: 10px;")
        layout.addWidget(desc)
        
        # Check if TM database is available
        if not hasattr(self, 'tm_database') or not self.tm_database:
            placeholder = QLabel("Translation Memories\n\nTM database not initialized.\nLoad a project to initialize TM system.")
            placeholder.setAlignment(Qt.AlignmentFlag.AlignCenter)
            placeholder.setStyleSheet("color: #888; font-size: 12px;")
            layout.addWidget(placeholder, stretch=1)
            return tab
        
        # TM Statistics
        stats_group = QGroupBox("TM Statistics")
        stats_layout = QVBoxLayout()
        
        try:
            # Get TM stats from database
            import sqlite3
            db_path = self.user_data_path / "Translation_Resources" / "supervertaler.db"
            conn = sqlite3.connect(str(db_path))
            cursor = conn.cursor()
            
            # Count total entries
            cursor.execute("SELECT COUNT(*) FROM translation_units")
            total_entries = cursor.fetchone()[0]
            
            # Count by TM ID
            cursor.execute("SELECT tm_id, COUNT(*) FROM translation_units GROUP BY tm_id")
            tm_stats = cursor.fetchall()
            
            conn.close()
            
            # Display stats
            stats_info = QLabel(f"Total Translation Units: {total_entries}")
            stats_info.setStyleSheet("font-weight: bold; color: #2563eb;")
            stats_layout.addWidget(stats_info)
            
            if tm_stats:
                stats_layout.addWidget(QLabel("TM Breakdown:"))
                for tm_id, count in tm_stats:
                    tm_label = QLabel(f"  • {tm_id}: {count} entries")
                    tm_label.setStyleSheet("margin-left: 15px; color: #555;")
                    stats_layout.addWidget(tm_label)
            
        except Exception as e:
            error_label = QLabel(f"Error loading TM statistics: {e}")
            error_label.setStyleSheet("color: #dc2626;")
            stats_layout.addWidget(error_label)
        
        stats_group.setLayout(stats_layout)
        layout.addWidget(stats_group)
        
        # TM Management Buttons
        buttons_group = QGroupBox("TM Management")
        buttons_layout = QVBoxLayout()
        
        # Import TMX button
        import_btn = QPushButton("📥 Import TMX File")
        import_btn.setToolTip("Import translation memory from TMX file")
        import_btn.clicked.connect(self.import_tmx_file)
        buttons_layout.addWidget(import_btn)
        
        # Export TMX button
        export_btn = QPushButton("📤 Export TM as TMX")
        export_btn.setToolTip("Export current translation memory to TMX file")
        export_btn.clicked.connect(self.export_tm_as_tmx)
        buttons_layout.addWidget(export_btn)
        
        # Clear TM button (with warning)
        clear_btn = QPushButton("🗑️ Clear All TM Entries")
        clear_btn.setToolTip("WARNING: This will delete all translation memory entries!")
        clear_btn.setStyleSheet("color: #dc2626; font-weight: bold;")
        clear_btn.clicked.connect(self.clear_tm_entries)
        buttons_layout.addWidget(clear_btn)
        
        buttons_group.setLayout(buttons_layout)
        layout.addWidget(buttons_group)
        
        layout.addStretch()
        
        return tab
    
    def import_tmx_file(self):
        """Import TMX file into translation memory"""
        try:
            file_path, _ = QFileDialog.getOpenFileName(
                self, 
                "Import TMX File", 
                "", 
                "TMX Files (*.tmx);;All Files (*.*)"
            )
            
            if file_path and self.tm_database:
                # Import TMX using the TM database
                self.tm_database.import_tmx(file_path)
                self.log(f"Successfully imported TMX file: {file_path}")
                QMessageBox.information(self, "Import Complete", f"TMX file imported successfully!\n\nFile: {file_path}")
            
        except Exception as e:
            self.log(f"Error importing TMX file: {e}")
            QMessageBox.critical(self, "Import Error", f"Failed to import TMX file:\n\n{e}")
    
    def export_tmx_from_grid(self):
        """Export all segments from current project grid as TMX file"""
        try:
            if not self.current_project or not self.current_project.segments:
                QMessageBox.warning(self, "No Project", "Please open a project first")
                return
            
            file_path, _ = QFileDialog.getSaveFileName(
                self, 
                "Export Grid as TMX", 
                "supervertaler_grid.tmx", 
                "TMX Files (*.tmx);;All Files (*.*)"
            )
            
            if not file_path:
                return
            
            source_segments = []
            target_segments = []
            entry_count = 0
            
            # Export all segments from grid (including empty translations)
            self.log("Exporting all segments from grid...")
            for segment in self.current_project.segments:
                source_text = segment.source
                target_text = segment.target if segment.target else ""
                
                if source_text:
                    source_segments.append(source_text)
                    target_segments.append(target_text)
                    entry_count += 1
            
            self.log(f"Found {entry_count} segments in grid")
            
            if entry_count == 0:
                QMessageBox.warning(self, "No Data", "No segments found in grid")
                return
            
            # Generate and save TMX
            from modules.tmx_generator import TMXGenerator
            tmx_generator = TMXGenerator(log_callback=self.log)
            
            source_lang = self.current_project.source_lang or "en"
            target_lang = self.current_project.target_lang or "nl"
            
            tmx_tree = tmx_generator.generate_tmx(
                source_segments=source_segments,
                target_segments=target_segments,
                source_lang=source_lang,
                target_lang=target_lang
            )
            
            if tmx_generator.save_tmx(tmx_tree, file_path):
                self.log(f"✓ Exported TMX from grid: {file_path} ({entry_count} segments)")
                QMessageBox.information(
                    self, 
                    "Export Complete", 
                    f"Grid exported to TMX successfully!\n\n"
                    f"File: {file_path}\n"
                    f"Segments: {entry_count}\n"
                    f"Language pair: {source_lang} → {target_lang}"
                )
            else:
                QMessageBox.warning(self, "Export Error", "Failed to save TMX file")
            
        except Exception as e:
            self.log(f"✗ Error exporting grid as TMX: {e}")
            import traceback
            self.log(traceback.format_exc())
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n\n{str(e)}")
    
    def export_tmx_from_selected(self):
        """Export only selected segments from grid as TMX file"""
        try:
            if not self.current_project or not self.current_project.segments:
                QMessageBox.warning(self, "No Project", "Please open a project first")
                return
            
            # Get selected rows from table
            selected_rows = set()
            for index in self.table.selectedIndexes():
                selected_rows.add(index.row())
            
            if not selected_rows:
                QMessageBox.warning(
                    self, 
                    "No Selection", 
                    "Please select one or more rows in the grid first.\n\n"
                    "Click on row numbers or use Ctrl+Click to select multiple rows."
                )
                return
            
            file_path, _ = QFileDialog.getSaveFileName(
                self, 
                "Export Selected Segments as TMX", 
                "supervertaler_selected.tmx", 
                "TMX Files (*.tmx);;All Files (*.*)"
            )
            
            if not file_path:
                return
            
            source_segments = []
            target_segments = []
            entry_count = 0
            
            # Export only selected segments
            self.log(f"Exporting {len(selected_rows)} selected segments...")
            for row in sorted(selected_rows):
                if row < len(self.current_project.segments):
                    segment = self.current_project.segments[row]
                    source_text = segment.source
                    target_text = segment.target if segment.target else ""
                    
                    if source_text:
                        source_segments.append(source_text)
                        target_segments.append(target_text)
                        entry_count += 1
            
            self.log(f"Selected {entry_count} segments to export")
            
            if entry_count == 0:
                QMessageBox.warning(self, "No Data", "No valid segments selected")
                return
            
            # Generate and save TMX
            from modules.tmx_generator import TMXGenerator
            tmx_generator = TMXGenerator(log_callback=self.log)
            
            source_lang = self.current_project.source_lang or "en"
            target_lang = self.current_project.target_lang or "nl"
            
            tmx_tree = tmx_generator.generate_tmx(
                source_segments=source_segments,
                target_segments=target_segments,
                source_lang=source_lang,
                target_lang=target_lang
            )
            
            if tmx_generator.save_tmx(tmx_tree, file_path):
                self.log(f"✓ Exported TMX from selected: {file_path} ({entry_count} segments)")
                QMessageBox.information(
                    self, 
                    "Export Complete", 
                    f"Selected segments exported to TMX successfully!\n\n"
                    f"File: {file_path}\n"
                    f"Segments: {entry_count}\n"
                    f"Language pair: {source_lang} → {target_lang}"
                )
            else:
                QMessageBox.warning(self, "Export Error", "Failed to save TMX file")
            
        except Exception as e:
            self.log(f"✗ Error exporting selected segments as TMX: {e}")
            import traceback
            self.log(traceback.format_exc())
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n\n{str(e)}")
    
    def export_tmx_from_tm_database(self):
        """Export translation memory entries as TMX file"""
        try:
            if not self.tm_database:
                QMessageBox.warning(self, "Error", "Translation memory database not available")
                return
            
            # Get TM entries from database
            tm_entries = self.tm_database.get_tm_entries(tm_id=None, limit=None)
            if not tm_entries:
                QMessageBox.warning(
                    self, 
                    "No Data", 
                    "No translation units in Translation Memory database.\n\n"
                    "Please import a TM or save translations to the TM first."
                )
                return
            
            file_path, _ = QFileDialog.getSaveFileName(
                self, 
                "Export TM Database as TMX", 
                "supervertaler_tm_database.tmx", 
                "TMX Files (*.tmx);;All Files (*.*)"
            )
            
            if not file_path:
                return
            
            source_segments = []
            target_segments = []
            
            # Extract segments from TM entries
            self.log(f"Exporting {len(tm_entries)} entries from TM database...")
            for entry in tm_entries:
                source_segments.append(entry.get('source_text', entry.get('source', '')))
                target_segments.append(entry.get('target_text', entry.get('target', '')))
            
            # Generate and save TMX
            from modules.tmx_generator import TMXGenerator
            tmx_generator = TMXGenerator(log_callback=self.log)
            
            # Use project languages or defaults
            source_lang = "en"
            target_lang = "nl"
            if self.current_project:
                source_lang = self.current_project.source_lang or "en"
                target_lang = self.current_project.target_lang or "nl"
            
            tmx_tree = tmx_generator.generate_tmx(
                source_segments=source_segments,
                target_segments=target_segments,
                source_lang=source_lang,
                target_lang=target_lang
            )
            
            if tmx_generator.save_tmx(tmx_tree, file_path):
                self.log(f"✓ Exported TMX from TM database: {file_path} ({len(tm_entries)} entries)")
                QMessageBox.information(
                    self, 
                    "Export Complete", 
                    f"Translation Memory exported to TMX successfully!\n\n"
                    f"File: {file_path}\n"
                    f"Entries: {len(tm_entries)}\n"
                    f"Language pair: {source_lang} → {target_lang}"
                )
            else:
                QMessageBox.warning(self, "Export Error", "Failed to save TMX file")
            
        except Exception as e:
            self.log(f"✗ Error exporting TM database as TMX: {e}")
            import traceback
            self.log(traceback.format_exc())
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n\n{str(e)}")
    
    def export_tm_as_tmx(self):
        """Legacy function - exports grid segments (for backward compatibility)"""
        self.export_tmx_from_grid()
    
    
    
    
    def clear_tm_entries(self):
        """Clear all translation memory entries (with confirmation)"""
        try:
            reply = QMessageBox.question(
                self, 
                "Clear Translation Memory", 
                "⚠️ WARNING: This will permanently delete ALL translation memory entries!\n\nThis action cannot be undone. Are you sure?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                # Double confirmation
                reply2 = QMessageBox.question(
                    self, 
                    "Final Confirmation", 
                    "Last chance! Are you absolutely sure you want to delete all TM entries?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                    QMessageBox.StandardButton.No
                )
                
                if reply2 == QMessageBox.StandardButton.Yes and self.tm_database:
                    # Clear TM entries
                    import sqlite3
                    db_path = self.user_data_path / "Translation_Resources" / "supervertaler.db"
                    conn = sqlite3.connect(str(db_path))
                    cursor = conn.cursor()
                    cursor.execute("DELETE FROM translation_units")
                    conn.commit()
                    conn.close()
                    
                    self.log("All translation memory entries cleared")
                    QMessageBox.information(self, "TM Cleared", "All translation memory entries have been deleted.")
            
        except Exception as e:
            self.log(f"Error clearing TM entries: {e}")
            QMessageBox.critical(self, "Clear Error", f"Failed to clear TM entries:\n\n{e}")
    
    def create_segmentation_rules_tab(self):
        """Create Segmentation Rules management tab"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(10, 10, 10, 10)
        
        # Header
        header = QLabel("📏 Segmentation Rules")
        header.setStyleSheet("font-size: 14px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(header)
        
        # Description
        desc = QLabel("Manage language-specific segmentation rules for accurate sentence/segment boundaries.")
        desc.setStyleSheet("color: #666; font-size: 11px; margin-bottom: 10px;")
        layout.addWidget(desc)
        
        # Language Selection
        lang_group = QGroupBox("Select Language")
        lang_layout = QHBoxLayout()
        
        lang_combo = QComboBox()
        languages = ["English (en)", "Dutch (nl)", "German (de)", "French (fr)", "Spanish (es)", "Italian (it)", "Portuguese (pt)", "Chinese (zh)", "Japanese (ja)", "Arabic (ar)"]
        lang_combo.addItems(languages)
        lang_layout.addWidget(QLabel("Language:"))
        lang_layout.addWidget(lang_combo, 1)
        
        lang_group.setLayout(lang_layout)
        layout.addWidget(lang_group)
        
        # Segmentation Rules
        rules_group = QGroupBox("Segmentation Rules")
        rules_layout = QVBoxLayout()
        
        # Current implementation info
        current_info = QLabel(
            "Current Implementation: SimpleSegmenter (language-agnostic)\n\n"
            "• Segments on: . ! ? (followed by space/newline)\n"
            "• Handles basic abbreviations: Mr. Dr. etc.\n"
            "• Preserves paragraph breaks\n"
            "• Treats each table cell as separate segment"
        )
        current_info.setStyleSheet("background-color: #f3f4f6; padding: 10px; border-radius: 4px; font-family: monospace;")
        rules_layout.addWidget(current_info)
        
        # Future implementation section
        future_label = QLabel("🚧 Language-Specific Rules (Planned):")
        future_label.setStyleSheet("font-weight: bold; margin-top: 15px;")
        rules_layout.addWidget(future_label)
        
        future_info = QLabel(
            "• German: Handle compound sentences, different abbreviations\n"
            "• Chinese/Japanese: Word boundary detection, different punctuation\n"
            "• Arabic: Right-to-left text handling\n"
            "• French: Quotation mark handling, spacing rules\n"
            "• Custom: User-defined regex patterns and exceptions"
        )
        future_info.setStyleSheet("color: #666; margin-left: 20px;")
        rules_layout.addWidget(future_info)
        
        rules_group.setLayout(rules_layout)
        layout.addWidget(rules_group)
        
        # Management buttons (for future implementation)
        buttons_group = QGroupBox("Rule Management")
        buttons_layout = QVBoxLayout()
        
        # Test segmentation button
        test_btn = QPushButton("🧪 Test Segmentation")
        test_btn.setToolTip("Test current segmentation rules on sample text")
        test_btn.clicked.connect(self.test_segmentation_rules)
        buttons_layout.addWidget(test_btn)
        
        # Import/Export buttons (disabled for now)
        import_btn = QPushButton("📥 Import Rules")
        import_btn.setToolTip("Import segmentation rules from file (Coming Soon)")
        import_btn.setEnabled(False)
        buttons_layout.addWidget(import_btn)
        
        export_btn = QPushButton("📤 Export Rules")
        export_btn.setToolTip("Export current segmentation rules (Coming Soon)")
        export_btn.setEnabled(False)
        buttons_layout.addWidget(export_btn)
        
        buttons_group.setLayout(buttons_layout)
        layout.addWidget(buttons_group)
        
        layout.addStretch()
        
        return tab
    
    def test_segmentation_rules(self):
        """Test current segmentation rules with sample text"""
        from PyQt6.QtWidgets import QInputDialog
        
        sample_text = (
            "Hello world. This is a test! How are you? "
            "Mr. Smith went to Dr. Jones. "
            "The U.S.A. is great. "
            "What about this... and that? "
            "End of test."
        )
        
        text, ok = QInputDialog.getMultiLineText(
            self, 
            "Test Segmentation", 
            "Enter text to test segmentation:\n(Default sample text provided)",
            sample_text
        )
        
        if ok and text:
            try:
                # Test with current segmenter
                from modules.simple_segmenter import SimpleSegmenter
                segmenter = SimpleSegmenter()
                
                # Create fake paragraph list for testing
                paragraphs = [(0, text)]
                segments = segmenter.segment_paragraphs(paragraphs)
                
                # Show results
                result_text = f"Input text:\n{text}\n\n"
                result_text += f"Segmentation results ({len(segments)} segments):\n"
                result_text += "=" * 50 + "\n"
                
                for i, (para_id, segment_text) in enumerate(segments, 1):
                    result_text += f"{i}. {segment_text}\n"
                
                QMessageBox.information(
                    self,
                    "Segmentation Test Results",
                    result_text
                )
                
            except Exception as e:
                QMessageBox.critical(
                    self,
                    "Segmentation Test Error",
                    f"Error testing segmentation:\n\n{e}"
                )
    
    def create_termbases_tab(self):
        """Create the Termbases tab - manage all termbases (global and project-specific)"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(10, 10, 10, 10)
        
        # Header
        header = QLabel("📚 Termbases")
        header.setStyleSheet("font-size: 14px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(header)
        
        # Description
        desc = QLabel("Manage termbases for terminology searching. Activate/deactivate for current project.")
        desc.setStyleSheet("color: #666; font-size: 11px; margin-bottom: 10px;")
        layout.addWidget(desc)
        
        # Check if database is available
        if not (hasattr(self, 'db_manager') and self.db_manager):
            placeholder = QLabel("Termbases Manager\n\nDatabase not initialized.")
            placeholder.setAlignment(Qt.AlignmentFlag.AlignCenter)
            placeholder.setStyleSheet("color: #888; font-size: 12px;")
            layout.addWidget(placeholder, stretch=1)
            return tab
        
        # Import here to avoid issues if database not available
        from modules.termbase_manager import TermbaseManager
        termbase_mgr = TermbaseManager(self.db_manager, self.log)
        self.termbase_mgr = termbase_mgr  # Store for later use in get_termbase_code
        
        # Search bar
        search_layout = QHBoxLayout()
        search_box = QLineEdit()
        search_box.setPlaceholderText("Search termbases...")
        search_box.setMaximumWidth(300)
        search_layout.addWidget(search_box)
        search_layout.addStretch()
        layout.addLayout(search_layout)
        
        # Termbase list with table
        termbase_table = QTableWidget()
        termbase_table.setColumnCount(6)
        termbase_table.setHorizontalHeaderLabels(["Active", "Name", "Languages", "Terms", "Priority", "Scope"])
        termbase_table.horizontalHeader().setStretchLastSection(False)
        termbase_table.setColumnWidth(0, 60)
        termbase_table.setColumnWidth(1, 200)
        termbase_table.setColumnWidth(2, 150)
        termbase_table.setColumnWidth(3, 80)
        termbase_table.setColumnWidth(4, 70)  # Priority column
        termbase_table.setColumnWidth(5, 100)
        
        # Get current project
        current_project = self.current_project if hasattr(self, 'current_project') else None
        # current_project is a Project object, not a dict
        project_id = current_project.id if (current_project and hasattr(current_project, 'id')) else None
        
        # Populate termbase list
        def refresh_termbase_list():
            termbases = termbase_mgr.get_all_termbases()
            termbase_table.setRowCount(len(termbases))
            
            for row, tb in enumerate(termbases):
                # Check if active for current project
                is_active = termbase_mgr.is_termbase_active(tb['id'], project_id) if project_id else True
                
                # Active checkbox
                checkbox = QCheckBox()
                checkbox.setChecked(is_active)
                checkbox.toggled.connect(lambda checked, tb_id=tb['id']: 
                    termbase_mgr.activate_termbase(tb_id, project_id) if checked and project_id else termbase_mgr.deactivate_termbase(tb_id, project_id) if project_id else None)
                termbase_table.setCellWidget(row, 0, checkbox)
                
                # Name (bold if active)
                name_item = QTableWidgetItem(tb['name'])
                if is_active:
                    font = name_item.font()
                    font.setBold(True)
                    name_item.setFont(font)
                termbase_table.setItem(row, 1, name_item)
                
                # Languages
                langs = f"{tb['source_lang'] or '?'} → {tb['target_lang'] or '?'}"
                termbase_table.setItem(row, 2, QTableWidgetItem(langs))
                
                # Term count
                termbase_table.setItem(row, 3, QTableWidgetItem(str(tb['term_count'])))
                
                # Priority (editable)
                priority = tb.get('priority', 50)  # Default 50 if not set
                priority_item = QTableWidgetItem(str(priority))
                priority_item.setFlags(priority_item.flags() | Qt.ItemFlag.ItemIsEditable)
                termbase_table.setItem(row, 4, priority_item)
                
                # Scope
                scope = "Global" if tb['is_global'] else "Project"
                termbase_table.setItem(row, 5, QTableWidgetItem(scope))
        
        # Handle priority changes - only process changes to Priority column (column 4)
        priority_changing = False  # Flag to prevent recursion
        
        def on_priority_changed(item: QTableWidgetItem):
            """Update termbase priority when edited"""
            # Only process Priority column (column 4)
            if item.column() != 4:
                return
            
            # Prevent recursion
            nonlocal priority_changing
            if priority_changing:
                return
            
            row = item.row()
            termbases_list = termbase_mgr.get_all_termbases()
            if row < len(termbases_list):
                termbase_id = termbases_list[row]['id']
                try:
                    new_priority = int(item.text())
                    # Clamp to valid range 1-99
                    new_priority = max(1, min(99, new_priority))
                    
                    # Temporarily disconnect to prevent recursion when setting text
                    priority_changing = True
                    termbase_table.itemChanged.disconnect(on_priority_changed)
                    item.setText(str(new_priority))
                    termbase_table.itemChanged.connect(on_priority_changed)
                    priority_changing = False
                    
                    # Update in database
                    cursor = self.db_manager.cursor
                    cursor.execute("UPDATE termbases SET priority = ?, modified_date = CURRENT_TIMESTAMP WHERE id = ?",
                                 (new_priority, termbase_id))
                    self.db_manager.connection.commit()
                    self.log(f"✓ Updated priority for termbase {termbase_id} to {new_priority}")
                except ValueError:
                    # Invalid input, revert to original value
                    priority_changing = True
                    termbase_table.itemChanged.disconnect(on_priority_changed)
                    termbases_list = termbase_mgr.get_all_termbases()
                    if row < len(termbases_list):
                        original_priority = termbases_list[row].get('priority', 50)
                        item.setText(str(original_priority))
                    termbase_table.itemChanged.connect(on_priority_changed)
                    priority_changing = False
                    self.log(f"⚠ Invalid priority value, reverted to original")
        
        termbase_table.itemChanged.connect(on_priority_changed)
        
        refresh_termbase_list()
        layout.addWidget(termbase_table, stretch=1)
        
        # Button bar
        button_layout = QHBoxLayout()
        
        create_btn = QPushButton("+ Create New")
        create_btn.clicked.connect(lambda: self._show_create_termbase_dialog(termbase_mgr, refresh_termbase_list, project_id))
        button_layout.addWidget(create_btn)
        
        import_btn = QPushButton("📥 Import")
        button_layout.addWidget(import_btn)
        
        export_btn = QPushButton("📤 Export")
        button_layout.addWidget(export_btn)
        
        delete_btn = QPushButton("🗑️ Delete")
        button_layout.addWidget(delete_btn)
        
        edit_btn = QPushButton("✏️ Edit Terms")
        edit_btn.clicked.connect(lambda: self._show_edit_terms_dialog(termbase_mgr, termbase_table, refresh_termbase_list))
        button_layout.addWidget(edit_btn)
        
        button_layout.addStretch()
        layout.addLayout(button_layout)
        
        return tab
    
    def _show_create_termbase_dialog(self, termbase_mgr, refresh_callback, project_id):
        """Show dialog to create new termbase"""
        from modules.termbase_manager import TermbaseManager
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Create New Termbase")
        dialog.setMinimumWidth(400)
        
        layout = QFormLayout()
        
        # Name
        name_field = QLineEdit()
        layout.addRow("Termbase Name:", name_field)
        
        # Source language
        source_lang_field = QLineEdit()
        source_lang_field.setPlaceholderText("e.g., en, nl, de")
        layout.addRow("Source Language:", source_lang_field)
        
        # Target language
        target_lang_field = QLineEdit()
        target_lang_field.setPlaceholderText("e.g., en, nl, de")
        layout.addRow("Target Language:", target_lang_field)
        
        # Description
        desc_field = QTextEdit()
        desc_field.setMaximumHeight(80)
        layout.addRow("Description:", desc_field)
        
        # Scope
        scope_group = QButtonGroup()
        global_radio = QRadioButton("Global (all projects)")
        global_radio.setChecked(True)
        project_radio = QRadioButton("Project-specific")
        scope_group.addButton(global_radio, 0)
        scope_group.addButton(project_radio, 1)
        
        scope_layout = QHBoxLayout()
        scope_layout.addWidget(global_radio)
        scope_layout.addWidget(project_radio)
        layout.addRow("Scope:", scope_layout)
        
        # Buttons
        button_layout = QHBoxLayout()
        create_btn = QPushButton("Create")
        cancel_btn = QPushButton("Cancel")
        
        def create_termbase():
            name = name_field.text().strip()
            if not name:
                QMessageBox.warning(dialog, "Error", "Please enter a termbase name")
                return
            
            source_lang = source_lang_field.text().strip() or None
            target_lang = target_lang_field.text().strip() or None
            description = desc_field.toPlainText().strip()
            is_global = scope_group.checkedId() == 0
            
            tb_project_id = None if is_global else project_id
            
            tb_id = termbase_mgr.create_termbase(
                name=name,
                source_lang=source_lang,
                target_lang=target_lang,
                project_id=tb_project_id,
                description=description,
                is_global=is_global
            )
            
            if tb_id:
                QMessageBox.information(dialog, "Success", f"Termbase '{name}' created successfully!")
                refresh_callback()
                dialog.accept()
            else:
                QMessageBox.critical(dialog, "Error", "Failed to create termbase")
        
        create_btn.clicked.connect(create_termbase)
        cancel_btn.clicked.connect(dialog.reject)
        
        button_layout.addStretch()
        button_layout.addWidget(create_btn)
        button_layout.addWidget(cancel_btn)
        layout.addRow("", button_layout)
        
        dialog.setLayout(layout)
        dialog.exec()
    
    def _show_edit_terms_dialog(self, termbase_mgr, termbase_table, refresh_callback):
        """Show dialog to edit terms in selected termbase"""
        selected_row = termbase_table.currentRow()
        if selected_row < 0:
            QMessageBox.warning(self, "Error", "Please select a termbase first")
            return
        
        # Get termbase ID from table
        tb_name = termbase_table.item(selected_row, 1).text()
        
        # Find termbase ID
        termbases = termbase_mgr.get_all_termbases()
        termbase = next((tb for tb in termbases if tb['name'] == tb_name), None)
        if not termbase:
            QMessageBox.warning(self, "Error", "Could not find selected termbase")
            return
        
        termbase_id = termbase['id']
        
        # Create dialog
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Edit Terms - {tb_name}")
        dialog.setMinimumSize(700, 500)
        
        layout = QVBoxLayout()
        
        # Terms table
        terms_table = QTableWidget()
        terms_table.setColumnCount(5)
        terms_table.setHorizontalHeaderLabels(["Source", "Target", "Domain", "Priority", "Forbidden"])
        terms_table.horizontalHeader().setStretchLastSection(False)
        terms_table.setColumnWidth(0, 150)
        terms_table.setColumnWidth(1, 150)
        terms_table.setColumnWidth(2, 150)
        terms_table.setColumnWidth(3, 80)
        terms_table.setColumnWidth(4, 80)
        
        # Load terms
        def refresh_terms_table():
            terms = termbase_mgr.get_terms(termbase_id)
            terms_table.setRowCount(len(terms))
            
            for row, term in enumerate(terms):
                terms_table.setItem(row, 0, QTableWidgetItem(term['source_term']))
                terms_table.setItem(row, 1, QTableWidgetItem(term['target_term']))
                terms_table.setItem(row, 2, QTableWidgetItem(term['domain'] or ""))
                terms_table.setItem(row, 3, QTableWidgetItem(str(term['priority'])))
                
                forbidden_check = QCheckBox()
                forbidden_check.setChecked(term['forbidden'])
                terms_table.setCellWidget(row, 4, forbidden_check)
        
        refresh_terms_table()
        layout.addWidget(QLabel(f"Terms in '{tb_name}':"), 0)
        layout.addWidget(terms_table, 1)
        
        # Add term section
        add_layout = QHBoxLayout()
        
        source_field = QLineEdit()
        source_field.setPlaceholderText("Source term")
        add_layout.addWidget(source_field)
        
        target_field = QLineEdit()
        target_field.setPlaceholderText("Target term")
        add_layout.addWidget(target_field)
        
        priority_spin = QSpinBox()
        priority_spin.setMinimum(1)
        priority_spin.setMaximum(99)
        priority_spin.setValue(50)
        priority_spin.setMaximumWidth(60)
        add_layout.addWidget(QLabel("Priority:"))
        add_layout.addWidget(priority_spin)
        
        add_btn = QPushButton("+ Add")
        def add_term():
            source = source_field.text().strip()
            target = target_field.text().strip()
            if not source or not target:
                QMessageBox.warning(dialog, "Error", "Both source and target terms are required")
                return
            
            termbase_mgr.add_term(
                termbase_id=termbase_id,
                source_term=source,
                target_term=target,
                priority=priority_spin.value()
            )
            
            source_field.clear()
            target_field.clear()
            refresh_terms_table()
        
        add_btn.clicked.connect(add_term)
        add_layout.addWidget(add_btn)
        layout.addLayout(add_layout)
        
        # Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.accept)
        layout.addWidget(close_btn)
        
        dialog.setLayout(layout)
        dialog.exec()
    
    def create_settings_tab(self):
        """Create the Settings tab - moved from Tools > Options dialog"""
        from PyQt6.QtWidgets import QTabWidget, QScrollArea
        
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        # Create nested tab widget
        settings_tabs = QTabWidget()
        self.settings_tabs = settings_tabs  # Store for reference
        
        # Scroll area wrapper for each tab (for long content)
        scroll_area_wrapper = lambda widget: self._wrap_in_scroll(widget)
        
        # ===== TAB 1: General Settings =====
        general_tab = self._create_general_settings_tab()
        settings_tabs.addTab(scroll_area_wrapper(general_tab), "⚙️ General")
        
        # ===== TAB 2: LLM Settings =====
        llm_tab = self._create_llm_settings_tab()
        settings_tabs.addTab(scroll_area_wrapper(llm_tab), "🤖 LLM Settings")
        
        # ===== TAB 3: Language Pair Settings =====
        lang_tab = self._create_language_pair_tab()
        settings_tabs.addTab(scroll_area_wrapper(lang_tab), "🌐 Language Pair")
        
        # ===== TAB 4: MT Settings =====
        mt_tab = self._create_mt_settings_tab()
        settings_tabs.addTab(scroll_area_wrapper(mt_tab), "🌐 MT Settings")
        
        # ===== TAB 5: View/Display Settings =====
        view_tab = self._create_view_settings_tab()
        settings_tabs.addTab(scroll_area_wrapper(view_tab), "🔍 View/Display")

        # ===== TAB 6: System Prompts (Layer 1) =====
        system_prompts_tab = self._create_system_prompts_tab()
        settings_tabs.addTab(scroll_area_wrapper(system_prompts_tab), "📝 System Prompts")

        # ===== TAB 7: Voice Dictation Settings =====
        dictation_tab = self._create_voice_dictation_settings_tab()
        settings_tabs.addTab(scroll_area_wrapper(dictation_tab), "🎤 Voice Dictation")

        # ===== TAB 8: Keyboard Shortcuts =====
        from modules.keyboard_shortcuts_widget import KeyboardShortcutsWidget
        shortcuts_tab = KeyboardShortcutsWidget(self)
        settings_tabs.addTab(shortcuts_tab, "⌨️ Keyboard Shortcuts")

        # ===== TAB 9: Log (moved from main tabs) =====
        log_tab = self.create_log_tab()
        settings_tabs.addTab(log_tab, "📋 Log")
        
        layout.addWidget(settings_tabs)
        
        return tab
    
    def _wrap_in_scroll(self, widget):
        """Wrap a widget in a scroll area"""
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setWidget(widget)
        scroll.setStyleSheet("QScrollArea { border: none; }")
        return scroll
    
    def _create_language_pair_tab(self):
        """Create Language Pair Settings tab content"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)
        
        # Language Pair Settings group
        lang_group = QGroupBox("Translation Language Pair")
        lang_layout = QVBoxLayout()
        
        info_label = QLabel(
            "Set your default source and target languages for translation projects."
        )
        info_label.setWordWrap(True)
        info_label.setStyleSheet("font-size: 9pt; color: #666; padding: 8px; background-color: #f3f4f6; border-radius: 3px;")
        lang_layout.addWidget(info_label)
        
        # Language selection layout
        lang_select_layout = QHBoxLayout()
        
        # Source language
        source_label = QLabel("Source Language:")
        source_combo = QComboBox()
        
        # Available languages (from Tkinter version)
        available_languages = [
            "Afrikaans", "Albanian", "Arabic", "Armenian", "Basque", "Bengali",
            "Bulgarian", "Catalan", "Chinese (Simplified)", "Chinese (Traditional)",
            "Croatian", "Czech", "Danish", "Dutch", "English", "Estonian",
            "Finnish", "French", "Galician", "Georgian", "German", "Greek",
            "Hebrew", "Hindi", "Hungarian", "Icelandic", "Indonesian", "Irish",
            "Italian", "Japanese", "Korean", "Latvian", "Lithuanian", "Macedonian",
            "Malay", "Norwegian", "Persian", "Polish", "Portuguese", "Romanian",
            "Russian", "Serbian", "Slovak", "Slovenian", "Spanish", "Swahili",
            "Swedish", "Thai", "Turkish", "Ukrainian", "Urdu", "Vietnamese", "Welsh"
        ]
        source_combo.addItems(available_languages)
        source_combo.setCurrentText(self.source_language)
        
        lang_select_layout.addWidget(source_label)
        lang_select_layout.addWidget(source_combo, 1)
        
        # Swap button
        swap_btn = QPushButton("🔄 Swap")
        swap_btn.setToolTip("Swap source and target languages")
        lang_select_layout.addWidget(swap_btn)
        
        # Target language
        target_label = QLabel("Target Language:")
        target_combo = QComboBox()
        target_combo.addItems(available_languages)
        target_combo.setCurrentText(self.target_language)
        
        lang_select_layout.addWidget(target_label)
        lang_select_layout.addWidget(target_combo, 1)
        
        lang_layout.addLayout(lang_select_layout)
        
        # Swap functionality
        def on_swap():
            source_val = source_combo.currentText()
            target_val = target_combo.currentText()
            source_combo.setCurrentText(target_val)
            target_combo.setCurrentText(source_val)
        
        swap_btn.clicked.connect(on_swap)
        
        lang_group.setLayout(lang_layout)
        layout.addWidget(lang_group)
        
        # Save button
        save_btn = QPushButton("💾 Save Language Settings")
        save_btn.setStyleSheet("font-weight: bold; padding: 8px;")
        save_btn.clicked.connect(lambda: self._save_language_settings_from_ui(source_combo, target_combo))
        layout.addWidget(save_btn)
        
        layout.addStretch()
        
        return tab
    
    def _create_llm_settings_tab(self):
        """Create LLM Settings tab content"""
        from PyQt6.QtWidgets import QCheckBox, QGroupBox, QPushButton
        
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)
        
        # Load current settings
        settings = self.load_llm_settings()
        enabled_providers = self.load_provider_enabled_states()
        
        # LLM Provider Selection
        provider_group = QGroupBox("LLM Provider")
        provider_layout = QVBoxLayout()
        
        provider_label = QLabel("Select your preferred translation provider:")
        provider_layout.addWidget(provider_label)
        
        # Provider radio buttons (custom styled)
        provider_button_group = QButtonGroup(tab)
        
        openai_radio = CustomRadioButton("OpenAI (GPT-4o, GPT-5, o1, o3)")
        openai_radio.setChecked(settings.get('provider', 'openai') == 'openai')
        provider_button_group.addButton(openai_radio)
        provider_layout.addWidget(openai_radio)
        
        claude_radio = CustomRadioButton("Anthropic Claude (Claude Sonnet 4.5)")
        claude_radio.setChecked(settings.get('provider', 'openai') == 'claude')
        provider_button_group.addButton(claude_radio)
        provider_layout.addWidget(claude_radio)
        
        gemini_radio = CustomRadioButton("Google Gemini (Gemini 2.5 Flash)")
        gemini_radio.setChecked(settings.get('provider', 'openai') == 'gemini')
        provider_button_group.addButton(gemini_radio)
        provider_layout.addWidget(gemini_radio)
        
        provider_group.setLayout(provider_layout)
        layout.addWidget(provider_group)
        
        # Model Selection
        model_group = QGroupBox("Model Selection")
        model_layout = QVBoxLayout()
        
        model_label = QLabel("Choose the specific model to use:")
        model_layout.addWidget(model_label)
        
        # OpenAI models
        openai_model_label = QLabel("<b>OpenAI Models:</b>")
        model_layout.addWidget(openai_model_label)
        
        openai_combo = QComboBox()
        openai_combo.addItems([
            "gpt-4o (Recommended)",
            "gpt-4o-mini (Fast & Economical)",
            "gpt-5 (Reasoning, Temperature 1.0)",
            "o3-mini (Reasoning, Temperature 1.0)",
            "o1 (Reasoning, Temperature 1.0)",
            "gpt-4-turbo"
        ])
        openai_combo.setToolTip(
            "GPT-4o (Recommended): Fast, reliable, excellent for general translation.\n"
            "  • Best for: 90% of translation tasks\n"
            "  • Speed: 2-10 seconds per segment\n"
            "  • Use for: Business documents, technical manuals, general content\n\n"
            "GPT-4o-mini: Faster and cheaper, good quality for simple text.\n"
            "  • Best for: High-volume simple translations\n"
            "  • Speed: 1-5 seconds per segment\n\n"
            "GPT-5 (Reasoning): Deep analysis, handles complex/ambiguous text.\n"
            "  • Best for: Legal contracts, literary translation, marketing copy\n"
            "  • Speed: 30 seconds to 5+ minutes per segment\n"
            "  • Use when: Accuracy > speed, complex context, cultural nuances\n"
            "  • Cost: Premium pricing\n\n"
            "o3-mini/o1: Similar to GPT-5 but different reasoning approaches."
        )
        # Set current selection
        current_openai_model = settings.get('openai_model', 'gpt-4o')
        for i in range(openai_combo.count()):
            if current_openai_model in openai_combo.itemText(i).lower():
                openai_combo.setCurrentIndex(i)
                break
        openai_combo.setEnabled(openai_radio.isChecked())
        model_layout.addWidget(openai_combo)
        
        model_layout.addSpacing(10)
        
        # Claude models
        claude_model_label = QLabel("<b>Claude Models:</b>")
        model_layout.addWidget(claude_model_label)
        
        claude_combo = QComboBox()
        claude_combo.addItems([
            "claude-sonnet-4-5-20250929 (Recommended - Best Balance)",
            "claude-haiku-4-5-20251001 (Fast & Affordable)",
            "claude-opus-4-1-20250805 (Premium - Complex Reasoning)"
        ])
        claude_combo.setToolTip(
            "Claude Sonnet 4.5 (Recommended): Best balance of speed, quality, and cost.\n"
            "  • Pricing: $3/$15 per million tokens (input/output)\n"
            "  • Best for: General translation, multilingual content\n"
            "  • Strengths: Fast, cost-effective, excellent quality\n\n"
            "Claude Haiku 4.5: 2x faster, 1/3 the cost of Sonnet.\n"
            "  • Pricing: $1/$5 per million tokens\n"
            "  • Best for: Large translation projects, high-volume batch processing\n"
            "  • Use when: Speed and budget are priorities\n\n"
            "Claude Opus 4.1: Premium model with deep reasoning.\n"
            "  • Pricing: $15/$75 per million tokens (5x more expensive)\n"
            "  • Best for: Legal documents, technical specifications, complex reasoning\n"
            "  • Use when: Highest accuracy is critical, budget is not a concern"
        )
        current_claude_model = settings.get('claude_model', 'claude-sonnet-4-5-20250929')
        for i in range(claude_combo.count()):
            if current_claude_model in claude_combo.itemText(i):
                claude_combo.setCurrentIndex(i)
                break
        claude_combo.setEnabled(claude_radio.isChecked())
        model_layout.addWidget(claude_combo)
        
        model_layout.addSpacing(10)
        
        # Gemini models
        gemini_model_label = QLabel("<b>Gemini Models:</b>")
        model_layout.addWidget(gemini_model_label)
        
        gemini_combo = QComboBox()
        gemini_combo.addItems([
            "gemini-2.5-flash (Recommended - Best Balance)",
            "gemini-2.5-flash-lite (Fastest & Most Economical)",
            "gemini-2.5-pro (Premium - Complex Reasoning)",
            "gemini-2.0-flash-exp (Experimental)"
        ])
        gemini_combo.setToolTip(
            "Gemini 2.5 Flash (Recommended): Best price-performance balance.\n"
            "  • Best for: General translation, high-volume tasks\n"
            "  • Context: 1M tokens input, 65K output\n"
            "  • Strengths: Fast, cost-effective, well-rounded\n"
            "  • Use for: Most translation tasks\n\n"
            "Gemini 2.5 Flash-Lite: Fastest and most economical option.\n"
            "  • Best for: High-throughput batch translation\n"
            "  • Context: 1M tokens input, 65K output\n"
            "  • Strengths: Maximum speed, lowest cost\n"
            "  • Use when: Budget and speed are top priorities\n\n"
            "Gemini 2.5 Pro: State-of-the-art reasoning for complex problems.\n"
            "  • Best for: Complex analytical translation, legal documents\n"
            "  • Context: 1M tokens input, 65K output\n"
            "  • Strengths: Highest accuracy, complex reasoning\n"
            "  • Use when: Maximum quality needed, regardless of cost\n\n"
            "Gemini 2.0 Flash Exp: Experimental model for testing."
        )
        current_gemini_model = settings.get('gemini_model', 'gemini-2.5-flash')
        for i in range(gemini_combo.count()):
            if current_gemini_model in gemini_combo.itemText(i):
                gemini_combo.setCurrentIndex(i)
                break
        gemini_combo.setEnabled(gemini_radio.isChecked())
        model_layout.addWidget(gemini_combo)
        
        # Connect radio buttons to enable/disable combos
        def update_combo_states():
            openai_combo.setEnabled(openai_radio.isChecked())
            claude_combo.setEnabled(claude_radio.isChecked())
            gemini_combo.setEnabled(gemini_radio.isChecked())
        
        openai_radio.toggled.connect(update_combo_states)
        claude_radio.toggled.connect(update_combo_states)
        gemini_radio.toggled.connect(update_combo_states)
        
        model_group.setLayout(model_layout)
        layout.addWidget(model_group)
        
        # Enable/Disable Providers
        provider_enable_group = QGroupBox("Enable/Disable LLM Providers")
        provider_enable_layout = QVBoxLayout()
        
        provider_enable_info = QLabel(
            "Uncheck providers you don't want to use. Only enabled providers will be available for translation."
        )
        provider_enable_info.setWordWrap(True)
        provider_enable_info.setStyleSheet("font-size: 9pt; color: #666; padding: 5px;")
        provider_enable_layout.addWidget(provider_enable_info)
        
        openai_enable_cb = CheckmarkCheckBox("Enable OpenAI")
        openai_enable_cb.setChecked(enabled_providers.get('llm_openai', True))
        provider_enable_layout.addWidget(openai_enable_cb)
        
        claude_enable_cb = CheckmarkCheckBox("Enable Claude")
        claude_enable_cb.setChecked(enabled_providers.get('llm_claude', True))
        provider_enable_layout.addWidget(claude_enable_cb)
        
        gemini_enable_cb = CheckmarkCheckBox("Enable Gemini")
        gemini_enable_cb.setChecked(enabled_providers.get('llm_gemini', True))
        provider_enable_layout.addWidget(gemini_enable_cb)
        
        provider_enable_group.setLayout(provider_enable_layout)
        layout.addWidget(provider_enable_group)
        
        # API Keys info
        api_keys_group = QGroupBox("API Keys")
        api_keys_layout = QVBoxLayout()
        
        api_keys_info = QLabel(
            f"Configure your API keys in:<br>"
            f"<code>{self.user_data_path / 'api_keys.txt'}</code><br><br>"
            f"See example file for format:<br>"
            f"<code>{self.user_data_path / 'api_keys.example.txt'}</code>"
        )
        api_keys_info.setWordWrap(True)
        api_keys_layout.addWidget(api_keys_info)
        
        # Button to open API keys file
        open_keys_btn = QPushButton("📝 Open API Keys File")
        open_keys_btn.clicked.connect(lambda: self.open_api_keys_file())
        api_keys_layout.addWidget(open_keys_btn)
        
        api_keys_group.setLayout(api_keys_layout)
        layout.addWidget(api_keys_group)
        
        # Translation Preferences
        prefs_group = QGroupBox("Translation Preferences")
        prefs_layout = QVBoxLayout()
        prefs_layout.setSpacing(10)
        
        # Load current preferences
        general_prefs = self.load_general_settings()
        batch_size = general_prefs.get('batch_size', 100)
        surrounding_segments = general_prefs.get('surrounding_segments', 5)
        use_full_context = general_prefs.get('use_full_context', True)
        auto_insert_100 = general_prefs.get('auto_insert_100', False)
        check_tm_before_api = general_prefs.get('check_tm_before_api', True)
        auto_propagate_100 = general_prefs.get('auto_propagate_100', True)
        
        # Auto-insert 100% TM matches
        auto_insert_cb = CheckmarkCheckBox("✨ Auto-insert 100% TM matches")
        auto_insert_cb.setChecked(auto_insert_100)
        prefs_layout.addWidget(auto_insert_cb)
        auto_insert_info = QLabel("  ⓘ Automatically fill target with 100% matches when entering untranslated segments")
        auto_insert_info.setStyleSheet("font-size: 9pt; color: #666; padding-left: 20px;")
        prefs_layout.addWidget(auto_insert_info)
        
        prefs_layout.addSpacing(5)
        
        # Batch Size
        batch_size_layout = QHBoxLayout()
        batch_size_label = QLabel("Batch Size (segments per API call):")
        batch_size_layout.addWidget(batch_size_label)
        batch_size_spin = QSpinBox()
        batch_size_spin.setMinimum(1)
        batch_size_spin.setMaximum(500)
        batch_size_spin.setValue(batch_size)
        batch_size_spin.setToolTip("Larger batches = faster but higher API cost per call")
        batch_size_layout.addWidget(batch_size_spin)
        batch_size_layout.addStretch()
        prefs_layout.addLayout(batch_size_layout)
        batch_size_info = QLabel("  ⓘ Larger batches = faster but higher API cost per call. Default: 100")
        batch_size_info.setStyleSheet("font-size: 9pt; color: #666; padding-left: 20px;")
        prefs_layout.addWidget(batch_size_info)
        
        prefs_layout.addSpacing(5)
        
        # Surrounding segments
        surrounding_layout = QHBoxLayout()
        surrounding_label = QLabel("Surrounding segments (single-segment translation):")
        surrounding_layout.addWidget(surrounding_label)
        surrounding_spin = QSpinBox()
        surrounding_spin.setMinimum(0)
        surrounding_spin.setMaximum(20)
        surrounding_spin.setValue(surrounding_segments)
        surrounding_spin.setToolTip("Send nearby segments for context without full document")
        surrounding_layout.addWidget(surrounding_spin)
        surrounding_segments_label = QLabel("segments before/after")
        surrounding_layout.addWidget(surrounding_segments_label)
        surrounding_layout.addStretch()
        prefs_layout.addLayout(surrounding_layout)
        surrounding_info = QLabel("  ⓘ Send nearby segments for context without full document. 0 = no context. Default: 5")
        surrounding_info.setStyleSheet("font-size: 9pt; color: #666; padding-left: 20px;")
        prefs_layout.addWidget(surrounding_info)
        
        prefs_layout.addSpacing(5)
        
        # Include full document context
        full_context_cb = CheckmarkCheckBox("Include full document context in batch translation (increases API usage)")
        full_context_cb.setChecked(use_full_context)
        prefs_layout.addWidget(full_context_cb)
        full_context_info = QLabel("  ⓘ Context helps with consistency but sends more data - use for technical docs")
        full_context_info.setStyleSheet("font-size: 9pt; color: #666; padding-left: 20px;")
        prefs_layout.addWidget(full_context_info)
        
        prefs_layout.addSpacing(5)
        
        # Check TM before API call
        check_tm_cb = CheckmarkCheckBox("Check TM before API call")
        check_tm_cb.setChecked(check_tm_before_api)
        prefs_layout.addWidget(check_tm_cb)
        
        prefs_layout.addSpacing(5)
        
        # Auto-propagate 100% TM matches
        auto_propagate_cb = CheckmarkCheckBox("Auto-propagate 100% TM matches")
        auto_propagate_cb.setChecked(auto_propagate_100)
        prefs_layout.addWidget(auto_propagate_cb)
        
        prefs_layout.addSpacing(5)
        
        # TM/Termbase lookup delay
        lookup_delay = general_prefs.get('lookup_delay', 1500)  # Default 1.5 seconds
        delay_layout = QHBoxLayout()
        delay_label = QLabel("TM/Termbase lookup delay:")
        delay_layout.addWidget(delay_label)
        delay_spin = QSpinBox()
        delay_spin.setMinimum(0)
        delay_spin.setMaximum(10000)
        delay_spin.setSingleStep(100)
        delay_spin.setValue(lookup_delay)
        delay_spin.setSuffix(" ms")
        delay_spin.setToolTip("Delay before searching TM/Termbase when selecting segment (prevents searches while navigating quickly)")
        delay_layout.addWidget(delay_spin)
        delay_layout.addStretch()
        prefs_layout.addLayout(delay_layout)
        delay_info = QLabel("  ⓘ Prevents searches while navigating quickly. 0 = instant, 1500 = 1.5 seconds")
        delay_info.setStyleSheet("font-size: 9pt; color: #666; padding-left: 20px;")
        prefs_layout.addWidget(delay_info)
        
        prefs_group.setLayout(prefs_layout)
        layout.addWidget(prefs_group)
        
        # Save button
        save_btn = QPushButton("💾 Save LLM Settings")
        save_btn.setStyleSheet("font-weight: bold; padding: 8px;")
        save_btn.clicked.connect(lambda: self._save_llm_settings_from_ui(
            openai_radio, claude_radio, gemini_radio, 
            openai_combo, claude_combo, gemini_combo,
            openai_enable_cb, claude_enable_cb, gemini_enable_cb,
            batch_size_spin, surrounding_spin, full_context_cb,
            auto_insert_cb, check_tm_cb, auto_propagate_cb,
            delay_spin
        ))
        layout.addWidget(save_btn)
        
        layout.addStretch()
        
        return tab
    
    def _create_mt_settings_tab(self):
        """Create MT Settings tab content"""
        from PyQt6.QtWidgets import QCheckBox, QGroupBox, QPushButton
        
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)
        
        enabled_providers = self.load_provider_enabled_states()
        
        # Enable/Disable MT Providers
        mt_provider_group = QGroupBox("Machine Translation Providers")
        mt_provider_layout = QVBoxLayout()
        
        mt_info = QLabel(
            "Enable or disable individual MT providers. Supervertaler will use the first available enabled provider in priority order:\n"
            "1. Google Translate, 2. DeepL, 3. Microsoft Translator, 4. Amazon Translate, 5. ModernMT, 6. MyMemory"
        )
        mt_info.setWordWrap(True)
        mt_info.setStyleSheet("font-size: 9pt; color: #666; padding: 5px;")
        mt_provider_layout.addWidget(mt_info)
        
        google_translate_enable_cb = QCheckBox("Enable Google Translate")
        google_translate_enable_cb.setChecked(enabled_providers.get('mt_google_translate', True))
        mt_provider_layout.addWidget(google_translate_enable_cb)
        
        deepl_enable_cb = QCheckBox("Enable DeepL")
        deepl_enable_cb.setChecked(enabled_providers.get('mt_deepl', True))
        mt_provider_layout.addWidget(deepl_enable_cb)
        
        microsoft_enable_cb = QCheckBox("Enable Microsoft Translator")
        microsoft_enable_cb.setChecked(enabled_providers.get('mt_microsoft', True))
        mt_provider_layout.addWidget(microsoft_enable_cb)
        
        amazon_enable_cb = QCheckBox("Enable Amazon Translate")
        amazon_enable_cb.setChecked(enabled_providers.get('mt_amazon', True))
        mt_provider_layout.addWidget(amazon_enable_cb)
        
        modernmt_enable_cb = QCheckBox("Enable ModernMT")
        modernmt_enable_cb.setChecked(enabled_providers.get('mt_modernmt', True))
        mt_provider_layout.addWidget(modernmt_enable_cb)
        
        mymemory_enable_cb = QCheckBox("Enable MyMemory (Free tier)")
        mymemory_enable_cb.setChecked(enabled_providers.get('mt_mymemory', True))
        mt_provider_layout.addWidget(mymemory_enable_cb)
        
        mt_provider_group.setLayout(mt_provider_layout)
        layout.addWidget(mt_provider_group)
        
        # API Keys info for MT
        mt_api_keys_group = QGroupBox("API Keys")
        mt_api_keys_layout = QVBoxLayout()
        
        mt_api_keys_info = QLabel(
            f"Configure your MT API keys in:<br>"
            f"<code>{self.user_data_path / 'api_keys.txt'}</code><br><br>"
            f"See example file for format:<br>"
            f"<code>{self.user_data_path / 'api_keys.example.txt'}</code>"
        )
        mt_api_keys_info.setWordWrap(True)
        mt_api_keys_layout.addWidget(mt_api_keys_info)
        
        # Button to open API keys file
        mt_open_keys_btn = QPushButton("📝 Open API Keys File")
        mt_open_keys_btn.clicked.connect(lambda: self.open_api_keys_file())
        mt_api_keys_layout.addWidget(mt_open_keys_btn)
        
        mt_api_keys_group.setLayout(mt_api_keys_layout)
        layout.addWidget(mt_api_keys_group)
        
        # Save button
        save_btn = QPushButton("💾 Save MT Settings")
        save_btn.setStyleSheet("font-weight: bold; padding: 8px;")
        save_btn.clicked.connect(lambda: self._save_mt_settings_from_ui(
            google_translate_enable_cb, deepl_enable_cb, microsoft_enable_cb,
            amazon_enable_cb, modernmt_enable_cb, mymemory_enable_cb
        ))
        layout.addWidget(save_btn)
        
        layout.addStretch()
        
        return tab
    
    def _create_general_settings_tab(self):
        """Create General Settings tab content"""
        from PyQt6.QtWidgets import QCheckBox, QGroupBox, QPushButton
        
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)
        
        general_settings = self.load_general_settings()
        
        # Startup Settings group
        startup_group = QGroupBox("Startup Settings")
        startup_layout = QVBoxLayout()
        
        restore_last_project_cb = QCheckBox("Restore last project on startup")
        restore_last_project_cb.setChecked(general_settings.get('restore_last_project', False))
        restore_last_project_cb.setToolTip(
            "When enabled, Supervertaler will automatically open the last project you were working on when the application starts."
        )
        startup_layout.addWidget(restore_last_project_cb)
        
        startup_group.setLayout(startup_layout)
        layout.addWidget(startup_group)
        
        # Translation memory and termbase settings section
        tm_termbase_group = QGroupBox("Translation memory and termbase settings")
        tm_termbase_layout = QVBoxLayout()
        
        # Auto-propagate exact TM matches
        auto_propagate_cb = QCheckBox("Auto-propagate exact TM matches (100%)")
        auto_propagate_cb.setChecked(general_settings.get('auto_propagate_exact_matches', True))
        auto_propagate_cb.setToolTip(
            "Automatically fill target with 100% TM matches when a segment is selected and empty.\n"
            "This saves time by applying exact matches without manual confirmation."
        )
        tm_termbase_layout.addWidget(auto_propagate_cb)
        
        # TM/Termbase matching toggle
        tm_matching_cb = QCheckBox("Enable TM && Termbase Matching")
        tm_matching_cb.setChecked(self.enable_tm_matching)  # Load current state
        tm_matching_cb.setToolTip(
            "When enabled, translation memory and termbase searches are performed automatically\n"
            "when you select a segment (after a 1.5 second delay). Disable to improve performance\n"
            "when navigating quickly through segments."
        )
        tm_matching_cb.toggled.connect(self.toggle_tm_termbase_matching)
        tm_termbase_layout.addWidget(tm_matching_cb)
        self.tm_matching_checkbox = tm_matching_cb  # Store reference for updates
        self.auto_propagate_checkbox = auto_propagate_cb  # Store reference for updates
        
        tm_termbase_group.setLayout(tm_termbase_layout)
        layout.addWidget(tm_termbase_group)

        # AI Assistant settings group
        ai_group = QGroupBox("AI Assistant Settings")
        ai_layout = QVBoxLayout()

        auto_markdown_cb = QCheckBox("Auto-generate markdown for imported documents")
        auto_markdown_cb.setChecked(general_settings.get('auto_generate_markdown', False))
        auto_markdown_cb.setToolTip(
            "When enabled, Supervertaler will automatically convert imported documents\n"
            "to markdown format and save them in user_data_private/AI_Assistant/current_document/.\n"
            "This allows the AI Assistant and you to access a markdown version of your documents."
        )
        ai_layout.addWidget(auto_markdown_cb)

        ai_group.setLayout(ai_layout)
        layout.addWidget(ai_group)

        # Find & Replace settings group
        find_replace_group = QGroupBox("Find && Replace Settings")
        find_replace_layout = QVBoxLayout()
        
        allow_replace_cb = QCheckBox("Allow Replace in Source Text")
        allow_replace_cb.setChecked(self.allow_replace_in_source)
        allow_replace_cb.setToolTip(
            "⚠️ WARNING: Enabling this allows replacing text in the source column.\n"
            "This can be dangerous as it modifies your original source text.\n"
            "Use with extreme caution!"
        )
        find_replace_layout.addWidget(allow_replace_cb)
        
        # Add warning label
        warning_label = QLabel(
            "⚠️ <b>Warning:</b> Replacing in source text modifies your original content.\n"
            "This feature is disabled by default for safety."
        )
        warning_label.setWordWrap(True)
        warning_label.setStyleSheet("color: #d97706; padding: 10px; background-color: #fef3c7; border-radius: 3px;")
        find_replace_layout.addWidget(warning_label)
        
        find_replace_group.setLayout(find_replace_layout)
        layout.addWidget(find_replace_group)
        
        # Save button
        save_btn = QPushButton("💾 Save General Settings")
        save_btn.setStyleSheet("font-weight: bold; padding: 8px;")
        save_btn.clicked.connect(lambda: self._save_general_settings_from_ui(
            restore_last_project_cb, allow_replace_cb, auto_propagate_cb, auto_markdown_cb
        ))
        layout.addWidget(save_btn)
        
        layout.addStretch()
        
        return tab
    
    def _create_view_settings_tab(self):
        """Create View/Display Settings tab content"""
        from PyQt6.QtWidgets import QGroupBox, QPushButton
        
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)
        
        font_settings = self.load_general_settings()
        
        # Grid Text Font Size section
        grid_group = QGroupBox("📊 Grid Text Font Size")
        grid_layout = QVBoxLayout()
        
        grid_size_info = QLabel(
            "Set the default font size for the grid (source and target columns).\n"
            "You can also adjust this using View menu → Grid Text Zoom."
        )
        grid_size_info.setStyleSheet("font-size: 8pt; color: #666; padding: 8px; background-color: #f3f4f6; border-radius: 2px;")
        grid_size_info.setWordWrap(True)
        grid_layout.addWidget(grid_size_info)
        
        grid_spin_layout = QHBoxLayout()
        grid_spin_layout.addWidget(QLabel("Font Size:"))
        grid_font_spin = QSpinBox()
        grid_font_spin.setMinimum(7)
        grid_font_spin.setMaximum(72)
        grid_font_spin.setValue(font_settings.get('grid_font_size', 11))
        grid_font_spin.setSuffix(" pt")
        grid_font_spin.setToolTip("Grid font size (7-72 pt)")
        grid_spin_layout.addWidget(grid_font_spin)
        grid_spin_layout.addStretch()
        grid_layout.addLayout(grid_spin_layout)
        
        grid_group.setLayout(grid_layout)
        layout.addWidget(grid_group)
        
        # Translation Results Pane Font Size section
        results_group = QGroupBox("📋 Translation Results Pane Font Size")
        results_layout = QVBoxLayout()
        
        results_size_info = QLabel(
            "Set the default font sizes for the translation results pane.\n"
            "You can also adjust these using View menu → Translation Results Pane."
        )
        results_size_info.setStyleSheet("font-size: 8pt; color: #666; padding: 8px; background-color: #f3f4f6; border-radius: 2px;")
        results_size_info.setWordWrap(True)
        results_layout.addWidget(results_size_info)
        
        # Match list font size
        match_spin_layout = QHBoxLayout()
        match_spin_layout.addWidget(QLabel("Match List Font Size:"))
        match_font_spin = QSpinBox()
        match_font_spin.setMinimum(7)
        match_font_spin.setMaximum(16)
        match_font_spin.setValue(font_settings.get('results_match_font_size', 9))
        match_font_spin.setSuffix(" pt")
        match_font_spin.setToolTip("Match list font size (7-16 pt)")
        match_spin_layout.addWidget(match_font_spin)
        match_spin_layout.addStretch()
        results_layout.addLayout(match_spin_layout)
        
        # Compare boxes font size
        compare_spin_layout = QHBoxLayout()
        compare_spin_layout.addWidget(QLabel("Compare Boxes Font Size:"))
        compare_font_spin = QSpinBox()
        compare_font_spin.setMinimum(7)
        compare_font_spin.setMaximum(14)
        compare_font_spin.setValue(font_settings.get('results_compare_font_size', 9))
        compare_font_spin.setSuffix(" pt")
        compare_font_spin.setToolTip("Compare boxes font size (7-14 pt)")
        compare_spin_layout.addWidget(compare_font_spin)
        compare_spin_layout.addStretch()
        results_layout.addLayout(compare_spin_layout)
        
        results_group.setLayout(results_layout)
        layout.addWidget(results_group)
        
        # Quick Reference section
        reference_group = QGroupBox("⌨️ Font Size Quick Reference")
        reference_layout = QVBoxLayout()
        
        reference_text = QLabel(
            "<b>All Zoom Controls:</b><br>"
            "View → Grid Text Zoom<br>"
            "• Ctrl++ (NumPad +) - Increase<br>"
            "• Ctrl+- (NumPad -) - Decrease<br><br>"
            "View → Translation Results Pane<br>"
            "• Ctrl+Shift++ (Ctrl+Shift+=) - Increase<br>"
            "• Ctrl+Shift+- - Decrease<br>"
            "• Results Zoom Reset - Back to default (9pt)<br>"
        )
        reference_text.setTextFormat(Qt.TextFormat.RichText)
        reference_text.setWordWrap(True)
        reference_layout.addWidget(reference_text)
        
        reference_group.setLayout(reference_layout)
        layout.addWidget(reference_group)
        
        # Save button
        save_btn = QPushButton("💾 Save View Settings")
        save_btn.setStyleSheet("font-weight: bold; padding: 8px;")
        save_btn.clicked.connect(lambda: self._save_view_settings_from_ui(
            grid_font_spin, match_font_spin, compare_font_spin
        ))
        layout.addWidget(save_btn)
        
        layout.addStretch()

        return tab

    def _create_voice_dictation_settings_tab(self):
        """Create Voice Dictation Settings tab content"""
        from PyQt6.QtWidgets import QGroupBox, QPushButton, QComboBox, QSpinBox

        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)

        # Load current dictation settings
        dictation_settings = self.load_dictation_settings()

        # Header info
        header_info = QLabel(
            "Configure voice dictation settings for hands-free translation input.\n"
            "Voice dictation allows you to speak translations instead of typing them."
        )
        header_info.setStyleSheet("font-size: 9pt; color: #444; padding: 10px; background-color: #E3F2FD; border-radius: 4px;")
        header_info.setWordWrap(True)
        layout.addWidget(header_info)

        # Whisper Model Settings group
        model_group = QGroupBox("🤖 Speech Recognition Model")
        model_layout = QVBoxLayout()

        model_info = QLabel(
            "Select the Whisper model size. Larger models are more accurate but slower.\n"
            "• tiny: Fastest, least accurate (~1 GB RAM)\n"
            "• base: Fast, good accuracy (~1 GB RAM) - Recommended\n"
            "• small: Slower, better accuracy (~2 GB RAM)\n"
            "• medium: Slow, very accurate (~5 GB RAM)\n"
            "• large: Slowest, best accuracy (~10 GB RAM)"
        )
        model_info.setStyleSheet("font-size: 8pt; color: #666; padding: 8px; background-color: #f3f4f6; border-radius: 2px;")
        model_info.setWordWrap(True)
        model_layout.addWidget(model_info)

        model_select_layout = QHBoxLayout()
        model_select_layout.addWidget(QLabel("Model:"))
        model_combo = QComboBox()
        model_combo.addItems(["tiny", "base", "small", "medium", "large"])
        model_combo.setCurrentText(dictation_settings.get('model', 'base'))
        model_combo.setToolTip("Select Whisper model size (base recommended)")
        model_select_layout.addWidget(model_combo)
        model_select_layout.addStretch()
        model_layout.addLayout(model_select_layout)

        model_group.setLayout(model_layout)
        layout.addWidget(model_group)

        # Recording Settings group
        recording_group = QGroupBox("🎤 Recording Settings")
        recording_layout = QVBoxLayout()

        # Max recording duration
        duration_info = QLabel(
            "Maximum recording duration per dictation session.\n"
            "Recording automatically stops after this time limit."
        )
        duration_info.setStyleSheet("font-size: 8pt; color: #666; padding: 8px; background-color: #f3f4f6; border-radius: 2px;")
        duration_info.setWordWrap(True)
        recording_layout.addWidget(duration_info)

        duration_layout = QHBoxLayout()
        duration_layout.addWidget(QLabel("Max Duration:"))
        duration_spin = QSpinBox()
        duration_spin.setMinimum(3)
        duration_spin.setMaximum(60)
        duration_spin.setValue(dictation_settings.get('max_duration', 10))
        duration_spin.setSuffix(" seconds")
        duration_spin.setToolTip("Maximum recording time (3-60 seconds)")
        duration_layout.addWidget(duration_spin)
        duration_layout.addStretch()
        recording_layout.addLayout(duration_layout)

        recording_group.setLayout(recording_layout)
        layout.addWidget(recording_group)

        # Language Settings group
        language_group = QGroupBox("🌐 Language Settings")
        language_layout = QVBoxLayout()

        language_info = QLabel(
            "By default, voice dictation uses your project's target language.\n"
            "You can override this to always use a specific language."
        )
        language_info.setStyleSheet("font-size: 8pt; color: #666; padding: 8px; background-color: #f3f4f6; border-radius: 2px;")
        language_info.setWordWrap(True)
        language_layout.addWidget(language_info)

        lang_select_layout = QHBoxLayout()
        lang_select_layout.addWidget(QLabel("Dictation Language:"))
        lang_combo = QComboBox()
        lang_combo.addItems([
            "Auto (use project target language)",
            "English", "Dutch", "German", "French", "Spanish",
            "Italian", "Portuguese", "Polish", "Russian",
            "Chinese", "Japanese", "Korean"
        ])
        lang_combo.setCurrentText(dictation_settings.get('language', 'Auto (use project target language)'))
        lang_combo.setToolTip("Language for speech recognition")
        lang_select_layout.addWidget(lang_combo)
        lang_select_layout.addStretch()
        language_layout.addLayout(lang_select_layout)

        language_group.setLayout(language_layout)
        layout.addWidget(language_group)

        # Save button
        save_btn = QPushButton("💾 Save Voice Dictation Settings")
        save_btn.setStyleSheet("background-color: #4CAF50; color: white; font-weight: bold; padding: 10px;")
        save_btn.clicked.connect(lambda: self.save_dictation_settings(
            model_combo.currentText(),
            duration_spin.value(),
            lang_combo.currentText()
        ))
        layout.addWidget(save_btn)

        # Store references for later access
        self.dictation_model_combo = model_combo
        self.dictation_duration_spin = duration_spin
        self.dictation_lang_combo = lang_combo

        layout.addStretch()

        return tab

    def _create_system_prompts_tab(self):
        """Create System Prompts (Layer 1) Settings tab content"""
        from PyQt6.QtWidgets import QGroupBox, QPushButton, QTextEdit, QComboBox

        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)

        # Header info
        header_group = QGroupBox("📝 System Prompts (Layer 1)")
        header_layout = QVBoxLayout()

        info_label = QLabel(
            "<b>System Prompts are the foundation of Supervertaler's 2-Layer Prompt Architecture.</b><br><br>"
            "These prompts are <b>always applied</b> and contain critical infrastructure instructions:<br>"
            "• CAT tool tag preservation (memoQ, Trados, CafeTran)<br>"
            "• Formatting rules and output standards<br>"
            "• Language-specific conventions (numbers, dates, typography)<br>"
            "• Professional translation context<br><br>"
            "<i>Layer 2 (Custom Prompts) are managed in the Prompts tab and can be freely attached/detached.</i>"
        )
        info_label.setTextFormat(Qt.TextFormat.RichText)
        info_label.setWordWrap(True)
        info_label.setStyleSheet("font-size: 9pt; color: #333; padding: 12px; background-color: #f3f4f6; border-radius: 3px;")
        header_layout.addWidget(info_label)

        header_group.setLayout(header_layout)
        layout.addWidget(header_group)

        # Mode selector
        mode_group = QGroupBox("🎯 Select System Prompt Mode")
        mode_layout = QVBoxLayout()

        mode_info = QLabel(
            "Supervertaler uses different system prompts for different translation modes:"
        )
        mode_info.setStyleSheet("font-size: 9pt; color: #666; padding: 8px;")
        mode_info.setWordWrap(True)
        mode_layout.addWidget(mode_info)

        mode_selector_layout = QHBoxLayout()
        mode_selector_layout.addWidget(QLabel("Mode:"))
        mode_combo = QComboBox()
        mode_combo.addItems([
            "Single Segment Translation",
            "Batch DOCX Translation",
            "Batch Bilingual Translation"
        ])
        mode_combo.setToolTip("Select which system prompt to view/edit")
        mode_selector_layout.addWidget(mode_combo)
        mode_selector_layout.addStretch()
        mode_layout.addLayout(mode_selector_layout)

        mode_group.setLayout(mode_layout)
        layout.addWidget(mode_group)

        # Editor
        editor_group = QGroupBox("✏️ Edit System Prompt")
        editor_layout = QVBoxLayout()

        editor_info = QLabel(
            "Edit the system prompt below. Use {{SOURCE_LANGUAGE}}, {{TARGET_LANGUAGE}}, and {{SOURCE_TEXT}} as placeholders."
        )
        editor_info.setStyleSheet("font-size: 8pt; color: #666; padding: 8px; background-color: #fff3cd; border-radius: 2px;")
        editor_info.setWordWrap(True)
        editor_layout.addWidget(editor_info)

        system_prompt_editor = QTextEdit()
        system_prompt_editor.setStyleSheet("font-family: 'Consolas', 'Courier New', monospace; font-size: 9pt;")
        system_prompt_editor.setMinimumHeight(400)
        editor_layout.addWidget(system_prompt_editor)

        editor_group.setLayout(editor_layout)
        layout.addWidget(editor_group)

        # Buttons
        buttons_layout = QHBoxLayout()

        reset_btn = QPushButton("🔄 Reset to Default")
        reset_btn.setToolTip("Restore the default system prompt for this mode")
        reset_btn.clicked.connect(lambda: self._reset_system_prompt(mode_combo, system_prompt_editor))
        buttons_layout.addWidget(reset_btn)

        buttons_layout.addStretch()

        save_btn = QPushButton("💾 Save System Prompt")
        save_btn.setStyleSheet("font-weight: bold; padding: 8px;")
        save_btn.clicked.connect(lambda: self._save_system_prompt_from_ui(mode_combo, system_prompt_editor))
        buttons_layout.addWidget(save_btn)

        layout.addLayout(buttons_layout)

        # Load initial prompt
        self._load_system_prompt_into_editor(mode_combo, system_prompt_editor)

        # Connect mode change to load new prompt
        mode_combo.currentIndexChanged.connect(lambda: self._load_system_prompt_into_editor(mode_combo, system_prompt_editor))

        return tab

    def _load_system_prompt_into_editor(self, mode_combo, editor):
        """Load the selected system prompt into the editor"""
        mode_map = {
            "Single Segment Translation": "single",
            "Batch DOCX Translation": "batch_docx",
            "Batch Bilingual Translation": "batch_bilingual"
        }

        selected_mode = mode_combo.currentText()
        mode_key = mode_map.get(selected_mode, "single")

        # Load from unified_prompt_manager if available
        if hasattr(self, 'unified_prompt_manager'):
            prompt_text = self.unified_prompt_manager.get_system_template(mode_key)
        else:
            # Fallback: load from JSON file
            system_prompts_file = self.user_data_path / "Prompt_Library" / "system_prompts_layer1.json"
            if system_prompts_file.exists():
                import json
                with open(system_prompts_file, 'r', encoding='utf-8') as f:
                    system_prompts = json.load(f)
                prompt_text = system_prompts.get(mode_key, "# SYSTEM PROMPT\n\nNo prompt defined.")
            else:
                prompt_text = "# SYSTEM PROMPT\n\nNo prompt defined."

        editor.setPlainText(prompt_text)

    def _save_system_prompt_from_ui(self, mode_combo, editor):
        """Save the edited system prompt"""
        mode_map = {
            "Single Segment Translation": "single",
            "Batch DOCX Translation": "batch_docx",
            "Batch Bilingual Translation": "batch_bilingual"
        }

        selected_mode = mode_combo.currentText()
        mode_key = mode_map.get(selected_mode, "single")
        prompt_text = editor.toPlainText()

        # Save to unified_prompt_manager if available
        if hasattr(self, 'unified_prompt_manager'):
            self.unified_prompt_manager.system_templates[mode_key] = prompt_text

        # Always save to JSON file
        import json
        system_prompts_file = self.user_data_path / "Prompt_Library" / "system_prompts_layer1.json"
        system_prompts_file.parent.mkdir(parents=True, exist_ok=True)

        # Load existing prompts
        if system_prompts_file.exists():
            with open(system_prompts_file, 'r', encoding='utf-8') as f:
                system_prompts = json.load(f)
        else:
            system_prompts = {}

        # Update and save
        system_prompts[mode_key] = prompt_text
        with open(system_prompts_file, 'w', encoding='utf-8') as f:
            json.dump(system_prompts, f, indent=2, ensure_ascii=False)

        self.log(f"✓ Saved system prompt: {selected_mode}")
        QMessageBox.information(
            self,
            "System Prompt Saved",
            f"System prompt for '{selected_mode}' has been saved successfully."
        )

    def _reset_system_prompt(self, mode_combo, editor):
        """Reset system prompt to default"""
        mode_map = {
            "Single Segment Translation": "single",
            "Batch DOCX Translation": "batch_docx",
            "Batch Bilingual Translation": "batch_bilingual"
        }

        selected_mode = mode_combo.currentText()
        mode_key = mode_map.get(selected_mode, "single")

        # Confirm reset
        reply = QMessageBox.question(
            self,
            "Reset System Prompt",
            f"Are you sure you want to reset the system prompt for '{selected_mode}' to its default value?\n\n"
            "This will discard any custom changes you've made.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            # Get default from unified_prompt_manager
            if hasattr(self, 'unified_prompt_manager'):
                default_prompt = self.unified_prompt_manager._get_default_system_template(mode_key)
                editor.setPlainText(default_prompt)
                self.log(f"✓ Reset system prompt to default: {selected_mode}")

    def _save_llm_settings_from_ui(self, openai_radio, claude_radio, gemini_radio,
                                   openai_combo, claude_combo, gemini_combo,
                                   openai_enable_cb, claude_enable_cb, gemini_enable_cb,
                                   batch_size_spin, surrounding_spin, full_context_cb,
                                   auto_insert_cb, check_tm_cb, auto_propagate_cb,
                                   delay_spin):
        """Save LLM settings from UI"""
        new_settings = {
            'provider': 'openai' if openai_radio.isChecked() else 
                       'claude' if claude_radio.isChecked() else 'gemini',
            'openai_model': openai_combo.currentText().split()[0],
            'claude_model': claude_combo.currentText().split()[0],
            'gemini_model': gemini_combo.currentText().split()[0]
        }
        self.save_llm_settings(new_settings)

        # Update current provider and model attributes for AI Assistant
        self.current_provider = new_settings['provider']
        provider_key = f"{self.current_provider}_model"
        self.current_model = new_settings.get(provider_key)

        enabled_states = {
            'llm_openai': openai_enable_cb.isChecked(),
            'llm_claude': claude_enable_cb.isChecked(),
            'llm_gemini': gemini_enable_cb.isChecked()
        }
        # Merge with existing MT settings
        existing = self.load_provider_enabled_states()
        enabled_states.update({k: v for k, v in existing.items() if k.startswith('mt_')})
        self.save_provider_enabled_states(enabled_states)
        
        # Save translation preferences
        general_prefs = self.load_general_settings()
        general_prefs['batch_size'] = batch_size_spin.value()
        general_prefs['surrounding_segments'] = surrounding_spin.value()
        general_prefs['use_full_context'] = full_context_cb.isChecked()
        general_prefs['auto_insert_100'] = auto_insert_cb.isChecked()
        general_prefs['check_tm_before_api'] = check_tm_cb.isChecked()
        general_prefs['auto_propagate_100'] = auto_propagate_cb.isChecked()
        general_prefs['lookup_delay'] = delay_spin.value()
        self.save_general_settings(general_prefs)
        
        self.log(f"✓ LLM settings saved: Provider={new_settings['provider']}, Batch Size={batch_size_spin.value()}")
        QMessageBox.information(self, "Settings Saved", "LLM settings have been saved successfully.")
    
    def _save_mt_settings_from_ui(self, google_cb, deepl_cb, microsoft_cb, amazon_cb, modernmt_cb, mymemory_cb):
        """Save MT settings from UI"""
        enabled_states = {
            'mt_google_translate': google_cb.isChecked(),
            'mt_deepl': deepl_cb.isChecked(),
            'mt_microsoft': microsoft_cb.isChecked(),
            'mt_amazon': amazon_cb.isChecked(),
            'mt_modernmt': modernmt_cb.isChecked(),
            'mt_mymemory': mymemory_cb.isChecked()
        }
        # Merge with existing LLM settings
        existing = self.load_provider_enabled_states()
        enabled_states.update({k: v for k, v in existing.items() if k.startswith('llm_')})
        self.save_provider_enabled_states(enabled_states)
        
        self.log("✓ MT settings saved")
        QMessageBox.information(self, "Settings Saved", "MT settings have been saved successfully.")
    
    def _save_general_settings_from_ui(self, restore_cb, allow_replace_cb, auto_propagate_cb, auto_markdown_cb=None):
        """Save general settings from UI"""
        self.allow_replace_in_source = allow_replace_cb.isChecked()
        self.update_warning_banner()

        # Update auto-propagate state
        self.auto_propagate_exact_matches = auto_propagate_cb.isChecked()

        # Update auto-markdown setting
        if auto_markdown_cb is not None:
            self.auto_generate_markdown = auto_markdown_cb.isChecked()

        general_settings = {
            'restore_last_project': restore_cb.isChecked(),
            'auto_propagate_exact_matches': self.auto_propagate_exact_matches,
            'auto_generate_markdown': self.auto_generate_markdown if hasattr(self, 'auto_generate_markdown') else False,
            'grid_font_size': self.default_font_size,  # Keep existing or update separately
            'results_match_font_size': 9,  # Keep existing
            'results_compare_font_size': 9  # Keep existing
        }
        self.save_general_settings(general_settings)
        
        self.log("✓ General settings saved")
        QMessageBox.information(self, "Settings Saved", "General settings have been saved successfully.")
    
    def _save_view_settings_from_ui(self, grid_spin, match_spin, compare_spin):
        """Save view settings from UI"""
        general_settings = {
            'restore_last_project': self.load_general_settings().get('restore_last_project', False),
            'auto_propagate_exact_matches': self.auto_propagate_exact_matches,  # Keep existing value
            'grid_font_size': grid_spin.value(),
            'results_match_font_size': match_spin.value(),
            'results_compare_font_size': compare_spin.value(),
            'enable_tm_termbase_matching': self.enable_tm_matching  # Save TM/termbase matching state
        }
        self.save_general_settings(general_settings)
        
        # Apply font sizes immediately
        if self.default_font_size != grid_spin.value():
            self.default_font_size = grid_spin.value()
            if hasattr(self, 'table') and self.table is not None:
                self.apply_font_to_grid()
                self.auto_resize_rows()
        
        # Apply results pane font sizes to all panels
        if hasattr(self, 'results_panels'):
            from modules.translation_results_panel import CompactMatchItem
            if CompactMatchItem.font_size_pt != match_spin.value():
                CompactMatchItem.set_font_size(match_spin.value())
                for panel in self.results_panels:
                    if hasattr(panel, 'set_font_size'):
                        panel.set_font_size(match_spin.value())
        
            from modules.translation_results_panel import TranslationResultsPanel
            if TranslationResultsPanel.compare_box_font_size != compare_spin.value():
                TranslationResultsPanel.compare_box_font_size = compare_spin.value()
                for panel in self.results_panels:
                    if hasattr(panel, 'set_compare_box_font_size'):
                        panel.set_compare_box_font_size(compare_spin.value())
        
        self.log("✓ View settings saved and applied")
        QMessageBox.information(self, "Settings Saved", "View settings have been saved and applied successfully.")
    
    def create_editor_tab(self):
        """Create the project editor tab with view switching (Grid/List/Document)"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # View switcher toolbar
        view_toolbar = QWidget()
        view_toolbar_layout = QHBoxLayout(view_toolbar)
        view_toolbar_layout.setContentsMargins(5, 5, 5, 5)
        view_toolbar_layout.setSpacing(5)
        
        view_label = QLabel("View:")
        view_label.setStyleSheet("font-weight: bold;")
        view_toolbar_layout.addWidget(view_label)
        
        self.grid_view_btn = QPushButton("📊 Grid")
        self.grid_view_btn.setCheckable(True)
        self.grid_view_btn.setChecked(True)
        self.grid_view_btn.clicked.connect(lambda: self.switch_view_mode(LayoutMode.GRID))
        view_toolbar_layout.addWidget(self.grid_view_btn)
        
        self.list_view_btn = QPushButton("📋 List")
        self.list_view_btn.setCheckable(True)
        self.list_view_btn.clicked.connect(lambda: self.switch_view_mode(LayoutMode.LIST))
        view_toolbar_layout.addWidget(self.list_view_btn)
        
        self.document_view_btn = QPushButton("📄 Document")
        self.document_view_btn.setCheckable(True)
        self.document_view_btn.clicked.connect(lambda: self.switch_view_mode(LayoutMode.DOCUMENT))
        view_toolbar_layout.addWidget(self.document_view_btn)
        
        view_toolbar_layout.addStretch()
        
        layout.addWidget(view_toolbar)
        
        # Create assistance panel (shared across all views)
        self.create_assistance_panel()
        
        # Stacked widget to hold different views
        self.view_stack = QStackedWidget()
        
        # Create all three views
        self.grid_view_widget = self.create_grid_view_widget()
        self.list_view_widget = self.create_list_view_widget()
        self.document_view_widget = self.create_document_view_widget()
        
        # Add views to stack
        self.view_stack.addWidget(self.grid_view_widget)
        self.view_stack.addWidget(self.list_view_widget)
        self.view_stack.addWidget(self.document_view_widget)
        
        layout.addWidget(self.view_stack)
        
        return tab
    
    def switch_view_mode(self, mode: str):
        """Switch between Grid/List/Document views"""
        self.current_view_mode = mode
        
        # Update button states
        self.grid_view_btn.setChecked(mode == LayoutMode.GRID)
        self.list_view_btn.setChecked(mode == LayoutMode.LIST)
        self.document_view_btn.setChecked(mode == LayoutMode.DOCUMENT)
        
        # Switch stack widget
        if mode == LayoutMode.GRID:
            self.view_stack.setCurrentIndex(0)
        
        elif mode == LayoutMode.LIST:
            self.view_stack.setCurrentIndex(1)
            # Refresh list view when switching to it
            if hasattr(self, 'list_tree') and self.current_project:
                self.refresh_list_view()
        
        elif mode == LayoutMode.DOCUMENT:
            self.view_stack.setCurrentIndex(2)
            # Refresh document view when switching to it
            if hasattr(self, 'document_container') and self.current_project:
                self.refresh_document_view()
        
        self.log(f"Switched to {mode} view")
    
    def create_grid_view_widget(self):
        """Create the Grid View widget (existing grid functionality)"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # Splitter for grid and assistance panel
        self.editor_splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # Left side: Grid container with filter boxes
        grid_container = QWidget()
        grid_layout = QVBoxLayout(grid_container)
        grid_layout.setContentsMargins(0, 0, 0, 0)
        grid_layout.setSpacing(5)
        
        # Warning banner for replace in source (hidden by default)
        self.warning_banner = QWidget()
        self.warning_banner.setStyleSheet(
            "background-color: #dc2626; color: white; padding: 8px; border-radius: 4px;"
        )
        warning_layout = QHBoxLayout(self.warning_banner)
        warning_layout.setContentsMargins(10, 5, 10, 5)
        
        warning_icon = QLabel("⚠️")
        warning_icon.setStyleSheet("font-size: 16px; background: transparent;")
        warning_layout.addWidget(warning_icon)
        
        warning_text = QLabel(
            "<b>WARNING:</b> Replace in Source Text is ENABLED. "
            "This allows modifying your original source segments. Use with extreme caution!"
        )
        warning_text.setStyleSheet("background: transparent; font-weight: bold;")
        warning_text.setWordWrap(True)
        warning_layout.addWidget(warning_text, stretch=1)
        
        settings_link = QPushButton("⚙️ Disable in Options")
        settings_link.setStyleSheet(
            "background-color: rgba(255, 255, 255, 0.2); color: white; "
            "border: 1px solid white; padding: 4px 12px; border-radius: 3px; font-weight: bold;"
        )
        settings_link.setCursor(Qt.CursorShape.PointingHandCursor)
        settings_link.clicked.connect(self._go_to_settings_tab)
        warning_layout.addWidget(settings_link)
        
        grid_layout.addWidget(self.warning_banner)
        self.warning_banner.hide()  # Hidden by default
        
        # Filter panel (like memoQ)
        filter_panel = QWidget()
        filter_layout = QHBoxLayout(filter_panel)
        filter_layout.setContentsMargins(0, 0, 0, 0)
        filter_layout.setSpacing(10)
        
        # Source filter
        source_filter_label = QLabel("Filter Source:")
        self.source_filter = QLineEdit()
        self.source_filter.setPlaceholderText("Type to filter source segments...")
        self.source_filter.textChanged.connect(self.apply_filters)
        self.source_filter.returnPressed.connect(self.apply_filters)
        
        # Target filter
        target_filter_label = QLabel("Filter Target:")
        self.target_filter = QLineEdit()
        self.target_filter.setPlaceholderText("Type to filter target segments...")
        self.target_filter.textChanged.connect(self.apply_filters)
        self.target_filter.returnPressed.connect(self.apply_filters)
        
        # Clear filters button
        clear_filters_btn = QPushButton("Clear Filters")
        clear_filters_btn.clicked.connect(self.clear_filters)
        clear_filters_btn.setMaximumWidth(100)
        
        filter_layout.addWidget(source_filter_label)
        filter_layout.addWidget(self.source_filter, stretch=1)
        filter_layout.addWidget(target_filter_label)
        filter_layout.addWidget(self.target_filter, stretch=1)
        filter_layout.addWidget(clear_filters_btn)
        
        grid_layout.addWidget(filter_panel)
        
        # Translation Grid
        self.create_translation_grid()
        grid_layout.addWidget(self.table)
        
        # Add grid container to splitter
        self.editor_splitter.addWidget(grid_container)
        
        # Set splitter proportions - Qt displays RIGHT to LEFT
        # Position 0 (grid/RIGHT) gets 1000, Position 1 (assistance/LEFT) gets 400
        self.editor_splitter.setSizes([1000, 400])
        
        layout.addWidget(self.editor_splitter)
        
        return widget
    
    def create_grid_view_widget_for_home(self):
        """Create Grid View widget with translation results panel at bottom"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # Main vertical splitter: Grid area (top) and Translation Results (bottom)
        self.home_grid_splitter = QSplitter(Qt.Orientation.Vertical)
        
        # Top: Grid container with filter boxes
        grid_container = QWidget()
        grid_layout = QVBoxLayout(grid_container)
        grid_layout.setContentsMargins(0, 0, 0, 0)
        grid_layout.setSpacing(5)
        
        # Warning banner for replace in source (hidden by default)
        if not hasattr(self, 'warning_banner'):
            self.warning_banner = QWidget()
            self.warning_banner.setStyleSheet(
                "background-color: #dc2626; color: white; padding: 8px; border-radius: 4px;"
            )
            warning_layout = QHBoxLayout(self.warning_banner)
            warning_layout.setContentsMargins(10, 5, 10, 5)
            
            warning_icon = QLabel("⚠️")
            warning_icon.setStyleSheet("font-size: 16px; background: transparent;")
            warning_layout.addWidget(warning_icon)
            
            warning_text = QLabel(
                "<b>WARNING:</b> Replace in Source Text is ENABLED. "
                "This allows modifying your original source segments. Use with extreme caution!"
            )
            warning_text.setStyleSheet("background: transparent; font-weight: bold;")
            warning_text.setWordWrap(True)
            warning_layout.addWidget(warning_text, stretch=1)
            
            settings_link = QPushButton("⚙️ Disable in Options")
            settings_link.setStyleSheet(
                "background-color: rgba(255, 255, 255, 0.2); color: white; "
                "border: 1px solid white; padding: 4px 12px; border-radius: 3px; font-weight: bold;"
            )
            settings_link.setCursor(Qt.CursorShape.PointingHandCursor)
            settings_link.clicked.connect(self._go_to_settings_tab)
            warning_layout.addWidget(settings_link)
        
        grid_layout.addWidget(self.warning_banner)
        self.warning_banner.hide()  # Hidden by default
        
        # Filter panel (like memoQ)
        filter_panel = QWidget()
        filter_layout = QHBoxLayout(filter_panel)
        filter_layout.setContentsMargins(0, 0, 0, 0)
        filter_layout.setSpacing(10)
        
        # Source filter
        source_filter_label = QLabel("Filter Source:")
        if not hasattr(self, 'source_filter'):
            self.source_filter = QLineEdit()
        self.source_filter.setPlaceholderText("Type to filter source segments... (Press Enter or click Filter)")
        self.source_filter.returnPressed.connect(self.apply_filters)
        
        # Target filter
        target_filter_label = QLabel("Filter Target:")
        if not hasattr(self, 'target_filter'):
            self.target_filter = QLineEdit()
        self.target_filter.setPlaceholderText("Type to filter target segments... (Press Enter or click Filter)")
        self.target_filter.returnPressed.connect(self.apply_filters)
        
        # Filter button (activates the filter)
        apply_filter_btn = QPushButton("Filter")
        apply_filter_btn.clicked.connect(self.apply_filters)
        apply_filter_btn.setMaximumWidth(80)
        apply_filter_btn.setStyleSheet("background-color: #4CAF50; color: white; font-weight: bold;")
        
        # Clear filters button
        clear_filters_btn = QPushButton("Clear Filters")
        clear_filters_btn.clicked.connect(self.clear_filters)
        clear_filters_btn.setMaximumWidth(100)
        
        filter_layout.addWidget(source_filter_label)
        filter_layout.addWidget(self.source_filter, stretch=1)
        filter_layout.addWidget(target_filter_label)
        filter_layout.addWidget(self.target_filter, stretch=1)
        filter_layout.addWidget(apply_filter_btn)
        filter_layout.addWidget(clear_filters_btn)
        
        grid_layout.addWidget(filter_panel)
        
        # Pagination controls (like Tkinter version)
        pagination_panel = QWidget()
        pagination_panel.setStyleSheet("background-color: #f8f8f8; padding: 5px;")
        pagination_layout = QHBoxLayout(pagination_panel)
        pagination_layout.setContentsMargins(10, 5, 10, 5)
        pagination_layout.setSpacing(10)
        
        # Pagination label (left side)
        if not hasattr(self, 'pagination_label'):
            self.pagination_label = QLabel("Segments 1-50 of 0")
        self.pagination_label.setStyleSheet("color: #555;")
        pagination_layout.addWidget(self.pagination_label)
        
        pagination_layout.addStretch()
        
        # Pagination controls (right side)
        # First page button
        if not hasattr(self, 'first_page_btn'):
            self.first_page_btn = QPushButton("⏮ First")
            self.first_page_btn.setMaximumWidth(70)
            if hasattr(self, 'go_to_first_page'):
                self.first_page_btn.clicked.connect(self.go_to_first_page)
        pagination_layout.addWidget(self.first_page_btn)
        
        # Previous page button
        if not hasattr(self, 'prev_page_btn'):
            self.prev_page_btn = QPushButton("◀ Prev")
            self.prev_page_btn.setMaximumWidth(70)
            if hasattr(self, 'go_to_prev_page'):
                self.prev_page_btn.clicked.connect(self.go_to_prev_page)
        pagination_layout.addWidget(self.prev_page_btn)
        
        # Page number input
        page_label = QLabel("Page:")
        pagination_layout.addWidget(page_label)
        
        if not hasattr(self, 'page_number_input'):
            self.page_number_input = QLineEdit()
            self.page_number_input.setMaximumWidth(50)
            self.page_number_input.setText("1")
            if hasattr(self, 'go_to_page'):
                self.page_number_input.returnPressed.connect(self.go_to_page)
        pagination_layout.addWidget(self.page_number_input)
        
        if not hasattr(self, 'total_pages_label'):
            self.total_pages_label = QLabel("of 1")
        pagination_layout.addWidget(self.total_pages_label)
        
        # Next page button
        if not hasattr(self, 'next_page_btn'):
            self.next_page_btn = QPushButton("Next ▶")
            self.next_page_btn.setMaximumWidth(70)
            if hasattr(self, 'go_to_next_page'):
                self.next_page_btn.clicked.connect(self.go_to_next_page)
        pagination_layout.addWidget(self.next_page_btn)
        
        # Last page button
        if not hasattr(self, 'last_page_btn'):
            self.last_page_btn = QPushButton("Last ⏭")
            self.last_page_btn.setMaximumWidth(70)
            if hasattr(self, 'go_to_last_page'):
                self.last_page_btn.clicked.connect(self.go_to_last_page)
        pagination_layout.addWidget(self.last_page_btn)
        
        # Page size selector
        page_size_label = QLabel("Per page:")
        pagination_layout.addWidget(page_size_label)
        
        if not hasattr(self, 'page_size_combo'):
            self.page_size_combo = QComboBox()
            self.page_size_combo.addItems(["25", "50", "100", "200", "All"])
            if not hasattr(self, 'grid_page_size'):
                self.grid_page_size = 50
            self.page_size_combo.setCurrentText(str(self.grid_page_size) if self.grid_page_size != 999999 else "All")
            if hasattr(self, 'on_page_size_changed'):
                self.page_size_combo.currentTextChanged.connect(self.on_page_size_changed)
        self.page_size_combo.setMaximumWidth(80)
        pagination_layout.addWidget(self.page_size_combo)
        
        grid_layout.addWidget(pagination_panel)
        
        # Note: Pagination methods (go_to_first_page, go_to_prev_page, etc.) will be implemented
        # when pagination functionality is fully added. For now, buttons are created but won't work.
        
        # Translation Grid (create it if it doesn't exist)
        if not hasattr(self, 'table') or self.table is None:
            self.create_translation_grid()
        grid_layout.addWidget(self.table)
        
        # Add grid container to top of vertical splitter
        self.home_grid_splitter.addWidget(grid_container)  # Grid on top
        
        # Add tabbed assistance panel at bottom (Translation Results, Segment Editor, Notes)
        # Always create fresh to avoid stale widget issues
        self.home_tabbed_panel = self.create_tabbed_assistance_panel()
        self.home_grid_splitter.addWidget(self.home_tabbed_panel)  # Tabs at bottom
        
        print(f"DEBUG: Grid view - Added tabbed panel to splitter")
        print(f"DEBUG: Tabbed panel has {self.home_tabbed_panel.count()} tabs")
        for i in range(self.home_tabbed_panel.count()):
            print(f"DEBUG:   Tab {i}: '{self.home_tabbed_panel.tabText(i)}'")
        
        # Set splitter proportions - Top to Bottom
        # Position 0 (grid) gets 700, Position 1 (tabbed panel) gets 300
        self.home_grid_splitter.setSizes([700, 300])
        
        print(f"DEBUG: Final splitter has {self.home_grid_splitter.count()} widgets:")
        for i in range(self.home_grid_splitter.count()):
            child_widget = self.home_grid_splitter.widget(i)
            print(f"  Position {i}: {child_widget.__class__.__name__}")
        
        # Ensure splitter handles are visible
        self.home_grid_splitter.setHandleWidth(8)
        self.home_grid_splitter.setChildrenCollapsible(False)
        
        layout.addWidget(self.home_grid_splitter)
        
        return widget
    
    def switch_home_view_mode(self, mode: str):
        """Switch between Grid/List/Document views in Home tab"""
        if not hasattr(self, 'home_view_stack'):
            return
        
        # Update button states
        if hasattr(self, 'home_grid_view_btn'):
            self.home_grid_view_btn.setChecked(mode == "grid")
        if hasattr(self, 'home_list_view_btn'):
            self.home_list_view_btn.setChecked(mode == "list")
        if hasattr(self, 'home_document_view_btn'):
            self.home_document_view_btn.setChecked(mode == "document")
        
        # Switch stack
        if mode == "grid":
            self.home_view_stack.setCurrentIndex(0)
        elif mode == "list":
            self.home_view_stack.setCurrentIndex(1)
            # Refresh list view when switching to it
            if hasattr(self, 'list_tree') and self.current_project:
                self.refresh_list_view()
        elif mode == "document":
            self.home_view_stack.setCurrentIndex(2)
            # CRITICAL: Refresh document view to render styles!
            if hasattr(self, 'document_container') and self.current_project:
                print(f"DEBUG: Refreshing document view from home tab switch...")
                self.refresh_document_view()
        
        self.log(f"Switched to {mode} view in Home tab")
    
    # Pagination methods (stubs for now - will be fully implemented later)
    def go_to_first_page(self):
        """Navigate to first page - stub implementation"""
        if not hasattr(self, 'grid_current_page'):
            self.grid_current_page = 0
        else:
            self.grid_current_page = 0
        # TODO: Implement full pagination logic
        pass
    
    def go_to_prev_page(self):
        """Navigate to previous page - stub implementation"""
        if not hasattr(self, 'grid_current_page'):
            self.grid_current_page = 0
        if self.grid_current_page > 0:
            self.grid_current_page -= 1
        # TODO: Implement full pagination logic
        pass
    
    def go_to_next_page(self):
        """Navigate to next page - stub implementation"""
        if not hasattr(self, 'grid_current_page'):
            self.grid_current_page = 0
        self.grid_current_page += 1
        # TODO: Implement full pagination logic
        pass
    
    def go_to_last_page(self):
        """Navigate to last page - stub implementation"""
        if not hasattr(self, 'grid_current_page'):
            self.grid_current_page = 0
        # TODO: Calculate total pages and set to last
        pass
    
    def go_to_page(self):
        """Navigate to specific page - stub implementation"""
        if not hasattr(self, 'page_number_input'):
            return
        try:
            page_num = int(self.page_number_input.text())
            if not hasattr(self, 'grid_current_page'):
                self.grid_current_page = 0
            self.grid_current_page = max(0, page_num - 1)  # Convert to 0-indexed
        except ValueError:
            pass
        # TODO: Implement full pagination logic
    
    def on_page_size_changed(self, text: str):
        """Handle page size change - stub implementation"""
        if not hasattr(self, 'grid_page_size'):
            self.grid_page_size = 50
        if text == "All":
            self.grid_page_size = 999999
        else:
            try:
                self.grid_page_size = int(text)
            except ValueError:
                self.grid_page_size = 50
        # Reset to first page
        if not hasattr(self, 'grid_current_page'):
            self.grid_current_page = 0
        else:
            self.grid_current_page = 0
        # TODO: Implement full pagination logic with reload
    
    def create_list_view_widget_for_home(self):
        """Create List View widget adapted for home tab (assistance panel at bottom)"""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        
        # Vertical splitter: List on top, Editor in middle, Assistance at bottom
        home_list_splitter = QSplitter(Qt.Orientation.Vertical)
        
        # Top: Segment list
        list_container = QWidget()
        list_layout = QVBoxLayout(list_container)
        list_layout.setContentsMargins(0, 0, 0, 0)
        
        # Filter panel
        filter_panel = QWidget()
        filter_layout = QHBoxLayout(filter_panel)
        filter_layout.setContentsMargins(5, 5, 5, 5)
        
        source_filter_label = QLabel("Filter Source:")
        if not hasattr(self, 'list_source_filter'):
            self.list_source_filter = QLineEdit()
        self.list_source_filter.setPlaceholderText("Type to filter...")
        self.list_source_filter.textChanged.connect(self.apply_list_filters)
        
        target_filter_label = QLabel("Filter Target:")
        if not hasattr(self, 'list_target_filter'):
            self.list_target_filter = QLineEdit()
        self.list_target_filter.setPlaceholderText("Type to filter...")
        self.list_target_filter.textChanged.connect(self.apply_list_filters)
        
        clear_btn = QPushButton("Clear")
        clear_btn.clicked.connect(self.clear_list_filters)
        clear_btn.setMaximumWidth(60)
        
        filter_layout.addWidget(source_filter_label)
        filter_layout.addWidget(self.list_source_filter, stretch=1)
        filter_layout.addWidget(target_filter_label)
        filter_layout.addWidget(self.list_target_filter, stretch=1)
        filter_layout.addWidget(clear_btn)
        
        list_layout.addWidget(filter_panel)
        
        # Segment tree (QTreeWidget for list view)
        if not hasattr(self, 'list_tree'):
            self.list_tree = QTreeWidget()
            self.list_tree.setHeaderLabels(["#", "Type", "Status", "Source", "Target"])
            self.list_tree.setAlternatingRowColors(True)
            self.list_tree.setRootIsDecorated(False)
            self.list_tree.setSelectionMode(QTreeWidget.SelectionMode.ExtendedSelection)  # Allow Ctrl/Shift multi-selection
            # Enable context menu for bulk operations
            self.list_tree.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
            self.list_tree.customContextMenuRequested.connect(self.show_list_context_menu)
            
            # Column widths - all resizable
            header = self.list_tree.header()
            header.setSectionResizeMode(0, QHeaderView.ResizeMode.Interactive)  # ID - resizable
            header.setSectionResizeMode(1, QHeaderView.ResizeMode.Interactive)  # Type - resizable
            header.setSectionResizeMode(2, QHeaderView.ResizeMode.Interactive)  # Match
            header.setSectionResizeMode(3, QHeaderView.ResizeMode.Interactive)  # Status - resizable
            header.setSectionResizeMode(4, QHeaderView.ResizeMode.Interactive)  # Source - resizable
            header.setSectionResizeMode(5, QHeaderView.ResizeMode.Interactive)  # Target - resizable
            
            self.list_tree.setColumnWidth(0, 50)
            self.list_tree.setColumnWidth(1, 80)
            self.list_tree.setColumnWidth(2, 80)
            # Set default widths for Source and Target columns
            self.list_tree.setColumnWidth(3, 80)
            self.list_tree.setColumnWidth(4, 350)
            self.list_tree.setColumnWidth(5, 350)
            
            # Connect selection
            self.list_tree.itemSelectionChanged.connect(self.on_list_segment_selected)
            self.list_tree.itemDoubleClicked.connect(lambda: self.focus_list_target_editor())
        
        list_layout.addWidget(self.list_tree)
        home_list_splitter.addWidget(list_container)
        
        # Bottom: Add tabbed assistance panel (Translation Results, Segment Editor, Notes)
        # Create separate instance for list view (can't share widgets between views)
        self.home_list_tabbed_panel = self.create_tabbed_assistance_panel()
        home_list_splitter.addWidget(self.home_list_tabbed_panel)
        home_list_splitter.setSizes([600, 400])
        
        print(f"DEBUG: List view - Added tabbed panel to splitter")
        print(f"DEBUG: Tabbed panel has {self.home_list_tabbed_panel.count()} tabs")
        for i in range(self.home_list_tabbed_panel.count()):
            print(f"DEBUG:   Tab {i}: '{self.home_list_tabbed_panel.tabText(i)}'")
        
        main_layout.addWidget(home_list_splitter)
        
        # Store current selected segment for list view
        if not hasattr(self, 'list_current_segment_id'):
            self.list_current_segment_id = None
        
        return widget
    
    def create_document_view_widget_for_home(self):
        """Create Document View widget adapted for home tab (tabbed panel at bottom)"""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        
        # Vertical splitter: Document on top, Tabbed panel at bottom
        home_doc_splitter = QSplitter(Qt.Orientation.Vertical)
        
        # Top: Document flow area
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setStyleSheet("background-color: white;")
        
        if not hasattr(self, 'document_container'):
            self.document_container = QWidget()
            self.document_layout = QVBoxLayout(self.document_container)
            self.document_layout.setContentsMargins(0, 0, 0, 0)  # No margins - text widget has padding
            self.document_layout.setSpacing(0)  # No spacing - content adds its own
        
        scroll_area.setWidget(self.document_container)
        home_doc_splitter.addWidget(scroll_area)
        
        # Bottom: Tabbed panel (Translation Results | Segment Editor | Notes)
        self.home_doc_tabbed_panel = self.create_tabbed_assistance_panel()
        home_doc_splitter.addWidget(self.home_doc_tabbed_panel)
        
        home_doc_splitter.setSizes([600, 250])
        
        print(f"DEBUG: Document view - Added tabbed panel to splitter")
        print(f"DEBUG: Tabbed panel has {self.home_doc_tabbed_panel.count()} tabs")
        for i in range(self.home_doc_tabbed_panel.count()):
            print(f"DEBUG:   Tab {i}: '{self.home_doc_tabbed_panel.tabText(i)}'")
        
        main_layout.addWidget(home_doc_splitter)
        
        # Store segment widgets for document view
        if not hasattr(self, 'doc_segment_widgets'):
            self.doc_segment_widgets = {}
        if not hasattr(self, 'doc_current_segment_id'):
            self.doc_current_segment_id = None
        
        return widget
    
    def create_list_view_widget(self):
        """Create the List View widget (segment list with editor panel)"""
        print("="*80)
        print("DEBUG: *** create_list_view_widget() CALLED ***")
        print("="*80)
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        
        # Main vertical splitter: List on top, Editor at bottom
        self.list_splitter = QSplitter(Qt.Orientation.Vertical)
        
        # Top: Segment list
        list_container = QWidget()
        list_layout = QVBoxLayout(list_container)
        list_layout.setContentsMargins(0, 0, 0, 0)
        
        # Filter panel
        filter_panel = QWidget()
        filter_layout = QHBoxLayout(filter_panel)
        filter_layout.setContentsMargins(5, 5, 5, 5)
        
        source_filter_label = QLabel("Filter Source:")
        self.list_source_filter = QLineEdit()
        self.list_source_filter.setPlaceholderText("Type to filter...")
        self.list_source_filter.textChanged.connect(self.apply_list_filters)
        
        target_filter_label = QLabel("Filter Target:")
        self.list_target_filter = QLineEdit()
        self.list_target_filter.setPlaceholderText("Type to filter...")
        self.list_target_filter.textChanged.connect(self.apply_list_filters)
        
        clear_btn = QPushButton("Clear")
        clear_btn.clicked.connect(self.clear_list_filters)
        clear_btn.setMaximumWidth(60)
        
        filter_layout.addWidget(source_filter_label)
        filter_layout.addWidget(self.list_source_filter, stretch=1)
        filter_layout.addWidget(target_filter_label)
        filter_layout.addWidget(self.list_target_filter, stretch=1)
        filter_layout.addWidget(clear_btn)
        
        list_layout.addWidget(filter_panel)
        
        # Segment tree (QTreeWidget for list view)
        self.list_tree = QTreeWidget()
        self.list_tree.setHeaderLabels(["#", "Type", "Status", "Source", "Target"])
        self.list_tree.setAlternatingRowColors(True)
        self.list_tree.setRootIsDecorated(False)
        self.list_tree.setSelectionMode(QTreeWidget.SelectionMode.SingleSelection)
        
        # Column widths - all resizable
        header = self.list_tree.header()
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.Interactive)
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.Interactive)
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.Interactive)
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.Interactive)
        header.setSectionResizeMode(4, QHeaderView.ResizeMode.Interactive)

        self.list_tree.setColumnWidth(0, 50)
        self.list_tree.setColumnWidth(1, 80)
        self.list_tree.setColumnWidth(2, 160)
        # Set default widths for Source and Target columns
        self.list_tree.setColumnWidth(3, 360)
        self.list_tree.setColumnWidth(4, 360)
        
        # Connect selection
        self.list_tree.itemSelectionChanged.connect(self.on_list_segment_selected)
        self.list_tree.itemDoubleClicked.connect(lambda: self.focus_list_target_editor())
        
        list_layout.addWidget(self.list_tree)
        
        self.list_splitter.addWidget(list_container)
        
        # Bottom: Tabbed assistance panel (Translation Results | Segment Editor | Notes)
        # Use the same tabbed panel approach as Grid View for consistency
        self.list_tabbed_panel = self.create_tabbed_assistance_panel()
        self.list_splitter.addWidget(self.list_tabbed_panel)
        
        self.list_splitter.setSizes([600, 250])
        
        main_layout.addWidget(self.list_splitter)
        
        # Store current selected segment for list view
        self.list_current_segment_id = None
        
        return widget
    
    def create_document_view_widget(self):
        """Create the Document View widget (natural document flow)"""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        
        # Main vertical splitter: Document on top, Tabbed panel at bottom
        self.doc_splitter = QSplitter(Qt.Orientation.Vertical)
        
        # Top: Document flow area
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setStyleSheet("background-color: white;")
        
        self.document_container = QWidget()
        self.document_layout = QVBoxLayout(self.document_container)
        self.document_layout.setContentsMargins(0, 0, 0, 0)  # No margins - text widget has padding
        self.document_layout.setSpacing(0)  # No spacing - content will add its own
        
        scroll_area.setWidget(self.document_container)
        
        self.doc_splitter.addWidget(scroll_area)
        
        # Bottom: Tabbed panel (Translation Results | Segment Editor | Notes)
        # Create tabbed assistance panel for document view
        self.doc_tabbed_panel = self.create_tabbed_assistance_panel()
        self.doc_splitter.addWidget(self.doc_tabbed_panel)
        
        self.doc_splitter.setSizes([600, 250])
        
        main_layout.addWidget(self.doc_splitter)
        
        # Store segment widgets and current selection
        self.doc_segment_widgets = {}
        self.doc_current_segment_id = None
        
        return widget
    
    def create_translation_grid(self):
        """Create the translation grid (QTableWidget)"""
        self.table = QTableWidget()
        
        # Configure columns
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels(["#", "Type", "Source", "Target", "Status"])
        
        # Column widths - Source and Target columns stretch to fill space, others are interactive
        header = self.table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.Interactive)  # ID
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.Interactive)  # Type
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)  # Source - stretch to fill space
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)  # Target - stretch to fill space
        header.setSectionResizeMode(4, QHeaderView.ResizeMode.Interactive)  # Status
        header.setStretchLastSection(False)  # Don't auto-stretch last section (we use Stretch mode for Source/Target)
        
        # Set initial column widths - give Source and Target equal space
        self.table.setColumnWidth(0, 50)   # ID
        self.table.setColumnWidth(1, 100)  # Type
        self.table.setColumnWidth(2, 400)  # Source
        self.table.setColumnWidth(3, 400)  # Target
        self.table.setColumnWidth(4, 120)  # Status
        
        # Enable word wrap in cells (both display and edit mode)
        self.table.setWordWrap(True)
        
        # Apply custom delegate for word wrap in edit mode
        # Pass None for assistance_panel (keyboard shortcuts disabled for now)
        self.table.setItemDelegate(WordWrapDelegate(None, self.table))
        
        # Row behavior - Enable multi-selection with full row selection (memoQ-style)
        self.table.verticalHeader().setVisible(False)
        self.table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)  # Select full rows for multi-selection
        self.table.setSelectionMode(QTableWidget.SelectionMode.ExtendedSelection)  # Allow Ctrl/Shift multi-selection
        
        # Alternating row colors (simplified view)
        self.table.setAlternatingRowColors(True)
        
        # Simplified grid styling - more subtle, review-focused (companion tool philosophy)
        self.table.setStyleSheet("""
            QTableWidget {
                border: 1px solid #ddd;
                background-color: white;
                gridline-color: #e0e0e0;
            }
            QTableWidget::item {
                padding: 4px;
                border: none;
                border-bottom: 1px solid rgba(0, 0, 0, 0.08);
                border-right: 1px solid rgba(0, 0, 0, 0.08);
            }
            QTableWidget::item:selected {
                background-color: #e3f2fd;  /* Light blue instead of bright blue */
                color: black;
            }
            QTableWidget::item:focus {
                border: 1px solid #2196F3;
            }
            QTableWidget::item:last-child {
                border-right: none;
            }
        """)
        
        # Simplified editing: Double-click only (no F2 key) - companion tool philosophy
        # Grid is primarily for viewing/reviewing, with minor edits allowed
        self.table.setEditTriggers(QTableWidget.EditTrigger.DoubleClicked)
        
        # Enable context menu for bulk operations
        self.table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.table.customContextMenuRequested.connect(self.show_grid_context_menu)
        
        # Connect signals for debugging
        self.table.itemChanged.connect(self.on_cell_changed)
        self.table.currentCellChanged.connect(self.on_cell_selected)
        self.table.itemClicked.connect(self.on_cell_clicked)
        
        # Add additional selection signal for row-based selection mode
        self.table.itemSelectionChanged.connect(self.on_selection_changed)
        
        # Debug: Confirm signal connections
        self.log("🔌 Table signals connected: currentCellChanged, itemClicked, itemSelectionChanged")
    
    def create_assistance_panel(self):
        """DEPRECATED: Old assistance panel creation - now using create_tabbed_assistance_panel"""
        # This function is kept for backward compatibility but is no longer used
        # The new system uses multiple results panels in tabs
        pass
    
    def create_tabbed_assistance_panel(self, parent=None):
        """
        Create a tabbed panel with Translation Results, Segment Editor, and Notes
        This is used in both Grid and List views for consistency
        """
        tabs = QTabWidget(parent)
        tabs.setTabPosition(QTabWidget.TabPosition.North)
        tabs.setDocumentMode(False)
        
        # Tab 1: Translation Results - Create a NEW instance for this panel
        # (Can't reuse widgets - they can only have one parent)
        try:
            from modules.translation_results_panel import TranslationResultsPanel
            results_panel = TranslationResultsPanel(self)
            results_panel.match_selected.connect(self.on_match_selected)
            results_panel.match_inserted.connect(self.on_match_inserted)
            tabs.addTab(results_panel, "🔍 Translation Results")
            
            # Store reference so we can update it when segments are selected
            if not hasattr(self, 'results_panels'):
                self.results_panels = []
            self.results_panels.append(results_panel)
        except ImportError as e:
            print(f"Warning: Could not import TranslationResultsPanel: {e}")
            placeholder = QWidget()
            tabs.addTab(placeholder, "🔍 Translation Results")
        
        # Tab 2: Segment Editor
        editor_widget = QWidget()
        editor_layout = QVBoxLayout(editor_widget)
        editor_layout.setContentsMargins(10, 10, 10, 10)
        
        # Segment info header
        info_layout = QHBoxLayout()
        tab_seg_info = QLabel("Select a segment to edit")
        tab_seg_info.setStyleSheet("font-weight: bold; font-size: 11pt;")
        info_layout.addWidget(tab_seg_info, stretch=1)
        
        # Status selector
        status_label = QLabel("Status:")
        tab_status_combo = QComboBox()
        for status_key in STATUS_ORDER:
            definition = get_status(status_key)
            tab_status_combo.addItem(definition.label, status_key)
        tab_status_combo.currentIndexChanged.connect(self.on_tab_status_combo_changed)
        info_layout.addWidget(status_label)
        info_layout.addWidget(tab_status_combo)
        editor_layout.addLayout(info_layout)
        
        # Source text (read-only)
        source_label = QLabel("Source:")
        source_label.setStyleSheet("font-weight: bold;")
        editor_layout.addWidget(source_label)
        tab_source_editor = QTextEdit()
        tab_source_editor.setReadOnly(True)
        tab_source_editor.setMaximumHeight(100)
        tab_source_editor.setStyleSheet("background-color: #f5f5f5;")
        editor_layout.addWidget(tab_source_editor)
        
        # Target text (editable)
        target_label = QLabel("Target:")
        target_label.setStyleSheet("font-weight: bold;")
        editor_layout.addWidget(target_label)
        tab_target_editor = QTextEdit()
        tab_target_editor.setMaximumHeight(100)
        tab_target_editor.textChanged.connect(self.on_tab_target_change)
        editor_layout.addWidget(tab_target_editor)
        
        # Store references to THIS panel's widgets so we can update them
        editor_widget.seg_info_label = tab_seg_info
        editor_widget.source_editor = tab_source_editor
        editor_widget.target_editor = tab_target_editor
        editor_widget.status_combo = tab_status_combo
        
        # Action buttons
        button_layout = QHBoxLayout()
        copy_btn = QPushButton("📋 Copy Source → Target")
        copy_btn.clicked.connect(self.copy_source_to_tab_target)
        clear_btn = QPushButton("🗑️ Clear Target")
        clear_btn.clicked.connect(self.clear_tab_target)

        # Voice dictation button
        dictate_btn = QPushButton("🎤 Dictate (F9)")
        dictate_btn.setStyleSheet("background-color: #4CAF50; color: white; font-weight: bold;")
        dictate_btn.clicked.connect(self.start_voice_dictation)
        dictate_btn.setToolTip("Click or press F9 to start voice dictation")

        # Store reference to dictate button for state updates
        editor_widget.dictate_btn = dictate_btn

        save_btn = QPushButton("💾 Save")
        save_btn.setStyleSheet("background-color: #4CAF50; color: white; font-weight: bold;")
        save_btn.clicked.connect(self.save_tab_segment)
        save_next_btn = QPushButton("💾 Save & Next (Ctrl+Enter)")
        save_next_btn.setStyleSheet("background-color: #2196F3; color: white; font-weight: bold;")
        save_next_btn.clicked.connect(self.save_tab_segment_and_next)

        button_layout.addWidget(copy_btn)
        button_layout.addWidget(clear_btn)
        button_layout.addWidget(dictate_btn)
        button_layout.addStretch()
        button_layout.addWidget(save_btn)
        button_layout.addWidget(save_next_btn)
        editor_layout.addLayout(button_layout)
        
        editor_layout.addStretch()
        tabs.addTab(editor_widget, "📝 Segment Editor")
        
        # Tab 3: Comments Editor
        notes_widget = QWidget()
        notes_layout = QVBoxLayout(notes_widget)
        notes_layout.setContentsMargins(10, 10, 10, 10)
        
        notes_header = QLabel("Segment Comments")
        notes_header.setStyleSheet("font-weight: bold; font-size: 11pt;")
        notes_layout.addWidget(notes_header)
        
        tab_notes_editor = QTextEdit()
        tab_notes_editor.setPlaceholderText("Add comments for this segment...\n\n"
                                            "Comments are saved per segment and can include:\n"
                                            "• Translation context\n"
                                            "• Client preferences\n"
                                            "• Terminology decisions\n"
                                            "• Questions or clarifications needed")
        tab_notes_editor.textChanged.connect(self.on_tab_notes_change)
        notes_layout.addWidget(tab_notes_editor)
        
        # Store reference to THIS panel's comments editor
        notes_widget.notes_editor = tab_notes_editor
        
        notes_button_layout = QHBoxLayout()
        save_notes_btn = QPushButton("💾 Save Comments")
        save_notes_btn.setStyleSheet("background-color: #4CAF50; color: white; font-weight: bold;")
        save_notes_btn.clicked.connect(self.save_tab_notes)
        notes_button_layout.addStretch()
        notes_button_layout.addWidget(save_notes_btn)
        notes_layout.addLayout(notes_button_layout)
        
        tabs.addTab(notes_widget, "📋 Comments")
        
        # Store references to tab widgets for updates
        tabs.editor_widget = editor_widget
        tabs.notes_widget = notes_widget
        
        # Keep track of all tabbed panels
        if not hasattr(self, 'tabbed_panels'):
            self.tabbed_panels = []
        self.tabbed_panels.append(tabs)
        
        return tabs
    
    def on_match_selected(self, match):
        """Handle match selection from TranslationResultsPanel"""
        # This will be called when user selects a match in the panel
        # Can be used for highlighting, compare boxes, etc.
        pass
    
    def on_match_inserted(self, match_text: str):
        """
        Handle match insertion (user pressed Ctrl+number or Spacebar in match pane)
        Insert the match text into the currently selected target cell
        WITHOUT confirming the segment or moving to next
        """
        try:
            if not self.current_project or not self.table:
                return
            
            # Get current cell from grid
            current_item = self.table.currentItem()
            if current_item:
                row = current_item.row()
                col = current_item.column()
                
                # Check if we're in a target column (column 3 for Target)
                if col == 3 and row < len(self.current_project.segments):
                    # Just insert the match text - don't confirm or move
                    segment = self.current_project.segments[row]
                    segment.target = match_text
                    
                    # Update the table view
                    self.table.item(row, col).setText(match_text)
                    self.log(f"✓ Match inserted into segment {row + 1}")
                    
                    # Stay in current segment - don't move or confirm
                    # User is still in edit mode and can continue editing
        except Exception as e:
            self.log(f"Error inserting match: {e}")
    
    # ========================================================================
    # PROJECT MANAGEMENT
    # ========================================================================
    
    def new_project(self):
        """Create a new project"""
        from PyQt6.QtWidgets import (
            QVBoxLayout, QHBoxLayout, QFormLayout, QGroupBox, QTextEdit,
            QComboBox, QRadioButton, QButtonGroup, QPushButton, QTabWidget
        )
        from modules.simple_segmenter import SimpleSegmenter
        
        # Create dialog
        dialog = QDialog(self)
        dialog.setWindowTitle("New Translation Project")
        dialog.setMinimumWidth(700)
        dialog.setMinimumHeight(500)
        
        layout = QVBoxLayout(dialog)
        
        # Project Settings Group
        settings_group = QGroupBox("Project Settings")
        settings_layout = QFormLayout()
        
        # Project name
        name_input = QLineEdit()
        name_input.setText("Untitled Project")
        name_input.selectAll()
        settings_layout.addRow("Project Name:", name_input)
        
        # Source language
        source_lang_combo = QComboBox()
        common_langs = [
            ("English", "en"),
            ("Dutch", "nl"),
            ("German", "de"),
            ("French", "fr"),
            ("Spanish", "es"),
            ("Italian", "it"),
            ("Portuguese", "pt"),
            ("Russian", "ru"),
            ("Chinese", "zh"),
            ("Japanese", "ja"),
        ]
        for lang_name, lang_code in common_langs:
            source_lang_combo.addItem(lang_name, lang_code)
        settings_layout.addRow("Source Language:", source_lang_combo)
        
        # Target language
        target_lang_combo = QComboBox()
        for lang_name, lang_code in common_langs:
            target_lang_combo.addItem(lang_name, lang_code)
        
        # Set defaults based on global language settings (if in common_langs)
        try:
            for lang_name, lang_code in common_langs:
                if lang_name == self.source_language:
                    source_lang_combo.setCurrentText(lang_name)
                if lang_name == self.target_language:
                    target_lang_combo.setCurrentText(lang_name)
        except:
            target_lang_combo.setCurrentIndex(1)  # Fallback to Dutch
        
        settings_layout.addRow("Target Language:", target_lang_combo)
        
        settings_group.setLayout(settings_layout)
        layout.addWidget(settings_group)
        
        # Import Options Group
        import_group = QGroupBox("Import Source Text")
        import_layout = QVBoxLayout()
        
        # Tab widget for different import methods
        import_tabs = QTabWidget()
        
        # Tab 1: Paste Text
        paste_tab = QWidget()
        paste_layout = QVBoxLayout(paste_tab)
        paste_layout.addWidget(QLabel("Paste or type your source text below:"))
        
        text_input = QTextEdit()
        text_input.setPlaceholderText(
            "Paste your text here...\n\n"
            "Text will be automatically segmented into sentences.\n"
            "Each sentence becomes a translatable segment."
        )
        paste_layout.addWidget(text_input)
        import_tabs.addTab(paste_tab, "📝 Paste Text")
        
        # Tab 2: Load from File
        file_tab = QWidget()
        file_layout = QVBoxLayout(file_tab)
        file_layout.addWidget(QLabel("Load text from a file:"))
        
        file_path_display = QLineEdit()
        file_path_display.setPlaceholderText("No file selected...")
        file_path_display.setReadOnly(True)
        
        browse_btn = QPushButton("📁 Browse...")
        
        def browse_file():
            file_path, _ = QFileDialog.getOpenFileName(
                dialog,
                "Select Text File",
                "",
                "Text Files (*.txt);;All Files (*.*)"
            )
            if file_path:
                file_path_display.setText(file_path)
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text_input.setText(f.read())
                    import_tabs.setCurrentIndex(0)  # Switch to paste tab to show loaded text
                except Exception as e:
                    QMessageBox.warning(dialog, "Error", f"Could not load file:\n{e}")
        
        browse_btn.clicked.connect(browse_file)
        
        file_btn_layout = QHBoxLayout()
        file_btn_layout.addWidget(file_path_display)
        file_btn_layout.addWidget(browse_btn)
        file_layout.addLayout(file_btn_layout)
        file_layout.addStretch()
        
        import_tabs.addTab(file_tab, "📄 Load File")
        
        # Tab 3: Start Empty
        empty_tab = QWidget()
        empty_layout = QVBoxLayout(empty_tab)
        empty_layout.addWidget(QLabel("Create an empty project:"))
        empty_layout.addWidget(QLabel("You can add segments manually after creation."))
        empty_layout.addStretch()
        import_tabs.addTab(empty_tab, "\u2b50 Start Empty")
        
        import_layout.addWidget(import_tabs)
        import_group.setLayout(import_layout)
        layout.addWidget(import_group)
        
        # Buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        create_btn = QPushButton("Create Project")
        create_btn.setDefault(True)
        create_btn.setMinimumWidth(120)
        create_btn.clicked.connect(dialog.accept)
        
        cancel_btn = QPushButton("Cancel")
        cancel_btn.setMinimumWidth(120)
        cancel_btn.clicked.connect(dialog.reject)
        
        button_layout.addWidget(create_btn)
        button_layout.addWidget(cancel_btn)
        layout.addLayout(button_layout)
        
        # Show dialog
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return
        
        # Create project
        project_name = name_input.text().strip() or "Untitled Project"
        source_lang = source_lang_combo.currentData()
        target_lang = target_lang_combo.currentData()
        
        # Validate languages
        if source_lang == target_lang:
            QMessageBox.warning(
                self,
                "Invalid Languages",
                "Source and target languages must be different!"
            )
            return
        
        # Create project
        self.current_project = Project(
            name=project_name,
            source_lang=source_lang,
            target_lang=target_lang,
            segments=[]
        )
        
        # Process source text if provided
        source_text = text_input.toPlainText().strip()
        if source_text:
            try:
                segmenter = SimpleSegmenter()
                sentences = segmenter.segment_text(source_text)
                
                # Create segments
                for idx, sentence in enumerate(sentences, start=1):
                    if sentence.strip():
                        segment = Segment(
                            id=idx,
                            source=sentence.strip(),
                            target="",
                            status="untranslated"
                        )
                        self.current_project.segments.append(segment)
                
                self.log(f"Created {len(self.current_project.segments)} segments from source text")
            except Exception as e:
                QMessageBox.warning(
                    self,
                    "Segmentation Error",
                    f"Could not segment text:\n{e}\n\nProject created but no segments added."
                )
        
        # Update UI
        self.project_file_path = None
        self.project_modified = True  # Mark as modified since it hasn't been saved
        self.update_window_title()
        self.load_segments_to_grid()
        self.initialize_tm_database()  # Initialize TM for this project
        
        self.log(f"Created new project: {project_name} ({source_lang} → {target_lang})")
        
        # Prompt to save
        if self.current_project.segments:
            reply = QMessageBox.question(
                self,
                "Save Project",
                "Would you like to save your new project now?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                self.save_project_as()
    
    def closeEvent(self, event):
        """Handle application close - cleanup AHK process"""
        try:
            # Terminate AutoHotkey process if running
            if hasattr(self.lookup_tab, 'ahk_process') and self.lookup_tab.ahk_process:
                try:
                    self.lookup_tab.ahk_process.terminate()
                    self.lookup_tab.ahk_process.wait(timeout=2)
                    print("[Universal Lookup] AHK process terminated")
                except:
                    # Force kill if terminate doesn't work
                    try:
                        self.lookup_tab.ahk_process.kill()
                    except:
                        pass
        except Exception as e:
            print(f"[Universal Lookup] Error terminating AHK: {e}")
        
        # Accept the close event
        event.accept()
    
    def open_project(self):
        """Open a project file"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Open Project",
            str(self.user_data_path / "Projects"),
            "JSON Files (*.json);;All Files (*.*)"
        )
        
        if file_path:
            self.load_project(file_path)
    
    def load_project(self, file_path: str):
        """Load project from file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # If no name in file, use filename
            if 'name' not in data:
                data['name'] = Path(file_path).stem
            
            self.current_project = Project.from_dict(data)
            self.project_file_path = file_path
            self.project_modified = False
            
            # Restore prompt settings if they exist (unified library)
            if hasattr(self.current_project, 'prompt_settings') and self.current_project.prompt_settings:
                prompt_settings = self.current_project.prompt_settings
                if hasattr(self, 'prompt_manager_qt') and self.prompt_manager_qt:
                    library = self.prompt_manager_qt.library
                    
                    # Restore primary prompt
                    primary_path = prompt_settings.get('active_primary_prompt_path')
                    if primary_path and primary_path in library.prompts:
                        library.set_primary_prompt(primary_path)
                        self.log(f"✓ Restored primary prompt: {primary_path}")
                    
                    # Restore attached prompts
                    attached_paths = prompt_settings.get('attached_prompt_paths', [])
                    for path in attached_paths:
                        if path in library.prompts:
                            library.attach_prompt(path)
                            self.log(f"✓ Restored attached prompt: {path}")
                    
                    # Restore mode
                    mode = prompt_settings.get('mode', 'single')
                    if hasattr(self.prompt_manager_qt, 'set_mode'):
                        self.prompt_manager_qt.set_mode(mode)
            
            self.load_segments_to_grid()
            self.initialize_tm_database()  # Initialize TM for this project
            self.update_window_title()
            self.add_to_recent_projects(file_path)
            
            self.log(f"✓ Loaded project: {self.current_project.name} ({len(self.current_project.segments)} segments)")
            
            # Start background batch processing of termbase matches for all segments
            # This pre-fills the cache while user works on the project
            self._start_termbase_batch_worker()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load project:\n{str(e)}")
            self.log(f"✗ Error loading project: {e}")
    
    def _start_termbase_batch_worker(self):
        """
        Start background thread to batch-process termbase matches for all segments.
        This pre-fills the cache while the user works on the project.
        """
        if not self.current_project or len(self.current_project.segments) == 0:
            return
        
        # Stop any existing worker thread
        self.termbase_batch_stop_event.set()
        if self.termbase_batch_worker_thread and self.termbase_batch_worker_thread.is_alive():
            self.log("⏹️  Stopping existing termbase batch worker...")
            self.termbase_batch_worker_thread.join(timeout=2)
        
        # Reset stop event for new worker
        self.termbase_batch_stop_event.clear()
        
        # Start new background worker thread
        segment_count = len(self.current_project.segments)
        self.log(f"🔄 Starting background termbase batch processor for {segment_count} segments...")
        
        self.termbase_batch_worker_thread = threading.Thread(
            target=self._termbase_batch_worker_run,
            args=(self.current_project.segments,),
            daemon=True  # Daemon thread - won't prevent program exit
        )
        self.termbase_batch_worker_thread.start()
    
    def _termbase_batch_worker_run(self, segments):
        """
        Background worker thread: process all segments and populate termbase cache.
        Runs in separate thread to not block UI.
        
        IMPORTANT: Creates its own database connection to avoid SQLite threading errors.
        """
        if not segments:
            return
        
        # Create a separate database connection for this thread
        # SQLite connections are thread-local and cannot be shared across threads
        import sqlite3
        try:
            thread_db_connection = sqlite3.connect(self.db_manager.db_path)
            thread_db_connection.row_factory = sqlite3.Row
            thread_db_cursor = thread_db_connection.cursor()
        except Exception as e:
            self.log(f"❌ Failed to create database connection in batch worker: {e}")
            return
        
        try:
            processed = 0
            cached = 0
            start_time = time.time()
            
            for segment in segments:
                # Check if stop event was signaled (user closed project or started new one)
                if self.termbase_batch_stop_event.is_set():
                    self.log(f"⏹️  Termbase batch worker stopped by user (processed {processed} segments)")
                    break
                
                segment_id = segment.id
                
                # Skip if already in cache (thread-safe check)
                with self.termbase_cache_lock:
                    if segment_id in self.termbase_cache:
                        cached += 1
                        continue
                
                # Search termbase for this segment using thread-local database connection
                try:
                    # Manually query the database using thread-local connection
                    matches = self._search_termbases_thread_safe(
                        segment.source,
                        thread_db_cursor,
                        source_lang=self.current_project.source_lang if self.current_project else None,
                        target_lang=self.current_project.target_lang if self.current_project else None
                    )
                    
                    if matches:
                        # Store in cache (thread-safe)
                        with self.termbase_cache_lock:
                            self.termbase_cache[segment_id] = matches
                        
                        processed += 1
                        
                        # Log progress every 100 segments
                        if processed % 100 == 0:
                            elapsed = time.time() - start_time
                            rate = processed / elapsed if elapsed > 0 else 0
                            remaining = len(segments) - processed
                            eta_seconds = remaining / rate if rate > 0 else 0
                            self.log(f"📊 Batch progress: {processed}/{len(segments)} cached " +
                                   f"({rate:.1f} seg/sec, ETA: {int(eta_seconds)}s)")
                
                except Exception as e:
                    self.log(f"❌ Error processing segment {segment_id} in batch worker: {e}")
                    continue
                
                # Small delay to prevent CPU saturation (let UI thread work)
                time.sleep(0.001)  # 1ms delay between segments
            
            elapsed = time.time() - start_time
            total_cached = len(self.termbase_cache)
            self.log(f"✅ Termbase batch worker complete: {processed} new + {cached} existing = " +
                   f"{total_cached} total cached in {elapsed:.1f}s")
            
        except Exception as e:
            self.log(f"❌ Termbase batch worker error: {e}")
            import traceback
            self.log(traceback.format_exc())
        
        finally:
            # Close thread-local database connection
            try:
                thread_db_cursor.close()
                thread_db_connection.close()
                self.log("✓ Closed thread-local database connection in batch worker")
            except:
                pass
    
    def _search_termbases_thread_safe(self, source_text: str, cursor, source_lang: str = None, target_lang: str = None) -> Dict[str, str]:
        """
        Search termbases using a provided cursor (thread-safe for background threads).
        This method allows background workers to query the database without SQLite threading errors.
        
        Args:
            source_text: The source text to search for
            cursor: A database cursor from a thread-local connection
            source_lang: Source language code
            target_lang: Target language code
        
        Returns:
            Dictionary of {term: translation} matches
        """
        if not source_text or not cursor:
            return {}
        
        try:
            # Convert language names to codes
            source_lang_code = self._convert_language_to_code(source_lang) if source_lang else None
            target_lang_code = self._convert_language_to_code(target_lang) if target_lang else None
            
            # Split source text into words and search for each
            words = source_text.split()
            matches = {}
            
            for word in words:
                # Remove punctuation and search
                clean_word = word.strip('.,!?;:')
                if len(clean_word) < 2:  # Skip short words
                    continue
                
                # Search termbases in database using provided cursor
                try:
                    # Build query based on language constraints
                    query = "SELECT source_term, target_term FROM termbase_terms WHERE 1=1"
                    params = []
                    
                    # Add word search constraint (case-insensitive)
                    query += " AND LOWER(source_term) LIKE ?"
                    params.append(f"%{clean_word.lower()}%")
                    
                    # Add language constraints if specified
                    if source_lang_code:
                        query += " AND source_lang = ?"
                        params.append(source_lang_code)
                    
                    if target_lang_code:
                        query += " AND target_lang = ?"
                        params.append(target_lang_code)
                    
                    # Limit results
                    query += " LIMIT 5"
                    
                    # Execute using thread-local cursor
                    cursor.execute(query, params)
                    results = cursor.fetchall()
                    
                    if results:
                        for row in results:
                            source_term = row[0] if isinstance(row, tuple) else row['source_term']
                            target_term = row[1] if isinstance(row, tuple) else row['target_term']
                            if source_term and target_term:
                                matches[source_term.strip()] = target_term.strip()
                
                except Exception as e:
                    # Log but continue with other words
                    pass
            
            return matches
            
        except Exception as e:
            return {}
    
    def stop_termbase_batch_worker(self):
        """Stop the background termbase batch worker gracefully"""
        if self.termbase_batch_worker_thread and self.termbase_batch_worker_thread.is_alive():
            self.log("⏹️  Stopping termbase batch worker...")
            self.termbase_batch_stop_event.set()
            self.termbase_batch_worker_thread.join(timeout=3)
            self.log("✓ Termbase batch worker stopped")
    
    def save_project(self):
        """Save current project"""
        if not self.current_project:
            return
        
        if not self.project_file_path:
            self.save_project_as()
        else:
            self.save_project_to_file(self.project_file_path)
    
    def save_project_as(self):
        """Save project with new filename"""
        if not self.current_project:
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Project As",
            str(self.user_data_path / "Projects" / f"{self.current_project.name}.json"),
            "JSON Files (*.json);;All Files (*.*)"
        )
        
        if file_path:
            self.save_project_to_file(file_path)
            self.project_file_path = file_path
            self.add_to_recent_projects(file_path)
    
    def save_project_to_file(self, file_path: str):
        """Save project to specified file"""
        try:
            self.current_project.modified = datetime.now().isoformat()
            
            # Save prompt settings if prompt manager is available
            if hasattr(self, 'prompt_manager_qt') and self.prompt_manager_qt:
                # Unified prompt library - save active primary and attached prompts
                library = self.prompt_manager_qt.library
                self.current_project.prompt_settings = {
                    'active_primary_prompt_path': library.active_primary_prompt_path,
                    'attached_prompt_paths': library.attached_prompt_paths.copy() if library.attached_prompt_paths else [],
                    'mode': getattr(self.prompt_manager_qt, 'current_mode', 'single'),
                }
            
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(self.current_project.to_dict(), f, indent=2, ensure_ascii=False)
            
            self.project_modified = False
            self.update_window_title()
            self.log(f"✓ Saved project: {Path(file_path).name}")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save project:\n{str(e)}")
            self.log(f"✗ Error saving project: {e}")
    
    def close_project(self):
        """Close current project"""
        if not self.current_project:
            QMessageBox.information(self, "No Project", "No project is currently open")
            return
        
        # Check for unsaved changes
        if self.project_modified:
            reply = QMessageBox.question(
                self, "Unsaved Changes",
                "Save current project before closing?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No | QMessageBox.StandardButton.Cancel
            )
            if reply == QMessageBox.StandardButton.Cancel:
                return
            elif reply == QMessageBox.StandardButton.Yes:
                self.save_project()
                # If save was cancelled or failed, project_modified will still be True
                if self.project_modified:
                    return
        
        # Stop background batch worker
        self.stop_termbase_batch_worker()
        
        # Clear termbase cache for this project
        self.termbase_cache.clear()
        
        # Clear project data
        self.current_project = None
        self.project_file_path = None
        self.project_modified = False
        
        # Clear the grid
        self.clear_grid()
        
        # Clear translation results in all panels
        if hasattr(self, 'results_panels'):
            for panel in self.results_panels:
                if hasattr(panel, 'clear_results'):
                    panel.clear_results()
                if hasattr(panel, 'tm_results_table'):
                    panel.tm_results_table.setRowCount(0)
                if hasattr(panel, 'termbase_results_table'):
                    panel.termbase_results_table.setRowCount(0)
        
        # Update window title
        self.update_window_title()
        
        self.log("✓ Project closed")
    
    def update_recent_menu(self):
        """Update recent projects menu"""
        self.recent_menu.clear()
        
        recent_projects = self.load_recent_projects()
        
        if not recent_projects:
            no_recent = QAction("No recent projects", self)
            no_recent.setEnabled(False)
            self.recent_menu.addAction(no_recent)
            return
        
        # Add menu items for each recent project
        for i, project_info in enumerate(recent_projects[:self.MAX_RECENT_PROJECTS]):
            path = project_info['path']
            name = project_info['name']
            
            # Show relative path if in workspace, otherwise just filename
            try:
                rel_path = Path(path).relative_to(Path.cwd())
                display_name = f"{i+1}. {name} ({rel_path.parent})"
            except ValueError:
                display_name = f"{i+1}. {name}"
            
            action = QAction(display_name, self)
            action.setStatusTip(path)
            action.triggered.connect(lambda checked, p=path: self.load_project(p))
            self.recent_menu.addAction(action)
        
        # Add separator and "Clear Recent" option
        if recent_projects:
            self.recent_menu.addSeparator()
            clear_action = QAction("Clear Recent Projects", self)
            clear_action.triggered.connect(self.clear_recent_projects)
            self.recent_menu.addAction(clear_action)
    
    def add_to_recent_projects(self, file_path: str):
        """Add project to recent projects list"""
        if not file_path or not os.path.exists(file_path):
            return
        
        # Load existing recent projects
        recent_projects = self.load_recent_projects()
        
        # Create new entry
        new_entry = {
            'path': os.path.abspath(file_path),
            'name': self.current_project.name if self.current_project else Path(file_path).stem,
            'last_opened': datetime.now().isoformat()
        }
        
        # Remove if already exists (to move to top)
        recent_projects = [p for p in recent_projects if p['path'] != new_entry['path']]
        
        # Add to beginning
        recent_projects.insert(0, new_entry)
        
        # Keep only MAX_RECENT_PROJECTS
        recent_projects = recent_projects[:self.MAX_RECENT_PROJECTS]
        
        # Save back to file
        self.save_recent_projects(recent_projects)
        
        # Update menu and home display
        self.update_recent_menu()
        self.update_recent_projects_display()
    
    def load_recent_projects(self) -> List[Dict[str, str]]:
        """Load recent projects from file"""
        if not self.recent_projects_file.exists():
            return []
        
        try:
            with open(self.recent_projects_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Handle both old dict format and new list format
            if isinstance(data, dict):
                # Old format: convert to list
                recent = []
                for key, value in data.items():
                    if isinstance(value, list):
                        # Extract path and name from old list format
                        for item in value:
                            if isinstance(item, dict):
                                recent.append(item)
                            else:
                                # String path
                                recent.append({
                                    'path': str(item),
                                    'name': Path(str(item)).stem,
                                    'last_opened': datetime.now().isoformat()
                                })
                    elif isinstance(value, str):
                        recent.append({
                            'path': value,
                            'name': Path(value).stem,
                            'last_opened': datetime.now().isoformat()
                        })
                return recent
            elif isinstance(data, list):
                # New format: already a list
                # Ensure all entries have required fields
                normalized = []
                for item in data:
                    if isinstance(item, dict) and 'path' in item:
                        # Ensure all required fields exist
                        if 'name' not in item:
                            item['name'] = Path(item['path']).stem
                        if 'last_opened' not in item:
                            item['last_opened'] = datetime.now().isoformat()
                        # Only include if file still exists
                        if os.path.exists(item['path']):
                            normalized.append(item)
                return normalized
            
            return []
        
        except Exception as e:
            self.log(f"Error loading recent projects: {e}")
            return []
    
    def save_recent_projects(self, recent_projects: List[Dict[str, str]]):
        """Save recent projects to file"""
        try:
            # Ensure directory exists
            self.user_data_path.mkdir(parents=True, exist_ok=True)
            
            with open(self.recent_projects_file, 'w', encoding='utf-8') as f:
                json.dump(recent_projects, f, indent=2, ensure_ascii=False)
        
        except Exception as e:
            self.log(f"Error saving recent projects: {e}")
    
    def update_recent_projects_display(self):
        """Update the recent projects display on the home screen"""
        if not hasattr(self, 'recent_projects_layout'):
            return
        
        # Clear existing widgets
        for i in reversed(range(self.recent_projects_layout.count())):
            item = self.recent_projects_layout.itemAt(i)
            if item and item.widget():
                item.widget().deleteLater()
        
        # Load recent projects
        recent_projects = self.load_recent_projects()
        
        if not recent_projects:
            no_recent = QLabel("No recent projects")
            no_recent.setStyleSheet("color: #888; font-style: italic; padding: 10px;")
            no_recent.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.recent_projects_layout.addWidget(no_recent)
            return
        
        # Add project buttons
        for project_info in recent_projects[:5]:  # Show max 5 on home screen
            path = project_info['path']
            name = project_info.get('name', Path(path).stem)
            
            btn = QPushButton(f"📄 {name}")
            btn.setStyleSheet("text-align: left; padding: 8px; background: white; border: 1px solid #ddd; border-radius: 3px;")
            btn.setCursor(Qt.CursorShape.PointingHandCursor)
            btn.clicked.connect(lambda checked, p=path: self.load_project(p))
            self.recent_projects_layout.addWidget(btn)
        
        self.recent_projects_layout.addStretch()
    
    def clear_recent_projects(self):
        """Clear all recent projects"""
        reply = QMessageBox.question(
            self,
            "Clear Recent Projects",
            "Are you sure you want to clear all recent projects?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            self.save_recent_projects([])
            self.update_recent_menu()
            if hasattr(self, 'recent_projects_layout'):
                self.update_recent_projects_display()
            self.log("Recent projects cleared")
    
    # ========================================================================
    # MEMOQ BILINGUAL DOCX IMPORT/EXPORT
    # ========================================================================
    
    def import_docx(self):
        """Import a monolingual DOCX document"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select DOCX File",
            "",
            "Word Documents (*.docx);;All Files (*.*)"
        )
        
        if not file_path:
            return
        
        QMessageBox.information(
            self,
            "Monolingual Import",
            "You selected the Monolingual DOCX import workflow.\n\n"
            "Projects created with this option can be exported as standard formats "
            "(DOCX, plain text, etc.), but they cannot be exported as memoQ bilingual "
            "DOCX files. If you need memoQ round-tripping, use 'Import memoQ bilingual document'."
        )

        self.import_docx_from_path(file_path)
    
    def import_docx_from_path(self, file_path):
        """Import a monolingual DOCX document from a given path"""
        try:
            self.log(f"Importing: {os.path.basename(file_path)}")
            
            # Import DOCX using the existing docx_handler
            if not hasattr(self, 'docx_handler'):
                from modules.docx_handler import DOCXHandler
                self.docx_handler = DOCXHandler()
            
            # Detect document language automatically
            detected_lang = self._detect_docx_language(file_path)
            if detected_lang:
                self.log(f"📍 Detected document language: {detected_lang.upper()}")
            
            paragraphs = self.docx_handler.import_docx(file_path)
            self.original_docx = file_path
            
            # Segment paragraphs
            self.log("Segmenting text...")
            
            # Use simple segmenter if available, otherwise create segments from paragraphs
            if hasattr(self, 'segmenter'):
                segmented = self.segmenter.segment_paragraphs(paragraphs)
            else:
                # Simple fallback segmentation
                from modules.simple_segmenter import SimpleSegmenter
                segmenter = SimpleSegmenter()
                segmented = segmenter.segment_paragraphs(paragraphs)
            
            # Create new project with the imported segments
            from dataclasses import dataclass
            from typing import List, Dict, Any, Optional
            from datetime import datetime
            
            @dataclass
            class ImportedSegment:
                id: int
                source: str
                target: str = ""
                status: str = "untranslated"
                notes: str = ""
                para_id: Optional[int] = None
                is_table: bool = False
                table_info: Optional[tuple] = None
                style: str = "Normal"
                doc_position: int = 0
                type: str = "#"  # Default segment type
            
            # Convert segments
            imported_segments = []
            for seg_id, (para_id, text) in enumerate(segmented, 1):
                # Get paragraph info if available
                para_info = None
                if hasattr(self.docx_handler, '_get_para_info'):
                    para_info = self.docx_handler._get_para_info(para_id)
                
                is_table = False
                table_info = None
                style = "Normal"
                doc_position = para_id if para_id else seg_id
                
                if para_info:
                    style = para_info.style or "Normal"
                    doc_position = getattr(para_info, 'document_position', para_id)
                    if getattr(para_info, 'is_table_cell', False):
                        is_table = True
                        table_info = (
                            getattr(para_info, 'table_index', 0),
                            getattr(para_info, 'row_index', 0),
                            getattr(para_info, 'cell_index', 0)
                        )
                
                segment = ImportedSegment(
                    id=seg_id,
                    source=text,
                    target="",
                    status="untranslated",
                    notes="",
                    para_id=para_id,
                    is_table=is_table,
                    table_info=table_info,
                    style=style,
                    doc_position=doc_position
                )
                imported_segments.append(segment)
            
            # Convert ImportedSegment objects to proper Segment objects
            segments = []
            for imported_seg in imported_segments:
                segment = Segment(
                    id=imported_seg.id,
                    source=imported_seg.source,
                    target=imported_seg.target,
                    status=imported_seg.status,
                    notes=imported_seg.notes,
                    type=imported_seg.type,  # Use actual type from import
                    paragraph_id=imported_seg.para_id if imported_seg.para_id else 0,
                    style=imported_seg.style,  # CRITICAL: Preserve style for heading detection!
                    document_position=imported_seg.doc_position,
                    is_table_cell=imported_seg.is_table,
                    table_info=imported_seg.table_info
                )
                segments.append(segment)
            
            # Create new project using the proper Project class
            # Use detected language if available, otherwise use defaults
            detected_source = detected_lang if detected_lang else getattr(self, 'source_language', 'en')
            target_lang = getattr(self, 'target_language', 'nl')
            
            project = Project(
                name=f"DOCX Import - {os.path.basename(file_path)}",
                source_lang=detected_source,
                target_lang=target_lang,
                segments=segments
            )
            
            # Set as current project and load into grid
            self.current_project = project
            self.current_document_path = file_path  # Store document path
            self.load_segments_to_grid()

            # Initialize TM for this project
            self.initialize_tm_database()

            # Update status
            self.log(f"✓ Loaded {len(segments)} segments from {len(paragraphs)} paragraphs")
            self.log(f"📍 Project language pair: {project.source_lang.upper()} → {project.target_lang.upper()}")
            self.update_window_title()  # Update window title to show project is loaded

            # Auto-generate markdown if enabled
            if hasattr(self, 'auto_generate_markdown') and self.auto_generate_markdown:
                if hasattr(self, 'prompt_manager_qt'):
                    self.prompt_manager_qt.generate_markdown_for_current_document()

            # Refresh AI Assistant context
            if hasattr(self, 'prompt_manager_qt'):
                self.prompt_manager_qt.refresh_context()
            
            # Build success message with language info
            lang_info = ""
            if detected_lang:
                lang_info = f"\n\n📍 Detected document language: {detected_lang.upper()}"
                if detected_lang != 'en':
                    lang_info += f"\n🔄 Language pair: {detected_lang.upper()} → {project.target_lang.upper()}"
            
            QMessageBox.information(
                self, 
                "Import Complete", 
                f"Successfully imported {len(segments)} segments from:\n{os.path.basename(file_path)}{lang_info}"
            )
            
        except Exception as e:
            self.log(f"✗ Import failed: {str(e)}")
            QMessageBox.critical(
                self, 
                "Import Error", 
                f"Failed to import DOCX:\n\n{str(e)}"
            )
    
    def _interpret_memoq_status(self, status_text: str, has_target: bool) -> Tuple[str, Optional[int]]:
        """Map memoQ status text to internal status and extract match percentage."""
        status_def, match_percent = match_memoq_status(status_text)
        if status_def.key == DEFAULT_STATUS.key and has_target:
            return STATUSES["pretranslated"].key, match_percent
        return status_def.key, match_percent

    def import_memoq_bilingual(self):
        """Import memoQ bilingual DOCX file (table format)"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select memoQ Bilingual DOCX File",
            "",
            "Word Documents (*.docx);;All Files (*.*)"
        )
        
        if not file_path:
            return
        
        QMessageBox.information(
            self,
            "memoQ Bilingual Import",
            "You selected the memoQ Bilingual DOCX import.\n\n"
            "This workflow is designed for memoQ round-tripping and can be exported "
            "back to memoQ bilingual DOCX. Other export formats (such as Monolingual DOCX) "
            "are not available for this project."
        )

        try:
            from docx import Document
            
            # Load the bilingual DOCX
            doc = Document(file_path)
            
            if not doc.tables:
                QMessageBox.critical(
                    self, "Error",
                    "No table found in the DOCX file.\n\nExpected memoQ bilingual format with a table."
                )
                return
            
            table = doc.tables[0]
            
            # Validate table structure (should have at least 3 rows: header, column names, data)
            if len(table.rows) < 3:
                QMessageBox.critical(
                    self, "Error",
                    f"Invalid table structure.\n\nExpected at least 3 rows, found {len(table.rows)}."
                )
                return
            
            # Extract source and target segments from columns 1 and 2 (skipping header rows 0 and 1)
            # Also extract formatting information (bold, italic, underline)
            segments_data = []  # List of dictionaries with segment metadata
            formatting_map = {}  # segment_index -> list of formatting info
            
            for row_idx in range(2, len(table.rows)):
                row = table.rows[row_idx]
                
                # Ensure we have at least 3 cells (0=segment#, 1=source, 2=target)
                if len(row.cells) >= 3:
                    source_cell = row.cells[1]
                    target_cell = row.cells[2]
                    
                    source_text = source_cell.text.strip()
                    target_text = target_cell.text.strip()

                    comment_text = row.cells[3].text.strip() if len(row.cells) >= 4 else ""
                    status_text = row.cells[4].text.strip() if len(row.cells) >= 5 else ""
                    mapped_status, match_percent = self._interpret_memoq_status(status_text, bool(target_text))
                    
                    # Always add the row to maintain alignment
                    segment_idx = len(segments_data)
                    segments_data.append({
                        'source': source_text,
                        'target': target_text,
                        'status': mapped_status,
                        'memoq_status': status_text,
                        'comment': comment_text,
                        'match_percent': match_percent
                    })
                    
                    # Extract formatting from runs in source cell
                    formatting_info = []
                    for paragraph in source_cell.paragraphs:
                        for run in paragraph.runs:
                            run_text = run.text
                            if run_text:  # Only store runs with actual text
                                formatting_info.append({
                                    'text': run_text,
                                    'bold': run.bold == True,
                                    'italic': run.italic == True,
                                    'underline': run.underline == True
                                })
                    
                    # Store formatting for this segment
                    if formatting_info:
                        formatting_map[segment_idx] = formatting_info
            
            if not segments_data:
                QMessageBox.warning(self, "Warning", "No segments found in the bilingual file.")
                return
            
            # Create new project with the imported segments
            project_name = Path(file_path).stem
            
            # Detect languages from table header (row 1, columns 1 and 2)
            header_row = table.rows[1]
            source_lang = "en"  # Default
            target_lang = "nl"  # Default
            
            if len(header_row.cells) >= 3:
                source_header = header_row.cells[1].text.strip().lower()
                target_header = header_row.cells[2].text.strip().lower()
                
                # Try to detect language from header
                lang_map = {
                    'english': 'en', 'dutch': 'nl', 'german': 'de', 'french': 'fr',
                    'spanish': 'es', 'italian': 'it', 'portuguese': 'pt', 'polish': 'pl'
                }
                
                for lang_name, lang_code in lang_map.items():
                    if lang_name in source_header:
                        source_lang = lang_code
                    if lang_name in target_header:
                        target_lang = lang_code
            
            # Create project
            self.current_project = Project(
                name=project_name,
                source_lang=source_lang,
                target_lang=target_lang,
                segments=[]
            )
            
            # Create segments
            for idx, segment_entry in enumerate(segments_data):
                segment = Segment(
                    id=idx + 1,
                    source=segment_entry['source'],
                    target=segment_entry['target'],
                    status=segment_entry['status'],
                    type="para",
                    notes=segment_entry.get('comment', ""),
                    match_percent=segment_entry.get('match_percent'),
                    memoQ_status=segment_entry.get('memoq_status', "")
                )
                self.current_project.segments.append(segment)
            
            # Store the original memoQ file info for export
            self.memoq_source_file = file_path
            self.memoq_formatting_map = formatting_map
            
            # Update UI
            self.project_file_path = None
            self.project_modified = True
            self.update_window_title()
            self.load_segments_to_grid()
            self.initialize_tm_database()
            
            self.log(f"✓ Imported memoQ bilingual DOCX: {len(segments_data)} segments from {Path(file_path).name}")

            # Store current document path for AI Assistant
            self.current_document_path = file_path

            # Generate markdown for AI Assistant if enabled
            if hasattr(self, 'auto_generate_markdown') and self.auto_generate_markdown:
                if hasattr(self, 'prompt_manager_qt'):
                    self.prompt_manager_qt.generate_markdown_for_current_document()

            # Refresh AI Assistant context
            if hasattr(self, 'prompt_manager_qt'):
                self.prompt_manager_qt.refresh_context()

            QMessageBox.information(
                self, "Import Successful",
                f"Imported {len(segments_data)} segment(s) from memoQ bilingual DOCX.\n\n"
                f"File: {Path(file_path).name}\n"
                f"Languages: {source_lang} → {target_lang}"
            )
            
        except ImportError:
            QMessageBox.critical(
                self, "Missing Dependency",
                "The 'python-docx' library is required for memoQ bilingual DOCX import.\n\n"
                "Install it with: pip install python-docx"
            )
        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"Failed to import memoQ bilingual DOCX:\n\n{str(e)}")
            self.log(f"✗ memoQ import failed: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def _detect_docx_language(self, docx_path):
        """
        Detect language from DOCX file metadata and content
        Returns detected language code (e.g., 'nl', 'en') or None
        """
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            languages_found = set()
            
            with zipfile.ZipFile(docx_path, 'r') as docx:
                # Method 1: Check document.xml for language attributes
                try:
                    doc_xml = docx.read('word/document.xml')
                    root = ET.fromstring(doc_xml)
                    
                    # Define XML namespace
                    namespaces = {
                        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    }
                    
                    # Look for w:lang elements with val attribute
                    for lang_elem in root.findall('.//w:lang', namespaces):
                        val = lang_elem.get('val')
                        if val and val != 'none':
                            # Extract language code (e.g., "nl-NL" -> "nl")
                            lang_code = val.split('-')[0].lower()
                            if len(lang_code) == 2 and lang_code.isalpha():
                                languages_found.add(lang_code)
                    
                    # Look for eastAsia and bidi attributes too
                    for lang_elem in root.findall('.//w:lang', namespaces):
                        for attr in ['eastAsia', 'bidi']:
                            val = lang_elem.get(attr)
                            if val and val != 'none':
                                lang_code = val.split('-')[0].lower()
                                if len(lang_code) == 2 and lang_code.isalpha():
                                    languages_found.add(lang_code)
                                    
                except Exception as e:
                    self.log(f"Language detection warning (document.xml): {e}")
                
                # Method 2: Check settings.xml for default language
                try:
                    settings_xml = docx.read('word/settings.xml')
                    root = ET.fromstring(settings_xml)
                    
                    namespaces = {
                        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    }
                    
                    # Look for default language settings
                    for lang_elem in root.findall('.//w:themeFontLang', namespaces):
                        val = lang_elem.get('val')
                        if val and val != 'none':
                            lang_code = val.split('-')[0].lower()
                            if len(lang_code) == 2 and lang_code.isalpha():
                                languages_found.add(lang_code)
                                
                except Exception as e:
                    self.log(f"Language detection warning (settings.xml): {e}")
            
            # Return the most common language found, or None if none detected
            if languages_found:
                # Remove 'en' if other languages are found (often default fallback)
                if len(languages_found) > 1 and 'en' in languages_found:
                    languages_found.discard('en')
                
                detected = list(languages_found)[0]  # Take first available
                return detected
            
            return None
            
        except Exception as e:
            self.log(f"Language detection error: {e}")
            return None
    
    def export_memoq_bilingual(self):
        """Export to memoQ bilingual DOCX format with translations"""
        # Check if we have segments
        if not self.current_project or not self.current_project.segments:
            QMessageBox.warning(self, "No Data", "No segments to export")
            return

        # Check if a memoQ source file was imported, or prompt for it
        if not hasattr(self, 'memoq_source_file') or not self.memoq_source_file:
            # Prompt user to select the original memoQ bilingual file
            reply = QMessageBox.question(
                self, "Select memoQ Source File",
                "To export to memoQ format, please select the original memoQ bilingual DOCX file.\n\n"
                "This is the file you originally imported from memoQ.\n\n"
                "Would you like to select it now?",
                QMessageBox.Yes | QMessageBox.No
            )

            if reply == QMessageBox.Yes:
                file_path, _ = QFileDialog.getOpenFileName(
                    self,
                    "Select Original memoQ Bilingual DOCX",
                    "",
                    "Word Documents (*.docx);;All Files (*.*)"
                )

                if file_path:
                    self.memoq_source_file = file_path
                    self.log(f"✓ memoQ source file set: {Path(file_path).name}")
                else:
                    self.log("Export cancelled - no source file selected")
                    return
            else:
                self.log("Export cancelled")
                return
        
        try:
            from docx import Document
            from docx.shared import RGBColor
            
            segments = list(self.current_project.segments)
            translations = [seg.target for seg in segments]

            if not translations or all(not t.strip() for t in translations):
                QMessageBox.warning(self, "Warning", "No translations found to export.")
                return
            
            # Load the original bilingual DOCX
            doc = Document(self.memoq_source_file)
            table = doc.tables[0]
            
            # Write translations to target column (column 2) and update status
            segments_updated = 0
            segments_with_formatting = 0
            
            for i, translation in enumerate(translations):
                segment = segments[i]
                row_idx = i + 2  # Skip header rows (0 and 1)

                # Safety check: ensure we don't go beyond available rows
                if row_idx >= len(table.rows):
                    self.log(f"⚠ Warning: Row {row_idx} exceeds table rows ({len(table.rows)}), stopping at segment {i}")
                    break

                row = table.rows[row_idx]
                num_cells = len(row.cells)

                # Write translation to column 2 (target) with formatting
                if num_cells >= 3:
                    target_cell = row.cells[2]

                    # Get formatting info for this segment (if available)
                    formatting_info = None
                    if hasattr(self, 'memoq_formatting_map') and i in self.memoq_formatting_map:
                        formatting_info = self.memoq_formatting_map[i]
                        if any(f['bold'] or f['italic'] or f['underline'] for f in formatting_info):
                            segments_with_formatting += 1

                    # Apply formatting to the target cell
                    self._apply_formatting_to_cell(target_cell, translation, formatting_info)
                    segments_updated += 1

                # Update comments in column 3 (if column exists)
                if num_cells >= 4:
                    if segment.notes and segment.notes.strip():
                        row.cells[3].text = segment.notes.strip()

                # Update status column using compose_memoq_status (if column exists)
                if num_cells >= 5:
                    existing = row.cells[4].text
                    row.cells[4].text = compose_memoq_status(segment.status, segment.match_percent, existing)
            
            # Prompt user to save the updated bilingual file
            default_name = Path(self.memoq_source_file).stem + "_translated.docx"
            save_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save memoQ Bilingual DOCX",
                default_name,
                "Word Documents (*.docx);;All Files (*.*)"
            )
            
            if save_path:
                doc.save(save_path)
                
                # Build success message
                success_msg = (
                    f"✓ Successfully exported {segments_updated} translation(s) to memoQ bilingual DOCX!\n\n"
                    f"File saved: {Path(save_path).name}\n\n"
                )
                
                if segments_with_formatting > 0:
                    success_msg += (
                        f"✓ Formatting preserved in {segments_with_formatting} segment(s)\n"
                        f"  (bold, italic, underline)\n\n"
                    )
                
                success_msg += (
                    f"✓ Status updated to 'Confirmed'\n"
                    f"✓ Table structure preserved\n\n"
                    f"You can now import this file back into memoQ."
                )
                
                self.log(f"✓ Exported {segments_updated} translations to memoQ bilingual DOCX: {Path(save_path).name}")
                
                QMessageBox.information(self, "Export Successful", success_msg)
        
        except ImportError:
            QMessageBox.critical(
                self, "Missing Dependency",
                "The 'python-docx' library is required for memoQ bilingual DOCX export.\n\n"
                "Install it with: pip install python-docx"
            )
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export memoQ bilingual DOCX:\n\n{str(e)}")
            self.log(f"✗ memoQ export failed: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def _apply_formatting_to_cell(self, cell, text, formatting_info=None):
        """Apply formatting to a cell based on formatting info from source"""
        # Clear existing paragraphs
        cell._element.clear_content()
        
        # Add new paragraph
        paragraph = cell.add_paragraph()
        
        if not formatting_info or not text:
            # No formatting - just add plain text
            paragraph.add_run(text)
            return
        
        # Apply formatting based on source formatting
        # For simplicity and safety, just apply the first formatting style found to the entire text
        # This prevents text truncation issues that can occur with complex formatting algorithms
        
        # Check if we have any formatting at all
        has_bold = any(fmt.get('bold', False) for fmt in formatting_info)
        has_italic = any(fmt.get('italic', False) for fmt in formatting_info)
        has_underline = any(fmt.get('underline', False) for fmt in formatting_info)
        
        # Add the full text with formatting
        run = paragraph.add_run(text)
        
        if has_bold:
            run.bold = True
        if has_italic:
            run.italic = True
        if has_underline:
            run.underline = True
    
    # ========================================================================
    # GRID MANAGEMENT
    # ========================================================================
    
    def load_segments_to_grid(self):
        """Load segments into the grid with termbase highlighting"""
        if not self.current_project or not self.current_project.segments:
            self.clear_grid()
            return
        
        self.table.setRowCount(len(self.current_project.segments))
        
        for row, segment in enumerate(self.current_project.segments):
            # Clear any previous cell widgets
            self.table.removeCellWidget(row, 2)  # Source
            self.table.removeCellWidget(row, 3)  # Target
            self.table.removeCellWidget(row, 4)  # Match
            self.table.removeCellWidget(row, 5)  # Status
            
            # ID - Segment number (starts with black foreground, will be highlighted orange when selected)
            id_item = QTableWidgetItem(str(segment.id))
            id_item.setFlags(id_item.flags() & ~Qt.ItemFlag.ItemIsEditable)  # Read-only
            id_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            # Explicitly set black foreground for all segment numbers (will be changed to orange when selected)
            id_item.setForeground(QColor("black"))
            id_item.setBackground(QColor())  # Default (white) background
            self.table.setItem(row, 0, id_item)
            
            # Type - show segment type based on style and content
            # Determine type display from style attribute and segment type
            style = getattr(segment, 'style', 'Normal')
            
            # Check for list items (bullets/numbering or <li> tags)
            source_text = segment.source.strip()
            is_list_item = (
                source_text.startswith('<li>') or  # Tagged list item
                source_text.lstrip().startswith(('• ', '- ', '* ', '· ')) or
                (len(source_text) > 2 and source_text[0].isdigit() and source_text[1:3] in ('. ', ') '))
            )
            
            # Determine type display
            if 'Title' in style:
                type_display = "Title"
            elif 'Heading 1' in style or 'Heading1' in style:
                type_display = "H1"
            elif 'Heading 2' in style or 'Heading2' in style:
                type_display = "H2"
            elif 'Heading 3' in style or 'Heading3' in style:
                type_display = "H3"
            elif 'Heading 4' in style or 'Heading4' in style:
                type_display = "H4"
            elif 'Subtitle' in style:
                type_display = "Sub"
            elif is_list_item:
                type_display = "li"
            elif segment.type and segment.type != "para":
                type_display = segment.type.upper()
            else:
                type_display = "¶"  # Paragraph symbol
            
            type_item = QTableWidgetItem(type_display)
            type_item.setFlags(type_item.flags() & ~Qt.ItemFlag.ItemIsEditable)  # Read-only
            type_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            
            # Color-code by type for better visibility
            if type_display in ("H1", "H2", "H3", "H4", "Title"):
                type_item.setForeground(QColor("#003366"))  # Dark blue for headings
                type_item.setBackground(QColor("#e6f3ff"))  # Light blue background
            elif type_display == "li":
                type_item.setForeground(QColor("#006600"))  # Dark green for list items
                type_item.setBackground(QColor("#f0f8f0"))  # Light green background
            
            self.table.setItem(row, 1, type_item)
            
            # Source - Use read-only QTextEdit widget for easy text selection
            source_editor = ReadOnlyGridTextEditor(segment.source, self.table)
            
            # Initialize empty termbase matches (will be populated lazily on segment selection or by background worker)
            source_editor.termbase_matches = {}
            
            # Set font to match grid
            font = QFont(self.default_font_family, self.default_font_size)
            source_editor.setFont(font)
            
            # Set as cell widget (allows easy text selection)
            self.table.setCellWidget(row, 2, source_editor)
            
            # Also set a placeholder item for row height calculation
            source_item = QTableWidgetItem()
            source_item.setFlags(Qt.ItemFlag.NoItemFlags)  # No interaction
            self.table.setItem(row, 2, source_item)
            
            # Target - Use editable QTextEdit widget for easy text selection and editing
            target_editor = EditableGridTextEditor(segment.target, self.table, row, self.table)
            target_editor.setFont(font)
            
            # Connect text changes to update segment
            # Use a factory function to create a proper closure that captures the current row, segment, and editor
            def make_target_changed_handler(target_row, target_segment, editor_widget):
                def on_target_text_changed():
                    new_text = editor_widget.toPlainText()
                    if target_row < len(self.current_project.segments):
                        target_segment.target = new_text
                        # Update status if translation was added
                        if target_segment.target and target_segment.status in (DEFAULT_STATUS.key, 'pretranslated', 'rejected'):
                            target_segment.status = 'translated'
                            self.update_status_icon(target_row, target_segment.status)
                        self.project_modified = True
                        self.update_window_title()
                        # Auto-resize the row
                        self.table.resizeRowToContents(target_row)
                return on_target_text_changed
            
            target_editor.textChanged.connect(make_target_changed_handler(row, segment, target_editor))
            
            # Set as cell widget
            self.table.setCellWidget(row, 3, target_editor)
            
            # Also set a placeholder item for row height calculation
            target_item = QTableWidgetItem()
            target_item.setFlags(Qt.ItemFlag.NoItemFlags)  # No interaction
            self.table.setItem(row, 3, target_item)

            # Pre-populate status cell item so gridlines render before widget assignment
            status_placeholder = QTableWidgetItem()
            status_placeholder.setFlags(Qt.ItemFlag.NoItemFlags)
            status_placeholder.setBackground(QColor(get_status(segment.status).color))
            self.table.setItem(row, 4, status_placeholder)

            # Status column (icon + match + comment)
            self._update_status_cell(row, segment)
        
        # Apply current font
        self.apply_font_to_grid()
        
        # Auto-resize rows
        self.auto_resize_rows()
        self._enforce_status_row_heights()
        
        self.log(f"✓ Loaded {len(self.current_project.segments)} segments to grid")

        # Also refresh List and Document views if they exist
        if hasattr(self, 'list_tree'):
            self.refresh_list_view()
        if hasattr(self, 'document_container'):
            self.refresh_document_view()

    def _create_status_cell_widget(self, segment: Segment) -> QWidget:
        widget = QWidget()
        layout = QHBoxLayout(widget)
        layout.setContentsMargins(4, 2, 4, 2)  # Slightly more vertical padding
        layout.setSpacing(6)
        layout.setAlignment(Qt.AlignmentFlag.AlignVCenter)  # Center content vertically
        # Remove fixed minimum height - let it adapt to row height
        widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        widget.setAttribute(Qt.WidgetAttribute.WA_TransparentForMouseEvents, True)

        # Make widget background transparent - let table item handle the background
        widget.setStyleSheet("background: transparent;")
        status_def = get_status(segment.status)
        status_label = QLabel(status_def.icon)
        status_label.setAlignment(Qt.AlignmentFlag.AlignVCenter | Qt.AlignmentFlag.AlignLeft)
        status_label.setToolTip(status_def.label)
        # Slightly smaller X for "not_started" to match other icons better
        font_size = "10px" if segment.status == "not_started" else "14px"
        # Make confirmed checkmark green
        color = "color: #2e7d32;" if segment.status == "confirmed" else ""
        status_label.setStyleSheet(f"font-size: {font_size}; {color} padding-right: 4px;")
        layout.addWidget(status_label)

        # Only add match label if there's a match percentage
        if segment.match_percent is not None:
            match_text = f"{segment.match_percent}%"
            match_label = QLabel(match_text)
            match_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            match_label.setMinimumWidth(40)
            match_label.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Preferred)
            if segment.match_percent >= 101:
                match_label.setStyleSheet("color: #1b5e20; font-weight: bold; padding-left: 4px; padding-right: 4px;")
                match_label.setToolTip("Context match from memoQ (101% or better)")
            elif segment.match_percent >= 100:
                match_label.setStyleSheet("color: #2e7d32; font-weight: bold; padding-left: 4px; padding-right: 4px;")
                match_label.setToolTip("Exact match from memoQ (100%)")
            elif segment.match_percent >= 90:
                match_label.setStyleSheet("color: #1565C0; font-weight: bold; padding-left: 4px; padding-right: 4px;")
                match_label.setToolTip(f"High fuzzy match {segment.match_percent}%")
            else:
                match_label.setStyleSheet("color: #0d47a1; padding-left: 4px; padding-right: 4px;")
                match_label.setToolTip(f"Fuzzy match {segment.match_percent}%")
            layout.addWidget(match_label)

        # Use 🗨️ (left speech bubble) with better contrast
        comment_label = QLabel("🗨️")
        comment_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        # Style comment label (text-shadow not supported in Qt, removed)
        if segment.notes and segment.notes.strip():
            comment_label.setStyleSheet("""
                color: #ff9800;
                font-size: 14px;
                font-weight: bold;
            """)
            comment_label.setToolTip(segment.notes.strip())
        else:
            comment_label.setStyleSheet("""
                color: #90A4AE;
                font-size: 14px;
            """)
            comment_label.setToolTip("No comments")
        layout.addWidget(comment_label)

        layout.addStretch(1)

        return widget

    def _update_status_cell(self, row: int, segment: Segment):
        status_widget = self._create_status_cell_widget(segment)
        self.table.setCellWidget(row, 4, status_widget)
        status_item = self.table.item(row, 4)
        if status_item is None:
            status_item = QTableWidgetItem()
            self.table.setItem(row, 4, status_item)
        status_item.setText("")
        status_item.setFlags(Qt.ItemFlag.ItemIsSelectable | Qt.ItemFlag.ItemIsEnabled)
        # Set background color on table item (widget is transparent)
        status_item.setBackground(QColor(get_status(segment.status).color))
        status_item.setData(Qt.ItemDataRole.DisplayRole, "")
        status_widget.updateGeometry()

    def _refresh_segment_status(self, segment: Segment):
        if not self.current_project:
            return
        for row, seg in enumerate(self.current_project.segments):
            if seg.id == segment.id:
                self._update_status_cell(row, seg)
                break

        # Update list view entry in place (if visible)
        if hasattr(self, 'list_tree') and self.list_tree:
            matches = self.list_tree.findItems(str(segment.id), Qt.MatchFlag.MatchExactly, 0)
            for item in matches:
                status_def = get_status(segment.status)
                match_text = f"{segment.match_percent}%" if segment.match_percent is not None else ""
                display = f"{status_def.icon}"
                if match_text:
                    display += f"  {match_text}"
                if segment.notes and segment.notes.strip():
                    display += "  💬"
                item.setText(2, display)
                status_tooltip = status_def.label
                if segment.match_percent is not None:
                    status_tooltip += f" | {segment.match_percent}% match"
                if segment.notes and segment.notes.strip():
                    status_tooltip += f"\nComment: {segment.notes.strip()}"
                item.setToolTip(2, status_tooltip)
                item.setBackground(2, QColor(status_def.color))
        self._enforce_status_row_heights()

    def _enforce_status_row_heights(self):
        """No longer needed - status widgets now adapt to row height."""
        # This method is kept for backward compatibility but does nothing
        pass

    def refresh_list_view(self):
        """Refresh the List View with current segments"""
        if not hasattr(self, 'list_tree') or not self.current_project:
            return
        
        self.list_tree.clear()
        
        if not self.current_project.segments:
            return
        
        # Apply filters if any
        source_filter = self.list_source_filter.text().lower() if hasattr(self, 'list_source_filter') else ""
        target_filter = self.list_target_filter.text().lower() if hasattr(self, 'list_target_filter') else ""
        
        for segment in self.current_project.segments:
            # Filter segments
            if source_filter and source_filter not in segment.source.lower():
                continue
            if target_filter and target_filter not in segment.target.lower():
                continue
            
            status_def = get_status(segment.status)
            match_text = f"{segment.match_percent}%" if segment.match_percent is not None else ""
            status_display = f"{status_def.icon}"
            if match_text:
                status_display += f"  {match_text}"
            if segment.notes and segment.notes.strip():
                status_display += "  💬"
            item = QTreeWidgetItem([
                str(segment.id),
                segment.type.upper() if segment.type != "para" else "#",
                status_display,
                segment.source[:100] + "..." if len(segment.source) > 100 else segment.source,
                segment.target[:100] + "..." if len(segment.target) > 100 else segment.target
            ])
            item.setData(0, Qt.ItemDataRole.UserRole, segment.id)  # Store segment ID
            
            # Color coding by status
            item.setBackground(2, QColor(status_def.color))

            # Tooltip for match/status columns
            status_tooltip = status_def.label
            if segment.match_percent is not None:
                status_tooltip += f" | {segment.match_percent}% match"
            if segment.notes and segment.notes.strip():
                status_tooltip += f"\nComment: {segment.notes.strip()}"
            item.setToolTip(2, status_tooltip)
            
            self.list_tree.addTopLevelItem(item)
        
        self.log(f"✓ Refreshed List View with {self.list_tree.topLevelItemCount()} segments")
    
    def refresh_document_view(self):
        """Refresh the Document View with current segments (Natural document flow like Tkinter)"""
        if not hasattr(self, 'document_container') or not self.current_project:
            return
        
        # Clear existing widgets
        while self.document_layout.count() > 0:
            item = self.document_layout.itemAt(0)
            if item:
                widget = item.widget()
                if widget:
                    widget.deleteLater()
                self.document_layout.removeItem(item)
        
        self.doc_segment_widgets.clear()
        
        if not self.current_project.segments:
            return
        
        # Apply filters if any
        source_filter = self.source_filter.text().lower() if hasattr(self, 'source_filter') else ""
        target_filter = self.target_filter.text().lower() if hasattr(self, 'target_filter') else ""
        
        # Create a single QTextEdit for the entire document (exactly like Tkinter's Text widget)
        doc_text = QTextEdit()
        doc_text.setReadOnly(True)
        
        # Simple styling - let it fill space naturally
        doc_text.setStyleSheet("""
            QTextEdit {
                background-color: white;
                border: none;
                padding: 30px 40px;
                font-size: 11pt;
            }
        """)
        
        # CRITICAL: Let the widget expand to fill available space (like Tkinter's pack(fill='both', expand=True))
        doc_text.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        
        cursor = doc_text.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.Start)
        
        # Smart paragraph grouping: Use paragraph_id if available, otherwise use heuristics
        paragraphs = {}
        current_para = 0
        prev_seg = None
        
        for seg in self.current_project.segments:
            # Check if segment has explicit paragraph_id
            if hasattr(seg, 'paragraph_id') and seg.paragraph_id > 0:
                para_id = seg.paragraph_id
            else:
                # Use heuristics to detect paragraph breaks
                if prev_seg:
                    prev_text = prev_seg.source.strip()
                    curr_text = seg.source.strip()
                    style = getattr(seg, 'style', 'Normal')
                    
                    # Check for paragraph break indicators
                    is_new_paragraph = (
                        (prev_text and prev_text[-1] in '.!?' and curr_text and curr_text[0].isupper()) or
                        'Heading' in style or
                        'Title' in style
                    )
                    
                    if is_new_paragraph:
                        current_para += 1
                    
                para_id = current_para
            
            if para_id not in paragraphs:
                paragraphs[para_id] = []
            paragraphs[para_id].append(seg)
            prev_seg = seg
        
        # Render each paragraph
        prev_para_id = None
        for para_id in sorted(paragraphs.keys()):
            para_segments = paragraphs[para_id]
            
            # Add paragraph break (double newline) between paragraphs
            if prev_para_id is not None:
                cursor.insertText("\n\n")
            prev_para_id = para_id
            
            # Check if this is a heading or special style
            first_seg = para_segments[0]
            style = getattr(first_seg, 'style', 'Normal')
            
            # DEBUG: Log first few paragraph styles to BOTH console and log
            if para_id < 5:
                debug_msg = f"DEBUG: Para {para_id} - Style: '{style}' - Text: {first_seg.source[:50]}..."
                print(debug_msg)
                self.log(debug_msg)
            
            # Determine formatting based on style
            char_format = QTextCharFormat()
            char_format.setFontFamily(self.default_font_family)
            
            if 'Heading 1' in style or 'Title' in style:
                char_format.setFontPointSize(16)
                char_format.setFontWeight(QFont.Weight.Bold)
                char_format.setForeground(QColor('#003366'))
            elif 'Heading 2' in style:
                char_format.setFontPointSize(14)
                char_format.setFontWeight(QFont.Weight.Bold)
                char_format.setForeground(QColor('#0066cc'))
            elif 'Heading 3' in style:
                char_format.setFontPointSize(12)
                char_format.setFontWeight(QFont.Weight.Bold)
                char_format.setForeground(QColor('#3399ff'))
            elif 'Subtitle' in style:
                char_format.setFontPointSize(12)
                char_format.setFontItalic(True)
                char_format.setForeground(QColor('#663399'))
            else:
                char_format.setFontPointSize(11)
                char_format.setForeground(QColor('#000000'))
            
            # Insert each segment in the paragraph
            for i, seg in enumerate(para_segments):
                # Apply filters
                if source_filter and source_filter not in seg.source.lower():
                    continue
                if target_filter and target_filter not in seg.target.lower():
                    continue
                
                # Add space between sentences (except for first in paragraph)
                if i > 0:
                    cursor.insertText(" ")
                
                # Determine what text to display
                if seg.target and seg.target.strip():
                    display_text = seg.target
                elif seg.source:
                    display_text = seg.source
                else:
                    display_text = f"[Segment {seg.id} - Empty]"
                
                # Set background color based on status
                if seg.status == 'untranslated':
                    char_format.setBackground(QColor('#ffe6e6'))
                elif seg.status == 'translated':
                    char_format.setBackground(QColor('#e6ffe6'))
                elif seg.status == 'approved':
                    char_format.setBackground(QColor('#e6f3ff'))
                else:
                    char_format.setBackground(QColor('white'))
                
                # Store position for click handling
                start_pos = cursor.position()
                
                # Parse and render inline formatting tags (<b>, <i>, <u>, <bi>, <li>)
                import re
                
                # Pattern to match formatting tags: <b>, </b>, <i>, </i>, <u>, </u>, <bi>, </bi>, <li>, </li>
                tag_pattern = re.compile(r'<(/?)([biu]|bi|li)>')
                
                # Check for list item tag at the start
                list_item_match = re.match(r'^<li>(.*)</li>$', display_text, re.DOTALL)
                if list_item_match:
                    # This is a list item - insert bullet and extract content
                    bullet_format = QTextCharFormat(char_format)
                    bullet_format.setFontWeight(QFont.Weight.Bold)
                    bullet_format.setForeground(QColor('#FF6600'))  # Orange bullet
                    cursor.insertText("• ", bullet_format)
                    
                    # Process the content inside <li> tags
                    display_text = list_item_match.group(1)
                
                # Split text into parts (text and tags)
                parts = []
                last_end = 0
                formatting_stack = []  # Stack to track nested formatting
                
                for match in tag_pattern.finditer(display_text):
                    # Add text before tag
                    if match.start() > last_end:
                        text_part = display_text[last_end:match.start()]
                        parts.append(('text', text_part, formatting_stack.copy()))
                    
                    # Process tag
                    is_closing = match.group(1) == '/'
                    tag_type = match.group(2)
                    
                    if is_closing:
                        # Remove from stack if present
                        if tag_type in formatting_stack:
                            formatting_stack.remove(tag_type)
                    else:
                        # Add to stack
                        if tag_type not in formatting_stack:
                            formatting_stack.append(tag_type)
                    
                    last_end = match.end()
                
                # Add remaining text
                if last_end < len(display_text):
                    parts.append(('text', display_text[last_end:], formatting_stack.copy()))
                
                # If no tags found, render as plain text
                if not parts:
                    cursor.insertText(display_text, char_format)
                else:
                    # Render each part with appropriate formatting
                    for part_type, text, format_tags in parts:
                        if part_type == 'text' and text:
                            # Create format with current formatting
                            part_format = QTextCharFormat(char_format)
                            
                            # Apply formatting based on tags
                            if 'bi' in format_tags or ('b' in format_tags and 'i' in format_tags):
                                part_format.setFontWeight(QFont.Weight.Bold)
                                part_format.setFontItalic(True)
                            elif 'b' in format_tags:
                                part_format.setFontWeight(QFont.Weight.Bold)
                            elif 'i' in format_tags:
                                part_format.setFontItalic(True)
                            
                            if 'u' in format_tags:
                                part_format.setFontUnderline(True)
                            
                            cursor.insertText(text, part_format)
                
                end_pos = cursor.position()
                
                # Store segment info for click handling
                self.doc_segment_widgets[seg.id] = {
                    'start': start_pos,
                    'end': end_pos,
                    'segment': seg
                }
        
        # Enable mouse tracking for click handling
        doc_text.setMouseTracking(True)
        doc_text.viewport().setCursor(Qt.CursorShape.PointingHandCursor)
        
        # Connect mouse click event
        def handle_click(event):
            cursor = doc_text.cursorForPosition(event.pos())
            pos = cursor.position()
            
            # Find which segment was clicked
            for seg_id, info in self.doc_segment_widgets.items():
                if info['start'] <= pos <= info['end']:
                    self.on_doc_segment_clicked(seg_id)
                    break
        
        doc_text.mousePressEvent = handle_click
        
        # Add the text widget to the layout - it will expand to fill available space
        self.document_layout.addWidget(doc_text)
        
        self.log(f"✓ Refreshed Document View with {len(self.doc_segment_widgets)} segments (natural flow)")
    
    def clear_grid(self):
        """Clear all rows from grid"""
        self.table.setRowCount(0)
    
    def auto_resize_rows(self):
        """Auto-resize all rows to fit content - Compact version"""
        if not hasattr(self, 'table') or not self.table:
            return
        
        # Manually calculate and set row heights for compact display
        for row in range(self.table.rowCount()):
            max_height = 1
            
            # Check source cell (column 2)
            source_widget = self.table.cellWidget(row, 2)
            if source_widget and isinstance(source_widget, ReadOnlyGridTextEditor):
                doc = source_widget.document()
                if doc:  # Safety check
                    col_width = self.table.columnWidth(2)
                    if col_width > 0:
                        doc.setTextWidth(col_width)
                        size = doc.size()
                        if size.isValid():  # Safety check
                            height = int(size.height())
                            max_height = max(max_height, height)
            
            # Check target cell (column 3)
            target_widget = self.table.cellWidget(row, 3)
            if target_widget and isinstance(target_widget, EditableGridTextEditor):
                doc = target_widget.document()
                if doc:  # Safety check
                    col_width = self.table.columnWidth(3)
                    if col_width > 0:
                        doc.setTextWidth(col_width)
                        size = doc.size()
                        if size.isValid():  # Safety check
                            height = int(size.height())
                            max_height = max(max_height, height)
            
            # Set row height with minimal padding (2px total)
            # Minimum 32px to accommodate status icons (16px) + match text + padding without any cutoff
            compact_height = max(max_height + 2, 32)
            self.table.setRowHeight(row, compact_height)
        
        self.log("✓ Auto-resized rows to fit content (compact)")
        self._enforce_status_row_heights()
    
    def apply_font_to_grid(self):
        """Apply selected font to all grid cells"""
        font = QFont(self.default_font_family, self.default_font_size)
        
        self.table.setFont(font)
        
        # Also update header font
        header_font = QFont(self.default_font_family, self.default_font_size, QFont.Weight.Bold)
        self.table.horizontalHeader().setFont(header_font)
        
        # Update fonts in QTextEdit widgets (source and target columns)
        if hasattr(self, 'table') and self.table:
            for row in range(self.table.rowCount()):
                # Source column (2) - ReadOnlyGridTextEditor
                source_widget = self.table.cellWidget(row, 2)
                if source_widget and isinstance(source_widget, ReadOnlyGridTextEditor):
                    source_widget.setFont(font)
                
                # Target column (3) - EditableGridTextEditor
                target_widget = self.table.cellWidget(row, 3)
                if target_widget and isinstance(target_widget, EditableGridTextEditor):
                    target_widget.setFont(font)
    
    def set_font_family(self, family_name: str):
        """Set font family from menu"""
        self.default_font_family = family_name
        self.apply_font_to_grid()
        self.auto_resize_rows()
        self.log(f"✓ Font changed to {family_name}")
    
    def increase_font_size(self):
        """Increase font size (Ctrl++)"""
        self.default_font_size = min(72, self.default_font_size + 1)
        self.apply_font_to_grid()
        self.auto_resize_rows()
        self.log(f"✓ Font size: {self.default_font_size}")
        # Save font size to preferences
        self.save_current_font_sizes()
    
    def decrease_font_size(self):
        """Decrease font size (Ctrl+-)"""
        self.default_font_size = max(7, self.default_font_size - 1)
        self.apply_font_to_grid()
        self.auto_resize_rows()
        self.log(f"✓ Font size: {self.default_font_size}")
        # Save font size to preferences
        self.save_current_font_sizes()
    
    def on_font_changed(self):
        """Handle font change - legacy method for compatibility"""
        self.apply_font_to_grid()
        self.auto_resize_rows()

    
    def zoom_in(self):
        """Increase font size (same as Ctrl++)"""
        self.increase_font_size()
    
    def zoom_out(self):
        """Decrease font size (same as Ctrl+-)"""
        self.decrease_font_size()
    
    def results_pane_zoom_in(self):
        """Increase font size in translation results pane"""
        if hasattr(self, 'results_panels'):
            for panel in self.results_panels:
                if hasattr(panel, 'zoom_in'):
                    panel.zoom_in()
            # Save font sizes to preferences
            self.save_current_font_sizes()
    
    def results_pane_zoom_out(self):
        """Decrease font size in translation results pane"""
        if hasattr(self, 'results_panels'):
            for panel in self.results_panels:
                if hasattr(panel, 'zoom_out'):
                    panel.zoom_out()
            # Save font sizes to preferences
            self.save_current_font_sizes()
    
    def results_pane_zoom_reset(self):
        """Reset font size in translation results pane to default"""
        if hasattr(self, 'results_panels'):
            for panel in self.results_panels:
                if hasattr(panel, 'reset_zoom'):
                    panel.reset_zoom()
            # Save font sizes to preferences
            self.save_current_font_sizes()
    
    def save_current_font_sizes(self):
        """Save current font sizes to preferences"""
        try:
            from modules.translation_results_panel import CompactMatchItem, TranslationResultsPanel
            general_settings = self.load_general_settings()
            general_settings['grid_font_size'] = self.default_font_size
            if hasattr(CompactMatchItem, 'font_size_pt'):
                general_settings['results_match_font_size'] = CompactMatchItem.font_size_pt
            if hasattr(TranslationResultsPanel, 'compare_box_font_size'):
                general_settings['results_compare_font_size'] = TranslationResultsPanel.compare_box_font_size
            # Preserve other settings
            if 'restore_last_project' not in general_settings:
                general_settings['restore_last_project'] = False
            self.save_general_settings(general_settings)
        except Exception as e:
            # Silently fail - don't interrupt user workflow
            pass
    
    def load_general_settings(self) -> Dict[str, Any]:
        """Load general application settings"""
        # Initialize auto-propagation from loaded settings
        settings = self._load_general_settings_from_file()
        if 'auto_propagate_exact_matches' in settings:
            self.auto_propagate_exact_matches = settings['auto_propagate_exact_matches']
        # Load TM/termbase matching setting
        if 'enable_tm_termbase_matching' in settings:
            self.enable_tm_matching = settings['enable_tm_termbase_matching']
            self.enable_termbase_matching = settings['enable_tm_termbase_matching']
        # Load auto-markdown setting
        self.auto_generate_markdown = settings.get('auto_generate_markdown', False)

        # Load LLM provider settings for AI Assistant
        llm_settings = self.load_llm_settings()
        self.current_provider = llm_settings.get('provider', 'openai')
        provider_key = f"{self.current_provider}_model"
        self.current_model = llm_settings.get(provider_key)

        return settings
    
    def _load_general_settings_from_file(self) -> Dict[str, Any]:
        """Load general settings from user preferences"""
        prefs_file = self.user_data_path / "ui_preferences.json"
        
        defaults = {
            'restore_last_project': False,
            'auto_propagate_exact_matches': True,
            'grid_font_size': 11,
            'results_match_font_size': 9,
            'results_compare_font_size': 9
        }
        
        if not prefs_file.exists():
            return defaults
        
        try:
            with open(prefs_file, 'r') as f:
                prefs = json.load(f)
                general = prefs.get('general_settings', {})
                # Merge with defaults to ensure all keys exist
                result = defaults.copy()
                result.update(general)
                return result
        except:
            return defaults
    
    def save_general_settings(self, settings: Dict[str, Any]):
        """Save general settings to user preferences"""
        prefs_file = self.user_data_path / "ui_preferences.json"
        
        # Load existing preferences
        prefs = {}
        if prefs_file.exists():
            try:
                with open(prefs_file, 'r') as f:
                    prefs = json.load(f)
            except:
                pass
        
        # Update general settings
        prefs['general_settings'] = settings
        
        # Save back
        try:
            with open(prefs_file, 'w') as f:
                json.dump(prefs, f, indent=2)
        except Exception as e:
            self.log(f"⚠ Could not save general settings: {str(e)}")

    def load_dictation_settings(self) -> Dict[str, Any]:
        """Load voice dictation settings"""
        prefs_file = self.user_data_path / "ui_preferences.json"

        defaults = {
            'model': 'base',
            'max_duration': 10,
            'language': 'Auto (use project target language)'
        }

        if not prefs_file.exists():
            return defaults

        try:
            with open(prefs_file, 'r') as f:
                prefs = json.load(f)
                dictation = prefs.get('dictation_settings', {})
                result = defaults.copy()
                result.update(dictation)
                return result
        except:
            return defaults

    def save_dictation_settings(self, model: str, duration: int, language: str):
        """Save voice dictation settings"""
        prefs_file = self.user_data_path / "ui_preferences.json"

        # Load existing preferences
        prefs = {}
        if prefs_file.exists():
            try:
                with open(prefs_file, 'r') as f:
                    prefs = json.load(f)
            except:
                pass

        # Update dictation settings
        prefs['dictation_settings'] = {
            'model': model,
            'max_duration': duration,
            'language': language
        }

        # Save back
        try:
            with open(prefs_file, 'w') as f:
                json.dump(prefs, f, indent=2)
            self.log(f"✓ Voice dictation settings saved: Model={model}, Duration={duration}s")
            QMessageBox.information(self, "Settings Saved",
                f"Voice dictation settings saved successfully!\n\n"
                f"Model: {model}\n"
                f"Max Duration: {duration} seconds\n"
                f"Language: {language}")
        except Exception as e:
            self.log(f"⚠ Could not save dictation settings: {str(e)}")
            QMessageBox.warning(self, "Save Error", f"Could not save settings:\n{str(e)}")

    def load_language_settings(self):
        """Load language settings from preferences"""
        prefs_file = self.user_data_path / "ui_preferences.json"

        defaults = {
            'source_language': 'English',
            'target_language': 'Dutch'
        }
        
        if not prefs_file.exists():
            return
        
        try:
            with open(prefs_file, 'r') as f:
                prefs = json.load(f)
                lang_settings = prefs.get('language_settings', {})
                self.source_language = lang_settings.get('source_language', defaults['source_language'])
                self.target_language = lang_settings.get('target_language', defaults['target_language'])
        except:
            pass
    
    def save_language_settings(self, source_lang: str, target_lang: str):
        """Save language settings to preferences"""
        prefs_file = self.user_data_path / "ui_preferences.json"
        
        # Load existing preferences
        prefs = {}
        if prefs_file.exists():
            try:
                with open(prefs_file, 'r') as f:
                    prefs = json.load(f)
            except:
                pass
        
        # Update language settings
        prefs['language_settings'] = {
            'source_language': source_lang,
            'target_language': target_lang
        }
        
        # Save back
        try:
            with open(prefs_file, 'w') as f:
                json.dump(prefs, f, indent=2)
        except Exception as e:
            self.log(f"⚠ Could not save language settings: {str(e)}")
    
    def _save_language_settings_from_ui(self, source_combo, target_combo):
        """Save language settings from UI"""
        self.source_language = source_combo.currentText()
        self.target_language = target_combo.currentText()
        self.save_language_settings(self.source_language, self.target_language)
        
        self.log(f"✓ Language settings saved: {self.source_language} → {self.target_language}")
        QMessageBox.information(self, "Settings Saved", 
                              f"Language settings have been saved:\n\n"
                              f"Source: {self.source_language}\n"
                              f"Target: {self.target_language}")
    
    def load_font_sizes_from_preferences(self):
        """Load and apply font sizes from preferences on startup"""
        try:
            general_settings = self.load_general_settings()
            
            # Load grid font size
            grid_size = general_settings.get('grid_font_size', 11)
            if 7 <= grid_size <= 72:
                self.default_font_size = grid_size
                if hasattr(self, 'table') and self.table is not None:
                    self.apply_font_to_grid()
            
            # Load results pane font sizes
            match_size = general_settings.get('results_match_font_size', 9)
            compare_size = general_settings.get('results_compare_font_size', 9)
            
            # Apply to all results panels
            if hasattr(self, 'results_panels'):
                if 7 <= match_size <= 16:
                    from modules.translation_results_panel import CompactMatchItem
                    CompactMatchItem.set_font_size(match_size)
                    for panel in self.results_panels:
                        if hasattr(panel, 'set_font_size'):
                            panel.set_font_size(match_size)
                
                if 7 <= compare_size <= 14:
                    from modules.translation_results_panel import TranslationResultsPanel
                    TranslationResultsPanel.compare_box_font_size = compare_size
                    for panel in self.results_panels:
                        if hasattr(panel, 'set_compare_box_font_size'):
                            panel.set_compare_box_font_size(compare_size)
                    
        except Exception as e:
            self.log(f"⚠ Could not load font sizes: {e}")
    
    def restore_last_project_if_enabled(self):
        """Restore the last opened project on startup if the setting is enabled"""
        try:
            # Check if setting is enabled
            general_settings = self.load_general_settings()
            if not general_settings.get('restore_last_project', False):
                return
            
            # Get recent projects
            recent_projects = self.load_recent_projects()
            if not recent_projects:
                return
            
            # Get the most recent project (first in the list)
            # Recent projects are sorted by most recent first
            last_project = recent_projects[0]
            project_path = last_project.get('path') if isinstance(last_project, dict) else last_project
            
            # Verify file exists
            if not project_path or not os.path.exists(project_path):
                self.log(f"⚠ Last project file not found: {project_path}")
                return
            
            # Load the project
            self.log(f"🔄 Restoring last project: {Path(project_path).name}")
            self.load_project(project_path)
            
        except Exception as e:
            # Silently fail - don't block startup if restoration fails
            self.log(f"⚠ Could not restore last project: {e}")
    
    
    def get_status_icon(self, status: str) -> str:
        """Get status icon for display"""
        return get_status(status).icon
    
    def update_status_icon(self, row: int, status: str):
        """Update status icon for a specific row"""
        if not self.current_project or row >= len(self.current_project.segments):
            return
        segment = self.current_project.segments[row]
        segment.status = status
        self._refresh_segment_status(segment)
    
    def on_cell_changed(self, item: QTableWidgetItem):
        """Handle cell content changes - now mainly for placeholder items"""
        # Target cell changes are handled by the EditableGridTextEditor's textChanged signal
        # This method is kept for compatibility but should rarely be called now
        # Placeholder items (columns 2 and 3) are non-interactive, so this shouldn't trigger
        pass
    
    def on_cell_selected(self, current_row, current_col, previous_row, previous_col):
        """Handle cell selection change"""
        self.log(f"🎯 on_cell_selected called: row {current_row}, col {current_col}")
        
        # DEBUG: Also log all connected signals for this table
        self.log(f"🔗 Table currentCellChanged signal connected: {self.table.currentCellChanged}")
        self.log(f"🔗 Table itemClicked signal connected: {self.table.itemClicked}")
        self.log(f"🔗 Table itemSelectionChanged signal connected: {self.table.itemSelectionChanged}")
        try:
            # Clear previous highlighting
            if previous_row >= 0 and previous_row < self.table.rowCount():
                prev_id_item = self.table.item(previous_row, 0)
                if prev_id_item:
                    prev_id_item.setBackground(QColor())  # Reset background to default
                    prev_id_item.setForeground(QColor("black"))  # Reset text color to black
            
            # Highlight current segment number in orange (like memoQ)
            if current_row >= 0 and current_row < self.table.rowCount():
                current_id_item = self.table.item(current_row, 0)
                if current_id_item:
                    current_id_item.setBackground(QColor("#FFA500"))  # Orange background
                    current_id_item.setForeground(QColor("white"))    # White text for contrast
            
            if not self.current_project or current_row < 0:
                return
            
            if current_row < len(self.current_project.segments):
                segment = self.current_project.segments[current_row]
                
                # Update tab segment editor
                if hasattr(self, 'update_tab_segment_editor'):
                    self.update_tab_segment_editor(
                        segment_id=segment.id,
                        source_text=segment.source,
                        target_text=segment.target,
                        status=segment.status,
                        notes=segment.notes
                    )
                
                # Get termbase matches (from cache or search on-demand)
                matches_dict = None  # Initialize at the top level
                try:
                    segment_id = segment.id
                    source_widget = self.table.cellWidget(current_row, 2)  # Source column is column 2
                    
                    # Check cache first (thread-safe)
                    stored_matches = {}
                    with self.termbase_cache_lock:
                        if segment_id in self.termbase_cache:
                            stored_matches = self.termbase_cache[segment_id]
                            self.log(f"✅ Cache HIT: Retrieved {len(stored_matches)} termbase matches for segment {segment_id}")
                    
                    # If not in cache, search on-demand (lazy loading)
                    if not stored_matches and source_widget:
                        self.log(f"🔍 Cache MISS: Searching termbases for segment {segment_id}...")
                        stored_matches = self.find_termbase_matches_in_source(segment.source)
                        
                        # Store in cache for future access (thread-safe)
                        if stored_matches:
                            with self.termbase_cache_lock:
                                self.termbase_cache[segment_id] = stored_matches
                            self.log(f"💾 Cached {len(stored_matches)} matches for segment {segment_id}")
                    
                    # Store in widget for backwards compatibility
                    if source_widget and hasattr(source_widget, 'termbase_matches'):
                        source_widget.termbase_matches = stored_matches
                    
                    if stored_matches:
                        # Convert stored matches to TranslationMatch objects
                        from modules.translation_results_panel import TranslationMatch
                        matches_dict = {
                            "LLM": [],
                            "NT": [],
                            "MT": [],
                            "TM": [],
                            "Termbases": []
                        }
                        
                        for source_term, target_term in stored_matches.items():
                            match_obj = TranslationMatch(
                                source=source_term,
                                target=target_term,
                                relevance=95,  # High relevance for termbase matches
                                metadata={
                                    'termbase_name': 'Default',
                                    'definition': '',
                                    'domain': ''
                                },
                                match_type='Termbase',
                                compare_source=source_term,
                                provider_code='TB'
                            )
                            matches_dict["Termbases"].append(match_obj)
                        
                        # Get current project languages for all translation services
                        source_lang = getattr(self.current_project, 'source_lang', None) if self.current_project else None
                        target_lang = getattr(self.current_project, 'target_lang', None) if self.current_project else None
                        
                        # Convert language names to codes if needed
                        if source_lang:
                            source_lang_code = self._convert_language_to_code(source_lang)
                        if target_lang:
                            target_lang_code = self._convert_language_to_code(target_lang)
                        
                        # Show immediate termbase matches, delay expensive TM/MT/LLM searches
                        self.log(f"🚀 Immediate display: {len(matches_dict['Termbases'])} termbase matches")
                        
                        # Update all tabbed results panels
                        if hasattr(self, 'results_panels'):
                            for panel in self.results_panels:
                                try:
                                    panel.set_matches(matches_dict)
                                except Exception as e:
                                    self.log(f"Error updating results panel: {e}")
                    else:
                        self.log("📋 No stored termbase matches found")
                except Exception as e:
                    self.log(f"Error retrieving stored termbase matches: {e}")
                    
                # Schedule expensive searches (TM, MT, LLM) with debouncing to prevent UI blocking
                # Pass the termbase matches to preserve them in delayed search
                if matches_dict:
                    termbase_matches = matches_dict.get('Termbases', [])
                    self._schedule_mt_and_llm_matches(segment, termbase_matches)
        except Exception as e:
            self.log(f"Critical error in on_cell_selected: {e}")
    
    def on_selection_changed(self):
        """Handle selection change for row-based selection mode"""
        try:
            # Get currently selected row
            current_row = self.table.currentRow()
            current_col = self.table.currentColumn()
            
            self.log(f"🔄 on_selection_changed called: row {current_row}, col {current_col}")
            
            # If we have a valid row selection, trigger the same logic as on_cell_selected
            if current_row >= 0:
                # Call on_cell_selected with appropriate parameters
                self.on_cell_selected(current_row, current_col, -1, -1)
        except Exception as e:
            self.log(f"Error in on_selection_changed: {e}")
    
    def test_cell_selection(self):
        """Test method to manually trigger cell selection for debugging"""
        self.log("🧪 TEST: Manually triggering cell selection for row 0")
        if self.table.rowCount() > 0:
            # Try to programmatically select first row
            self.table.selectRow(0)
            self.table.setCurrentCell(0, 2)  # Select source column
            self.on_cell_selected(0, 2, -1, -1)  # Manually trigger
        else:
            self.log("🧪 TEST: No rows available for testing")
    
    def get_selected_segments_from_grid(self):
        """Get list of selected segments from grid view"""
        if not self.current_project or not hasattr(self, 'table'):
            return []
        
        selected_rows = set()
        for item in self.table.selectedItems():
            selected_rows.add(item.row())
        
        segments = []
        for row in sorted(selected_rows):
            if 0 <= row < len(self.current_project.segments):
                segments.append(self.current_project.segments[row])
        
        return segments
    
    def get_selected_segments_from_list(self):
        """Get list of selected segments from list view"""
        if not self.current_project or not hasattr(self, 'list_tree'):
            return []
        
        selected_items = self.list_tree.selectedItems()
        segments = []
        
        for item in selected_items:
            segment_id = item.data(0, Qt.ItemDataRole.UserRole)
            if segment_id:
                # Find segment by ID
                for segment in self.current_project.segments:
                    if segment.id == segment_id:
                        segments.append(segment)
                        break
        
        return segments
    
    def show_grid_context_menu(self, position):
        """Show context menu for grid view with bulk operations"""
        selected_segments = self.get_selected_segments_from_grid()
        
        if not selected_segments:
            return
        
        menu = QMenu(self)
        
        # Clear translations action
        clear_action = menu.addAction("🗑️ Clear Translations")
        clear_action.setToolTip(f"Clear translations for {len(selected_segments)} selected segment(s)")
        clear_action.triggered.connect(lambda: self.clear_selected_translations(selected_segments, 'grid'))
        
        menu.addSeparator()
        
        # Select all action
        select_all_action = menu.addAction("📋 Select All (Ctrl+A)")
        select_all_action.triggered.connect(lambda: self.table.selectAll())
        
        menu.exec(self.table.viewport().mapToGlobal(position))
    
    def show_list_context_menu(self, position):
        """Show context menu for list view with bulk operations"""
        selected_segments = self.get_selected_segments_from_list()
        
        if not selected_segments:
            return
        
        menu = QMenu(self)
        
        # Clear translations action
        clear_action = menu.addAction("🗑️ Clear Translations")
        clear_action.setToolTip(f"Clear translations for {len(selected_segments)} selected segment(s)")
        clear_action.triggered.connect(lambda: self.clear_selected_translations(selected_segments, 'list'))
        
        menu.addSeparator()
        
        # Select all action
        select_all_action = menu.addAction("📋 Select All (Ctrl+A)")
        select_all_action.triggered.connect(lambda: self.list_tree.selectAll())
        
        menu.exec(self.list_tree.viewport().mapToGlobal(position))
    
    def clear_selected_translations(self, segments, view_type='grid'):
        """Clear translations for selected segments"""
        if not segments:
            return
        
        count = len(segments)
        cleared_count = 0
        
        for segment in segments:
            if segment.target:  # Only clear if there's something to clear
                segment.target = ""
                if segment.status != 'untranslated':
                    segment.status = 'untranslated'
                cleared_count += 1
        
        if cleared_count > 0:
            self.project_modified = True
            self.update_window_title()
            
            # Refresh the appropriate view
            if view_type == 'grid':
                self.load_segments_to_grid()
            elif view_type == 'list':
                self.refresh_list_view()
            
            self.log(f"✓ Cleared translations for {cleared_count} segment(s)")
        else:
            self.log(f"ℹ No translations to clear for selected segments")
    
    def clear_selected_translations_from_menu(self):
        """Clear translations for selected segments (called from Edit menu)"""
        # Determine which view is active
        if hasattr(self, 'home_view_stack') and self.home_view_stack:
            current_index = self.home_view_stack.currentIndex()
            if current_index == 0:  # Grid view
                selected_segments = self.get_selected_segments_from_grid()
                if selected_segments:
                    self.clear_selected_translations(selected_segments, 'grid')
                else:
                    QMessageBox.information(self, "No Selection", "Please select one or more segments to clear translations.")
            elif current_index == 1:  # List view
                selected_segments = self.get_selected_segments_from_list()
                if selected_segments:
                    self.clear_selected_translations(selected_segments, 'list')
                else:
                    QMessageBox.information(self, "No Selection", "Please select one or more segments to clear translations.")
            else:
                QMessageBox.information(self, "Not Available", "Bulk operations are only available in Grid and List views.")
        else:
            QMessageBox.information(self, "Not Available", "Please load a project first.")
    
    def _restore_active_prompt(self, prompt_type: str, prompt_name: str):
        """Restore an active prompt by name"""
        if not hasattr(self, 'prompt_manager_qt') or not self.prompt_manager_qt:
            return
        
        try:
            if prompt_type == 'translate':
                # Find and activate domain prompt
                for i in range(self.prompt_manager_qt.domain_tree.topLevelItemCount()):
                    item = self.prompt_manager_qt.domain_tree.topLevelItem(i)
                    filename = item.data(0, Qt.ItemDataRole.UserRole)
                    if filename:
                        prompt_data = self.prompt_manager_qt.prompt_library.get_prompt(filename)
                        if prompt_data and prompt_data.get('name') == prompt_name:
                            self.prompt_manager_qt.domain_tree.setCurrentItem(item)
                            self.prompt_manager_qt._activate_domain_expertise('translate')
                            break
            elif prompt_type == 'proofread':
                # Find and activate domain prompt
                for i in range(self.prompt_manager_qt.domain_tree.topLevelItemCount()):
                    item = self.prompt_manager_qt.domain_tree.topLevelItem(i)
                    filename = item.data(0, Qt.ItemDataRole.UserRole)
                    if filename:
                        prompt_data = self.prompt_manager_qt.prompt_library.get_prompt(filename)
                        if prompt_data and prompt_data.get('name') == prompt_name:
                            self.prompt_manager_qt.domain_tree.setCurrentItem(item)
                            self.prompt_manager_qt._activate_domain_expertise('proofread')
                            break
            elif prompt_type == 'project':
                # Find and activate project prompt
                for i in range(self.prompt_manager_qt.project_tree.topLevelItemCount()):
                    item = self.prompt_manager_qt.project_tree.topLevelItem(i)
                    filename = item.data(0, Qt.ItemDataRole.UserRole)
                    if filename:
                        prompt_data = self.prompt_manager_qt.prompt_library.get_prompt(filename)
                        if prompt_data and prompt_data.get('name') == prompt_name:
                            self.prompt_manager_qt.project_tree.setCurrentItem(item)
                            self.prompt_manager_qt._activate_project_guideline()
                            break
        except Exception as e:
            self.log(f"⚠ Could not restore {prompt_type} prompt '{prompt_name}': {e}")
    
    def _restore_active_style_guide(self, guide_name: str, language: str):
        """Restore an active style guide by name and language
        
        NOTE: Temporarily disabled for unified prompt library migration.
        TODO: Implement session restoration for unified library (save/restore active prompts)
        """
        if not hasattr(self, 'prompt_manager_qt') or not self.prompt_manager_qt:
            return
        
        # TODO: Adapt this for unified prompt library
        # The new system uses multi-attach, so we'd need to:
        # 1. Find the style guide in library.prompts
        # 2. Attach it using library.attach_prompt(path)
        # For now, skip restoration - user can re-select prompts
        self.log(f"⚠ Style guide restoration not yet implemented in unified library: '{guide_name}' ({language})")
    
    def _schedule_delayed_lookup(self, segment, current_row):
        """Schedule TM and termbase lookup with delay (cancel if user moves to another segment)"""
        self.log(f"🚀 _schedule_delayed_lookup called for segment {current_row + 1}: '{segment.source[:50]}...'")
        
        # Cancel any pending lookup
        if self.lookup_timer:
            self.lookup_timer.stop()
            self.lookup_timer = None
        
        # Store current segment ID for checking if user moved away
        segment_id = id(segment)
        self.current_lookup_segment_id = segment_id
        
        # If matching is disabled, skip
        self.log(f"🔧 SCHEDULING CHECK: TM enabled={self.enable_tm_matching}, Termbase enabled={self.enable_termbase_matching}")
        if not (self.enable_tm_matching or self.enable_termbase_matching):
            self.log(f"🔧 SCHEDULING BLOCKED: Both TM and termbase matching are disabled")
            return
        self.log(f"🔧 SCHEDULING OK: At least one matching type is enabled")
        
        # Get user-configured delay (default 1.5 seconds)
        general_prefs = self.load_general_settings()
        lookup_delay = general_prefs.get('lookup_delay', 1500)
        
        # Schedule lookup after user-configured delay (allows user to navigate quickly without triggering searches)
        from PyQt6.QtCore import QTimer
        self.lookup_timer = QTimer()
        self.lookup_timer.setSingleShot(True)
        self.lookup_timer.timeout.connect(lambda: self._perform_delayed_lookup(segment, current_row, segment_id))
        self.lookup_timer.start(lookup_delay)
        self.log(f"🚀 _schedule_delayed_lookup: Will search segment {current_row + 1} in {lookup_delay}ms")
    
    def _perform_delayed_lookup(self, segment, current_row, expected_segment_id):
        """Perform TM and termbase lookup after delay (only if user is still on same segment)"""
        self.log(f"� DELAYED LOOKUP STARTED for segment {current_row + 1}: '{segment.source[:50]}...'")
        self.log(f"🚨 DELAYED LOOKUP: expected_id={expected_segment_id}, current_id={self.current_lookup_segment_id}")
        
        # Check if user moved to another segment
        if self.current_lookup_segment_id != expected_segment_id:
            self.log(f"� DELAYED LOOKUP CANCELLED: User moved to different segment")
            return  # User moved away, cancel lookup
            
        self.log(f"🚨 DELAYED LOOKUP CONTINUING: User still on same segment")
        
        # Check if still on the same row
        if current_row < 0 or current_row >= self.table.rowCount():
            return
        
        current_segment = self.table.item(current_row, 2)  # Source column
        if not current_segment or current_segment.text() != segment.source:
            return  # Segment changed
        
        try:
            # Update TM matches (if enabled)
            if self.enable_tm_matching:
                try:
                    self.search_and_display_tm_matches(segment.source)
                except Exception as e:
                    self.log(f"Error searching TM: {e}")
            
            # Update TranslationResultsPanel if available
            if hasattr(self, 'assistance_widget') and hasattr(self.assistance_widget, 'set_segment_info'):
                try:
                    # Generate matches dictionary for the panel
                    from modules.translation_results_panel import TranslationMatch
                    
                    matches_dict = {
                        "LLM": [],     # LLM Translation (appears first)
                        "NT": [],      # No Translation
                        "MT": [],      # Machine Translation
                        "TM": [],      # Translation Memory
                        "Termbases": [] # Terminology
                    }
                    
                    # Generate TM matches from database if available (and enabled)
                    self.log(f"� DEBUG: Reached TM search section in _perform_delayed_lookup")
                    self.log(f"�🔍 TM SEARCH CHECK: TM enabled={self.enable_tm_matching}, has_db_manager={hasattr(self, 'db_manager')}, db_manager_valid={self.db_manager is not None if hasattr(self, 'db_manager') else False}")
                    if self.enable_tm_matching and hasattr(self, 'db_manager') and self.db_manager:
                        try:
                            self.log(f"🟢 TM SEARCH: Searching for matches for '{segment.source[:50]}...'")
                            
                            # Get current project languages
                            source_lang = getattr(self.current_project, 'source_lang', None) if self.current_project else None
                            target_lang = getattr(self.current_project, 'target_lang', None) if self.current_project else None
                            
                            # Convert language names to codes if needed
                            if source_lang:
                                source_lang = self._convert_language_to_code(source_lang)
                            if target_lang:
                                target_lang = self._convert_language_to_code(target_lang)
                            
                            self.log(f"🟢 TM SEARCH: Project languages: {source_lang} → {target_lang}")
                            
                            # Search in primary direction (current project direction)
                            tm_matches = self.db_manager.search_all(segment.source, max_results=10)
                            
                            # Also search in reverse direction (bidirectional TM matching)
                            # This is crucial for reusing TM in opposite direction
                            reverse_matches = []
                            if source_lang and target_lang and source_lang != target_lang:
                                try:
                                    # Search for segment.source as target text in reverse direction
                                    reverse_query = """
                                        SELECT source_text, target_text, source_lang, target_lang, tm_id, usage_count
                                        FROM translation_units 
                                        WHERE target_text = ? AND source_lang = ? AND target_lang = ?
                                        ORDER BY usage_count DESC
                                        LIMIT 5
                                    """
                                    self.db_manager.cursor.execute(reverse_query, (segment.source, target_lang, source_lang))
                                    reverse_rows = self.db_manager.cursor.fetchall()
                                    
                                    for row in reverse_rows:
                                        reverse_matches.append({
                                            'source': row['target_text'],  # Swap source/target for reverse match
                                            'target': row['source_text'],
                                            'match_pct': 95,  # High relevance for reverse exact match
                                            'tm_name': f"{row['tm_id'].replace('_', ' ').title()} (Reverse)",
                                            'tm_id': row['tm_id']
                                        })
                                    
                                    self.log(f"🟢 TM SEARCH: Found {len(reverse_matches)} reverse matches")
                                except Exception as e:
                                    self.log(f"Error in reverse TM search: {e}")
                            
                            # Combine primary and reverse matches
                            all_tm_matches = tm_matches + reverse_matches
                            
                            self.log(f"🟢 TM SEARCH: Total matches found: {len(all_tm_matches)} ({len(tm_matches)} primary + {len(reverse_matches)} reverse)")
                            
                            for match in all_tm_matches:
                                match_obj = TranslationMatch(
                                    source=match.get('source', ''),
                                    target=match.get('target', ''),
                                    relevance=match.get('match_pct', 0),
                                    metadata={
                                        'context': match.get('context', ''),
                                        'tm_name': match.get('tm_name', ''),
                                        'timestamp': match.get('created_at', ''),
                                        'direction': 'reverse' if 'Reverse' in match.get('tm_name', '') else 'primary'
                                    },
                                    match_type='TM',
                                    compare_source=match.get('source', ''),
                                    provider_code='TM'
                                )
                                matches_dict["TM"].append(match_obj)
                        except Exception as e:
                            self.log(f"Error loading TM matches: {e}")
                    
                    # Add termbase matches to the results panel (if enabled)
                    if self.enable_termbase_matching:
                        try:
                            if hasattr(self, 'db_manager') and self.db_manager:
                                # Search termbases directly to get full information including termbase_id
                                source_lang = getattr(self.current_project, 'source_lang', None) if self.current_project else 'en'  # Default to English
                                target_lang = getattr(self.current_project, 'target_lang', None) if self.current_project else 'nl'  # Default to Dutch
                                
                                # Debug logging
                                self.log(f"🔍 Termbase matching: segment='{segment.source}', langs={source_lang}->{target_lang}, project={self.current_project is not None}")
                                self.log(f"🔍 TM matching enabled: {self.enable_tm_matching}, TM database: {hasattr(self, 'tm_database') and self.tm_database is not None}")
                                project_id_raw = getattr(self.current_project, 'id', None) if (self.current_project and hasattr(self.current_project, 'id')) else None
                                
                                # Convert project_id to string if needed (database stores as TEXT)
                                project_id = str(project_id_raw) if project_id_raw is not None else None
                                
                                # Search for all terms in the source text
                                if segment.source:
                                    # First try exact phrase match, then fall back to word-by-word
                                    termbase_results_by_term = {}
                                    
                                    # Try exact phrase match first (whole source text)
                                    try:
                                        self.log("🟢 TERMBASE CALL: From delayed lookup (_delayed_lookup)")
                                        self.log(f"🔍 Searching termbases for exact phrase: '{segment.source.strip()}'")
                                        exact_results = self.db_manager.search_termbases(
                                            segment.source.strip(),
                                            source_lang=source_lang,
                                            target_lang=target_lang,
                                            project_id=project_id,
                                            min_length=len(segment.source.strip())
                                        )
                                        self.log(f"🔍 Exact phrase results: {len(exact_results)} matches")
                                        for tb_match in exact_results:
                                            source_term = tb_match.get('source_term', '').strip()
                                            target_term = tb_match.get('target_term', '').strip()
                                            termbase_id_raw = tb_match.get('termbase_id')
                                            
                                            try:
                                                termbase_id = int(termbase_id_raw) if termbase_id_raw else None
                                            except (ValueError, TypeError):
                                                termbase_id = termbase_id_raw
                                            
                                            if source_term and target_term and termbase_id:
                                                key = (source_term, target_term, termbase_id)
                                                if key not in termbase_results_by_term:
                                                    termbase_results_by_term[key] = tb_match
                                    except Exception as exact_error:
                                        self.log(f"Error in exact phrase search: {exact_error}")
                                    
                                    # Then search word-by-word for partial matches
                                    words = segment.source.split()
                                    for word in words:
                                        clean_word = word.strip('.,!?;:')
                                        if len(clean_word) < 2:
                                            continue
                                        
                                        try:
                                            tb_results = self.db_manager.search_termbases(
                                                clean_word,
                                                source_lang=source_lang,
                                                target_lang=target_lang,
                                                project_id=project_id,
                                                min_length=2
                                            )
                                            
                                            for tb_match in tb_results:
                                                source_term = tb_match.get('source_term', '').strip()
                                                target_term = tb_match.get('target_term', '').strip()
                                                termbase_id_raw = tb_match.get('termbase_id')
                                                
                                                # Convert termbase_id to int if it's stored as string
                                                try:
                                                    termbase_id = int(termbase_id_raw) if termbase_id_raw else None
                                                except (ValueError, TypeError):
                                                    termbase_id = termbase_id_raw  # Keep as-is if conversion fails
                                                
                                                if source_term and target_term and termbase_id:
                                                    # Store with termbase_id as key to avoid duplicates
                                                    key = (source_term, target_term, termbase_id)
                                                    if key not in termbase_results_by_term:
                                                        termbase_results_by_term[key] = tb_match
                                        except Exception as search_error:
                                            self.log(f"Error searching termbases for '{clean_word}': {search_error}")
                                            continue
                                    
                                    # Get termbase names for IDs found
                                    termbase_code_map = self.get_termbase_code_map()
                                    
                                    # Create TranslationMatch objects
                                    for (source_term, target_term, termbase_id), tb_match in termbase_results_by_term.items():
                                        try:
                                            # Get termbase code (or generate from name)
                                            termbase_code = self.get_termbase_code(termbase_id, termbase_code_map)
                                            
                                            # Get termbase priority (termbase-level, not term-level)
                                            # We need to query the termbase table for this
                                            termbase_priority = 50  # Default
                                            try:
                                                if self.db_manager and self.db_manager.cursor:
                                                    cursor = self.db_manager.cursor
                                                    cursor.execute("SELECT priority FROM termbases WHERE id = ?", (termbase_id,))
                                                    row = cursor.fetchone()
                                                    if row and row[0] is not None:
                                                        termbase_priority = row[0]
                                            except Exception as prio_error:
                                                self.log(f"Warning: Could not fetch termbase priority for {termbase_id}: {prio_error}")
                                                # Continue with default priority
                                            
                                            match_obj = TranslationMatch(
                                                source=source_term,
                                                target=target_term,
                                                relevance=100 - tb_match.get('priority', 99),  # Lower priority = higher relevance
                                                metadata={
                                                    'termbase_id': termbase_id,
                                                    'termbase_name': tb_match.get('termbase_name', 'Unknown'),
                                                    'termbase_priority': termbase_priority,  # Termbase-level priority for color shading
                                                    'term_priority': tb_match.get('priority', 99),  # Term-level priority for sorting
                                                    'definition': tb_match.get('definition', ''),
                                                    'domain': tb_match.get('domain', '')
                                                },
                                                match_type='Termbase',
                                                compare_source=source_term,
                                                provider_code=termbase_code
                                            )
                                            matches_dict["Termbases"].append(match_obj)
                                        except Exception as match_error:
                                            import traceback
                                            self.log(f"❌ Error creating termbase match: {match_error}")
                                            self.log(f"Traceback: {traceback.format_exc()}")
                                            continue
                                    
                                    if termbase_results_by_term:
                                        num_created = len(matches_dict["Termbases"])
                                        self.log(f"✓ Found {len(termbase_results_by_term)} termbase matches, created {num_created} TranslationMatch objects")
                                    else:
                                        self.log(f"ℹ️ No termbase matches found for segment: '{segment.source[:50]}...'")
                                else:
                                    self.log("⚠️ No source text in segment for termbase search")
                        
                            self.log("🎯 CHECKPOINT: Termbase search completed, about to continue to set_matches...")
                        except Exception as e:
                            import traceback
                            self.log(f"❌ Error adding termbase matches: {e}")
                            self.log(f"Traceback: {traceback.format_exc()}")
                    
                    self.log("🎯 CHECKPOINT: About to fetch LLM and MT, then call set_matches...")
                    
                    # Fetch LLM translation for current segment (async, updates panel when ready)
                    self._fetch_llm_translation_async(segment.source, segment, current_row)
                    
                    # Fetch MT translation for current segment (async, updates panel when ready)
                    self._fetch_mt_translation_async(segment.source, segment, current_row)
                    
                    # Display matches (LLM and MT will be added when ready)
                    total_matches = sum(len(matches) for matches in matches_dict.values())
                    self.log(f"🔍 Calling set_matches() with {total_matches} total matches:")
                    for match_type, matches in matches_dict.items():
                        if matches:
                            self.log(f"  {match_type}: {len(matches)} matches")
                            if match_type == "Termbases" and matches:
                                for i, match in enumerate(matches[:3]):  # Show first 3 for debugging
                                    self.log(f"    [{i}] {match.source} → {match.target} (relevance: {match.relevance})")
                    
                    # Update all tabbed results panels
                    if hasattr(self, 'results_panels') and self.results_panels:
                        self.log(f"✅ Updating {len(self.results_panels)} results panels")
                        for panel in self.results_panels:
                            try:
                                panel.set_matches(matches_dict)
                            except Exception as e:
                                self.log(f"Error updating panel: {e}")
                        self.log(f"✅ All panels updated successfully")
                    else:
                        self.log(f"❌ results_panels missing or empty: {getattr(self, 'results_panels', 'NOT_SET')}")
                except Exception as e:
                    self.log(f"Error updating TranslationResultsPanel: {e}")
        except Exception as e:
            self.log(f"Error in delayed lookup: {e}")
    
    def toggle_tm_termbase_matching(self, enabled: bool):
        """Toggle TM and termbase matching on/off"""
        self.enable_tm_matching = enabled
        self.enable_termbase_matching = enabled
        
        # Update checkbox in settings if it exists (prevents circular updates)
        if hasattr(self, 'tm_matching_checkbox') and self.tm_matching_checkbox:
            # Temporarily disconnect to prevent signal loop
            self.tm_matching_checkbox.blockSignals(True)
            self.tm_matching_checkbox.setChecked(enabled)
            self.tm_matching_checkbox.blockSignals(False)
        
        if enabled:
            self.log("✓ TM and Termbase matching enabled")
            # If a segment is currently selected, trigger lookup
            if hasattr(self, 'table') and self.table and hasattr(self, 'current_project') and self.current_project:
                current_row = self.table.currentRow()
                if current_row >= 0 and current_row < len(self.current_project.segments):
                    segment = self.current_project.segments[current_row]
                    self._schedule_delayed_lookup(segment, current_row)
        else:
            self.log("⚠ TM and Termbase matching disabled")
            # Cancel any pending lookup
            if self.lookup_timer:
                self.lookup_timer.stop()
                self.lookup_timer = None
    
    
    def on_cell_clicked(self, item: QTableWidgetItem):
        """Handle cell click - Status column clicks are disabled (use Segment Editor instead)"""
        self.log(f"🖱️ on_cell_clicked called: row {item.row()}, col {item.column()}")
        if not self.current_project:
            return

        row = item.row()
        col = item.column()

        # Status column (4) - clicks disabled, use Segment Editor to change status
        # This prevents visual issues and unwanted status changes
        if col == 4:
            self.log(f"Status column click ignored - use Segment Editor to change status")
            return
    
    # ========================================================================
    # TRANSLATION MEMORY
    # ========================================================================
    
    def initialize_tm_database(self):
        """Initialize TM database when project is loaded"""
        if not self.current_project:
            return
        
        try:
            from modules.translation_memory import TMDatabase
            
            # Get database path
            db_path = self.user_data_path / "Translation_Resources" / "supervertaler.db"
            
            # Initialize TM database
            self.tm_database = TMDatabase(
                source_lang=self.current_project.source_lang,
                target_lang=self.current_project.target_lang,
                db_path=str(db_path),
                log_callback=self.log
            )
            
            # Update Universal Lookup tab with TM database
            if hasattr(self, 'lookup_tab') and self.lookup_tab:
                self.lookup_tab.set_tm_database(self.tm_database)
            
            self.log(f"TM database initialized ({self.current_project.source_lang} → {self.current_project.target_lang})")
        except Exception as e:
            self.log(f"Warning: Could not initialize TM database: {e}")
            self.tm_database = None
    
    def search_and_display_tm_matches(self, source_text: str):
        """Search TM and Termbases and display matches with visual diff for fuzzy matches"""
        self.log(f"🚨 search_and_display_tm_matches called with source_text: '{source_text[:50]}...'")
        if not source_text or not source_text.strip():
            if hasattr(self, 'tm_display'):
                self.tm_display.clear()
            return
        
        # Initialize TM if not already done
        if not self.tm_database and self.current_project:
            self.initialize_tm_database()
        
        if not self.tm_database:
            if hasattr(self, 'tm_display'):
                self.tm_display.setHtml(
                    "<p style='color: #999;'><i>Translation Memory not available</i></p>"
                )
            return
        
        try:
            # Search for matches
            matches = self.tm_database.search_all(source_text, max_matches=5)
            
            # Auto-propagate exact match if enabled (check both settings)
            general_prefs = self.load_general_settings()
            auto_propagate_100 = general_prefs.get('auto_propagate_100', True)
            auto_propagate_enabled = self.auto_propagate_exact_matches or auto_propagate_100
            
            if matches and auto_propagate_enabled:
                first_match = matches[0]
                match_pct = first_match.get('match_pct', 0)
                
                # Check if we have a 100% match
                if match_pct == 100:
                    # Check if current segment is empty/untranslated
                    if hasattr(self, 'table') and hasattr(self, 'current_project'):
                        current_row = self.table.currentRow()
                        if current_row >= 0 and current_row < len(self.current_project.segments):
                            segment = self.current_project.segments[current_row]
                            
                            # Only auto-propagate if segment is empty or untranslated
                            if (not segment.target or not segment.target.strip() or 
                                segment.status in ["untranslated", "not_started", ""]):
                                
                                # Fill the segment
                                segment.target = first_match.get('target', '')
                                segment.status = "translated"
                                segment.modified = True
                                self.project_modified = True
                                
                                # Update grid (using cell widget)
                                target_widget = self.table.cellWidget(current_row, 3)
                                if target_widget and isinstance(target_widget, EditableGridTextEditor):
                                    target_widget.setPlainText(segment.target)
                                
                                # Update status column
                                status_item = self.table.item(current_row, 1)
                                if status_item:
                                    status_item.setText("Translated")
                                    status_item.setBackground(QColor("#d1fae5"))  # Light green for translated
                                
                                # Update project modified status (progress is handled by status bar)
                                self.log(f"✨ Auto-propagated 100% TM match for segment #{current_row + 1}")
            
            if not matches:
                if hasattr(self, 'tm_display'):
                    self.tm_display.setHtml(
                        f"<p style='color: #666;'><b>Source:</b> {source_text}</p>"
                        f"<p style='color: #999;'><i>No translation memory matches found</i></p>"
                    )
                return
            
            # If using TranslationResultsPanel, populate it with TM and Termbase results
            if hasattr(self.assistance_widget, 'set_matches'):
                try:
                    from modules.translation_results_panel import TranslationMatch
                    
                    # Convert TM matches
                    tm_matches = []
                    for match in matches:
                        tm_match = TranslationMatch(
                            source=match.get('source', ''),
                            target=match.get('target', ''),
                            relevance=int(match.get('match_pct', 0)),
                            metadata={
                                'tm_name': match.get('tm_name', 'Unknown TM'),
                                'context': match.get('context', '')
                            },
                            match_type='TM',
                            compare_source=match.get('source', '')
                        )
                        tm_matches.append(tm_match)
                    
                    # Search for termbase matches
                    termbase_matches = []
                    try:
                        if hasattr(self, 'db_manager') and self.db_manager:
                            # Get current project language pair
                            source_lang = None
                            target_lang = None
                            project_id = None
                            
                            if self.current_project:
                                # current_project is a Project object, not a dict
                                project_id = self.current_project.id if hasattr(self.current_project, 'id') else None
                                source_lang = self.current_project.source_lang if hasattr(self.current_project, 'source_lang') else None
                                target_lang = self.current_project.target_lang if hasattr(self.current_project, 'target_lang') else None
                            
                            # Search for individual words in the source text
                            words = source_text.split()
                            searched_terms = set()  # Track what we've already searched
                            
                            for word in words:
                                # Remove punctuation
                                clean_word = word.strip('.,!?;:')
                                if len(clean_word) < 2 or clean_word in searched_terms:
                                    continue
                                
                                searched_terms.add(clean_word)
                                
                                # Search termbases for this word
                                tb_results = self.db_manager.search_termbases(
                                    clean_word,
                                    source_lang=source_lang,
                                    target_lang=target_lang,
                                    project_id=project_id,
                                    min_length=2
                                )
                                
                                # Convert termbase results to TranslationMatch objects
                                for tb_match in tb_results:
                                    termbase_match = TranslationMatch(
                                        source=tb_match.get('source_term', ''),
                                        target=tb_match.get('target_term', ''),
                                        relevance=100 - tb_match.get('priority', 99),  # Lower priority = higher relevance
                                        metadata={
                                            'termbase_id': tb_match.get('termbase_id'),
                                            'priority': tb_match.get('priority', 99),
                                            'definition': tb_match.get('definition', ''),
                                            'domain': tb_match.get('domain', ''),
                                            'forbidden': tb_match.get('forbidden', False)
                                        },
                                        match_type='Termbase',
                                        compare_source=tb_match.get('source_term', '')
                                    )
                                    termbase_matches.append(termbase_match)
                    except Exception as e:
                        self.log(f"Error searching termbases: {e}")

                    
                    # Update the panel with both TM and Termbase matches
                    matches_dict = {}
                    if tm_matches:
                        matches_dict['TM'] = tm_matches
                    if termbase_matches:
                        matches_dict['Termbases'] = termbase_matches
                    
                    if matches_dict and hasattr(self, 'results_panels'):
                        for panel in self.results_panels:
                            try:
                                panel.set_matches(matches_dict)
                            except Exception as e:
                                self.log(f"Error updating panel: {e}")
                    return
                except Exception as e:
                    self.log(f"Error populating TranslationResultsPanel: {e}")
            
            # Fallback: display in tm_display if it exists
            if not hasattr(self, 'tm_display'):
                return
            
            # Build HTML display with diff highlighting
            html = f"<div style='font-family: {self.default_font_family}; font-size: {self.default_font_size}pt;'>"
            html += f"<p style='color: #666; margin-bottom: 10px;'><b>🔍 Searching for:</b><br>{source_text}</p>"
            html += "<hr style='border: 1px solid #ddd;'>"
            
            for idx, match in enumerate(matches, start=1):
                match_pct = match.get('match_pct', 0)
                tm_name = match.get('tm_name', 'Unknown TM')
                source_match = match.get('source', '')
                target_match = match.get('target', '')
                
                # Color based on match percentage
                if match_pct == 100:
                    color = "#22c55e"  # Green for exact match
                    badge_style = "background-color: #22c55e; color: white;"
                elif match_pct >= 95:
                    color = "#3b82f6"  # Blue for high match
                    badge_style = "background-color: #3b82f6; color: white;"
                elif match_pct >= 85:
                    color = "#f59e0b"  # Orange for medium match
                    badge_style = "background-color: #f59e0b; color: white;"
                else:
                    color = "#ef4444"  # Red for low match
                    badge_style = "background-color: #ef4444; color: white;"
                
                html += f"<div style='margin: 10px 0; padding: 10px; background-color: #f9fafb; border-left: 4px solid {color}; border-radius: 4px;'>"
                html += f"<div style='margin-bottom: 5px;'>"
                html += f"<span style='padding: 2px 8px; border-radius: 3px; font-weight: bold; font-size: 10pt; {badge_style}'>{match_pct}%</span> "
                html += f"<span style='color: #666; font-size: 9pt;'>{tm_name}</span>"
                html += "</div>"
                
                # Display source with diff if not 100% match
                if match_pct < 100:
                    diff_html = self.create_diff_html(source_text, source_match)
                    html += f"<p style='margin: 5px 0; color: #444;'><b>TM Source:</b><br>{diff_html}</p>"
                else:
                    html += f"<p style='margin: 5px 0; color: #444;'><b>TM Source:</b><br>{source_match}</p>"
                
                # Display target (clickable to insert)
                html += f"<p style='margin: 5px 0; color: #059669; font-weight: bold;'><b>TM Target:</b><br>"
                html += f"<span style='cursor: pointer; text-decoration: underline;' "
                html += f"onclick='alert(\"{target_match}\")'>{target_match}</span></p>"
                html += "</div>"
            
            html += "</div>"
            self.tm_display.setHtml(html)
            
        except Exception as e:
            self.log(f"Error searching TM: {e}")
            if hasattr(self, 'tm_display'):
                self.tm_display.setHtml(
                    f"<p style='color: #ef4444;'><b>Error:</b> {e}</p>"
                )
    
    def create_diff_html(self, current_text: str, tm_text: str) -> str:
        """Create HTML with visual diff highlighting between current and TM source
        
        Shows changes inline like memoQ:
        - Red strikethrough for deleted text (in TM but not in current)
        - Green underline + bold for added text (in current but not in TM)
        - Sequential display: deletions then additions in order
        """
        from difflib import SequenceMatcher
        
        s = SequenceMatcher(None, tm_text, current_text)
        html_parts = []
        
        for tag, i1, i2, j1, j2 in s.get_opcodes():
            tm_chunk = tm_text[i1:i2]
            current_chunk = current_text[j1:j2]
            
            if tag == 'equal':
                # Matching text - show normally
                html_parts.append(tm_chunk)
            elif tag == 'delete':
                # Text removed from TM - show with red strikethrough (like memoQ deletion)
                html_parts.append(f'<span style="color: #dc2626; text-decoration: line-through;">{tm_chunk}</span>')
            elif tag == 'insert':
                # Text added in current - show with green underline + bold (like memoQ insertion)
                html_parts.append(f'<span style="color: #22c55e; text-decoration: underline; font-weight: bold;">{current_chunk}</span>')
            elif tag == 'replace':
                # Text changed - show deletion (strikethrough) followed by insertion (underline)
                # This matches memoQ's behavior: old text strikethrough, new text underlined
                html_parts.append(f'<span style="color: #dc2626; text-decoration: line-through;">{tm_chunk}</span>')
                html_parts.append(' ')  # Add space between deletion and insertion
                html_parts.append(f'<span style="color: #22c55e; text-decoration: underline; font-weight: bold;">{current_chunk}</span>')
        
        return ''.join(html_parts)
    
    # ========================================================================
    # TERMBASE HIGHLIGHTING & MATCHING
    # ========================================================================
    
    def get_termbase_code_map(self) -> Dict[int, str]:
        """Get mapping of termbase_id to custom codes (from user preferences)"""
        prefs_file = self.user_data_path / "ui_preferences.json"
        
        if not prefs_file.exists():
            return {}
        
        try:
            with open(prefs_file, 'r') as f:
                prefs = json.load(f)
                return prefs.get('termbase_code_map', {})  # {termbase_id: "CODE"}
        except:
            return {}
    
    def save_termbase_code_map(self, code_map: Dict[int, str]):
        """Save termbase code mappings to user preferences"""
        prefs_file = self.user_data_path / "ui_preferences.json"
        
        # Load existing preferences
        prefs = {}
        if prefs_file.exists():
            try:
                with open(prefs_file, 'r') as f:
                    prefs = json.load(f)
            except:
                pass
        
        # Update termbase code map
        prefs['termbase_code_map'] = code_map
        
        # Save back
        try:
            with open(prefs_file, 'w') as f:
                json.dump(prefs, f, indent=2)
        except Exception as e:
            self.log(f"⚠ Could not save termbase code map: {str(e)}")
    
    def get_termbase_code(self, termbase_id: int, code_map: Optional[Dict[int, str]] = None) -> str:
        """
        Get code for a termbase (custom code or auto-generated from name)
        
        Args:
            termbase_id: The termbase ID
            code_map: Optional pre-loaded code map (for performance)
            
        Returns:
            Code string (e.g., "MED", "LEG", "TECH") or empty string
        """
        if code_map is None:
            code_map = self.get_termbase_code_map()
        
        # Check if user has set a custom code
        if termbase_id in code_map:
            return code_map[termbase_id]
        
        # Auto-generate code from termbase name
        try:
            termbase_name = None
            
            if hasattr(self, 'termbase_mgr') and self.termbase_mgr:
                termbase = self.termbase_mgr.get_termbase(termbase_id)
                if termbase:
                    termbase_name = termbase.get('name', '') if isinstance(termbase, dict) else getattr(termbase, 'name', '')
            
            # Fallback: query database directly
            if not termbase_name and hasattr(self, 'db_manager') and self.db_manager:
                cursor = self.db_manager.cursor
                cursor.execute("SELECT name FROM termbases WHERE id = ?", (termbase_id,))
                row = cursor.fetchone()
                if row:
                    termbase_name = row[0]
            
            if not termbase_name:
                return ""
            
            # Generate code from name: take first 3-4 uppercase letters
            # Examples: "Medical Dictionary" -> "MED", "Legal Terms" -> "LEG", "Technical Glossary" -> "TECH"
            words = termbase_name.split()
            if len(words) >= 2:
                # Take first letter of first two words, then add more if needed
                code = (words[0][0] + words[1][0]).upper()
                if len(code) < 3 and len(words[0]) >= 2:
                    code = words[0][:3].upper()
            else:
                # Single word: take first 3 uppercase letters
                code = termbase_name[:3].upper()
            
            # Ensure code is 2-4 characters
            if len(code) > 4:
                code = code[:4]
            
            return code
        except Exception as e:
            self.log(f"Error generating termbase code: {e}")
            return ""
    
    def _convert_language_to_code(self, language: str) -> str:
        """Convert full language name to language code for termbase search"""
        language_map = {
            'Dutch': 'nl',
            'English': 'en', 
            'French': 'fr',
            'German': 'de',
            'Spanish': 'es',
            'Italian': 'it',
            'Portuguese': 'pt',
            'Russian': 'ru',
            'Chinese': 'zh',
            'Japanese': 'ja',
            'Korean': 'ko',
            'Arabic': 'ar'
        }
        return language_map.get(language, language.lower() if language else None)

    def find_termbase_matches_in_source(self, source_text: str) -> Dict[str, str]:
        """
        Find all termbase matches in source text
        Returns dict of {term: translation} for all matches found
        """
        if not source_text or not hasattr(self, 'db_manager') or not self.db_manager:
            return {}

        try:
            source_lang = self.current_project.source_lang if self.current_project else None
            target_lang = self.current_project.target_lang if self.current_project else None
            
            # Convert language names to codes for termbase search
            source_lang_code = self._convert_language_to_code(source_lang) if source_lang else None
            target_lang_code = self._convert_language_to_code(target_lang) if target_lang else None            # Debug logging
            self.log(f"🔍 Searching termbases for: '{source_text}' ({source_lang} → {target_lang})")
            if self.current_project:
                self.log(f"  ✓ Current project: {self.current_project.name}")
                self.log(f"  ✓ Project languages: {self.current_project.source_lang} → {self.current_project.target_lang}")
                self.log(f"  ✓ Language codes for search: {source_lang_code} → {target_lang_code}")
            else:
                self.log(f"  ❌ No current project loaded!")
            
            # First, check if we have any termbases at all
            query = "SELECT COUNT(*) FROM termbase_terms"
            self.db_manager.cursor.execute(query)
            total_terms = self.db_manager.cursor.fetchone()[0]
            self.log(f"  Total terms in database: {total_terms}")
            
            if total_terms == 0:
                self.log(f"  ⚠️  No termbases loaded in database")
                return {}
            
            # Search termbases for all terms that appear in the source text
            # Split source text into words and search for each one
            words = source_text.split()
            matches = {}
            self.log(f"  Processing {len(words)} words: {words[:10]}{'...' if len(words) > 10 else ''}")
            
            for word in words:
                # Remove punctuation and search
                clean_word = word.strip('.,!?;:')
                if len(clean_word) < 2:  # Skip short words
                    self.log(f"    Skipping short word: '{clean_word}'")
                    continue
                
                self.log(f"    Searching for: '{clean_word}'")
                
                termbase_results = self.db_manager.search_termbases(
                    clean_word,
                    source_lang=source_lang_code,
                    target_lang=target_lang_code,
                    min_length=2
                )
                
                if termbase_results:
                    self.log(f"  ✓ Found {len(termbase_results)} match(es) for '{clean_word}'")
                    for result in termbase_results:
                        source_term = result.get('source_term', '').strip()
                        target_term = result.get('target_term', '').strip()
                        if source_term and target_term:
                            matches[source_term] = target_term
                            self.log(f"    → {source_term} = {target_term}")
            
            self.log(f"🔍 Total unique matches: {len(matches)}")
            return matches
            
        except Exception as e:
            self.log(f"❌ Error finding termbase matches: {e}")
            import traceback
            self.log(f"Traceback: {traceback.format_exc()}")
            return {}
    
    def highlight_source_with_termbase(self, row: int, source_text: str):
        """
        Highlight termbase matches in the source column of the grid
        Creates a custom widget with clickable term highlights
        """
        if not self.table or row < 0:
            return
        
        try:
            # Find termbase matches
            self.log("🟡 TERMBASE CALL: From highlighting (highlight_termbase_matches)")
            termbase_matches = self.find_termbase_matches_in_source(source_text)
            
            if not termbase_matches:
                # No matches - display plain text
                return
            
            # Create highlighted widget
            widget = TermbaseHighlightWidget(source_text, termbase_matches, parent=self.table)
            
            # Store reference to insert translation on double-click
            def on_term_clicked(term, translation):
                self.insert_term_translation(row, translation)
            
            widget.term_double_clicked = on_term_clicked
            
            # Set the widget in the source column (column 2)
            self.table.setCellWidget(row, 2, widget)
            
        except Exception as e:
            self.log(f"Error highlighting termbase matches: {e}")
    
    def insert_term_translation(self, row: int, translation: str):
        """
        Double-click handler: insert termbase translation into target cell
        """
        if not self.table or row < 0 or row >= len(self.current_project.segments):
            return
        
        try:
            # Get the target cell
            target_item = self.table.item(row, 3)
            if not target_item:
                return
            
            # Insert translation (using cell widget)
            segment = self.current_project.segments[row]
            segment.target = translation
            target_widget = self.table.cellWidget(row, 3)
            if target_widget and isinstance(target_widget, EditableGridTextEditor):
                target_widget.setPlainText(translation)
            
            # Update status
            if segment.status in (DEFAULT_STATUS.key, 'pretranslated', 'rejected'):
                segment.status = 'translated'
                self.update_status_icon(row, 'translated')
            else:
                self._refresh_segment_status(segment)
            
            self.project_modified = True
            self.update_window_title()
            self.log(f"✓ Term translation inserted in segment {row + 1}")
            
        except Exception as e:
            self.log(f"Error inserting term translation: {e}")
    
    # ========================================================================
    # NAVIGATION & SEARCH
    # ========================================================================
    
    def show_find_replace_dialog(self):
        """Show unified Find & Replace dialog (Ctrl+F and Ctrl+H both open same dialog)
        Filters grid to show only matching segments, like filter boxes.
        """
        if not self.current_project or not self.current_project.segments:
            QMessageBox.information(self, "No Project", "Please open a project first.")
            return
        
        from PyQt6.QtWidgets import QCheckBox, QGroupBox
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Find and Replace")
        dialog.setMinimumWidth(700)
        dialog.setMinimumHeight(450)
        
        # Main horizontal layout - left side (options) and right side (buttons)
        main_layout = QHBoxLayout(dialog)
        
        # Left side - all options
        left_layout = QVBoxLayout()
        
        # Left side - all options
        left_layout = QVBoxLayout()
        
        # Find what
        find_layout = QHBoxLayout()
        find_layout.addWidget(QLabel("Find what:"))
        self.find_input = QLineEdit()
        find_layout.addWidget(self.find_input, stretch=1)
        left_layout.addLayout(find_layout)
        
        # Replace with
        replace_layout = QHBoxLayout()
        replace_layout.addWidget(QLabel("Replace with:"))
        self.replace_input = QLineEdit()
        replace_layout.addWidget(self.replace_input, stretch=1)
        left_layout.addLayout(replace_layout)
        
        # Search in options
        search_in_group = QGroupBox("Search in")
        search_in_layout = QVBoxLayout()
        
        self.search_source_cb = QCheckBox("Source text")
        self.search_target_cb = QCheckBox("Target text")
        self.search_target_cb.setChecked(True)  # Default to target
        
        search_in_layout.addWidget(self.search_source_cb)
        search_in_layout.addWidget(self.search_target_cb)
        search_in_group.setLayout(search_in_layout)
        left_layout.addWidget(search_in_group)
        
        # Match options
        match_group = QGroupBox("Match")
        match_layout = QVBoxLayout()
        
        self.match_group = QButtonGroup(dialog)
        match_anything = QRadioButton("Anything")
        match_anything.setChecked(True)
        match_whole_words = QRadioButton("Only whole words")
        match_entire = QRadioButton("Entire segment")
        
        self.match_group.addButton(match_anything, 0)
        self.match_group.addButton(match_whole_words, 1)
        self.match_group.addButton(match_entire, 2)
        
        match_layout.addWidget(match_anything)
        match_layout.addWidget(match_whole_words)
        match_layout.addWidget(match_entire)
        
        self.case_sensitive_cb = QCheckBox("Case sensitive")
        match_layout.addWidget(self.case_sensitive_cb)
        
        match_group.setLayout(match_layout)
        left_layout.addWidget(match_group)
        
        left_layout.addStretch()
        
        # Right side - buttons
        right_layout = QVBoxLayout()
        
        find_next_btn = QPushButton("Find next")
        find_next_btn.setMinimumWidth(150)
        find_next_btn.clicked.connect(lambda: self.find_next_match())
        right_layout.addWidget(find_next_btn)
        
        find_all_btn = QPushButton("Find all")
        find_all_btn.setMinimumWidth(150)
        find_all_btn.clicked.connect(lambda: self.find_all_matches())
        right_layout.addWidget(find_all_btn)
        
        # Replace buttons
        right_layout.addSpacing(10)
        
        replace_this_btn = QPushButton("Replace this")
        replace_this_btn.setMinimumWidth(150)
        replace_this_btn.clicked.connect(lambda: self.replace_current_match())
        right_layout.addWidget(replace_this_btn)
        
        replace_all_btn = QPushButton("Replace all")
        replace_all_btn.setMinimumWidth(150)
        replace_all_btn.clicked.connect(lambda: self.replace_all_matches())
        right_layout.addWidget(replace_all_btn)
        
        right_layout.addSpacing(10)
        
        highlight_all_btn = QPushButton("Highlight all")
        highlight_all_btn.setMinimumWidth(150)
        highlight_all_btn.clicked.connect(lambda: self.highlight_all_matches())
        right_layout.addWidget(highlight_all_btn)
        
        clear_highlight_btn = QPushButton("Clear highlighting")
        clear_highlight_btn.setMinimumWidth(150)
        clear_highlight_btn.clicked.connect(lambda: self.clear_search_highlights())
        right_layout.addWidget(clear_highlight_btn)
        
        right_layout.addStretch()
        
        close_btn = QPushButton("Close")
        close_btn.setMinimumWidth(150)
        close_btn.clicked.connect(dialog.close)
        right_layout.addWidget(close_btn)
        
        # Add left and right to main layout
        main_layout.addLayout(left_layout, stretch=2)
        main_layout.addLayout(right_layout, stretch=1)
        
        # Store dialog reference and show
        self.find_replace_dialog = dialog
        self.current_match_index = -1
        self.find_matches = []
        
        dialog.show()
    
    def find_next_match(self):
        """Find next occurrence of search term, filtering grid to show only matching rows"""
        find_text = self.find_input.text()
        if not find_text:
            return
        
        search_source = self.search_source_cb.isChecked()
        search_target = self.search_target_cb.isChecked()
        case_sensitive = self.case_sensitive_cb.isChecked()
        match_mode = self.match_group.checkedId()
        
        # Find all matches if not already done
        if not self.find_matches:
            self.find_all_matches_internal(find_text, search_source, search_target, case_sensitive, match_mode)
            
            if self.find_matches:
                # First search - reload grid and filter rows
                self.load_segments_to_grid()
                
                # Get unique rows that have matches
                matching_rows = set(row for row, col in self.find_matches)
                
                # Hide all non-matching rows
                for row in range(len(self.current_project.segments)):
                    self.table.setRowHidden(row, row not in matching_rows)
        
        if not self.find_matches:
            QMessageBox.information(self.find_replace_dialog, "Find", "No matches found.")
            return
        
        # Move to next match
        self.current_match_index = (self.current_match_index + 1) % len(self.find_matches)
        row, col = self.find_matches[self.current_match_index]
        
        # Highlight the search term in the matched cell
        segment = self.current_project.segments[row]
        text = segment.source if col == 2 else segment.target
        self.highlight_search_term(row, col, text, find_text)
        
        self.table.setCurrentCell(row, col)
        self.table.scrollToItem(self.table.item(row, col))
        self.log(f"Match {self.current_match_index + 1} of {len(self.find_matches)}")
    
    def find_all_matches(self):
        """Find and highlight all matches, filtering grid to show only matching rows"""
        find_text = self.find_input.text()
        if not find_text:
            return
        
        search_source = self.search_source_cb.isChecked()
        search_target = self.search_target_cb.isChecked()
        case_sensitive = self.case_sensitive_cb.isChecked()
        match_mode = self.match_group.checkedId()
        
        self.find_all_matches_internal(find_text, search_source, search_target, case_sensitive, match_mode)
        
        if self.find_matches:
            # First, reload grid to clear previous highlights
            self.load_segments_to_grid()
            
            # Get unique rows that have matches
            matching_rows = set(row for row, col in self.find_matches)
            
            # Hide all non-matching rows (like filter boxes do)
            for row in range(len(self.current_project.segments)):
                self.table.setRowHidden(row, row not in matching_rows)
            
            # Highlight all matches with yellow
            for row, col in self.find_matches:
                segment = self.current_project.segments[row]
                text = segment.source if col == 2 else segment.target
                self.highlight_search_term(row, col, text, find_text)
            
            # Jump to first match
            first_row, first_col = self.find_matches[0]
            self.table.setCurrentCell(first_row, first_col)
            self.table.scrollToItem(self.table.item(first_row, first_col))
            self.current_match_index = 0
            
            self.log(f"Found {len(self.find_matches)} match(es) in {len(matching_rows)} segment(s)")
        else:
            QMessageBox.information(self.find_replace_dialog, "Find", "No matches found.")
    
    def find_all_matches_internal(self, find_text, search_source, search_target, case_sensitive, match_mode):
        """Internal method to find all matches"""
        import re
        
        self.find_matches = []
        
        for row, segment in enumerate(self.current_project.segments):
            texts_to_search = []
            if search_source:
                texts_to_search.append((segment.source, 2))
            if search_target:
                texts_to_search.append((segment.target, 3))
            
            for text, col in texts_to_search:
                if self.text_matches(text, find_text, case_sensitive, match_mode):
                    self.find_matches.append((row, col))
    
    def text_matches(self, text, find_text, case_sensitive, match_mode):
        """Check if text matches search criteria"""
        import re
        
        if match_mode == 2:  # Entire segment
            if case_sensitive:
                return text == find_text
            else:
                return text.lower() == find_text.lower()
        
        elif match_mode == 1:  # Whole words
            if case_sensitive:
                pattern = r'\b' + re.escape(find_text) + r'\b'
                return bool(re.search(pattern, text))
            else:
                pattern = r'\b' + re.escape(find_text) + r'\b'
                return bool(re.search(pattern, text, re.IGNORECASE))
        
        else:  # Anything (0)
            if case_sensitive:
                return find_text in text
            else:
                return find_text.lower() in text.lower()
    
    def replace_current_match(self):
        """Replace the currently selected match"""
        if self.current_match_index < 0 or self.current_match_index >= len(self.find_matches):
            QMessageBox.information(self.find_replace_dialog, "Replace", "No match selected. Use 'Find next' first.")
            return
        
        find_text = self.find_input.text()
        replace_text = self.replace_input.text()
        row, col = self.find_matches[self.current_match_index]
        
        # Safety check: prevent replacing in source if disabled
        if col == 2 and not self.allow_replace_in_source:  # Source column
            QMessageBox.warning(
                self.find_replace_dialog,
                "Replace in Source Disabled",
                "Replacing in source text is disabled for safety.\n\n"
                "To enable this feature, go to Settings > General and check 'Allow Replace in Source Text'."
            )
            return
        
        segment = self.current_project.segments[row]
        case_sensitive = self.case_sensitive_cb.isChecked()
        match_mode = self.match_group.checkedId()
        
        # Determine which field to update
        if col == 2:  # Source column
            field_text = segment.source
        else:  # col == 3, Target column
            field_text = segment.target
        
        if match_mode == 2:  # Entire segment
            new_text = replace_text
        else:
            # Replace using the appropriate method
            if case_sensitive:
                new_text = field_text.replace(find_text, replace_text, 1)
            else:
                import re
                pattern = re.escape(find_text)
                new_text = re.sub(pattern, replace_text, field_text, count=1, flags=re.IGNORECASE)
        
        # Update the appropriate field
        if col == 2:
            segment.source = new_text
        else:
            segment.target = new_text
        
        # Update table
        item = self.table.item(row, col)
        if item:
            item.setText(segment.target)
        
        self.project_modified = True
        self.update_window_title()
        
        # Refresh matches
        self.find_matches = []
        self.find_next_match()
    
    def replace_all_matches(self):
        """Replace all matches in target segments"""
        find_text = self.find_input.text()
        replace_text = self.replace_input.text()
        
        if not find_text:
            return
        
        search_source = self.search_source_cb.isChecked()
        search_target = self.search_target_cb.isChecked()
        case_sensitive = self.case_sensitive_cb.isChecked()
        match_mode = self.match_group.checkedId()
        
        # Safety check: warn if trying to replace in source when disabled
        if search_source and not self.allow_replace_in_source:
            reply = QMessageBox.warning(
                self.find_replace_dialog,
                "Replace in Source Disabled",
                "Replacing in source text is disabled for safety.\n\n"
                "Only target text will be replaced.\n\n"
                "To enable replacing in source, go to Settings > General and check 'Allow Replace in Source Text'.",
                QMessageBox.StandardButton.Ok | QMessageBox.StandardButton.Cancel
            )
            if reply == QMessageBox.StandardButton.Cancel:
                return
            # Continue with target only
            search_source = False
            if not search_target:
                return
        
        # Find all matches
        self.find_all_matches_internal(find_text, search_source, search_target, case_sensitive, match_mode)
        
        if not self.find_matches:
            QMessageBox.information(self.find_replace_dialog, "Replace All", "No matches found.")
            return
        
        # Count matches by column
        source_count = sum(1 for _, col in self.find_matches if col == 2)
        target_count = sum(1 for _, col in self.find_matches if col == 3)
        
        # Filter out source matches if not allowed
        if not self.allow_replace_in_source:
            self.find_matches = [(row, col) for row, col in self.find_matches if col == 3]
            if not self.find_matches:
                QMessageBox.information(self.find_replace_dialog, "Replace All", "No matches found in Target column.")
                return
        
        # Build confirmation message
        msg_parts = []
        if source_count > 0 and self.allow_replace_in_source:
            msg_parts.append(f"{source_count} in source")
        if target_count > 0:
            msg_parts.append(f"{target_count} in target")
        
        confirmation_msg = f"Replace {len(self.find_matches)} occurrence(s) of '{find_text}' with '{replace_text}'?\n\n"
        if msg_parts:
            confirmation_msg += f"({', '.join(msg_parts)})"
        
        # Extra warning if replacing in source
        if source_count > 0 and self.allow_replace_in_source:
            confirmation_msg += "\n\n⚠️ WARNING: This will modify source text!"
        
        # Confirm
        reply = QMessageBox.question(
            self.find_replace_dialog,
            "Replace All",
            confirmation_msg,
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply != QMessageBox.StandardButton.Yes:
            return
        
        # Perform replacements
        import re
        replaced_count = 0
        
        for row, col in self.find_matches:
            segment = self.current_project.segments[row]
            
            # Get the appropriate field
            if col == 2:  # Source
                old_text = segment.source
            else:  # col == 3, Target
                old_text = segment.target
            
            # Perform replacement
            if match_mode == 2:  # Entire segment
                new_text = replace_text
            else:
                if case_sensitive:
                    new_text = old_text.replace(find_text, replace_text)
                else:
                    pattern = re.escape(find_text)
                    new_text = re.sub(pattern, replace_text, old_text, flags=re.IGNORECASE)
            
            if new_text != old_text:
                replaced_count += 1
                # Update the appropriate field
                if col == 2:
                    segment.source = new_text
                else:
                    segment.target = new_text
                
                # Update table
                item = self.table.item(row, col)
                if item:
                    item.setText(new_text)
        
        self.project_modified = True
        self.update_window_title()
        
        # Clear matches and reload
        self.find_matches = []
        self.load_segments_to_grid()
        
        QMessageBox.information(self.find_replace_dialog, "Replace All", f"Replaced {replaced_count} occurrence(s).")
        self.log(f"✓ Replaced {replaced_count} occurrence(s) of '{find_text}'")
    
    def highlight_all_matches(self):
        """Highlight all matches using the filter system"""
        find_text = self.find_input.text()
        if not find_text:
            return
        
        search_source = self.search_source_cb.isChecked()
        search_target = self.search_target_cb.isChecked()
        
        # Use filter boxes to highlight
        if search_source:
            self.source_filter.setText(find_text)
        else:
            self.source_filter.clear()
        
        if search_target:
            self.target_filter.setText(find_text)
        else:
            self.target_filter.clear()
    
    def clear_search_highlights(self):
        """Clear all search highlights and unhide all rows (for Find & Replace dialog)"""
        self.load_segments_to_grid()
        
        # Unhide all rows
        for row in range(self.table.rowCount()):
            self.table.setRowHidden(row, False)
        
        self.find_matches = []
        self.current_match_index = -1
        self.log("Search highlights cleared")
    
    def show_goto_dialog(self):
        """Show dialog to jump to a specific segment"""
        if not self.current_project or not self.current_project.segments:
            QMessageBox.information(self, "No Project", "Please open a project first.")
            return
        
        max_segment = len(self.current_project.segments)
        segment_num, ok = QInputDialog.getInt(
            self,
            "Go to Segment",
            f"Enter segment number (1-{max_segment}):",
            value=self.table.currentRow() + 1,
            min=1,
            max=max_segment
        )
        
        if ok:
            row = segment_num - 1
            self.table.setCurrentCell(row, 3)  # Jump to Target column
            self.table.scrollToItem(self.table.item(row, 3))
            self.log(f"Jumped to segment {segment_num}")
    
    def show_search_dialog(self):
        """Show enhanced search dialog with source/target options"""
        if not self.current_project or not self.current_project.segments:
            QMessageBox.information(self, "No Project", "Please open a project first.")
            return
        
        from PyQt6.QtWidgets import QDialog, QLineEdit, QRadioButton, QButtonGroup, QDialogButtonBox
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Find in Segments")
        dialog.setMinimumWidth(400)
        
        layout = QVBoxLayout(dialog)
        
        # Search text input
        search_layout = QHBoxLayout()
        search_layout.addWidget(QLabel("Search for:"))
        search_input = QLineEdit()
        search_layout.addWidget(search_input)
        layout.addLayout(search_layout)
        
        # Search scope options
        scope_label = QLabel("Search in:")
        layout.addWidget(scope_label)
        
        radio_layout = QHBoxLayout()
        search_scope_group = QButtonGroup(dialog)
        
        both_radio = QRadioButton("Both Source and Target")
        both_radio.setChecked(True)
        source_radio = QRadioButton("Source only")
        target_radio = QRadioButton("Target only")
        
        search_scope_group.addButton(both_radio, 0)
        search_scope_group.addButton(source_radio, 1)
        search_scope_group.addButton(target_radio, 2)
        
        radio_layout.addWidget(both_radio)
        radio_layout.addWidget(source_radio)
        radio_layout.addWidget(target_radio)
        layout.addLayout(radio_layout)
        
        # Dialog buttons
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)
        layout.addWidget(button_box)
        
        # Show dialog and process
        if dialog.exec() == QDialog.DialogCode.Accepted:
            search_text = search_input.text()
            if search_text:
                scope = search_scope_group.checkedId()
                self.search_segments(search_text, scope)
    
    def search_segments(self, search_text: str, scope: int = 0):
        """Search for text in segments and highlight ONLY the exact search term
        
        Args:
            search_text: Text to search for
            scope: 0=both, 1=source only, 2=target only
        """
        if not self.current_project:
            return
        
        # Populate filter boxes based on scope
        if scope == 0:  # Both
            self.source_filter.setText(search_text)
            self.target_filter.setText(search_text)
        elif scope == 1:  # Source only
            self.source_filter.setText(search_text)
            self.target_filter.clear()
        elif scope == 2:  # Target only
            self.source_filter.clear()
            self.target_filter.setText(search_text)
        
        # The filter boxes will trigger apply_filters() automatically
        # which handles highlighting and hiding
        
        # Find first match and jump to it
        search_lower = search_text.lower()
        for row, segment in enumerate(self.current_project.segments):
            source_match = scope != 2 and search_lower in segment.source.lower()
            target_match = scope != 1 and search_lower in segment.target.lower()
            
            if source_match or target_match:
                self.table.setCurrentCell(row, 3)
                self.table.scrollToItem(self.table.item(row, 3))
                break
    
    def highlight_search_term(self, row: int, col: int, text: str, search_term: str):
        """Highlight only the search term within the cell text using HTML"""
        # Create a custom widget with HTML formatting
        label = QLabel()
        label.setWordWrap(True)
        label.setTextFormat(Qt.TextFormat.RichText)
        
        # Build HTML with highlighted search term (case-insensitive)
        search_lower = search_term.lower()
        text_lower = text.lower()
        
        # Find all occurrences and build highlighted HTML
        html_text = ""
        last_pos = 0
        
        while True:
            pos = text_lower.find(search_lower, last_pos)
            if pos == -1:
                # Add remaining text
                html_text += text[last_pos:]
                break
            
            # Add text before match
            html_text += text[last_pos:pos]
            # Add highlighted match
            html_text += f'<span style="background-color: yellow;">{text[pos:pos+len(search_term)]}</span>'
            last_pos = pos + len(search_term)
        
        label.setText(html_text)
        self.table.setCellWidget(row, col, label)
    
    def apply_filters(self):
        """Apply source and target filters to show/hide rows and highlight matches"""
        if not self.current_project:
            return
        
        # Safety check: ensure table and filter widgets exist
        if not hasattr(self, 'table') or self.table is None:
            return
        
        if not hasattr(self, 'source_filter') or not hasattr(self, 'target_filter'):
            return
        
        source_filter_text = self.source_filter.text().strip()
        target_filter_text = self.target_filter.text().strip()
        
        # If both empty, clear everything
        if not source_filter_text and not target_filter_text:
            self.clear_filters()
            return
        
        # Clear previous highlights by reloading
        self.load_segments_to_grid()
        
        # Safety check: ensure table has correct number of rows after reload
        if self.table.rowCount() != len(self.current_project.segments):
            self.log("⚠ Warning: Table row count mismatch after reload")
            return
        
        visible_count = 0
        
        for row, segment in enumerate(self.current_project.segments):
            # Safety check: ensure row is valid
            if row >= self.table.rowCount():
                break
                
            source_match = not source_filter_text or source_filter_text.lower() in segment.source.lower()
            target_match = not target_filter_text or target_filter_text.lower() in segment.target.lower()
            
            show_row = source_match and target_match
            
            # Hide/show row
            self.table.setRowHidden(row, not show_row)
            
            if show_row:
                visible_count += 1
                
                # Highlight matching terms
                if source_filter_text and source_filter_text.lower() in segment.source.lower():
                    self.highlight_search_term(row, 2, segment.source, source_filter_text)
                
                if target_filter_text and target_filter_text.lower() in segment.target.lower():
                    self.highlight_search_term(row, 3, segment.target, target_filter_text)
        
        # Update status
        if source_filter_text or target_filter_text:
            self.log(f"Filter applied: showing {visible_count} of {len(self.current_project.segments)} segments")
    
    def clear_filters(self):
        """Clear all filter boxes, highlighting, and show all rows"""
        # Safety check: ensure filter widgets exist
        if not hasattr(self, 'source_filter') or not hasattr(self, 'target_filter'):
            return
        
        # Block signals to prevent recursion
        self.source_filter.blockSignals(True)
        self.target_filter.blockSignals(True)
        
        self.source_filter.clear()
        self.target_filter.clear()
        
        # Re-enable signals
        self.source_filter.blockSignals(False)
        self.target_filter.blockSignals(False)
        
        # Reload grid to remove all highlighting
        if self.current_project:
            # Safety check: ensure table exists
            if not hasattr(self, 'table') or self.table is None:
                return
                
            self.load_segments_to_grid()
            
            # Explicitly show all rows (unhide them)
            for row in range(self.table.rowCount()):
                self.table.setRowHidden(row, False)
        
        self.log("Filters cleared")
    
    # ========================================================================
    # LIST VIEW METHODS
    # ========================================================================
    
    def apply_list_filters(self):
        """Apply filters to List View"""
        self.refresh_list_view()
    
    def clear_list_filters(self):
        """Clear List View filters"""
        if hasattr(self, 'list_source_filter'):
            self.list_source_filter.clear()
        if hasattr(self, 'list_target_filter'):
            self.list_target_filter.clear()
        self.refresh_list_view()
    
    def on_list_segment_selected(self):
        """Handle segment selection in List View"""
        selected_items = self.list_tree.selectedItems()
        if not selected_items:
            self.list_current_segment_id = None
            # Update tab segment editor to show nothing selected
            if hasattr(self, 'update_tab_segment_editor'):
                self.update_tab_segment_editor(
                    segment_id=None,
                    source_text="",
                    target_text="",
                    status="untranslated",
                    notes=""
                )
            return
        
        item = selected_items[0]
        segment_id = item.data(0, Qt.ItemDataRole.UserRole)
        self.list_current_segment_id = segment_id
        
        # Find segment
        segment = next((s for s in self.current_project.segments if s.id == segment_id), None)
        if not segment:
            return
        
        # Update tab segment editor with new tabbed panel approach
        if hasattr(self, 'update_tab_segment_editor'):
            self.update_tab_segment_editor(
                segment_id=segment.id,
                source_text=segment.source,
                target_text=segment.target,
                status=segment.status,
                notes=segment.notes if hasattr(segment, 'notes') else ""
            )
        
        # Trigger assistance panel update (same as grid view)
        # Find row index for on_cell_selected
        row = next((i for i, s in enumerate(self.current_project.segments) if s.id == segment_id), -1)
        if row >= 0:
            # Call on_cell_selected to update translation results
            self.on_cell_selected(row, 3, -1, -1)
    
    def on_list_status_change(self, status: str):
        """Handle status change in List View (deprecated - now handled by tab editor)"""
        # This is kept for backwards compatibility but actual editing is done in tab panel
        pass
    
    def on_list_target_change(self):
        """Handle target text change in List View (deprecated - now handled by tab editor)"""
        # This is kept for backwards compatibility but actual editing is done in tab panel
        pass
    
    def copy_source_to_list_target(self):
        """Copy source to target in List View (deprecated - now handled by tab editor)"""
        # This is kept for backwards compatibility but actual editing is done in tab panel
        pass
    
    def clear_list_target(self):
        """Clear target in List View (deprecated - now handled by tab editor)"""
        # This is kept for backwards compatibility but actual editing is done in tab panel
        pass
    
    def focus_list_target_editor(self):
        """Focus the target editor in List View"""
        # Focus the tab segment editor instead
        if hasattr(self, 'list_tabbed_panel'):
            # Switch to Segment Editor tab
            self.list_tabbed_panel.setCurrentIndex(1)  # Index 1 is Segment Editor
    
    def save_list_segment_and_next(self):
        """Save current segment and move to next in List View"""
        if not self.list_current_segment_id:
            return
        
        # Find current item
        current_item = self.list_tree.currentItem()
        if not current_item:
            return
        
        # Find next item
        current_index = self.list_tree.indexOfTopLevelItem(current_item)
        next_index = current_index + 1
        
        if next_index < self.list_tree.topLevelItemCount():
            next_item = self.list_tree.topLevelItem(next_index)
            self.list_tree.setCurrentItem(next_item)
            self.list_tree.scrollToItem(next_item)
        else:
            self.log("✓ Last segment - no next segment")
    
    # ========================================================================
    # TABBED SEGMENT EDITOR METHODS (for Grid/List views)
    # ========================================================================
    
    def on_tab_target_change(self):
        """Handle target text change in tab editor - updates all panels"""
        if not hasattr(self, 'tab_current_segment_id') or not self.tab_current_segment_id:
            return
        
        # Get the new text from whichever panel triggered this
        new_text = None
        if hasattr(self, 'tabbed_panels'):
            for panel in self.tabbed_panels:
                try:
                    if hasattr(panel, 'editor_widget') and hasattr(panel.editor_widget, 'target_editor'):
                        new_text = panel.editor_widget.target_editor.toPlainText()
                        break  # Use text from first panel found
                except:
                    continue
        
        if new_text is None:
            return
        
        # Auto-save to current_project data structure
        if self.current_project and hasattr(self.current_project, 'segments'):
            for seg in self.current_project.segments:
                if seg.id == self.tab_current_segment_id:
                    seg.target = new_text
                    self.project_modified = True
                    
                    # Update all other panels to keep them in sync
                    if hasattr(self, 'tabbed_panels'):
                        for panel in self.tabbed_panels:
                            try:
                                if hasattr(panel, 'editor_widget') and hasattr(panel.editor_widget, 'target_editor'):
                                    # Temporarily disconnect to avoid infinite loop
                                    panel.editor_widget.target_editor.blockSignals(True)
                                    panel.editor_widget.target_editor.setPlainText(new_text)
                                    panel.editor_widget.target_editor.blockSignals(False)
                            except:
                                pass
                    break
    
    def on_tab_status_combo_changed(self, index: int):
        combo = self.sender()
        if isinstance(combo, QComboBox) and index >= 0:
            status_key = combo.itemData(index)
            if status_key:
                self.on_tab_segment_status_change(status_key)

    def on_tab_segment_status_change(self, status_key: str):
        """Handle status change in tab editor - updates all panels"""
        if not hasattr(self, 'tab_current_segment_id') or not self.tab_current_segment_id:
            return

        status_key = status_key or DEFAULT_STATUS.key
        status_def = get_status(status_key)

        if self.current_project and hasattr(self.current_project, 'segments'):
            for seg in self.current_project.segments:
                if seg.id == self.tab_current_segment_id:
                    seg.status = status_key
                    self.project_modified = True
                    self.log(f"✓ Status changed to: {status_def.label}")

                    # Update all panels to keep them in sync
                    if hasattr(self, 'tabbed_panels'):
                        for panel in self.tabbed_panels:
                            try:
                                if hasattr(panel, 'editor_widget') and hasattr(panel.editor_widget, 'status_combo'):
                                    combo = panel.editor_widget.status_combo
                                    combo.blockSignals(True)
                                    idx = combo.findData(status_key)
                                    if idx >= 0:
                                        combo.setCurrentIndex(idx)
                                    combo.blockSignals(False)
                            except Exception as e:
                                self.log(f"Error syncing status combo: {e}")
                    self._refresh_segment_status(seg)
                    break
    
    def on_tab_notes_change(self):
        """Handle comments change in tab editor - updates all panels"""
        if not hasattr(self, 'tab_current_segment_id') or not self.tab_current_segment_id:
            return
        
        # Get the new comments from whichever panel triggered this
        new_notes = None
        if hasattr(self, 'tabbed_panels'):
            for panel in self.tabbed_panels:
                try:
                    if hasattr(panel, 'notes_widget') and hasattr(panel.notes_widget, 'notes_editor'):
                        new_notes = panel.notes_widget.notes_editor.toPlainText()
                        break
                except:
                    continue
        
        if new_notes is None:
            return
        
        if self.current_project and hasattr(self.current_project, 'segments'):
            for seg in self.current_project.segments:
                if seg.id == self.tab_current_segment_id:
                    seg.notes = new_notes
                    self.project_modified = True
                    
                    # Update all other panels to keep them in sync
                    if hasattr(self, 'tabbed_panels'):
                        for panel in self.tabbed_panels:
                            try:
                                if hasattr(panel, 'notes_widget') and hasattr(panel.notes_widget, 'notes_editor'):
                                    panel.notes_widget.notes_editor.blockSignals(True)
                                    panel.notes_widget.notes_editor.setPlainText(new_notes)
                                    panel.notes_widget.notes_editor.blockSignals(False)
                            except:
                                pass
                    self._refresh_segment_status(seg)
                    break
    
    def copy_source_to_tab_target(self):
        """Copy source to target in tab editor - works with all panels"""
        if hasattr(self, 'tabbed_panels'):
            for panel in self.tabbed_panels:
                try:
                    if hasattr(panel, 'editor_widget'):
                        source_text = panel.editor_widget.source_editor.toPlainText()
                        panel.editor_widget.target_editor.setPlainText(source_text)
                except:
                    pass
    
    def clear_tab_target(self):
        """Clear target in tab editor - works with all panels"""
        if hasattr(self, 'tabbed_panels'):
            for panel in self.tabbed_panels:
                try:
                    if hasattr(panel, 'editor_widget'):
                        panel.editor_widget.target_editor.clear()
                except:
                    pass

    def start_voice_dictation(self):
        """Start voice dictation into target field"""
        if hasattr(self, 'dictation_thread') and self.dictation_thread and self.dictation_thread.isRunning():
            # Already recording
            return

        try:
            # Load dictation settings
            dictation_settings = self.load_dictation_settings()
            model_name = dictation_settings.get('model', 'base')
            max_duration = dictation_settings.get('max_duration', 10)
            lang_setting = dictation_settings.get('language', 'Auto (use project target language)')

            # Determine language
            if lang_setting == 'Auto (use project target language)':
                # Use project's target language
                target_lang = getattr(self, 'target_language', 'English')
            else:
                # Use override language from settings
                target_lang = lang_setting

            # Map language names to Whisper codes
            lang_map = {
                'English': 'en',
                'Dutch': 'nl',
                'German': 'de',
                'French': 'fr',
                'Spanish': 'es',
                'Italian': 'it',
                'Portuguese': 'pt',
                'Polish': 'pl',
                'Russian': 'ru',
                'Chinese': 'zh',
                'Japanese': 'ja',
                'Korean': 'ko'
            }
            lang_code = lang_map.get(target_lang, 'auto')

            # Create dictation thread with user settings
            self.dictation_thread = QuickDictationThread(
                model_name=model_name,
                language=lang_code,
                duration=max_duration
            )

            # Connect signals
            self.dictation_thread.transcription_ready.connect(self.on_dictation_complete)
            self.dictation_thread.status_update.connect(self.on_dictation_status)
            self.dictation_thread.error_occurred.connect(self.on_dictation_error)
            self.dictation_thread.finished.connect(self.on_dictation_finished)

            # Change button appearance
            self._set_dictation_button_recording(True)

            # Start recording
            self.dictation_thread.start()

        except Exception as e:
            self._set_dictation_button_recording(False)
            QMessageBox.critical(self, "Error", f"Failed to start dictation:\n{str(e)}")

    def on_dictation_complete(self, text):
        """Handle completed dictation"""
        # Try to insert into currently focused target field (grid or editor)
        focused_widget = QApplication.focusWidget()

        # Check if focused widget is a grid target cell
        if isinstance(focused_widget, EditableGridTextEditor):
            current_text = focused_widget.toPlainText()
            if current_text:
                focused_widget.setPlainText(current_text + " " + text)
            else:
                focused_widget.setPlainText(text)
            # Move cursor to end
            cursor = focused_widget.textCursor()
            cursor.movePosition(cursor.MoveOperation.End)
            focused_widget.setTextCursor(cursor)
        # Otherwise insert into segment editor below grid
        elif hasattr(self, 'tabbed_panels'):
            for panel in self.tabbed_panels:
                try:
                    if hasattr(panel, 'editor_widget'):
                        current_text = panel.editor_widget.target_editor.toPlainText()
                        if current_text:
                            panel.editor_widget.target_editor.setPlainText(current_text + " " + text)
                        else:
                            panel.editor_widget.target_editor.setPlainText(text)
                        cursor = panel.editor_widget.target_editor.textCursor()
                        cursor.movePosition(cursor.MoveOperation.End)
                        panel.editor_widget.target_editor.setTextCursor(cursor)
                except:
                    pass

        # Show notification
        self.status_bar.showMessage(f"✅ Dictation: {text[:50]}...", 3000)

    def on_dictation_status(self, message):
        """Show dictation status"""
        self.status_bar.showMessage(message, 2000)

    def on_dictation_error(self, error_msg):
        """Handle dictation error"""
        self._set_dictation_button_recording(False)
        self.status_bar.showMessage(f"❌ Voice dictation error", 3000)

        # Show detailed error dialog for FFmpeg issues
        if "FFmpeg" in error_msg or "ffmpeg" in error_msg:
            QMessageBox.warning(self, "FFmpeg Required", error_msg)

    def on_dictation_finished(self):
        """Handle dictation thread finishing"""
        self._set_dictation_button_recording(False)

    def _set_dictation_button_recording(self, is_recording):
        """Change dictate button appearance based on recording state"""
        if hasattr(self, 'tabbed_panels'):
            for panel in self.tabbed_panels:
                try:
                    # Use stored reference to dictate button
                    if hasattr(panel, 'editor_widget') and hasattr(panel.editor_widget, 'dictate_btn'):
                        button = panel.editor_widget.dictate_btn
                        if is_recording:
                            button.setText("🔴 Recording...")
                            button.setStyleSheet("background-color: #D32F2F; color: white; font-weight: bold;")
                        else:
                            button.setText("🎤 Dictate (F9)")
                            button.setStyleSheet("background-color: #4CAF50; color: white; font-weight: bold;")
                except:
                    pass

    def save_tab_segment(self):
        """Save current segment in tab editor"""
        if not hasattr(self, 'tab_current_segment_id') or not self.tab_current_segment_id:
            return
        self.log(f"✓ Saved segment {self.tab_current_segment_id}")
    
    def save_tab_segment_and_next(self):
        """Save current segment and move to next in tab editor"""
        self.save_tab_segment()
        
        # Move to next segment based on current view
        if self.current_view_mode == LayoutMode.GRID or (hasattr(self, 'home_view_stack') and 
                                                          self.home_view_stack.currentIndex() == 0):
            # Grid view - move to next row in table
            if hasattr(self, 'table') and self.table:
                current_row = self.table.currentRow()
                if current_row < self.table.rowCount() - 1:
                    self.table.setCurrentCell(current_row + 1, self.table.currentColumn())
        elif self.current_view_mode == LayoutMode.LIST or (hasattr(self, 'home_view_stack') and 
                                                             self.home_view_stack.currentIndex() == 1):
            # List view - move to next item in list
            if hasattr(self, 'list_tree') and self.list_tree:
                current_item = self.list_tree.currentItem()
                if current_item:
                    current_index = self.list_tree.indexOfTopLevelItem(current_item)
                    if current_index < self.list_tree.topLevelItemCount() - 1:
                        next_item = self.list_tree.topLevelItem(current_index + 1)
                        self.list_tree.setCurrentItem(next_item)
    
    def save_tab_notes(self):
        """Save comments in tab editor"""
        if not hasattr(self, 'tab_current_segment_id') or not self.tab_current_segment_id:
            return
        self.log(f"✓ Saved comments for segment {self.tab_current_segment_id}")
    
    def update_tab_segment_editor(self, segment_id: int, source_text: str, target_text: str, 
                                   status: str = "untranslated", notes: str = ""):
        """Update the tab segment editor with current segment data"""
        # Update ALL tabbed panels (grid view and list view)
        if hasattr(self, 'tabbed_panels'):
            for panel in self.tabbed_panels:
                try:
                    # Update segment editor tab
                    if hasattr(panel, 'editor_widget'):
                        editor = panel.editor_widget
                        editor.seg_info_label.setText(f"Segment {segment_id}")
                        editor.source_editor.setPlainText(source_text)
                        editor.target_editor.setPlainText(target_text)
                        idx = editor.status_combo.findData(status)
                        if idx >= 0:
                            editor.status_combo.blockSignals(True)
                            editor.status_combo.setCurrentIndex(idx)
                            editor.status_combo.blockSignals(False)
                    
                    # Update notes tab
                    if hasattr(panel, 'notes_widget'):
                        panel.notes_widget.notes_editor.setPlainText(notes)
                except Exception as e:
                    self.log(f"Error updating tabbed panel: {e}")
        
        # Store current segment ID
        self.tab_current_segment_id = segment_id
    
    # ========================================================================
    # DOCUMENT VIEW METHODS
    # ========================================================================
    
    def on_doc_segment_clicked(self, segment_id: int):
        """Handle segment click in Document View"""
        self.doc_current_segment_id = segment_id
        
        # Find segment
        segment = next((s for s in self.current_project.segments if s.id == segment_id), None)
        if not segment:
            return
        
        # Update all tabbed panels (like Grid and List views)
        self.update_tab_segment_editor(
            segment_id=segment.id,
            source_text=segment.source,
            target_text=segment.target,
            status=segment.status,
            notes=segment.notes if hasattr(segment, 'notes') else ""
        )
        
        # Also update main assistance widget with translation matches
        # Trigger assistance panel update (search for matches)
        row = next((i for i, s in enumerate(self.current_project.segments) if s.id == segment_id), -1)
        if row >= 0:
            # Trigger the match search like we do in grid/list view
            self.on_cell_selected(row, 3, -1, -1)
    
    def on_doc_status_change(self, status: str):
        """Handle status change in Document View"""
        if not self.doc_current_segment_id:
            return
        
        segment = next((s for s in self.current_project.segments if s.id == self.doc_current_segment_id), None)
        if segment:
            segment.status = status
            self.project_modified = True
            self.refresh_document_view()
            # Reselect segment after refresh
            if self.doc_current_segment_id and self.doc_current_segment_id in self.doc_segment_widgets:
                seg_id = self.doc_current_segment_id  # Store in local variable for lambda
                QTimer.singleShot(100, lambda: self.on_doc_segment_clicked(seg_id))
    
    def on_doc_target_change(self):
        """Handle target text change in Document View"""
        if not self.doc_current_segment_id:
            return
        
        segment = next((s for s in self.current_project.segments if s.id == self.doc_current_segment_id), None)
        if segment:
            new_target = self.doc_target_editor.toPlainText()
            if segment.target != new_target:
                segment.target = new_target
                self.project_modified = True
                # Update document view display
                if self.doc_current_segment_id in self.doc_segment_widgets:
                    widget = self.doc_segment_widgets[self.doc_current_segment_id]
                    layout = widget.layout()
                    if layout:
                        for i in range(layout.count()):
                            item = layout.itemAt(i)
                            if item and item.widget():
                                w = item.widget()
                                if isinstance(w, QLabel) and "Target:" in w.text():
                                    if new_target:
                                        w.setText(f"<b>Target:</b> {new_target}")
                                        w.setStyleSheet("color: #0066cc;")
                                    else:
                                        w.setText("<i>Not translated</i>")
                                        w.setStyleSheet("color: #999; font-style: italic;")
                                    break
    
    def copy_source_to_doc_target(self):
        """Copy source to target in Document View"""
        source_text = self.doc_source_editor.toPlainText()
        self.doc_target_editor.setPlainText(source_text)
        self.on_doc_target_change()
    
    def clear_doc_target(self):
        """Clear target in Document View"""
        self.doc_target_editor.clear()
        self.on_doc_target_change()
    
    def save_doc_segment_and_next(self):
        """Save current segment and move to next in Document View"""
        if not self.doc_current_segment_id:
            return
        
        # Find current segment index
        current_idx = next((i for i, s in enumerate(self.current_project.segments) if s.id == self.doc_current_segment_id), -1)
        if current_idx < 0:
            return
        
        # Find next segment
        next_idx = current_idx + 1
        if next_idx < len(self.current_project.segments):
            next_segment = self.current_project.segments[next_idx]
            # Scroll to and select next segment
            if next_segment.id in self.doc_segment_widgets:
                self.on_doc_segment_clicked(next_segment.id)
                # Scroll to widget
                widget = self.doc_segment_widgets[next_segment.id]
                # Navigate up to find QScrollArea
                parent = widget.parent()
                while parent and not isinstance(parent, QScrollArea):
                    parent = parent.parent()
                if isinstance(parent, QScrollArea):
                    parent.ensureWidgetVisible(widget)
            self.log(f"✓ Moved to segment {next_segment.id}")
        else:
            self.log("✓ Last segment - no next segment")
    
    # ========================================================================
    # UTILITY
    # ========================================================================
    
    def update_window_title(self):
        """Update window title with project name and modified state"""
        title = "Supervertaler Qt v1.3.2"
        if ENABLE_PRIVATE_FEATURES:
            title += " [🛠️ DEV MODE]"
        if self.current_project:
            title += f" - {self.current_project.name}"
            if self.project_modified:
                title += " *"
        self.setWindowTitle(title)
    
    def log(self, message: str):
        """Log message to status bar and session log"""
        if hasattr(self, 'status_bar'):
            self.status_bar.showMessage(message)
        print(f"[LOG] {message}")
        
        # Also append to session log tab if it exists
        if hasattr(self, 'session_log') and self.session_log:
            from datetime import datetime
            timestamp = datetime.now().strftime("%H:%M:%S")
            formatted_message = f"[{timestamp}] {message}\n"
            try:
                self.session_log.appendPlainText(formatted_message)
                # Auto-scroll to bottom
                scrollbar = self.session_log.verticalScrollBar()
                if scrollbar:
                    scrollbar.setValue(scrollbar.maximum())
            except Exception:
                pass  # Silently fail if widget not ready
    
    def show_options_dialog(self):
        """Show application options dialog - DEPRECATED: Redirects to Settings tab for backwards compatibility"""
        # Redirect to Settings tab instead of showing dialog
        self._go_to_settings_tab()
    
    def open_api_keys_file(self):
        """Open API keys file in system text editor"""
        api_keys_file = self.user_data_path / "api_keys.txt"
        
        # Create file if it doesn't exist
        if not api_keys_file.exists():
            try:
                # Copy from example file
                example_file = self.user_data_path / "api_keys.example.txt"
                if example_file.exists():
                    import shutil
                    shutil.copy(example_file, api_keys_file)
                    self.log(f"✓ Created api_keys.txt from example")
                else:
                    # Create basic file
                    with open(api_keys_file, 'w') as f:
                        f.write("# Add your API keys here\n")
                        f.write("# openai=sk-your-key\n")
                        f.write("# claude=sk-ant-your-key\n")
                        f.write("# gemini=your-key\n")
            except Exception as e:
                QMessageBox.warning(self, "Error", f"Could not create api_keys.txt: {str(e)}")
                return
        
        # Open in system editor
        try:
            import subprocess
            import platform
            if platform.system() == 'Windows':
                os.startfile(api_keys_file)  # type: ignore
            elif platform.system() == 'Darwin':  # macOS
                subprocess.run(['open', str(api_keys_file)])
            else:  # Linux
                subprocess.run(['xdg-open', str(api_keys_file)])
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Could not open api_keys.txt: {str(e)}")
    
    def _go_to_settings_tab(self):
        """Navigate to Settings tab (from menu)"""
        if hasattr(self, 'right_tabs'):
            # Right tabs: Prompt Manager=0, Resources=1, Modules=2, Settings=3
            self.right_tabs.setCurrentIndex(3)
    
    def _go_to_universal_lookup(self):
        """Navigate to Universal Lookup in Modules tab"""
        if hasattr(self, 'right_tabs'):
            # Right tabs: Prompt Manager=0, Resources=1, Modules=2, Settings=3
            self.right_tabs.setCurrentIndex(2)  # Switch to Modules tab
            # Then switch to Universal Lookup sub-tab
            if hasattr(self, 'modules_tabs'):
                # Find Universal Lookup index in modules tabs
                for i in range(self.modules_tabs.count()):
                    if "Universal Lookup" in self.modules_tabs.tabText(i):
                        self.modules_tabs.setCurrentIndex(i)
                        break
    
    def open_api_keys_file(self):
        """Open API keys file in system text editor"""
        api_keys_file = self.user_data_path / "api_keys.txt"
        
        # Create file if it doesn't exist
        if not api_keys_file.exists():
            try:
                # Copy from example file
                example_file = self.user_data_path / "api_keys.example.txt"
                if example_file.exists():
                    import shutil
                    shutil.copy(example_file, api_keys_file)
                    self.log(f"✓ Created api_keys.txt from example")
                else:
                    # Create basic file
                    with open(api_keys_file, 'w') as f:
                        f.write("# Add your API keys here\n")
                        f.write("# openai=sk-your-key\n")
                        f.write("# claude=sk-ant-your-key\n")
                        f.write("# gemini=your-key\n")
            except Exception as e:
                QMessageBox.warning(self, "Error", f"Could not create api_keys.txt: {str(e)}")
                return
        
        # Open in system editor
        try:
            import subprocess
            import platform
            
            if platform.system() == 'Windows':
                os.startfile(api_keys_file)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.call(['open', api_keys_file])
            else:  # Linux
                subprocess.call(['xdg-open', api_keys_file])
                
            self.log(f"✓ Opened {api_keys_file}")
        except Exception as e:
            QMessageBox.information(
                self, "API Keys File",
                f"Please edit this file manually:\n\n{api_keys_file}"
            )
    
    def load_llm_settings(self) -> Dict[str, str]:
        """Load LLM settings from user preferences"""
        prefs_file = self.user_data_path / "ui_preferences.json"
        
        if not prefs_file.exists():
            return {
                'provider': 'openai',
                'openai_model': 'gpt-4o',
                'claude_model': 'claude-sonnet-4-5-20250929',
                'gemini_model': 'gemini-2.5-flash'
            }
        
        try:
            with open(prefs_file, 'r') as f:
                prefs = json.load(f)
                return prefs.get('llm_settings', {
                    'provider': 'openai',
                    'openai_model': 'gpt-4o',
                    'claude_model': 'claude-sonnet-4-5-20250929',
                    'gemini_model': 'gemini-2.5-flash'
                })
        except:
            return {
                'provider': 'openai',
                'openai_model': 'gpt-4o',
                'claude_model': 'claude-sonnet-4-5-20250929',
                'gemini_model': 'gemini-2.5-flash'
            }
    
    def save_llm_settings(self, settings: Dict[str, str]):
        """Save LLM settings to user preferences"""
        prefs_file = self.user_data_path / "ui_preferences.json"
        
        # Load existing preferences
        prefs = {}
        if prefs_file.exists():
            try:
                with open(prefs_file, 'r') as f:
                    prefs = json.load(f)
            except:
                pass
        
        # Update LLM settings
        prefs['llm_settings'] = settings
        
        # Save back
        try:
            with open(prefs_file, 'w') as f:
                json.dump(prefs, f, indent=2)
        except Exception as e:
            self.log(f"⚠ Could not save LLM settings: {str(e)}")
    
    def load_provider_enabled_states(self) -> Dict[str, bool]:
        """Load provider enable/disable states from user preferences"""
        prefs_file = self.user_data_path / "ui_preferences.json"
        
        if not prefs_file.exists():
            # Default: all providers enabled
            return {
                'llm_openai': True,
                'llm_claude': True,
                'llm_gemini': True,
                'mt_google_translate': True,
                'mt_deepl': True,
                'mt_microsoft': True,
                'mt_amazon': True,
                'mt_modernmt': True,
                'mt_mymemory': True
            }
        
        try:
            with open(prefs_file, 'r') as f:
                prefs = json.load(f)
                return prefs.get('provider_enabled_states', {
                    'llm_openai': True,
                    'llm_claude': True,
                    'llm_gemini': True,
                    'mt_google_translate': True,
                    'mt_deepl': True,
                    'mt_microsoft': True,
                    'mt_amazon': True,
                    'mt_modernmt': True,
                    'mt_mymemory': True
                })
        except:
            # Default: all providers enabled
            return {
                'llm_openai': True,
                'llm_claude': True,
                'llm_gemini': True,
                'mt_google_translate': True,
                'mt_deepl': True,
                'mt_microsoft': True,
                'mt_amazon': True,
                'mt_modernmt': True,
                'mt_mymemory': True
            }
    
    def save_provider_enabled_states(self, states: Dict[str, bool]):
        """Save provider enable/disable states to user preferences"""
        prefs_file = self.user_data_path / "ui_preferences.json"
        
        # Load existing preferences
        prefs = {}
        if prefs_file.exists():
            try:
                with open(prefs_file, 'r') as f:
                    prefs = json.load(f)
            except:
                pass
        
        # Update provider enabled states
        prefs['provider_enabled_states'] = states
        
        # Save back
        try:
            with open(prefs_file, 'w') as f:
                json.dump(prefs, f, indent=2)
        except Exception as e:
            self.log(f"⚠ Could not save provider enabled states: {str(e)}")
    
    def update_warning_banner(self):
        """Show/hide warning banner based on allow_replace_in_source setting"""
        if self.allow_replace_in_source:
            self.warning_banner.show()
        else:
            self.warning_banner.hide()
    
    def show_tm_manager(self):
        """Show Translation Memory Manager dialog"""
        from modules.translation_memory import TMDatabase
        
        if not self.current_project:
            QMessageBox.warning(
                self,
                "No Project Loaded",
                "Please load or create a project first."
            )
            return
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Translation Memory Manager")
        dialog.setMinimumSize(900, 600)
        
        layout = QVBoxLayout(dialog)
        
        # Header
        header_label = QLabel("<h2>📚 Translation Memory Manager</h2>")
        layout.addWidget(header_label)
        
        # Tab widget for different TM management tasks
        tabs = QTabWidget()
        
        # ===== Tab 1: TM Database Overview =====
        overview_tab = QWidget()
        overview_layout = QVBoxLayout(overview_tab)
        
        # Current project info
        info_group = QGroupBox("Current Project")
        info_layout = QGridLayout()
        info_layout.addWidget(QLabel("<b>Name:</b>"), 0, 0)
        info_layout.addWidget(QLabel(self.current_project.name), 0, 1)
        info_layout.addWidget(QLabel("<b>Language Pair:</b>"), 1, 0)
        info_layout.addWidget(QLabel(f"{self.current_project.source_lang} → {self.current_project.target_lang}"), 1, 1)
        info_group.setLayout(info_layout)
        overview_layout.addWidget(info_group)
        
        # TM Statistics
        stats_group = QGroupBox("Translation Memory Statistics")
        stats_layout = QVBoxLayout()
        
        if self.tm_database:
            # Get TM count from database
            try:
                # Query database for entry count
                cursor = self.tm_database.db.connection.cursor()
                cursor.execute(
                    "SELECT COUNT(*) FROM translation_units WHERE source_lang=? AND target_lang=?",
                    (self.current_project.source_lang, self.current_project.target_lang)
                )
                entry_count = cursor.fetchone()[0]
                
                stats_text = f"""<ul>
                    <li><b>Total Entries:</b> {entry_count:,}</li>
                    <li><b>Database Path:</b> {self.tm_database.db.db_path}</li>
                    <li><b>Active Language Pair:</b> {self.current_project.source_lang} → {self.current_project.target_lang}</li>
                </ul>"""
                stats_label = QLabel(stats_text)
            except Exception as e:
                stats_label = QLabel(f"<span style='color: #ef4444;'>Error reading database: {e}</span>")
        else:
            stats_label = QLabel("<span style='color: #f59e0b;'>No TM database initialized</span>")
        
        stats_layout.addWidget(stats_label)
        stats_group.setLayout(stats_layout)
        overview_layout.addWidget(stats_group)
        
        # Refresh button
        refresh_btn = QPushButton("🔄 Refresh Statistics")
        refresh_btn.clicked.connect(lambda: self.show_tm_manager())  # Reopen dialog
        overview_layout.addWidget(refresh_btn)
        
        overview_layout.addStretch()
        tabs.addTab(overview_tab, "📊 Overview")
        
        # ===== Tab 2: Add Entry Manually =====
        add_tab = QWidget()
        add_layout = QVBoxLayout(add_tab)
        
        add_label = QLabel("<h3>Add Translation Entry</h3>")
        add_layout.addWidget(add_label)
        
        form_layout = QGridLayout()
        
        form_layout.addWidget(QLabel(f"<b>Source ({self.current_project.source_lang}):</b>"), 0, 0)
        source_input = QTextEdit()
        source_input.setMaximumHeight(100)
        source_input.setPlaceholderText("Enter source text...")
        form_layout.addWidget(source_input, 1, 0)
        
        form_layout.addWidget(QLabel(f"<b>Target ({self.current_project.target_lang}):</b>"), 2, 0)
        target_input = QTextEdit()
        target_input.setMaximumHeight(100)
        target_input.setPlaceholderText("Enter target translation...")
        form_layout.addWidget(target_input, 3, 0)
        
        add_layout.addLayout(form_layout)
        
        # Add button
        def add_entry():
            source = source_input.toPlainText().strip()
            target = target_input.toPlainText().strip()
            
            if not source or not target:
                QMessageBox.warning(dialog, "Missing Data", "Please enter both source and target text.")
                return
            
            if self.tm_database:
                try:
                    self.tm_database.add_entry(source, target)
                    QMessageBox.information(dialog, "Success", "Translation entry added to TM database!")
                    source_input.clear()
                    target_input.clear()
                    self.log(f"Added TM entry: {source[:50]}... → {target[:50]}...")
                except Exception as e:
                    QMessageBox.critical(dialog, "Error", f"Failed to add entry: {e}")
            else:
                QMessageBox.warning(dialog, "No Database", "TM database not initialized.")
        
        add_btn = QPushButton("✅ Add to Translation Memory")
        add_btn.clicked.connect(add_entry)
        add_btn.setStyleSheet("background-color: #22c55e; color: white; padding: 10px; font-weight: bold;")
        add_layout.addWidget(add_btn)
        
        add_layout.addStretch()
        tabs.addTab(add_tab, "➕ Add Entry")
        
        # ===== Tab 3: Import TMX =====
        import_tab = QWidget()
        import_layout = QVBoxLayout(import_tab)
        
        import_label = QLabel("<h3>Import TMX File</h3>")
        import_layout.addWidget(import_label)
        
        import_info = QLabel(
            "Import translation memories from TMX files. "
            "Entries will be added to the database for the current language pair."
        )
        import_info.setWordWrap(True)
        import_layout.addWidget(import_info)
        
        # File selection
        file_group = QGroupBox("TMX File")
        file_layout = QHBoxLayout()
        
        file_path_label = QLabel("No file selected")
        file_path_label.setStyleSheet("color: #6b7280;")
        file_layout.addWidget(file_path_label, 1)
        
        selected_file = [None]  # Use list to make it mutable in nested function
        
        def browse_tmx():
            file_path, _ = QFileDialog.getOpenFileName(
                dialog,
                "Select TMX File",
                str(self.user_data_path / "Translation_Resources"),
                "TMX Files (*.tmx);;All Files (*.*)"
            )
            if file_path:
                selected_file[0] = file_path
                file_path_label.setText(file_path)
                file_path_label.setStyleSheet("color: #22c55e;")
        
        browse_btn = QPushButton("📁 Browse...")
        browse_btn.clicked.connect(browse_tmx)
        file_layout.addWidget(browse_btn)
        
        file_group.setLayout(file_layout)
        import_layout.addWidget(file_group)
        
        # Import button
        def import_tmx():
            if not selected_file[0]:
                QMessageBox.warning(dialog, "No File", "Please select a TMX file first.")
                return
            
            if not self.tm_database:
                QMessageBox.warning(dialog, "No Database", "TM database not initialized.")
                return
            
            try:
                # Import TMX file
                count = self.tm_database.import_tmx(
                    selected_file[0],
                    self.current_project.source_lang,
                    self.current_project.target_lang
                )
                
                QMessageBox.information(
                    dialog,
                    "Import Successful",
                    f"Imported {count} translation entries from TMX file!"
                )
                self.log(f"Imported {count} entries from {selected_file[0]}")
                
                # Clear selection
                selected_file[0] = None
                file_path_label.setText("No file selected")
                file_path_label.setStyleSheet("color: #6b7280;")
                
            except Exception as e:
                QMessageBox.critical(dialog, "Import Error", f"Failed to import TMX file:\n{e}")
        
        import_btn = QPushButton("📥 Import TMX File")
        import_btn.clicked.connect(import_tmx)
        import_btn.setStyleSheet("background-color: #3b82f6; color: white; padding: 10px; font-weight: bold;")
        import_layout.addWidget(import_btn)
        
        import_layout.addStretch()
        tabs.addTab(import_tab, "📥 Import TMX")
        
        # ===== Tab 4: Browse Entries =====
        browse_tab = QWidget()
        browse_layout = QVBoxLayout(browse_tab)
        
        browse_label = QLabel("<h3>Browse TM Entries</h3>")
        browse_layout.addWidget(browse_label)
        
        # Search box
        search_layout = QHBoxLayout()
        search_layout.addWidget(QLabel("🔍 Search:"))
        search_input = QLineEdit()
        search_input.setPlaceholderText("Enter search term...")
        search_layout.addWidget(search_input, 1)
        browse_layout.addLayout(search_layout)
        
        # Results table
        results_table = QTableWidget()
        results_table.setColumnCount(3)
        results_table.setHorizontalHeaderLabels(["Source", "Target", "Match %"])
        results_table.horizontalHeader().setStretchLastSection(False)
        results_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        results_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        results_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Fixed)
        results_table.setColumnWidth(2, 80)
        results_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        browse_layout.addWidget(results_table)
        
        def search_tm():
            search_term = search_input.text().strip()
            if not search_term or not self.tm_database:
                results_table.setRowCount(0)
                return
            
            try:
                matches = self.tm_database.search_all(search_term, max_matches=50)
                
                results_table.setRowCount(len(matches))
                for i, match in enumerate(matches):
                    source_item = QTableWidgetItem(match.get('source', ''))
                    target_item = QTableWidgetItem(match.get('target', ''))
                    match_pct_item = QTableWidgetItem(f"{match.get('match_pct', 0)}%")
                    
                    # Color code by match percentage
                    match_pct = match.get('match_pct', 0)
                    if match_pct == 100:
                        color = "#22c55e"
                    elif match_pct >= 95:
                        color = "#3b82f6"
                    elif match_pct >= 85:
                        color = "#f59e0b"
                    else:
                        color = "#ef4444"
                    
                    match_pct_item.setForeground(QColor(color))
                    match_pct_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                    
                    results_table.setItem(i, 0, source_item)
                    results_table.setItem(i, 1, target_item)
                    results_table.setItem(i, 2, match_pct_item)
                
            except Exception as e:
                QMessageBox.critical(dialog, "Search Error", f"Failed to search TM: {e}")
        
        search_input.returnPressed.connect(search_tm)
        search_btn = QPushButton("🔍 Search")
        search_btn.clicked.connect(search_tm)
        search_layout.addWidget(search_btn)
        
        browse_layout.addWidget(QLabel("<i>Tip: Press Enter to search</i>"))
        
        tabs.addTab(browse_tab, "🔍 Browse")
        
        # Add tabs to dialog
        layout.addWidget(tabs)
        
        # Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.accept)
        layout.addWidget(close_btn)
        
        dialog.exec()
    
    def show_about(self):
        """Show about dialog with clickable website link"""
        dialog = QDialog(self)
        dialog.setWindowTitle("About Supervertaler Qt")
        dialog.setMinimumWidth(400)
        
        layout = QVBoxLayout(dialog)
        layout.setSpacing(10)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Title
        title = QLabel("<h2>Supervertaler Qt v1.3.2</h2>")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)
        
        # Description
        desc = QLabel("<p><b>AI-powered tool for translators & writers</b></p>")
        desc.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(desc)
        
        # Website link
        website_label = QLabel(
            '<p style="text-align: center;">'
            '<a href="https://supervertaler.com/" style="color: #1976D2; text-decoration: none;">'
            'https://supervertaler.com/'
            '</a></p>'
        )
        website_label.setOpenExternalLinks(True)
        website_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(website_label)
        
        # Additional info
        info = QLabel(
            "<p><b>Author:</b> Michael Beijer</p>"
            "<p><b>License:</b> MIT</p>"
            "<hr>"
            "<p><i>v1.2.2 - Translation Results Panels, Document Formatting & Tag System</i></p>"
        )
        info.setWordWrap(True)
        layout.addWidget(info)
        
        # OK button
        ok_btn = QPushButton("OK")
        ok_btn.clicked.connect(dialog.accept)
        ok_btn.setDefault(True)
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        button_layout.addWidget(ok_btn)
        button_layout.addStretch()
        layout.addLayout(button_layout)
        
        dialog.exec()
    
    def closeEvent(self, event):
        """Handle window close event"""
        if self.project_modified:
            reply = QMessageBox.question(
                self,
                "Unsaved Changes",
                "You have unsaved changes. Do you want to save before closing?",
                QMessageBox.StandardButton.Save |
                QMessageBox.StandardButton.Discard |
                QMessageBox.StandardButton.Cancel
            )
            
            if reply == QMessageBox.StandardButton.Save:
                self.save_project()
                event.accept()
            elif reply == QMessageBox.StandardButton.Discard:
                event.accept()
            else:
                event.ignore()
        else:
            event.accept()
    
    # =========================================================================
    # LLM TRANSLATION INTEGRATION
    # =========================================================================
    
    def translate_current_segment(self):
        """Translate currently selected segment using LLM (Ctrl+T)"""
        current_row = self.table.currentRow()
        if current_row < 0 or not self.current_project:
            QMessageBox.warning(self, "No Selection", "Please select a segment to translate.")
            return
        
        segment = self.current_project.segments[current_row]
        
        if not segment.source.strip():
            QMessageBox.warning(self, "Empty Source", "Cannot translate empty source text.")
            return
        
        # Load API keys
        api_keys = self.load_api_keys()
        if not api_keys:
            reply = QMessageBox.question(
                self, "API Keys Missing",
                "No API keys found. Would you like to configure them now?\n\n"
                f"Keys should be in: {self.user_data_path / 'api_keys.txt'}",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                self._go_to_settings_tab()
            return
        
        # Load LLM settings
        settings = self.load_llm_settings()
        provider = settings.get('provider', 'openai')
        
        # Get model based on provider
        model_key = f'{provider}_model'
        model = settings.get(model_key, 'gpt-4o')
        
        # Check if API key exists for selected provider
        if provider not in api_keys:
            reply = QMessageBox.question(
                self, f"{provider.title()} API Key Missing",
                f"{provider.title()} API key not found in api_keys.txt\n\n"
                f"Would you like to configure it now?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                self._go_to_settings_tab()
            return
        
        try:
            self.status_bar.showMessage(f"Translating segment #{segment.id} with {provider} ({model})...")
            QApplication.processEvents()  # Update UI
            
            # Use modular LLM client with user's settings
            from modules.llm_clients import LLMClient
            
            client = LLMClient(
                api_key=api_keys[provider],
                provider=provider,
                model=model
            )
            
            # Check TM before API call if enabled
            general_prefs = self.load_general_settings()
            check_tm_before_api = general_prefs.get('check_tm_before_api', True)
            tm_match = None
            
            if check_tm_before_api and self.tm_database:
                try:
                    matches = self.tm_database.search_all(segment.source, max_matches=1)
                    if matches and matches[0].get('match_pct', 0) == 100:
                        tm_match = matches[0].get('target', '')
                        self.log(f"✓ Found 100% TM match for segment #{segment.id}")
                except Exception as e:
                    self.log(f"⚠ TM check error: {e}")
            
            # If we have a 100% match and auto-insert is enabled, use it
            auto_insert_100 = general_prefs.get('auto_insert_100', False)
            if tm_match and auto_insert_100:
                # Auto-insert the TM match
                segment.target = tm_match
                segment.status = "translated"
                self.project_modified = True
                self.update_window_title()
                
                # Update grid (using cell widget)
                target_widget = self.table.cellWidget(current_row, 3)
                if target_widget and isinstance(target_widget, EditableGridTextEditor):
                    target_widget.setPlainText(tm_match)
                else:
                    # Fallback if widget doesn't exist
                    self.table.setItem(current_row, 3, QTableWidgetItem(tm_match))
                self.update_status_icon(current_row, "translated")
                
                # Add to TM if not already there
                if self.tm_database:
                    try:
                        self.tm_database.add_to_project_tm(segment.source, tm_match)
                    except:
                        pass
                
                self.log(f"✓ Auto-inserted 100% TM match for segment #{segment.id}")
                self.status_bar.showMessage(f"✓ Segment #{segment.id} translated from TM (100% match)", 3000)
                return
            
            # Build full prompt using prompt manager (includes UICONTROL tags, etc.)
            custom_prompt = None
            if hasattr(self, 'prompt_manager_qt') and self.prompt_manager_qt:
                try:
                    # Get surrounding segments if enabled
                    surrounding_segments = general_prefs.get('surrounding_segments', 5)
                    surrounding_context = ""
                    
                    if surrounding_segments > 0:
                        try:
                            current_idx = current_row
                            start_idx = max(0, current_idx - surrounding_segments)
                            end_idx = min(len(self.current_project.segments), current_idx + surrounding_segments + 1)
                            surrounding = self.current_project.segments[start_idx:end_idx]
                            
                            if len(surrounding) > 1:  # Only add if we have surrounding segments
                                context_parts = []
                                for i, seg in enumerate(surrounding):
                                    actual_idx = start_idx + i
                                    if actual_idx == current_idx:
                                        context_parts.append(f">>> {seg.id}. {seg.source} <<<  [TRANSLATE THIS]")
                                    else:
                                        context_parts.append(f"{seg.id}. {seg.source}")
                                        if seg.target:
                                            context_parts.append(f"    → {seg.target}")
                                
                                surrounding_context = "\n\n**SURROUNDING SEGMENTS FOR CONTEXT:**\n"
                                surrounding_context += "(The segment marked with >>> <<< is the one to translate)\n\n"
                                surrounding_context += "\n".join(context_parts)
                                self.log(f"  Including {len(surrounding)} surrounding segments ({surrounding_segments} before/after)")
                        except Exception as e:
                            self.log(f"⚠ Could not add surrounding segments: {e}")
                    
                    custom_prompt = self.prompt_manager_qt.build_final_prompt(
                        source_text=segment.source,
                        source_lang=self.current_project.source_lang,
                        target_lang=self.current_project.target_lang,
                        mode="single"
                    )
                    
                    # Add surrounding context before the translation delimiter
                    if surrounding_context:
                        # Insert before the "YOUR TRANSLATION" delimiter
                        if "**YOUR TRANSLATION" in custom_prompt:
                            custom_prompt = custom_prompt.replace(
                                "**YOUR TRANSLATION",
                                surrounding_context + "\n\n**YOUR TRANSLATION"
                            )
                        else:
                            custom_prompt += surrounding_context
                            
                except Exception as e:
                    self.log(f"⚠ Could not build prompt from manager: {e}")
            
            # Translate using the module
            translation = client.translate(
                text=segment.source,
                source_lang=self.current_project.source_lang,
                target_lang=self.current_project.target_lang,
                custom_prompt=custom_prompt
            )
            
            if translation:
                # Update segment
                segment.target = translation
                segment.status = "translated"
                
                # Update grid - Column 3 is Target (using cell widget)
                target_widget = self.table.cellWidget(current_row, 3)
                if target_widget and isinstance(target_widget, EditableGridTextEditor):
                    target_widget.setPlainText(translation)
                else:
                    # Fallback: create new widget if none exists
                    self.table.setItem(current_row, 3, QTableWidgetItem(translation))
                self.update_status_icon(current_row, "translated")
                
                # Add to Translation Memory
                if self.tm_database:
                    try:
                        self.tm_database.add_to_project_tm(segment.source, translation)
                        self.log(f"✓ Added to TM: {segment.source[:30]}... → {translation[:30]}...")
                    except Exception as tm_error:
                        self.log(f"⚠ Could not add to TM: {str(tm_error)}")
                
                # Mark project as modified
                self.project_modified = True
                self.update_window_title()
                
                self.log(f"✓ Segment #{segment.id} translated with {provider}/{model}")
                self.status_bar.showMessage(f"✓ Segment #{segment.id} translated", 3000)
            else:
                self.log(f"✗ Translation failed for segment #{segment.id}")
                QMessageBox.warning(self, "Translation Failed", "No translation received from LLM.")
                
        except Exception as e:
            self.log(f"✗ Translation error: {str(e)}")
            QMessageBox.critical(self, "Translation Error", f"Failed to translate segment:\n\n{str(e)}")
            self.status_bar.showMessage("Translation failed", 3000)
    
    def translate_multiple_segments(self, scope: str):
        """Translate segments by scope (selected, status-based, or all)."""
        if not self.current_project:
            QMessageBox.warning(self, "No Project", "Please load or create a project first.")
            return

        segments_with_rows: List[Tuple[int, Segment]] = []
        description = ""

        if scope == "selected_not_started":
            selected = self._get_selected_segments_with_rows()
            if not selected:
                QMessageBox.information(
                    self,
                    "No Selection",
                    "Please select one or more segments in the grid or list view first."
                )
                return
            target_statuses = {"not_started"}
            segments_with_rows = [(idx, seg) for idx, seg in selected if seg.status in target_statuses]
            if not segments_with_rows:
                QMessageBox.information(
                    self,
                    "Nothing to Translate",
                    "The selected segments are not marked as 'Not started'."
                )
                return
            description = "selected not-started segment(s)"
        elif scope == "all_not_started":
            target_statuses = {"not_started"}
            segments_with_rows = [
                (idx, seg) for idx, seg in enumerate(self.current_project.segments)
                if seg.status in target_statuses
            ]
            description = "not-started segment(s)"
        elif scope == "all_pretranslated":
            target_statuses = {"pretranslated"}
            segments_with_rows = [
                (idx, seg) for idx, seg in enumerate(self.current_project.segments)
                if seg.status in target_statuses
            ]
            description = "pre-translated segment(s)"
        elif scope == "all_not_started_pretranslated":
            target_statuses = {"not_started", "pretranslated"}
            segments_with_rows = [
                (idx, seg) for idx, seg in enumerate(self.current_project.segments)
                if seg.status in target_statuses
            ]
            description = "not-started and pre-translated segment(s)"
        elif scope == "all_translatable":
            target_statuses = TRANSLATABLE_STATUSES
            segments_with_rows = [
                (idx, seg) for idx, seg in enumerate(self.current_project.segments)
                if seg.status in target_statuses
            ]
            labels = ", ".join(sorted({get_status(s).label for s in target_statuses}))
            description = f"{labels} segment(s)"
        elif scope == "all_segments":
            segments_with_rows = list(enumerate(self.current_project.segments))
            description = "segment(s)"
        else:
            return

        if not segments_with_rows:
            QMessageBox.information(
                self,
                "Nothing to Translate",
                f"No {description} available for translation."
            )
            return

        segments_with_rows = sorted(segments_with_rows, key=lambda item: item[0])
        self.translate_batch(segments_with_rows=segments_with_rows, scope_description=description)

    def _get_selected_segments_with_rows(self) -> List[Tuple[int, Segment]]:
        """Return selected segments across active views as (row_index, segment)."""
        if not self.current_project:
            return []

        selected_ids = set()

        try:
            for seg in self.get_selected_segments_from_grid():
                if hasattr(seg, 'id'):
                    selected_ids.add(seg.id)
        except Exception:
            pass

        try:
            for seg in self.get_selected_segments_from_list():
                if hasattr(seg, 'id'):
                    selected_ids.add(seg.id)
        except Exception:
            pass

        if not selected_ids:
            return []

        segments_with_rows: List[Tuple[int, Segment]] = []
        for idx, seg in enumerate(self.current_project.segments):
            if seg.id in selected_ids:
                segments_with_rows.append((idx, seg))

        return segments_with_rows

    def translate_batch(self, segments_with_rows: Optional[List[Tuple[int, Segment]]] = None, scope_description: Optional[str] = None):
        """Translate multiple segments with progress dialog"""
        if not self.current_project:
            QMessageBox.warning(self, "No Project", "Please load or create a project first.")
            return

        if segments_with_rows is None:
            segments_with_rows = [
                (idx, seg) for idx, seg in enumerate(self.current_project.segments)
                if seg.status == DEFAULT_STATUS.key
            ]
            scope_description = scope_description or "untranslated segment(s)"
        else:
            valid_segments: List[Tuple[int, Segment]] = []
            for row_index, segment in segments_with_rows:
                if 0 <= row_index < len(self.current_project.segments):
                    valid_segments.append((row_index, segment))
            segments_with_rows = valid_segments
            scope_description = scope_description or "segment(s)"

        if not segments_with_rows:
            QMessageBox.information(
                self,
                "Nothing to Translate",
                f"No {scope_description} available for translation."
            )
            return

        segments_to_translate = sorted(segments_with_rows, key=lambda item: item[0])
        total_segments = len(segments_to_translate)

        reply = QMessageBox.question(
            self,
            "Batch Translation",
            f"Found {total_segments} {scope_description}.\n\n"
            f"Translate them using your configured LLM provider?\n\n"
            f"This may take several minutes and consume API credits.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply != QMessageBox.StandardButton.Yes:
            return

        # Load API keys and settings
        api_keys = self.load_api_keys()
        if not api_keys:
            QMessageBox.critical(
                self, "API Keys Missing",
                "Please configure your API keys in Settings first."
            )
            return

        settings = self.load_llm_settings()
        provider = settings.get('provider', 'openai')
        model_key = f'{provider}_model'
        model = settings.get(model_key, 'gpt-4o')

        if provider not in api_keys:
            QMessageBox.critical(
                self, f"{provider.title()} API Key Missing",
                f"Please configure your {provider.title()} API key in Settings."
            )
            return

        # Create progress dialog
        progress = QDialog(self)
        progress.setWindowTitle("Batch Translation Progress")
        progress.setMinimumWidth(600)
        progress.setMinimumHeight(250)
        progress.setModal(True)

        layout = QVBoxLayout(progress)

        # Header
        header_label = QLabel(f"<h3>🚀 Translating {total_segments} segments</h3>")
        layout.addWidget(header_label)

        # Provider info
        info_label = QLabel(f"<b>Provider:</b> {provider.title()} | <b>Model:</b> {model}")
        layout.addWidget(info_label)

        # Progress bar
        progress_bar = QProgressBar()
        progress_bar.setMaximum(total_segments)
        progress_bar.setValue(0)
        layout.addWidget(progress_bar)

        # Current segment label
        current_label = QLabel("Starting...")
        layout.addWidget(current_label)

        # Statistics
        stats_label = QLabel(f"Translated: 0 | Failed: 0 | Remaining: {total_segments}")
        layout.addWidget(stats_label)

        # Close button (initially disabled)
        close_btn = QPushButton("Close")
        close_btn.setEnabled(False)
        close_btn.clicked.connect(progress.accept)
        layout.addWidget(close_btn)

        # Show progress dialog
        progress.show()
        QApplication.processEvents()

        translated_count = 0
        failed_count = 0
        segment_idx = 0

        try:
            from modules.llm_clients import LLMClient

            client = LLMClient(
                api_key=api_keys[provider],
                provider=provider,
                model=model
            )

            source_lang = getattr(self.current_project, 'source_lang', 'en')
            target_lang = getattr(self.current_project, 'target_lang', 'nl')

            self.log(f"Batch translation: {source_lang} → {target_lang}")

            general_prefs = self.load_general_settings()
            batch_size = general_prefs.get('batch_size', 100)
            total_batches = (total_segments + batch_size - 1) // batch_size

            for batch_num in range(total_batches):
                batch_start = batch_num * batch_size
                batch_end = min(batch_start + batch_size, total_segments)
                batch_segments = segments_to_translate[batch_start:batch_end]

                current_label.setText(f"Translating batch {batch_num + 1}/{total_batches} ({len(batch_segments)} segments)...")
                progress_bar.setValue(segment_idx)
                QApplication.processEvents()

                try:
                    batch_prompt_parts = []

                    base_prompt = None
                    if hasattr(self, 'prompt_manager_qt') and self.prompt_manager_qt and batch_segments:
                        try:
                            first_segment = batch_segments[0][1]
                            base_prompt = self.prompt_manager_qt.build_final_prompt(
                                source_text=first_segment.source,
                                source_lang=source_lang,
                                target_lang=target_lang,
                                mode="single"
                            )
                            if "**SOURCE TEXT:**" in base_prompt:
                                base_prompt = base_prompt.split("**SOURCE TEXT:**")[0].strip()
                            elif "Translate the following" in base_prompt:
                                base_prompt = base_prompt.split("Translate the following")[0].strip()
                        except Exception:
                            base_prompt = None

                    if base_prompt:
                        batch_prompt_parts.append(base_prompt)
                    else:
                        batch_prompt_parts.append(f"Translate the following text segments from {source_lang} to {target_lang}.")

                    use_full_context = general_prefs.get('use_full_context', True)
                    if use_full_context and self.current_project:
                        try:
                            context_parts = []
                            for seg in self.current_project.segments:
                                if seg.source:
                                    context_parts.append(f"{seg.id}. {seg.source}")
                                    if seg.target:
                                        context_parts.append(f"    → {seg.target}")

                            if context_parts:
                                batch_prompt_parts.append("\n**FULL DOCUMENT CONTEXT:**")
                                batch_prompt_parts.append("(For reference - segments to translate are marked below)\n")
                                batch_prompt_parts.append("\n".join(context_parts))
                                batch_prompt_parts.append("")
                                self.log(f"  Including full document context ({len(self.current_project.segments)} segments)")
                        except Exception as e:
                            self.log(f"⚠ Could not add full document context: {e}")

                    batch_prompt_parts.append(f"\n**SEGMENTS TO TRANSLATE ({len(batch_segments)} segments):**")
                    batch_prompt_parts.append("Provide ONLY the translations, one per line, in the same order. NO explanations, NO segment numbers, NO labels.\n")

                    for row_index, seg in batch_segments:
                        batch_prompt_parts.append(f"{seg.id}. {seg.source}")

                    batch_prompt_parts.append("\n**YOUR TRANSLATIONS (one per line, in order):**")

                    batch_prompt = "\n".join(batch_prompt_parts)

                    self.log(f"🤖 Translating batch {batch_num + 1}/{total_batches} ({len(batch_segments)} segments)...")

                    first_segment_text = batch_segments[0][1].source if batch_segments else ""
                    batch_response = client.translate(
                        text=first_segment_text,
                        source_lang=source_lang,
                        target_lang=target_lang,
                        custom_prompt=batch_prompt
                    )

                    import re
                    translation_lines = []
                    for line in batch_response.strip().split('\n'):
                        cleaned_line = re.sub(r'^[\d\-\*\)\.\s]+', '', line.strip())
                        if cleaned_line:
                            translation_lines.append(cleaned_line)

                    if len(translation_lines) != len(batch_segments):
                        self.log(f"⚠ Warning: Expected {len(batch_segments)} translations, got {len(translation_lines)}")
                        while len(translation_lines) < len(batch_segments):
                            translation_lines.append("")
                        translation_lines = translation_lines[:len(batch_segments)]

                    for i, (row_index, segment) in enumerate(batch_segments):
                        translation = translation_lines[i] if i < len(translation_lines) else ""

                        if translation:
                            segment.target = translation
                            segment.status = "draft"

                            if row_index < self.table.rowCount():
                                target_widget = self.table.cellWidget(row_index, 3)
                                if target_widget and isinstance(target_widget, EditableGridTextEditor):
                                    target_widget.setPlainText(translation)
                                else:
                                    self.table.setItem(row_index, 3, QTableWidgetItem(translation))
                                    self.update_status_icon(row_index, "translated")

                            if self.tm_database:
                                try:
                                    self.tm_database.add_to_project_tm(segment.source, translation)
                                except Exception:
                                    pass

                            translated_count += 1
                            self.log(f"✓ Batch: Segment #{segment.id} translated")
                        else:
                            failed_count += 1
                            self.log(f"✗ Batch: Segment #{segment.id} - empty translation")

                        segment_idx += 1
                        progress_bar.setValue(segment_idx)
                        remaining = total_segments - segment_idx
                        stats_label.setText(
                            f"Translated: {translated_count} | Failed: {failed_count} | Remaining: {remaining}"
                        )
                        QApplication.processEvents()

                except Exception as e:
                    for row_index, segment in batch_segments:
                        failed_count += 1
                        segment_idx += 1
                        self.log(f"✗ Batch: Segment #{segment.id} - {str(e)}")
                        progress_bar.setValue(segment_idx)
                        remaining = total_segments - segment_idx
                        stats_label.setText(
                            f"Translated: {translated_count} | Failed: {failed_count} | Remaining: {remaining}"
                        )
                        QApplication.processEvents()

            if translated_count > 0:
                self.project_modified = True
                self.update_window_title()

            progress_bar.setValue(total_segments)
            current_label.setText(
                f"<b>✓ Batch translation complete!</b><br>"
                f"Successfully translated: {translated_count}<br>"
                f"Failed: {failed_count}"
            )

            close_btn.setEnabled(True)

            self.log(f"✓ Batch translation complete: {translated_count} translated, {failed_count} failed")

        except Exception as e:
            QMessageBox.critical(
                progress,
                "Batch Translation Error",
                f"Batch translation failed:\n\n{str(e)}"
            )
            self.log(f"✗ Batch translation error: {str(e)}")
            close_btn.setEnabled(True)

        progress.exec()
    
    def _fetch_llm_translation_async(self, source_text: str, segment, row_index: int):
        """Fetch LLM translation asynchronously and add to results panel"""
        if not source_text or not source_text.strip():
            return
        
        # Check if we have API keys and LLM settings configured
        api_keys = self.load_api_keys()
        if not api_keys:
            return  # No API keys configured, skip LLM

        settings = self.load_llm_settings()
        provider = settings.get('provider', 'openai')
        
        # Check if provider is enabled
        enabled_providers = self.load_provider_enabled_states()
        provider_enabled_map = {
            "openai": "llm_openai",
            "claude": "llm_claude",
            "gemini": "llm_gemini"
        }
        enabled_key = provider_enabled_map.get(provider)
        if not enabled_providers.get(enabled_key, True):
            return  # Provider is disabled, skip LLM
        
        provider_key_map = {
            "openai": "openai",
            "claude": "claude",
            "gemini": "gemini"
        }
        api_key_name = provider_key_map.get(provider)
        if not api_keys.get(api_key_name):
            return  # No API key for selected provider
        
        # Get model based on provider
        model_map = {
            "openai": settings.get('openai_model', 'gpt-4o'),
            "claude": settings.get('claude_model', 'claude-sonnet-4-5-20250929'),
            "gemini": settings.get('gemini_model', 'gemini-2.5-flash')
        }
        model = model_map.get(provider, 'gpt-4o')
        
        # Get languages
        source_lang = getattr(self.current_project, 'source_lang', 'en')
        target_lang = getattr(self.current_project, 'target_lang', 'nl')
        
        # Use QTimer to run in background (non-blocking)
        def fetch_and_update():
            try:
                from modules.llm_clients import LLMClient
                from modules.translation_results_panel import TranslationMatch
                
                client = LLMClient(
                    api_key=api_keys[api_key_name],
                    provider=provider,
                    model=model
                )
                
                # Build full prompt using prompt manager
                custom_prompt = None
                try:
                    # Access parent through closure
                    parent = self
                    if hasattr(parent, 'prompt_manager_qt') and parent.prompt_manager_qt:
                        custom_prompt = parent.prompt_manager_qt.build_final_prompt(
                            source_text=source_text,
                            source_lang=source_lang,
                            target_lang=target_lang,
                            mode="single"
                        )
                except Exception as e:
                    self.log(f"⚠ Could not build LLM prompt from manager: {e}")
                
                translation = client.translate(
                    text=source_text,
                    source_lang=source_lang,
                    target_lang=target_lang,
                    custom_prompt=custom_prompt
                )
                
                if translation and hasattr(self, 'assistance_widget') and self.assistance_widget:
                    # Check if this segment is still selected
                    current_row = self.table.currentRow()
                    if current_row == row_index:
                        # Map provider to code
                        provider_code_map = {
                            'openai': 'GPT',
                            'claude': 'CL',
                            'gemini': 'GEM'
                        }
                        provider_code = provider_code_map.get(provider, 'LLM')
                        
                        # Create LLM match
                        llm_match = TranslationMatch(
                            source=source_text,
                            target=translation,
                            relevance=98,  # High relevance for LLM
                            metadata={
                                'model': model,
                                'provider': provider,
                                'timestamp': datetime.now().isoformat()
                            },
                            match_type='LLM',
                            compare_source=source_text,
                            provider_code=provider_code
                        )
                        
                        # Update all results panels with LLM match
                        if hasattr(self, 'results_panels'):
                            current_matches = {"LLM": [llm_match]}
                            for panel in self.results_panels:
                                try:
                                    # Get current matches and add LLM
                                    if hasattr(panel, 'matches_by_type'):
                                        existing_matches = panel.matches_by_type.copy()
                                        existing_matches["LLM"] = [llm_match]
                                        panel.set_matches(existing_matches)
                                    else:
                                        panel.set_matches(current_matches)
                                except Exception as e:
                                    self.log(f"Error updating panel with LLM: {e}")
                        self.log(f"✓ LLM translation added to results pane")
            except Exception as e:
                # Silently fail - don't interrupt workflow
                self.log(f"LLM translation failed (silent): {str(e)}")
        
        # Use QTimer to run async (small delay to avoid blocking)
        QTimer.singleShot(100, fetch_and_update)
    
    def _fetch_mt_translation_async(self, source_text: str, segment, row_index: int):
        """Fetch Machine Translation asynchronously and add to results panel"""
        if not source_text or not source_text.strip():
            return
        
        # Check if we have MT API keys configured
        api_keys = self.load_api_keys()
        if not api_keys:
            return  # No API keys configured, skip MT
        
        # Check enabled providers
        enabled_providers = self.load_provider_enabled_states()
        
        # Try MT providers in priority order (first available and enabled wins)
        # Priority: Google Translate > DeepL > Microsoft > Amazon > ModernMT > MyMemory
        mt_provider = None
        mt_api_key = None
        
        if enabled_providers.get('mt_google_translate', True) and api_keys.get('google_translate'):
            mt_provider = 'google_translate'
            mt_api_key = api_keys['google_translate']
        elif enabled_providers.get('mt_deepl', True) and api_keys.get('deepl'):
            mt_provider = 'deepl'
            mt_api_key = api_keys['deepl']
        elif enabled_providers.get('mt_microsoft', True) and (api_keys.get('microsoft_translate') or api_keys.get('azure_translate')):
            mt_provider = 'microsoft'
            mt_api_key = api_keys.get('microsoft_translate') or api_keys.get('azure_translate')
        elif enabled_providers.get('mt_amazon', True) and (api_keys.get('amazon_translate') or api_keys.get('aws_translate')):
            mt_provider = 'amazon'
            mt_api_key = api_keys.get('amazon_translate') or api_keys.get('aws_translate')
        elif enabled_providers.get('mt_modernmt', True) and api_keys.get('modernmt'):
            mt_provider = 'modernmt'
            mt_api_key = api_keys['modernmt']
        elif enabled_providers.get('mt_mymemory', True) and api_keys.get('mymemory'):
            mt_provider = 'mymemory'
            mt_api_key = api_keys.get('mymemory')  # Optional, MyMemory works without key
        # Note: MyMemory can work without a key, but we only use it if explicitly configured
        else:
            return  # No MT API keys available or all providers disabled
        
        # Get languages
        source_lang = getattr(self.current_project, 'source_lang', 'en')
        target_lang = getattr(self.current_project, 'target_lang', 'nl')
        
        # Use QTimer to run in background (non-blocking)
        def fetch_and_update():
            try:
                from modules.translation_results_panel import TranslationMatch
                
                # Call appropriate MT service
                if mt_provider == 'google_translate':
                    translation = self.call_google_translate(source_text, source_lang, target_lang, mt_api_key)
                    provider_name = 'Google Translate'
                elif mt_provider == 'deepl':
                    translation = self.call_deepl(source_text, source_lang, target_lang, mt_api_key)
                    provider_name = 'DeepL'
                elif mt_provider == 'microsoft':
                    translation = self.call_microsoft_translate(source_text, source_lang, target_lang, mt_api_key)
                    provider_name = 'Microsoft Translator'
                elif mt_provider == 'amazon':
                    # Amazon requires region, check for amazon_translate_region or use default
                    region = api_keys.get('amazon_translate_region', 'us-east-1')
                    translation = self.call_amazon_translate(source_text, source_lang, target_lang, mt_api_key, region)
                    provider_name = 'Amazon Translate'
                elif mt_provider == 'modernmt':
                    translation = self.call_modernmt(source_text, source_lang, target_lang, mt_api_key)
                    provider_name = 'ModernMT'
                elif mt_provider == 'mymemory':
                    translation = self.call_mymemory(source_text, source_lang, target_lang, mt_api_key)
                    provider_name = 'MyMemory'
                else:
                    return
                
                if translation and not translation.startswith('['):  # Skip error messages
                    if hasattr(self, 'assistance_widget') and self.assistance_widget:
                        # Check if this segment is still selected
                        current_row = self.table.currentRow()
                        if current_row == row_index:
                            # Map provider to code
                            provider_code_map = {
                                'Google Translate': 'GT',
                                'DeepL': 'DL',
                                'Microsoft Translator': 'MS',
                                'Amazon Translate': 'AT',
                                'ModernMT': 'MMT',
                                'MyMemory': 'MM'
                            }
                            provider_code = provider_code_map.get(provider_name, 'MT')
                            
                            # Create MT match
                            mt_match = TranslationMatch(
                                source=source_text,
                                target=translation,
                                relevance=95,  # High relevance for MT
                                metadata={
                                    'provider': provider_name,
                                    'timestamp': datetime.now().isoformat()
                                },
                                match_type='MT',
                                compare_source=source_text,
                                provider_code=provider_code
                            )
                            
                            # Update all results panels with MT match
                            if hasattr(self, 'results_panels'):
                                current_matches = {"MT": [mt_match]}
                                for panel in self.results_panels:
                                    try:
                                        # Get current matches and add MT
                                        if hasattr(panel, 'matches_by_type'):
                                            existing_matches = panel.matches_by_type.copy()
                                            existing_matches["MT"] = [mt_match]
                                            panel.set_matches(existing_matches)
                                        else:
                                            panel.set_matches(current_matches)
                                    except Exception as e:
                                        self.log(f"Error updating panel with MT: {e}")
                            self.log(f"✓ MT translation ({provider_name}) added to results pane")
            except Exception as e:
                # Silently fail - don't interrupt workflow
                self.log(f"MT translation failed (silent): {str(e)}")
        
        # Use QTimer to run async (small delay to avoid blocking, slightly after LLM)
        QTimer.singleShot(150, fetch_and_update)
    
    def call_google_translate(self, text: str, source_lang: str, target_lang: str, api_key: str = None) -> str:
        """Call Google Cloud Translation API using REST API"""
        try:
            import requests
            
            if not api_key:
                api_keys = self.load_api_keys()
                api_key = api_keys.get("google_translate")
            
            if not api_key:
                return "[Google Cloud Translation requires API key]"
            
            # Convert language codes (handle locale codes like "en-US" -> "en")
            src_code = source_lang.split('-')[0].split('_')[0].lower()
            tgt_code = target_lang.split('-')[0].split('_')[0].lower()
            
            # Call REST API directly
            url = "https://translation.googleapis.com/language/translate/v2"
            params = {
                'key': api_key,
                'q': text,
                'source': src_code,
                'target': tgt_code,
                'format': 'text'
            }
            
            response = requests.post(url, params=params, timeout=10)
            response.raise_for_status()
            
            result = response.json()
            return result['data']['translations'][0]['translatedText']
            
        except ImportError:
            return "[Google Translate requires: pip install requests]"
        except Exception as e:
            return f"[Google Translate error: {str(e)}]"
    
    def call_deepl(self, text: str, source_lang: str, target_lang: str, api_key: str = None) -> str:
        """Call DeepL API"""
        try:
            import deepl
            
            if not api_key:
                api_keys = self.load_api_keys()
                api_key = api_keys.get("deepl")
            
            if not api_key:
                return "[DeepL requires API key]"
            
            translator = deepl.Translator(api_key)
            
            # Convert language codes (DeepL uses uppercase)
            src_code = source_lang.split('-')[0].split('_')[0].upper()
            tgt_code = target_lang.split('-')[0].split('_')[0].upper()
            
            result = translator.translate_text(text, source_lang=src_code, target_lang=tgt_code)
            return result.text
            
        except ImportError:
            return "[DeepL requires: pip install deepl]"
        except Exception as e:
            return f"[DeepL error: {str(e)}]"
    
    def call_microsoft_translate(self, text: str, source_lang: str, target_lang: str, api_key: str = None, region: str = None) -> str:
        """Call Microsoft Azure Translator API"""
        try:
            import requests
            
            api_keys = self.load_api_keys()
            if not api_key:
                api_key = api_keys.get("microsoft_translate") or api_keys.get("azure_translate")
            if not region:
                region = api_keys.get("microsoft_translate_region") or api_keys.get("azure_region") or "global"
            
            if not api_key:
                return "[Microsoft Translator requires API key]"
            
            # Convert language codes (Azure uses standard codes)
            src_code = source_lang.split('-')[0].split('_')[0].lower()
            tgt_code = target_lang.split('-')[0].split('_')[0].lower()
            
            # Microsoft Translator API v3.0
            endpoint = f"https://api.cognitive.microsofttranslator.com/translate"
            params = {
                'api-version': '3.0',
                'from': src_code,
                'to': tgt_code
            }
            headers = {
                'Ocp-Apim-Subscription-Key': api_key,
                'Ocp-Apim-Subscription-Region': region,
                'Content-Type': 'application/json'
            }
            body = [{'text': text}]
            
            response = requests.post(endpoint, params=params, headers=headers, json=body, timeout=10)
            response.raise_for_status()
            
            result = response.json()
            return result[0]['translations'][0]['text']
            
        except ImportError:
            return "[Microsoft Translator requires: pip install requests]"
        except Exception as e:
            return f"[Microsoft Translator error: {str(e)}]"
    
    def call_amazon_translate(self, text: str, source_lang: str, target_lang: str, api_key: str = None, region: str = 'us-east-1') -> str:
        """Call Amazon Translate API (AWS)"""
        try:
            import boto3
            from botocore.exceptions import ClientError
            
            api_keys = self.load_api_keys()
            if not api_key:
                api_key = api_keys.get("amazon_translate") or api_keys.get("aws_translate")
            
            # AWS also needs secret key
            secret_key = api_keys.get("amazon_translate_secret") or api_keys.get("aws_secret_key")
            
            if not api_key or not secret_key:
                return "[Amazon Translate requires API key and secret key]"
            
            # Convert language codes (AWS uses standard codes)
            src_code = source_lang.split('-')[0].split('_')[0]
            tgt_code = target_lang.split('-')[0].split('_')[0]
            
            translate_client = boto3.client(
                'translate',
                aws_access_key_id=api_key,
                aws_secret_access_key=secret_key,
                region_name=region
            )
            
            result = translate_client.translate_text(
                Text=text,
                SourceLanguageCode=src_code,
                TargetLanguageCode=tgt_code
            )
            
            return result['TranslatedText']
            
        except ImportError:
            return "[Amazon Translate requires: pip install boto3]"
        except Exception as e:
            return f"[Amazon Translate error: {str(e)}]"
    
    def call_modernmt(self, text: str, source_lang: str, target_lang: str, api_key: str = None) -> str:
        """Call ModernMT API"""
        try:
            import requests
            
            api_keys = self.load_api_keys()
            if not api_key:
                api_key = api_keys.get("modernmt")
            
            if not api_key:
                return "[ModernMT requires API key]"
            
            # Convert language codes
            src_code = source_lang.split('-')[0].split('_')[0]
            tgt_code = target_lang.split('-')[0].split('_')[0]
            
            # ModernMT API endpoint
            url = "https://api.modernmt.com/translate"
            headers = {
                'MMT-ApiKey': api_key,
                'Content-Type': 'application/json'
            }
            data = {
                'q': text,
                'source': src_code,
                'target': tgt_code
            }
            
            response = requests.post(url, headers=headers, json=data, timeout=10)
            response.raise_for_status()
            
            result = response.json()
            return result['data']['translation']
            
        except ImportError:
            return "[ModernMT requires: pip install requests]"
        except Exception as e:
            return f"[ModernMT error: {str(e)}]"
    
    def call_mymemory(self, text: str, source_lang: str, target_lang: str, api_key: str = None) -> str:
        """Call MyMemory Translation API (free tier available, simple REST API)"""
        try:
            import requests
            
            # MyMemory is free, but API key provides higher limits
            api_keys = self.load_api_keys()
            api_key = api_key or api_keys.get("mymemory")  # Optional, works without key
            
            # Convert language codes (MyMemory uses 2-letter codes)
            src_code = source_lang.split('-')[0].split('_')[0].lower()
            tgt_code = target_lang.split('-')[0].split('_')[0].lower()
            
            # MyMemory API endpoint
            url = "https://api.mymemory.translated.net/get"
            params = {
                'q': text,
                'langpair': f"{src_code}|{tgt_code}"
            }
            if api_key:
                params['key'] = api_key
            
            response = requests.get(url, params=params, timeout=10)
            response.raise_for_status()
            
            result = response.json()
            if result.get('responseStatus') == 200:
                return result['responseData']['translatedText']
            else:
                return f"[MyMemory error: {result.get('responseDetails', 'Unknown error')}]"
            
        except ImportError:
            return "[MyMemory requires: pip install requests]"
        except Exception as e:
            return f"[MyMemory error: {str(e)}]"
    
    def load_api_keys(self) -> Dict[str, str]:
        """Load API keys from user_data folder"""
        api_keys = {}
        api_keys_file = self.user_data_path / "api_keys.txt"
        
        if not api_keys_file.exists():
            return api_keys
        
        try:
            with open(api_keys_file, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line and '=' in line and not line.startswith('#'):
                        key, value = line.split('=', 1)
                        api_keys[key.strip().lower()] = value.strip()
        except Exception as e:
            self.log(f"⚠ Error loading API keys: {str(e)}")
        
        return api_keys
    
    def ensure_example_api_keys(self):
        """Create example API keys file on first launch for new users"""
        example_file = self.user_data_path / "api_keys.example.txt"
        
        # Only create if it doesn't exist
        if example_file.exists():
            return
        
        example_content = """# Supervertaler API Keys Configuration
# =======================================
# 
# Add your API keys below. Remove the # to activate a key.
# Never commit this file with real keys to version control!
#
# For actual use:
# 1. Copy this file to "api_keys.txt" in the same folder
# 2. Add your real API keys (remove the # from the lines you use)
# 3. Save the file
#
# Available providers:

# OpenAI (GPT-4, GPT-4o, GPT-5, o1, o3 models)
# Get your key at: https://platform.openai.com/api-keys
#openai = YOUR_OPENAI_KEY_HERE

# Anthropic Claude (Claude Sonnet 4.5, Haiku 4.5, Opus 4.1)
# Get your key at: https://console.anthropic.com/settings/keys
#claude = YOUR_CLAUDE_KEY_HERE

# Google Gemini (Gemini 2.0 Flash, Pro models)
# Get your key at: https://aistudio.google.com/app/apikey
#google = YOUR_GOOGLE_GEMINI_KEY_HERE

# ===== Machine Translation APIs (for MT preview) =====

# Supervertaler will use MT providers in this priority order (first available):
# 1. Google Translate
# 2. DeepL
# 3. Microsoft Translator
# 4. Amazon Translate
# 5. ModernMT
# 6. MyMemory (free fallback, works without key)

# Google Cloud Translation API (Priority: Highest)
# Get your key at: https://console.cloud.google.com/apis/credentials
# Enable "Cloud Translation API" first, then create API key
#google_translate = YOUR_GOOGLE_CLOUD_TRANSLATE_API_KEY

# DeepL API Pro
# Get your key at: https://www.deepl.com/pro-api
# Note: Free DeepL web version keys don't work, you need API Pro
#deepl = YOUR_DEEPL_API_PRO_KEY

# Microsoft Azure Translator
# Get your key at: https://azure.microsoft.com/en-us/services/cognitive-services/translator/
# Requires Azure subscription
#microsoft_translate = YOUR_AZURE_TRANSLATOR_KEY
#microsoft_translate_region = global
# (Region examples: "global", "eastus", "westus")

# Amazon Translate (AWS)
# Get your keys at: https://aws.amazon.com/translate/
# Requires AWS account
#amazon_translate = YOUR_AWS_ACCESS_KEY_ID
#amazon_translate_secret = YOUR_AWS_SECRET_ACCESS_KEY
#amazon_translate_region = us-east-1
# (AWS region examples: "us-east-1", "eu-west-1", etc.)

# ModernMT
# Get your key at: https://www.modernmt.com/api
#modernmt = YOUR_MODERNMT_API_KEY

# MyMemory Translation (Free tier available, no key required)
# Get API key (optional, for higher limits): https://mymemory.translated.net/
# Works without key but with lower rate limits
#mymemory = YOUR_MYMEMORY_KEY

# ===== NOTES =====

# Temperature settings for reasoning models:
# - GPT-5, o1, o3: Temperature parameter OMITTED (not supported by these models)
# - Standard models: Use temperature=0.3 (automatically applied)

# TROUBLESHOOTING:
# ---------------
# If you see "API Key Missing" errors:
# 1. Check that your file is named "api_keys.txt" (not "api_keys.example.txt")
# 2. Make sure the # is removed from the beginning of the line
# 3. Verify there are no extra spaces around the = sign
# 4. Confirm your key is valid (test at the provider's website)
#
# If keys still don't work:
# 1. Check console output to see which file was loaded
# 2. Look for errors: "Error reading api_keys.txt"
# 3. Restart Supervertaler after editing api_keys.txt
#
# DeepL Authorization Error:
# - Your key must be for "DeepL API Pro" (not free web version)
# - Get API key from: https://www.deepl.com/pro-api
#
# Google Translate Error:
# - Enable "Cloud Translation API" in Google Cloud Console
# - Create API key at: https://console.cloud.google.com/apis/credentials
"""
        
        try:
            with open(example_file, 'w', encoding='utf-8') as f:
                f.write(example_content)
            self.log(f"✓ Created example API keys file: {example_file}")
        except Exception as e:
            self.log(f"⚠ Could not create example API keys file: {str(e)}")
    
    def show_autofingers(self):
        """Show AutoFingers by switching to the AutoFingers tab"""
        # Find the AutoFingers tab index and activate it
        # AutoFingers is in Modules tab (right_tabs index 2)
        if hasattr(self, 'right_tabs'):
            self.right_tabs.setCurrentIndex(2)  # Switch to Modules tab
            # Then switch to AutoFingers sub-tab
            if hasattr(self, 'modules_tabs'):
                for i in range(self.modules_tabs.count()):
                    if "AutoFingers" in self.modules_tabs.tabText(i):
                        self.modules_tabs.setCurrentIndex(i)
                        break
    
    def show_theme_editor(self):
        """Show Theme Editor dialog"""
        dialog = ThemeEditorDialog(self, self.theme_manager)
        if dialog.exec():
            # Theme may have been changed, reapply
            self.theme_manager.apply_theme(QApplication.instance())
    
    def _schedule_mt_and_llm_matches(self, segment, termbase_matches=None):
        """Schedule MT and LLM matches with debouncing - only call APIs when user stops clicking"""
        try:
            # Cancel any previous MT/LLM requests
            if hasattr(self, '_mt_llm_timer'):
                self._mt_llm_timer.stop()
                self.log("🚀 Cancelled previous MT/LLM request (user moved to new segment)")
            
            # Store segment and termbase matches for delayed lookup
            self._pending_mt_llm_segment = segment
            self._pending_termbase_matches = termbase_matches or []
            
            # Start debounced timer - only call APIs after user stops clicking for 1.5 seconds
            from PyQt6.QtCore import QTimer
            self._mt_llm_timer = QTimer()
            self._mt_llm_timer.setSingleShot(True)
            self._mt_llm_timer.timeout.connect(lambda: self._execute_mt_llm_lookup())
            self._mt_llm_timer.start(1500)  # Wait 1.5 seconds of inactivity
            
            self.log(f"🚀 Scheduled MT/LLM lookup for '{segment.source[:50]}...' (will execute in 1.5s if user doesn't move)")
                
        except Exception as e:
            self.log(f"Error scheduling MT/LLM search: {e}")
    
    def _execute_mt_llm_lookup(self):
        """Execute the actual MT/LLM lookup after debounce delay"""
        try:
            if not hasattr(self, '_pending_mt_llm_segment'):
                self.log("🚀 DELAYED LOOKUP CANCELLED: No _pending_mt_llm_segment attribute")
                return
                
            segment = self._pending_mt_llm_segment
            self.log(f"🚀 EXECUTING delayed MT/LLM lookup for '{segment.source[:50]}...'")
            self.log(f"🚀 DELAYED LOOKUP: LLM matching enabled: {self.enable_llm_matching}")
            self.log(f"🚀 DELAYED LOOKUP: MT matching enabled: {self.enable_mt_matching}")
            
            # Call the actual lookup method (now includes TM, MT, and LLM)
            self._search_mt_and_llm_matches(segment)
                
        except Exception as e:
            self.log(f"Error executing MT/LLM lookup: {e}")
            import traceback
            self.log(f"Delayed lookup error traceback: {traceback.format_exc()}")
    
    def _search_mt_and_llm_matches(self, segment):
        """Search for TM, MT and LLM matches - called only after debounce delay"""
        try:
            from modules.translation_results_panel import TranslationMatch
            
            # Get current project languages for all translation services
            source_lang = getattr(self.current_project, 'source_lang', None) if self.current_project else None
            target_lang = getattr(self.current_project, 'target_lang', None) if self.current_project else None
            
            # Convert language names to codes if needed
            source_lang_code = None
            target_lang_code = None
            if source_lang:
                source_lang_code = self._convert_language_to_code(source_lang)
            if target_lang:
                target_lang_code = self._convert_language_to_code(target_lang)
            
            # Prepare matches dict - use the stored termbase matches from immediate display
            matches_dict = {
                "LLM": [],
                "NT": [],
                "MT": [],
                "TM": [],
                "Termbases": getattr(self, '_pending_termbase_matches', [])
            }
            
            # 🔥 DELAYED TM SEARCH: Search TM database (moved from UI thread to prevent blocking)
            if self.enable_tm_matching and hasattr(self, 'db_manager') and self.db_manager:
                try:
                    self.log(f"🚀 DELAYED TM SEARCH: Searching for '{segment.source[:50]}...'")
                    self.log(f"🚀 DELAYED TM SEARCH: Project languages: {source_lang_code} → {target_lang_code}")
                    
                    # Search in primary direction (current project direction)
                    tm_matches = self.db_manager.search_all(segment.source, max_results=10)
                    
                    # Also search in reverse direction (bidirectional TM matching)
                    reverse_matches = []
                    if source_lang_code and target_lang_code and source_lang_code != target_lang_code:
                        try:
                            # Search for text as target text in reverse direction
                            reverse_query = """
                                SELECT source_text, target_text, source_lang, target_lang, tm_id, usage_count
                                FROM translation_units 
                                WHERE target_text = ? AND source_lang = ? AND target_lang = ?
                                ORDER BY usage_count DESC
                                LIMIT 5
                            """
                            self.db_manager.cursor.execute(reverse_query, (segment.source, target_lang_code, source_lang_code))
                            reverse_rows = self.db_manager.cursor.fetchall()
                            
                            for row in reverse_rows:
                                reverse_matches.append({
                                    'source': row['target_text'],  # Swap source/target for reverse match
                                    'target': row['source_text'],
                                    'match_pct': 95,  # High relevance for reverse exact match
                                    'tm_name': f"{row['tm_id'].replace('_', ' ').title()} (Reverse)",
                                    'tm_id': row['tm_id']
                                })
                            
                            self.log(f"🚀 DELAYED TM SEARCH: Found {len(reverse_matches)} reverse matches")
                        except Exception as e:
                            self.log(f"Error in reverse TM search: {e}")
                    
                    # Combine primary and reverse matches
                    all_tm_matches = tm_matches + reverse_matches
                    self.log(f"🚀 DELAYED TM SEARCH: Total matches found: {len(all_tm_matches)} ({len(tm_matches)} primary + {len(reverse_matches)} reverse)")
                    
                    for match in all_tm_matches:
                        match_obj = TranslationMatch(
                            source=match.get('source', ''),
                            target=match.get('target', ''),
                            relevance=match.get('match_pct', 0),
                            metadata={
                                'context': match.get('context', ''),
                                'tm_name': match.get('tm_name', ''),
                                'timestamp': match.get('created_at', ''),
                                'direction': 'reverse' if 'Reverse' in match.get('tm_name', '') else 'primary'
                            },
                            match_type='TM',
                            compare_source=match.get('source', ''),
                            provider_code='TM'
                        )
                        matches_dict["TM"].append(match_obj)
                except Exception as e:
                    self.log(f"Error in delayed TM search: {e}")
            
            # Add MT and LLM matches
            self._add_mt_and_llm_matches(segment, matches_dict, source_lang, target_lang, source_lang_code, target_lang_code)
            
            # Update display with all delayed matches (TM, MT, LLM)
            total_new_matches = len(matches_dict["TM"]) + len(matches_dict["MT"]) + len(matches_dict["LLM"])
            if total_new_matches > 0:
                self.log(f"🚀 DELAYED SEARCH COMPLETE: Added {len(matches_dict['TM'])} TM + {len(matches_dict['MT'])} MT + {len(matches_dict['LLM'])} LLM matches")
                # Update all results panels
                if hasattr(self, 'results_panels') and self.results_panels:
                    self.log(f"✅ Updating {len(self.results_panels)} results panels with delayed matches")
                    for panel in self.results_panels:
                        try:
                            panel.set_matches(matches_dict)
                        except Exception as e:
                            self.log(f"Error updating panel: {e}")
            else:
                self.log("🚀 DELAYED SEARCH COMPLETE: No TM, MT or LLM matches found")
                
        except Exception as e:
            self.log(f"Error in MT/LLM search: {e}")
    def _add_mt_and_llm_matches(self, segment, matches_dict, source_lang, target_lang, source_lang_code, target_lang_code):
        """Add MT and LLM matches to the matches dictionary"""
        from modules.translation_results_panel import TranslationMatch
        
        # 🤖 DIRECT MT SEARCH: Machine Translation matches
        if self.enable_mt_matching:
            try:
                self.log(f"🤖 DIRECT MT SEARCH: Getting machine translation for '{segment.source[:50]}...'")
                
                # Use Google Translate via the wrapper function
                from modules.llm_clients import get_google_translation
                
                self.log(f"🤖 DIRECT MT SEARCH: Calling get_google_translation with source_lang='{source_lang_code or 'auto'}', target_lang='{target_lang_code or 'en'}'")
                
                google_result = get_google_translation(
                    segment.source, 
                    source_lang_code or 'auto',  # Auto-detect if no source lang 
                    target_lang_code or 'en'     # Default to English if no target lang
                )
                
                self.log(f"🤖 DIRECT MT SEARCH: Google Translate result: {google_result}")
                
                if google_result and google_result.get('translation'):
                    match = TranslationMatch(
                        source=segment.source,
                        target=google_result['translation'],
                        relevance=95,  # High confidence for MT
                        metadata=google_result.get('metadata', {}),
                        match_type="MT",
                        provider_code='GT'
                    )
                    matches_dict["MT"].append(match)
                    self.log(f"✓ Added Google Translate match: {google_result['translation'][:50]}...")
                elif google_result and google_result.get('error'):
                    self.log(f"⚠ Google Translate error: {google_result.get('error')}")
                else:
                    self.log(f"⚠ Google Translate returned no translation")
                    
            except Exception as e:
                self.log(f"⚠ Error getting Google Translate: {e}")
                import traceback
                self.log(f"⚠ Traceback: {traceback.format_exc()}")
        
        # 🧠 DIRECT LLM SEARCH: Large Language Model matches  
        if self.enable_llm_matching:
            try:
                self.log(f"🧠 DIRECT LLM SEARCH: Getting LLM translations for '{segment.source[:50]}...'")
                self.log(f"🧠 DIRECT LLM SEARCH: LLM matching is enabled: {self.enable_llm_matching}")
                
                # Import LLM wrapper functions
                from modules.llm_clients import get_openai_translation, get_claude_translation
                self.log("🧠 DIRECT LLM SEARCH: Successfully imported LLM wrapper functions")
                
                # Use existing LLM infrastructure from Prompt Manager
                api_keys = self.load_api_keys()
                self.log(f"🧠 DIRECT LLM SEARCH: Loaded API keys: {list(api_keys.keys()) if api_keys else 'None'}")
                if not api_keys:
                    self.log("⚠ No API keys found for LLM translation")
                    return
                
                settings = self.load_llm_settings()
                
                # Check if we should use prompts for grid translations (configurable)
                use_prompts_in_grid = settings.get('use_prompts_in_grid', False)  # Default: simple translations
                self.log(f"🧠 LLM Grid Translations - Use Prompts: {use_prompts_in_grid}")
                
                # Get OpenAI translation using existing infrastructure
                if 'openai' in api_keys:
                    try:
                        self.log("🧠 DIRECT LLM SEARCH: Attempting OpenAI translation...")
                        from modules.llm_clients import LLMClient
                        self.log("🧠 DIRECT LLM SEARCH: LLMClient imported for OpenAI")
                        
                        openai_model = settings.get('openai_model', 'gpt-4o')
                        self.log(f"🧠 DIRECT LLM SEARCH: Using OpenAI model: {openai_model}")
                        client = LLMClient(
                            api_key=api_keys['openai'],
                            provider='openai',
                            model=openai_model
                        )
                        self.log("🧠 DIRECT LLM SEARCH: LLMClient initialized for OpenAI")
                        
                        translation = client.translate(
                            text=segment.source,
                            source_lang=source_lang_code or 'nl',
                            target_lang=target_lang_code or 'en'
                        )
                        self.log(f"🧠 DIRECT LLM SEARCH: OpenAI translation result: {translation[:100] if translation else 'None'}...")
                        
                        if translation and translation.strip():
                            match = TranslationMatch(
                                source=segment.source,
                                target=translation.strip(),
                                relevance=88,  # Good confidence for LLM
                                metadata={'model': openai_model, 'provider': 'OpenAI'},
                                match_type="LLM",
                                compare_source=segment.source,
                                provider_code='OA'
                            )
                            matches_dict["LLM"].append(match)
                            self.log(f"✓ Added OpenAI match: {translation[:50]}...")
                        else:
                            self.log("⚠ OpenAI translation was empty or None")
                            
                    except Exception as e:
                        self.log(f"⚠ Error getting OpenAI translation: {e}")
                        import traceback
                        self.log(f"⚠ OpenAI error traceback: {traceback.format_exc()}")
                
                # Get Claude translation using existing infrastructure
                if 'claude' in api_keys:
                    try:
                        from modules.llm_clients import LLMClient
                        
                        claude_model = settings.get('claude_model', 'claude-sonnet-4-5-20250929')
                        client = LLMClient(
                            api_key=api_keys['claude'],
                            provider='claude',
                            model=claude_model
                        )
                        
                        translation = client.translate(
                            text=segment.source,
                            source_lang=source_lang_code or 'nl',
                            target_lang=target_lang_code or 'en'
                        )
                        
                        if translation and translation.strip():
                            match = TranslationMatch(
                                source=segment.source,
                                target=translation.strip(),
                                relevance=90,  # High confidence for Claude
                                metadata={'model': claude_model, 'provider': 'Claude'},
                                match_type="LLM",
                                compare_source=segment.source,
                                provider_code='CL'
                            )
                            matches_dict["LLM"].append(match)
                            self.log(f"✓ Added Claude match: {translation[:50]}...")
                            
                    except Exception as e:
                        self.log(f"⚠ Error getting Claude translation: {e}")
                
                # Get Gemini translation using existing infrastructure
                # Check for both 'gemini' and 'google' keys for backward compatibility
                gemini_api_key = api_keys.get('gemini') or api_keys.get('google')
                if gemini_api_key:
                    try:
                        self.log("🧠 DIRECT LLM SEARCH: Attempting Gemini translation...")
                        from modules.llm_clients import LLMClient
                        self.log("🧠 DIRECT LLM SEARCH: LLMClient imported for Gemini")
                        
                        gemini_model = settings.get('gemini_model', 'gemini-2.5-flash')
                        self.log(f"🧠 DIRECT LLM SEARCH: Using Gemini model: {gemini_model}")
                        client = LLMClient(
                            api_key=gemini_api_key,
                            provider='gemini',
                            model=gemini_model
                        )
                        self.log("🧠 DIRECT LLM SEARCH: LLMClient initialized for Gemini")
                        
                        translation = client.translate(
                            text=segment.source,
                            source_lang=source_lang_code or 'nl',
                            target_lang=target_lang_code or 'en'
                        )
                        
                        if translation and translation.strip():
                            match = TranslationMatch(
                                source=segment.source,
                                target=translation.strip(),
                                relevance=85,  # Good confidence for Gemini
                                metadata={'model': gemini_model, 'provider': 'Gemini'},
                                match_type="LLM",
                                compare_source=segment.source,
                                provider_code='GM'
                            )
                            matches_dict["LLM"].append(match)
                            self.log(f"✓ Added Gemini match: {translation[:50]}...")
                            
                    except Exception as e:
                        self.log(f"⚠ Error getting Gemini translation: {e}")
                    
            except Exception as e:
                self.log(f"⚠ Error in LLM translation process: {e}")
    
    def _clean_provider_prefix(self, translation):
        """Remove provider prefixes like [OpenAI], [Claude], etc. from translations"""
        import re
        # Remove patterns like [OpenAI], [Claude], [ChatGPT], [Anthropic] at the start
        cleaned = re.sub(r'^\[(?:OpenAI|Claude|ChatGPT|Anthropic|GPT-\d+)\]\s*', '', translation.strip())
        return cleaned


# ============================================================================
# UNIVERSAL LOOKUP TAB
# ============================================================================

class UniversalLookupTab(QWidget):
    """
    Universal Lookup - System-wide translation lookup
    Works anywhere on your computer: in CAT tools, browsers, Word, any text box
    """
    
    def __init__(self, parent=None):
        super().__init__(parent)
        
        # Import lookup engine
        try:
            from modules.universal_lookup import UniversalLookupEngine, LookupResult
            self.UniversalLookupEngine = UniversalLookupEngine
            self.LookupResult = LookupResult
        except ImportError:
            QMessageBox.critical(
                self,
                "Missing Module",
                "Could not import universal_lookup module.\nPlease ensure modules/universal_lookup.py exists."
            )
            self.UniversalLookupEngine = None
            self.LookupResult = None
            return
        
        # Initialize engine
        self.engine = None
        self.tm_database = None
        self.hotkey_registered = False
        
        # UI setup
        self.init_ui()
        
        # Register global hotkey
        self.register_global_hotkey()
    
    def init_ui(self):
        """Initialize the UI"""
        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(5)  # Reduced from 10 to 5 for consistency
        
        # Header
        header = QLabel("🔍 Universal Lookup")
        header.setStyleSheet("font-size: 16pt; font-weight: bold; color: #1976D2;")
        layout.addWidget(header, 0)  # 0 = no stretch, stays compact
        
        # Description
        if os.name == 'nt':
            # Windows - full functionality
            description_text = (
                "Look up translations anywhere on your computer.\n"
                "Press Ctrl+Alt+L or paste text manually to search your translation memory."
            )
        else:
            # Mac/Linux - manual mode only
            description_text = (
                "Look up translations in your translation memory.\n"
                "⚠️ Global hotkey not available on this platform. Paste text manually to search."
            )
        
        description = QLabel(description_text)
        description.setWordWrap(True)
        description.setStyleSheet("color: #666; padding: 5px; background-color: #E3F2FD; border-radius: 3px;")
        layout.addWidget(description, 0)  # 0 = no stretch, stays compact
        
        # Mode selector (using label instead of group box)
        mode_label_header = QLabel("⚙️ Operating Mode")
        mode_label_header.setStyleSheet("font-weight: bold; font-size: 10pt; margin-top: 10px;")
        layout.addWidget(mode_label_header, 0)  # 0 = no stretch
        
        mode_layout = QHBoxLayout()
        
        mode_label = QLabel("Mode:")
        mode_layout.addWidget(mode_label)
        
        self.mode_combo = QComboBox()
        self.mode_combo.addItems(["Universal (Any Text Box)", "memoQ", "Trados", "CafeTran"])
        self.mode_combo.currentTextChanged.connect(self.on_mode_changed)
        mode_layout.addWidget(self.mode_combo, stretch=1)
        
        layout.addLayout(mode_layout)
        
        # Source text area (using label instead of group box)
        source_label_header = QLabel("📝 Source Text")
        source_label_header.setStyleSheet("font-weight: bold; font-size: 10pt; margin-top: 10px;")
        layout.addWidget(source_label_header)
        
        self.source_text = QTextEdit()
        self.source_text.setPlaceholderText("Click 'Capture Text' or paste text here to search...")
        self.source_text.setMaximumHeight(100)
        layout.addWidget(self.source_text)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        if os.name == 'nt':
            # Windows - show capture button (though hotkey is preferred)
            capture_btn = QPushButton("📥 Manual Capture")
            capture_btn.setToolTip("Manually trigger text capture (Ctrl+Alt+L is recommended)")
            capture_btn.setStyleSheet("background-color: #4CAF50; color: white; padding: 8px;")
            capture_btn.clicked.connect(self.capture_text)
            button_layout.addWidget(capture_btn)
        
        search_btn = QPushButton("🔍 Search")
        search_btn.setStyleSheet("font-weight: bold; background-color: #2196F3; color: white; padding: 8px;")
        search_btn.clicked.connect(self.perform_lookup)
        button_layout.addWidget(search_btn)
        
        clear_btn = QPushButton("🗑️ Clear")
        clear_btn.clicked.connect(self.clear_all)
        button_layout.addWidget(clear_btn)
        
        layout.addLayout(button_layout)
        
        # Results area (with tabs for TM, termbase, MT)
        self.results_tabs = QTabWidget()
        
        # TM Results tab
        tm_tab = self.create_tm_results_tab()
        self.results_tabs.addTab(tm_tab, "📖 TM Matches")
        
        # Termbase Results tab
        termbase_tab = self.create_termbase_results_tab()
        self.results_tabs.addTab(termbase_tab, "📚 Termbase Terms")
        
        # MT Results tab
        mt_tab = self.create_mt_results_tab()
        self.results_tabs.addTab(mt_tab, "🤖 Machine Translation")
        
        layout.addWidget(self.results_tabs, stretch=1)
        
        # Status bar
        self.status_label = QLabel("Ready. Select a mode and capture text to begin.")
        self.status_label.setStyleSheet("padding: 5px; background-color: #f5f5f5; border-radius: 3px;")
        layout.addWidget(self.status_label, 0)  # 0 = no stretch, stays compact
    
    def create_tm_results_tab(self):
        """Create the TM results tab"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        # Results table
        self.tm_results_table = QTableWidget()
        self.tm_results_table.setColumnCount(4)
        self.tm_results_table.setHorizontalHeaderLabels(["Match %", "Source", "Target", "Type"])
        self.tm_results_table.horizontalHeader().setStretchLastSection(False)
        self.tm_results_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        self.tm_results_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)
        self.tm_results_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.tm_results_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.tm_results_table.doubleClicked.connect(self.on_tm_result_double_click)
        
        layout.addWidget(self.tm_results_table)
        
        # Action buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        copy_btn = QPushButton("📋 Copy Target")
        copy_btn.clicked.connect(self.copy_selected_tm_target)
        button_layout.addWidget(copy_btn)
        
        insert_btn = QPushButton("📥 Insert Target")
        insert_btn.setToolTip("Insert selected translation (Ctrl+V)")
        insert_btn.clicked.connect(self.insert_selected_tm_target)
        button_layout.addWidget(insert_btn)
        
        layout.addLayout(button_layout)
        
        return tab
    
    def create_termbase_results_tab(self):
        """Create the termbase results tab"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        # Results table
        self.termbase_results_table = QTableWidget()
        self.termbase_results_table.setColumnCount(2)
        self.termbase_results_table.setHorizontalHeaderLabels(["Term (Source)", "Translation (Target)"])
        self.termbase_results_table.horizontalHeader().setStretchLastSection(True)
        self.termbase_results_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self.termbase_results_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.termbase_results_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        
        layout.addWidget(self.termbase_results_table)
        
        # Info label
        info = QLabel("💡 Tip: Double-click a term to copy it to clipboard")
        info.setStyleSheet("color: #666; font-size: 9pt; padding: 5px;")
        layout.addWidget(info)
        
        return tab
    
    def create_mt_results_tab(self):
        """Create the MT results tab"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        # Results list
        self.mt_results_layout = QVBoxLayout()
        
        # Placeholder
        placeholder = QLabel("🤖 Machine Translation\n\nComing soon: DeepL, OpenAI, Google Translate integration")
        placeholder.setAlignment(Qt.AlignmentFlag.AlignCenter)
        placeholder.setStyleSheet("color: #999; padding: 40px;")
        self.mt_results_layout.addWidget(placeholder)
        
        layout.addLayout(self.mt_results_layout)
        layout.addStretch()
        
        return tab
    
    def on_mode_changed(self, mode_text):
        """Handle mode change"""
        mode_map = {
            "Universal (Any Text Box)": "universal",
            "memoQ": "memoq",
            "Trados": "trados",
            "CafeTran": "cafetran"
        }
        
        mode = mode_map.get(mode_text, "universal")
        
        if self.engine:
            self.engine.mode = mode
            self.status_label.setText(f"Mode changed to: {mode_text}")
    
    def __del__(self):
        """Destructor - cleanup AHK process when widget is destroyed"""
        try:
            self.unregister_global_hotkey()
        except:
            pass
    
    def capture_text(self, delay=True):
        """
        Capture text from the active application
        
        Args:
            delay: If True, wait 1.5 seconds before capturing (for manual button clicks)
                   If False, capture immediately (for global hotkey)
        """
        if not self.UniversalLookupEngine:
            return
        
        # Initialize engine if needed
        if not self.engine:
            mode_text = self.mode_combo.currentText()
            mode_map = {
                "Universal (Any Text Box)": "universal",
                "memoQ": "memoq",
                "Trados": "trados",
                "CafeTran": "cafetran"
            }
            mode = mode_map.get(mode_text, "universal")
            self.engine = self.UniversalLookupEngine(mode=mode)
            
            # Set TM database if available
            if self.tm_database:
                self.engine.set_tm_database(self.tm_database)
        
        if delay:
            # Manual button click - give time to switch windows
            self.status_label.setText("⏳ Capturing text... Switch to target application now!")
            QApplication.processEvents()
            time.sleep(1.5)
        else:
            # Global hotkey - capture immediately
            self.status_label.setText("⏳ Capturing text...")
            QApplication.processEvents()
        
        # Capture text
        text = self.engine.capture_text()
        
        if text:
            self.source_text.setPlainText(text)
            self.status_label.setText(f"✓ Captured {len(text)} characters. Searching...")
            # Auto-search after capture
            self.perform_lookup()
        else:
            self.status_label.setText("✗ No text captured. Try again.")
    
    def perform_lookup(self):
        """Perform lookup on the source text"""
        text = self.source_text.toPlainText().strip()
        
        if not text:
            self.status_label.setText("⚠️ No text to search. Enter or capture text first.")
            return
        
        if not self.engine:
            # Initialize engine
            mode_text = self.mode_combo.currentText()
            mode_map = {
                "Universal (Any Text Box)": "universal",
                "memoQ": "memoq",
                "Trados": "trados",
                "CafeTran": "cafetran"
            }
            mode = mode_map.get(mode_text, "universal")
            self.engine = UniversalLookupEngine(mode=mode)
            
            # Set TM database if available
            if self.tm_database:
                self.engine.set_tm_database(self.tm_database)
        
        self.status_label.setText("🔍 Searching...")
        QApplication.processEvents()
        
        # Perform lookups
        results = self.engine.lookup_all(text)
        
        # Display TM results
        self.display_tm_results(results.get('tm', []))
        
        # Display termbase results
        self.display_glossary_results(results.get('termbase', []))
        
        # Display MT results
        self.display_mt_results(results.get('mt', []))
        
        total_results = len(results.get('tm', [])) + len(results.get('termbase', [])) + len(results.get('mt', []))
        self.status_label.setText(f"✓ Found {total_results} results")
    
    def display_tm_results(self, results):
        """Display TM results in the table"""
        self.tm_results_table.setRowCount(0)
        
        for result in results:
            row = self.tm_results_table.rowCount()
            self.tm_results_table.insertRow(row)
            
            # Match percentage
            match_item = QTableWidgetItem(f"{result.match_percent}%")
            if result.match_percent == 100:
                match_item.setBackground(QColor("#C8E6C9"))  # Green for exact
            elif result.match_percent >= 95:
                match_item.setBackground(QColor("#FFF9C4"))  # Yellow for high
            self.tm_results_table.setItem(row, 0, match_item)
            
            # Source
            self.tm_results_table.setItem(row, 1, QTableWidgetItem(result.source))
            
            # Target
            self.tm_results_table.setItem(row, 2, QTableWidgetItem(result.target))
            
            # Type
            match_type = result.metadata.get('match_type', 'unknown')
            self.tm_results_table.setItem(row, 3, QTableWidgetItem(match_type))
        
        self.tm_results_table.resizeRowsToContents()
    
    def display_termbase_results(self, results):
        """Display termbase results"""
        self.termbase_results_table.setRowCount(0)
        
        for result in results:
            row = self.termbase_results_table.rowCount()
            self.termbase_results_table.insertRow(row)
            
            self.termbase_results_table.setItem(row, 0, QTableWidgetItem(result.source))
            self.termbase_results_table.setItem(row, 1, QTableWidgetItem(result.target))
        
        self.termbase_results_table.resizeRowsToContents()
    
    def display_mt_results(self, results):
        """Display MT results"""
        # Clear existing
        while self.mt_results_layout.count():
            item = self.mt_results_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        
        if results:
            for result in results:
                label = QLabel(f"{result.metadata.get('provider', 'MT')}: {result.target}")
                label.setWordWrap(True)
                self.mt_results_layout.addWidget(label)
        else:
            placeholder = QLabel("🤖 No MT results yet\n\n(MT integration coming soon)")
            placeholder.setAlignment(Qt.AlignmentFlag.AlignCenter)
            placeholder.setStyleSheet("color: #999; padding: 20px;")
            self.mt_results_layout.addWidget(placeholder)
    
    def on_tm_result_double_click(self, index):
        """Handle double-click on TM result"""
        self.copy_selected_tm_target()
    
    def copy_selected_tm_target(self):
        """Copy selected TM target to clipboard"""
        selected = self.tm_results_table.selectedItems()
        if selected:
            row = selected[0].row()
            target_item = self.tm_results_table.item(row, 2)
            if target_item:
                pyperclip.copy(target_item.text())
                self.status_label.setText(f"✓ Copied to clipboard: {target_item.text()[:50]}...")
    
    def insert_selected_tm_target(self):
        """Insert selected TM target into active application"""
        selected = self.tm_results_table.selectedItems()
        if selected:
            row = selected[0].row()
            target_item = self.tm_results_table.item(row, 2)
            if target_item:
                pyperclip.copy(target_item.text())
                self.status_label.setText("✓ Copied to clipboard. Press Ctrl+V to paste.")
                # Could auto-paste here if we wanted to be aggressive
                # pyautogui.hotkey('ctrl', 'v')
    
    def clear_all(self):
        """Clear all text and results"""
        self.source_text.clear()
        self.tm_results_table.setRowCount(0)
        self.termbase_results_table.setRowCount(0)
        self.status_label.setText("Cleared. Ready for new lookup.")
    
    def set_tm_database(self, tm_db):
        """Set TM database (called from main window)"""
        self.tm_database = tm_db
        if self.engine:
            self.engine.set_tm_database(tm_db)
    
    def register_global_hotkey(self):
        """Register global hotkey for Universal Lookup"""
        global _ahk_process
        try:
            # Kill any existing instances of the AHK script first
            if os.name == 'nt':
                try:
                    # Use multiple methods to ensure cleanup
                    # Method 1: Kill by window title
                    subprocess.run(['taskkill', '/F', '/FI', 'WINDOWTITLE eq universal_lookup_hotkey.ahk*'],
                                 capture_output=True, creationflags=subprocess.CREATE_NO_WINDOW)
                    
                    # Method 2: Kill AutoHotkey processes more aggressively
                    subprocess.run(['taskkill', '/F', '/IM', 'AutoHotkey.exe'],
                                 capture_output=True, creationflags=subprocess.CREATE_NO_WINDOW)
                    
                    # Method 3: Kill by process name pattern
                    import psutil
                    try:
                        for proc in psutil.process_iter(['pid', 'name', 'cmdline']):
                            try:
                                if 'universal_lookup_hotkey' in ' '.join(proc.cmdline() or []):
                                    proc.kill()
                            except:
                                pass
                    except:
                        pass
                except:
                    pass
            
            # Find AutoHotkey executable
            username = os.environ.get('USERNAME', '')
            ahk_paths = [
                r"C:\Program Files\AutoHotkey\v2\AutoHotkey.exe",
                r"C:\Program Files\AutoHotkey\v2\AutoHotkey64.exe",
                r"C:\Program Files\AutoHotkey\AutoHotkey.exe",
                r"C:\Program Files (x86)\AutoHotkey\AutoHotkey.exe",
                fr"C:\Users\{username}\AppData\Local\Programs\AutoHotkey\AutoHotkey.exe"
            ]
            
            ahk_exe = None
            for path in ahk_paths:
                if os.path.exists(path):
                    ahk_exe = path
                    break
            
            if not ahk_exe:
                print("[Universal Lookup] AutoHotkey not found. Please install from https://www.autohotkey.com/")
                print("[Universal Lookup] Searched paths:", ahk_paths)
                self.hotkey_registered = False
                return
            
            ahk_script = Path(__file__).parent / "universal_lookup_hotkey.ahk"
            if ahk_script.exists():
                # Start AHK script in background (hidden)
                self.ahk_process = subprocess.Popen([ahk_exe, str(ahk_script)],
                                                   creationflags=subprocess.CREATE_NO_WINDOW if os.name == 'nt' else 0)
                # Store in global variable for atexit cleanup
                _ahk_process = self.ahk_process
                print(f"[Universal Lookup] AHK hotkey registered: Ctrl+Alt+L")
                
                # Start file watcher
                self.start_file_watcher()
                self.hotkey_registered = True
            else:
                print(f"[Universal Lookup] AHK script not found: {ahk_script}")
                self.hotkey_registered = False
        except Exception as e:
            print(f"[Universal Lookup] Could not start AHK hotkey: {e}")
            self.hotkey_registered = False
    
    def start_file_watcher(self):
        """Watch for signal file from AHK"""
        self.signal_file = Path(__file__).parent / "lookup_signal.txt"
        self.capture_file = Path(__file__).parent / "temp_capture.txt"
        
        # Create timer to check for signal file
        self.file_check_timer = QTimer()
        self.file_check_timer.timeout.connect(self.check_for_signal)
        self.file_check_timer.start(100)  # Check every 100ms
    
    def check_for_signal(self):
        """Check if AHK wrote a signal file"""
        if self.signal_file.exists():
            try:
                # Delete signal file
                self.signal_file.unlink()
                
                # Get text from clipboard (AHK already copied it)
                time.sleep(0.1)  # Give clipboard a moment
                text = pyperclip.paste()
                
                # Trigger lookup
                if text:
                    self.on_ahk_capture(text)
            except Exception as e:
                print(f"[Universal Lookup] Error reading capture: {e}")
    
    def on_ahk_capture(self, text):
        """Handle text captured by AHK"""
        try:
            # Bring Supervertaler to foreground
            main_window = self.window()
            if main_window:
                # Check if window was maximized before restoring
                was_maximized = main_window.isMaximized()
                
                # Restore if minimized or hidden
                if main_window.isMinimized():
                    main_window.showNormal()
                elif main_window.isHidden():
                    main_window.show()
                else:
                    # If already visible, just activate
                    main_window.show()
                
                # Restore maximized state if it was maximized
                if was_maximized:
                    main_window.showMaximized()
                    
                # Bring to front and activate
                main_window.raise_()
                main_window.activateWindow()
                
            # Show universal lookup dialog
            self.show_universal_lookup(text)
            
        except Exception as e:
            print(f"[Universal Lookup] Error handling capture: {e}")
    
    def show_universal_lookup(self, text):
        """Show universal lookup with pre-filled text"""
        try:
            # Switch to Universal Lookup tab
            if hasattr(self, 'right_tabs'):
                # Find Universal Lookup tab
                for i in range(self.right_tabs.count()):
                    if "Universal Lookup" in self.right_tabs.tabText(i):
                        self.right_tabs.setCurrentIndex(i)
                        break
            
            # Fill in text if Universal Lookup tab exists
            universal_lookup_tab = getattr(self, 'universal_lookup_tab', None)
            if universal_lookup_tab and hasattr(universal_lookup_tab, 'input_text'):
                universal_lookup_tab.input_text.setPlainText(text)
                # Trigger lookup
                if hasattr(universal_lookup_tab, 'lookup_button'):
                    universal_lookup_tab.lookup_button.click()
                    
        except Exception as e:
            print(f"[Universal Lookup] Error showing lookup: {e}")




# ============================================================================
# AUTOFINGERS DIALOG
# ============================================================================

class CheckmarkCheckBox(QCheckBox):
    """Custom checkbox with green background and white checkmark when checked"""
    
    def __init__(self, text="", parent=None):
        super().__init__(text, parent)
        self.setCheckable(True)
        self.setEnabled(True)
        self.setStyleSheet("""
            QCheckBox {
                font-size: 9pt;
                spacing: 6px;
            }
            QCheckBox::indicator {
                width: 18px;
                height: 18px;
                border: 2px solid #999;
                border-radius: 3px;
                background-color: white;
            }
            QCheckBox::indicator:checked {
                background-color: #4CAF50;
                border-color: #4CAF50;
            }
            QCheckBox::indicator:hover {
                border-color: #666;
            }
            QCheckBox::indicator:checked:hover {
                background-color: #45a049;
                border-color: #45a049;
            }
        """)
    
    def paintEvent(self, event):
        """Override paint event to draw white checkmark when checked"""
        super().paintEvent(event)
        
        if self.isChecked():
            # Get the indicator rectangle using QStyle
            from PyQt6.QtWidgets import QStyleOptionButton
            from PyQt6.QtGui import QPainter, QPen, QColor
            from PyQt6.QtCore import QPointF, QRect
            
            opt = QStyleOptionButton()
            self.initStyleOption(opt)
            indicator_rect = self.style().subElementRect(
                self.style().SubElement.SE_CheckBoxIndicator,
                opt,
                self
            )
            
            if indicator_rect.isValid():
                # Draw white checkmark
                painter = QPainter(self)
                painter.setRenderHint(QPainter.RenderHint.Antialiasing)
                pen_width = max(2.0, min(indicator_rect.width(), indicator_rect.height()) * 0.12)
                painter.setPen(QPen(QColor(255, 255, 255), pen_width, Qt.PenStyle.SolidLine, Qt.PenCapStyle.RoundCap, Qt.PenJoinStyle.RoundJoin))
                painter.setBrush(QColor(255, 255, 255))
                
                # Draw checkmark (✓ shape) - coordinates relative to indicator
                x = indicator_rect.x()
                y = indicator_rect.y()
                w = indicator_rect.width()
                h = indicator_rect.height()
                
                # Add padding (15% on all sides)
                padding = min(w, h) * 0.15
                x += padding
                y += padding
                w -= padding * 2
                h -= padding * 2
                
                # Checkmark path
                check_x1 = x + w * 0.10
                check_y1 = y + h * 0.50
                check_x2 = x + w * 0.35
                check_y2 = y + h * 0.70
                check_x3 = x + w * 0.90
                check_y3 = y + h * 0.25
                
                # Draw two lines forming the checkmark
                painter.drawLine(QPointF(check_x2, check_y2), QPointF(check_x3, check_y3))
                painter.drawLine(QPointF(check_x1, check_y1), QPointF(check_x2, check_y2))
                
                painter.end()


class CustomRadioButton(QRadioButton):
    """Custom radio button with square indicator, green when checked, white checkmark"""
    
    def __init__(self, text="", parent=None):
        super().__init__(text, parent)
        self.setCheckable(True)
        self.setEnabled(True)
        self.setStyleSheet("""
            QRadioButton {
                font-size: 9pt;
                spacing: 6px;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
                border: 2px solid #999;
                border-radius: 3px;
                background-color: white;
            }
            QRadioButton::indicator:checked {
                background-color: #4CAF50;
                border-color: #4CAF50;
            }
            QRadioButton::indicator:hover {
                border-color: #666;
            }
            QRadioButton::indicator:checked:hover {
                background-color: #45a049;
                border-color: #45a049;
            }
        """)
    
    def paintEvent(self, event):
        """Override paint event to draw white checkmark when checked"""
        super().paintEvent(event)
        
        if self.isChecked():
            # Get the indicator rectangle using QStyle
            from PyQt6.QtWidgets import QStyleOptionButton
            from PyQt6.QtGui import QPainter, QPen, QColor
            from PyQt6.QtCore import QPointF, QRect
            
            opt = QStyleOptionButton()
            self.initStyleOption(opt)
            indicator_rect = self.style().subElementRect(
                self.style().SubElement.SE_RadioButtonIndicator,
                opt,
                self
            )
            
            if indicator_rect.isValid():
                # Draw white checkmark
                painter = QPainter(self)
                painter.setRenderHint(QPainter.RenderHint.Antialiasing)
                # Slightly thinner pen for better fit on smaller displays
                pen_width = max(2.0, min(indicator_rect.width(), indicator_rect.height()) * 0.12)
                painter.setPen(QPen(QColor(255, 255, 255), pen_width, Qt.PenStyle.SolidLine, Qt.PenCapStyle.RoundCap, Qt.PenJoinStyle.RoundJoin))
                painter.setBrush(QColor(255, 255, 255))
                
                # Draw checkmark (✓ shape) - coordinates relative to indicator
                # Add padding to prevent clipping on smaller displays
                x = indicator_rect.x()
                y = indicator_rect.y()
                w = indicator_rect.width()
                h = indicator_rect.height()
                
                # Add padding (15% on all sides) to ensure checkmark doesn't get cut off on smaller displays
                padding = min(w, h) * 0.15
                x += padding
                y += padding
                w -= padding * 2
                h -= padding * 2
                
                # Checkmark path: bottom-left to middle, then middle to top-right
                # Using proportions that create a nice checkmark shape with proper padding
                check_x1 = x + w * 0.10  # Left point (more padding from left)
                check_y1 = y + h * 0.50  # Bottom point (centered vertically)
                check_x2 = x + w * 0.35  # Middle-bottom point
                check_y2 = y + h * 0.70  # Bottom point (with padding from bottom)
                check_x3 = x + w * 0.90  # Right point (more padding from right)
                check_y3 = y + h * 0.25  # Top point (with padding from top)
                
                # Draw two lines forming the checkmark
                painter.drawLine(QPointF(check_x2, check_y2), QPointF(check_x3, check_y3))
                painter.drawLine(QPointF(check_x1, check_y1), QPointF(check_x2, check_y2))
                
                painter.end()


class AutoFingersWidget(QWidget):
    """
    AutoFingers - CAT Tool Automation Widget
    Provides UI for translation automation in tools like memoQ
    Now integrated as a tab in the main Supervertaler window
    """
    
    def __init__(self, parent=None):
        super().__init__(parent)
        
        # Import AutoFingers engine
        try:
            from modules.autofingers_engine import AutoFingersEngine
            self.AutoFingersEngine = AutoFingersEngine
        except ImportError:
            QMessageBox.critical(
                self,
                "Import Error",
                "Could not import AutoFingers engine.\n"
                "Make sure modules/autofingers_engine.py exists."
            )
            return
        
        # Initialize engine
        self.engine = None
        self.is_running = False
        
        # Get default TMX path from user_data
        if ENABLE_PRIVATE_FEATURES:
            default_tmx = "user_data_private/autofingers_tm.tmx"
        else:
            default_tmx = "user_data/autofingers_tm.tmx"
        
        self.tmx_file = default_tmx
        
        self.setup_ui()
        self.load_settings()
        
    def setup_ui(self):
        """Setup the user interface"""
        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(5)
        
        # Header (matches Universal Lookup / PDF Rescue / TMX Editor style)
        header = QLabel("🤖 AutoFingers")
        header.setStyleSheet("font-size: 16pt; font-weight: bold; color: #1976D2;")
        layout.addWidget(header, 0)  # 0 = no stretch, stays compact
        
        # Description box (matches Universal Lookup / PDF Rescue / TMX Editor style)
        info = QLabel(
            "Automated Translation Pasting for memoQ.\n"
            "AutoFingers reads from a TMX file and pastes translations automatically."
        )
        info.setWordWrap(True)
        info.setStyleSheet("color: #666; padding: 5px; background-color: #E3F2FD; border-radius: 3px;")
        layout.addWidget(info, 0)  # 0 = no stretch, stays compact
        
        # Tabs
        tabs = QTabWidget()
        
        # TAB 1: Control Panel (with integrated settings)
        control_tab = self.create_control_tab()
        tabs.addTab(control_tab, "🎮 Control Panel")
        
        # TAB 2: TMX Manager
        tmx_tab = self.create_tmx_tab()
        tabs.addTab(tmx_tab, "📚 TMX Manager")
        
        layout.addWidget(tabs, 1)  # 1 = stretch to fill space
        
        # Close button
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(self.close)
        button_layout.addWidget(close_btn)
        
        layout.addLayout(button_layout, 0)  # 0 = no stretch, stays compact
        
        # Setup global keyboard shortcuts (matching AutoHotkey version)
        self.setup_shortcuts()
    
    def setup_shortcuts(self):
        """Setup GLOBAL keyboard shortcuts for AutoFingers actions"""
        import keyboard
        
        try:
            # Store hotkey references for later removal
            self.hotkeys = []
            
            # Register global hotkeys that work even when memoQ has focus
            # Ctrl+Alt+P - Process single segment
            self.hotkeys.append(keyboard.add_hotkey('ctrl+alt+p', self.process_single_safe))
            
            # Ctrl+Shift+L - Toggle loop mode
            self.hotkeys.append(keyboard.add_hotkey('ctrl+shift+l', self.toggle_loop_safe))
            
            # Ctrl+Alt+S - Stop loop
            self.hotkeys.append(keyboard.add_hotkey('ctrl+alt+s', self.stop_loop_safe))
            
            # Ctrl+Alt+R - Reload TMX
            self.hotkeys.append(keyboard.add_hotkey('ctrl+alt+r', self.reload_tmx_safe))
            
            self.log("✓ Global hotkeys registered: Ctrl+Alt+P (single), Ctrl+Shift+L (loop), Ctrl+Alt+S (stop), Ctrl+Alt+R (reload)")
            self.log("ℹ️ Hotkeys work globally - even when memoQ has focus!")
            
        except Exception as e:
            self.log(f"⚠️ Could not register global hotkeys: {str(e)}")
            self.log("ℹ️ You may need to run as Administrator for global hotkeys")
    
    def process_single_safe(self):
        """Safe wrapper for process_single (called from global hotkey - runs in hotkey thread)"""
        try:
            # Dispatch to main Qt thread to avoid threading issues
            QTimer.singleShot(0, self.process_single)
        except Exception as e:
            print(f"Error in process_single_safe: {e}")
    
    def toggle_loop_safe(self):
        """Safe wrapper for toggle_loop (called from global hotkey - runs in hotkey thread)"""
        try:
            # Dispatch to main Qt thread to avoid threading issues
            QTimer.singleShot(0, self.toggle_loop)
        except Exception as e:
            print(f"Error in toggle_loop_safe: {e}")
    
    def stop_loop_safe(self):
        """Safe wrapper for stop_loop (called from global hotkey)"""
        try:
            self.stop_loop()
        except Exception as e:
            print(f"Error in stop_loop: {e}")
    
    def reload_tmx_safe(self):
        """Safe wrapper for reload_tmx (called from global hotkey)"""
        try:
            self.reload_tmx()
        except Exception as e:
            print(f"Error in reload_tmx: {e}")
    
    def cleanup_hotkeys(self):
        """Cleanup AutoFingers hotkeys when widget is closed/hidden"""
        # Unregister ONLY AutoFingers hotkeys
        try:
            import keyboard
            if hasattr(self, 'hotkeys'):
                for hotkey in self.hotkeys:
                    try:
                        keyboard.remove_hotkey(hotkey)
                    except:
                        pass
                self.log("AutoFingers hotkeys unregistered")
        except Exception as e:
            print(f"Error unregistering hotkeys: {e}")
    
    def create_control_tab(self):
        """Create the main control panel with horizontal layout"""
        tab = QWidget()
        main_layout = QVBoxLayout(tab)
        main_layout.setContentsMargins(8, 8, 8, 8)
        main_layout.setSpacing(8)
        
        # TOP ROW: Actions and TMX File side by side
        top_row = QHBoxLayout()
        top_row.setSpacing(8)
        
        # === LEFT: ACTIONS GROUP ===
        actions_group = QGroupBox("🎯 Actions")
        actions_layout = QHBoxLayout()
        actions_layout.setSpacing(8)
        
        # Single button
        single_col = QVBoxLayout()
        single_btn = QPushButton("▶️ Single")
        single_btn.setMinimumHeight(40)
        single_btn.setMinimumWidth(100)
        single_btn.setStyleSheet("font-size: 10pt; font-weight: bold;")
        single_btn.clicked.connect(self.process_single)
        single_col.addWidget(single_btn)
        single_info = QLabel("(Ctrl+Alt+P)")
        single_info.setStyleSheet("color: #666; font-size: 8pt;")
        single_info.setAlignment(Qt.AlignmentFlag.AlignCenter)
        single_col.addWidget(single_info)
        actions_layout.addLayout(single_col)
        
        # Loop button
        loop_col = QVBoxLayout()
        self.loop_btn = QPushButton("▶ Loop Mode")
        self.loop_btn.setMinimumHeight(40)
        self.loop_btn.setMinimumWidth(110)
        self.loop_btn.setStyleSheet("font-size: 10pt; font-weight: bold; background-color: #4CAF50; color: white;")
        self.loop_btn.clicked.connect(self.toggle_loop)
        loop_col.addWidget(self.loop_btn)
        loop_info = QLabel("(Ctrl+Shift+L)")
        loop_info.setStyleSheet("color: #666; font-size: 8pt;")
        loop_info.setAlignment(Qt.AlignmentFlag.AlignCenter)
        loop_col.addWidget(loop_info)
        actions_layout.addLayout(loop_col)
        
        # Segments spin
        segs_col = QVBoxLayout()
        segs_label = QLabel("Segments:")
        segs_label.setStyleSheet("font-size: 9pt; font-weight: bold;")
        segs_col.addWidget(segs_label)
        self.loop_segments_spin = QSpinBox()
        self.loop_segments_spin.setMinimum(0)
        self.loop_segments_spin.setMaximum(9999)
        self.loop_segments_spin.setValue(0)
        self.loop_segments_spin.setSuffix(" (0=∞)")
        self.loop_segments_spin.setMinimumHeight(28)
        self.loop_segments_spin.setStyleSheet("font-size: 9pt;")
        segs_col.addWidget(self.loop_segments_spin)
        actions_layout.addLayout(segs_col)
        
        # Progress
        progress_col = QVBoxLayout()
        self.progress_label = QLabel("Ready")
        self.progress_label.setStyleSheet("font-weight: bold; font-size: 9pt;")
        progress_col.addWidget(self.progress_label)
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setMinimumHeight(28)
        progress_col.addWidget(self.progress_bar)
        actions_layout.addLayout(progress_col)
        
        actions_group.setLayout(actions_layout)
        top_row.addWidget(actions_group, 2)
        
        # === RIGHT: TMX FILE GROUP ===
        tmx_group = QGroupBox("📁 TMX File")
        tmx_layout = QVBoxLayout()
        tmx_layout.setSpacing(4)
        
        # File path row
        path_row = QHBoxLayout()
        self.tmx_path_label = QLabel(self.tmx_file)
        self.tmx_path_label.setStyleSheet("padding: 4px; background-color: #F5F5F5; font-size: 8pt;")
        path_row.addWidget(self.tmx_path_label, 1)
        
        browse_btn = QPushButton("Browse")
        browse_btn.setMaximumWidth(70)
        browse_btn.clicked.connect(self.browse_tmx)
        path_row.addWidget(browse_btn)
        
        reload_btn = QPushButton("Reload")
        reload_btn.setMaximumWidth(70)
        reload_btn.clicked.connect(self.reload_tmx)
        path_row.addWidget(reload_btn)
        
        tmx_layout.addLayout(path_row)
        
        # Status
        self.tmx_status_label = QLabel("No TMX loaded")
        self.tmx_status_label.setStyleSheet("padding: 4px; font-weight: bold; font-size: 9pt;")
        tmx_layout.addWidget(self.tmx_status_label)
        
        tmx_group.setLayout(tmx_layout)
        top_row.addWidget(tmx_group, 1)
        
        main_layout.addLayout(top_row)
        
        # MIDDLE ROW: Settings in a 2-column grid layout for better organization
        settings_group = QGroupBox("⚙️ Settings")
        settings_layout = QGridLayout()
        settings_layout.setSpacing(15)
        settings_layout.setColumnStretch(0, 1)  # Left column flexible
        settings_layout.setColumnStretch(1, 1)  # Right column flexible
        
        # LEFT COLUMN: Languages and Timing (input fields grouped together)
        left_col = QVBoxLayout()
        left_col.setSpacing(12)
        
        # Languages section
        lang_group = QVBoxLayout()
        lang_label = QLabel("Languages:")
        lang_label.setStyleSheet("font-weight: bold; font-size: 9pt;")
        lang_group.addWidget(lang_label)
        lang_inputs = QHBoxLayout()
        lang_inputs.addWidget(QLabel("Source:"))
        self.source_lang_edit = QLineEdit("en")
        self.source_lang_edit.setMaximumWidth(50)
        lang_inputs.addWidget(self.source_lang_edit)
        lang_inputs.addSpacing(10)
        lang_inputs.addWidget(QLabel("Target:"))
        self.target_lang_edit = QLineEdit("nl")
        self.target_lang_edit.setMaximumWidth(50)
        lang_inputs.addWidget(self.target_lang_edit)
        lang_inputs.addStretch()
        lang_group.addLayout(lang_inputs)
        left_col.addLayout(lang_group)
        
        # Timing section
        timing_group = QVBoxLayout()
        timing_label = QLabel("Timing (ms):")
        timing_label.setStyleSheet("font-weight: bold; font-size: 9pt;")
        timing_group.addWidget(timing_label)
        timing_inputs = QHBoxLayout()
        timing_inputs.addWidget(QLabel("Loop:"))
        self.loop_delay_spin = QSpinBox()
        self.loop_delay_spin.setRange(500, 10000)
        self.loop_delay_spin.setValue(4000)
        self.loop_delay_spin.setSuffix(" ms")
        self.loop_delay_spin.setMaximumWidth(90)
        timing_inputs.addWidget(self.loop_delay_spin)
        timing_inputs.addSpacing(10)
        timing_inputs.addWidget(QLabel("Confirm:"))
        self.confirm_delay_spin = QSpinBox()
        self.confirm_delay_spin.setRange(100, 5000)
        self.confirm_delay_spin.setValue(900)
        self.confirm_delay_spin.setSuffix(" ms")
        self.confirm_delay_spin.setMaximumWidth(90)
        timing_inputs.addWidget(self.confirm_delay_spin)
        timing_inputs.addStretch()
        timing_group.addLayout(timing_inputs)
        left_col.addLayout(timing_group)
        
        left_col.addStretch()  # Push content to top
        
        # RIGHT COLUMN: Behavior checkboxes and Save button
        right_col = QVBoxLayout()
        right_col.setSpacing(12)
        
        # Behavior section
        behavior_group = QVBoxLayout()
        behavior_label = QLabel("Behavior:")
        behavior_label.setStyleSheet("font-weight: bold; font-size: 9pt;")
        behavior_group.addWidget(behavior_label)
        
        # Use custom checkboxes with green background and white checkmark when checked
        self.auto_confirm_check = CheckmarkCheckBox("Confirm segments")
        self.auto_confirm_check.setChecked(True)
        self.auto_confirm_check.setToolTip("When checked: Confirm segment with Ctrl+Enter before moving to next. When unchecked: Move to next with Alt+N without confirming")
        behavior_group.addWidget(self.auto_confirm_check)
        self.skip_no_match_check = CheckmarkCheckBox("Skip no match")
        self.skip_no_match_check.setChecked(True)
        behavior_group.addWidget(self.skip_no_match_check)
        right_col.addLayout(behavior_group)

        # Tag Cleaning section
        tag_cleaning_group = QVBoxLayout()
        tag_cleaning_label = QLabel("Tag Cleaning:")
        tag_cleaning_label.setStyleSheet("font-weight: bold; font-size: 9pt; margin-top: 8px;")
        tag_cleaning_group.addWidget(tag_cleaning_label)

        # Master switch
        self.tag_cleaning_enabled_check = CheckmarkCheckBox("Enable tag cleaning")
        self.tag_cleaning_enabled_check.setChecked(False)
        self.tag_cleaning_enabled_check.setToolTip("Remove CAT tool tags from translations before pasting")
        tag_cleaning_group.addWidget(self.tag_cleaning_enabled_check)

        # Granular controls (indented)
        tag_types_layout = QVBoxLayout()
        tag_types_layout.setContentsMargins(20, 0, 0, 0)  # Indent

        self.clean_memoq_index_tags_check = CheckmarkCheckBox("memoQ index tags ([1} {2])")
        self.clean_memoq_index_tags_check.setChecked(True)
        self.clean_memoq_index_tags_check.setToolTip("Remove memoQ index tags like [1} {2] [3} etc.")
        tag_types_layout.addWidget(self.clean_memoq_index_tags_check)

        self.clean_trados_tags_check = CheckmarkCheckBox("Trados Studio tags")
        self.clean_trados_tags_check.setChecked(False)
        self.clean_trados_tags_check.setEnabled(False)  # Not implemented yet
        self.clean_trados_tags_check.setToolTip("Coming soon")
        tag_types_layout.addWidget(self.clean_trados_tags_check)

        self.clean_cafetran_tags_check = CheckmarkCheckBox("CafeTran tags")
        self.clean_cafetran_tags_check.setChecked(False)
        self.clean_cafetran_tags_check.setEnabled(False)  # Not implemented yet
        self.clean_cafetran_tags_check.setToolTip("Coming soon")
        tag_types_layout.addWidget(self.clean_cafetran_tags_check)

        self.clean_wordfast_tags_check = CheckmarkCheckBox("Wordfast tags")
        self.clean_wordfast_tags_check.setChecked(False)
        self.clean_wordfast_tags_check.setEnabled(False)  # Not implemented yet
        self.clean_wordfast_tags_check.setToolTip("Coming soon")
        tag_types_layout.addWidget(self.clean_wordfast_tags_check)

        tag_cleaning_group.addLayout(tag_types_layout)
        right_col.addLayout(tag_cleaning_group)
        
        # Save button - centered at bottom of right column
        right_col.addStretch()
        save_btn = QPushButton("💾 Save Settings")
        save_btn.setMinimumHeight(35)
        save_btn.setStyleSheet("font-size: 9pt; font-weight: bold;")
        save_btn.clicked.connect(self.save_settings)
        right_col.addWidget(save_btn)
        
        # Add columns to grid layout
        left_widget = QWidget()
        left_widget.setLayout(left_col)

        # Wrap right column in scroll area for smaller screens
        right_widget = QWidget()
        right_widget.setLayout(right_col)

        right_scroll = QScrollArea()
        right_scroll.setWidget(right_widget)
        right_scroll.setWidgetResizable(True)
        right_scroll.setFrameShape(QFrame.Shape.NoFrame)
        right_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        right_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        settings_layout.addWidget(left_widget, 0, 0)
        settings_layout.addWidget(right_scroll, 0, 1)
        
        settings_group.setLayout(settings_layout)
        
        # Create horizontal container for Settings and Activity Log side-by-side
        settings_log_row = QHBoxLayout()
        settings_log_row.setSpacing(15)
        settings_log_row.addWidget(settings_group, 1)  # Settings takes flexible space
        
        # Activity Log on the right
        log_group = QGroupBox("📋 Activity Log")
        log_layout = QVBoxLayout()
        log_layout.setContentsMargins(5, 5, 5, 5)
        
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMinimumHeight(200)  # Give it a reasonable minimum height
        self.log_text.setStyleSheet("""
            font-family: 'Consolas', monospace;
            padding: 4px;
            line-height: 1.3;
            font-size: 8pt;
        """)
        log_layout.addWidget(self.log_text)
        
        log_group.setLayout(log_layout)
        settings_log_row.addWidget(log_group, 1)  # Activity Log takes equal flexible space
        
        main_layout.addLayout(settings_log_row)
        
        return tab
    
    def create_tmx_tab(self):
        """Create the TMX manager tab"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        # Info
        info = QLabel(
            "Manage your AutoFingers translation memory file.\n"
            "The TMX file stores source-target translation pairs."
        )
        info.setWordWrap(True)
        layout.addWidget(info)
        
        # Stats
        self.tmx_stats_label = QLabel("No TMX loaded")
        self.tmx_stats_label.setStyleSheet("padding: 10px; background-color: #F5F5F5; font-weight: bold;")
        layout.addWidget(self.tmx_stats_label)
        
        # Actions
        btn_layout = QVBoxLayout()
        
        create_btn = QPushButton("➕ Create New Empty TMX")
        create_btn.clicked.connect(self.create_empty_tmx)
        btn_layout.addWidget(create_btn)
        
        open_btn = QPushButton("📂 Open TMX in Editor")
        open_btn.clicked.connect(self.open_tmx_in_editor)
        btn_layout.addWidget(open_btn)
        
        import_btn = QPushButton("📥 Import from Supervertaler TM")
        import_btn.clicked.connect(self.import_from_tm)
        btn_layout.addWidget(import_btn)
        
        layout.addLayout(btn_layout)
        layout.addStretch()
        
        return tab
    
    def log(self, message: str):
        """Add message to activity log (thread-safe)"""
        # Skip logging if not on main thread to prevent QTextDocument crashes
        from threading import current_thread, main_thread
        if current_thread() is main_thread():
            timestamp = datetime.now().strftime("%H:%M:%S")
            self._append_log(f"[{timestamp}] {message}")
        else:
            # Queue for main thread
            timestamp = datetime.now().strftime("%H:%M:%S")
            QTimer.singleShot(0, lambda msg=f"[{timestamp}] {message}": self._append_log(msg))
    
    def _append_log(self, message: str):
        """Actually append to log (runs on main thread)"""
        self.log_text.append(message)
        # Auto-scroll to bottom
        scrollbar = self.log_text.verticalScrollBar()
        if scrollbar:
            scrollbar.setValue(scrollbar.maximum())
    
    def update_progress_label(self, text: str):
        """Thread-safe progress label update"""
        QTimer.singleShot(0, lambda t=text: self.progress_label.setText(t))
    
    def update_progress_bar(self, value: int):
        """Thread-safe progress bar update"""
        QTimer.singleShot(0, lambda v=value: self.progress_bar.setValue(v))
    
    def browse_tmx(self):
        """Browse for TMX file"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select TMX File",
            "",
            "TMX Files (*.tmx);;All Files (*.*)"
        )
        
        if file_path:
            self.tmx_file = file_path
            self.tmx_path_label.setText(file_path)
            self.reload_tmx()
    
    def reload_tmx(self):
        """Reload TMX file"""
        try:
            self.engine = self.AutoFingersEngine(
                tmx_file=self.tmx_file,
                source_lang=self.source_lang_edit.text(),
                target_lang=self.target_lang_edit.text()
            )
            
            # Apply settings
            self.engine.loop_delay = self.loop_delay_spin.value()
            self.engine.confirm_delay = self.confirm_delay_spin.value()
            self.engine.auto_confirm = self.auto_confirm_check.isChecked()
            self.engine.skip_no_match = self.skip_no_match_check.isChecked()

            # Apply tag cleaner settings
            if self.tag_cleaning_enabled_check.isChecked():
                self.engine.tag_cleaner.enable()
            else:
                self.engine.tag_cleaner.disable()

            if self.clean_memoq_index_tags_check.isChecked():
                self.engine.tag_cleaner.enable_memoq_index_tags()
            else:
                self.engine.tag_cleaner.disable_memoq_index_tags()

            success, message = self.engine.load_tmx()
            
            if success:
                self.log(f"✓ {message}")
                self.tmx_status_label.setText(f"✓ {message}")
                self.tmx_status_label.setStyleSheet("padding: 10px; font-weight: bold; color: green;")
                self.tmx_stats_label.setText(
                    f"📊 TMX Statistics:\n"
                    f"  • Translation Units: {self.engine.tm_count}\n"
                    f"  • Source Language: {self.engine.source_lang}\n"
                    f"  • Target Language: {self.engine.target_lang}"
                )
            else:
                self.log(f"✗ {message}")
                self.tmx_status_label.setText(f"✗ {message}")
                self.tmx_status_label.setStyleSheet("padding: 10px; font-weight: bold; color: red;")
                
        except Exception as e:
            self.log(f"✗ Error: {str(e)}")
            QMessageBox.critical(self, "Error", f"Failed to load TMX:\n{str(e)}")
    
    def sync_settings_to_engine(self):
        """Sync current UI settings to the engine"""
        if self.engine:
            self.engine.loop_delay = self.loop_delay_spin.value()
            self.engine.confirm_delay = self.confirm_delay_spin.value()
            self.engine.auto_confirm = self.auto_confirm_check.isChecked()
            self.engine.skip_no_match = self.skip_no_match_check.isChecked()
    
    def process_single(self):
        """Process a single segment"""
        if not self.engine:
            QMessageBox.warning(self, "No TMX", "Please load a TMX file first.")
            return
        
        # Sync current settings to engine before processing
        self.sync_settings_to_engine()
        
        self.log("▶️ Processing single segment...")
        self.progress_label.setText("Processing...")
        
        # Give user time to switch to memoQ
        QApplication.processEvents()
        import time
        time.sleep(1)
        
        try:
            success, message = self.engine.process_single_segment()
            
            # Log match result
            if self.engine.last_source and self.engine.last_match:
                self.log(f"✓ Match: {self.engine.last_match.match_type} ({self.engine.last_match.match_percent}%)")
            else:
                self.log(f"⚠️ No match found for segment")
            
            if success:
                self.log(f"✓ {message}")
                self.progress_label.setText("✓ Segment processed successfully")
            else:
                self.log(f"✗ {message}")
                self.progress_label.setText(f"✗ {message}")
        except Exception as e:
            self.log(f"❌ Error in process_single: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
    
    def toggle_loop(self):
        """Toggle loop mode on/off"""
        if self.is_running:
            # Stop loop
            self.engine.stop()
            self.is_running = False
            self.loop_btn.setText("▶ Loop Mode")
            self.loop_btn.setStyleSheet("font-size: 9pt; font-weight: bold; background-color: #4CAF50; color: white; padding: 2px;")
            self.progress_bar.setVisible(False)
            self.log("⏹️ Loop mode stopped")
        else:
            # Start loop
            if not self.engine:
                QMessageBox.warning(self, "No TMX", "Please load a TMX file first.")
                return
            
            # Sync current settings to engine before starting loop
            self.sync_settings_to_engine()
            
            max_segments = self.loop_segments_spin.value()
            self.is_running = True
            self.loop_btn.setText("⏹ Stop Loop")
            self.loop_btn.setStyleSheet("font-size: 9pt; font-weight: bold; background-color: #F44336; color: white; padding: 2px;")
            
            if max_segments > 0:
                self.progress_bar.setMaximum(max_segments)
                self.progress_bar.setValue(0)
                self.progress_bar.setVisible(True)
            else:
                self.progress_bar.setVisible(False)
            
            self.log(f"▶️ Starting loop mode ({max_segments if max_segments > 0 else '∞'} segments)...")
            
            # Start loop in background thread
            self.loop_thread = threading.Thread(target=self.run_loop, args=(max_segments,), daemon=True)
            self.loop_thread.start()
    
    def run_loop(self, max_segments):
        """Run the loop mode in background thread"""
        import time
        
        segment_count = 0
        
        while self.is_running:
            # Check if reached limit
            if max_segments > 0 and segment_count >= max_segments:
                self.is_running = False
                self.log(f"✓ Completed {segment_count} segments")
                self.update_progress_label(f"✓ Completed {segment_count} segments")
                QTimer.singleShot(0, self.reset_loop_ui)
                break
            
            # Process one segment
            try:
                success, message = self.engine.process_single_segment()
                
                if success:
                    segment_count += 1
                    self.log(f"✓ {message}")
                    self.update_progress_label(f"Processing... ({segment_count} completed)")
                    if max_segments > 0:
                        self.update_progress_bar(segment_count)
                else:
                    self.log(f"✗ {message}")
                    
                    # Stop if no match and not skipping
                    if not self.engine.skip_no_match:
                        self.is_running = False
                        self.update_progress_label(f"Stopped - no translation found")
                        QTimer.singleShot(0, self.reset_loop_ui)
                        break
                
                # Wait between segments
                if self.is_running:
                    time.sleep(self.engine.loop_delay / 1000)
                    
            except Exception as e:
                self.log(f"✗ Error: {str(e)}")
                self.is_running = False
                QTimer.singleShot(0, self.reset_loop_ui)
                break
    
    def reset_loop_ui(self):
        """Reset UI after loop stops"""
        self.loop_btn.setText("▶ Loop Mode")
        self.loop_btn.setStyleSheet("font-size: 9pt; font-weight: bold; background-color: #4CAF50; color: white; padding: 2px;")
        if self.progress_bar.maximum() > 0:
            self.progress_bar.setVisible(True)  # Keep visible to show final progress
        else:
            self.progress_bar.setVisible(False)
    
    def stop_loop(self):
        """Stop loop mode (separate method for keyboard shortcut)"""
        if self.is_running:
            self.toggle_loop()  # Reuse toggle logic when running
        else:
            self.log("Loop mode is not running")
    
    def create_empty_tmx(self):
        """Create a new empty TMX file"""
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Create New TMX File",
            self.tmx_file,
            "TMX Files (*.tmx)"
        )
        
        if file_path:
            temp_engine = self.AutoFingersEngine(
                tmx_file=file_path,
                source_lang=self.source_lang_edit.text(),
                target_lang=self.target_lang_edit.text()
            )
            
            if temp_engine.create_empty_tmx():
                self.log(f"✓ Created empty TMX: {file_path}")
                self.tmx_file = file_path
                self.tmx_path_label.setText(file_path)
                self.reload_tmx()
            else:
                QMessageBox.critical(self, "Error", "Failed to create TMX file")
    
    def open_tmx_in_editor(self):
        """Open TMX file in external editor"""
        import subprocess
        try:
            if sys.platform == 'win32':
                os.startfile(self.tmx_file)
            elif sys.platform == 'darwin':
                subprocess.call(['open', self.tmx_file])
            else:
                subprocess.call(['xdg-open', self.tmx_file])
            self.log(f"📂 Opened TMX in editor: {self.tmx_file}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Could not open TMX file:\n{str(e)}")
    
    def import_from_tm(self):
        """Import translations from Supervertaler TM database"""
        QMessageBox.information(
            self,
            "Feature Coming Soon",
            "TM import functionality will be added in the next update!\n\n"
            "For now, you can manually edit the TMX file or use the\n"
            "AutoHotkey version to populate your translation memory."
        )
    
    def load_settings(self):
        """Load saved settings from file"""
        try:
            settings_file = Path("user_data_private" if ENABLE_PRIVATE_FEATURES else "user_data") / "autofingers_settings.json"
            if settings_file.exists():
                with open(settings_file, 'r', encoding='utf-8') as f:
                    settings = json.load(f)
                
                # Apply settings to UI
                self.loop_delay_spin.setValue(settings.get('loop_delay', 4000))
                self.confirm_delay_spin.setValue(settings.get('confirm_delay', 900))
                self.auto_confirm_check.setChecked(settings.get('auto_confirm', True))
                self.skip_no_match_check.setChecked(settings.get('skip_no_match', True))
                # Backward compatibility: if old settings had use_down_arrow=True, set auto_confirm=False
                if settings.get('use_down_arrow', False):
                    self.auto_confirm_check.setChecked(False)

                # Load tag cleaner settings
                tag_cleaner_settings = settings.get('tag_cleaner', {})
                self.tag_cleaning_enabled_check.setChecked(tag_cleaner_settings.get('enabled', False))
                memoq_settings = tag_cleaner_settings.get('memoq', {})
                self.clean_memoq_index_tags_check.setChecked(
                    memoq_settings.get('index_tags', {}).get('enabled', True)
                )

                self.log("✓ Settings loaded")
        except Exception as e:
            self.log(f"⚠️ Could not load settings: {e}")
    
    def save_settings(self):
        """Save current settings to file"""
        try:
            settings = {
                'loop_delay': self.loop_delay_spin.value(),
                'confirm_delay': self.confirm_delay_spin.value(),
                'auto_confirm': self.auto_confirm_check.isChecked(),
                'skip_no_match': self.skip_no_match_check.isChecked(),
                'tag_cleaner': {
                    'enabled': self.tag_cleaning_enabled_check.isChecked(),
                    'memoq': {
                        'index_tags': {
                            'enabled': self.clean_memoq_index_tags_check.isChecked()
                        }
                    },
                    'trados': {},
                    'cafetran': {},
                    'wordfast': {}
                }
            }

            settings_file = Path("user_data_private" if ENABLE_PRIVATE_FEATURES else "user_data") / "autofingers_settings.json"
            settings_file.parent.mkdir(parents=True, exist_ok=True)

            with open(settings_file, 'w', encoding='utf-8') as f:
                json.dump(settings, f, indent=2)

            self.log("💾 Settings saved")
            QMessageBox.information(self, "Saved", "Settings saved successfully!")
        except Exception as e:
            self.log(f"⚠️ Could not save settings: {e}")
            QMessageBox.critical(self, "Error", f"Failed to save settings:\n{str(e)}")


# ============================================================================
# APPLICATION ENTRY POINT
# ============================================================================

def main():
    """Application entry point"""
    app = QApplication(sys.argv)
    app.setApplicationName("Supervertaler Qt")
    app.setOrganizationName("Supervertaler")
    
    window = SupervertalerQt()
    window.show()
    
    sys.exit(app.exec())


if __name__ == '__main__':
    main()
