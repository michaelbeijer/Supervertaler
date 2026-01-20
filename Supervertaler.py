"""   
Supervertaler
=============
The Ultimate Translation Workbench.
Modern PyQt6 interface with specialised modules to handle any problem.
Version: 1.9.136 (Glossary matching fix for punctuation)
Release Date: January 20, 2026
Framework: PyQt6

This is the modern edition of Supervertaler using PyQt6 framework.
For the classic tkinter edition, see Supervertaler_tkinter.py

Key Features:
- Complete Translation Matching: Termbase + TM + MT + Multi-LLM
- Project Termbases: Dedicated terminology per project with automatic extraction
- Supervoice: AI-powered voice dictation (100+ languages via OpenAI Whisper)
- Superimage: Extract images from DOCX files with preview
- Google Cloud Translation API integration
- Multi-LLM Support: OpenAI GPT, Claude, Google Gemini
- 2-Layer Prompt Architecture (System + Custom Prompts) with AI Assistant
- AI Assistant with conversational interface for document analysis
- Superlookup with global hotkey (Ctrl+Alt+L)
- Detachable Log window for multi-monitor setups
- Modern theme system (6 themes + custom editor)
- AutoFingers automation for memoQ with TagCleaner module
- memoQ bilingual DOCX import/export
- Bilingual Table export/import for review workflow
- SQLite-based translation memory with FTS5 search
- Professional TMX editor
- Spellcheck integration with Hunspell and custom dictionary support

Author: Michael Beijer
License: MIT
"""

# Version Information.
__version__ = "1.9.145"
__phase__ = "0.9"
__release_date__ = "2026-01-20"
__edition__ = "Qt"

import sys
import json
import os
import subprocess
import atexit
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple, Callable
from dataclasses import dataclass, asdict
from datetime import datetime


def get_resource_path(relative_path: str) -> Path:
    """Get absolute path to resource, works for dev and for PyInstaller bundled app."""
    if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
        # Running as compiled executable (PyInstaller)
        base_path = Path(sys._MEIPASS)
    else:
        # Running in development
        base_path = Path(__file__).parent
    return base_path / relative_path


def get_user_data_path() -> Path:
    """
    Get the path to user data directory.
    
    In frozen builds (EXE): 'user_data' folder next to the EXE
    In development: 'user_data' folder next to the script
    
    The build process copies user_data directly next to the EXE for easy access.
    """
    if getattr(sys, 'frozen', False):
        # Frozen build: user_data is next to the EXE (copied by build script)
        return Path(sys.executable).parent / "user_data"
    else:
        # Development: user_data next to script
        return Path(__file__).parent / "user_data"

import threading
import time  # For delays in Superlookup
import re

# Fix encoding for Windows console (UTF-8 support)
# Only set if stdout/stderr exist (they're None in PyInstaller --windowed mode)
if sys.platform == 'win32':
    import io
    if sys.stdout is not None and hasattr(sys.stdout, 'buffer'):
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    if sys.stderr is not None and hasattr(sys.stderr, 'buffer'):
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

# External dependencies
import pyperclip  # For clipboard operations in Superlookup
from modules.superlookup import SuperlookupEngine  # Superlookup engine
from modules.voice_dictation_lite import QuickDictationThread  # Voice dictation
from modules.voice_commands import VoiceCommandManager, VoiceCommand, ContinuousVoiceListener  # Voice commands (Talon-style)
from modules.statuses import (
    STATUSES,
    DEFAULT_STATUS,
    StatusDefinition,
    get_status,
    match_memoq_status,
    compose_memoq_status,
)
from modules import file_dialog_helper as fdh  # File dialog helper with last directory memory
from modules.spellcheck_manager import SpellcheckManager, get_spellcheck_manager  # Spellcheck with Hunspell
from modules.find_replace_qt import (
    FindReplaceHistory,
    FindReplaceOperation,
    FindReplaceSet,
    FindReplaceSetsManager,
    HistoryComboBox,
)  # F&R History and Sets
from modules.shortcut_manager import ShortcutManager  # Keyboard shortcut management


STATUS_ORDER = [
    "not_started",
    "pretranslated",
    "translated",
    "confirmed",
    "tr_confirmed",
    "proofread",
    "approved",
    "rejected",
]

STATUS_CYCLE = STATUS_ORDER

TRANSLATABLE_STATUSES = {"not_started", "pretranslated", "translated"}

# Check for PyQt6 and offer to install if missing
try:
    from PyQt6.QtWidgets import (
        QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
        QTableWidget, QTableWidgetItem, QTableWidgetSelectionRange, QHeaderView, QMenuBar, QMenu,
        QFileDialog, QMessageBox, QToolBar, QLabel, QComboBox,
        QPushButton, QSpinBox, QSplitter, QTextEdit, QStatusBar,
        QStyledItemDelegate, QInputDialog, QDialog, QLineEdit, QRadioButton,
        QButtonGroup, QDialogButtonBox, QTabWidget, QGroupBox, QGridLayout, QCheckBox,
        QProgressBar, QProgressDialog, QFormLayout, QTabBar, QPlainTextEdit, QAbstractItemDelegate,
        QFrame, QListWidget, QListWidgetItem, QStackedWidget, QTreeWidget, QTreeWidgetItem,
        QScrollArea, QSizePolicy, QSlider, QToolButton, QAbstractItemView
    )
    from PyQt6.QtCore import Qt, QSize, QTimer, pyqtSignal, QObject, QUrl
    from PyQt6.QtGui import QFont, QAction, QKeySequence, QIcon, QTextOption, QColor, QDesktopServices, QTextCharFormat, QTextCursor, QBrush, QSyntaxHighlighter, QPalette, QTextBlockFormat, QCursor, QFontMetrics
    from PyQt6.QtWidgets import QStyleOptionViewItem, QStyle
    from PyQt6.QtCore import QRectF
    from PyQt6.QtNetwork import QNetworkAccessManager, QNetworkRequest, QNetworkReply
    from PyQt6 import sip
except ImportError:
    print("PyQt6 not found. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "PyQt6"])
    print("✓ PyQt6 installed. Please restart the application.")
    sys.exit(0)


# ============================================================================
# PRIVATE DATA PROTECTION SYSTEM
# ============================================================================

# Check for .supervertaler.local file to enable private features (for developers only)
# This ensures that private data (API keys, personal projects, etc.) stays in user_data_private/
# which is .gitignored, preventing accidental upload to GitHub
ENABLE_PRIVATE_FEATURES = os.path.exists(
    os.path.join(os.path.dirname(os.path.abspath(__file__)), ".supervertaler.local")
)
if ENABLE_PRIVATE_FEATURES:
    print("[DEV MODE] Private features enabled (.supervertaler.local found)")
    print("[DEV MODE] Using 'user_data_private/' folder (git-ignored)")
else:
    print("[USER MODE] Using 'user_data/' folder")


# ============================================================================
# GLOBAL AHK PROCESS CLEANUP
# ============================================================================

# Global variable to track AHK process
_ahk_process = None

def cleanup_ahk_process():
    """Cleanup function to kill AHK process on exit"""
    global _ahk_process
    if _ahk_process:
        try:
            _ahk_process.terminate()
            _ahk_process.wait(timeout=1)
            print("[Hotkeys] AHK process terminated on exit")
        except:
            try:
                _ahk_process.kill()
                print("[Hotkeys] AHK process killed on exit")
            except:
                pass

# Register cleanup function to run on Python exit
atexit.register(cleanup_ahk_process)


# ============================================================================
# ============================================================================
# DATA MODELS
# ============================================================================

# ============================================================================
# INLINE FORMATTING TAG UTILITIES
# ============================================================================

def runs_to_tagged_text(paragraphs) -> str:
    """
    Convert Word document paragraphs with run formatting to HTML-tagged text.
    
    Args:
        paragraphs: List of python-docx Paragraph objects
        
    Returns:
        String with HTML tags for formatting (e.g., "<b>bold</b> normal <i>italic</i>")
    """
    result_parts = []
    
    for paragraph in paragraphs:
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # Determine which tags to apply
            is_bold = run.bold == True
            is_italic = run.italic == True
            is_underline = run.underline == True
            
            # Build tagged text
            if is_bold or is_italic or is_underline:
                # Open tags (order: bold, italic, underline)
                if is_bold:
                    text = f"<b>{text}"
                if is_italic:
                    text = f"<i>{text}" if not is_bold else text.replace("<b>", "<b><i>", 1)
                if is_underline:
                    if is_bold and is_italic:
                        text = text.replace("<b><i>", "<b><i><u>", 1)
                    elif is_bold:
                        text = text.replace("<b>", "<b><u>", 1)
                    elif is_italic:
                        text = text.replace("<i>", "<i><u>", 1)
                    else:
                        text = f"<u>{text}"
                
                # Close tags (reverse order: underline, italic, bold)
                if is_underline:
                    text = f"{text}</u>"
                if is_italic:
                    text = f"{text}</i>"
                if is_bold:
                    text = f"{text}</b>"
            
            result_parts.append(text)
    
    return ''.join(result_parts)


def strip_formatting_tags(text: str) -> str:
    """
    Remove HTML formatting tags from text, leaving plain text.
    
    Args:
        text: Text with HTML tags like <b>, </b>, <i>, </i>, <u>, </u>
        
    Returns:
        Plain text without tags
    """
    import re
    # Remove <b>, </b>, <i>, </i>, <u>, </u> tags
    return re.sub(r'</?[biu]>', '', text)


def has_formatting_tags(text: str) -> bool:
    """
    Check if text contains any formatting tags.
    
    Args:
        text: Text to check
        
    Returns:
        True if text contains <b>, <i>, or <u> tags
    """
    import re
    return bool(re.search(r'</?[biu]>', text))


def apply_formatting_tags(text: str, tag: str) -> str:
    """
    Wrap text with the specified formatting tag.
    
    Args:
        text: Text to wrap
        tag: Tag name ('b', 'i', or 'u')
        
    Returns:
        Tagged text like "<b>text</b>"
    """
    if tag in ('b', 'i', 'u'):
        return f"<{tag}>{text}</{tag}>"
    return text


def get_formatted_html_display(text: str) -> str:
    """
    Convert our simple tags to HTML for rich text display.
    
    Args:
        text: Text with <b>, <i>, <u> tags
        
    Returns:
        HTML string suitable for QTextEdit.setHtml()
    """
    # Escape HTML entities first (except our tags)
    import html
    
    # Temporarily replace our tags with placeholders
    text = text.replace('<b>', '\x00B_OPEN\x00')
    text = text.replace('</b>', '\x00B_CLOSE\x00')
    text = text.replace('<i>', '\x00I_OPEN\x00')
    text = text.replace('</i>', '\x00I_CLOSE\x00')
    text = text.replace('<u>', '\x00U_OPEN\x00')
    text = text.replace('</u>', '\x00U_CLOSE\x00')
    
    # Escape other HTML
    text = html.escape(text)
    
    # Restore our tags as real HTML
    text = text.replace('\x00B_OPEN\x00', '<b>')
    text = text.replace('\x00B_CLOSE\x00', '</b>')
    text = text.replace('\x00I_OPEN\x00', '<i>')
    text = text.replace('\x00I_CLOSE\x00', '</i>')
    text = text.replace('\x00U_OPEN\x00', '<u>')
    text = text.replace('\x00U_CLOSE\x00', '</u>')
    
    return text


def tagged_text_to_runs(text: str) -> list:
    """
    Parse text with HTML formatting tags and return a list of runs with formatting info.
    
    Args:
        text: Text with <b>, <i>, <u> tags (can be nested)
        
    Returns:
        List of dicts: [{'text': str, 'bold': bool, 'italic': bool, 'underline': bool}, ...]
    
    Example:
        "Hello <b>bold</b> and <i>italic</i> world"
        -> [{'text': 'Hello ', 'bold': False, 'italic': False, 'underline': False},
            {'text': 'bold', 'bold': True, 'italic': False, 'underline': False},
            {'text': ' and ', 'bold': False, 'italic': False, 'underline': False},
            {'text': 'italic', 'bold': False, 'italic': True, 'underline': False},
            {'text': ' world', 'bold': False, 'italic': False, 'underline': False}]
    """
    import re
    
    runs = []
    
    # Track current formatting state
    is_bold = False
    is_italic = False
    is_underline = False
    
    # Pattern to match opening/closing tags
    tag_pattern = re.compile(r'(</?[biu]>)')
    
    # Split text by tags, keeping the tags as delimiters
    parts = tag_pattern.split(text)
    
    current_text = ""
    
    for part in parts:
        if part == '<b>':
            # Save any accumulated text before changing state
            if current_text:
                runs.append({
                    'text': current_text,
                    'bold': is_bold,
                    'italic': is_italic,
                    'underline': is_underline
                })
                current_text = ""
            is_bold = True
        elif part == '</b>':
            if current_text:
                runs.append({
                    'text': current_text,
                    'bold': is_bold,
                    'italic': is_italic,
                    'underline': is_underline
                })
                current_text = ""
            is_bold = False
        elif part == '<i>':
            if current_text:
                runs.append({
                    'text': current_text,
                    'bold': is_bold,
                    'italic': is_italic,
                    'underline': is_underline
                })
                current_text = ""
            is_italic = True
        elif part == '</i>':
            if current_text:
                runs.append({
                    'text': current_text,
                    'bold': is_bold,
                    'italic': is_italic,
                    'underline': is_underline
                })
                current_text = ""
            is_italic = False
        elif part == '<u>':
            if current_text:
                runs.append({
                    'text': current_text,
                    'bold': is_bold,
                    'italic': is_italic,
                    'underline': is_underline
                })
                current_text = ""
            is_underline = True
        elif part == '</u>':
            if current_text:
                runs.append({
                    'text': current_text,
                    'bold': is_bold,
                    'italic': is_italic,
                    'underline': is_underline
                })
                current_text = ""
            is_underline = False
        else:
            # Regular text - accumulate it
            current_text += part
    
    # Don't forget any remaining text
    if current_text:
        runs.append({
            'text': current_text,
            'bold': is_bold,
            'italic': is_italic,
            'underline': is_underline
        })
    
    return runs


# ============================================================================
# MEMOQ TAG UTILITIES (for tag insertion shortcuts)
# ============================================================================

def extract_memoq_tags(text: str) -> list:
    """
    Extract all memoQ-style tags from text in order of appearance.
    
    memoQ uses several tag types:
    - Paired opening tags: [1}, [2}, [3} etc.
    - Paired closing tags: {1], {2], {3] etc.
    - Standalone tags: [1], [2], [3] etc. (e.g., for tabs, special characters)
    
    Args:
        text: Source text containing tags
        
    Returns:
        List of tag strings in order of appearance: ['[1}', '{1]', '[2]', ...]
    """
    import re
    # Match:
    # - Opening paired tags: [N}
    # - Closing paired tags: {N]
    # - Standalone tags: [N]
    pattern = r'(\[\d+\}|\{\d+\]|\[\d+\])'
    return re.findall(pattern, text)


def extract_html_tags(text: str) -> list:
    """
    Extract all HTML/XML tags from text in order of appearance.
    
    Supports common formatting tags used in translation:
    - Opening tags: <b>, <i>, <u>, <li>, <p>, <span>, etc.
    - Closing tags: </b>, </i>, </u>, </li>, </p>, </span>, etc.
    - Self-closing tags: <br/>, <hr/>, etc.
    
    Args:
        text: Source text containing HTML tags
        
    Returns:
        List of tag strings in order of appearance: ['<li>', '</li>', '<b>', '</b>', ...]
    """
    import re
    # Match HTML/XML tags: <tagname>, </tagname>, <tagname/>, <tagname attr="value">
    pattern = r'(</?[a-zA-Z][a-zA-Z0-9]*(?:\s+[^>]*)?>)'
    return re.findall(pattern, text)


def extract_all_tags(text: str) -> list:
    """
    Extract all tags (memoQ and HTML) from text in order of appearance.
    
    Args:
        text: Source text containing tags
        
    Returns:
        List of all tag strings in order of appearance
    """
    import re
    # Combined pattern for both memoQ tags and HTML tags
    # memoQ: [N}, {N], [N]
    # HTML: <tag>, </tag>, <tag/>, <tag attr="value"> - includes hyphenated tags like li-o, li-b
    pattern = r'(\[\d+\}|\{\d+\]|\[\d+\]|</?[a-zA-Z][a-zA-Z0-9-]*(?:\s+[^>]*)?>)'
    return re.findall(pattern, text)


def count_pipe_symbols(text: str) -> int:
    """Count the number of CafeTran pipe symbols in text."""
    return text.count('|')


def get_next_pipe_count_needed(source_text: str, target_text: str) -> int:
    """
    Get how many more pipe symbols are needed in target to match source.
    
    Args:
        source_text: Source segment text with pipe symbols
        target_text: Current target text
        
    Returns:
        Number of additional pipe symbols needed (0 if target has enough or more)
    """
    source_pipes = count_pipe_symbols(source_text)
    target_pipes = count_pipe_symbols(target_text)
    return max(0, source_pipes - target_pipes)


def get_tag_pair(tag_number: int) -> tuple:
    """
    Get opening and closing tag pair for a given number.
    
    Args:
        tag_number: The tag number (1, 2, 3, etc.)
        
    Returns:
        Tuple of (opening_tag, closing_tag) e.g. ('[1}', '{1]')
    """
    return (f'[{tag_number}}}', f'{{{tag_number}]')


def find_next_unused_tag(source_text: str, target_text: str) -> str:
    """
    Find the next tag from source that hasn't been used in target yet.
    
    Supports both memoQ tags ([1}, {1], [1]) and HTML tags (<li>, </li>, <b>, etc.)
    
    Args:
        source_text: Source segment text with tags
        target_text: Current target text (may have some tags already)
        
    Returns:
        The next tag to insert, or empty string if all tags are used
    """
    # Use combined extraction for both memoQ and HTML 	
    source_tags = extract_all_tags(source_text)
    target_tags = extract_all_tags(target_text)
    
    # Count occurrences in target
    from collections import Counter
    target_tag_counts = Counter(target_tags)
    source_tag_counts = Counter(source_tags)
    
    # Find first tag that needs more occurrences in target
    for tag in source_tags:
        source_count = source_tag_counts[tag]
        target_count = target_tag_counts.get(tag, 0)
        if target_count < source_count:
            return tag
    
    return ""  # All tags already in target


def get_wrapping_tag_pair(source_text: str, target_text: str) -> tuple:
    """
    Get the next available tag pair for wrapping selected text.
    
    Finds the first tag number from source where either opening or closing
    tag is not yet in target.
    
    Args:
        source_text: Source segment text with tags
        target_text: Current target text
        
    Returns:
        Tuple of (opening_tag, closing_tag) or (None, None) if no pairs available
    """
    import re
    
    # Extract all tag numbers from source
    source_tags = extract_memoq_tags(source_text)
    if not source_tags:
        return (None, None)
    
    # Get unique tag numbers in order
    tag_numbers = []
    for tag in source_tags:
        match = re.search(r'\d+', tag)
        if match:
            num = int(match.group())
            if num not in tag_numbers:
                tag_numbers.append(num)
    
    target_tags = extract_memoq_tags(target_text)
    
    # Find first tag number where pair is not complete in target
    for num in tag_numbers:
        opening, closing = get_tag_pair(num)
        if opening not in target_tags or closing not in target_tags:
            return (opening, closing)
    
    return (None, None)


@dataclass
class Segment:
    """Translation segment (matches tkinter version format)"""
    id: int
    source: str
    target: str = ""
    status: str = DEFAULT_STATUS.key
    type: str = "para"  # para, heading, list_item, table_cell
    notes: str = ""  # Stored as “Comments” in UI
    match_percent: Optional[int] = None  # memoQ match score if provided
    memoQ_status: str = ""  # Raw memoQ status text
    locked: bool = False  # For compatibility with tkinter version
    paragraph_id: int = 0  # Group segments by paragraph for document flow
    style: str = "Normal"  # Heading 1, Heading 2, Title, Subtitle, Normal, etc.
    document_position: int = 0  # Position in original document
    is_table_cell: bool = False  # Whether this segment is in a table
    table_info: Optional[tuple] = None  # (table_idx, row_idx, cell_idx) if is_table_cell
    modified: bool = False  # Track if segment has been edited
    created_at: str = ""  # Creation timestamp
    modified_at: str = ""  # Last modification timestamp
    list_number: Optional[int] = None  # For numbered lists: 1, 2, 3, etc. None for bullets or non-lists
    list_type: str = ""  # "numbered", "bullet", or "" for non-list items
    file_id: Optional[int] = None  # ID of the file this segment belongs to (for multi-file projects)
    file_name: str = ""  # Name of the file this segment belongs to (for multi-file projects)
    dejavu_segment_id: str = ""  # Déjà Vu segment ID for round-trip export
    dejavu_row_index: Optional[int] = None  # Déjà Vu row index for export mapping
    
    def __post_init__(self):
        """Initialize timestamps if not provided"""
        if not self.created_at:
            self.created_at = datetime.now().isoformat()
        if not self.modified_at:
            self.modified_at = datetime.now().isoformat()
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for JSON serialization"""
        return asdict(self)
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'Segment':
        """Create Segment from dictionary, ignoring unknown fields"""
        # Only use fields that the dataclass knows about
        valid_fields = {f.name for f in cls.__dataclass_fields__.values()}
        filtered_data = {k: v for k, v in data.items() if k in valid_fields}
        return cls(**filtered_data)


@dataclass
class Project:
    """Translation project"""
    name: str
    source_lang: str = "en"
    target_lang: str = "nl"
    segments: List[Segment] = None
    created: str = ""
    modified: str = ""
    prompt_settings: Dict[str, Any] = None  # Store active prompt settings
    tm_settings: Dict[str, Any] = None  # Store activated TM settings
    termbase_settings: Dict[str, Any] = None  # Store activated termbase settings
    nt_settings: Dict[str, Any] = None  # Store activated non-translatables settings
    spellcheck_settings: Dict[str, Any] = None  # Store spellcheck settings {enabled, language}
    ui_settings: Dict[str, Any] = None  # Store UI settings (results pane zoom, etc.)
    general_settings_overrides: Dict[str, Any] = None  # Per-project overrides for general settings
    id: int = None  # Unique project ID for TM activation tracking
    original_docx_path: str = None  # Path to original DOCX for structure-preserving export
    trados_source_path: str = None  # Path to original Trados bilingual DOCX for round-trip export
    memoq_source_path: str = None  # Path to original memoQ bilingual DOCX for round-trip export
    mqxliff_source_path: str = None  # Path to original memoQ XLIFF for round-trip export
    cafetran_source_path: str = None  # Path to original CafeTran bilingual DOCX for round-trip export
    sdlppx_source_path: str = None  # Path to original Trados SDLPPX package for SDLRPX export
    original_txt_path: str = None  # Path to original simple text file for round-trip export
    dejavu_source_path: str = None  # Path to original Déjà Vu bilingual RTF for round-trip export
    concordance_geometry: Dict[str, int] = None  # Window geometry for Concordance Search {x, y, width, height}
    # Multi-file project support
    files: List[Dict[str, Any]] = None  # List of files in project: [{id, name, path, type, segment_count, ...}]
    is_multifile: bool = False  # True if this is a multi-file project
    
    def __post_init__(self):
        if self.segments is None:
            self.segments = []
        if self.files is None:
            self.files = []
        if not self.created:
            self.created = datetime.now().isoformat()
        if not self.modified:
            self.modified = datetime.now().isoformat()
        if self.prompt_settings is None:
            self.prompt_settings = {}
        if self.tm_settings is None:
            self.tm_settings = {}
        if self.termbase_settings is None:
            self.termbase_settings = {}
        if self.nt_settings is None:
            self.nt_settings = {}
        if self.spellcheck_settings is None:
            self.spellcheck_settings = {}
        if self.ui_settings is None:
            self.ui_settings = {}
        if self.general_settings_overrides is None:
            self.general_settings_overrides = {}
        # Generate ID if not set (for backward compatibility with old projects)
        if self.id is None:
            import hashlib
            # Create stable ID from project name + created timestamp
            id_source = f"{self.name}_{self.created}"
            self.id = int(hashlib.md5(id_source.encode()).hexdigest()[:8], 16)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for JSON serialization.
        
        Structure is organized for human readability when opened in a text editor:
        1. Project identification (name, languages, dates, id)
        2. Settings (prompts, TM, termbases, spellcheck, etc.)
        3. Source file paths
        4. UI state (concordance geometry)
        5. Segments (at the end - the actual translation content)
        """
        # Start with core project metadata
        result = {
            'name': self.name,
            'source_lang': self.source_lang,
            'target_lang': self.target_lang,
            'created': self.created,
            'modified': self.modified,
            'id': self.id  # Save project ID
        }
        
        # Add settings (prompts, TM, termbases, etc.)
        if hasattr(self, 'prompt_settings'):
            result['prompt_settings'] = self.prompt_settings
        if self.tm_settings:
            result['tm_settings'] = self.tm_settings
        if self.termbase_settings:
            result['termbase_settings'] = self.termbase_settings
        if self.nt_settings:
            result['nt_settings'] = self.nt_settings
        if self.spellcheck_settings:
            result['spellcheck_settings'] = self.spellcheck_settings
        if self.ui_settings:
            result['ui_settings'] = self.ui_settings
        
        # Add source file paths
        if self.original_docx_path:
            result['original_docx_path'] = self.original_docx_path
        if self.trados_source_path:
            result['trados_source_path'] = self.trados_source_path
        if self.memoq_source_path:
            result['memoq_source_path'] = self.memoq_source_path
        if self.mqxliff_source_path:
            result['mqxliff_source_path'] = self.mqxliff_source_path
        if self.cafetran_source_path:
            result['cafetran_source_path'] = self.cafetran_source_path
        if self.sdlppx_source_path:
            result['sdlppx_source_path'] = self.sdlppx_source_path
        if self.original_txt_path:
            result['original_txt_path'] = self.original_txt_path
        if self.dejavu_source_path:
            result['dejavu_source_path'] = self.dejavu_source_path
        
        # Add UI state
        if self.concordance_geometry:
            result['concordance_geometry'] = self.concordance_geometry
        
        # Add multi-file project data
        if self.is_multifile:
            result['is_multifile'] = self.is_multifile
        if self.files:
            result['files'] = self.files
        
        # Add segments LAST (so they appear at the end of the file)
        result['segments'] = [seg.to_dict() for seg in self.segments]
        
        return result
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'Project':
        """Create Project from dictionary"""
        segments = [Segment.from_dict(seg) for seg in data.get('segments', [])]
        
        # Handle missing name field (use filename or default)
        name = data.get('name', 'Untitled Project')
        
        project = cls(
            name=name,
            source_lang=data.get('source_lang', 'en'),
            target_lang=data.get('target_lang', 'nl'),
            segments=segments,
            created=data.get('created', ''),
            modified=data.get('modified', ''),
            id=data.get('id', None)  # Load project ID (will auto-generate if missing)
        )
        # Store prompt settings if they exist
        if 'prompt_settings' in data:
            project.prompt_settings = data['prompt_settings']
        # Store TM settings if they exist
        if 'tm_settings' in data:
            project.tm_settings = data['tm_settings']
        # Store termbase settings if they exist
        if 'termbase_settings' in data:
            project.termbase_settings = data['termbase_settings']
        # Store non-translatables settings if they exist
        if 'nt_settings' in data:
            project.nt_settings = data['nt_settings']
        # Store spellcheck settings if they exist
        if 'spellcheck_settings' in data:
            project.spellcheck_settings = data['spellcheck_settings']
        # Store UI settings if they exist
        if 'ui_settings' in data:
            project.ui_settings = data['ui_settings']
        # Store original DOCX path if it exists
        if 'original_docx_path' in data:
            project.original_docx_path = data['original_docx_path']
        # Store Trados source path if it exists
        if 'trados_source_path' in data:
            project.trados_source_path = data['trados_source_path']
        # Store memoQ source path if it exists
        if 'memoq_source_path' in data:
            project.memoq_source_path = data['memoq_source_path']
        # Store memoQ XLIFF source path if it exists
        if 'mqxliff_source_path' in data:
            project.mqxliff_source_path = data['mqxliff_source_path']
        # Store CafeTran source path if it exists
        if 'cafetran_source_path' in data:
            project.cafetran_source_path = data['cafetran_source_path']
        # Store SDLPPX source path if it exists
        if 'sdlppx_source_path' in data:
            project.sdlppx_source_path = data['sdlppx_source_path']
        # Store original TXT path if it exists
        if 'original_txt_path' in data:
            project.original_txt_path = data['original_txt_path']
        # Store Déjà Vu source path if it exists
        if 'dejavu_source_path' in data:
            project.dejavu_source_path = data['dejavu_source_path']
        # Store concordance window geometry if it exists
        if 'concordance_geometry' in data:
            project.concordance_geometry = data['concordance_geometry']
        # Store multi-file project data if it exists
        if 'is_multifile' in data:
            project.is_multifile = data['is_multifile']
        if 'files' in data:
            project.files = data['files']
        return project


# ============================================================================
# CUSTOM DELEGATES AND EDITORS
# ============================================================================


class _CtrlReturnEventFilter(QObject):
    """App-level event filter to catch Ctrl+Return/Ctrl+Enter for grid confirm.

    This is a workaround for cases where the main keyboard Return key (Key_Return)
    is swallowed before it reaches the source cell editor or QShortcut.
    """

    def __init__(self, main_window):
        super().__init__(main_window)
        self._main_window = main_window

    def eventFilter(self, obj, event):
        from PyQt6.QtCore import QEvent
        from PyQt6.QtWidgets import QApplication

        if event.type() != QEvent.Type.KeyPress:
            return False

        # Only handle Ctrl+Return/Ctrl+Enter (ignore Shift/Alt variants)
        mods = event.modifiers()
        if not (mods & Qt.KeyboardModifier.ControlModifier):
            return False
        if mods & (Qt.KeyboardModifier.ShiftModifier | Qt.KeyboardModifier.AltModifier):
            return False
        if event.key() not in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
            return False

        mw = self._main_window
        if not mw or not mw.isActiveWindow():
            return False

        table = getattr(mw, 'table', None)
        if table is None:
            return False

        focus = QApplication.focusWidget()
        if focus is None:
            return False

        # Allow triggering when focus is in the grid itself OR in the grid's filter boxes.
        # This keeps the shortcut useful while typing filters (e.g. "Filter Source" box).
        source_filter = getattr(mw, 'source_filter', None)
        target_filter = getattr(mw, 'target_filter', None)
        if focus in (source_filter, target_filter) or focus is table or focus is table.viewport() or table.isAncestorOf(focus) or isinstance(focus, (ReadOnlyGridTextEditor, EditableGridTextEditor)):
            if hasattr(mw, 'confirm_selected_or_next'):
                mw.confirm_selected_or_next()
                return True

        return False


class GridTableEventFilter:
    """Mixin to pass keyboard shortcuts from editor to table"""
    pass


class GridTextEditor(QTextEdit):
    """Custom QTextEdit for grid cells that passes special shortcuts to parent"""
    
    table_widget = None  # Will be set by delegate
    assistance_panel = None  # Will be set by delegate
    current_row = None  # Track which row this editor is in
    
    def __init__(self, parent=None):
        super().__init__(parent)
        # Prevent Tab from changing focus - we handle it manually
        self.setTabChangesFocus(False)
    
    def keyPressEvent(self, event):
        """Override keyPressEvent to handle Tab - this is called AFTER event()"""
        # Ctrl+Tab: Insert actual tab character
        if event.key() == Qt.Key.Key_Tab and event.modifiers() == Qt.KeyboardModifier.ControlModifier:
            self.insertPlainText('\t')
            event.accept()
            return
        
        # Tab (without Ctrl): Cycle to source cell
        if event.key() == Qt.Key.Key_Tab:
            if self.table_widget and self.current_row is not None:
                # Close current editor
                self.table_widget.closeEditor(self, QAbstractItemDelegate.EndEditHint.NoHint)
                # Open editor for source cell in same row (column 2)
                source_index = self.table_widget.model().index(self.current_row, 2)
                self.table_widget.setCurrentIndex(source_index)
                self.table_widget.edit(source_index)
                event.accept()
                return
        
        # Handle other Ctrl+ shortcuts
        if event.modifiers() & Qt.KeyboardModifier.ControlModifier:
            # Ctrl+1 through Ctrl+9: Insert match by number
            # TODO: Update to work with new results_panels system
            if event.key() >= Qt.Key.Key_1 and event.key() <= Qt.Key.Key_9:
                # Disabled for now - needs update to work with results_panels
                pass
            # Ctrl+Up/Down: Send to grid for grid navigation
            elif event.key() in (Qt.Key.Key_Up, Qt.Key.Key_Down):
                if self.table_widget:
                    self.table_widget.keyPressEvent(event)
                    event.accept()
                    return
            # Ctrl+Enter: Confirm & Next (call main window method directly)
            elif event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
                if self.main_window and hasattr(self.main_window, 'confirm_selected_or_next'):
                    self.main_window.confirm_selected_or_next()
                event.accept()
                return
            # Ctrl+, (comma): Insert next memoQ tag or wrap selection with tag pair
            elif event.key() == Qt.Key.Key_Comma:
                self._insert_next_tag_or_wrap_selection()
                event.accept()
                return
        
        # Shift+Enter: Insert line break (for multi-line content)
        if event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
            if event.modifiers() & Qt.KeyboardModifier.ShiftModifier:
                super().keyPressEvent(event)
                event.accept()
                return
        
        # Enter/Return: Don't insert, let delegate handle it (will close editor)
        if event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
            # Just ignore it - the table's default behavior will close the editor
            event.ignore()
            return
        
        # All other keys: Handle normally (including ESC for closing editor)
        print(f"🔑 GridTextEditor: Passing to super().keyPressEvent")
        super().keyPressEvent(event)
    
    def _insert_next_tag_or_wrap_selection(self):
        """
        Insert the next memoQ tag, HTML tag, or CafeTran pipe symbol from source, or wrap selection.
        
        Behavior:
        - If text is selected: Wrap it with the next available tag pair [N}selection{N] or |selection|
        - If no selection: Insert the next unused tag/pipe from source at cursor position
        
        Supports:
        - memoQ tags: [1}, {1], [2}, {2], etc.
        - HTML/XML tags: <li>, </li>, <b>, </b>, <i>, </i>, etc.
        - CafeTran pipe symbols: |
        
        Shortcut: Ctrl+, (comma)
        """
        # Get the main window and current segment
        if not self.table_widget or self.current_row is None:
            return
        
        # Navigate up to find main window
        main_window = self.table_widget.parent()
        while main_window and not hasattr(main_window, 'current_project'):
            main_window = main_window.parent()
        
        if not main_window or not hasattr(main_window, 'current_project'):
            return
        
        if not main_window.current_project or self.current_row >= len(main_window.current_project.segments):
            return
        
        segment = main_window.current_project.segments[self.current_row]
        source_text = segment.source
        current_target = self.toPlainText()
        
        # Check what type of tags are in the source
        has_memoq_tags = bool(extract_memoq_tags(source_text))
        has_html_tags = bool(extract_html_tags(source_text))
        has_any_tags = has_memoq_tags or has_html_tags
        has_pipe_symbols = '|' in source_text
        
        # Check if there's a selection
        cursor = self.textCursor()
        if cursor.hasSelection():
            selected_text = cursor.selectedText()
            
            # Try memoQ tag pair first
            if has_memoq_tags:
                opening_tag, closing_tag = get_wrapping_tag_pair(source_text, current_target)
                if opening_tag and closing_tag:
                    wrapped_text = f"{opening_tag}{selected_text}{closing_tag}"
                    cursor.insertText(wrapped_text)
                    if hasattr(main_window, 'log'):
                        main_window.log(f"🏷️ Wrapped selection with {opening_tag}...{closing_tag}")
                    return
            
            # Try CafeTran pipe symbols
            if has_pipe_symbols:
                pipes_needed = get_next_pipe_count_needed(source_text, current_target)
                if pipes_needed >= 2:
                    # Wrap with pipes
                    wrapped_text = f"|{selected_text}|"
                    cursor.insertText(wrapped_text)
                    if hasattr(main_window, 'log'):
                        main_window.log(f"🏷️ Wrapped selection with |...|")
                    return
            
            if hasattr(main_window, 'log'):
                main_window.log("⚠️ No tag pairs available from source")
        else:
            # No selection - insert next unused tag or pipe at cursor
            
            # Try memoQ tags and HTML tags (find_next_unused_tag handles both)
            if has_any_tags:
                next_tag = find_next_unused_tag(source_text, current_target)
                if next_tag:
                    cursor.insertText(next_tag)
                    if hasattr(main_window, 'log'):
                        main_window.log(f"🏷️ Inserted tag: {next_tag}")
                    return
            
            # Try CafeTran pipe symbols
            if has_pipe_symbols:
                pipes_needed = get_next_pipe_count_needed(source_text, current_target)
                if pipes_needed > 0:
                    cursor.insertText('|')
                    if hasattr(main_window, 'log'):
                        main_window.log(f"🏷️ Inserted pipe symbol (|)")
                    return
            
            if hasattr(main_window, 'log'):
                main_window.log("✓ All tags from source already in target")


class ReadOnlyGridTextEditor(QTextEdit):
    """Read-only QTextEdit for source cells - allows easy text selection"""
    
    # Class variable for tag highlight color (shared across all instances)
    tag_highlight_color = '#7f0001'  # Default memoQ dark red
    
    table_widget = None  # Will be set by delegate
    current_row = None  # Track which row this editor is in
    allow_source_edit = False  # Will be set by delegate based on settings
    
    def __init__(self, text: str = "", parent=None, row: int = -1):
        super().__init__(parent)
        self.row = row  # Store row number for Tab cycling
        self.table_ref = parent  # Store table reference (parent is the table)
        self.setReadOnly(True)  # Prevent typing but allow selection
        self.setPlainText(text)

        # CRITICAL: Enable keyboard focus and text selection
        self.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
        # Enable text interaction: selection with keyboard and mouse
        self.setTextInteractionFlags(
            Qt.TextInteractionFlag.TextSelectableByKeyboard |
            Qt.TextInteractionFlag.TextSelectableByMouse
        )

        # Make inactive selections stay visible with same color
        from PyQt6.QtGui import QPalette
        palette = self.palette()
        palette.setColor(QPalette.ColorGroup.Active, QPalette.ColorRole.Highlight, QColor("#D0E7FF"))
        palette.setColor(QPalette.ColorGroup.Active, QPalette.ColorRole.HighlightedText, QColor("black"))
        palette.setColor(QPalette.ColorGroup.Inactive, QPalette.ColorRole.Highlight, QColor("#D0E7FF"))
        palette.setColor(QPalette.ColorGroup.Inactive, QPalette.ColorRole.HighlightedText, QColor("black"))
        self.setPalette(palette)

        # Style to look like a normal cell with subtle selection
        # Background and text colors now managed by theme system
        self.setStyleSheet("""
            QTextEdit {
                border: none;
                padding: 0px 4px 0px 0px;
            }
            QTextEdit:focus {
                border: 1px solid #2196F3;
            }
            QTextEdit::selection {
                background-color: #D0E7FF;
                color: black;
            }
        """)

        # Set document margins to 0 for compact display
        doc = self.document()
        doc.setDocumentMargin(0)

        # Configure text option for minimal line spacing
        # With zero-width spaces inserted after invisible character markers,
        # normal WordWrap will work correctly
        wrap_mode = QTextOption.WrapMode.WordWrap

        self.setWordWrapMode(wrap_mode)
        self.setAcceptRichText(False)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)

        text_option = QTextOption()
        text_option.setWrapMode(wrap_mode)
        doc.setDefaultTextOption(text_option)
        
        # Set minimum height to 0 - let content determine size
        self.setMinimumHeight(0)
        self.setMaximumHeight(16777215)  # Qt's max int
        
        # Store termbase matches for this cell (for tooltip and double-click)
        self.termbase_matches = {}
        
        # Enable mouse tracking for hover tooltips
        self.setMouseTracking(True)

        # Add syntax highlighter for tags (no spellcheck for source cells)
        # Get invisible char color from main window if available
        main_window = self._get_main_window()
        invisible_char_color = main_window.invisible_char_color if main_window and hasattr(main_window, 'invisible_char_color') else '#999999'
        self.highlighter = TagHighlighter(self.document(), self.tag_highlight_color, invisible_char_color, enable_spellcheck=False)

        # Store raw text (with tags) for mode switching
        self._raw_text = text
    
    def keyPressEvent(self, event):
        """Override to fix clipboard and word navigation when invisible characters shown"""
        from PyQt6.QtCore import Qt
        from PyQt6.QtWidgets import QApplication
        from PyQt6.QtGui import QTextCursor
        
        # Check for Ctrl+C (copy)
        if event.modifiers() == Qt.KeyboardModifier.ControlModifier and event.key() == Qt.Key.Key_C:
            # Get selected text
            cursor = self.textCursor()
            if cursor.hasSelection():
                selected_text = cursor.selectedText()
                # Reverse invisible character replacements before copying
                main_window = self._get_main_window()
                if main_window and hasattr(main_window, 'reverse_invisible_replacements'):
                    clean_text = main_window.reverse_invisible_replacements(selected_text)
                    # Also replace paragraph separator with newline (Qt uses U+2029)
                    clean_text = clean_text.replace('\u2029', '\n')
                    # Set clipboard with clean text
                    clipboard = QApplication.clipboard()
                    clipboard.setText(clean_text)
                    return  # Don't call parent - we handled it
        
        # Handle Ctrl+Arrow word navigation when invisibles are shown
        main_window = self._get_main_window()
        if main_window and hasattr(main_window, 'invisible_display_settings'):
            if main_window.invisible_display_settings.get('spaces', False):
                ctrl_only = event.modifiers() == Qt.KeyboardModifier.ControlModifier
                ctrl_shift = event.modifiers() == (Qt.KeyboardModifier.ControlModifier | Qt.KeyboardModifier.ShiftModifier)
                
                if (ctrl_only or ctrl_shift) and event.key() in (Qt.Key.Key_Left, Qt.Key.Key_Right):
                    cursor = self.textCursor()
                    text = self.toPlainText()
                    pos = cursor.position()
                    word_chars = ('·', '\u200B', '\n', '\t', '→', '°', '¶', ' ')
                    
                    if event.key() == Qt.Key.Key_Right:
                        # Move to end of current word, then skip delimiters to start of next word
                        # First skip any delimiters we're on
                        while pos < len(text) and text[pos] in word_chars:
                            pos += 1
                        # Then skip to end of word (next delimiter)
                        while pos < len(text) and text[pos] not in word_chars:
                            pos += 1
                    else:  # Key_Left
                        # Move backwards: skip delimiters, then find start of word
                        if pos > 0:
                            pos -= 1
                        # Skip any delimiters
                        while pos > 0 and text[pos] in word_chars:
                            pos -= 1
                        # Find start of word
                        while pos > 0 and text[pos - 1] not in word_chars:
                            pos -= 1
                    
                    # Apply cursor movement
                    if ctrl_shift:
                        cursor.setPosition(pos, QTextCursor.MoveMode.KeepAnchor)
                    else:
                        cursor.setPosition(pos)
                    self.setTextCursor(cursor)
                    return
        
        # Arrow Up/Down: memoQ-style segment navigation at cell boundaries
        # When cursor is at top line and Up is pressed, go to previous segment
        # When cursor is at bottom line and Down is pressed, go to next segment
        if event.key() in (Qt.Key.Key_Up, Qt.Key.Key_Down) and event.modifiers() == Qt.KeyboardModifier.NoModifier:
            cursor = self.textCursor()
            
            # Get current cursor position info
            current_block = cursor.block()
            doc = self.document()
            first_block = doc.firstBlock()
            last_block = doc.lastBlock()
            
            # Get table reference
            table = self.table_ref if hasattr(self, 'table_ref') else self.parent()
            
            if event.key() == Qt.Key.Key_Up:
                # Check if we're on the first line
                if current_block == first_block:
                    # Navigate to previous segment
                    main_window = self._get_main_window()
                    if main_window and hasattr(main_window, 'go_to_previous_segment'):
                        # Get cursor column position for smart positioning
                        col_in_line = cursor.positionInBlock()
                        # Navigate and let go_to_previous_segment position cursor in target cell
                        main_window.go_to_previous_segment(target_column=col_in_line, to_last_line=True)
                        
                        # Now move focus from target to source cell, preserving column position
                        if table:
                            current_row = table.currentRow()
                            source_widget = table.cellWidget(current_row, 2)  # Column 2 is Source
                            if source_widget and hasattr(source_widget, 'textCursor'):
                                new_cursor = source_widget.textCursor()
                                src_doc = source_widget.document()
                                src_last_block = src_doc.lastBlock()
                                new_cursor.setPosition(src_last_block.position())
                                line_length = src_last_block.length() - 1
                                target_pos = min(col_in_line, max(0, line_length))
                                new_cursor.movePosition(QTextCursor.MoveOperation.Right, QTextCursor.MoveMode.MoveAnchor, target_pos)
                                source_widget.setTextCursor(new_cursor)
                                source_widget.setFocus()
                        return
            
            elif event.key() == Qt.Key.Key_Down:
                # Check if we're on the last line
                if current_block == last_block:
                    # Navigate to next segment
                    main_window = self._get_main_window()
                    if main_window and hasattr(main_window, 'go_to_next_segment'):
                        # Get cursor column position for smart positioning
                        col_in_line = cursor.positionInBlock()
                        # Navigate and let go_to_next_segment position cursor in target cell
                        main_window.go_to_next_segment(target_column=col_in_line, to_first_line=True)
                        
                        # Now move focus from target to source cell, preserving column position
                        if table:
                            current_row = table.currentRow()
                            source_widget = table.cellWidget(current_row, 2)  # Column 2 is Source
                            if source_widget and hasattr(source_widget, 'textCursor'):
                                new_cursor = source_widget.textCursor()
                                src_doc = source_widget.document()
                                src_first_block = src_doc.firstBlock()
                                new_cursor.setPosition(src_first_block.position())
                                line_length = src_first_block.length() - 1
                                target_pos = min(col_in_line, max(0, line_length))
                                new_cursor.movePosition(QTextCursor.MoveOperation.Right, QTextCursor.MoveMode.MoveAnchor, target_pos)
                                source_widget.setTextCursor(new_cursor)
                                source_widget.setFocus()
                        return
        
        # Ctrl+Enter: Confirm & Next (same behavior as in target cell)
        if event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
            mods = event.modifiers()
            if (mods & Qt.KeyboardModifier.ControlModifier) and not (mods & (Qt.KeyboardModifier.ShiftModifier | Qt.KeyboardModifier.AltModifier)):
                main_window = self._get_main_window()
                if main_window and hasattr(main_window, 'confirm_selected_or_next'):
                    main_window.confirm_selected_or_next()
                    event.accept()
                    return
        
        # Tab key: Cycle to target cell (column 3) in same row
        if event.key() == Qt.Key.Key_Tab and event.modifiers() == Qt.KeyboardModifier.NoModifier:
            table = self.table_ref if hasattr(self, 'table_ref') else self.parent()
            if table and self.row >= 0:
                target_widget = table.cellWidget(self.row, 3)  # Column 3 is Target
                if target_widget:
                    target_widget.setFocus()
                    table.setCurrentCell(self.row, 3)
                    return
        
        super().keyPressEvent(event)
    
    def mouseDoubleClickEvent(self, event):
        """Handle double-click to select words properly when invisibles are shown"""
        from PyQt6.QtCore import Qt
        from PyQt6.QtGui import QTextCursor
        
        main_window = self._get_main_window()
        # Check if invisible spaces are being shown
        if main_window and hasattr(main_window, 'invisible_display_settings'):
            if main_window.invisible_display_settings.get('spaces', False):
                # Get cursor position at click
                cursor = self.cursorForPosition(event.pos())
                pos = cursor.position()
                text = self.toPlainText()
                
                # Find word boundaries using middle dot (·) or zero-width space as delimiters
                # Find start of word (search backwards for · or start of text)
                start = pos
                while start > 0 and text[start - 1] not in ('·', '\u200B', '\n', '\t', '→', '°', '¶'):
                    start -= 1
                
                # Find end of word (search forwards for · or end of text)
                end = pos
                while end < len(text) and text[end] not in ('·', '\u200B', '\n', '\t', '→', '°', '¶'):
                    end += 1
                
                # Select the word
                cursor.setPosition(start)
                cursor.setPosition(end, QTextCursor.MoveMode.KeepAnchor)
                self.setTextCursor(cursor)
                return
        
        # Default behavior
        super().mouseDoubleClickEvent(event)
    
    def _get_main_window(self):
        """Get the main application window by traversing the parent hierarchy"""
        from PyQt6.QtWidgets import QTableWidget
        
        # Try multiple ways to get table reference (different creation paths set different attributes)
        # Check table_widget FIRST - delegate explicitly sets this to the correct table
        # table_ref may be set to viewport (wrong) when delegate creates the editor
        table = None
        if hasattr(self, 'table_widget') and self.table_widget and isinstance(self.table_widget, QTableWidget):
            table = self.table_widget
        elif hasattr(self, 'table_ref') and self.table_ref and isinstance(self.table_ref, QTableWidget):
            table = self.table_ref
        else:
            # Walk up parent hierarchy to find QTableWidget
            widget = self.parent()
            while widget and not isinstance(widget, QTableWidget):
                widget = widget.parent() if hasattr(widget, 'parent') else None
            table = widget
        
        if not table:
            return None
        main_window = table.parent()
        while main_window and not hasattr(main_window, 'go_to_first_segment'):
            main_window = main_window.parent()
        return main_window
    
    def update_display_mode(self, text: str, show_tags: bool):
        """
        Update the display based on tag view mode.
        
        Args:
            text: The raw text (with HTML tags like <b>bold</b>)
            show_tags: If True, show raw tags. If False, show formatted WYSIWYG.
        """
        self._raw_text = text
        
        if show_tags:
            # Show raw tags as plain text
            self.setPlainText(text)
        else:
            # Show WYSIWYG - convert tags to actual formatting
            html = get_formatted_html_display(text)
            self.setHtml(html)
    
    def get_raw_text(self) -> str:
        """Get the raw text with tags, regardless of display mode."""
        return getattr(self, '_raw_text', self.toPlainText())

    def highlight_termbase_matches(self, matches_dict: Dict):
        """
        Highlight termbase matches in the text using the configured style.
        Does NOT change the widget - just adds formatting to existing text.
        
        Supported styles (configured in Settings > View Settings):
        - 'background': Pastel green background colors based on priority (default)
        - 'dotted': Subtle dotted underline (IDE/code editor style)
        - 'semibold': Slightly bolder text with tinted color (typographic)
        
        Args:
            matches_dict: Dictionary of {term: {'translation': str, 'priority': int}} or {term: str}
        """
        from PyQt6.QtGui import QTextCursor, QTextCharFormat, QColor, QFont
        
        # Get the document and create a cursor
        doc = self.document()
        text = self.toPlainText()
        text_lower = text.lower()
        
        # IMPORTANT: Always clear all previous formatting first to prevent inconsistent highlighting
        cursor = QTextCursor(doc)
        cursor.select(QTextCursor.SelectionType.Document)
        default_fmt = QTextCharFormat()
        cursor.setCharFormat(default_fmt)
        
        # If no matches, we're done (highlighting has been cleared)
        if not matches_dict:
            return
        
        # Get highlight style from main window settings
        highlight_style = 'background'  # default
        dotted_color = '#808080'  # default medium gray (more visible)
        parent = self.parent()
        while parent:
            if hasattr(parent, 'termbase_highlight_style'):
                highlight_style = getattr(parent, 'termbase_highlight_style', 'background')
                dotted_color = getattr(parent, 'termbase_dotted_color', '#808080')
                break
            elif hasattr(parent, 'load_general_settings'):
                settings = parent.load_general_settings()
                highlight_style = settings.get('termbase_highlight_style', 'semibold')
                dotted_color = settings.get('termbase_dotted_color', '#808080')
                break
            parent = parent.parent() if hasattr(parent, 'parent') else None
        
        # Sort matches by source term length (longest first) to avoid partial matches
        # Since dict keys are now term_ids, we need to extract source terms first
        term_entries = []
        for term_id, match_info in matches_dict.items():
            if isinstance(match_info, dict):
                source_term = match_info.get('source', '')
                if source_term:
                    term_entries.append((source_term, term_id, match_info))
        
        # Sort by source term length (longest first)
        term_entries.sort(key=lambda x: len(x[0]), reverse=True)
        
        # Track positions we've already highlighted to avoid overlaps
        highlighted_ranges = []
        
        for term, term_id, match_info in term_entries:
            # Get ranking, forbidden status, and termbase type
            ranking = match_info.get('ranking', None)
            forbidden = match_info.get('forbidden', False)
            is_project_termbase = match_info.get('is_project_termbase', False)
            translation = match_info.get('target', match_info.get('translation', ''))
            notes = match_info.get('notes', '')
            termbase_name = match_info.get('termbase_name', '')
            
            # IMPORTANT: Treat ranking #1 as project termbase (even if flag not set)
            is_effective_project = is_project_termbase or (ranking == 1)
            
            # Build tooltip text
            tooltip_text = ""
            if translation:
                tooltip_text = f"Glossary: {translation}"
                if termbase_name:
                    tooltip_text += f"\nFrom: {termbase_name}"
                if notes:
                    tooltip_text += f"\nNote: {notes}"
            
            # Find all occurrences of this term (case-insensitive)
            term_lower = term.lower()
            start = 0
            while True:
                idx = text_lower.find(term_lower, start)
                if idx == -1:
                    break
                
                end_idx = idx + len(term)
                
                # Check if this range overlaps with already highlighted text
                overlaps = any(
                    (idx < h_end and end_idx > h_start)
                    for h_start, h_end in highlighted_ranges
                )
                
                if not overlaps:
                    # Create cursor for this position
                    cursor = QTextCursor(doc)
                    cursor.setPosition(idx)
                    cursor.setPosition(end_idx, QTextCursor.MoveMode.KeepAnchor)
                    
                    # Create format based on style
                    fmt = QTextCharFormat()
                    
                    if highlight_style == 'dotted':
                        # DOTTED UNDERLINE STYLE (IDE/code editor approach)
                        # Simple dotted line like the Gemini example - clean and unobtrusive
                        fmt.setUnderlineStyle(QTextCharFormat.UnderlineStyle.DotLine)
                        if forbidden:
                            fmt.setUnderlineColor(QColor(0, 0, 0))  # Black for forbidden
                        else:
                            # Higher priority = red (more attention), lower = gray (subtle)
                            if ranking == 1:
                                fmt.setUnderlineColor(QColor('#CC0000'))  # Red for priority 1 (project termbase)
                            elif ranking == 2:
                                fmt.setUnderlineColor(QColor('#505050'))  # Dark gray for priority 2
                            elif ranking == 3:
                                fmt.setUnderlineColor(QColor('#707070'))  # Medium gray
                            else:
                                fmt.setUnderlineColor(QColor(dotted_color))  # User-configured color
                        # Add tooltip with translation and notes
                        if tooltip_text:
                            fmt.setToolTip(tooltip_text)
                    
                    elif highlight_style == 'semibold':
                        # SEMIBOLD TEXT STYLE (typographic approach)
                        fmt.setFontWeight(QFont.Weight.DemiBold)
                        if forbidden:
                            fmt.setForeground(QColor(180, 0, 0))  # Dark red for forbidden
                        else:
                            # Tinted dark color based on ranking
                            if ranking == 1:
                                fmt.setForeground(QColor('#1B5E20'))  # Dark green for priority 1
                            elif ranking == 2:
                                fmt.setForeground(QColor('#2E7D32'))  # Medium dark green
                            elif ranking == 3:
                                fmt.setForeground(QColor('#388E3C'))  # Medium green
                            else:
                                fmt.setForeground(QColor('#43A047'))  # Lighter green
                        # Add tooltip with translation and notes
                        if tooltip_text:
                            fmt.setToolTip(tooltip_text)
                    
                    else:
                        # BACKGROUND COLOR STYLE (default - current behavior)
                        if forbidden:
                            color = QColor(0, 0, 0)  # Black for forbidden terms
                            fmt.setForeground(QColor("white"))
                        else:
                            # Use ranking to determine soft green shade
                            if ranking is not None:
                                if ranking == 1:
                                    color = QColor(165, 214, 167)  # Soft medium green (Green 200)
                                elif ranking == 2:
                                    color = QColor(200, 230, 201)  # Soft light green (Green 100)
                                elif ranking == 3:
                                    color = QColor(220, 237, 200)  # Very soft light green
                                else:
                                    color = QColor(232, 245, 233)  # Extremely soft pastel green
                            else:
                                color = QColor(200, 230, 201)  # Green 100 (fallback)
                            fmt.setForeground(QColor("black"))
                        fmt.setBackground(color)
                        # Add tooltip with translation and notes
                        if tooltip_text:
                            fmt.setToolTip(tooltip_text)
                    
                    # Apply format
                    cursor.setCharFormat(fmt)
                    
                    # Track this range as highlighted
                    highlighted_ranges.append((idx, end_idx))
                
                start = end_idx
    
    def highlight_non_translatables(self, nt_matches: list, highlighted_ranges: list = None):
        """
        Highlight non-translatable matches in the text using pastel yellow background.
        
        Args:
            nt_matches: List of dicts with 'text', 'start', 'end' keys from NT manager
            highlighted_ranges: Optional list of already-highlighted ranges to avoid overlap
        """
        from PyQt6.QtGui import QTextCursor, QTextCharFormat, QColor
        
        if not nt_matches:
            return
        
        doc = self.document()
        
        # Track ranges we've highlighted (to avoid overlap with termbase matches)
        if highlighted_ranges is None:
            highlighted_ranges = []
        
        # Pastel yellow for non-translatables
        nt_color = QColor(255, 253, 208)  # Pastel yellow (#FFFDD0)
        
        for match in nt_matches:
            start_pos = match.get('start', 0)
            end_pos = match.get('end', 0)
            
            if start_pos >= end_pos:
                continue
            
            # Check for overlap with existing highlights
            overlaps = any(
                (start_pos < h_end and end_pos > h_start)
                for h_start, h_end in highlighted_ranges
            )
            
            if overlaps:
                continue
            
            # Create cursor for this position
            cursor = QTextCursor(doc)
            cursor.setPosition(start_pos)
            cursor.setPosition(end_pos, QTextCursor.MoveMode.KeepAnchor)
            
            # Create format with pastel yellow background
            fmt = QTextCharFormat()
            fmt.setBackground(nt_color)
            fmt.setForeground(QColor("#5D4E37"))  # Dark brown text for contrast
            
            cursor.setCharFormat(fmt)
            highlighted_ranges.append((start_pos, end_pos))
    
    def event(self, event):
        """Override event() to catch Tab and Ctrl+T keys before Qt's default handling"""
        # Catch Tab key at event level (before keyPressEvent)
        if event.type() == event.Type.KeyPress:
            key_event = event
            
            # Ctrl+E: Add selected terms to termbase (with dialog)
            if key_event.key() == Qt.Key.Key_E and key_event.modifiers() == Qt.KeyboardModifier.ControlModifier:
                self._handle_add_to_termbase()
                return True  # Event handled
            
            # Alt+Left: Quick add selected terms to last-used termbase (no dialog)
            if key_event.key() == Qt.Key.Key_Left and key_event.modifiers() == Qt.KeyboardModifier.AltModifier:
                self._handle_quick_add_to_termbase()
                return True  # Event handled
            
            # Ctrl+Alt+N: Add selected text to non-translatables
            if key_event.key() == Qt.Key.Key_N and key_event.modifiers() == (Qt.KeyboardModifier.ControlModifier | Qt.KeyboardModifier.AltModifier):
                self._handle_add_to_nt()
                return True  # Event handled
            
            # Ctrl+Home: Navigate to first segment (pass to main window)
            if key_event.key() == Qt.Key.Key_Home and key_event.modifiers() == Qt.KeyboardModifier.ControlModifier:
                main_window = self._get_main_window()
                if main_window and hasattr(main_window, 'go_to_first_segment'):
                    main_window.go_to_first_segment()
                return True  # Event handled
            
            # Ctrl+End: Navigate to last segment (pass to main window)
            if key_event.key() == Qt.Key.Key_End and key_event.modifiers() == Qt.KeyboardModifier.ControlModifier:
                main_window = self._get_main_window()
                if main_window and hasattr(main_window, 'go_to_last_segment'):
                    main_window.go_to_last_segment()
                return True  # Event handled
            
            if key_event.key() == Qt.Key.Key_Tab:
                # Ctrl+Tab: Insert actual tab character (if editing is allowed)
                if key_event.modifiers() == Qt.KeyboardModifier.ControlModifier:
                    if self.allow_source_edit:
                        self.insertPlainText('\t')
                        return True  # Event handled
                
                # Tab alone: Cycle to target cell (column 3) in same row
                if key_event.modifiers() == Qt.KeyboardModifier.NoModifier:
                    if self.table_ref and self.row >= 0:
                        # Get target cell widget (column 3)
                        target_widget = self.table_ref.cellWidget(self.row, 3)
                        if target_widget:
                            target_widget.setFocus()
                            self.table_ref.setCurrentCell(self.row, 3)
                            return True  # Event handled, don't propagate
        
        # Let base class handle all other events
        return super().event(event)
    
    def keyPressEvent(self, event):
        """Override to handle other keys (Tab is handled in event())"""
        # All keys: Handle normally (Tab already handled in event())
        super().keyPressEvent(event)
    
    def sizeHint(self):
        """Return compact size based on content"""
        doc = self.document()
        ideal_width = self.width() if self.width() > 0 else 200
        doc.setTextWidth(ideal_width)
        height = int(doc.size().height())
        # Add minimal padding (2px total)
        height = max(height + 2, 1)
        current_width = self.width() if self.width() > 0 else 200
        return QSize(current_width, height)
    
    def _handle_add_to_termbase(self):
        """Handle Ctrl+T: Add selected source and target terms to termbase"""
        if not self.table_ref or self.row < 0:
            return
        
        # Get source selection (from this widget)
        source_text = self.textCursor().selectedText().strip()
        
        # Get target cell widget and its selection
        target_widget = self.table_ref.cellWidget(self.row, 3)
        target_text = ""
        if target_widget and hasattr(target_widget, 'textCursor'):
            target_text = target_widget.textCursor().selectedText().strip()
        
        # Validate we have both selections
        if not source_text or not target_text:
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.warning(
                self,
                "Selection Required",
                "Please select text in both Source and Target cells before adding to glossary.\\n\\n"
                "Workflow:\\n"
                "1. Select term in source cell\\n"
                "2. Press Tab to cycle to target cell\\n"
                "3. Select corresponding translation\\n"
                "4. Press Ctrl+E (or right-click) to add to glossary"
            )
            return
        
        # Find main window and call add_to_termbase method
        main_window = self.table_ref.parent()
        while main_window and not hasattr(main_window, 'add_term_pair_to_termbase'):
            main_window = main_window.parent()
        
        if main_window and hasattr(main_window, 'add_term_pair_to_termbase'):
            main_window.add_term_pair_to_termbase(source_text, target_text)
    
    def _handle_quick_add_to_termbase(self):
        """Handle Ctrl+R: Quick add selected source and target terms to termbase (no dialog)"""
        if not self.table_ref or self.row < 0:
            return
        
        # Get source selection (from this widget)
        source_text = self.textCursor().selectedText().strip()
        
        # Get target cell widget and its selection
        target_widget = self.table_ref.cellWidget(self.row, 3)
        target_text = ""
        if target_widget and hasattr(target_widget, 'textCursor'):
            target_text = target_widget.textCursor().selectedText().strip()
        
        # Validate we have both selections
        if not source_text or not target_text:
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.warning(
                self,
                "Selection Required",
                "Please select text in both Source and Target cells before quick-adding to termbase.\n\n"
                "Tip: Use Ctrl+E to add with a dialog where you can choose termbase and add metadata."
            )
            return
        
        # Find main window and call quick_add_to_termbase method
        main_window = self.table_ref.parent()
        while main_window and not hasattr(main_window, 'quick_add_term_pair_to_termbase'):
            main_window = main_window.parent()
        
        if main_window and hasattr(main_window, 'quick_add_term_pair_to_termbase'):
            main_window.quick_add_term_pair_to_termbase(source_text, target_text)
    
    def _handle_superlookup_search(self):
        """Handle right-click: Search selected text in Superlookup"""
        # Get selected text
        selected_text = self.textCursor().selectedText().strip()
        
        if not selected_text:
            return
        
        # Find main window
        main_window = self._get_main_window()
        if not main_window:
            return
        
        # Get project languages
        source_lang = None
        target_lang = None
        if hasattr(main_window, 'current_project') and main_window.current_project:
            source_lang = main_window.current_project.source_lang
            target_lang = main_window.current_project.target_lang
        
        # Navigate to Superlookup
        main_window._go_to_superlookup()
        
        # Trigger search
        if hasattr(main_window, 'lookup_tab') and main_window.lookup_tab:
            main_window.lookup_tab.search_with_query(
                selected_text,
                switch_to_vertical=True,
                source_lang=source_lang,
                target_lang=target_lang
            )
    
    def mouseMoveEvent(self, event):
        """Show tooltip when hovering over highlighted termbase matches"""
        super().mouseMoveEvent(event)
        
        # Get cursor position
        cursor = self.cursorForPosition(event.pos())
        cursor_pos = cursor.position()
        
        # Check if cursor is over a termbase match
        for term_id, match_info in self.termbase_matches.items():
            # Extract source term from match_info
            term = match_info.get('source', '') if isinstance(match_info, dict) else str(term_id)
            # Find term position in text
            text = self.toPlainText()
            text_lower = text.lower()
            term_lower = term.lower()
            
            start = 0
            while True:
                idx = text_lower.find(term_lower, start)
                if idx == -1:
                    break
                
                end_idx = idx + len(term)
                
                # Check if cursor is within this term's range
                if idx <= cursor_pos <= end_idx:
                    # Get translation and other info
                    if isinstance(match_info, dict):
                        translation = match_info.get('translation', '')
                        priority = match_info.get('priority', 50)
                        forbidden = match_info.get('forbidden', False)
                        
                        # Build tooltip (show only target translation)
                        tooltip = f"<b>{translation}</b>"
                        if forbidden:
                            tooltip += "<br><span style='color: red;'>⚠️ FORBIDDEN TERM</span>"
                        tooltip += f"<br><span style='color: #666;'>Priority: {priority}</span>"
                        tooltip += "<br><span style='color: #666;'><i>Double-click to insert</i></span>"
                    else:
                        tooltip = f"<b>{match_info}</b><br><span style='color: #666;'><i>Double-click to insert</i></span>"
                    
                    self.setToolTip(tooltip)
                    return
                
                start = end_idx
        
        # No match found at cursor position
        self.setToolTip("")
    
    def mouseDoubleClickEvent(self, event):
        """Insert termbase translation on double-click"""
        # Get cursor position
        cursor = self.cursorForPosition(event.pos())
        cursor_pos = cursor.position()
        
        # Check if double-clicked on a termbase match
        for term_id, match_info in self.termbase_matches.items():
            # Extract source term from match_info
            term = match_info.get('source', '') if isinstance(match_info, dict) else str(term_id)
            text = self.toPlainText()
            text_lower = text.lower()
            term_lower = term.lower()
            
            start = 0
            while True:
                idx = text_lower.find(term_lower, start)
                if idx == -1:
                    break
                
                end_idx = idx + len(term)
                
                # Check if cursor is within this term's range
                if idx <= cursor_pos <= end_idx:
                    # Get translation
                    if isinstance(match_info, dict):
                        translation = match_info.get('translation', '')
                    else:
                        translation = match_info
                    
                    # Insert translation into target cell at cursor position
                    if self.table_ref and self.row >= 0:
                        target_widget = self.table_ref.cellWidget(self.row, 3)
                        if target_widget and hasattr(target_widget, 'textCursor'):
                            # Insert at current cursor position in target
                            target_cursor = target_widget.textCursor()
                            target_cursor.insertText(translation)
                            target_widget.setFocus()
                    return
                
                start = end_idx
        
        # If not on a termbase match, allow normal double-click selection
        super().mouseDoubleClickEvent(event)
    
    def mousePressEvent(self, event):
        """Allow text selection on click and trigger row selection"""
        super().mousePressEvent(event)

        # Use stored table reference and row number
        if self.table_ref and self.row >= 0:
            try:
                self.table_ref.selectRow(self.row)
                self.table_ref.setCurrentCell(self.row, 2)

                # CRITICAL: Manually trigger on_cell_selected since signals aren't firing
                # Find the main window and call the method directly
                main_window = self.table_ref.parent()
                while main_window and not hasattr(main_window, 'on_cell_selected'):
                    main_window = main_window.parent()
                if main_window and hasattr(main_window, 'on_cell_selected'):
                    main_window.on_cell_selected(self.row, 2, -1, -1)
            except Exception as e:
                print(f"Error triggering manual cell selection: {e}")

    def mouseReleaseEvent(self, event):
        """Smart word selection - expand partial selections to full words
        
        Works across multiple lines - if you select text spanning several lines,
        partial words at the START and END of your selection will be expanded.
        """
        super().mouseReleaseEvent(event)

        # Check if smart selection is enabled
        main_window = self._get_main_window()
        if main_window and hasattr(main_window, 'enable_smart_word_selection'):
            if not main_window.enable_smart_word_selection:
                return  # Feature disabled

        # Get the current cursor
        cursor = self.textCursor()

        # Only expand if there's a selection
        if cursor.hasSelection():
            # Get selection boundaries
            start = cursor.selectionStart()
            end = cursor.selectionEnd()

            # Get the full text
            text = self.toPlainText()

            # Helper function to check if character is part of a word
            # Includes alphanumeric, underscore, hyphen, and apostrophe
            def is_word_char(char):
                return char.isalnum() or char in "_-'"

            # Track if we need to update the selection
            selection_changed = False

            # Expand START boundary if we're in the middle of a word
            # (i.e., the character before the selection is a word character)
            if start > 0 and is_word_char(text[start - 1]):
                # Also check that the first selected character is a word char
                # (to avoid expanding when selecting from whitespace)
                if start < len(text) and is_word_char(text[start]):
                    while start > 0 and is_word_char(text[start - 1]):
                        start -= 1
                    selection_changed = True

            # Expand END boundary if we're in the middle of a word
            # (i.e., the character after the selection is a word character)
            if end < len(text) and is_word_char(text[end]):
                # Also check that the last selected character is a word char
                # (to avoid expanding when selecting to whitespace)
                if end > 0 and is_word_char(text[end - 1]):
                    while end < len(text) and is_word_char(text[end]):
                        end += 1
                    selection_changed = True

            # Set the new selection if boundaries changed
            if selection_changed:
                cursor.setPosition(start)
                cursor.setPosition(end, cursor.MoveMode.KeepAnchor)
                self.setTextCursor(cursor)

    def _get_main_window(self):
        """Get the main application window by traversing the parent hierarchy"""
        if not self.parent():
            return None
        main_window = self.parent()
        while main_window and not hasattr(main_window, 'go_to_first_segment'):
            main_window = main_window.parent()
        return main_window
    
    def focusInEvent(self, event):
        """Select text when focused for easy copying and trigger row selection"""
        super().focusInEvent(event)
        # Don't auto-select - let user select manually
        
        # Use stored table reference and row number
        if self.table_ref and self.row >= 0:
            try:
                self.table_ref.selectRow(self.row)
                self.table_ref.setCurrentCell(self.row, 2)
                
                # CRITICAL: Manually trigger on_cell_selected since signals aren't firing
                # Find the main window and call the method directly
                main_window = self.table_ref.parent()
                while main_window and not hasattr(main_window, 'on_cell_selected'):
                    main_window = main_window.parent()
                if main_window and hasattr(main_window, 'on_cell_selected'):
                    main_window.on_cell_selected(self.row, 2, -1, -1)
            except Exception as e:
                print(f"Error triggering manual cell selection: {e}")

    def contextMenuEvent(self, event):
        """Show context menu with Add to Glossary and Add to Non-Translatables options"""
        from PyQt6.QtWidgets import QMenu
        from PyQt6.QtGui import QAction
        
        menu = QMenu(self)
        
        # Add standard actions
        if self.textCursor().hasSelection():
            copy_action = QAction("Copy", self)
            copy_action.triggered.connect(self.copy)
            menu.addAction(copy_action)
            menu.addSeparator()
        
        # Superlookup search action
        if self.textCursor().hasSelection():
            superlookup_action = QAction("🔍 Search in Superlookup (Ctrl+K)", self)
            superlookup_action.triggered.connect(self._handle_superlookup_search)
            menu.addAction(superlookup_action)
            menu.addSeparator()

        # QuickMenu (prompt-based actions)
        try:
            main_window = self._get_main_window()
            quickmenu_items = []
            if main_window and hasattr(main_window, 'prompt_manager_qt') and main_window.prompt_manager_qt:
                lib = getattr(main_window.prompt_manager_qt, 'library', None)
                if lib and hasattr(lib, 'get_quickmenu_grid_prompts'):
                    quickmenu_items = lib.get_quickmenu_grid_prompts() or []

            if quickmenu_items:
                qm_menu = menu.addMenu("⚡ QuickMenu")
                for rel_path, label in sorted(quickmenu_items, key=lambda x: (x[1] or x[0]).lower()):
                    prompt_menu = qm_menu.addMenu(label or rel_path)

                    run_show = QAction("▶ Run (show response)…", self)
                    run_show.triggered.connect(
                        lambda checked=False, p=rel_path: main_window.run_grid_quickmenu_prompt(p, origin_widget=self, behavior="show")
                    )
                    prompt_menu.addAction(run_show)

                    run_replace = QAction("↺ Run and replace target selection", self)
                    run_replace.triggered.connect(
                        lambda checked=False, p=rel_path: main_window.run_grid_quickmenu_prompt(p, origin_widget=self, behavior="replace")
                    )
                    prompt_menu.addAction(run_replace)

                menu.addSeparator()
        except Exception:
            # Never break the normal context menu due to QuickMenu errors
            pass
        
        # Add to glossary action (with dialog)
        add_to_tb_action = QAction("📖 Add to Glossary (Ctrl+E)", self)
        add_to_tb_action.triggered.connect(self._handle_add_to_termbase)
        menu.addAction(add_to_tb_action)
        
        # Quick add to glossary action (no dialog) - uses last-selected glossary from Ctrl+E
        quick_add_action = QAction("⚡ Quick Add to Glossary (Alt+Left)", self)
        quick_add_action.triggered.connect(self._handle_quick_add_to_termbase)
        menu.addAction(quick_add_action)
        
        # Add to non-translatables action
        add_to_nt_action = QAction("🚫 Add to Non-Translatables (Ctrl+Alt+N)", self)
        add_to_nt_action.triggered.connect(self._handle_add_to_nt)
        menu.addAction(add_to_nt_action)
        
        menu.exec(event.globalPos())
    
    def _handle_add_to_nt(self):
        """Handle Ctrl+Alt+N: Add selected text to active non-translatable list(s)"""
        # Get selected text
        selected_text = self.textCursor().selectedText().strip()
        
        if not selected_text:
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.warning(
                self,
                "Selection Required",
                "Please select text in the Source cell before adding to non-translatables."
            )
            return
        
        # Find main window and call add_to_nt method
        table = self.table_ref if hasattr(self, 'table_ref') else self.parent()
        if table:
            main_window = table.parent()
            while main_window and not hasattr(main_window, 'add_text_to_non_translatables'):
                main_window = main_window.parent()
            
            if main_window and hasattr(main_window, 'add_text_to_non_translatables'):
                main_window.add_text_to_non_translatables(selected_text)
            else:
                from PyQt6.QtWidgets import QMessageBox
                QMessageBox.warning(
                    self,
                    "Feature Not Available",
                    "Non-translatables functionality not available."
                )

    def set_background_color(self, color: str):
        """Set the background color for this text editor (for alternating row colors)"""
        self.setStyleSheet(f"""
            QTextEdit {{
                border: none;
                background-color: {color};
                padding: 0px;
            }}
        """)


class TagHighlighter(QSyntaxHighlighter):
    """Syntax highlighter for HTML/XML tags, CAT tool tags, CafeTran pipe symbols, and spell checking in text editors"""

    # Class-level reference to spellcheck manager (shared across all instances)
    _spellcheck_manager = None
    _spellcheck_enabled = False
    _is_cafetran_project = False  # Only highlight pipe symbols for CafeTran projects

    def __init__(self, document, tag_color='#7f0001', invisible_char_color='#999999', enable_spellcheck=False):
        super().__init__(document)
        self.tag_color = tag_color
        self.invisible_char_color = invisible_char_color
        self._local_spellcheck_enabled = enable_spellcheck
        self.update_tag_format()

    @classmethod
    def set_spellcheck_manager(cls, manager):
        """Set the shared spellcheck manager instance"""
        cls._spellcheck_manager = manager

    @classmethod
    def set_spellcheck_enabled(cls, enabled: bool):
        """Enable or disable spellchecking globally"""
        cls._spellcheck_enabled = enabled

    def update_tag_format(self):
        """Update the tag format with current color"""
        from PyQt6.QtGui import QTextCharFormat, QColor
        self.tag_format = QTextCharFormat()
        self.tag_format.setForeground(QColor(self.tag_color))

        # CafeTran pipe symbols - red and bold like in CafeTran
        self.pipe_format = QTextCharFormat()
        self.pipe_format.setForeground(QColor('#FF0000'))  # Red
        self.pipe_format.setFontWeight(700)  # Bold

        # Invisible character symbols - use configured color
        self.invisible_format = QTextCharFormat()
        self.invisible_format.setForeground(QColor(self.invisible_char_color))

        # Spellcheck format - red wavy underline
        self.spellcheck_format = QTextCharFormat()
        self.spellcheck_format.setUnderlineColor(QColor('#FF0000'))
        self.spellcheck_format.setUnderlineStyle(QTextCharFormat.UnderlineStyle.WaveUnderline)

    def set_tag_color(self, color: str):
        """Update tag highlight color"""
        self.tag_color = color
        self.update_tag_format()
        self.rehighlight()

    def set_invisible_char_color(self, color: str):
        """Update invisible character color"""
        self.invisible_char_color = color
        self.update_tag_format()
        self.rehighlight()

    def set_local_spellcheck(self, enabled: bool):
        """Enable or disable spellcheck for this specific highlighter"""
        self._local_spellcheck_enabled = enabled
        self.rehighlight()
    
    def highlightBlock(self, text):
        """Highlight all tags, pipe symbols, invisible chars, and misspelled words in the text block"""
        import re
        # Combined pattern for ALL CAT tool tag types:
        # 1. HTML/XML: <tag>, </tag>, <tag/>, <tag attr="val">
        # 2. Trados numeric: <1>, </1>
        # 3. memoQ numeric bracket tags:
        #    - Opening: [1}, [2} etc.
        #    - Closing: {1], {2] etc.
        #    - Standalone: [1], [2] etc.
        # 4. memoQ content tags with text/attributes (from bilingual DOCX):
        #    - [uicontrol id="GUID-..."], [image cid="..." href="..."], etc.
        #    - {uicontrol}, {image}, etc. (closing tags)
        #    NOTE: Opening [tag] MUST have attributes (space+content) to avoid matching
        #          placeholders like [Company] or [Bedrijf]. Closing {tag} doesn't need attrs.
        tag_patterns = [
            r'</?[a-zA-Z][a-zA-Z0-9-]*/?(?:\s[^>]*)?>',  # HTML/XML tags
            r'</?\d+>',                                   # Trados numeric: <1>, </1>
            r'\[\d+[}\]]',                                # memoQ numeric: [1}, [1]
            r'\{\d+[}\]]',                                # memoQ numeric: {1}, {1]
            r'\[[^}\]]+\}',                               # memoQ mixed: [anything} (exclude } and ])
            r'\{[^\[\]]+\]',                              # memoQ mixed: {anything] (exclude [ and ])
            r'\[[a-zA-Z][^}\]]*\s[^}\]]*\]',              # memoQ content: [tag attr...] (exclude } and ])
            r'\{[a-zA-Z][a-zA-Z0-9_-]*\}',                # memoQ closing: {uicontrol}, {MQ}
            r'\{\d{5}\}',                                 # Déjà Vu tags: {00108}, {00109}, etc.
        ]
        combined_pattern = re.compile('|'.join(tag_patterns))

        matches_found = list(combined_pattern.finditer(text))

        for match in matches_found:
            start = match.start()
            length = match.end() - start
            self.setFormat(start, length, self.tag_format)

        # Match invisible character symbols (light blue)
        for i, char in enumerate(text):
            if char in '·→°¶':  # Invisible character replacement symbols
                self.setFormat(i, 1, self.invisible_format)
        
        # CafeTran pipe symbols (red and bold) - ONLY for CafeTran projects
        if TagHighlighter._is_cafetran_project:
            for i, char in enumerate(text):
                if char == '|':
                    self.setFormat(i, 1, self.pipe_format)

        # Spell checking - only for target editors when enabled
        if self._local_spellcheck_enabled and TagHighlighter._spellcheck_enabled and TagHighlighter._spellcheck_manager:
            self._highlight_misspelled_words(text)

    def _highlight_misspelled_words(self, text):
        """Highlight misspelled words with red wavy underline"""
        # Safety check - if spellcheck manager is not available or disabled, skip
        if not TagHighlighter._spellcheck_manager:
            return
        
        # Check if spellcheck manager detected a crash - if so, skip entirely
        if hasattr(TagHighlighter._spellcheck_manager, '_crash_detected') and TagHighlighter._spellcheck_manager._crash_detected:
            return
        
        import re
        # Find all words (letters only, including accented characters)
        # Skip words inside tags or that look like technical content
        word_pattern = re.compile(r'\b([a-zA-ZÀ-ÿ\']+)\b', re.UNICODE)
        
        try:
            for match in word_pattern.finditer(text):
                word = match.group(1)
                start = match.start(1)
                length = len(word)
                
                # Skip very short words and words with apostrophes at start/end
                if len(word) < 2:
                    continue
                
                # Check if this word is inside ANY type of tag:
                # - HTML/XML: < ... >
                # - memoQ/DITA: [ ... } or [ ... ] or { ... ]
                before_text = text[:start]
                
                # Check for HTML tags
                last_html_open = before_text.rfind('<')
                last_html_close = before_text.rfind('>')
                if last_html_open > last_html_close:
                    continue  # Inside HTML tag
                
                # Check for bracket tags: [ ... } or [ ... ]
                last_square_open = before_text.rfind('[')
                last_square_close = max(before_text.rfind('}'), before_text.rfind(']'))
                if last_square_open > last_square_close:
                    continue  # Inside [tag...} or [tag...]
                
                # Check for curly tags: { ... ] or { ... }
                last_curly_open = before_text.rfind('{')
                last_curly_close = max(before_text.rfind(']'), before_text.rfind('}'))
                if last_curly_open > last_curly_close:
                    continue  # Inside {tag...] or {tag...}
                
                # Check spelling
                if not TagHighlighter._spellcheck_manager.check_word(word):
                    self.setFormat(start, length, self.spellcheck_format)
        except Exception as e:
            # If anything goes wrong during spellcheck highlighting, disable it
            print(f"Spellcheck highlighting error: {e}")
            if TagHighlighter._spellcheck_manager:
                TagHighlighter._spellcheck_manager._crash_detected = True
                TagHighlighter._spellcheck_manager.enabled = False


class EditableGridTextEditor(QTextEdit):
    """Editable QTextEdit for target cells - allows text selection and editing"""
    
    # Class variable for tag highlight color (shared across all instances)
    tag_highlight_color = '#7f0001'  # Default memoQ dark red
    
    # Class variables for focus border customization
    focus_border_color = '#f1b79a'  # Default peach/salmon
    focus_border_thickness = 2  # Default 2px (slightly thicker than before)
    
    def __init__(self, text: str = "", parent=None, row: int = -1, table=None):
        super().__init__(parent)
        self.row = row  # Store row number for auto-selection
        self.table = table  # Store table reference
        self.setReadOnly(False)

        # CRITICAL: Track initial load state to ignore spurious textChanged events
        # Qt queues document change events during setPlainText() even with blockSignals(True).
        # When signals are unblocked later, Qt delivers these queued events, causing false
        # textChanged signals. This flag allows us to ignore the first event after unblocking.
        self._initial_load_complete = False

        # CRITICAL: Block ALL signals during initialization to prevent textChanged from firing
        # when setText is called. Signals will be unblocked AFTER signal handler is connected
        # in load_segments_to_grid. This prevents ANY textChanged events during grid loading.
        self.blockSignals(True)
        self.setPlainText(text)
        # DO NOT unblock signals here - they will be unblocked after handler connection

        # CRITICAL: Enable strong focus to receive Tab key events
        self.setFocusPolicy(Qt.FocusPolicy.StrongFocus)

        # Make inactive selections stay visible with same color
        from PyQt6.QtGui import QPalette
        palette = self.palette()
        palette.setColor(QPalette.ColorGroup.Active, QPalette.ColorRole.Highlight, QColor("#D0E7FF"))
        palette.setColor(QPalette.ColorGroup.Active, QPalette.ColorRole.HighlightedText, QColor("black"))
        palette.setColor(QPalette.ColorGroup.Inactive, QPalette.ColorRole.Highlight, QColor("#D0E7FF"))
        palette.setColor(QPalette.ColorGroup.Inactive, QPalette.ColorRole.HighlightedText, QColor("black"))
        self.setPalette(palette)

        # Add syntax highlighter for tags (with spellcheck enabled for target cells)
        # Get invisible char color from main window if available
        main_window = self._get_main_window()
        invisible_char_color = main_window.invisible_char_color if main_window and hasattr(main_window, 'invisible_char_color') else '#999999'
        self.highlighter = TagHighlighter(self.document(), self.tag_highlight_color, invisible_char_color, enable_spellcheck=True)

        # Style to look like a normal cell with subtle selection
        # Background and text colors now managed by theme system
        # Border color and thickness use class variables for user customization
        border_color = EditableGridTextEditor.focus_border_color
        border_thickness = EditableGridTextEditor.focus_border_thickness
        self.setStyleSheet(f"""
            QTextEdit {{
                border: none;
                padding: 0px 4px 0px 0px;
            }}
            QTextEdit:focus {{
                border: {border_thickness}px solid {border_color};
            }}
            QTextEdit::selection {{
                background-color: #D0E7FF;
                color: black;
            }}
        """)

        # Set document margins to 0 for compact display
        doc = self.document()
        doc.setDocumentMargin(0)

        # Configure text option for minimal line spacing
        # With zero-width spaces inserted after invisible character markers,
        # normal WordWrap will work correctly
        wrap_mode = QTextOption.WrapMode.WordWrap

        self.setWordWrapMode(wrap_mode)
        self.setAcceptRichText(False)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)

        text_option = QTextOption()
        text_option.setWrapMode(wrap_mode)
        doc.setDefaultTextOption(text_option)
        
        # Set minimum height to 0 - let content determine size
        self.setMinimumHeight(0)
        self.setMaximumHeight(16777215)  # Qt's max int
    
    def mouseDoubleClickEvent(self, event):
        """Handle double-click to select words properly when invisibles are shown"""
        from PyQt6.QtGui import QTextCursor
        
        main_window = self._get_main_window()
        # Check if invisible spaces are being shown
        if main_window and hasattr(main_window, 'invisible_display_settings'):
            if main_window.invisible_display_settings.get('spaces', False):
                # Get cursor position at click
                cursor = self.cursorForPosition(event.pos())
                pos = cursor.position()
                text = self.toPlainText()
                
                # Find word boundaries using middle dot (·) or zero-width space as delimiters
                # Find start of word (search backwards for · or start of text)
                start = pos
                while start > 0 and text[start - 1] not in ('·', '\u200B', '\n', '\t', '→', '°', '¶'):
                    start -= 1
                
                # Find end of word (search forwards for · or end of text)
                end = pos
                while end < len(text) and text[end] not in ('·', '\u200B', '\n', '\t', '→', '°', '¶'):
                    end += 1
                
                # Select the word
                cursor.setPosition(start)
                cursor.setPosition(end, QTextCursor.MoveMode.KeepAnchor)
                self.setTextCursor(cursor)
                return
        
        # Default behavior
        super().mouseDoubleClickEvent(event)
    
    def _get_main_window(self):
        """Get the main application window by traversing the parent hierarchy"""
        if not self.table:
            return None
        main_window = self.table.parent()
        while main_window and not hasattr(main_window, 'go_to_first_segment'):
            main_window = main_window.parent()
        return main_window
    
    def _handle_add_to_termbase(self):
        """Handle Ctrl+T: Add selected source and target terms to termbase"""
        if not self.table or self.row < 0:
            return
        
        # Get target selection (from this widget)
        target_text = self.textCursor().selectedText().strip()
        
        # Get source cell widget and its selection
        source_widget = self.table.cellWidget(self.row, 2)
        source_text = ""
        if source_widget and hasattr(source_widget, 'textCursor'):
            source_text = source_widget.textCursor().selectedText().strip()
        
        # Validate we have both selections
        if not source_text or not target_text:
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.warning(
                self,
                "Selection Required",
                "Please select text in both Source and Target cells before adding to glossary.\\n\\n"
                "Workflow:\\n"
                "1. Select term in source/target cell\\n"
                "2. Press Tab to cycle to other cell\\n"
                "3. Select corresponding term\\n"
                "4. Press Ctrl+E (or right-click) to add to glossary"
            )
            return
        
        # Find main window and call add_to_termbase method
        main_window = self.table.parent()
        while main_window and not hasattr(main_window, 'add_term_pair_to_termbase'):
            main_window = main_window.parent()
        
        if main_window and hasattr(main_window, 'add_term_pair_to_termbase'):
            main_window.add_term_pair_to_termbase(source_text, target_text)
    
    def _handle_quick_add_to_termbase(self):
        """Handle Ctrl+R: Quick add selected source and target terms to glossary (no dialog)"""
        if not self.table or self.row < 0:
            return
        
        # Get target selection (from this widget)
        target_text = self.textCursor().selectedText().strip()
        
        # Get source cell widget and its selection
        source_widget = self.table.cellWidget(self.row, 2)
        source_text = ""
        if source_widget and hasattr(source_widget, 'textCursor'):
            source_text = source_widget.textCursor().selectedText().strip()
        
        # Validate we have both selections
        if not source_text or not target_text:
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.warning(
                self,
                "Selection Required",
                "Please select text in both Source and Target cells before quick-adding to glossary.\n\n"
                "Tip: Use Ctrl+E to add with a dialog where you can choose glossary and add metadata."
            )
            return
        
        # Find main window and call quick_add_to_termbase method
        main_window = self.table.parent()
        while main_window and not hasattr(main_window, 'quick_add_term_pair_to_termbase'):
            main_window = main_window.parent()
        
        if main_window and hasattr(main_window, 'quick_add_term_pair_to_termbase'):
            main_window.quick_add_term_pair_to_termbase(source_text, target_text)
    
    def _handle_superlookup_search(self):
        """Handle right-click: Search selected text in Superlookup"""
        # Get selected text
        selected_text = self.textCursor().selectedText().strip()
        
        if not selected_text:
            return
        
        # Find main window
        main_window = self.table.parent() if self.table else None
        while main_window and not hasattr(main_window, '_go_to_superlookup'):
            main_window = main_window.parent()
        
        if not main_window:
            return
        
        # Get project languages
        source_lang = None
        target_lang = None
        if hasattr(main_window, 'current_project') and main_window.current_project:
            source_lang = main_window.current_project.source_lang
            target_lang = main_window.current_project.target_lang
        
        # Navigate to Superlookup
        main_window._go_to_superlookup()
        
        # Trigger search
        if hasattr(main_window, 'lookup_tab') and main_window.lookup_tab:
            main_window.lookup_tab.search_with_query(
                selected_text,
                switch_to_vertical=True,
                source_lang=source_lang,
                target_lang=target_lang
            )
    
    def contextMenuEvent(self, event):
        """Show context menu with Add to Glossary, Non-Translatables, and Spellcheck options"""
        from PyQt6.QtWidgets import QMenu
        from PyQt6.QtGui import QAction
        
        menu = QMenu(self)
        
        # Check if cursor is on a misspelled word for spellcheck suggestions
        cursor_pos = self.cursorForPosition(event.pos())
        misspelled_word, word_start, word_end = self._get_misspelled_word_at_cursor(cursor_pos)
        
        if misspelled_word:
            # Get spellcheck suggestions
            if TagHighlighter._spellcheck_manager:
                suggestions = TagHighlighter._spellcheck_manager.get_suggestions(misspelled_word)
                
                if suggestions:
                    # Add suggestions at the top
                    for suggestion in suggestions[:5]:  # Limit to 5 suggestions
                        suggestion_action = QAction(f"✓ {suggestion}", self)
                        # Use default argument to capture current value
                        suggestion_action.triggered.connect(
                            lambda checked, s=suggestion, start=word_start, end=word_end: 
                            self._replace_word(start, end, s)
                        )
                        suggestion_action.setFont(menu.font())
                        menu.addAction(suggestion_action)
                    menu.addSeparator()
                
                # Add to dictionary action
                add_to_dict_action = QAction(f"📖 Add '{misspelled_word}' to Dictionary (Alt+D)", self)
                add_to_dict_action.triggered.connect(
                    lambda checked, w=misspelled_word: self._add_to_dictionary(w)
                )
                menu.addAction(add_to_dict_action)
                
                # Ignore this session action
                ignore_action = QAction(f"🔇 Ignore '{misspelled_word}' (this session)", self)
                ignore_action.triggered.connect(
                    lambda checked, w=misspelled_word: self._ignore_word(w)
                )
                menu.addAction(ignore_action)
                menu.addSeparator()
        
        # Add standard edit actions
        if self.textCursor().hasSelection():
            copy_action = QAction("Copy", self)
            copy_action.triggered.connect(self.copy)
            menu.addAction(copy_action)
            
            cut_action = QAction("Cut", self)
            cut_action.triggered.connect(self.cut)
            menu.addAction(cut_action)
            menu.addSeparator()
        
        paste_action = QAction("Paste", self)
        paste_action.triggered.connect(self.paste)
        menu.addAction(paste_action)
        menu.addSeparator()
        
        # Superlookup search action
        if self.textCursor().hasSelection():
            superlookup_action = QAction("🔍 Search in Superlookup (Ctrl+K)", self)
            superlookup_action.triggered.connect(self._handle_superlookup_search)
            menu.addAction(superlookup_action)
            menu.addSeparator()

        # QuickMenu (prompt-based actions)
        try:
            main_window = self.table.parent() if self.table else None
            while main_window and not hasattr(main_window, 'run_grid_quickmenu_prompt'):
                main_window = main_window.parent()

            quickmenu_items = []
            if main_window and hasattr(main_window, 'prompt_manager_qt') and main_window.prompt_manager_qt:
                lib = getattr(main_window.prompt_manager_qt, 'library', None)
                if lib and hasattr(lib, 'get_quickmenu_grid_prompts'):
                    quickmenu_items = lib.get_quickmenu_grid_prompts() or []

            if quickmenu_items:
                qm_menu = menu.addMenu("⚡ QuickMenu")
                for rel_path, label in sorted(quickmenu_items, key=lambda x: (x[1] or x[0]).lower()):
                    prompt_menu = qm_menu.addMenu(label or rel_path)

                    run_show = QAction("▶ Run (show response)…", self)
                    run_show.triggered.connect(
                        lambda checked=False, p=rel_path: main_window.run_grid_quickmenu_prompt(p, origin_widget=self, behavior="show")
                    )
                    prompt_menu.addAction(run_show)

                    run_replace = QAction("↺ Run and replace selection", self)
                    run_replace.triggered.connect(
                        lambda checked=False, p=rel_path: main_window.run_grid_quickmenu_prompt(p, origin_widget=self, behavior="replace")
                    )
                    prompt_menu.addAction(run_replace)

                menu.addSeparator()
        except Exception:
            # Never break the normal context menu due to QuickMenu errors
            pass
        
        # Add to termbase action (with dialog)
        add_to_tb_action = QAction("📖 Add to Glossary (Ctrl+E)", self)
        add_to_tb_action.triggered.connect(self._handle_add_to_termbase)
        menu.addAction(add_to_tb_action)
        
        # Quick add to termbase action (no dialog) - uses last-selected termbase from Ctrl+E
        quick_add_action = QAction("⚡ Quick Add to Glossary (Alt+Left)", self)
        quick_add_action.triggered.connect(self._handle_quick_add_to_termbase)
        menu.addAction(quick_add_action)
        
        # Add to non-translatables action
        add_to_nt_action = QAction("🚫 Add to Non-Translatables (Ctrl+Alt+N)", self)
        add_to_nt_action.triggered.connect(self._handle_add_to_nt)
        menu.addAction(add_to_nt_action)
        
        menu.exec(event.globalPos())

    def _get_misspelled_word_at_cursor(self, cursor):
        """Get the misspelled word at the cursor position, if any"""
        import re
        
        if not TagHighlighter._spellcheck_enabled or not TagHighlighter._spellcheck_manager:
            return None, 0, 0
        
        text = self.toPlainText()
        pos = cursor.position()
        
        # Find word boundaries around cursor position
        word_pattern = re.compile(r'\b([a-zA-ZÀ-ÿ\']+)\b', re.UNICODE)
        
        for match in word_pattern.finditer(text):
            start = match.start(1)
            end = match.end(1)
            
            if start <= pos <= end:
                word = match.group(1)
                # Check if this word is misspelled
                if not TagHighlighter._spellcheck_manager.check_word(word):
                    return word, start, end
                break
        
        return None, 0, 0

    def _replace_word(self, start: int, end: int, replacement: str):
        """Replace a word at the given position with the replacement"""
        cursor = self.textCursor()
        cursor.setPosition(start)
        cursor.setPosition(end, cursor.MoveMode.KeepAnchor)
        cursor.insertText(replacement)

    def _add_to_dictionary(self, word: str):
        """Add word to the custom dictionary"""
        if TagHighlighter._spellcheck_manager:
            TagHighlighter._spellcheck_manager.add_to_dictionary(word)
            
            # Show confirmation in main window log
            main_window = self._get_main_window()
            if main_window and hasattr(main_window, 'log'):
                main_window.log(f"✓ Added '{word}' to custom dictionary")

            # Refresh all highlighters to update other occurrences
            if main_window and hasattr(main_window, '_refresh_all_highlighters'):
                main_window._refresh_all_highlighters()
            else:
                # Fallback to local rehighlight
                self.highlighter.rehighlight()

    def _ignore_word(self, word: str):
        """Ignore word for this session"""
        if TagHighlighter._spellcheck_manager:
            TagHighlighter._spellcheck_manager.ignore_word(word)
            
            # Show confirmation in main window log
            main_window = self._get_main_window()
            if main_window and hasattr(main_window, 'log'):
                main_window.log(f"🔇 Ignoring '{word}' for this session")

            # Refresh all highlighters to update other occurrences
            if main_window and hasattr(main_window, '_refresh_all_highlighters'):
                main_window._refresh_all_highlighters()
            else:
                # Fallback to local rehighlight
                self.highlighter.rehighlight()
    
    def sizeHint(self):
        """Return compact size based on content"""
        doc = self.document()
        ideal_width = self.width() if self.width() > 0 else 200
        doc.setTextWidth(ideal_width)
        height = int(doc.size().height())
        # Add minimal padding (2px total)
        height = max(height + 2, 1)
        current_width = self.width() if self.width() > 0 else 200
        return QSize(current_width, height)
    
    def mousePressEvent(self, event):
        """Allow text selection on click and auto-select row"""
        super().mousePressEvent(event)
        # Auto-select the row when clicking in the target cell
        if self.table and self.row >= 0:
            self.table.selectRow(self.row)
            self.table.setCurrentCell(self.row, 3)  # Column 3 is Target

            # CRITICAL: Manually trigger on_cell_selected since signals aren't firing
            # Find the main window and call the method directly
            try:
                main_window = self.table.parent()
                while main_window and not hasattr(main_window, 'on_cell_selected'):
                    main_window = main_window.parent()
                if main_window and hasattr(main_window, 'on_cell_selected'):
                    main_window.on_cell_selected(self.row, 3, -1, -1)
            except Exception as e:
                print(f"Error triggering manual cell selection: {e}")

    def mouseReleaseEvent(self, event):
        """Smart word selection - expand partial selections to full words
        
        Works across multiple lines - if you select text spanning several lines,
        partial words at the START and END of your selection will be expanded.
        """
        super().mouseReleaseEvent(event)

        # Check if smart selection is enabled
        main_window = self._get_main_window()
        if main_window and hasattr(main_window, 'enable_smart_word_selection'):
            if not main_window.enable_smart_word_selection:
                return  # Feature disabled

        # Get the current cursor
        cursor = self.textCursor()

        # Only expand if there's a selection
        if cursor.hasSelection():
            # Get selection boundaries
            start = cursor.selectionStart()
            end = cursor.selectionEnd()

            # Get the full text
            text = self.toPlainText()

            # Helper function to check if character is part of a word
            # Includes alphanumeric, underscore, hyphen, and apostrophe
            def is_word_char(char):
                return char.isalnum() or char in "_-'"

            # Track if we need to update the selection
            selection_changed = False

            # Expand START boundary if we're in the middle of a word
            # (i.e., the character before the selection is a word character)
            if start > 0 and is_word_char(text[start - 1]):
                # Also check that the first selected character is a word char
                # (to avoid expanding when selecting from whitespace)
                if start < len(text) and is_word_char(text[start]):
                    while start > 0 and is_word_char(text[start - 1]):
                        start -= 1
                    selection_changed = True

            # Expand END boundary if we're in the middle of a word
            # (i.e., the character after the selection is a word character)
            if end < len(text) and is_word_char(text[end]):
                # Also check that the last selected character is a word char
                # (to avoid expanding when selecting to whitespace)
                if end > 0 and is_word_char(text[end - 1]):
                    while end < len(text) and is_word_char(text[end]):
                        end += 1
                    selection_changed = True

            # Set the new selection if boundaries changed
            if selection_changed:
                cursor.setPosition(start)
                cursor.setPosition(end, cursor.MoveMode.KeepAnchor)
                self.setTextCursor(cursor)

    def focusInEvent(self, event):
        """Ensure text remains visible when focused and auto-select row"""
        super().focusInEvent(event)
        # Ensure the widget is properly visible
        self.setVisible(True)
        self.show()
        # Auto-select the row when focusing the target cell
        if self.table and self.row >= 0:
            self.table.selectRow(self.row)
            self.table.setCurrentCell(self.row, 3)  # Column 3 is Target
            
            # CRITICAL: Manually trigger on_cell_selected since signals aren't firing
            # Find the main window and call the method directly
            try:
                main_window = self.table.parent()
                while main_window and not hasattr(main_window, 'on_cell_selected'):
                    main_window = main_window.parent()
                if main_window and hasattr(main_window, 'on_cell_selected'):
                    main_window.on_cell_selected(self.row, 3, -1, -1)
            except Exception as e:
                print(f"Error triggering manual cell selection: {e}")
    
    def keyPressEvent(self, event):
        """Handle Tab and Ctrl+E keys to cycle between source and target cells"""
        from PyQt6.QtWidgets import QApplication
        from PyQt6.QtGui import QTextCursor
        
        # Ctrl+C: Fix clipboard when copying with invisible characters shown
        if event.key() == Qt.Key.Key_C and event.modifiers() == Qt.KeyboardModifier.ControlModifier:
            cursor = self.textCursor()
            if cursor.hasSelection():
                selected_text = cursor.selectedText()
                # Reverse invisible character replacements before copying
                main_window = self._get_main_window()
                if main_window and hasattr(main_window, 'reverse_invisible_replacements'):
                    clean_text = main_window.reverse_invisible_replacements(selected_text)
                    # Also replace paragraph separator with newline (Qt uses U+2029)
                    clean_text = clean_text.replace('\u2029', '\n')
                    # Set clipboard with clean text
                    clipboard = QApplication.clipboard()
                    clipboard.setText(clean_text)
                    event.accept()
                    return
        
        # Handle Ctrl+Arrow word navigation when invisibles are shown
        main_window = self._get_main_window()
        if main_window and hasattr(main_window, 'invisible_display_settings'):
            if main_window.invisible_display_settings.get('spaces', False):
                ctrl_only = event.modifiers() == Qt.KeyboardModifier.ControlModifier
                ctrl_shift = event.modifiers() == (Qt.KeyboardModifier.ControlModifier | Qt.KeyboardModifier.ShiftModifier)
                
                if (ctrl_only or ctrl_shift) and event.key() in (Qt.Key.Key_Left, Qt.Key.Key_Right):
                    cursor = self.textCursor()
                    text = self.toPlainText()
                    pos = cursor.position()
                    word_chars = ('·', '\u200B', '\n', '\t', '→', '°', '¶', ' ')
                    
                    if event.key() == Qt.Key.Key_Right:
                        # Move to end of current word, then skip delimiters to start of next word
                        # First skip any delimiters we're on
                        while pos < len(text) and text[pos] in word_chars:
                            pos += 1
                        # Then skip to end of word (next delimiter)
                        while pos < len(text) and text[pos] not in word_chars:
                            pos += 1
                    else:  # Key_Left
                        # Move backwards: skip delimiters, then find start of word
                        if pos > 0:
                            pos -= 1
                        # Skip any delimiters
                        while pos > 0 and text[pos] in word_chars:
                            pos -= 1
                        # Find start of word
                        while pos > 0 and text[pos - 1] not in word_chars:
                            pos -= 1
                    
                    # Apply cursor movement
                    if ctrl_shift:
                        cursor.setPosition(pos, QTextCursor.MoveMode.KeepAnchor)
                    else:
                        cursor.setPosition(pos)
                    self.setTextCursor(cursor)
                    event.accept()
                    return
        
        # Ctrl+E: Add selected terms to termbase (with dialog)
        if event.key() == Qt.Key.Key_E and event.modifiers() == Qt.KeyboardModifier.ControlModifier:
            self._handle_add_to_termbase()
            event.accept()
            return
        
        # Alt+Left: Quick add selected terms to last-used termbase (no dialog)
        if event.key() == Qt.Key.Key_Left and event.modifiers() == Qt.KeyboardModifier.AltModifier:
            self._handle_quick_add_to_termbase()
            event.accept()
            return
        
        # Ctrl+Alt+N: Add selected text to non-translatables
        if event.key() == Qt.Key.Key_N and event.modifiers() == (Qt.KeyboardModifier.ControlModifier | Qt.KeyboardModifier.AltModifier):
            self._handle_add_to_nt()
            event.accept()
            return
        
        # Ctrl+Home: Navigate to first segment (pass to main window)
        if event.key() == Qt.Key.Key_Home and event.modifiers() == Qt.KeyboardModifier.ControlModifier:
            main_window = self._get_main_window()
            if main_window and hasattr(main_window, 'go_to_first_segment'):
                main_window.go_to_first_segment()
            event.accept()
            return
        
        # Ctrl+End: Navigate to last segment (pass to main window)
        if event.key() == Qt.Key.Key_End and event.modifiers() == Qt.KeyboardModifier.ControlModifier:
            main_window = self._get_main_window()
            if main_window and hasattr(main_window, 'go_to_last_segment'):
                main_window.go_to_last_segment()
            event.accept()
            return
        
        # Ctrl+, (comma): Insert next memoQ tag or wrap selection with tag pair
        if event.key() == Qt.Key.Key_Comma and event.modifiers() == Qt.KeyboardModifier.ControlModifier:
            self._insert_next_tag_or_wrap_selection()
            event.accept()
            return
        
        # Ctrl+Shift+S: Copy source text to target
        if event.key() == Qt.Key.Key_S and event.modifiers() == (Qt.KeyboardModifier.ControlModifier | Qt.KeyboardModifier.ShiftModifier):
            self._copy_source_to_target()
            event.accept()
            return
        
        # Ctrl+B: Apply bold formatting to selection
        if event.key() == Qt.Key.Key_B and event.modifiers() == Qt.KeyboardModifier.ControlModifier:
            self._apply_formatting_tag('b')
            event.accept()
            return
        
        # Ctrl+I: Apply italic formatting to selection
        if event.key() == Qt.Key.Key_I and event.modifiers() == Qt.KeyboardModifier.ControlModifier:
            self._apply_formatting_tag('i')
            event.accept()
            return
        
        # Ctrl+U: Apply underline formatting to selection
        if event.key() == Qt.Key.Key_U and event.modifiers() == Qt.KeyboardModifier.ControlModifier:
            self._apply_formatting_tag('u')
            event.accept()
            return
        
        # Ctrl+Enter: Confirm & Next (call main window method directly)
        if event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter) and event.modifiers() == Qt.KeyboardModifier.ControlModifier:
            main_window = self._get_main_window()
            if main_window and hasattr(main_window, 'confirm_selected_or_next'):
                main_window.confirm_selected_or_next()
            event.accept()
            return
        
        # Shift+Enter: Insert line break (for multi-line content)
        if event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter) and event.modifiers() == Qt.KeyboardModifier.ShiftModifier:
            super().keyPressEvent(event)
            event.accept()
            return
        
        # Plain Enter: Don't insert newline, just accept
        if event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter) and event.modifiers() == Qt.KeyboardModifier.NoModifier:
            event.accept()
            return
        
        # Ctrl+Tab: Insert actual tab character
        if event.key() == Qt.Key.Key_Tab and event.modifiers() == Qt.KeyboardModifier.ControlModifier:
            self.insertPlainText('\t')
            event.accept()
            return
        
        # Tab alone: Cycle to source cell (column 2) in same row
        if event.key() == Qt.Key.Key_Tab and event.modifiers() == Qt.KeyboardModifier.NoModifier:
            if self.table and self.row >= 0:
                # Get source cell widget (column 2)
                source_widget = self.table.cellWidget(self.row, 2)
                if source_widget:
                    source_widget.setFocus()
                    self.table.setCurrentCell(self.row, 2)
                    event.accept()
                    return
        
        # Arrow Up/Down: memoQ-style segment navigation at cell boundaries
        # When cursor is at top VISUAL line and Up is pressed, go to previous segment
        # When cursor is at bottom VISUAL line and Down is pressed, go to next segment
        if event.key() in (Qt.Key.Key_Up, Qt.Key.Key_Down) and event.modifiers() == Qt.KeyboardModifier.NoModifier:
            cursor = self.textCursor()
            current_block = cursor.block()
            doc = self.document()
            first_block = doc.firstBlock()
            last_block = doc.lastBlock()
            
            # Get the visual line number within the current block
            layout = current_block.layout()
            if layout:
                pos_in_block = cursor.positionInBlock()
                current_visual_line = layout.lineForTextPosition(pos_in_block)
                current_line_num = current_visual_line.lineNumber() if current_visual_line.isValid() else 0
                total_lines_in_block = layout.lineCount()
            else:
                current_line_num = 0
                total_lines_in_block = 1
            
            if event.key() == Qt.Key.Key_Up:
                # Only navigate to previous segment if we're on the FIRST visual line of the FIRST block
                is_first_visual_line = (current_block == first_block and current_line_num == 0)
                if is_first_visual_line:
                    main_window = self._get_main_window()
                    if main_window and hasattr(main_window, 'go_to_previous_segment'):
                        # Get cursor column within current visual line
                        if layout and current_visual_line.isValid():
                            col_in_line = pos_in_block - int(current_visual_line.textStart())
                        else:
                            col_in_line = cursor.positionInBlock()
                        # Signal rapid navigation for performance optimization
                        main_window._arrow_key_navigation = True
                        main_window.go_to_previous_segment(target_column=col_in_line, to_last_line=True)
                        event.accept()
                        return
            
            elif event.key() == Qt.Key.Key_Down:
                # Only navigate to next segment if we're on the LAST visual line of the LAST block
                is_last_visual_line = (current_block == last_block and current_line_num >= total_lines_in_block - 1)
                if is_last_visual_line:
                    main_window = self._get_main_window()
                    if main_window and hasattr(main_window, 'go_to_next_segment'):
                        # Get cursor column within current visual line
                        if layout and current_visual_line.isValid():
                            col_in_line = pos_in_block - int(current_visual_line.textStart())
                        else:
                            col_in_line = cursor.positionInBlock()
                        # Signal rapid navigation for performance optimization
                        main_window._arrow_key_navigation = True
                        main_window.go_to_next_segment(target_column=col_in_line, to_first_line=True)
                        event.accept()
                        return
        
        # All other keys: Handle normally
        super().keyPressEvent(event)
    
    def _handle_add_to_nt(self):
        """Handle Ctrl+Alt+N: Add selected text to active non-translatable list(s)"""
        # Get selected text from source cell (for NT, we typically add from source)
        # But if this is target and source is available, use source
        selected_text = self.textCursor().selectedText().strip()
        
        # If no selection in target, try getting from source
        if not selected_text and self.table and self.row >= 0:
            source_widget = self.table.cellWidget(self.row, 2)
            if source_widget and hasattr(source_widget, 'textCursor'):
                selected_text = source_widget.textCursor().selectedText().strip()
        
        if not selected_text:
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.warning(
                self,
                "Selection Required",
                "Please select text before adding to non-translatables."
            )
            return
        
        # Find main window and call add_to_nt method
        if self.table:
            main_window = self.table.parent()
            while main_window and not hasattr(main_window, 'add_text_to_non_translatables'):
                main_window = main_window.parent()
            
            if main_window and hasattr(main_window, 'add_text_to_non_translatables'):
                main_window.add_text_to_non_translatables(selected_text)
            else:
                from PyQt6.QtWidgets import QMessageBox
                QMessageBox.warning(
                    self,
                    "Feature Not Available",
                    "Non-translatables functionality not available."
                )

    def _insert_next_tag_or_wrap_selection(self):
        """
        Insert the next memoQ tag, HTML tag, or CafeTran pipe symbol from source, or wrap selection.
        
        Behavior:
        - If text is selected: Wrap it with the next available tag pair [N}selection{N] or |selection|
        - If no selection: Insert the next unused tag/pipe from source at cursor position
        
        Supports:
        - memoQ tags: [1}, {1], [2}, {2], etc.
        - HTML/XML tags: <li>, </li>, <b>, </b>, <i>, </i>, etc.
        - CafeTran pipe symbols: |
        
        Shortcut: Ctrl+, (comma)
        """
        if not self.table or self.row < 0:
            return
        
        # Navigate up to find main window
        main_window = self.table.parent()
        while main_window and not hasattr(main_window, 'current_project'):
            main_window = main_window.parent()
        
        if not main_window or not hasattr(main_window, 'current_project'):
            return
        
        if not main_window.current_project or self.row >= len(main_window.current_project.segments):
            return
        
        segment = main_window.current_project.segments[self.row]
        source_text = segment.source
        current_target = self.toPlainText()
        
        # Check what type of tags are in the source
        has_memoq_tags = bool(extract_memoq_tags(source_text))
        has_html_tags = bool(extract_html_tags(source_text))
        has_any_tags = has_memoq_tags or has_html_tags
        has_pipe_symbols = '|' in source_text
        
        # Check if there's a selection
        cursor = self.textCursor()
        if cursor.hasSelection():
            selected_text = cursor.selectedText()
            
            # Try memoQ tag pair first
            if has_memoq_tags:
                opening_tag, closing_tag = get_wrapping_tag_pair(source_text, current_target)
                if opening_tag and closing_tag:
                    wrapped_text = f"{opening_tag}{selected_text}{closing_tag}"
                    cursor.insertText(wrapped_text)
                    if hasattr(main_window, 'log'):
                        main_window.log(f"🏷️ Wrapped selection with {opening_tag}...{closing_tag}")
                    return
            
            # Try CafeTran pipe symbols
            if has_pipe_symbols:
                pipes_needed = get_next_pipe_count_needed(source_text, current_target)
                if pipes_needed >= 2:
                    # Wrap with pipes
                    wrapped_text = f"|{selected_text}|"
                    cursor.insertText(wrapped_text)
                    if hasattr(main_window, 'log'):
                        main_window.log(f"🏷️ Wrapped selection with |...|")
                    return
            
            if hasattr(main_window, 'log'):
                main_window.log("⚠️ No tag pairs available from source")
        else:
            # No selection - insert next unused tag or pipe at cursor
            
            # Try memoQ tags and HTML tags (find_next_unused_tag handles both)
            if has_any_tags:
                next_tag = find_next_unused_tag(source_text, current_target)
                if next_tag:
                    cursor.insertText(next_tag)
                    if hasattr(main_window, 'log'):
                        main_window.log(f"🏷️ Inserted tag: {next_tag}")
                    return
            
            # Try CafeTran pipe symbols
            if has_pipe_symbols:
                pipes_needed = get_next_pipe_count_needed(source_text, current_target)
                if pipes_needed > 0:
                    cursor.insertText('|')
                    if hasattr(main_window, 'log'):
                        main_window.log(f"🏷️ Inserted pipe symbol (|)")
                    return
            
            if hasattr(main_window, 'log'):
                main_window.log("✓ All tags from source already in target")
    
    def _copy_source_to_target(self):
        """
        Copy source text to target cell.
        
        Shortcut: Ctrl+Shift+S
        """
        if not self.table or self.row < 0:
            return
        
        # Navigate up to find main window
        main_window = self.table.parent()
        while main_window and not hasattr(main_window, 'current_project'):
            main_window = main_window.parent()
        
        if not main_window or not hasattr(main_window, 'current_project'):
            return
        
        if not main_window.current_project or self.row >= len(main_window.current_project.segments):
            return
        
        segment = main_window.current_project.segments[self.row]
        source_text = segment.source
        
        # Set the target text
        self.setPlainText(source_text)
        
        if hasattr(main_window, 'log'):
            main_window.log(f"📋 Copied source to target (segment {self.row + 1})")

    def _apply_formatting_tag(self, tag: str):
        """
        Apply or toggle a formatting tag on the selected text.
        
        Args:
            tag: The tag to apply ('b', 'i', or 'u')
        """
        cursor = self.textCursor()
        if not cursor.hasSelection():
            return
        
        selected_text = cursor.selectedText()
        
        # Check if the text is already wrapped with this tag
        open_tag = f"<{tag}>"
        close_tag = f"</{tag}>"
        
        # Get the full text and selection position
        full_text = self.toPlainText()
        start = cursor.selectionStart()
        end = cursor.selectionEnd()
        
        # Check for existing tags just before/after selection
        prefix_has_tag = start >= len(open_tag) and full_text[start - len(open_tag):start] == open_tag
        suffix_has_tag = end + len(close_tag) <= len(full_text) and full_text[end:end + len(close_tag)] == close_tag
        
        if prefix_has_tag and suffix_has_tag:
            # Remove the tags (toggle off)
            new_text = full_text[:start - len(open_tag)] + selected_text + full_text[end + len(close_tag):]
            cursor.setPosition(start - len(open_tag))
            cursor.setPosition(end + len(close_tag), cursor.MoveMode.KeepAnchor)
            cursor.insertText(selected_text)
        else:
            # Add the tags (toggle on)
            wrapped_text = f"{open_tag}{selected_text}{close_tag}"
            cursor.insertText(wrapped_text)
            
            # Re-select the wrapped text (including tags)
            cursor.setPosition(start)
            cursor.setPosition(start + len(wrapped_text), cursor.MoveMode.KeepAnchor)
            self.setTextCursor(cursor)

    def update_display_mode(self, text: str, show_tags: bool):
        """
        Update the display based on tag view mode.
        
        Args:
            text: The raw text (with HTML tags like <b>bold</b>)
            show_tags: If True, show raw tags. If False, show formatted WYSIWYG.
        """
        # Store raw text as property for later retrieval
        self._raw_text = text
        
        self.blockSignals(True)
        if show_tags:
            # Tag view: Show plain text with visible tags
            # The TagHighlighter will colorize the tags
            self.setPlainText(text)
        else:
            # WYSIWYG view: Apply formatting
            if has_formatting_tags(text):
                html = get_formatted_html_display(text)
                self.setHtml(html)
            else:
                # No tags, just plain text
                self.setPlainText(text)
        self.blockSignals(False)
    
    def get_raw_text(self) -> str:
        """Get the raw text with tags, regardless of display mode."""
        return getattr(self, '_raw_text', self.toPlainText())

    def set_background_color(self, color: str):
        """Set the background color for this text editor (for alternating row colors)"""
        # Use class variables for border settings to respect user customization
        border_color = EditableGridTextEditor.focus_border_color
        border_thickness = EditableGridTextEditor.focus_border_thickness
        self.setStyleSheet(f"""
            QTextEdit {{
                border: none;
                background-color: {color};
                padding: 0px 4px 0px 0px;
            }}
            QTextEdit:focus {{
                border: {border_thickness}px solid {border_color};
            }}
            QTextEdit::selection {{
                background-color: #D0E7FF;
                color: black;
            }}
        """)


class SearchHighlightDelegate(QStyledItemDelegate):
    """
    Custom delegate that highlights search terms in cells while preserving full editability.
    Uses custom painting to draw highlights underneath the text.
    """

    def __init__(self, parent=None):
        super().__init__(parent)
        self.search_terms = {}  # Dict of {(row, col): search_term}
        self.highlight_color = QColor(255, 255, 0, 180)  # Yellow with some transparency

    def set_highlight(self, row: int, col: int, search_term: str):
        """Set a search term to highlight for a specific cell"""
        self.search_terms[(row, col)] = search_term

    def clear_highlight(self, row: int, col: int):
        """Clear highlight for a specific cell"""
        if (row, col) in self.search_terms:
            del self.search_terms[(row, col)]

    def clear_all_highlights(self):
        """Clear all highlights"""
        self.search_terms.clear()

    def paint(self, painter, option, index):
        """Custom paint method that highlights search terms"""
        row = index.row()
        col = index.column()

        # Check if this cell has a search term to highlight
        if (row, col) in self.search_terms:
            search_term = self.search_terms[(row, col)]
            text = index.data(Qt.ItemDataRole.DisplayRole) or ""

            if text and search_term:
                # Save painter state
                painter.save()

                # Draw the background first (handles selection, etc.)
                option_copy = QStyleOptionViewItem(option)
                self.initStyleOption(option_copy, index)

                # Get the style
                style = option.widget.style() if option.widget else QApplication.style()

                # Draw background
                style.drawPrimitive(QStyle.PrimitiveElement.PE_PanelItemViewItem, option_copy, painter, option.widget)

                # Calculate text rect
                text_rect = style.subElementRect(QStyle.SubElement.SE_ItemViewItemText, option_copy, option.widget)

                # Get font metrics
                fm = painter.fontMetrics()

                # Find and highlight all occurrences of the search term
                text_lower = text.lower()
                search_lower = search_term.lower()
                pos = 0

                while True:
                    found = text_lower.find(search_lower, pos)
                    if found == -1:
                        break

                    # Calculate the position of this occurrence
                    prefix = text[:found]
                    prefix_width = fm.horizontalAdvance(prefix)
                    term_width = fm.horizontalAdvance(text[found:found + len(search_term)])

                    # Draw highlight rectangle
                    highlight_rect = QRectF(
                        text_rect.left() + prefix_width,
                        text_rect.top(),
                        term_width,
                        text_rect.height()
                    )
                    painter.fillRect(highlight_rect, self.highlight_color)

                    pos = found + len(search_term)

                # Draw the text on top
                painter.setPen(option_copy.palette.color(QPalette.ColorRole.Text))
                painter.drawText(text_rect, Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter, text)

                painter.restore()
                return

        # Default painting for non-highlighted cells
        super().paint(painter, option, index)


class ClickableHighlightLabel(QLabel):
    """
    Custom QLabel for highlighted search results that allows double-click to edit.
    When double-clicked, it removes itself from the table cell and triggers editing.

    Mouse events are transparent by default - the table's viewport handles them.
    We use installEventFilter on the table's viewport to catch double-clicks.
    """

    def __init__(self, table, row: int, col: int, plain_text: str, parent=None):
        super().__init__(parent)
        self.table = table
        self.row = row
        self.col = col
        self.plain_text = plain_text  # Store the original plain text for restoration
        # Make mouse events pass through to the table underneath
        self.setAttribute(Qt.WidgetAttribute.WA_TransparentForMouseEvents, True)


class TermbaseHighlightWidget(QLabel):
    """Custom label widget that displays source text with termbase highlights and tooltips"""
    
    # Signal for term double-click
    term_double_clicked = None
    
    def __init__(self, source_text: str, termbase_matches: Optional[Dict[str, str]] = None, parent=None):
        """
        Args:
            source_text: The full source text to display
            termbase_matches: Dict of {term: translation} for highlighted terms
            parent: Parent widget
        """
        super().__init__(parent)
        self.source_text = source_text
        self.termbase_matches = termbase_matches if termbase_matches is not None else {}
        self.term_ranges = {}  # Store {term: (start_pos, end_pos)}
        self.setWordWrap(True)
        self.setup_display()
        self.setMouseTracking(True)
    
    def setup_display(self):
        """Create rich text display with highlighted terms"""
        html = self._create_highlighted_html()
        self.setText(html)
        self.setTextFormat(Qt.TextFormat.RichText)
    
    def _create_highlighted_html(self) -> str:
        """Create HTML with highlighted termbase matches using priority colors"""
        text = self.source_text
        
        # Sort matches by source term length (longest first) to avoid overlaps
        # Since dict keys are now term_ids, extract source terms first
        term_entries = []
        for term_id, match_info in self.termbase_matches.items():
            if isinstance(match_info, dict):
                source_term = match_info.get('source', '')
                if source_term:
                    term_entries.append((source_term, match_info))
        
        term_entries.sort(key=lambda x: len(x[0]), reverse=True)
        
        # Build HTML with highlights
        html = text
        offset = 0
        
        for term, match_info in term_entries:
            # Case-insensitive search
            search_term = term.lower()
            search_html = html.lower()
            
            # Get priority-based color (if termbase_matches contains priority info)
            if isinstance(match_info, dict) and 'priority' in match_info:
                priority = match_info.get('priority', 50)
                translation = match_info.get('translation', '')
            else:
                # Backward compatibility: if just string, use default priority
                priority = 50
                translation = match_info if isinstance(match_info, str) else ''
            
            # Calculate color based on priority (1-99, higher = darker blue)
            # Priority 99 = darkest (#0066CC), Priority 1 = lightest (#99CCFF)
            darkness = int(255 - (priority * 1.5))  # Higher priority = lower RGB value = darker
            darkness = max(0, min(darkness, 200))  # Clamp between 0-200
            color = f"rgb(0, {darkness}, 255)"
            
            # Find all occurrences
            start = 0
            while True:
                idx = search_html.find(search_term, start)
                if idx == -1:
                    break
                
                # Extract original casing
                original_term = html[idx:idx + len(term)]
                
                # Replace with highlighted version using priority color
                highlighted = f'<span style="background-color: {color}; color: white; padding: 1px 3px; border-radius: 2px; font-weight: bold; cursor: pointer;" class="termbase-match" data-term="{term}" title="{translation} (Priority: {priority})">{original_term}</span>'
                
                html = html[:idx] + highlighted + html[idx + len(term):]
                search_html = html.lower()
                
                start = idx + len(highlighted)
        
        return html
    
    def mouseMoveEvent(self, ev):
        """Show tooltip when hovering over highlighted terms"""
        # Find if cursor is over a highlighted term using text format
        # This is a simplified approach - shows generic tooltip
        super().mouseMoveEvent(ev)
    
    def mouseDoubleClickEvent(self, ev):
        """Handle double-click on highlighted terms"""
        # Find which term was clicked by getting the word at the position
        # Since QLabel doesn't provide cursor position easily, we look at the click
        # and try to extract the clicked word from the text
        
        # Get position in text (approximate)
        pos = ev.pos()
        
        # Try to find which term is near the click point
        # This is a simplified approach - in production would need better handling
        for term_id, match_info in self.termbase_matches.items():
            term = match_info.get('source', '') if isinstance(match_info, dict) else str(term_id)
            translation = match_info.get('translation', match_info) if isinstance(match_info, dict) else match_info
            # Check if this term contains the click area (very approximate)
            if self.term_double_clicked is not None and callable(self.term_double_clicked):
                # For now, if we get a double-click and there are matches,
                # this is a signal to insert the first/only match
                # Better implementation would need proper hit detection
                pass
        
        super().mouseDoubleClickEvent(ev)


class WordWrapDelegate(QStyledItemDelegate):
    """Custom delegate to enable word wrap when editing cells, with search term highlighting"""

    def __init__(self, assistance_panel=None, table_widget=None, allow_source_edit=False):
        super().__init__()
        self.assistance_panel = assistance_panel
        self.table_widget = table_widget
        self.allow_source_edit = allow_source_edit  # Controls whether source can be edited
        self.search_terms = {}  # Dict of {(row, col): search_term} for highlighting
        self.highlight_color = QColor(255, 255, 0, 180)  # Yellow with some transparency

    def set_highlight(self, row: int, col: int, search_term: str):
        """Set a search term to highlight for a specific cell"""
        self.search_terms[(row, col)] = search_term

    def clear_highlight(self, row: int, col: int):
        """Clear highlight for a specific cell"""
        if (row, col) in self.search_terms:
            del self.search_terms[(row, col)]

    def clear_all_highlights(self):
        """Clear all highlights"""
        self.search_terms.clear()

    def paint(self, painter, option, index):
        """Custom paint method that highlights search terms while keeping cells editable.

        Uses QTextDocument with HTML for reliable cross-line highlighting.
        Instead of storing per-cell terms, we store a global search term and check
        if the cell's text contains it.
        """
        row = index.row()
        col = index.column()

        # Check if this cell has a search term to highlight
        # Use global search term if set, otherwise check per-cell dict
        search_term = None
        if col == 3 and hasattr(self, 'global_search_term') and self.global_search_term:
            # For target column, use global search term
            search_term = self.global_search_term
        elif col == 2 and hasattr(self, 'global_source_search_term') and self.global_source_search_term:
            # For source column, use global source search term
            search_term = self.global_source_search_term
        elif (row, col) in self.search_terms:
            search_term = self.search_terms[(row, col)]

        # Debug: Log once per column 3 cell when global_search_term is set
        if col == 3 and hasattr(self, 'global_search_term') and self.global_search_term and row == 0:
            print(f"[PAINT DEBUG] Row 0, Col 3: global_search_term='{self.global_search_term}', search_term='{search_term}'")

        if search_term:
            text = index.data(Qt.ItemDataRole.DisplayRole) or ""

            # Only highlight if the search term is actually in the text
            if text and search_term.lower() in text.lower():
                from PyQt6.QtGui import QColor, QTextDocument, QAbstractTextDocumentLayout, QPalette
                from PyQt6.QtCore import QRectF, QSizeF
                import html as html_module

                # Get style and text rect
                option_copy = QStyleOptionViewItem(option)
                self.initStyleOption(option_copy, index)
                style = option.widget.style() if option.widget else QApplication.style()

                # Draw background first (handles selection state etc.)
                style.drawPrimitive(QStyle.PrimitiveElement.PE_PanelItemViewItem, option_copy, painter, option.widget)

                # Get the text rectangle
                text_rect = style.subElementRect(QStyle.SubElement.SE_ItemViewItemText, option_copy, option.widget)

                # Add small padding like default delegate
                text_rect = text_rect.adjusted(3, 0, -3, 0)

                # Create HTML with highlighted search terms
                # Escape HTML special characters first
                escaped_text = html_module.escape(text)
                escaped_term = html_module.escape(search_term)

                # Case-insensitive replacement with highlighting
                import re
                pattern = re.compile(re.escape(escaped_term), re.IGNORECASE)

                def replace_with_highlight(match):
                    return f'<span style="background-color: #FFFF00;">{match.group(0)}</span>'

                html_text = pattern.sub(replace_with_highlight, escaped_text)

                # Create QTextDocument with the highlighted HTML
                doc = QTextDocument()
                doc.setDefaultFont(option.font)
                doc.setTextWidth(text_rect.width())
                doc.setHtml(html_text)

                painter.save()
                painter.translate(text_rect.topLeft())

                # Clip to text rect
                painter.setClipRect(QRectF(0, 0, text_rect.width(), text_rect.height()))

                # Draw the document
                doc.drawContents(painter)

                painter.restore()
                return

        # Default painting for non-highlighted cells
        super().paint(painter, option, index)
    
    def createEditor(self, parent, option, index):
        """Create a QTextEdit for multi-line editing with word wrap"""
        # Target column (column 3) - always editable
        if index.column() == 3:
            editor = GridTextEditor(parent)
            editor.assistance_panel = self.assistance_panel
            editor.table_widget = self.table_widget
            editor.current_row = index.row()
            editor.setWordWrapMode(QTextOption.WrapMode.WordWrap)
            editor.setAcceptRichText(False)  # Plain text only
            editor.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
            editor.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            
            # Ensure the row is tall enough for editing
            table = parent.parent()
            if hasattr(table, 'resizeRowToContents'):
                # Schedule row resize after editor is shown
                from PyQt6.QtCore import QTimer
                QTimer.singleShot(0, lambda: table.resizeRowToContents(index.row()))
            
            return editor
        
        # Source column (column 2) - editable for selection/dual selection, optionally editable for content
        elif index.column() == 2:
            editor = ReadOnlyGridTextEditor("", parent)
            editor.table_widget = self.table_widget
            editor.current_row = index.row()
            editor.allow_source_edit = self.allow_source_edit
            
            # If allow_source_edit is True, make it actually editable (with warning color)
            if self.allow_source_edit:
                editor.setReadOnly(False)
                editor.setStyleSheet("""
                    QTextEdit {
                        border: none;
                        background-color: #FFF9E6;
                        padding: 0px;
                        color: black;
                    }
                    QTextEdit:focus {
                        border: 1px solid #FF9800;
                        background-color: #FFFACD;
                        color: black;
                    }
                    QTextEdit::selection {
                        background-color: #D0E7FF;
                        color: black;
                    }
                """)
            
            # Ensure the row is tall enough for editing
            table = parent.parent()
            if hasattr(table, 'resizeRowToContents'):
                from PyQt6.QtCore import QTimer
                QTimer.singleShot(0, lambda: table.resizeRowToContents(index.row()))
            
            return editor
        else:
            return super().createEditor(parent, option, index)
    
    def setEditorData(self, editor, index):
        """Load data into the editor"""
        if isinstance(editor, (QTextEdit, ReadOnlyGridTextEditor)):
            text = index.model().data(index, Qt.ItemDataRole.EditRole)
            editor.setPlainText(text or "")
            # Ensure cursor is at start
            cursor = editor.textCursor()
            cursor.movePosition(cursor.MoveOperation.Start)
            editor.setTextCursor(cursor)
        else:
            super().setEditorData(editor, index)
    
    def setModelData(self, editor, model, index):
        """Save data from editor back to model"""
        if isinstance(editor, ReadOnlyGridTextEditor):
            # Only save if editing is allowed
            if editor.allow_source_edit:
                text = editor.toPlainText()
                model.setData(index, text, Qt.ItemDataRole.EditRole)
            # Otherwise, just close without saving (was for selection only)
        elif isinstance(editor, QTextEdit):
            text = editor.toPlainText()
            model.setData(index, text, Qt.ItemDataRole.EditRole)
        else:
            super().setModelData(editor, model, index)
    
    def updateEditorGeometry(self, editor, option, index):
        """Set the editor geometry to match the cell size"""
        if isinstance(editor, (QTextEdit, ReadOnlyGridTextEditor)):
            # Make the editor fill the cell properly with some padding
            rect = option.rect
            editor.setGeometry(rect)
        else:
            super().updateEditorGeometry(editor, option, index)


# ============================================================================
# THEME EDITOR DIALOG
# ============================================================================

class ThemeEditorDialog(QDialog):
    """Dialog for editing and managing themes"""
    
    def __init__(self, parent, theme_manager):
        super().__init__(parent)
        self.theme_manager = theme_manager
        self.setWindowTitle("Theme Editor")
        self.setModal(True)
        self.resize(700, 600)
        
        self.setup_ui()
        self.load_themes()
    
    def setup_ui(self):
        """Create the UI"""
        layout = QVBoxLayout(self)
        
        # Theme selection
        theme_group = QGroupBox("Select Theme")
        theme_layout = QHBoxLayout()
        
        self.theme_combo = QComboBox()
        self.theme_combo.currentTextChanged.connect(self.on_theme_selected)
        theme_layout.addWidget(QLabel("Theme:"))
        theme_layout.addWidget(self.theme_combo, 1)
        
        self.apply_btn = QPushButton("✓ Apply")
        self.apply_btn.clicked.connect(self.apply_theme)
        theme_layout.addWidget(self.apply_btn)
        
        theme_group.setLayout(theme_layout)
        layout.addWidget(theme_group)
        
        # Color customization
        colors_group = QGroupBox("Customize Colors")
        colors_layout = QGridLayout()
        
        # Create color pickers for main colors
        self.color_buttons = {}
        color_configs = [
            ("window_bg", "Window Background", 0, 0),
            ("base", "Input Fields", 0, 1),
            ("button", "Buttons", 0, 2),
            ("text", "Text", 1, 0),
            ("highlight", "Highlight", 1, 1),
            ("border", "Borders", 1, 2),
            ("grid_header", "Table Headers", 2, 0),
            ("alternate_bg", "Alternate Rows", 2, 1),
            ("tm_exact", "100% TM Match", 3, 0),
            ("tm_high", "95-99% TM Match", 3, 1),
        ]
        
        for attr, label, row, col in color_configs:
            lbl = QLabel(label + ":")
            btn = QPushButton()
            btn.setMinimumHeight(30)
            btn.setProperty("color_attr", attr)
            btn.clicked.connect(lambda checked, a=attr: self.pick_color(a))
            self.color_buttons[attr] = btn
            
            colors_layout.addWidget(lbl, row * 2, col)
            colors_layout.addWidget(btn, row * 2 + 1, col)
        
        colors_group.setLayout(colors_layout)
        layout.addWidget(colors_group)
        
        # Custom theme actions
        custom_group = QGroupBox("Custom Themes")
        custom_layout = QHBoxLayout()
        
        save_btn = QPushButton("💾 Save as Custom Theme")
        save_btn.clicked.connect(self.save_custom_theme)
        custom_layout.addWidget(save_btn)
        
        delete_btn = QPushButton("🗑️ Delete Custom Theme")
        delete_btn.clicked.connect(self.delete_custom_theme)
        custom_layout.addWidget(delete_btn)
        
        custom_group.setLayout(custom_layout)
        layout.addWidget(custom_group)
        
        # Dialog buttons
        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | 
            QDialogButtonBox.StandardButton.Cancel
        )
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
    
    def load_themes(self):
        """Load available themes into combo box"""
        self.theme_combo.clear()
        themes = self.theme_manager.get_all_themes()
        for theme_name in themes.keys():
            self.theme_combo.addItem(theme_name)
        
        # Select current theme
        current_idx = self.theme_combo.findText(self.theme_manager.current_theme.name)
        if current_idx >= 0:
            self.theme_combo.setCurrentIndex(current_idx)
    
    def on_theme_selected(self, theme_name):
        """When theme is selected from combo box"""
        if not theme_name:
            return
        
        theme = self.theme_manager.get_theme(theme_name)
        if theme:
            self.update_color_buttons(theme)
    
    def update_color_buttons(self, theme):
        """Update color button backgrounds"""
        for attr, btn in self.color_buttons.items():
            color = getattr(theme, attr, "#FFFFFF")
            btn.setStyleSheet(f"background-color: {color}; border: 1px solid #999;")
            btn.setText(color)
    
    def pick_color(self, attr):
        """Open color picker for an attribute"""
        from PyQt6.QtWidgets import QColorDialog
        
        theme_name = self.theme_combo.currentText()
        theme = self.theme_manager.get_theme(theme_name)
        
        if not theme:
            return
        
        current_color = QColor(getattr(theme, attr))
        color = QColorDialog.getColor(current_color, self, f"Pick {attr}")
        
        if color.isValid():
            # Update button
            hex_color = color.name()
            self.color_buttons[attr].setStyleSheet(
                f"background-color: {hex_color}; border: 1px solid #999;"
            )
            self.color_buttons[attr].setText(hex_color)
            
            # Update theme
            setattr(theme, attr, hex_color)
    
    def apply_theme(self):
        """Apply the selected theme"""
        theme_name = self.theme_combo.currentText()
        if self.theme_manager.set_theme(theme_name):
            self.theme_manager.apply_theme(QApplication.instance())

            # Call the main app's refresh method to update all UI elements
            if hasattr(self.parent(), 'refresh_theme_colors'):
                self.parent().refresh_theme_colors()

            QMessageBox.information(self, "Theme Applied",
                                   f"Theme '{theme_name}' has been applied.")
    
    def save_custom_theme(self):
        """Save current settings as a custom theme"""
        from PyQt6.QtWidgets import QInputDialog
        
        name, ok = QInputDialog.getText(self, "Save Theme", "Enter theme name:")
        if ok and name:
            # Create new theme from current settings
            theme_name = self.theme_combo.currentText()
            base_theme = self.theme_manager.get_theme(theme_name)
            
            if base_theme:
                # Create copy with new name
                from modules.theme_manager import Theme
                new_theme = Theme(
                    name=name,
                    **{k: v for k, v in base_theme.to_dict().items() if k != 'name'}
                )
                
                # Update with any modified colors
                for attr, btn in self.color_buttons.items():
                    color = btn.text()
                    if color.startswith('#'):
                        setattr(new_theme, attr, color)
                
                self.theme_manager.save_custom_theme(new_theme)
                self.load_themes()
                
                # Select the new theme
                idx = self.theme_combo.findText(name)
                if idx >= 0:
                    self.theme_combo.setCurrentIndex(idx)
                
                QMessageBox.information(self, "Theme Saved", 
                                       f"Custom theme '{name}' has been saved.")
    
    def delete_custom_theme(self):
        """Delete the selected custom theme"""
        theme_name = self.theme_combo.currentText()
        
        # Can't delete predefined themes
        from modules.theme_manager import ThemeManager
        if theme_name in ThemeManager.PREDEFINED_THEMES:
            QMessageBox.warning(self, "Cannot Delete", 
                               "Cannot delete predefined themes.")
            return
        
        reply = QMessageBox.question(self, "Delete Theme",
                                    f"Delete custom theme '{theme_name}'?",
                                    QMessageBox.StandardButton.Yes | 
                                    QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.Yes:
            if self.theme_manager.delete_custom_theme(theme_name):
                self.load_themes()
                QMessageBox.information(self, "Theme Deleted", 
                                       f"Theme '{theme_name}' has been deleted.")


# ============================================================================
# DETACHED LOG WINDOW
# ============================================================================

class DetachedLogWindow(QWidget):
    """Detached log window that can be moved to another screen"""

    def __init__(self, parent):
        super().__init__()
        self.parent = parent
        self.setWindowTitle("Supervertaler - Session Log")
        self.setWindowIcon(self.parent.windowIcon())
        self.resize(800, 600)

        # Create layout
        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 10, 10, 10)

        # Top toolbar
        toolbar = QWidget()
        toolbar_layout = QHBoxLayout(toolbar)
        toolbar_layout.setContentsMargins(0, 0, 0, 0)

        # Info label
        info_label = QLabel("📋 This is a detached log window. It will update in real-time.")
        info_label.setStyleSheet("color: #666; font-style: italic;")
        toolbar_layout.addWidget(info_label)

        toolbar_layout.addStretch()

        # Re-attach button
        reattach_btn = QPushButton("↩️ Close")
        reattach_btn.setToolTip("Close this detached window")
        reattach_btn.clicked.connect(self.close)
        toolbar_layout.addWidget(reattach_btn)

        layout.addWidget(toolbar)

        # Log display
        self.log_display = QPlainTextEdit()
        self.log_display.setReadOnly(True)
        self.log_display.setStyleSheet("""
            QPlainTextEdit {
                font-family: 'Consolas', 'Courier New', monospace;
                font-size: 10px;
                border: 1px solid #ccc;
            }
        """)
        layout.addWidget(self.log_display)

        # Copy existing log content
        if hasattr(parent, 'session_log') and parent.session_log:
            self.log_display.setPlainText(parent.session_log.toPlainText())
            # Scroll to bottom
            scrollbar = self.log_display.verticalScrollBar()
            if scrollbar:
                scrollbar.setValue(scrollbar.maximum())

    def closeEvent(self, event):
        """Handle window close"""
        # Remove from parent's list
        try:
            if self in self.parent.detached_log_windows:
                self.parent.detached_log_windows.remove(self)
        except:
            pass
        event.accept()


# ============================================================================
# TERM METADATA DIALOG
# ============================================================================

class TermMetadataDialog(QDialog):
    """Dialog for adding/editing term metadata before saving to termbase"""
    
    def __init__(self, source_term: str, target_term: str, active_termbases: list, parent=None, user_data_path=None):
        super().__init__(parent)
        self.source_term = source_term
        self.target_term = target_term
        self.active_termbases = active_termbases
        self.termbase_checkboxes = {}  # Store checkbox references
        self.user_data_path = user_data_path
        self.saved_selections = self._load_termbase_selections()
        self.setup_ui()
    
    def _load_termbase_selections(self):
        """Load saved termbase selections from preferences"""
        if not self.user_data_path:
            return None
        
        prefs_file = self.user_data_path / "ui_preferences.json"
        if not prefs_file.exists():
            return None
        
        try:
            with open(prefs_file, 'r') as f:
                prefs = json.load(f)
                return prefs.get('add_term_termbase_selections', None)
        except:
            return None
    
    def _save_termbase_selections(self):
        """Save current termbase selections to preferences"""
        if not self.user_data_path:
            return
        
        prefs_file = self.user_data_path / "ui_preferences.json"
        
        # Load existing preferences
        prefs = {}
        if prefs_file.exists():
            try:
                with open(prefs_file, 'r') as f:
                    prefs = json.load(f)
            except:
                pass
        
        # Save the selected termbase IDs
        selected_ids = [tb_id for tb_id, cb in self.termbase_checkboxes.items() if cb.isChecked()]
        prefs['add_term_termbase_selections'] = selected_ids
        
        try:
            with open(prefs_file, 'w') as f:
                json.dump(prefs, f, indent=2)
        except:
            pass
        
    def setup_ui(self):
        self.setWindowTitle("Add Term to Glossary")
        self.setMinimumWidth(550)

        # Auto-resize to fit screen (max 85% of screen height)
        screen = QApplication.primaryScreen().availableGeometry()
        max_height = int(screen.height() * 0.85)
        self.setMaximumHeight(max_height)

        # Start with very compact size for laptops
        self.resize(600, min(550, max_height))

        # Create main layout
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)

        # Create scroll area for all content
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)

        content_widget = QWidget()
        layout = QVBoxLayout(content_widget)
        layout.setSpacing(4)
        layout.setContentsMargins(6, 6, 6, 6)
        
        # Header
        header = QLabel("Add term pair to glossary")
        header.setStyleSheet("font-size: 12px; font-weight: bold; margin-bottom: 5px; padding: 4px;")
        layout.addWidget(header)
        
        # Term pair display (read-only)
        term_group = QGroupBox("Term Pair")
        term_layout = QFormLayout()
        
        source_label = QLabel(self.source_term)
        source_label.setStyleSheet("padding: 5px; border-radius: 3px;")
        source_label.setWordWrap(True)
        term_layout.addRow("Source:", source_label)
        
        target_label = QLabel(self.target_term)
        target_label.setStyleSheet("padding: 5px; border-radius: 3px;")
        target_label.setWordWrap(True)
        term_layout.addRow("Target:", target_label)
        
        term_group.setLayout(term_layout)
        layout.addWidget(term_group)
        
        # Termbase selection
        tb_group = QGroupBox("Save to Glossary(s)")
        tb_layout = QVBoxLayout()
        
        if not self.active_termbases:
            no_tb_label = QLabel("⚠️ No active glossaries found. Please activate at least one glossary first.")
            no_tb_label.setStyleSheet("color: #d32f2f; padding: 10px;")
            no_tb_label.setWordWrap(True)
            tb_layout.addWidget(no_tb_label)
        else:
            # Header with select all/none buttons
            header_layout = QHBoxLayout()
            select_all_btn = QPushButton("Select All")
            select_all_btn.setMaximumWidth(100)
            select_none_btn = QPushButton("Select None")
            select_none_btn.setMaximumWidth(100)
            
            def select_all():
                for cb in self.termbase_checkboxes.values():
                    cb.setChecked(True)
            
            def select_none():
                for cb in self.termbase_checkboxes.values():
                    cb.setChecked(False)
            
            select_all_btn.clicked.connect(select_all)
            select_none_btn.clicked.connect(select_none)
            
            header_layout.addWidget(select_all_btn)
            header_layout.addWidget(select_none_btn)
            header_layout.addStretch()
            tb_layout.addLayout(header_layout)
            
            # Checkboxes for each active termbase
            for tb in self.active_termbases:
                is_project_tb = tb.get('is_project_termbase', False)
                
                # Use pink checkbox for project termbase, green for others
                if is_project_tb:
                    cb = PinkCheckmarkCheckBox(f"📌 {tb['name']} (Project)")
                else:
                    cb = CheckmarkCheckBox(tb['name'])
                
                # Use saved selection if available, otherwise default to all selected
                if self.saved_selections is not None:
                    cb.setChecked(tb['id'] in self.saved_selections)
                else:
                    cb.setChecked(True)  # Default: all selected
                cb.setToolTip(f"Languages: {tb.get('source_lang', '?')} → {tb.get('target_lang', '?')}")
                
                self.termbase_checkboxes[tb['id']] = cb
                tb_layout.addWidget(cb)
        
        tb_group.setLayout(tb_layout)
        layout.addWidget(tb_group)
        
        # Metadata fields
        meta_group = QGroupBox("Metadata (Optional)")
        meta_layout = QFormLayout()
        
        # Domain
        self.domain_edit = QLineEdit()
        self.domain_edit.setPlaceholderText("e.g., Patents, Legal, Medical, IT...")
        meta_layout.addRow("Domain:", self.domain_edit)
        
        # Notes
        self.notes_edit = QTextEdit()
        self.notes_edit.setMaximumHeight(45)
        self.notes_edit.setPlaceholderText("Usage notes, context, definition, URLs...")
        self.notes_edit.setStyleSheet("padding: 3px; font-size: 10px;")
        meta_layout.addRow("Notes:", self.notes_edit)
        
        # Project
        self.project_edit = QLineEdit()
        self.project_edit.setPlaceholderText("Optional project name...")
        meta_layout.addRow("Project:", self.project_edit)
        
        # Client
        self.client_edit = QLineEdit()
        self.client_edit.setPlaceholderText("Optional client name...")
        meta_layout.addRow("Client:", self.client_edit)
        
        # Forbidden term checkbox
        self.forbidden_check = CheckmarkCheckBox("Mark as forbidden term")
        self.forbidden_check.setToolTip("Forbidden terms trigger warnings when used")
        meta_layout.addRow("", self.forbidden_check)
        
        meta_group.setLayout(meta_layout)
        layout.addWidget(meta_group)
        
        # Source Synonyms section (collapsible)
        source_syn_group = QGroupBox()
        source_syn_main_layout = QVBoxLayout()

        # Header with collapse button
        source_syn_header = QHBoxLayout()
        self.source_syn_toggle = QToolButton()
        self.source_syn_toggle.setText("▼")
        self.source_syn_toggle.setStyleSheet("QToolButton { border: none; font-weight: bold; }")
        self.source_syn_toggle.setFixedSize(20, 20)
        self.source_syn_toggle.setCheckable(True)
        self.source_syn_toggle.setChecked(False)
        source_syn_header.addWidget(self.source_syn_toggle)

        source_syn_label = QLabel("Source Synonyms (Optional)")
        source_syn_label.setStyleSheet("font-weight: bold;")
        source_syn_header.addWidget(source_syn_label)
        source_syn_header.addStretch()
        source_syn_main_layout.addLayout(source_syn_header)

        # Collapsible content
        self.source_syn_content = QWidget()
        source_syn_layout = QVBoxLayout(self.source_syn_content)
        source_syn_layout.setContentsMargins(0, 0, 0, 0)
        self.source_syn_content.setVisible(False)

        # Instructions
        source_syn_info = QLabel("Add alternative source terms. First item = preferred term:")
        source_syn_info.setStyleSheet("color: #666; font-size: 10px;")
        source_syn_layout.addWidget(source_syn_info)
        
        # Input field + Add button + Forbidden checkbox
        source_add_layout = QHBoxLayout()
        self.source_synonym_edit = QLineEdit()
        self.source_synonym_edit.setPlaceholderText("Enter source synonym and press Add or Enter...")
        source_add_layout.addWidget(self.source_synonym_edit)
        
        self.source_synonym_forbidden_check = CheckmarkCheckBox("Forbidden")
        self.source_synonym_forbidden_check.setToolTip("Mark this source synonym as forbidden")
        source_add_layout.addWidget(self.source_synonym_forbidden_check)
        
        source_add_syn_btn = QPushButton("Add")
        source_add_syn_btn.setMaximumWidth(60)
        source_add_syn_btn.clicked.connect(self.add_source_synonym)
        source_add_layout.addWidget(source_add_syn_btn)
        source_syn_layout.addLayout(source_add_layout)
        
        # Connect Enter key to add synonym
        self.source_synonym_edit.returnPressed.connect(self.add_source_synonym)
        
        # List of source synonyms with control buttons
        source_list_layout = QHBoxLayout()
        
        self.source_synonym_list = QListWidget()
        self.source_synonym_list.setMaximumHeight(100)
        self.source_synonym_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.source_synonym_list.customContextMenuRequested.connect(self.show_source_synonym_context_menu)
        source_list_layout.addWidget(self.source_synonym_list)
        
        # Up/Down buttons for source synonyms
        source_button_col = QVBoxLayout()
        source_move_up_btn = QPushButton("▲")
        source_move_up_btn.setToolTip("Move synonym up (higher priority)")
        source_move_up_btn.setMaximumWidth(30)
        source_move_up_btn.clicked.connect(self.move_source_synonym_up)
        source_button_col.addWidget(source_move_up_btn)
        
        source_move_down_btn = QPushButton("▼")
        source_move_down_btn.setToolTip("Move synonym down (lower priority)")
        source_move_down_btn.setMaximumWidth(30)
        source_move_down_btn.clicked.connect(self.move_source_synonym_down)
        source_button_col.addWidget(source_move_down_btn)
        
        source_button_col.addStretch()
        
        source_delete_btn = QPushButton("✗")
        source_delete_btn.setToolTip("Delete synonym")
        source_delete_btn.setMaximumWidth(30)
        source_delete_btn.clicked.connect(self.delete_selected_source_synonym)
        source_button_col.addWidget(source_delete_btn)
        
        source_list_layout.addLayout(source_button_col)
        source_syn_layout.addLayout(source_list_layout)

        # Add collapsible content to main layout
        source_syn_main_layout.addWidget(self.source_syn_content)
        source_syn_group.setLayout(source_syn_main_layout)

        # Connect toggle button
        self.source_syn_toggle.clicked.connect(lambda: self.toggle_section(self.source_syn_toggle, self.source_syn_content))

        layout.addWidget(source_syn_group)
        
        # Target Synonyms section (collapsible)
        target_syn_group = QGroupBox()
        target_syn_main_layout = QVBoxLayout()

        # Header with collapse button
        target_syn_header = QHBoxLayout()
        self.target_syn_toggle = QToolButton()
        self.target_syn_toggle.setText("▼")
        self.target_syn_toggle.setStyleSheet("QToolButton { border: none; font-weight: bold; }")
        self.target_syn_toggle.setFixedSize(20, 20)
        self.target_syn_toggle.setCheckable(True)
        self.target_syn_toggle.setChecked(False)
        target_syn_header.addWidget(self.target_syn_toggle)

        target_syn_label = QLabel("Target Synonyms (Optional)")
        target_syn_label.setStyleSheet("font-weight: bold;")
        target_syn_header.addWidget(target_syn_label)
        target_syn_header.addStretch()
        target_syn_main_layout.addLayout(target_syn_header)

        # Collapsible content
        self.target_syn_content = QWidget()
        target_syn_layout = QVBoxLayout(self.target_syn_content)
        target_syn_layout.setContentsMargins(0, 0, 0, 0)
        self.target_syn_content.setVisible(False)

        # Instructions
        target_syn_info = QLabel("Add alternative translations (synonyms). First item = preferred term:")
        target_syn_info.setStyleSheet("color: #666; font-size: 10px;")
        target_syn_layout.addWidget(target_syn_info)
        
        # Input field + Add button + Forbidden checkbox
        target_add_layout = QHBoxLayout()
        self.target_synonym_edit = QLineEdit()
        self.target_synonym_edit.setPlaceholderText("Enter synonym and press Add or Enter...")
        target_add_layout.addWidget(self.target_synonym_edit)
        
        self.target_synonym_forbidden_check = CheckmarkCheckBox("Forbidden")
        self.target_synonym_forbidden_check.setToolTip("Mark this synonym as forbidden (warning when used)")
        target_add_layout.addWidget(self.target_synonym_forbidden_check)
        
        target_add_syn_btn = QPushButton("Add")
        target_add_syn_btn.setMaximumWidth(60)
        target_add_syn_btn.clicked.connect(self.add_target_synonym)
        target_add_layout.addWidget(target_add_syn_btn)
        target_syn_layout.addLayout(target_add_layout)
        
        # Connect Enter key to add synonym
        self.target_synonym_edit.returnPressed.connect(self.add_target_synonym)
        
        # List of target synonyms with control buttons
        target_list_layout = QHBoxLayout()
        
        self.target_synonym_list = QListWidget()
        self.target_synonym_list.setMaximumHeight(100)
        self.target_synonym_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.target_synonym_list.customContextMenuRequested.connect(self.show_target_synonym_context_menu)
        target_list_layout.addWidget(self.target_synonym_list)
        
        # Up/Down buttons for target synonyms
        target_button_col = QVBoxLayout()
        target_move_up_btn = QPushButton("▲")
        target_move_up_btn.setToolTip("Move synonym up (higher priority)")
        target_move_up_btn.setMaximumWidth(30)
        target_move_up_btn.clicked.connect(self.move_target_synonym_up)
        target_button_col.addWidget(target_move_up_btn)
        
        target_move_down_btn = QPushButton("▼")
        target_move_down_btn.setToolTip("Move synonym down (lower priority)")
        target_move_down_btn.setMaximumWidth(30)
        target_move_down_btn.clicked.connect(self.move_target_synonym_down)
        target_button_col.addWidget(target_move_down_btn)
        
        target_button_col.addStretch()
        
        target_delete_btn = QPushButton("✗")
        target_delete_btn.setToolTip("Delete synonym")
        target_delete_btn.setMaximumWidth(30)
        target_delete_btn.clicked.connect(self.delete_selected_target_synonym)
        target_button_col.addWidget(target_delete_btn)
        
        target_list_layout.addLayout(target_button_col)
        target_syn_layout.addLayout(target_list_layout)

        # Add collapsible content to main layout
        target_syn_main_layout.addWidget(self.target_syn_content)
        target_syn_group.setLayout(target_syn_main_layout)

        # Connect toggle button
        self.target_syn_toggle.clicked.connect(lambda: self.toggle_section(self.target_syn_toggle, self.target_syn_content))

        layout.addWidget(target_syn_group)

        # Buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(cancel_btn)
        
        save_btn = QPushButton("Add to Glossary")
        save_btn.setStyleSheet("background-color: #4CAF50; color: white; font-weight: bold; padding: 5px 15px; border: none; outline: none;")
        save_btn.clicked.connect(self._accept_and_save)
        save_btn.setDefault(True)
        button_layout.addWidget(save_btn)

        layout.addLayout(button_layout)

        # Set the scroll area content
        scroll.setWidget(content_widget)
        main_layout.addWidget(scroll)

    def toggle_section(self, toggle_btn, content_widget):
        """Toggle visibility of a collapsible section"""
        is_visible = content_widget.isVisible()
        content_widget.setVisible(not is_visible)
        toggle_btn.setText("▼" if is_visible else "▲")

    # ========================================================================
    # SOURCE SYNONYM METHODS
    # ========================================================================
    
    def add_source_synonym(self):
        """Add a source synonym to the list"""
        synonym = self.source_synonym_edit.text().strip()
        if synonym:
            # Check for duplicates
            for i in range(self.source_synonym_list.count()):
                item = self.source_synonym_list.item(i)
                item_text = item.data(Qt.ItemDataRole.UserRole).get('text', '')
                if item_text == synonym:
                    QMessageBox.warning(self, "Duplicate", f"Source synonym '{synonym}' already added.")
                    return
            
            # Don't allow the main source term as a synonym
            if synonym.lower() == self.source_term.lower():
                QMessageBox.warning(self, "Invalid Synonym", "Cannot add the main source term as a synonym.")
                return
            
            # Create list item with stored data
            is_forbidden = self.source_synonym_forbidden_check.isChecked()
            display_text = f"{'🚫 ' if is_forbidden else ''}{synonym}"
            
            item = QListWidgetItem(display_text)
            item.setData(Qt.ItemDataRole.UserRole, {
                'text': synonym,
                'forbidden': is_forbidden
            })
            
            if is_forbidden:
                item.setForeground(QColor('#d32f2f'))
            
            self.source_synonym_list.addItem(item)
            self.source_synonym_edit.clear()
            self.source_synonym_forbidden_check.setChecked(False)
            self.source_synonym_edit.setFocus()
    
    def move_source_synonym_up(self):
        """Move selected source synonym up in the list"""
        current_row = self.source_synonym_list.currentRow()
        if current_row > 0:
            item = self.source_synonym_list.takeItem(current_row)
            self.source_synonym_list.insertItem(current_row - 1, item)
            self.source_synonym_list.setCurrentRow(current_row - 1)
    
    def move_source_synonym_down(self):
        """Move selected source synonym down in the list"""
        current_row = self.source_synonym_list.currentRow()
        if current_row < self.source_synonym_list.count() - 1 and current_row >= 0:
            item = self.source_synonym_list.takeItem(current_row)
            self.source_synonym_list.insertItem(current_row + 1, item)
            self.source_synonym_list.setCurrentRow(current_row + 1)
    
    def delete_selected_source_synonym(self):
        """Delete selected source synonym"""
        current_row = self.source_synonym_list.currentRow()
        if current_row >= 0:
            self.source_synonym_list.takeItem(current_row)
    
    def show_source_synonym_context_menu(self, position):
        """Show context menu for source synonym list"""
        if self.source_synonym_list.count() == 0:
            return
        
        current_item = self.source_synonym_list.currentItem()
        if not current_item:
            return
        
        menu = QMenu()
        
        # Toggle forbidden status
        data = current_item.data(Qt.ItemDataRole.UserRole)
        is_forbidden = data.get('forbidden', False)
        
        if is_forbidden:
            toggle_action = menu.addAction("Mark as Allowed")
        else:
            toggle_action = menu.addAction("Mark as Forbidden")
        
        menu.addSeparator()
        delete_action = menu.addAction("Delete")
        
        action = menu.exec(self.source_synonym_list.mapToGlobal(position))
        
        if action == toggle_action:
            # Toggle forbidden status
            data['forbidden'] = not is_forbidden
            text = data['text']
            display_text = f"{'🚫 ' if data['forbidden'] else ''}{text}"
            current_item.setText(display_text)
            current_item.setData(Qt.ItemDataRole.UserRole, data)
            
            if data['forbidden']:
                current_item.setForeground(QColor('#d32f2f'))
            else:
                current_item.setForeground(QColor('#000000'))
                
        elif action == delete_action:
            self.source_synonym_list.takeItem(self.source_synonym_list.row(current_item))
    
    def get_source_synonyms(self):
        """Return list of source synonym dictionaries with text, forbidden flag, and order"""
        synonyms = []
        for i in range(self.source_synonym_list.count()):
            item = self.source_synonym_list.item(i)
            data = item.data(Qt.ItemDataRole.UserRole)
            synonyms.append({
                'text': data['text'],
                'forbidden': data['forbidden'],
                'order': i
            })
        return synonyms
    
    # ========================================================================
    # TARGET SYNONYM METHODS
    # ========================================================================
    
    def add_target_synonym(self):
        """Add a target synonym to the list"""
        synonym = self.target_synonym_edit.text().strip()
        if synonym:
            # Check for duplicates
            for i in range(self.target_synonym_list.count()):
                item = self.target_synonym_list.item(i)
                item_text = item.data(Qt.ItemDataRole.UserRole).get('text', '')
                if item_text == synonym:
                    QMessageBox.warning(self, "Duplicate", f"Synonym '{synonym}' already added.")
                    return
            
            # Don't allow the main target term as a synonym
            if synonym.lower() == self.target_term.lower():
                QMessageBox.warning(self, "Invalid Synonym", "Cannot add the main target term as a synonym.")
                return
            
            # Create list item with stored data
            is_forbidden = self.target_synonym_forbidden_check.isChecked()
            display_text = f"{'🚫 ' if is_forbidden else ''}{synonym}"
            
            item = QListWidgetItem(display_text)
            item.setData(Qt.ItemDataRole.UserRole, {
                'text': synonym,
                'forbidden': is_forbidden
            })
            
            if is_forbidden:
                item.setForeground(QColor('#d32f2f'))
            
            self.target_synonym_list.addItem(item)
            self.target_synonym_edit.clear()
            self.target_synonym_forbidden_check.setChecked(False)
            self.target_synonym_edit.setFocus()
    
    def move_target_synonym_up(self):
        """Move selected target synonym up in the list"""
        current_row = self.target_synonym_list.currentRow()
        if current_row > 0:
            item = self.target_synonym_list.takeItem(current_row)
            self.target_synonym_list.insertItem(current_row - 1, item)
            self.target_synonym_list.setCurrentRow(current_row - 1)
    
    def move_target_synonym_down(self):
        """Move selected target synonym down in the list"""
        current_row = self.target_synonym_list.currentRow()
        if current_row < self.target_synonym_list.count() - 1 and current_row >= 0:
            item = self.target_synonym_list.takeItem(current_row)
            self.target_synonym_list.insertItem(current_row + 1, item)
            self.target_synonym_list.setCurrentRow(current_row + 1)
    
    def delete_selected_target_synonym(self):
        """Delete selected target synonym"""
        current_row = self.target_synonym_list.currentRow()
        if current_row >= 0:
            self.target_synonym_list.takeItem(current_row)
    
    def show_target_synonym_context_menu(self, position):
        """Show context menu for target synonym list"""
        if self.target_synonym_list.count() == 0:
            return
        
        current_item = self.target_synonym_list.currentItem()
        if not current_item:
            return
        
        menu = QMenu()
        
        # Toggle forbidden status
        data = current_item.data(Qt.ItemDataRole.UserRole)
        is_forbidden = data.get('forbidden', False)
        
        if is_forbidden:
            toggle_action = menu.addAction("Mark as Allowed")
        else:
            toggle_action = menu.addAction("Mark as Forbidden")
        
        menu.addSeparator()
        delete_action = menu.addAction("Delete")
        
        action = menu.exec(self.target_synonym_list.mapToGlobal(position))
        
        if action == toggle_action:
            # Toggle forbidden status
            data['forbidden'] = not is_forbidden
            text = data['text']
            display_text = f"{'🚫 ' if data['forbidden'] else ''}{text}"
            current_item.setText(display_text)
            current_item.setData(Qt.ItemDataRole.UserRole, data)
            
            if data['forbidden']:
                current_item.setForeground(QColor('#d32f2f'))
            else:
                current_item.setForeground(QColor('#000000'))
                
        elif action == delete_action:
            self.target_synonym_list.takeItem(self.target_synonym_list.row(current_item))
    
    def get_target_synonyms(self):
        """Return list of target synonym dictionaries with text, forbidden flag, and order"""
        synonyms = []
        for i in range(self.target_synonym_list.count()):
            item = self.target_synonym_list.item(i)
            data = item.data(Qt.ItemDataRole.UserRole)
            synonyms.append({
                'text': data['text'],
                'forbidden': data['forbidden'],
                'order': i
            })
        return synonyms
    
    def get_metadata(self):
        """Return dictionary of metadata fields"""
        return {
            'domain': self.domain_edit.text().strip(),
            'notes': self.notes_edit.toPlainText().strip(),
            'project': self.project_edit.text().strip(),
            'client': self.client_edit.text().strip(),
            'forbidden': self.forbidden_check.isChecked()
        }
    
    def get_selected_termbases(self):
        """Return list of selected termbase IDs"""
        return [tb_id for tb_id, cb in self.termbase_checkboxes.items() if cb.isChecked()]
    
    def _accept_and_save(self):
        """Save termbase selections and accept the dialog"""
        self._save_termbase_selections()
        self.accept()


class VoiceCommandEditDialog(QDialog):
    """Dialog for adding/editing voice commands"""
    
    CATEGORIES = ["navigation", "editing", "translation", "lookup", "file", "view", "dictation", "memoq", "trados", "custom"]
    ACTION_TYPES = [
        ("internal", "Internal Action (Supervertaler)"),
        ("keystroke", "Keystroke (e.g., ctrl+s)"),
        ("ahk_inline", "AutoHotkey Code"),
        ("ahk_script", "AutoHotkey Script File"),
    ]
    
    def __init__(self, parent=None, command: VoiceCommand = None):
        super().__init__(parent)
        self.command = command
        self.setup_ui()
        
        if command:
            self.populate_from_command(command)
    
    def setup_ui(self):
        self.setWindowTitle("Edit Voice Command" if self.command else "Add Voice Command")
        self.setMinimumWidth(500)
        
        layout = QVBoxLayout(self)
        
        # Phrase
        phrase_layout = QHBoxLayout()
        phrase_layout.addWidget(QLabel("Phrase:"))
        self.phrase_edit = QLineEdit()
        self.phrase_edit.setPlaceholderText("e.g., confirm segment")
        phrase_layout.addWidget(self.phrase_edit)
        layout.addLayout(phrase_layout)
        
        # Aliases
        aliases_layout = QHBoxLayout()
        aliases_layout.addWidget(QLabel("Aliases:"))
        self.aliases_edit = QLineEdit()
        self.aliases_edit.setPlaceholderText("e.g., confirm, done, okay (comma-separated)")
        aliases_layout.addWidget(self.aliases_edit)
        layout.addLayout(aliases_layout)
        
        # Action Type
        type_layout = QHBoxLayout()
        type_layout.addWidget(QLabel("Type:"))
        self.type_combo = QComboBox()
        for value, label in self.ACTION_TYPES:
            self.type_combo.addItem(label, value)
        self.type_combo.currentIndexChanged.connect(self._on_type_changed)
        type_layout.addWidget(self.type_combo)
        layout.addLayout(type_layout)
        
        # Action
        action_layout = QVBoxLayout()
        action_label = QLabel("Action:")
        action_layout.addWidget(action_label)
        self.action_edit = QTextEdit()
        self.action_edit.setMaximumHeight(100)
        self.action_edit.setPlaceholderText("For internal: action_name\nFor keystroke: ctrl+s\nFor AHK: Send, ^s")
        action_layout.addWidget(self.action_edit)
        layout.addLayout(action_layout)
        
        # Internal actions dropdown (for internal type)
        self.internal_actions_layout = QHBoxLayout()
        self.internal_actions_layout.addWidget(QLabel("Preset:"))
        self.internal_combo = QComboBox()
        self.internal_combo.addItems([
            "navigate_next", "navigate_previous", "navigate_first", "navigate_last",
            "confirm_segment", "copy_source_to_target", "clear_target",
            "translate_segment", "batch_translate",
            "open_superlookup", "concordance_search",
            "show_log", "show_editor",
            "start_dictation", "stop_listening"
        ])
        self.internal_combo.currentTextChanged.connect(lambda t: self.action_edit.setPlainText(t))
        self.internal_actions_layout.addWidget(self.internal_combo)
        self.internal_actions_layout.addStretch()
        layout.addLayout(self.internal_actions_layout)
        
        # Description
        desc_layout = QHBoxLayout()
        desc_layout.addWidget(QLabel("Description:"))
        self.desc_edit = QLineEdit()
        self.desc_edit.setPlaceholderText("e.g., Confirm current segment")
        desc_layout.addWidget(self.desc_edit)
        layout.addLayout(desc_layout)
        
        # Category
        cat_layout = QHBoxLayout()
        cat_layout.addWidget(QLabel("Category:"))
        self.cat_combo = QComboBox()
        self.cat_combo.addItems(self.CATEGORIES)
        self.cat_combo.setEditable(True)
        cat_layout.addWidget(self.cat_combo)
        layout.addLayout(cat_layout)
        
        # Buttons
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(self.reject)
        btn_layout.addWidget(cancel_btn)
        
        save_btn = QPushButton("Save")
        save_btn.clicked.connect(self.accept)
        save_btn.setDefault(True)
        btn_layout.addWidget(save_btn)
        
        layout.addLayout(btn_layout)
        
        # Initial type setup
        self._on_type_changed()
    
    def _on_type_changed(self):
        """Show/hide internal actions dropdown based on type"""
        is_internal = self.type_combo.currentData() == "internal"
        for i in range(self.internal_actions_layout.count()):
            widget = self.internal_actions_layout.itemAt(i).widget()
            if widget:
                widget.setVisible(is_internal)
    
    def populate_from_command(self, cmd: VoiceCommand):
        """Populate dialog from existing command"""
        self.phrase_edit.setText(cmd.phrase)
        self.aliases_edit.setText(", ".join(cmd.aliases))
        
        # Find and set action type
        for i in range(self.type_combo.count()):
            if self.type_combo.itemData(i) == cmd.action_type:
                self.type_combo.setCurrentIndex(i)
                break
        
        self.action_edit.setPlainText(cmd.action)
        self.desc_edit.setText(cmd.description)
        
        # Set category
        idx = self.cat_combo.findText(cmd.category)
        if idx >= 0:
            self.cat_combo.setCurrentIndex(idx)
        else:
            self.cat_combo.setCurrentText(cmd.category)
    
    def get_command(self) -> VoiceCommand:
        """Get the command from dialog inputs"""
        aliases_text = self.aliases_edit.text().strip()
        aliases = [a.strip() for a in aliases_text.split(",") if a.strip()] if aliases_text else []
        
        return VoiceCommand(
            phrase=self.phrase_edit.text().strip(),
            aliases=aliases,
            action_type=self.type_combo.currentData(),
            action=self.action_edit.toPlainText().strip(),
            description=self.desc_edit.text().strip(),
            category=self.cat_combo.currentText().strip(),
            enabled=True
        )


class AdvancedFiltersDialog(QDialog):
    """Dialog for advanced filtering options"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setup_ui()
    
    def setup_ui(self):
        self.setWindowTitle("Advanced Filters")
        self.setMinimumWidth(600)
        self.setMinimumHeight(500)
        layout = QVBoxLayout(self)
        
        # Create scroll area for filters
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.Shape.NoFrame)
        
        content = QWidget()
        content_layout = QVBoxLayout(content)
        content_layout.setSpacing(15)
        
        # Match Rate Filter
        match_group = QGroupBox("Match Rate (%)")
        match_layout = QHBoxLayout()
        
        self.match_rate_check = CheckmarkCheckBox("Enable")
        match_layout.addWidget(self.match_rate_check)
        
        match_layout.addWidget(QLabel("From:"))
        self.match_min_spin = QSpinBox()
        self.match_min_spin.setRange(0, 100)
        self.match_min_spin.setValue(0)
        self.match_min_spin.setSuffix("%")
        match_layout.addWidget(self.match_min_spin)
        
        match_layout.addWidget(QLabel("To:"))
        self.match_max_spin = QSpinBox()
        self.match_max_spin.setRange(0, 100)
        self.match_max_spin.setValue(100)
        self.match_max_spin.setSuffix("%")
        match_layout.addWidget(self.match_max_spin)
        match_layout.addStretch()
        
        match_group.setLayout(match_layout)
        content_layout.addWidget(match_group)
        
        # Row Status Filter
        status_group = QGroupBox("Row Status")
        status_layout = QVBoxLayout()
        
        self.status_not_started = CheckmarkCheckBox("Not started")
        self.status_edited = CheckmarkCheckBox("Edited")
        self.status_translated = CheckmarkCheckBox("Translated")
        self.status_confirmed = CheckmarkCheckBox("Confirmed")
        self.status_draft = CheckmarkCheckBox("Draft")
        
        status_layout.addWidget(self.status_not_started)
        status_layout.addWidget(self.status_edited)
        status_layout.addWidget(self.status_translated)
        status_layout.addWidget(self.status_confirmed)
        status_layout.addWidget(self.status_draft)
        
        status_group.setLayout(status_layout)
        content_layout.addWidget(status_group)
        
        # Locked/Unlocked Filter
        locked_group = QGroupBox("Locked Status")
        locked_layout = QVBoxLayout()
        
        self.locked_both = CheckmarkRadioButton("Both locked and unlocked rows")
        self.locked_only = CheckmarkRadioButton("Only locked rows")
        self.locked_unlocked_only = CheckmarkRadioButton("Only unlocked rows")
        self.locked_both.setChecked(True)
        
        locked_layout.addWidget(self.locked_both)
        locked_layout.addWidget(self.locked_only)
        locked_layout.addWidget(self.locked_unlocked_only)
        
        locked_group.setLayout(locked_layout)
        content_layout.addWidget(locked_group)
        
        # Other Properties
        other_group = QGroupBox("Other Properties")
        other_layout = QVBoxLayout()
        
        self.has_comments_check = CheckmarkCheckBox("Has comments/notes")
        self.has_proofreading_check = CheckmarkCheckBox("Has proofreading issues")
        self.repetitions_check = CheckmarkCheckBox("Repetitions only")
        self.auto_propagated_check = CheckmarkCheckBox("Auto-propagated")
        
        other_layout.addWidget(self.has_comments_check)
        other_layout.addWidget(self.has_proofreading_check)
        other_layout.addWidget(self.repetitions_check)
        other_layout.addWidget(self.auto_propagated_check)
        
        other_group.setLayout(other_layout)
        content_layout.addWidget(other_group)
        
        content_layout.addStretch()
        scroll.setWidget(content)
        layout.addWidget(scroll)
        
        # Buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        reset_btn = QPushButton("Reset All")
        reset_btn.clicked.connect(self.reset_filters)
        button_layout.addWidget(reset_btn)
        
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(cancel_btn)
        
        apply_btn = QPushButton("Apply Filters")
        apply_btn.setStyleSheet("background-color: #2196F3; color: white; font-weight: bold; padding: 5px 15px; border: none; outline: none;")
        apply_btn.clicked.connect(self.accept)
        apply_btn.setDefault(True)
        button_layout.addWidget(apply_btn)
        
        layout.addLayout(button_layout)
    
    def reset_filters(self):
        """Reset all filters to default"""
        self.match_rate_check.setChecked(False)
        self.match_min_spin.setValue(0)
        self.match_max_spin.setValue(100)
        
        self.status_not_started.setChecked(False)
        self.status_edited.setChecked(False)
        self.status_translated.setChecked(False)
        self.status_confirmed.setChecked(False)
        self.status_draft.setChecked(False)
        
        self.locked_both.setChecked(True)
        
        self.has_comments_check.setChecked(False)
        self.has_proofreading_check.setChecked(False)
        self.repetitions_check.setChecked(False)
        self.auto_propagated_check.setChecked(False)
    
    def get_filters(self):
        """Return dictionary of filter settings"""
        filters = {}
        
        # Match rate
        filters['match_rate_enabled'] = self.match_rate_check.isChecked()
        filters['match_rate_min'] = self.match_min_spin.value()
        filters['match_rate_max'] = self.match_max_spin.value()
        
        # Row status
        row_status = []
        if self.status_not_started.isChecked():
            row_status.append('not_started')
        if self.status_edited.isChecked():
            row_status.append('edited')
        if self.status_translated.isChecked():
            row_status.append('translated')
        if self.status_confirmed.isChecked():
            row_status.append('confirmed')
        if self.status_draft.isChecked():
            row_status.append('draft')
        filters['row_status'] = row_status
        
        # Locked filter
        if self.locked_only.isChecked():
            filters['locked_filter'] = 'locked'
        elif self.locked_unlocked_only.isChecked():
            filters['locked_filter'] = 'unlocked'
        else:
            filters['locked_filter'] = None
        
        # Other properties
        filters['has_comments'] = self.has_comments_check.isChecked()
        filters['has_proofreading'] = self.has_proofreading_check.isChecked()
        filters['repetitions_only'] = self.repetitions_check.isChecked()
        filters['auto_propagated'] = self.auto_propagated_check.isChecked()
        
        return filters


# ============================================================================
# MAIN WINDOW
# ============================================================================

class SupervertalerQt(QMainWindow):
    """Main application window"""
    
    MAX_RECENT_PROJECTS = 10  # Maximum number of recent projects to track
    
    # Signal for thread-safe logging (background threads emit, main thread handles)
    _log_signal = pyqtSignal(str)
    
    def __init__(self):
        super().__init__()
        
        # Connect thread-safe log signal (must be done first for logging to work from threads)
        self._log_signal.connect(self._log_to_ui)
        
        # Application state
        self.current_project: Optional[Project] = None
        self.project_file_path: Optional[str] = None
        self.project_modified = False
        
        # memoQ bilingual DOCX import tracking
        self.memoq_source_file = None
        
        # UI Configuration
        self.default_font_family = "Calibri"
        self.default_font_size = 11
        
        # Application settings
        self.allow_replace_in_source = False  # Safety: don't allow replace in source by default
        self.auto_propagate_exact_matches = True  # Auto-fill 100% TM matches for empty segments
        self.auto_insert_100_percent_matches = True  # Auto-insert 100% TM matches when segment selected
        self.auto_confirm_100_percent_matches = False  # Auto-confirm 100% matches when navigating with Ctrl+Enter
        self.auto_confirm_overwrite_existing = False  # Allow auto-confirm to overwrite existing target content
        self.tm_save_mode = 'latest'  # 'all' = keep all translations with timestamps, 'latest' = only keep most recent (DEFAULT)
        
        # Tab position setting
        self.tabs_above_grid = False  # Whether to show Termview/Session Log tabs above grid
        
        # Right panel visibility settings
        self.show_translation_results_pane = False  # Show Translation Results tab (hidden by default for new users)
        self.show_compare_panel = True  # Show Compare Panel tab
        
        # TM and Termbase matching toggle (default: enabled)
        self.enable_tm_matching = True
        self.enable_termbase_matching = True
        self.enable_mt_matching = True  # Machine Translation enabled
        self.enable_llm_matching = False  # LLM Translation enabled (DISABLED by default - too slow for navigation!)
        self.enable_termbase_grid_highlighting = True  # Highlight termbase matches in source cells

        # Termbase display settings
        self.termbase_display_order = 'appearance'  # Options: 'alphabetical', 'appearance', 'length'
        self.termbase_hide_shorter_matches = False  # Hide shorter terms included in longer ones

        # Invisible character display settings
        self.showing_invisible_spaces = False  # Track whether spaces are shown as middle dots
        self.invisible_char_color = '#999999'  # Light gray color for invisible characters
        self.invisible_display_settings = {
            'spaces': False,
            'tabs': False,
            'nbsp': False,
            'linebreaks': False
        }

        # Grid row color settings (memoQ-style alternating row colors)
        self.enable_alternating_row_colors = True  # Enable alternating row colors by default
        self.even_row_color = '#FFFFFF'  # White for even rows
        self.odd_row_color = '#F0F0F0'  # Light gray for odd rows
        
        # Termbase highlight style settings
        self.termbase_highlight_style = 'semibold'  # 'background', 'dotted', or 'semibold'
        self.termbase_dotted_color = '#808080'  # Medium gray for dotted underline (more visible)
        
        # Focus border settings for target cells
        self.focus_border_color = '#f1b79a'  # Peach/salmon
        self.focus_border_thickness = 2  # 2px

        # Debug mode settings (for troubleshooting performance issues)
        self.debug_mode_enabled = False  # Enables verbose debug logging
        self.debug_auto_export = False  # Auto-export debug logs to file
        self.debug_log_buffer = []  # Buffer for debug logs (for export)
        
        # Precision scroll settings (for fine-tuned grid navigation)
        self.precision_scroll_divisor = 3  # Divide row height by this (higher = finer increments)
        self.auto_center_active_segment = True  # Auto-scroll to keep active segment centered (default ON)
        
        # Translation service availability flags (would be set from config/API keys)
        self.google_translate_enabled = True  # For demo purposes
        self.deepl_enabled = False  # Not implemented yet
        self.openai_enabled = True  # For demo purposes
        self.claude_enabled = True  # For demo purposes
        
        # Timer for delayed lookup (cancel if user moves to another segment)
        self.lookup_timer = None
        self.current_lookup_segment_id = None
        
        # Termbase cache for performance optimization
        # Maps segment ID → {term: translation} dictionary
        self.termbase_cache = {}
        self.termbase_cache_lock = threading.Lock()  # Thread-safe cache access
        self.termbase_batch_worker_thread = None  # Background worker thread
        self.termbase_batch_stop_event = threading.Event()  # Signal to stop background worker
        
        # TM/MT/LLM prefetch cache for instant segment switching (like memoQ)
        # Maps segment ID → {"TM": [...], "MT": [...], "LLM": [...]}
        self.translation_matches_cache = {}
        self.translation_matches_cache_lock = threading.Lock()
        self.prefetch_worker_thread = None
        self.prefetch_stop_event = threading.Event()
        self.prefetch_queue = []  # List of segment IDs to prefetch
        
        # Undo/Redo stack for grid edits
        self.undo_stack = []  # List of (segment_id, old_target, new_target, old_status, new_status)
        self.redo_stack = []  # List of undone actions that can be redone
        self.max_undo_levels = 100  # Maximum number of undo levels to keep
        
        # Global language settings (defaults)
        self.source_language = "English"
        self.target_language = "Dutch"

        # Supervoice model download tracking
        self.is_loading_model = False
        self.loading_model_name = None
        
        # Target editor signal suppression (prevents load-time churn)
        self._suppress_target_change_handlers = False
        self.warning_banners: Dict[str, QWidget] = {}
        
        # Superlookup detached window
        self.lookup_detached_window = None
        
        # ============================================================================
        # USER DATA PATH INITIALIZATION
        # ============================================================================
        # Set up user data paths - handles both development and frozen (EXE) builds
        # In frozen builds, user_data is copied directly next to the EXE by the build script.
        from modules.database_manager import DatabaseManager
        
        if ENABLE_PRIVATE_FEATURES:
            # Developer mode: use private folder (git-ignored)
            self.user_data_path = Path(__file__).parent / "user_data_private"
        else:
            # Normal mode: use the helper function
            self.user_data_path = get_user_data_path()
        
        # Ensure user_data directory exists (creates empty folder if missing)
        self.user_data_path.mkdir(parents=True, exist_ok=True)
        
        print(f"[Data Paths] User data: {self.user_data_path}")
        
        # Database Manager for Termbases
        self.db_manager = DatabaseManager(
            db_path=str(self.user_data_path / "Translation_Resources" / "supervertaler.db"),
            log_callback=self.log
        )
        self.db_manager.connect()
        
        # TM Database - Initialize early so Superlookup works without a project loaded
        from modules.translation_memory import TMDatabase
        self.tm_database = TMDatabase(
            source_lang=None,  # Will be set when project is loaded
            target_lang=None,  # Will be set when project is loaded
            db_path=str(self.user_data_path / "Translation_Resources" / "supervertaler.db"),
            log_callback=self.log
        )
        
        # TM Metadata Manager - needed for TM list in Superlookup
        from modules.tm_metadata_manager import TMMetadataManager
        self.tm_metadata_mgr = TMMetadataManager(self.db_manager, self.log)
        
        # Spellcheck Manager for target language spell checking
        self.spellcheck_manager = get_spellcheck_manager(str(self.user_data_path))
        # Note: spellcheck_enabled will be loaded from preferences later in _load_spellcheck_settings()
        # For now set to False, it gets updated when UI is created
        self.spellcheck_enabled = False
        # Set up the shared spellcheck manager for TagHighlighter instances
        TagHighlighter.set_spellcheck_manager(self.spellcheck_manager)
        TagHighlighter.set_spellcheck_enabled(self.spellcheck_enabled)
        
        # Find & Replace History Manager
        self.fr_history = FindReplaceHistory(str(self.user_data_path))
        
        # Shortcut Manager for keyboard shortcuts (including enable/disable)
        self.shortcut_manager = ShortcutManager(Path(self.user_data_path) / "shortcuts.json")
        
        # Voice Command Manager for Talon-style voice commands
        self.voice_command_manager = VoiceCommandManager(self.user_data_path, main_window=self)
        
        # Continuous Voice Listener (always-on mode) - initialized on demand
        self.voice_listener = None  # Will be ContinuousVoiceListener when enabled
        
        # Figure Context Manager for multimodal AI translation
        from modules.figure_context_manager import FigureContextManager
        self.figure_context = FigureContextManager(self)
        
        # Theme Manager
        from modules.theme_manager import ThemeManager
        self.theme_manager = None  # Will be initialized after UI setup
        
        # User data path - uses safety system to prevent private data leaks
        # If .supervertaler.local exists: uses "user_data_private" (git-ignored)
        # Otherwise: uses "user_data" (safe to commit)
        base_folder = "user_data_private" if ENABLE_PRIVATE_FEATURES else "user_data"
        self.recent_projects_file = self.user_data_path / "recent_projects.json"
        
        # Initialize UI
        self.init_ui()
        
        # Initialize theme manager and apply theme
        self.theme_manager = ThemeManager(self.user_data_path)
        self.theme_manager.apply_theme(QApplication.instance())
        
        # Update widgets that were created before theme_manager existed
        if hasattr(self, 'termview_widget') and self.termview_widget:
            self.termview_widget.theme_manager = self.theme_manager
            if hasattr(self.termview_widget, 'apply_theme'):
                self.termview_widget.apply_theme()
        
        if hasattr(self, 'translation_results_panel') and self.translation_results_panel:
            self.translation_results_panel.theme_manager = self.theme_manager
            # Also update class-level theme_manager for CompactMatchItem
            from modules.translation_results_panel import CompactMatchItem
            CompactMatchItem.theme_manager = self.theme_manager
            print(f"🎨 DEBUG: Calling apply_theme on translation_results_panel", flush=True)
            # Write to file for debugging
            with open("theme_debug.txt", "w") as f:
                f.write(f"apply_theme called, theme={self.theme_manager.current_theme.name}\n")
                f.write(f"compare_text_edits count: {len(self.translation_results_panel.compare_text_edits)}\n")
                f.flush()
            if hasattr(self.translation_results_panel, 'apply_theme'):
                self.translation_results_panel.apply_theme()
                print(f"🎨 DEBUG: apply_theme() called successfully", flush=True)
        else:
            print(f"🎨 DEBUG: translation_results_panel NOT FOUND!", flush=True)
            with open("theme_debug.txt", "w") as f:
                f.write("translation_results_panel NOT FOUND!\n")
                f.flush()
        
        # Schedule theme refresh after UI is fully initialized
        # This ensures all widgets are properly themed at startup
        from PyQt6.QtCore import QTimer
        QTimer.singleShot(100, self.refresh_theme_colors)
        
        # Create example API keys file on first launch (after UI is ready)
        self.ensure_example_api_keys()
        
        self.log(f"Welcome to Supervertaler v{__version__}")
        self.log("Supervertaler: The Ultimate Translation Workbench.")
        
        # Load general settings (including auto-propagation)
        self.load_general_settings()
        
        # Load language settings
        self.load_language_settings()
        
        # Restore last project if enabled in settings
        self.restore_last_project_if_enabled()
        
        # Auto-open log window if enabled in settings
        general_settings = self.load_general_settings()
        if general_settings.get('auto_open_log', False):
            # Use QTimer to open log window after UI fully initializes
            from PyQt6.QtCore import QTimer
            QTimer.singleShot(500, self.detach_log_window)  # 500ms delay to ensure UI is ready

        # Initialize auto backup timer
        self.auto_backup_timer = None
        self.last_backup_time = None
        self.restart_auto_backup_timer()

        # Load font sizes from preferences (after UI is fully initialized)
        QApplication.instance().processEvents()  # Allow UI to finish initializing
        self.load_font_sizes_from_preferences()
        
        # Auto-check for new models if enabled in settings
        if general_settings.get('auto_check_models', True):
            from PyQt6.QtCore import QTimer
            QTimer.singleShot(2000, lambda: self._check_for_new_models(force=False))  # 2 second delay
        
        # First-run check - show Features tab to new users
        if not general_settings.get('first_run_completed', False):
            from PyQt6.QtCore import QTimer
            QTimer.singleShot(500, self._show_first_run_welcome)
    
    def _show_first_run_welcome(self):
        """Show welcome message and Features tab on first run."""
        try:
            # Create a custom dialog with checkbox
            from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QCheckBox, QDialogButtonBox
            
            dialog = QDialog(self)
            dialog.setWindowTitle("Welcome to Supervertaler!")
            dialog.setMinimumWidth(450)
            
            layout = QVBoxLayout(dialog)
            layout.setSpacing(15)
            
            # Icon and title
            title_label = QLabel("<h2>Welcome to Supervertaler! 🎉</h2>")
            layout.addWidget(title_label)
            
            # Message
            msg_label = QLabel(
                "Supervertaler uses a <b>modular architecture</b> - you can install "
                "only the features you need to save disk space.<br><br>"
                "We'll now show you the <b>Features</b> tab where you can see which "
                "optional components are installed and how to add more.<br><br>"
                "💡 <b>Tip:</b> You can always access this from Settings → Features."
            )
            msg_label.setWordWrap(True)
            layout.addWidget(msg_label)
            
            # Checkbox - use our standard green checkmark style
            dont_show_checkbox = CheckmarkCheckBox("Don't show this again")
            dont_show_checkbox.setChecked(True)  # Default to not showing again
            layout.addWidget(dont_show_checkbox)
            
            # OK button
            button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok)
            button_box.accepted.connect(dialog.accept)
            layout.addWidget(button_box)
            
            dialog.exec()
            
            # Navigate to Settings → Features tab
            self.main_tabs.setCurrentIndex(4)  # Settings tab
            if hasattr(self, 'settings_tabs'):
                # Find the Features tab index
                for i in range(self.settings_tabs.count()):
                    if "Features" in self.settings_tabs.tabText(i):
                        self.settings_tabs.setCurrentIndex(i)
                        break
            
            # Save the preference to general_settings.json (where load_general_settings reads from)
            if dont_show_checkbox.isChecked():
                settings = self.load_general_settings()
                settings['first_run_completed'] = True
                self.save_general_settings(settings)
                self.log("✅ First-run welcome completed (won't show again)")
            else:
                self.log("✅ First-run welcome shown (will show again next time)")
        except Exception as e:
            self.log(f"⚠️ First-run welcome error: {e}")
    
    def _check_for_new_models(self, force: bool = False):
        """
        Check for new LLM models from providers

        Args:
            force: Force check even if checked recently
        """
        try:
            from modules.model_version_checker import ModelVersionChecker
            from modules.model_update_dialog import ModelUpdateDialog, NoNewModelsDialog
            from modules.llm_clients import load_api_keys

            # Load API keys
            api_keys = load_api_keys()

            # Initialize checker with cache in user_data
            cache_path = self.user_data_path / "model_version_cache.json"
            checker = ModelVersionChecker(cache_path=str(cache_path))

            # Check if we should run (unless forced)
            if not force and not checker.should_check():
                # Already checked recently
                return

            self.log("🔍 Checking for new LLM models...")

            # Run the check
            results = checker.check_all_providers(
                openai_key=api_keys.get("OPENAI_API_KEY"),
                anthropic_key=api_keys.get("ANTHROPIC_API_KEY"),
                google_key=api_keys.get("GOOGLE_API_KEY"),
                force=force
            )

            # If this was a forced manual check, log the results
            if force:
                self.log(f"  OpenAI: {len(results.get('openai', {}).get('new_models', []))} new models")
                self.log(f"  Claude: {len(results.get('claude', {}).get('new_models', []))} new models")
                self.log(f"  Gemini: {len(results.get('gemini', {}).get('new_models', []))} new models")

            # Check if any new models found
            if checker.has_new_models(results):
                self.log(f"✨ New models detected! Opening dialog...")

                # Show dialog with new models
                dialog = ModelUpdateDialog(results, parent=self)
                dialog.models_selected.connect(self._on_new_models_selected)
                dialog.exec()

            else:
                # No new models
                if force:
                    # Only show "no new models" dialog for manual checks
                    cache_info = checker.get_cache_info()
                    dialog = NoNewModelsDialog(
                        last_check=cache_info.get('last_check'),
                        parent=self
                    )
                    dialog.exec()
                else:
                    # Silent for automatic checks
                    self.log("✓ No new models detected")

        except ImportError as e:
            self.log(f"⚠ Could not check for new models: Missing dependencies ({e})")
        except Exception as e:
            self.log(f"⚠ Error checking for new models: {e}")
            if force:
                # Show error dialog for manual checks
                from PyQt6.QtWidgets import QMessageBox
                QMessageBox.warning(
                    self,
                    "Model Check Error",
                    f"Failed to check for new models:\n\n{str(e)}"
                )

    def _on_new_models_selected(self, selected_models: dict):
        """
        Handle user selection of new models to add

        Args:
            selected_models: Dict of {provider: [model_ids]}
        """
        try:
            # Update the known models in llm_clients.py
            # For now, just log what would be added
            self.log("📦 Adding selected models to Supervertaler:")

            for provider, models in selected_models.items():
                if models:
                    self.log(f"  {provider.capitalize()}: {len(models)} model(s)")
                    for model in models:
                        self.log(f"    • {model}")

            # TODO: Actually add the models to the configuration
            # This would require modifying llm_clients.py or a separate config file
            # For now, just show a message
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.information(
                self,
                "Models Added",
                f"Selected models have been noted:\n\n"
                f"{sum(len(m) for m in selected_models.values())} model(s) from "
                f"{len(selected_models)} provider(s)\n\n"
                f"These models will be available after restarting Supervertaler.\n\n"
                f"Note: You may need to manually add them to Settings → AI Settings for now."
            )

        except Exception as e:
            self.log(f"❌ Error adding models: {e}")

    def init_ui(self):
        """Initialize the user interface"""
        # Build window title with dev mode indicator
        title = f"Supervertaler v{__version__}"
        if ENABLE_PRIVATE_FEATURES:
            title += " [🛠️ DEV MODE]"
        self.setWindowTitle(title)
        self.setGeometry(100, 100, 1400, 800)

        # Set application icon
        from PyQt6.QtGui import QIcon
        icon_path = get_resource_path("assets/icon.ico")
        if icon_path.exists():
            self.setWindowIcon(QIcon(str(icon_path)))

        # Ensure window can be resized (no minimum size constraint)
        self.setMinimumSize(400, 300)  # Very small minimum to allow resizing
        
        # Create menu bar (ribbon removed - using traditional menus)
        self.create_menus()
        
        # Ribbon removed - all functionality moved to menu bar
        # self.create_ribbon()
        
        # Create main layout
        self.create_main_layout()
        
        # Create status bar with progress indicators
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        self.status_bar.showMessage("Ready")
        
        # Add permanent progress widgets to status bar (right side)
        self._setup_progress_indicators()

        # Setup global shortcuts
        self.setup_global_shortcuts()

    def setup_global_shortcuts(self):
        """Setup application-wide keyboard shortcuts"""
        from PyQt6.QtGui import QShortcut
        
        # Store all shortcuts with their IDs and default key sequences for enable/disable management
        self.global_shortcuts = {}  # shortcut_id -> QShortcut
        self.global_shortcut_keys = {}  # shortcut_id -> current key sequence string (from manager)
        
        def create_shortcut(shortcut_id: str, default_key: str, handler):
            """Create a shortcut, using custom key from manager if set, respecting enabled state"""
            # Get the actual key sequence (custom or default) from the manager
            key_sequence = self.shortcut_manager.get_shortcut(shortcut_id)
            # Important: an empty string is a valid "unassigned" default.
            # Only fall back to the passed default_key if the shortcut ID is unknown to the manager.
            if key_sequence == "":
                default_ids = getattr(self.shortcut_manager, 'DEFAULT_SHORTCUTS', {})
                custom_ids = getattr(self.shortcut_manager, 'custom_shortcuts', {})
                if shortcut_id not in default_ids and shortcut_id not in custom_ids:
                    key_sequence = default_key
            
            shortcut = QShortcut(QKeySequence(key_sequence), self)
            shortcut.activated.connect(handler)
            # Store the current key sequence for re-enabling later
            self.global_shortcut_keys[shortcut_id] = key_sequence
            # If disabled, clear the key sequence to fully release the key combination
            if not self.shortcut_manager.is_enabled(shortcut_id):
                shortcut.setKey(QKeySequence())  # Clear key to release combination
            self.global_shortcuts[shortcut_id] = shortcut
            return shortcut

        # F9 - Voice dictation
        create_shortcut("voice_dictate", "F9", self.start_voice_dictation)
        
        # Ctrl+Up/Down - Cycle through translation matches
        create_shortcut("match_cycle_previous", "Ctrl+Up", self.select_previous_match)
        create_shortcut("match_cycle_next", "Ctrl+Down", self.select_next_match)
        
        # Ctrl+1 through Ctrl+9 - Insert match by number
        self.match_shortcuts = []
        for i in range(1, 10):
            shortcut_id = f"match_insert_{i}"
            default_key = f"Ctrl+{i}"
            key_sequence = self.shortcut_manager.get_shortcut(shortcut_id)
            # Respect intentionally blank defaults (unassigned shortcuts)
            if key_sequence == "":
                default_ids = getattr(self.shortcut_manager, 'DEFAULT_SHORTCUTS', {})
                custom_ids = getattr(self.shortcut_manager, 'custom_shortcuts', {})
                if shortcut_id not in default_ids and shortcut_id not in custom_ids:
                    key_sequence = default_key
            shortcut = QShortcut(QKeySequence(key_sequence), self)
            shortcut.activated.connect(lambda num=i: self.insert_match_by_number(num))
            self.global_shortcut_keys[shortcut_id] = key_sequence
            if not self.shortcut_manager.is_enabled(shortcut_id):
                shortcut.setKey(QKeySequence())  # Clear key to release combination
            self.global_shortcuts[shortcut_id] = shortcut
            self.match_shortcuts.append(shortcut)

        # Compare Panel insertion shortcut
        # Alt+0 inserts MT; double-tap Alt+0,0 inserts TM Target.
        self._compare_panel_last_key = None
        self._compare_panel_last_time = 0
        create_shortcut("compare_insert_alt0", "Alt+0", self._handle_compare_panel_alt0_shortcut)

        # Compare Panel navigation shortcuts
        # MT prev/next: Ctrl+Alt+Left / Ctrl+Alt+Right
        # TM prev/next: Ctrl+Alt+Up / Ctrl+Alt+Down
        create_shortcut(
            "compare_nav_mt_prev",
            "Ctrl+Alt+Left",
            lambda: self._compare_panel_nav_mt(-1) if self._get_active_match_shortcut_mode() == 'compare' else None,
        )
        create_shortcut(
            "compare_nav_mt_next",
            "Ctrl+Alt+Right",
            lambda: self._compare_panel_nav_mt(1) if self._get_active_match_shortcut_mode() == 'compare' else None,
        )
        create_shortcut(
            "compare_nav_tm_prev",
            "Ctrl+Alt+Up",
            lambda: self._compare_panel_nav_tm(-1) if self._get_active_match_shortcut_mode() == 'compare' else None,
        )
        create_shortcut(
            "compare_nav_tm_next",
            "Ctrl+Alt+Down",
            lambda: self._compare_panel_nav_tm(1) if self._get_active_match_shortcut_mode() == 'compare' else None,
        )
        
        # Alt+0 through Alt+9 - Insert term from TermView by number
        # Supports double-tap for terms 11-20 (00, 11, 22, ..., 99)
        self.termview_shortcuts = []
        self._termview_last_key = None  # Track last key for double-tap detection
        self._termview_last_time = 0    # Track timing for double-tap
        
        # Double-tap Shift detection for context menu
        self._shift_last_press_time = 0
        self._shift_double_tap_threshold = 0.35  # 350ms window
        
        for i in range(0, 10):  # 0-9
            shortcut_id = f"termview_insert_{i}"
            default_key = f"Alt+{i}"
            key_sequence = self.shortcut_manager.get_shortcut(shortcut_id)
            # Respect intentionally blank defaults (unassigned shortcuts)
            if key_sequence == "":
                default_ids = getattr(self.shortcut_manager, 'DEFAULT_SHORTCUTS', {})
                custom_ids = getattr(self.shortcut_manager, 'custom_shortcuts', {})
                if shortcut_id not in default_ids and shortcut_id not in custom_ids:
                    key_sequence = default_key
            shortcut = QShortcut(QKeySequence(key_sequence), self)
            shortcut.activated.connect(lambda num=i: self._handle_termview_shortcut(num))
            self.global_shortcut_keys[shortcut_id] = key_sequence
            if not self.shortcut_manager.is_enabled(shortcut_id):
                shortcut.setKey(QKeySequence())  # Clear key to release combination
            self.global_shortcuts[shortcut_id] = shortcut
            self.termview_shortcuts.append(shortcut)
        
        # Ctrl+Space - Insert currently selected match
        create_shortcut("match_insert_selected_ctrl", "Ctrl+Space", self.insert_selected_match)
        
        # Alt+Up/Down - Navigate to previous/next segment
        create_shortcut("segment_previous", "Alt+Up", self.go_to_previous_segment)
        create_shortcut("segment_next", "Alt+Down", self.go_to_next_segment)
        
        # Ctrl+Home/End - Navigate to first/last segment
        create_shortcut("segment_go_to_top", "Ctrl+Home", self.go_to_first_segment)
        create_shortcut("segment_go_to_bottom", "Ctrl+End", self.go_to_last_segment)
        
        # Ctrl+Enter - Confirm segment(s) and go to next unconfirmed
        # If multiple segments selected: confirm all selected
        # If single segment: confirm and go to next unconfirmed
        ctrl_enter_shortcut = create_shortcut("editor_save_and_next", "Ctrl+Return", self.confirm_selected_or_next)
        # CRITICAL: Use ApplicationShortcut context so it works even when focus is in QTextEdit widgets
        ctrl_enter_shortcut.setContext(Qt.ShortcutContext.ApplicationShortcut)

        # Workaround: ensure Ctrl+Return (main keyboard) is caught even if swallowed before QShortcut
        # (Observed: Ctrl+Enter on numpad works, Ctrl+Return does not when focus is in source cell.)
        if not hasattr(self, '_ctrl_return_event_filter'):
            from PyQt6.QtWidgets import QApplication
            self._ctrl_return_event_filter = _CtrlReturnEventFilter(self)
            QApplication.instance().installEventFilter(self._ctrl_return_event_filter)
        
        # Note: Double-tap Shift for context menu is handled by AutoHotkey (double_shift_menu.ahk)
        # Qt's event system makes reliable double-tap detection difficult in Python
        
        # Ctrl+Shift+Enter - Always confirm all selected segments
        create_shortcut("editor_confirm_selected", "Ctrl+Shift+Return", self.confirm_selected_segments)
        
        # Note: Ctrl+Shift+S (Copy source to target) is handled in EditableGridTextEditor.keyPressEvent
        
        # Ctrl+K - Concordance Search
        create_shortcut("tools_concordance_search", "Ctrl+K", self.show_concordance_search)
        
        # Ctrl+Shift+F - Filter on selected text / Clear filter (toggle)
        create_shortcut("filter_selected_text", "Ctrl+Shift+F", self.filter_on_selected_text)
        
        # Ctrl+Alt+T - Toggle Tag View
        create_shortcut("view_toggle_tags", "Ctrl+Alt+T", self._toggle_tag_view_via_shortcut)
        
        # Page Up/Down - Navigate pagination pages
        create_shortcut("page_prev", "PgUp", self.go_to_prev_page)
        create_shortcut("page_next", "PgDown", self.go_to_next_page)
        
        # Shift+Page Up/Down - Select range of segments
        create_shortcut("select_range_up", "Shift+PgUp", self.select_range_page_up)
        create_shortcut("select_range_down", "Shift+PgDown", self.select_range_page_down)
        
        # Ctrl+G - Go to segment
        create_shortcut("edit_goto", "Ctrl+G", self.show_goto_dialog)
        
        # F5 - Force refresh matches (clear all caches and re-search)
        create_shortcut("tools_force_refresh", "F5", self.force_refresh_matches)
        
        # Ctrl+Shift+1 - Quick add term with Priority 1
        create_shortcut("editor_quick_add_priority_1", "Ctrl+Shift+1", lambda: self._quick_add_term_with_priority(1))
        
        # Ctrl+Shift+2 - Quick add term with Priority 2
        create_shortcut("editor_quick_add_priority_2", "Ctrl+Shift+2", lambda: self._quick_add_term_with_priority(2))
        
        # Alt+D - Add word at cursor to dictionary
        create_shortcut("editor_add_to_dictionary", "Alt+D", self.add_word_to_dictionary_shortcut)
        
        # Ctrl+N - Focus Segment Note tab
        create_shortcut("editor_focus_notes", "Ctrl+N", self.focus_segment_notes)
        
        # Alt+K - Open QuickMenu directly
        create_shortcut("editor_open_quickmenu", "Alt+K", self.open_quickmenu)
    
    def focus_segment_notes(self):
        """Switch to Segment Note tab and focus the notes editor so user can start typing immediately"""
        if not hasattr(self, 'bottom_tabs'):
            return
        
        # Switch to Segment note tab (index 1)
        self.bottom_tabs.setCurrentIndex(1)
        
        # Focus the notes editor so user can start typing
        if hasattr(self, 'bottom_notes_edit'):
            self.bottom_notes_edit.setFocus()
    
    def open_quickmenu(self):
        """Open QuickMenu popup at current cursor position for quick AI prompt selection.
        
        User can navigate with arrow keys and press Enter to select a prompt.
        """
        try:
            # Get QuickMenu items from prompt library
            quickmenu_items = []
            if hasattr(self, 'prompt_manager_qt') and self.prompt_manager_qt:
                lib = getattr(self.prompt_manager_qt, 'library', None)
                if lib and hasattr(lib, 'get_quickmenu_grid_prompts'):
                    quickmenu_items = lib.get_quickmenu_grid_prompts() or []
            
            if not quickmenu_items:
                self.log("⚠️ No QuickMenu prompts available. Add prompts with 'Show in Supervertaler QuickMenu' enabled.")
                return
            
            # Find the currently focused widget (source or target cell)
            focus_widget = QApplication.focusWidget()
            
            # Build the menu
            menu = QMenu(self)
            menu.setTitle("⚡ QuickMenu")
            
            for rel_path, label in sorted(quickmenu_items, key=lambda x: (x[1] or x[0]).lower()):
                prompt_menu = menu.addMenu(label or rel_path)
                
                run_show = QAction("▶ Run (show response)…", self)
                run_show.triggered.connect(
                    lambda checked=False, p=rel_path, w=focus_widget: self.run_grid_quickmenu_prompt(p, origin_widget=w, behavior="show")
                )
                prompt_menu.addAction(run_show)
                
                run_replace = QAction("↺ Run and replace target selection", self)
                run_replace.triggered.connect(
                    lambda checked=False, p=rel_path, w=focus_widget: self.run_grid_quickmenu_prompt(p, origin_widget=w, behavior="replace")
                )
                prompt_menu.addAction(run_replace)
            
            # Show menu at cursor position (or center of focused widget)
            if focus_widget:
                # Get cursor rectangle if it's a text editor
                if hasattr(focus_widget, 'cursorRect'):
                    cursor_rect = focus_widget.cursorRect()
                    pos = focus_widget.mapToGlobal(cursor_rect.bottomLeft())
                else:
                    # Fallback to center of widget
                    pos = focus_widget.mapToGlobal(focus_widget.rect().center())
            else:
                # Fallback to mouse cursor position
                pos = QCursor.pos()
            
            menu.exec(pos)
            
        except Exception as e:
            self.log(f"❌ Error opening QuickMenu: {e}")
    
    def refresh_shortcut_enabled_states(self):
        """Refresh enabled/disabled states and key bindings of all global shortcuts from shortcut manager.
        
        When disabled: clears the key sequence to fully release the key combination.
        When enabled: restores the key sequence from the manager (custom or default).
        """
        if not hasattr(self, 'global_shortcuts'):
            return
        for shortcut_id, shortcut in self.global_shortcuts.items():
            is_enabled = self.shortcut_manager.is_enabled(shortcut_id)
            if is_enabled:
                # Get the current key sequence from manager (may have been customized)
                key_sequence = self.shortcut_manager.get_shortcut(shortcut_id)
                if key_sequence:
                    self.global_shortcut_keys[shortcut_id] = key_sequence
                    shortcut.setKey(QKeySequence(key_sequence))
            else:
                # Clear the key to release the combination for other uses
                shortcut.setKey(QKeySequence())

    def _setup_progress_indicators(self):
        """Setup permanent progress indicator widgets in the status bar"""
        from PyQt6.QtWidgets import QLabel, QFrame

        # Create a container frame for all progress info
        progress_frame = QFrame()
        progress_layout = QHBoxLayout(progress_frame)
        progress_layout.setContentsMargins(0, 0, 10, 0)
        progress_layout.setSpacing(20)

        # LLM Provider/Model indicator (leftmost in the permanent area)
        self.llm_indicator_label = QLabel("")
        self.llm_indicator_label.setStyleSheet("color: #888; font-size: 11px;")
        self.llm_indicator_label.setToolTip("Current LLM provider and model")
        progress_layout.addWidget(self.llm_indicator_label)
        
        # Update LLM indicator on startup
        self._update_llm_indicator()

        # Words translated label
        self.progress_words_label = QLabel("Words: --")
        self.progress_words_label.setStyleSheet("color: #555; font-size: 11px;")
        self.progress_words_label.setToolTip("Words with translation / Total words (percentage)")
        progress_layout.addWidget(self.progress_words_label)

        # Segments confirmed label
        self.progress_confirmed_label = QLabel("Confirmed: --")
        self.progress_confirmed_label.setStyleSheet("color: #555; font-size: 11px;")
        self.progress_confirmed_label.setToolTip("Confirmed segments / Total segments (percentage)")
        progress_layout.addWidget(self.progress_confirmed_label)

        # Remaining segments label
        self.progress_remaining_label = QLabel("Remaining: --")
        self.progress_remaining_label.setStyleSheet("color: #555; font-size: 11px;")
        self.progress_remaining_label.setToolTip("Segments still requiring work")
        progress_layout.addWidget(self.progress_remaining_label)
        
        # Files indicator (for multi-file projects)
        self.progress_files_label = QLabel("")
        self.progress_files_label.setStyleSheet("color: #2196F3; font-size: 11px;")
        self.progress_files_label.setToolTip("Files in multi-file project (click for details)")
        self.progress_files_label.setCursor(Qt.CursorShape.PointingHandCursor)
        self.progress_files_label.mousePressEvent = lambda e: self.show_file_progress_dialog()
        self.progress_files_label.hide()  # Hidden by default, shown for multi-file projects
        progress_layout.addWidget(self.progress_files_label)
        
        # Always-on voice indicator
        self.alwayson_indicator_label = QLabel("")
        self.alwayson_indicator_label.setStyleSheet("font-size: 11px; font-weight: bold;")
        self.alwayson_indicator_label.setToolTip("Always-on voice listening status\nClick to toggle")
        self.alwayson_indicator_label.setCursor(Qt.CursorShape.PointingHandCursor)
        self.alwayson_indicator_label.mousePressEvent = lambda e: self._toggle_alwayson_from_statusbar()
        self.alwayson_indicator_label.hide()  # Hidden until enabled
        progress_layout.addWidget(self.alwayson_indicator_label)

        # Add as permanent widget (stays on right side)
        self.status_bar.addPermanentWidget(progress_frame)
        
        # Initialize Ollama keep-warm timer
        self.ollama_keepwarm_timer = None
        self._setup_ollama_keepwarm()
    
    def _update_llm_indicator(self):
        """Update the LLM provider/model indicator in the status bar"""
        try:
            settings = self.load_llm_settings()
            provider = settings.get('provider', 'openai')
            model_key = f'{provider}_model'
            model = settings.get(model_key, 'gpt-4o')
            
            # Format nicely
            if provider == 'ollama':
                icon = "🖥️"
                display = f"{icon} Local: {model}"
            elif provider == 'openai':
                icon = "🤖"
                display = f"{icon} {model}"
            elif provider == 'claude':
                icon = "🟣"
                display = f"{icon} {model}"
            elif provider == 'gemini':
                icon = "💎"
                display = f"{icon} {model}"
            else:
                display = f"{provider}: {model}"
            
            if hasattr(self, 'llm_indicator_label'):
                self.llm_indicator_label.setText(display)
                self.llm_indicator_label.setToolTip(f"LLM Provider: {provider.title()}\nModel: {model}")
        except Exception as e:
            if hasattr(self, 'llm_indicator_label'):
                self.llm_indicator_label.setText("")
    
    def _setup_ollama_keepwarm(self):
        """Setup timer to keep Ollama model warm (loaded in memory)"""
        from PyQt6.QtCore import QTimer
        
        # Check settings to see if keep-warm is enabled
        general_settings = self.load_general_settings()
        keepwarm_enabled = general_settings.get('ollama_keepwarm', False)
        
        if keepwarm_enabled:
            self._start_ollama_keepwarm_timer()
    
    def _start_ollama_keepwarm_timer(self):
        """Start the Ollama keep-warm timer (pings every 4 minutes)"""
        from PyQt6.QtCore import QTimer
        
        if self.ollama_keepwarm_timer is None:
            self.ollama_keepwarm_timer = QTimer(self)
            self.ollama_keepwarm_timer.timeout.connect(self._ping_ollama_keepwarm)
        
        # Ping every 4 minutes (Ollama unloads after 5 min of inactivity)
        self.ollama_keepwarm_timer.start(4 * 60 * 1000)  # 4 minutes in ms
        self.log("🔥 Ollama keep-warm enabled (pinging every 4 minutes)")
    
    def _stop_ollama_keepwarm_timer(self):
        """Stop the Ollama keep-warm timer"""
        if self.ollama_keepwarm_timer:
            self.ollama_keepwarm_timer.stop()
            self.log("❄️ Ollama keep-warm disabled")
    
    def _ping_ollama_keepwarm(self):
        """Send a minimal request to Ollama to keep the model loaded"""
        import threading
        
        def ping():
            try:
                import requests
                settings = self.load_llm_settings()
                if settings.get('provider') != 'ollama':
                    return
                
                model = settings.get('ollama_model', 'qwen2.5:7b')
                endpoint = "http://localhost:11434"
                
                # Send minimal request to keep model warm
                response = requests.post(
                    f"{endpoint}/api/generate",
                    json={
                        "model": model,
                        "prompt": "hi",
                        "options": {"num_predict": 1}  # Generate just 1 token
                    },
                    timeout=30
                )
                if response.status_code == 200:
                    print(f"🔥 Ollama keep-warm ping successful ({model})")
            except Exception as e:
                print(f"⚠️ Ollama keep-warm ping failed: {e}")
        
        # Run in background thread to not block UI
        threading.Thread(target=ping, daemon=True).start()

    def update_progress_stats(self):
        """Update the progress indicator labels in the status bar"""
        try:
            if not self.current_project or not self.current_project.segments:
                self.progress_words_label.setText("Words: --")
                self.progress_confirmed_label.setText("Confirmed: --")
                self.progress_remaining_label.setText("Remaining: --")
                if hasattr(self, 'progress_files_label'):
                    self.progress_files_label.hide()
                return

            segments = self.current_project.segments
            total_segments = len(segments)

            # Count words in source text for each segment
            total_words = 0
            translated_words = 0
            confirmed_count = 0
            remaining_count = 0

            # Statuses that indicate "done" (confirmed or higher)
            confirmed_statuses = {'confirmed', 'tr_confirmed', 'proofread', 'approved'}
            # Statuses that need work
            unfinished_statuses = {'not_started', 'pretranslated', 'rejected'}

            for segment in segments:
                # Count words in source (simple split on whitespace)
                source_words = len(segment.source.split()) if segment.source else 0
                total_words += source_words

                # If segment has a non-empty translation, count those words as translated
                if segment.target and segment.target.strip():
                    translated_words += source_words

                # Count confirmed (or higher status)
                if segment.status in confirmed_statuses:
                    confirmed_count += 1

                # Count remaining (not_started, pretranslated, or rejected)
                if segment.status in unfinished_statuses:
                    remaining_count += 1

            # Calculate percentages
            word_percent = (translated_words / total_words * 100) if total_words > 0 else 0
            confirmed_percent = (confirmed_count / total_segments * 100) if total_segments > 0 else 0

            # Update labels with color-coding based on progress
            word_color = self._get_progress_color(word_percent)
            confirmed_color = self._get_progress_color(confirmed_percent)
            remaining_color = "#c00" if remaining_count > 0 else "#080"

            self.progress_words_label.setText(f"Words: {translated_words}/{total_words} ({word_percent:.0f}%)")
            self.progress_words_label.setStyleSheet(f"color: {word_color}; font-size: 11px;")

            self.progress_confirmed_label.setText(f"Confirmed: {confirmed_count}/{total_segments} ({confirmed_percent:.0f}%)")
            self.progress_confirmed_label.setStyleSheet(f"color: {confirmed_color}; font-size: 11px;")

            self.progress_remaining_label.setText(f"Remaining: {remaining_count}")
            self.progress_remaining_label.setStyleSheet(f"color: {remaining_color}; font-size: 11px;")
            
            # Multi-file project indicator
            if hasattr(self, 'progress_files_label'):
                is_multifile = getattr(self.current_project, 'is_multifile', False)
                files = getattr(self.current_project, 'files', [])
                
                if is_multifile and files:
                    # Count completed files
                    completed_files = 0
                    for file_info in files:
                        file_id = file_info['id']
                        file_segs = [s for s in segments if getattr(s, 'file_id', None) == file_id]
                        if file_segs:
                            conf_count = sum(1 for s in file_segs if s.status in confirmed_statuses)
                            if conf_count == len(file_segs):
                                completed_files += 1
                    
                    self.progress_files_label.setText(f"📁 Files: {completed_files}/{len(files)}")
                    self.progress_files_label.setToolTip(
                        f"Multi-file project: {len(files)} files\n"
                        f"Completed: {completed_files} files\n\n"
                        "Click for detailed file progress"
                    )
                    self.progress_files_label.show()
                else:
                    self.progress_files_label.hide()

        except Exception as e:
            self.log(f"⚠️ Error updating progress stats: {e}")

    def _get_progress_color(self, percent: float) -> str:
        """Get color for progress percentage: red < 50%, orange < 80%, green >= 80%"""
        if percent < 50:
            return "#c00"  # Red
        elif percent < 80:
            return "#c60"  # Orange
        else:
            return "#080"  # Green

    def create_menus(self):
        """Create application menus"""
        menubar = self.menuBar()
        
        # File Menu
        file_menu = menubar.addMenu("&File")
        
        new_action = QAction("&New Project", self)
        new_action.setShortcut(QKeySequence.StandardKey.New)
        new_action.triggered.connect(self.new_project)
        file_menu.addAction(new_action)
        
        open_action = QAction("&Open Project...", self)
        open_action.setShortcut(QKeySequence.StandardKey.Open)
        open_action.triggered.connect(self.open_project)
        file_menu.addAction(open_action)
        
        # Recent projects submenu
        self.recent_menu = file_menu.addMenu("Open &Recent")
        self.update_recent_menu()
        
        file_menu.addSeparator()
        
        save_action = QAction("&Save", self)
        save_action.setShortcut(QKeySequence.StandardKey.Save)
        save_action.triggered.connect(self.save_project)
        file_menu.addAction(save_action)
        
        save_as_action = QAction("Save &As...", self)
        # No keyboard shortcut - Ctrl+Shift+S is used for Copy Source to Target in editor
        save_as_action.triggered.connect(self.save_project_as)
        file_menu.addAction(save_as_action)
        
        file_menu.addSeparator()
        
        close_action = QAction("&Close Project", self)
        close_action.triggered.connect(self.close_project)
        file_menu.addAction(close_action)
        
        file_menu.addSeparator()
        
        # Import/Export submenu
        import_menu = file_menu.addMenu("&Import")
        
        import_docx_action = QAction("&Monolingual Document (DOCX)...", self)
        import_docx_action.triggered.connect(self.import_docx)
        import_docx_action.setShortcut("Ctrl+O")
        import_menu.addAction(import_docx_action)
        
        import_txt_action = QAction("Simple &Text File (TXT)...", self)
        import_txt_action.triggered.connect(self.import_simple_txt)
        import_menu.addAction(import_txt_action)
        
        import_menu.addSeparator()
        
        # Multi-file folder import
        import_folder_action = QAction("📁 &Folder (Multiple Files)...", self)
        import_folder_action.triggered.connect(self.import_folder_multifile)
        import_menu.addAction(import_folder_action)
        
        import_menu.addSeparator()  # Separate monolingual from bilingual tools
        
        import_memoq_action = QAction("memoQ &Bilingual Table (DOCX)...", self)
        import_memoq_action.triggered.connect(self.import_memoq_bilingual)
        import_menu.addAction(import_memoq_action)
        
        import_memoq_xliff_action = QAction("memoQ &XLIFF (.mqxliff)...", self)
        import_memoq_xliff_action.triggered.connect(self.import_memoq_xliff)
        import_menu.addAction(import_memoq_xliff_action)
        
        import_cafetran_action = QAction("&CafeTran Bilingual Table (DOCX)...", self)
        import_cafetran_action.triggered.connect(self.import_cafetran_bilingual)
        import_menu.addAction(import_cafetran_action)
        
        # Trados submenu - group all Trados imports together
        trados_submenu = import_menu.addMenu("&Trados Studio")

        import_trados_bilingual_action = QAction("Bilingual &Review (DOCX)...", self)
        import_trados_bilingual_action.triggered.connect(self.import_trados_bilingual)
        trados_submenu.addAction(import_trados_bilingual_action)

        import_sdlppx_action = QAction("&Package (SDLPPX)...", self)
        import_sdlppx_action.triggered.connect(self.import_sdlppx_package)
        trados_submenu.addAction(import_sdlppx_action)

        # Phrase (Memsource) import
        import_phrase_bilingual_action = QAction("&Phrase (Memsource) Bilingual (DOCX)...", self)
        import_phrase_bilingual_action.triggered.connect(self.import_phrase_bilingual)
        import_menu.addAction(import_phrase_bilingual_action)

        # Déjà Vu X3 import
        import_dejavu_action = QAction("&Déjà Vu X3 Bilingual (RTF)...", self)
        import_dejavu_action.triggered.connect(self.import_dejavu_bilingual)
        import_menu.addAction(import_dejavu_action)

        import_menu.addSeparator()
        
        import_review_table_action = QAction("&Bilingual Table (DOCX) - Update Project...", self)
        import_review_table_action.triggered.connect(self.import_review_table)
        import_menu.addAction(import_review_table_action)
        
        export_menu = file_menu.addMenu("&Export")
        
        export_memoq_action = QAction("memoQ &Bilingual Table - Translated (DOCX)...", self)
        export_memoq_action.triggered.connect(self.export_memoq_bilingual)
        export_menu.addAction(export_memoq_action)
        
        export_memoq_xliff_action = QAction("memoQ &XLIFF - Translated (.mqxliff)...", self)
        export_memoq_xliff_action.triggered.connect(self.export_memoq_xliff)
        export_menu.addAction(export_memoq_xliff_action)
        
        export_cafetran_action = QAction("&CafeTran Bilingual Table - Translated (DOCX)...", self)
        export_cafetran_action.triggered.connect(self.export_cafetran_bilingual)
        export_menu.addAction(export_cafetran_action)
        
        # Trados submenu - group all Trados exports together
        trados_export_submenu = export_menu.addMenu("&Trados Studio")

        export_trados_bilingual_action = QAction("Bilingual &Review - Translated (DOCX)...", self)
        export_trados_bilingual_action.triggered.connect(self.export_trados_bilingual)
        trados_export_submenu.addAction(export_trados_bilingual_action)

        export_sdlrpx_action = QAction("Return &Package (SDLRPX)...", self)
        export_sdlrpx_action.triggered.connect(self.export_sdlrpx_package)
        trados_export_submenu.addAction(export_sdlrpx_action)

        # Phrase (Memsource) export
        export_phrase_bilingual_action = QAction("&Phrase (Memsource) Bilingual - Translated (DOCX)...", self)
        export_phrase_bilingual_action.triggered.connect(self.export_phrase_bilingual)
        export_menu.addAction(export_phrase_bilingual_action)

        # Déjà Vu X3 export
        export_dejavu_action = QAction("&Déjà Vu X3 Bilingual - Translated (RTF)...", self)
        export_dejavu_action.triggered.connect(self.export_dejavu_bilingual)
        export_menu.addAction(export_dejavu_action)
        
        export_target_docx_action = QAction("&Target Only (DOCX)...", self)
        export_target_docx_action.triggered.connect(self.export_target_only_docx)
        export_menu.addAction(export_target_docx_action)
        
        export_txt_action = QAction("Simple &Text File - Translated (TXT)...", self)
        export_txt_action.triggered.connect(self.export_simple_txt)
        export_menu.addAction(export_txt_action)
        
        export_ai_action = QAction("📄 &AI-Readable Markdown (.md)...", self)
        export_ai_action.triggered.connect(self.export_bilingual_table_markdown)
        export_ai_action.setToolTip("Export segments in [SEGMENT] format for AI translation/review")
        export_menu.addAction(export_ai_action)
        
        export_menu.addSeparator()
        
        # Multi-file folder export
        export_folder_action = QAction("📁 &Folder (Multiple Files)...", self)
        export_folder_action.triggered.connect(self.export_folder_multifile)
        export_folder_action.setToolTip("Export multi-file project to folder with separate files")
        export_menu.addAction(export_folder_action)
        
        export_menu.addSeparator()
        
        # Relocate source folder for multi-file projects
        relocate_source_action = QAction("🔗 &Relocate Source Folder...", self)
        relocate_source_action.triggered.connect(self.relocate_source_folder)
        relocate_source_action.setToolTip("Repoint to moved/renamed source folder for multi-file project")
        export_menu.addAction(relocate_source_action)
        
        export_menu.addSeparator()
        
        # Supervertaler Bilingual Table exports
        export_review_table_action = QAction("Supervertaler Bilingual Table - With &Tags (DOCX)...", self)
        export_review_table_action.triggered.connect(self.export_review_table_with_tags)
        export_menu.addAction(export_review_table_action)
        
        export_review_table_formatted_action = QAction("Supervertaler Bilingual Table - &Formatted (DOCX)...", self)
        export_review_table_formatted_action.triggered.connect(self.export_review_table_formatted)
        export_menu.addAction(export_review_table_formatted_action)
        
        export_menu.addSeparator()
        
        export_grid_action = QAction("TMX from &Grid (all segments)...", self)
        export_grid_action.triggered.connect(self.export_tmx_from_grid)
        export_menu.addAction(export_grid_action)
        
        export_selected_action = QAction("TMX from &Selected Segments...", self)
        export_selected_action.triggered.connect(self.export_tmx_from_selected)
        export_menu.addAction(export_selected_action)
        
        export_tm_action = QAction("TMX from &TM(s) for Current Project...", self)
        export_tm_action.triggered.connect(self.export_tmx_from_tm_database)
        export_menu.addAction(export_tm_action)
        
        file_menu.addSeparator()
        
        # Project Info
        project_info_action = QAction("📋 Project &Info...", self)
        project_info_action.triggered.connect(self.show_project_info_dialog)
        file_menu.addAction(project_info_action)
        
        file_menu.addSeparator()
        
        exit_action = QAction("E&xit", self)
        exit_action.setShortcut(QKeySequence.StandardKey.Quit)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        # Edit Menu
        edit_menu = menubar.addMenu("&Edit")
        
        self.undo_action = QAction("&Undo", self)
        self.undo_action.setShortcut(QKeySequence.StandardKey.Undo)
        self.undo_action.triggered.connect(self.undo_action_handler)
        self.undo_action.setEnabled(False)
        edit_menu.addAction(self.undo_action)
        
        self.redo_action = QAction("&Redo", self)
        self.redo_action.setShortcut(QKeySequence.StandardKey.Redo)
        self.redo_action.triggered.connect(self.redo_action_handler)
        self.redo_action.setEnabled(False)
        edit_menu.addAction(self.redo_action)
        
        edit_menu.addSeparator()
        
        find_action = QAction("&Find...", self)
        find_action.setShortcut(QKeySequence.StandardKey.Find)
        find_action.triggered.connect(self.show_find_replace_dialog)
        edit_menu.addAction(find_action)
        
        replace_action = QAction("&Replace...", self)
        replace_action.setShortcut(QKeySequence.StandardKey.Replace)
        replace_action.triggered.connect(self.show_find_replace_dialog)
        edit_menu.addAction(replace_action)
        
        edit_menu.addSeparator()
        
        goto_action = QAction("&Go to Segment...\tCtrl+G", self)
        goto_action.triggered.connect(self.show_goto_dialog)
        edit_menu.addAction(goto_action)
        
        edit_menu.addSeparator()
        
        translate_action = QAction("&Translate Segment", self)
        translate_action.setShortcut("Ctrl+T")
        translate_action.triggered.connect(self.translate_current_segment)
        edit_menu.addAction(translate_action)

        translate_menu = edit_menu.addMenu("Batch &Translate")

        translate_selected_not_started_action = QAction("Translate selected not-started segments", self)
        translate_selected_not_started_action.triggered.connect(
            lambda checked=False: self.translate_multiple_segments("selected_not_started")
        )
        translate_menu.addAction(translate_selected_not_started_action)

        translate_all_not_started_action = QAction("Translate all not-started segments", self)
        translate_all_not_started_action.triggered.connect(
            lambda checked=False: self.translate_multiple_segments("all_not_started")
        )
        translate_menu.addAction(translate_all_not_started_action)

        translate_all_pretranslated_action = QAction("Translate all pre-translated segments", self)
        translate_all_pretranslated_action.triggered.connect(
            lambda checked=False: self.translate_multiple_segments("all_pretranslated")
        )
        translate_menu.addAction(translate_all_pretranslated_action)

        translate_pending_action = QAction("Translate all not-started & pre-translated", self)
        translate_pending_action.setShortcut("Ctrl+Shift+T")
        translate_pending_action.triggered.connect(
            lambda checked=False: self.translate_multiple_segments("all_not_started_pretranslated")
        )
        translate_menu.addAction(translate_pending_action)

        translate_translatable_action = QAction("Translate all translatable segments", self)
        translate_translatable_action.triggered.connect(
            lambda checked=False: self.translate_multiple_segments("all_translatable")
        )
        translate_menu.addAction(translate_translatable_action)

        translate_all_segments_action = QAction("Translate all segments (all statuses)", self)
        translate_all_segments_action.triggered.connect(
            lambda checked=False: self.translate_multiple_segments("all_segments")
        )
        translate_menu.addAction(translate_all_segments_action)
        
        translate_menu.addSeparator()
        
        # NEW: Translate only empty segments (for re-running failed batches)
        translate_empty_action = QAction("Translate all empty segments", self)
        translate_empty_action.setToolTip("Translate segments with empty target (useful after partial batch translation)")
        translate_empty_action.triggered.connect(
            lambda checked=False: self.translate_multiple_segments("all_empty")
        )
        translate_menu.addAction(translate_empty_action)
        
        translate_filtered_action = QAction("Translate all filtered segments", self)
        translate_filtered_action.setToolTip("Translate only segments currently visible after filtering")
        translate_filtered_action.triggered.connect(
            lambda checked=False: self.translate_multiple_segments("filtered_segments")
        )
        translate_menu.addAction(translate_filtered_action)
        
        edit_menu.addSeparator()
        
        # Bulk Operations submenu
        bulk_menu = edit_menu.addMenu("Bulk &Operations")
        
        confirm_selected_action = QAction("✅ &Confirm Selected Segments", self)
        confirm_selected_action.setShortcut("Ctrl+Shift+Return")
        confirm_selected_action.setToolTip("Confirm all selected segments (Ctrl+Shift+Enter)")
        confirm_selected_action.triggered.connect(self.confirm_selected_segments_from_menu)
        bulk_menu.addAction(confirm_selected_action)
        
        clear_translations_action = QAction("🗑️ &Clear Translations", self)
        clear_translations_action.setToolTip("Clear translations for selected segments")
        clear_translations_action.triggered.connect(self.clear_selected_translations_from_menu)
        bulk_menu.addAction(clear_translations_action)
        
        copy_source_to_target_action = QAction("📋 Copy &Source to Target", self)
        copy_source_to_target_action.setToolTip("Copy source text to target for selected/filtered segments")
        copy_source_to_target_action.triggered.connect(self.copy_source_to_target_bulk)
        bulk_menu.addAction(copy_source_to_target_action)
        
        send_to_tm_action = QAction("💾 &Send Segments to TM...", self)
        send_to_tm_action.setToolTip("Send confirmed segments to a writable Translation Memory")
        send_to_tm_action.triggered.connect(self.send_segments_to_tm_dialog)
        bulk_menu.addAction(send_to_tm_action)
        
        bulk_menu.addSeparator()
        
        clean_tags_action = QAction("🧹 Clean &Tags...", self)
        clean_tags_action.setToolTip("Remove formatting tags from selected segments")
        clean_tags_action.triggered.connect(self.show_clean_tags_dialog)
        bulk_menu.addAction(clean_tags_action)
        
        proofread_action = QAction("✅ &Proofread Translation...", self)
        proofread_action.setToolTip("Use AI to proofread and verify translation quality")
        proofread_action.triggered.connect(self.show_proofread_dialog)
        bulk_menu.addAction(proofread_action)
        
        clear_proofread_notes_action = QAction("🗑️ Clear All &Proofreading Notes", self)
        clear_proofread_notes_action.setToolTip("Remove all AI proofreading notes from entire project (preserves your personal notes)")
        clear_proofread_notes_action.triggered.connect(self._bulk_clear_proofreading_notes)
        bulk_menu.addAction(clear_proofread_notes_action)
        
        edit_menu.addSeparator()
        
        # Superlookup
        superlookup_action = QAction("🔍 &Superlookup...", self)
        superlookup_action.setShortcut("Ctrl+Alt+L")
        # Tab indices: Grid=0, Project resources=1, Tools=2, Settings=3
        superlookup_action.triggered.connect(lambda: self._go_to_superlookup() if hasattr(self, 'main_tabs') else None)  # Navigate to Superlookup
        edit_menu.addAction(superlookup_action)
        
        # View Menu
        view_menu = menubar.addMenu("&View")
        
        # Navigation submenu
        nav_menu = view_menu.addMenu("📑 &Navigate To")
        
        go_editor_action = QAction("📝 &Grid", self)
        go_editor_action.triggered.connect(lambda: self.main_tabs.setCurrentIndex(0) if hasattr(self, 'main_tabs') else None)
        nav_menu.addAction(go_editor_action)
        
        go_resources_action = QAction("🗂️ Project &resources", self)
        go_resources_action.triggered.connect(lambda: self.main_tabs.setCurrentIndex(1) if hasattr(self, 'main_tabs') else None)
        nav_menu.addAction(go_resources_action)
        
        go_prompt_manager_action = QAction("🤖 &Prompt Manager", self)
        go_prompt_manager_action.triggered.connect(lambda: self.main_tabs.setCurrentIndex(2) if hasattr(self, 'main_tabs') else None)
        nav_menu.addAction(go_prompt_manager_action)
        
        go_tools_action = QAction("🛠️ &Tools", self)
        go_tools_action.triggered.connect(lambda: self.main_tabs.setCurrentIndex(3) if hasattr(self, 'main_tabs') else None)
        nav_menu.addAction(go_tools_action)
        
        go_settings_action = QAction("⚙️ &Settings", self)
        go_settings_action.triggered.connect(lambda: self.main_tabs.setCurrentIndex(4) if hasattr(self, 'main_tabs') else None)
        nav_menu.addAction(go_settings_action)
        
        view_menu.addSeparator()
        
        # Grid Text section
        grid_zoom_menu = view_menu.addMenu("📊 &Grid Text Zoom")
        
        grid_zoom_in = QAction("Grid Zoom &In", self)
        grid_zoom_in.setShortcut(QKeySequence.StandardKey.ZoomIn)
        grid_zoom_in.triggered.connect(self.zoom_in)
        grid_zoom_menu.addAction(grid_zoom_in)
        
        grid_zoom_out = QAction("Grid Zoom &Out", self)
        grid_zoom_out.setShortcut(QKeySequence.StandardKey.ZoomOut)
        grid_zoom_out.triggered.connect(self.zoom_out)
        grid_zoom_menu.addAction(grid_zoom_out)

        # Note: Do not add separate Ctrl++/Ctrl+- actions here.
        # QKeySequence.StandardKey.ZoomIn/ZoomOut already maps to these keys on many layouts,
        # and duplicates cause: "QAction::event: Ambiguous shortcut overload".
        grid_zoom_menu.addSeparator()
        
        grid_font_family_menu = grid_zoom_menu.addMenu("Grid Font &Family")
        font_families = ["Calibri", "Segoe UI", "Arial", "Consolas", "Verdana", 
                        "Times New Roman", "Georgia", "Courier New"]
        for font_name in font_families:
            font_action = QAction(font_name, self)
            font_action.triggered.connect(lambda checked, f=font_name: self.set_font_family(f))
            grid_font_family_menu.addAction(font_action)
        
        view_menu.addSeparator()
        
        # Translation Results Pane section
        results_zoom_menu = view_menu.addMenu("📋 Translation &Results Pane")
        
        results_zoom_in_action = QAction("Results Zoom &In", self)
        results_zoom_in_action.setShortcut("Ctrl+Shift+=")
        results_zoom_in_action.triggered.connect(self.results_pane_zoom_in)
        results_zoom_menu.addAction(results_zoom_in_action)
        
        results_zoom_out_action = QAction("Results Zoom &Out", self)
        results_zoom_out_action.setShortcut("Ctrl+Shift+-")
        results_zoom_out_action.triggered.connect(self.results_pane_zoom_out)
        results_zoom_menu.addAction(results_zoom_out_action)
        
        results_zoom_reset_action = QAction("Results Zoom &Reset", self)
        results_zoom_reset_action.triggered.connect(self.results_pane_zoom_reset)
        results_zoom_menu.addAction(results_zoom_reset_action)
        
        results_zoom_menu.addSeparator()
        
        results_note = QAction("(Includes match list + compare boxes)", self)
        results_note.setEnabled(False)
        results_zoom_menu.addAction(results_note)
        
        # Compare Panel section
        compare_zoom_menu = view_menu.addMenu("🔍 &Compare Panel")
        
        compare_zoom_in_action = QAction("Compare Panel Zoom &In", self)
        compare_zoom_in_action.setShortcut("Ctrl+Alt+=")
        compare_zoom_in_action.triggered.connect(self.compare_panel_zoom_in)
        compare_zoom_menu.addAction(compare_zoom_in_action)
        
        compare_zoom_out_action = QAction("Compare Panel Zoom &Out", self)
        compare_zoom_out_action.setShortcut("Ctrl+Alt+-")
        compare_zoom_out_action.triggered.connect(self.compare_panel_zoom_out)
        compare_zoom_menu.addAction(compare_zoom_out_action)
        
        compare_zoom_reset_action = QAction("Compare Panel Zoom &Reset", self)
        compare_zoom_reset_action.triggered.connect(self.compare_panel_zoom_reset)
        compare_zoom_menu.addAction(compare_zoom_reset_action)

        view_menu.addSeparator()

        auto_resize_action = QAction("📐 &Auto-Resize Rows", self)
        auto_resize_action.triggered.connect(self.auto_resize_rows)
        auto_resize_action.setToolTip("Automatically resize all rows to fit content")
        view_menu.addAction(auto_resize_action)

        view_menu.addSeparator()
        
        # Multi-file project progress
        file_progress_action = QAction("📁 &File Progress...", self)
        file_progress_action.triggered.connect(self.show_file_progress_dialog)
        file_progress_action.setToolTip("View translation progress per file (multi-file projects)")
        view_menu.addAction(file_progress_action)
        
        # Proofreading results
        proofread_results_action = QAction("✅ &Proofreading Results...", self)
        proofread_results_action.triggered.connect(self.show_proofreading_results_dialog)
        proofread_results_action.setToolTip("View and manage proofreading issues")
        view_menu.addAction(proofread_results_action)

        view_menu.addSeparator()

        theme_action = QAction("🎨 &Theme Editor...", self)
        theme_action.triggered.connect(self.show_theme_editor)
        view_menu.addAction(theme_action)

        # Tools Menu
        tools_menu = menubar.addMenu("&Tools")
        
        # Tools in same order as Tools tab
        autofingers_action = QAction("✋ &AutoFingers...", self)
        autofingers_action.setShortcut("Ctrl+Shift+A")
        autofingers_action.triggered.connect(self.show_autofingers)
        tools_menu.addAction(autofingers_action)
        
        superconverter_action = QAction("🔄 Super&converter...", self)
        superconverter_action.triggered.connect(lambda: self._navigate_to_tool("Superconverter"))
        tools_menu.addAction(superconverter_action)
        
        pdf_rescue_action = QAction("📄 &PDF Rescue...", self)
        pdf_rescue_action.triggered.connect(lambda: self._navigate_to_tool("PDF Rescue"))
        tools_menu.addAction(pdf_rescue_action)
        
        superbench_action = QAction("📊 Super&bench...", self)
        superbench_action.triggered.connect(lambda: self._navigate_to_tool("Superbench"))
        tools_menu.addAction(superbench_action)
        
        superbrowser_action = QAction("🌐 Super&browser...", self)
        superbrowser_action.triggered.connect(lambda: self._navigate_to_tool("Superbrowser"))
        tools_menu.addAction(superbrowser_action)
        
        supercleaner_action = QAction("🧹 Supercleaner...", self)
        supercleaner_action.triggered.connect(lambda: self._navigate_to_tool("Supercleaner"))
        tools_menu.addAction(supercleaner_action)
        
        superlookup_action = QAction("🔍 Super&lookup (Ctrl+K)...", self)
        # Note: Actual Ctrl+K shortcut handled by QShortcut in setup_global_shortcuts()
        # which calls show_concordance_search() for proper selection capture
        superlookup_action.triggered.connect(self.show_concordance_search)
        tools_menu.addAction(superlookup_action)
        
        supervoice_action = QAction("🎤 Super&voice...", self)
        supervoice_action.triggered.connect(lambda: self._navigate_to_tool("Supervoice"))
        tools_menu.addAction(supervoice_action)
        
        encoding_action = QAction("🔧 &Text Encoding Repair...", self)
        encoding_action.triggered.connect(lambda: self._navigate_to_tool("Text Encoding Repair"))
        tools_menu.addAction(encoding_action)
        
        tmx_editor_action = QAction("✏️ T&MX Editor...", self)
        tmx_editor_action.triggered.connect(lambda: self._navigate_to_tool("TMX Editor"))
        tools_menu.addAction(tmx_editor_action)
        
        tracked_changes_action = QAction("🔄 Tracked &Changes...", self)
        tracked_changes_action.triggered.connect(lambda: self._navigate_to_tool("Tracked Changes"))
        tools_menu.addAction(tracked_changes_action)
        
        tools_menu.addSeparator()
        
        image_extractor_action = QAction("🖼️ &Image Extractor (Superimage)...", self)
        image_extractor_action.triggered.connect(self.show_image_extractor_from_tools)
        image_extractor_action.setToolTip("Extract images from DOCX files")
        tools_menu.addAction(image_extractor_action)
        
        tools_menu.addSeparator()

        settings_action = QAction("&Settings...", self)
        settings_action.triggered.connect(lambda: self._go_to_settings_tab())
        tools_menu.addAction(settings_action)
        
        # Help Menu
        help_menu = menubar.addMenu("&Help")

        # Documentation links (GitHub URLs for universal access)
        # Removed internal manual link — documentation migrated to GitBook

        # Place Supervertaler Help at the top of the Help menu
        superdocs_action = QAction("Supervertaler Help", self)
        superdocs_action.setToolTip("Online documentation (GitBook)")
        superdocs_action.triggered.connect(lambda: self._open_url("https://supervertaler.gitbook.io/superdocs/"))
        help_menu.addAction(superdocs_action)

        help_menu.addSeparator()

        shortcuts_action = QAction("⌨️ Keyboard Shortcuts", self)
        shortcuts_action.triggered.connect(lambda: self._open_url("https://github.com/michaelbeijer/Supervertaler/blob/main/docs/guides/KEYBOARD_SHORTCUTS.md"))
        help_menu.addAction(shortcuts_action)

        changelog_action = QAction("📝 Changelog", self)
        changelog_action.triggered.connect(lambda: self._open_url("https://github.com/michaelbeijer/Supervertaler/blob/main/CHANGELOG.md"))
        help_menu.addAction(changelog_action)

        update_check_action = QAction("🔄 Check for Updates...", self)
        update_check_action.setToolTip("Check whether you are running the latest Supervertaler release")
        update_check_action.triggered.connect(self.check_for_updates)
        help_menu.addAction(update_check_action)

        copy_version_info_action = QAction("📋 Copy Version Info", self)
        copy_version_info_action.setToolTip("Copy version and system info to clipboard (useful for support)")
        copy_version_info_action.triggered.connect(self.copy_version_info_to_clipboard)
        help_menu.addAction(copy_version_info_action)

        help_menu.addSeparator()

        github_action = QAction("🔗 GitHub Repository", self)
        github_action.triggered.connect(lambda: self._open_url("https://github.com/michaelbeijer/Supervertaler"))
        help_menu.addAction(github_action)

        help_menu.addSeparator()
        
        # AutoHotkey setup (Windows only)
        if os.name == 'nt':
            ahk_setup_action = QAction("⌨️ Setup AutoHotkey (Global Hotkey)", self)
            ahk_setup_action.setToolTip("Configure AutoHotkey for Superlookup global hotkey (Ctrl+Alt+L)")
            ahk_setup_action.triggered.connect(self._show_ahk_setup_from_menu)
            help_menu.addAction(ahk_setup_action)
            help_menu.addSeparator()

        about_action = QAction("ℹ️ About", self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)
    
    def record_undo_state(self, segment_id, old_target, new_target, old_status, new_status):
        """Record an undo state when grid cells are edited"""
        # Don't record if nothing actually changed
        if old_target == new_target and old_status == new_status:
            return
        
        # Add to undo stack
        undo_entry = {
            "segment_id": segment_id,
            "old_target": old_target,
            "new_target": new_target,
            "old_status": old_status,
            "new_status": new_status
        }
        self.undo_stack.append(undo_entry)
        
        # Trim undo stack to max levels
        if len(self.undo_stack) > self.max_undo_levels:
            self.undo_stack.pop(0)
        
        # Clear redo stack (can't redo after new edit)
        self.redo_stack.clear()
        
        # Update menu actions
        self.update_undo_redo_actions()
    
    def undo_action_handler(self):
        """Handle Undo (Ctrl+Z) action"""
        if not self.undo_stack:
            return
        
        # Pop last action from undo stack
        action = self.undo_stack.pop()
        
        # Find the segment
        segment_id = action["segment_id"]
        segment = None
        for seg in self.current_project.segments:
            if seg.segment_id == segment_id:
                segment = seg
                break
        
        if not segment:
            return
        
        # Revert to old values
        segment.target_text = action["old_target"]
        segment.status = action["old_status"]
        
        # Update grid display
        row = self.find_grid_row_by_segment_id(segment_id)
        if row is not None:
            # Update target text cell
            target_item = self.grid.item(row, 2)
            if target_item:
                target_item.setText(action["old_target"])
            
            # Update status cell
            status_item = self.grid.item(row, 3)
            if status_item:
                status_item.setText(action["old_status"])
        
        # Move action to redo stack
        self.redo_stack.append(action)
        
        # Update menu actions
        self.update_undo_redo_actions()
    
    def redo_action_handler(self):
        """Handle Redo (Ctrl+Shift+Z / Ctrl+Y) action"""
        if not self.redo_stack:
            return
        
        # Pop last action from redo stack
        action = self.redo_stack.pop()
        
        # Find the segment
        segment_id = action["segment_id"]
        segment = None
        for seg in self.current_project.segments:
            if seg.segment_id == segment_id:
                segment = seg
                break
        
        if not segment:
            return
        
        # Reapply new values
        segment.target_text = action["new_target"]
        segment.status = action["new_status"]
        
        # Update grid display
        row = self.find_grid_row_by_segment_id(segment_id)
        if row is not None:
            # Update target text cell
            target_item = self.grid.item(row, 2)
            if target_item:
                target_item.setText(action["new_target"])
            
            # Update status cell
            status_item = self.grid.item(row, 3)
            if status_item:
                status_item.setText(action["new_status"])
        
        # Move action back to undo stack
        self.undo_stack.append(action)
        
        # Update menu actions
        self.update_undo_redo_actions()
    
    def update_undo_redo_actions(self):
        """Update enabled/disabled state of undo/redo menu actions"""
        self.undo_action.setEnabled(len(self.undo_stack) > 0)
        self.redo_action.setEnabled(len(self.redo_stack) > 0)
    
    def find_grid_row_by_segment_id(self, segment_id):
        """Find the grid row index for a given segment ID"""
        for row in range(self.grid.rowCount()):
            id_item = self.grid.item(row, 0)
            if id_item and id_item.text() == str(segment_id):
                return row
        return None
    
    def create_quick_access_toolbar(self):
        """Create Quick Access Toolbar above ribbon"""
        from PyQt6.QtWidgets import QToolBar, QWidget, QHBoxLayout
        from PyQt6.QtCore import QSize
        
        qat = QToolBar("Quick Access")
        qat.setMovable(False)
        qat.setFloatable(False)
        qat.setIconSize(QSize(20, 20))
        qat.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonIconOnly)
        qat.setMaximumHeight(28)
        qat.setMinimumWidth(170)  # Fixed width to align ribbon tabs
        qat.setMaximumWidth(170)
        
        # Styling for compact appearance
        qat.setStyleSheet("""
            QToolBar {
                background: transparent;
                border: none;
                spacing: 2px;
                padding: 2px;
            }
            QToolButton {
                background: transparent;
                border: 1px solid transparent;
                border-radius: 3px;
                padding: 2px;
                margin: 0px;
                font-size: 14pt;
            }
            QToolButton:hover {
                background: rgba(255, 255, 255, 0.1);
                border: 1px solid rgba(255, 255, 255, 0.2);
            }
            QToolButton:pressed {
                background: rgba(0, 0, 0, 0.1);
            }
        """)
        
        # Empty toolbar - all buttons removed, minimize moved to ribbon
        # Add empty separator just to maintain spacing
        qat.addSeparator()
        
        return qat
    
    def create_ribbon(self):
        """
        Create modern ribbon interface - DISABLED
        Ribbon functionality has been moved to traditional menu bar.
        This code is kept for potential future use.
        """
        # Ribbon removed - all actions are now in the menu bar
        # Keeping this method for potential future re-enablement
        return
        
        # from modules.ribbon_widget import RibbonWidget, RibbonTab, RibbonGroup, RibbonButton
        # 
        # # Create ribbon widget
        # self.ribbon = RibbonWidget(self)
        # 
        # # Connect ribbon actions to methods
        # self.ribbon.action_triggered.connect(self.handle_ribbon_action)
        
        # Ribbon code commented out - all functionality moved to menu bar
        # HOME TAB - Project management and navigation
        # home_tab = RibbonTab()
        # project_group = RibbonGroup("Project")
        # project_group.add_button(self.ribbon.create_button("New", "📄", "new", "New project (Ctrl+N)"))
        # ... (all ribbon tab creation code commented out)
        # All actions are now available in the traditional menu bar
    
    def create_ribbon_toolbar(self):
        """Create a toolbar to hold the ribbon - DISABLED"""
        # Ribbon removed - using traditional menu bar instead
        # This method kept for potential future use
        return None
    
    def toggle_ribbon_minimized(self, minimized: bool):
        """Toggle ribbon between full and minimized (tabs-only) mode - DISABLED"""
        # Ribbon removed - no action needed
        self.ribbon_minimized = minimized  # Store state for backwards compatibility
        pass
    
    def show_ribbon_temporarily(self, index: int):
        """Show ribbon temporarily when tab clicked in minimized mode - DISABLED"""
        # Ribbon removed - no action needed
        pass
    
    def on_main_tab_changed(self, index: int):
        """Handle main tab change"""
        # Ribbon removed - menu bar remains constant across all tabs
        # No context switching needed
        pass
    
    def create_main_layout(self):
        """Create main application layout with all-tab interface"""
        # Central widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # Main layout
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(5, 0, 5, 5)  # Reduced top margin to minimize space below menu bar
        
        # ===== SIMPLIFIED TAB-BASED UI =====
        # Single tab widget with all functionality
        from modules.unified_prompt_manager_qt import UnifiedPromptManagerQt
        
        # Create main tab widget
        self.main_tabs = QTabWidget()
        self.main_tabs.tabBar().setFocusPolicy(Qt.FocusPolicy.NoFocus)
        self.main_tabs.tabBar().setDrawBase(False)
        self.main_tabs.setStyleSheet("""
            QTabBar::tab { padding: 8px 15px; outline: 0; }
            QTabBar::tab:focus { outline: none; }
            QTabBar::tab:selected { 
                border-bottom: 1px solid #2196F3; 
                background-color: rgba(33, 150, 243, 0.08);
            }
        """)
        
        # ===== 1. GRID TAB =====
        # Contains the translation grid
        grid_widget = self.create_grid_view_widget_for_home()
        self.main_tabs.addTab(grid_widget, "📝 Grid")
        
        # ===== 2. PROJECT RESOURCES TAB =====
        # Contains TM, Termbases, Non-Translatables
        resources_tab = self.create_resources_tab()
        self.main_tabs.addTab(resources_tab, "🗂️ Resources")
        
        # ===== 3. PROMPT MANAGER TAB =====
        # Unified Prompt Library + AI Assistant
        from modules.unified_prompt_manager_qt import UnifiedPromptManagerQt
        prompt_widget = QWidget()
        self.prompt_manager_qt = UnifiedPromptManagerQt(self, standalone=False)
        self.prompt_manager_qt.create_tab(prompt_widget)
        self.main_tabs.addTab(prompt_widget, "🤖 Prompt Manager")
        
        # Keep backward compatibility reference
        self.document_views_widget = self.main_tabs
        
        # 4. TOOLS
        tools_tab = self.create_specialised_tools_tab()
        self.main_tabs.addTab(tools_tab, "🛠️ Tools")

        # 5. SETTINGS
        settings_tab = self.create_settings_tab()
        self.main_tabs.addTab(settings_tab, "⚙️ Settings")
        
        # Set startup tab to Grid (index 0)
        self.main_tabs.setCurrentIndex(0)
        
        main_layout.addWidget(self.main_tabs)
        
        # Connect tab changes to handle view refreshes
        self.main_tabs.currentChanged.connect(self._on_main_tab_changed)

    
    def _create_placeholder_tab(self, title: str, description: str) -> QWidget:
        """Create a simple placeholder tab"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(10, 10, 10, 10)
        
        header = QLabel(title)
        header.setStyleSheet("font-size: 14px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(header)
        
        placeholder = QLabel(description)
        placeholder.setAlignment(Qt.AlignmentFlag.AlignCenter)
        placeholder.setStyleSheet("color: #888; font-size: 12px;")
        layout.addWidget(placeholder, stretch=1)
        
        return tab
    
    def create_non_translatables_tab(self) -> QWidget:
        """Create the Non-Translatables tab - manage non-translatable content"""
        from modules.non_translatables_manager import NonTranslatablesManager, NonTranslatableList
        
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(10, 10, 10, 10)
        
        # Header
        header = QLabel("🚫 Non-Translatables")
        header.setStyleSheet("font-size: 14px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(header)
        
        # Description
        desc = QLabel("Manage terms and phrases that should not be translated (brand names, product names, technical identifiers, etc.).\n"
                      "Non-translatables are highlighted in pastel yellow in the source text and Translation Results panel.")
        desc.setWordWrap(True)
        desc.setStyleSheet("color: #666; font-size: 11px; margin-bottom: 10px;")
        layout.addWidget(desc)
        
        # Initialize NT manager
        nt_path = self.user_data_path / "Translation_Resources" / "Non-translatables"
        nt_path.mkdir(parents=True, exist_ok=True)
        self.nt_manager = NonTranslatablesManager(str(nt_path), self.log)
        
        # Try to load existing lists or convert from plain text
        existing_nt_files = list(nt_path.glob("*.svntl")) + list(nt_path.glob("*.ntl"))
        if not existing_nt_files:
            # Check for plain text file to convert
            txt_file = nt_path / "non-translatables.txt"
            if txt_file.exists():
                self.log("📄 Found plain text NT file, converting to .svntl format...")
                nt_list = self.nt_manager.load_from_plain_text(str(txt_file), "Default Non-Translatables")
                if nt_list:
                    self.nt_manager.save_list(nt_list)
                    self.nt_manager.lists[nt_list.name] = nt_list
                    self.nt_manager.active_lists.append(nt_list.name)
        else:
            self.nt_manager.load_all_lists()
        
        # Toolbar
        toolbar = QHBoxLayout()
        
        # List selection combo
        list_combo = QComboBox()
        list_combo.setMinimumWidth(200)
        list_combo.setToolTip("Select a non-translatable list")
        toolbar.addWidget(QLabel("List:"))
        toolbar.addWidget(list_combo)
        
        # Buttons
        new_btn = QPushButton("➕ New List")
        new_btn.setToolTip("Create a new non-translatable list")
        toolbar.addWidget(new_btn)
        
        import_btn = QPushButton("📥 Import")
        import_btn.setToolTip("Import from file (.svntl, .ntl, .txt, or memoQ .mqres)")
        toolbar.addWidget(import_btn)
        
        export_btn = QPushButton("📤 Export")
        export_btn.setToolTip("Export selected list")
        toolbar.addWidget(export_btn)
        
        delete_btn = QPushButton("🗑️ Delete")
        delete_btn.setToolTip("Delete selected list")
        toolbar.addWidget(delete_btn)
        
        toolbar.addStretch()
        
        # Active checkbox
        active_checkbox = CheckmarkCheckBox("Active")
        active_checkbox.setToolTip("Toggle whether this list is active for matching")
        toolbar.addWidget(active_checkbox)
        
        layout.addLayout(toolbar)
        
        # Split view: entry list on left, add/edit on right
        splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # Left: Entry list
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(0, 0, 0, 0)
        
        # Search box
        search_box = QLineEdit()
        search_box.setPlaceholderText("Search entries...")
        left_layout.addWidget(search_box)
        
        # Entry table
        entry_table = QTableWidget()
        entry_table.setColumnCount(2)
        entry_table.setHorizontalHeaderLabels(["Entry", "Category"])
        entry_table.horizontalHeader().setStretchLastSection(True)
        entry_table.setColumnWidth(0, 250)
        entry_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        entry_table.setAlternatingRowColors(True)
        entry_table.setSortingEnabled(True)  # Enable column sorting
        entry_table.horizontalHeader().setSortIndicatorShown(True)
        entry_table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)  # Enable right-click menu
        left_layout.addWidget(entry_table)
        
        # Entry count
        count_label = QLabel("0 entries")
        count_label.setStyleSheet("color: #666; font-size: 10px;")
        left_layout.addWidget(count_label)
        
        splitter.addWidget(left_widget)
        
        # Right: Add/Edit panel
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(10, 0, 0, 0)
        
        right_layout.addWidget(QLabel("<b>Add/Edit Entry</b>"))
        
        entry_input = QLineEdit()
        entry_input.setPlaceholderText("Enter non-translatable term...")
        right_layout.addWidget(entry_input)
        
        category_input = QLineEdit()
        category_input.setPlaceholderText("Category (optional)")
        right_layout.addWidget(category_input)
        
        notes_input = QTextEdit()
        notes_input.setPlaceholderText("Notes (optional)")
        notes_input.setMaximumHeight(80)
        right_layout.addWidget(notes_input)
        
        add_btn = QPushButton("➕ Add Entry")
        right_layout.addWidget(add_btn)
        
        remove_btn = QPushButton("🗑️ Remove Selected")
        right_layout.addWidget(remove_btn)
        
        right_layout.addStretch()
        
        # Bulk operations
        bulk_group = QGroupBox("Bulk Import")
        bulk_layout = QVBoxLayout(bulk_group)
        
        bulk_text = QTextEdit()
        bulk_text.setPlaceholderText("Paste multiple entries (one per line) for bulk import...")
        bulk_text.setMaximumHeight(100)
        bulk_layout.addWidget(bulk_text)
        
        bulk_btn = QPushButton("Import All")
        bulk_layout.addWidget(bulk_btn)
        
        right_layout.addWidget(bulk_group)
        
        splitter.addWidget(right_widget)
        splitter.setSizes([500, 300])
        
        layout.addWidget(splitter, stretch=1)
        
        # Store references for callbacks
        self.nt_list_combo = list_combo
        self.nt_entry_table = entry_table
        self.nt_count_label = count_label
        self.nt_active_checkbox = active_checkbox
        
        # Populate list combo
        def refresh_list_combo():
            list_combo.clear()
            for name, nt_list in self.nt_manager.lists.items():
                active_marker = "✓ " if nt_list.is_active else ""
                list_combo.addItem(f"{active_marker}{name}", name)
            
            if list_combo.count() == 0:
                list_combo.addItem("(No lists - create or import one)", "")
        
        # Populate entry table for selected list
        def refresh_entry_table():
            # Disable sorting while populating to avoid performance issues
            entry_table.setSortingEnabled(False)
            entry_table.setRowCount(0)
            current_name = list_combo.currentData()
            
            if not current_name or current_name not in self.nt_manager.lists:
                count_label.setText("0 entries")
                active_checkbox.setChecked(False)
                entry_table.setSortingEnabled(True)
                return
            
            nt_list = self.nt_manager.lists[current_name]
            active_checkbox.blockSignals(True)
            active_checkbox.setChecked(nt_list.is_active)
            active_checkbox.blockSignals(False)
            
            search_term = search_box.text().lower()
            
            filtered = [e for e in nt_list.entries 
                       if not search_term or search_term in e.text.lower()]
            
            entry_table.setRowCount(len(filtered))
            for row, entry in enumerate(filtered):
                entry_table.setItem(row, 0, QTableWidgetItem(entry.text))
                entry_table.setItem(row, 1, QTableWidgetItem(entry.category))
            
            # Re-enable sorting after populating
            entry_table.setSortingEnabled(True)
            
            count_label.setText(f"{len(nt_list.entries)} entries ({len(filtered)} shown)")
        
        # Connect signals
        list_combo.currentIndexChanged.connect(refresh_entry_table)
        search_box.textChanged.connect(refresh_entry_table)
        
        # Helper to save NT settings to project
        def save_nt_settings_to_project():
            """Save active NT list names to project settings"""
            if hasattr(self, 'current_project') and self.current_project:
                active_list_names = [name for name, lst in self.nt_manager.lists.items() if lst.is_active]
                if not hasattr(self.current_project, 'nt_settings'):
                    self.current_project.nt_settings = {}
                self.current_project.nt_settings['active_lists'] = active_list_names
                self.log(f"💾 Saved NT settings to project: {len(active_list_names)} active list(s)")
        
        # Active toggle
        def on_active_toggle(checked):
            current_name = list_combo.currentData()
            if current_name and current_name in self.nt_manager.lists:
                self.nt_manager.set_list_active(current_name, checked)
                refresh_list_combo()
                # Keep selection
                for i in range(list_combo.count()):
                    if list_combo.itemData(i) == current_name:
                        list_combo.setCurrentIndex(i)
                        break
                # Save to project settings
                save_nt_settings_to_project()
        
        active_checkbox.toggled.connect(on_active_toggle)
        
        # New list
        def on_new_list():
            name, ok = QInputDialog.getText(tab, "New Non-Translatable List", "List name:")
            if ok and name:
                if name in self.nt_manager.lists:
                    QMessageBox.warning(tab, "Error", f"List '{name}' already exists.")
                    return
                nt_list = self.nt_manager.create_list(name)
                self.nt_manager.save_list(nt_list)
                refresh_list_combo()
                # Select new list
                for i in range(list_combo.count()):
                    if list_combo.itemData(i) == name:
                        list_combo.setCurrentIndex(i)
                        break
        
        new_btn.clicked.connect(on_new_list)
        
        # Add entry
        def on_add_entry():
            text = entry_input.text().strip()
            if not text:
                return
            
            current_name = list_combo.currentData()
            if not current_name:
                QMessageBox.warning(tab, "Error", "Please select or create a list first.")
                return
            
            self.nt_manager.add_entry(current_name, text, 
                                      notes_input.toPlainText(), 
                                      category_input.text())
            self.nt_manager.save_list(self.nt_manager.lists[current_name])
            
            entry_input.clear()
            notes_input.clear()
            category_input.clear()
            refresh_entry_table()
        
        add_btn.clicked.connect(on_add_entry)
        entry_input.returnPressed.connect(on_add_entry)
        
        # Remove entry
        def on_remove_entry():
            current_name = list_combo.currentData()
            if not current_name:
                return
            
            selected = entry_table.selectedItems()
            if not selected:
                return
            
            # Get unique row indices
            rows = set(item.row() for item in selected)
            
            # Get entry texts to remove
            for row in sorted(rows, reverse=True):
                text_item = entry_table.item(row, 0)
                if text_item:
                    self.nt_manager.remove_entry(current_name, text_item.text())
            
            self.nt_manager.save_list(self.nt_manager.lists[current_name])
            refresh_entry_table()
        
        remove_btn.clicked.connect(on_remove_entry)
        
        # Right-click context menu for entry table
        def on_entry_context_menu(pos):
            selected = entry_table.selectedItems()
            if not selected:
                return
            
            menu = QMenu(entry_table)
            
            # Get selected row count
            rows = set(item.row() for item in selected)
            count_text = f"Delete {len(rows)} entr{'ies' if len(rows) > 1 else 'y'}"
            
            delete_action = QAction(f"🗑️ {count_text}", menu)
            delete_action.triggered.connect(on_remove_entry)
            menu.addAction(delete_action)
            
            menu.exec(entry_table.viewport().mapToGlobal(pos))
        
        entry_table.customContextMenuRequested.connect(on_entry_context_menu)
        
        # Delete key shortcut for entry table
        def on_entry_key_press(event):
            if event.key() == Qt.Key.Key_Delete:
                on_remove_entry()
            else:
                QTableWidget.keyPressEvent(entry_table, event)
        
        entry_table.keyPressEvent = on_entry_key_press
        
        # Bulk import
        def on_bulk_import():
            current_name = list_combo.currentData()
            if not current_name:
                QMessageBox.warning(tab, "Error", "Please select or create a list first.")
                return
            
            text = bulk_text.toPlainText()
            lines = [l.strip() for l in text.splitlines() if l.strip()]
            
            if not lines:
                return
            
            added = 0
            for line in lines:
                if line not in [e.text for e in self.nt_manager.lists[current_name].entries]:
                    self.nt_manager.add_entry(current_name, line)
                    added += 1
            
            self.nt_manager.save_list(self.nt_manager.lists[current_name])
            bulk_text.clear()
            refresh_entry_table()
            
            QMessageBox.information(tab, "Bulk Import", f"Added {added} entries ({len(lines) - added} duplicates skipped)")
        
        bulk_btn.clicked.connect(on_bulk_import)
        
        # Import from file
        def on_import():
            filepath, _ = QFileDialog.getOpenFileName(
                tab, "Import Non-Translatables",
                str(nt_path),
                "All Supported (*.svntl *.ntl *.txt *.mqres);;Supervertaler NT List (*.svntl *.ntl);;Plain Text (*.txt);;memoQ Non-Translatables (*.mqres)"
            )
            
            if not filepath:
                return
            
            filepath_lower = filepath.lower()
            
            # Load based on file type
            if filepath_lower.endswith('.mqres'):
                imported_list = self.nt_manager.import_memoq_mqres(filepath)
            elif filepath_lower.endswith('.svntl') or filepath_lower.endswith('.ntl'):
                imported_list = self.nt_manager.load_list(filepath)
            else:
                imported_list = self.nt_manager.load_from_plain_text(filepath)
            
            if not imported_list:
                QMessageBox.warning(tab, "Import Error", "Failed to import file. Check the log for details.")
                return
            
            # Ask user whether to create new list or merge
            current_name = list_combo.currentData()
            if current_name and current_name in self.nt_manager.lists:
                # Create custom dialog with clear button labels
                msg_box = QMessageBox(tab)
                msg_box.setWindowTitle("Import Options")
                msg_box.setText(f"Import {len(imported_list.entries)} entries")
                msg_box.setInformativeText(f"Choose how to import into your non-translatables:")
                msg_box.setIcon(QMessageBox.Icon.Question)
                
                create_btn = msg_box.addButton(f"Create New List", QMessageBox.ButtonRole.YesRole)
                merge_btn = msg_box.addButton(f"Merge into '{current_name}'", QMessageBox.ButtonRole.NoRole)
                cancel_btn = msg_box.addButton(QMessageBox.StandardButton.Cancel)
                
                msg_box.exec()
                clicked = msg_box.clickedButton()
                
                if clicked == cancel_btn:
                    return
                elif clicked == merge_btn:
                    # Merge into current list
                    added, skipped = self.nt_manager.merge_into_list(current_name, imported_list)
                    self.nt_manager.save_list(self.nt_manager.lists[current_name])
                    refresh_entry_table()
                    QMessageBox.information(tab, "Import Complete", 
                                          f"Merged {added} entries into '{current_name}'\n({skipped} duplicates skipped)")
                    return
            
            # Create new list
            # Check for name collision
            new_name = imported_list.name
            counter = 1
            while new_name in self.nt_manager.lists:
                new_name = f"{imported_list.name} ({counter})"
                counter += 1
            
            imported_list.name = new_name
            self.nt_manager.lists[new_name] = imported_list
            self.nt_manager.active_lists.append(new_name)
            self.nt_manager.save_list(imported_list)
            
            refresh_list_combo()
            
            # Select imported list
            for i in range(list_combo.count()):
                if list_combo.itemData(i) == new_name:
                    list_combo.setCurrentIndex(i)
                    break
            
            QMessageBox.information(tab, "Import Complete", 
                                  f"Created list '{new_name}' with {len(imported_list.entries)} entries")
        
        import_btn.clicked.connect(on_import)
        
        # Export
        def on_export():
            current_name = list_combo.currentData()
            if not current_name or current_name not in self.nt_manager.lists:
                QMessageBox.warning(tab, "Error", "Please select a list to export.")
                return
            
            filepath, selected_filter = QFileDialog.getSaveFileName(
                tab, "Export Non-Translatables",
                str(nt_path / f"{current_name}.svntl"),
                "Supervertaler NT List (*.svntl);;Plain Text (*.txt)"
            )
            
            if not filepath:
                return
            
            if "Plain Text" in selected_filter or filepath.lower().endswith('.txt'):
                success = self.nt_manager.export_to_plain_text(current_name, filepath)
            else:
                success = self.nt_manager.export_list(current_name, filepath)
            
            if success:
                QMessageBox.information(tab, "Export Complete", f"Exported to:\n{filepath}")
            else:
                QMessageBox.warning(tab, "Export Error", "Failed to export. Check the log for details.")
        
        export_btn.clicked.connect(on_export)
        
        # Delete list
        def on_delete():
            current_name = list_combo.currentData()
            if not current_name or current_name not in self.nt_manager.lists:
                return
            
            confirm = QMessageBox.question(
                tab, "Delete List",
                f"Are you sure you want to delete '{current_name}'?\n\nThis cannot be undone.",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if confirm == QMessageBox.StandardButton.Yes:
                self.nt_manager.delete_list(current_name)
                refresh_list_combo()
                refresh_entry_table()
        
        delete_btn.clicked.connect(on_delete)
        
        # Initial population
        refresh_list_combo()
        if list_combo.count() > 0:
            refresh_entry_table()
        
        return tab
    
    def create_prompt_manager_tab(self) -> QWidget:
        """Create the Unified Prompt Library tab - Simplified 2-Layer Architecture"""
        from modules.unified_prompt_manager_qt import UnifiedPromptManagerQt
        
        # Create Unified Prompt Manager widget (embedded mode, not standalone)
        prompt_widget = QWidget()
        self.prompt_manager_qt = UnifiedPromptManagerQt(self, standalone=False)
        self.prompt_manager_qt.create_tab(prompt_widget)
        
        return prompt_widget
    
    
    def create_superconverter_tab(self) -> QWidget:
        """Create the Superconverter tab - Format conversion tools"""
        container = QWidget()
        main_layout = QVBoxLayout(container)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(10)
        
        # Header
        header = QLabel("🔄 Superconverter")
        header.setStyleSheet("font-size: 16pt; font-weight: bold; color: #1976D2;")
        main_layout.addWidget(header)
        
        # Description
        description = QLabel(
            "Convert translation data between different formats - perfect for AI workflows, CAT tool exchanges, and data processing."
        )
        description.setWordWrap(True)
        description.setStyleSheet("color: #666; padding: 5px; background-color: #E3F2FD; border-radius: 3px;")
        main_layout.addWidget(description)
        
        main_layout.addSpacing(5)
        
        # Create tabbed interface for different conversion types
        tabs = QTabWidget()
        tabs.setStyleSheet("QTabBar::tab { outline: 0; } QTabBar::tab:focus { outline: none; }")
        
        # Tab 1: Bilingual Export (current project)
        bilingual_tab = QWidget()
        bilingual_layout = QVBoxLayout(bilingual_tab)
        bilingual_layout.setContentsMargins(15, 15, 15, 15)
        
        bilingual_info = QLabel(
            "<b>Export Current Project</b><br><br>"
            "Export your translation project as a Markdown table - perfect for AI chat interfaces like ChatGPT and Claude.<br><br>"
            "The table format renders beautifully and makes it easy for AI to understand and process your segments."
        )
        bilingual_info.setWordWrap(True)
        bilingual_info.setTextFormat(Qt.TextFormat.RichText)
        bilingual_layout.addWidget(bilingual_info)
        
        bilingual_layout.addSpacing(15)
        
        bilingual_btn = QPushButton("📄 Export as Markdown Table")
        bilingual_btn.setMinimumHeight(40)
        bilingual_btn.setStyleSheet(
            "QPushButton { background-color: #2196F3; color: white; font-weight: bold; "
            "border: none; border-radius: 5px; padding: 10px; outline: none; }"
            "QPushButton:hover { background-color: #1976D2; }"
        )
        bilingual_btn.clicked.connect(self.export_bilingual_table_markdown)
        bilingual_layout.addWidget(bilingual_btn)
        
        bilingual_layout.addStretch()
        tabs.addTab(bilingual_tab, "📊 Bilingual Table")
        
        # Tab 2: Document Converter (monolingual docs to Markdown)
        doc_tab = QWidget()
        doc_layout = QVBoxLayout(doc_tab)
        doc_layout.setContentsMargins(15, 15, 15, 15)
        
        doc_info = QLabel(
            "<b>Convert Documents to Markdown</b><br><br>"
            "Convert DOCX or TXT documents to Markdown format, preserving structure (headings, lists, paragraphs).<br><br>"
            "Perfect for preparing documents for AI processing or publishing on web platforms."
        )
        doc_info.setWordWrap(True)
        doc_info.setTextFormat(Qt.TextFormat.RichText)
        doc_layout.addWidget(doc_info)
        
        doc_layout.addSpacing(15)
        
        # Single file button
        single_doc_btn = QPushButton("📄 Convert Single Document")
        single_doc_btn.setMinimumHeight(40)
        single_doc_btn.setStyleSheet(
            "QPushButton { background-color: #4CAF50; color: white; font-weight: bold; "
            "border: none; border-radius: 5px; padding: 10px; outline: none; }"
            "QPushButton:hover { background-color: #388E3C; }"
        )
        single_doc_btn.clicked.connect(self.convert_document_to_markdown)
        doc_layout.addWidget(single_doc_btn)
        
        doc_layout.addSpacing(10)
        
        # Batch conversion button
        batch_doc_btn = QPushButton("📁 Batch Convert Multiple Documents")
        batch_doc_btn.setMinimumHeight(40)
        batch_doc_btn.setStyleSheet(
            "QPushButton { background-color: #FF9800; color: white; font-weight: bold; "
            "border: none; border-radius: 5px; padding: 10px; outline: none; }"
            "QPushButton:hover { background-color: #F57C00; }"
        )
        batch_doc_btn.clicked.connect(self.batch_convert_documents_to_markdown)
        doc_layout.addWidget(batch_doc_btn)
        
        doc_layout.addStretch()
        tabs.addTab(doc_tab, "📝 Document → Markdown")
        
        # Tab 3: TMX Tools (placeholder for future)
        tmx_tab = QWidget()
        tmx_layout = QVBoxLayout(tmx_tab)
        tmx_layout.setContentsMargins(15, 15, 15, 15)
        
        tmx_info = QLabel(
            "<b>TMX Conversion Tools</b><br><br>"
            "Convert Translation Memory eXchange (TMX) files to and from tab-delimited format.<br><br>"
            "<i>Coming soon...</i>"
        )
        tmx_info.setWordWrap(True)
        tmx_info.setTextFormat(Qt.TextFormat.RichText)
        tmx_info.setStyleSheet("color: #888;")
        tmx_layout.addWidget(tmx_info)
        
        tmx_layout.addSpacing(15)
        
        tmx_to_tsv_btn = QPushButton("TMX → Tab-delimited")
        tmx_to_tsv_btn.setMinimumHeight(40)
        tmx_to_tsv_btn.setEnabled(False)
        tmx_to_tsv_btn.setStyleSheet("QPushButton:disabled { color: #999; background-color: #E0E0E0; }")
        tmx_layout.addWidget(tmx_to_tsv_btn)
        
        tmx_layout.addSpacing(10)
        
        tsv_to_tmx_btn = QPushButton("Tab-delimited → TMX")
        tsv_to_tmx_btn.setMinimumHeight(40)
        tsv_to_tmx_btn.setEnabled(False)
        tsv_to_tmx_btn.setStyleSheet("QPushButton:disabled { color: #999; background-color: #E0E0E0; }")
        tmx_layout.addWidget(tsv_to_tmx_btn)
        
        tmx_layout.addStretch()
        tabs.addTab(tmx_tab, "🔄 TMX Tools")
        
        main_layout.addWidget(tabs)
        
        return container
    
    def create_supercleaner_tab(self) -> QWidget:
        """Create the Supercleaner tab - Clean DOCX documents"""
        from modules.supercleaner_ui import SupercleanerUI
        
        # Create container widget with header
        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(5)
        
        # Header (matches TMX Editor / AutoFingers / PDF Rescue style)
        header = QLabel("🧹 Supercleaner")
        header.setStyleSheet("font-size: 16pt; font-weight: bold; color: #1976D2;")
        layout.addWidget(header, 0)
        
        # Description box (matches other tools style)
        description = QLabel(
            "Clean DOCX documents before translation - removes formatting issues, excessive tags, and OCR artifacts.\n"
            "Inspired by TransTools Document Cleaner, Unbreaker, and CodeZapper."
        )
        description.setWordWrap(True)
        description.setStyleSheet("color: #666; padding: 5px; background-color: #E3F2FD; border-radius: 3px;")
        layout.addWidget(description, 0)
        
        # Create Supercleaner UI widget
        supercleaner = SupercleanerUI(parent=self)
        layout.addWidget(supercleaner, 1)  # 1 = stretch factor
        
        # Store reference
        self.supercleaner_embedded = supercleaner
        
        return container
    
    def create_tmx_editor_tab(self) -> QWidget:
        """Create the TMX Editor tab - Edit TMs"""
        from modules.tmx_editor_qt import TmxEditorUIQt
        
        # Create container widget with Superlookup style header
        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(5)  # Reduced from 10 to 5 for tighter spacing
        
        # Header (matches Superlookup / AutoFingers / PDF Rescue style)
        header = QLabel("📝 TMX Editor")
        header.setStyleSheet("font-size: 16pt; font-weight: bold; color: #1976D2;")
        layout.addWidget(header, 0)  # 0 = no stretch, stays compact
        
        # Description box (matches Superlookup / AutoFingers / PDF Rescue style)
        description = QLabel(
            "Edit translation memory files directly - inspired by Heartsome TMX Editor.\n"
            "Open, edit, filter, and manage your TMX translation memories."
        )
        description.setWordWrap(True)
        description.setStyleSheet("color: #666; padding: 5px; background-color: #E3F2FD; border-radius: 3px;")
        layout.addWidget(description, 0)  # 0 = no stretch, stays compact
        
        # Create TMX Editor widget (embedded mode) - pass database manager
        tmx_editor = TmxEditorUIQt(parent=None, standalone=False, db_manager=self.db_manager)
        layout.addWidget(tmx_editor, 1)  # 1 = stretch factor, expands to fill space
        
        # Store reference for potential future use
        self.tmx_editor_embedded = tmx_editor
        
        return container
    
    def create_reference_images_tab(self) -> QWidget:
        """Create the Image Context tab - Load images as visual context for AI translation"""
        from modules.image_extractor import ImageExtractor
        
        tab = QWidget()
        main_layout = QVBoxLayout(tab)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(5)
        
        # === IMAGE CONTEXT SECTION (TOP) ===
        context_group = QGroupBox("🎯 Image Context - Load Images for AI Translation")
        context_layout = QVBoxLayout()
        
        # Description
        context_desc = QLabel(
            "Load figure images to automatically include with AI translations when text references them (e.g., 'Figure 1', 'see fig 2A').\n"
            "The AI will 'see' the images and better translate technical descriptions and part references."
        )
        context_desc.setWordWrap(True)
        context_desc.setStyleSheet("font-size: 10px; padding: 5px; border-radius: 3px;")
        context_layout.addWidget(context_desc)
        
        # Context controls row
        context_controls = QHBoxLayout()
        
        load_context_btn = QPushButton("📁 Load Images Folder")
        load_context_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                font-weight: bold;
                padding: 8px 16px;
                border-radius: 3px;
                border: none;
                outline: none;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QPushButton:focus {
                outline: none;
                border: none;
            }
        """)
        load_context_btn.clicked.connect(self._on_load_image_context_folder)
        context_controls.addWidget(load_context_btn)
        
        clear_context_btn = QPushButton("🗑️ Clear")
        clear_context_btn.clicked.connect(self._on_clear_image_context)
        clear_context_btn.setMaximumWidth(80)
        context_controls.addWidget(clear_context_btn)
        
        # Status label
        self.image_context_status_label = QLabel("No images loaded")
        self.image_context_status_label.setStyleSheet("color: #999; font-size: 11px; padding: 5px;")
        context_controls.addWidget(self.image_context_status_label, 1)
        
        context_layout.addLayout(context_controls)
        context_group.setLayout(context_layout)
        main_layout.addWidget(context_group)
        
        # === IMAGE EXTRACTOR SECTION (BOTTOM) ===
        extractor_group = QGroupBox("🛠️ Image Extractor - Extract Images from DOCX Files")
        extractor_layout = QVBoxLayout()
        
        # Compact header with title and extract button in one row
        header_layout = QHBoxLayout()
        
        title = QLabel("Extract images to save them for later use as context")
        title.setStyleSheet("font-size: 14px; font-weight: bold; color: #2c3e50;")
        header_layout.addWidget(title)
        
        header_layout.addStretch()
        
        # Extract button (moved to top)
        extract_btn = QPushButton("🖼️ Extract Images")
        extract_btn.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                font-weight: bold;
                font-size: 12px;
                padding: 6px 12px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        extract_btn.clicked.connect(self._on_extract_images)
        header_layout.addWidget(extract_btn)
        
        main_layout.addLayout(header_layout)
        
        # Compact controls in a single row
        controls_layout = QHBoxLayout()
        
        # Input files
        add_file_btn = QPushButton("📄 Add File")
        add_file_btn.clicked.connect(self._on_add_docx_file_for_extraction)
        add_file_btn.setMaximumWidth(100)
        controls_layout.addWidget(add_file_btn)
        
        add_folder_btn = QPushButton("📁 Folder")
        add_folder_btn.clicked.connect(self._on_add_docx_folder_for_extraction)
        add_folder_btn.setMaximumWidth(80)
        controls_layout.addWidget(add_folder_btn)
        
        clear_list_btn = QPushButton("🗑️")
        clear_list_btn.clicked.connect(lambda: self.image_extractor_file_list.clear())
        clear_list_btn.setMaximumWidth(40)
        clear_list_btn.setToolTip("Clear file list")
        controls_layout.addWidget(clear_list_btn)
        
        # File list (compact, inline)
        self.image_extractor_file_list = QListWidget()
        self.image_extractor_file_list.setMaximumHeight(60)
        self.image_extractor_file_list.setStyleSheet("font-size: 9px;")
        controls_layout.addWidget(self.image_extractor_file_list, 1)
        
        # Auto-folder checkbox
        self.image_extractor_auto_folder = CheckmarkCheckBox("📁 Auto-folder")
        self.image_extractor_auto_folder.setChecked(False)
        self.image_extractor_auto_folder.setToolTip("Create 'Images' folder next to each DOCX file")
        self.image_extractor_auto_folder.toggled.connect(self._on_auto_folder_toggled)
        controls_layout.addWidget(self.image_extractor_auto_folder)
        
        # Output directory
        self.image_extractor_output_dir = QLineEdit()
        self.image_extractor_output_dir.setPlaceholderText("Output directory...")
        self.image_extractor_output_dir.setMaximumWidth(200)
        controls_layout.addWidget(self.image_extractor_output_dir)
        
        browse_btn = QPushButton("Browse...")
        browse_btn.clicked.connect(self._on_browse_output_dir_for_extraction)
        browse_btn.setMaximumWidth(80)
        controls_layout.addWidget(browse_btn)
        
        # Filename prefix
        controls_layout.addWidget(QLabel("Prefix:"))
        self.image_extractor_prefix = QLineEdit("Fig.")
        self.image_extractor_prefix.setMaximumWidth(60)
        controls_layout.addWidget(self.image_extractor_prefix)
        
        main_layout.addLayout(controls_layout)
        
        # Main horizontal splitter (left: operations, right: preview)
        results_splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # Left panel: Status and file list
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(0, 0, 0, 0)
        
        status_label = QLabel("📋 Results")
        status_label.setStyleSheet("font-weight: bold; font-size: 9px; color: #666;")
        left_layout.addWidget(status_label)
        
        # Status text area (resizable)
        self.image_extractor_status = QTextEdit()
        self.image_extractor_status.setReadOnly(True)
        self.image_extractor_status.setMinimumHeight(50)
        self.image_extractor_status.setStyleSheet("font-size: 9px;")
        self.image_extractor_status.setPlaceholderText("Extraction status...")
        left_layout.addWidget(self.image_extractor_status)
        
        # Extracted files list
        files_label = QLabel("📂 Extracted Files (click to preview)")
        files_label.setStyleSheet("font-weight: bold; font-size: 9px; color: #666; margin-top: 3px;")
        left_layout.addWidget(files_label)
        
        self.image_extractor_files_list = QListWidget()
        self.image_extractor_files_list.itemClicked.connect(self._on_file_list_item_clicked)
        self.image_extractor_files_list.setStyleSheet("""
            QListWidget {
                border: 1px solid #ccc;
                font-size: 9px;
            }
            QListWidget::item {
                padding: 3px;
                border-bottom: 1px solid #eee;
            }
            QListWidget::item:selected {
                background-color: #e3f2fd;
                color: black;
            }
        """)
        left_layout.addWidget(self.image_extractor_files_list)
        
        results_splitter.addWidget(left_widget)
        
        # Right panel: Full-height image preview
        preview_widget = QWidget()
        preview_layout = QVBoxLayout(preview_widget)
        preview_layout.setContentsMargins(0, 0, 0, 0)
        
        preview_label = QLabel("🖼️ Image Preview")
        preview_label.setStyleSheet("font-weight: bold; font-size: 9px; color: #666;")
        preview_layout.addWidget(preview_label)
        
        # Large image display area with scroll
        preview_scroll = QScrollArea()
        preview_scroll.setWidgetResizable(True)
        preview_scroll.setStyleSheet("QScrollArea { border: 1px solid #ccc; }")
        preview_scroll.setMinimumWidth(300)
        
        self.image_extractor_preview = QLabel()
        self.image_extractor_preview.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.image_extractor_preview.setStyleSheet("padding: 10px;")
        self.image_extractor_preview.setText("No image selected\n\nClick on a file in the list to preview\nor\nExtract images to see them here")
        self.image_extractor_preview.setWordWrap(True)
        
        preview_scroll.setWidget(self.image_extractor_preview)
        preview_layout.addWidget(preview_scroll)
        
        # Preview navigation buttons at bottom
        nav_layout = QHBoxLayout()
        
        self.preview_prev_btn = QPushButton("◀ Previous")
        self.preview_prev_btn.clicked.connect(self._on_preview_prev)
        self.preview_prev_btn.setEnabled(False)
        nav_layout.addWidget(self.preview_prev_btn)
        
        self.preview_image_label = QLabel("No images")
        self.preview_image_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.preview_image_label.setStyleSheet("font-size: 8px; color: #666;")
        nav_layout.addWidget(self.preview_image_label)
        
        self.preview_next_btn = QPushButton("Next ▶")
        self.preview_next_btn.clicked.connect(self._on_preview_next)
        self.preview_next_btn.setEnabled(False)
        nav_layout.addWidget(self.preview_next_btn)
        
        preview_layout.addLayout(nav_layout)
        
        results_splitter.addWidget(preview_widget)
        
        # Set initial splitter sizes (40% left, 60% right)
        results_splitter.setSizes([400, 600])
        
        main_layout.addWidget(results_splitter)
        
        # Initialize extractor and preview state
        self.image_extractor = ImageExtractor()
        self.extracted_image_files = []
        self.current_preview_index = 0
        
        return tab
    
    def create_pdf_rescue_tab(self) -> QWidget:
        """Create the PDF Rescue tab - AI OCR"""
        from modules.pdf_rescue_Qt import PDFRescueQt
        
        # Create PDF Rescue widget (embedded mode, not standalone)
        pdf_rescue_widget = QWidget()
        self.pdf_rescue_qt = PDFRescueQt(self, standalone=False)
        self.pdf_rescue_qt.create_tab(pdf_rescue_widget)
        
        return pdf_rescue_widget
    
    def create_encoding_repair_tab(self) -> QWidget:
        """Create the Encoding Repair tab - Text Encoding Tool"""
        from modules.encoding_repair_Qt import EncodingRepairQt
        
        # Create Encoding Repair widget (embedded mode, not standalone)
        encoding_repair_widget = QWidget()
        self.encoding_repair_qt = EncodingRepairQt(self, standalone=False)
        self.encoding_repair_qt.create_tab(encoding_repair_widget)
        
        return encoding_repair_widget
    
    def create_tracked_changes_tab(self) -> QWidget:
        """Create the Tracked Changes tab - Post-Translation Analysis"""
        return self._create_placeholder_tab(
            "📊 Tracked Changes",
            "Tracked Changes - Coming Soon\n\nFeatures:\n• Track translation changes\n• Version history\n• Comparison reports"
        )
    
    # ═══════════════════════════════════════════════════════════════════════════
    # Image Extractor Helper Methods
    # ═══════════════════════════════════════════════════════════════════════════
    
    def _on_add_docx_file_for_extraction(self):
        """Add a single DOCX file to the extraction list"""
        file_path, _ = fdh.get_open_file_name(
            self,
            "Select DOCX File",
            "Word Documents (*.docx)"
        )
        
        if file_path:
            # Avoid duplicates
            items = [self.image_extractor_file_list.item(i).text() 
                    for i in range(self.image_extractor_file_list.count())]
            if file_path not in items:
                self.image_extractor_file_list.addItem(file_path)
    
    def _on_add_docx_folder_for_extraction(self):
        """Add all DOCX files from a folder to the extraction list"""
        folder_path = fdh.get_existing_directory(
            self,
            "Select Folder Containing DOCX Files"
        )
        
        if folder_path:
            import glob
            docx_files = glob.glob(os.path.join(folder_path, "*.docx"))
            
            # Get existing items to avoid duplicates
            items = [self.image_extractor_file_list.item(i).text() 
                    for i in range(self.image_extractor_file_list.count())]
            
            added = 0
            for docx_file in docx_files:
                if docx_file not in items:
                    self.image_extractor_file_list.addItem(docx_file)
                    added += 1
            
            if added > 0:
                self.image_extractor_status.append(f"✅ Added {added} DOCX file(s) from folder")
    
    def _on_browse_output_dir_for_extraction(self):
        """Browse for output directory"""
        folder_path = fdh.get_existing_directory(
            self,
            "Select Output Directory for Extracted Images"
        )
        
        if folder_path:
            self.image_extractor_output_dir.setText(folder_path)
    
    def _on_auto_folder_toggled(self, checked):
        """Handle auto-folder checkbox toggle"""
        # Disable/enable output directory selection based on auto-folder
        self.image_extractor_output_dir.setEnabled(not checked)
        
        if checked:
            self.image_extractor_output_dir.setPlaceholderText("Auto: 'Images' folder next to each DOCX")
        else:
            self.image_extractor_output_dir.setPlaceholderText("Choose output directory...")
    
    def _on_file_list_item_clicked(self, item):
        """Handle click on extracted file list item"""
        if not item:
            return
        
        # Get the file path stored in the item
        file_path = item.data(Qt.ItemDataRole.UserRole)
        
        if file_path and file_path in self.extracted_image_files:
            # Update current index and preview
            self.current_preview_index = self.extracted_image_files.index(file_path)
            self._update_preview()
    
    def _on_extract_images(self):
        """Extract images from all DOCX files in the list"""
        # Validate inputs
        if self.image_extractor_file_list.count() == 0:
            QMessageBox.warning(
                self,
                "No Files",
                "Please add at least one DOCX file to extract images from."
            )
            return
        
        # Check if using auto-folder mode
        use_auto_folder = self.image_extractor_auto_folder.isChecked()
        
        if not use_auto_folder:
            output_dir = self.image_extractor_output_dir.text().strip()
            if not output_dir:
                QMessageBox.warning(
                    self,
                    "No Output Directory",
                    "Please select an output directory for the extracted images."
                )
                return
        
        prefix = self.image_extractor_prefix.text().strip()
        if not prefix:
            prefix = "Fig."
        
        # Get list of files
        docx_files = [self.image_extractor_file_list.item(i).text() 
                     for i in range(self.image_extractor_file_list.count())]
        
        # Clear status, file list, and preview
        self.image_extractor_status.clear()
        self.image_extractor_status.append("🔄 Starting image extraction...\n")
        self.image_extractor_files_list.clear()
        self.extracted_image_files = []
        self.current_preview_index = 0
        self.image_extractor_preview.setText("No image selected\n\nClick on a file in the list to preview\nor\nExtract images to see them here")
        QApplication.processEvents()
        
        try:
            total_count = 0
            all_extracted_files = []
            
            if use_auto_folder:
                # Extract to "Images" folder next to each DOCX
                self.image_extractor_status.append("📁 Mode: Auto-folder (Images subfolder per DOCX)\n")
                
                for docx_file in docx_files:
                    docx_dir = os.path.dirname(docx_file)
                    docx_name = os.path.splitext(os.path.basename(docx_file))[0]
                    auto_output_dir = os.path.join(docx_dir, "Images")
                    
                    self.image_extractor_status.append(f"📄 Processing: {os.path.basename(docx_file)}")
                    
                    count, files = self.image_extractor.extract_images_from_docx(
                        docx_file,
                        auto_output_dir,
                        prefix
                    )
                    
                    total_count += count
                    all_extracted_files.extend(files)
                    
                    if count > 0:
                        self.image_extractor_status.append(f"   ✅ {count} image(s) → {auto_output_dir}")
                    else:
                        self.image_extractor_status.append(f"   ⚠️  No images found")
                    
                    QApplication.processEvents()
                
                output_msg = f"Images saved in 'Images' subfolders next to each DOCX file"
            else:
                # Extract all to single directory
                output_dir = self.image_extractor_output_dir.text().strip()
                self.image_extractor_status.append(f"📁 Output directory: {output_dir}\n")
                
                total_count, all_extracted_files = self.image_extractor.extract_from_multiple_docx(
                    docx_files, 
                    output_dir, 
                    prefix
                )
                
                output_msg = output_dir
            
            # Show results
            self.image_extractor_status.append(f"\n✅ Successfully extracted {total_count} images!")
            
            if all_extracted_files:
                # Store files for preview
                self.extracted_image_files = all_extracted_files
                self.current_preview_index = 0
                
                # Populate file list
                self.image_extractor_files_list.clear()
                for i, file_path in enumerate(all_extracted_files, 1):
                    item = QListWidgetItem(f"{i}. {os.path.basename(file_path)}")
                    item.setData(Qt.ItemDataRole.UserRole, file_path)  # Store full path
                    item.setToolTip(file_path)  # Show full path on hover
                    self.image_extractor_files_list.addItem(item)
                
                # Enable preview buttons
                if len(self.extracted_image_files) > 0:
                    self.preview_prev_btn.setEnabled(len(self.extracted_image_files) > 1)
                    self.preview_next_btn.setEnabled(len(self.extracted_image_files) > 1)
                    
                    # Select and show first image
                    self.image_extractor_files_list.setCurrentRow(0)
                    self._update_preview()
            
            # Success message
            QMessageBox.information(
                self,
                "Extraction Complete",
                f"Successfully extracted {total_count} images!\n\n{output_msg}"
            )
            
        except Exception as e:
            self.image_extractor_status.append(f"\n❌ Error: {str(e)}")
            QMessageBox.critical(
                self,
                "Extraction Error",
                f"An error occurred during extraction:\n{str(e)}"
            )
    
    def _update_preview(self):
        """Update the image preview with the current image"""
        if not self.extracted_image_files or self.current_preview_index >= len(self.extracted_image_files):
            self.image_extractor_preview.setText("No image selected\n\nClick on a file in the list to preview\nor\nExtract images to see them here")
            self.preview_image_label.setText("No images")
            return
        
        image_path = self.extracted_image_files[self.current_preview_index]
        
        # Update list selection to match current preview
        if hasattr(self, 'image_extractor_files_list'):
            self.image_extractor_files_list.setCurrentRow(self.current_preview_index)
        
        try:
            from PyQt6.QtGui import QPixmap
            
            pixmap = QPixmap(image_path)
            
            if not pixmap.isNull():
                # Get the scroll area size for better scaling
                scroll_area = self.image_extractor_preview.parent()
                if scroll_area and hasattr(scroll_area, 'viewport'):
                    viewport_size = scroll_area.viewport().size()
                    max_width = viewport_size.width() - 40  # Account for padding
                    max_height = viewport_size.height() - 40
                else:
                    # Fallback to larger default size
                    max_width = 800
                    max_height = 600
                
                # Scale image to fit preview area while maintaining aspect ratio
                scaled_pixmap = pixmap.scaled(
                    max_width, max_height,
                    Qt.AspectRatioMode.KeepAspectRatio,
                    Qt.TransformationMode.SmoothTransformation
                )
                
                self.image_extractor_preview.setPixmap(scaled_pixmap)
                
                # Update label with size info
                self.preview_image_label.setText(
                    f"Image {self.current_preview_index + 1} of {len(self.extracted_image_files)}: "
                    f"{os.path.basename(image_path)} ({pixmap.width()}×{pixmap.height()}px)"
                )
            else:
                self.image_extractor_preview.setText(f"Failed to load image:\n{os.path.basename(image_path)}")
                
        except Exception as e:
            self.image_extractor_preview.setText(f"Error loading image:\n{str(e)}")
    
    def _on_preview_prev(self):
        """Show previous image in preview"""
        if self.extracted_image_files and self.current_preview_index > 0:
            self.current_preview_index -= 1
            self._update_preview()
    
    def _on_preview_next(self):
        """Show next image in preview"""
        if self.extracted_image_files and self.current_preview_index < len(self.extracted_image_files) - 1:
            self.current_preview_index += 1
            self._update_preview()

    # === IMAGE CONTEXT METHODS ===
    
    def _on_load_image_context_folder(self):
        """Load a folder of figure images for AI translation context"""
        from PyQt6.QtWidgets import QFileDialog
        
        folder = QFileDialog.getExistingDirectory(
            self,
            "Select Folder with Figure Images",
            "",
            QFileDialog.Option.ShowDirsOnly
        )
        
        if folder:
            try:
                count = self.figure_context.load_from_folder(folder)
                
                if count > 0:
                    self.image_context_status_label.setText(
                        f"✅ {count} image{'s' if count != 1 else ''} loaded from: {os.path.basename(folder)}"
                    )
                    self.image_context_status_label.setStyleSheet("color: #4CAF50; font-weight: bold; font-size: 11px; padding: 5px;")
                    self.log(f"[Image Context] Loaded {count} images from {folder}")
                    
                    # Update Prompt Manager display
                    if hasattr(self, 'prompt_manager_qt') and self.prompt_manager_qt:
                        self.prompt_manager_qt.update_image_context_display()
                    
                    QMessageBox.information(
                        self,
                        "Images Loaded",
                        f"Successfully loaded {count} figure image{'s' if count != 1 else ''}.\n\n"
                        f"These images will automatically be included with AI translations when the text "
                        f"references figures (e.g., 'Figure 1', 'see fig 2A')."
                    )
                else:
                    self.image_context_status_label.setText("⚠️ No valid images found in folder")
                    self.image_context_status_label.setStyleSheet("color: #FF9800; font-size: 11px; padding: 5px;")
                    QMessageBox.warning(
                        self,
                        "No Images Found",
                        "The selected folder does not contain any valid image files.\n\n"
                        "Supported formats: .png, .jpg, .jpeg, .gif, .bmp, .tiff\n"
                        "Filename examples: 'Figure 1.png', 'fig2a.jpg', 'Fig. 3-B.png'"
                    )
            except Exception as e:
                self.image_context_status_label.setText(f"❌ Error loading images")
                self.image_context_status_label.setStyleSheet("color: #F44336; font-size: 11px; padding: 5px;")
                self.log(f"[Image Context] Error loading images: {e}")
                QMessageBox.critical(
                    self,
                    "Error Loading Images",
                    f"Failed to load images from folder:\n\n{str(e)}"
                )
    
    def _on_clear_image_context(self):
        """Clear all loaded image context"""
        if self.figure_context.has_images():
            reply = QMessageBox.question(
                self,
                "Clear Image Context",
                f"Clear {self.figure_context.get_image_count()} loaded image{'s' if self.figure_context.get_image_count() != 1 else ''}?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                self.figure_context.clear()
                self.image_context_status_label.setText("No images loaded")
                self.image_context_status_label.setStyleSheet("color: #999; font-size: 11px; padding: 5px;")
                self.log("[Image Context] Cleared all images")
                
                # Update Prompt Manager display
                if hasattr(self, 'prompt_manager_qt') and self.prompt_manager_qt:
                    self.prompt_manager_qt.update_image_context_display()
        else:
            QMessageBox.information(
                self,
                "No Images Loaded",
                "There are no images currently loaded."
            )

    def create_llm_leaderboard_tab(self) -> QWidget:
        """Create the Superbench tab - Benchmark LLM translation quality"""
        from modules.llm_superbench_ui import LLMLeaderboardUI

        # Create LLM client factory that uses existing API keys
        def llm_client_factory(provider: str, model_id: str):
            from modules.llm_clients import LLMClient
            api_keys = self._get_api_keys()

            # Map provider names to API key names
            # (gemini uses "google" as the key name in api_keys.txt)
            provider_to_key = {
                "openai": "openai",
                "claude": "claude",
                "gemini": "google"  # Gemini uses Google API key
            }

            key_name = provider_to_key.get(provider, provider)
            api_key = api_keys.get(key_name, "")

            if not api_key:
                raise ValueError(f"No API key configured for {provider}. Please add it in Settings.")

            return LLMClient(api_key=api_key, provider=provider, model=model_id)

        # Create and return the leaderboard UI widget
        leaderboard_widget = LLMLeaderboardUI(
            parent=self,
            llm_client_factory=llm_client_factory
        )

        return leaderboard_widget

    def create_superbrowser_tab(self) -> QWidget:
        """Create the Superbrowser tab - Multi-Chat AI Browser"""
        from modules.superbrowser import SuperbrowserWidget

        # Create and return the Superbrowser widget
        superbrowser_widget = SuperbrowserWidget(parent=self)

        return superbrowser_widget

    def create_superdocs_tab(self) -> QWidget:
        """Create the Help tab - Online Documentation Viewer"""
        # The embedded docs viewer was removed in favor of online documentation.
        placeholder = QWidget()
        layout = QVBoxLayout(placeholder)
        label = QLabel("📚 Supervertaler Help is now available online.\n\nVisit https://supervertaler.gitbook.io/superdocs/ to view the documentation.")
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label.setStyleSheet("color: #888; font-size: 12px;")
        layout.addWidget(label)
        return placeholder

    def _get_api_keys(self) -> dict:
        """Get API keys from settings"""
        from modules.llm_clients import load_api_keys
        return load_api_keys()
    
    def show_concordance_search(self, initial_query: str = None):
        """
        Show concordance search (Ctrl+K) - now uses Superlookup instead of separate dialog.
        
        Superlookup provides concordance search plus:
        - Termbase search
        - Machine Translation
        - Web Resources
        
        The view can be toggled between Horizontal (table) and Vertical (list) modes.
        """
        try:
            # Get selected text if available and no initial query
            if not initial_query:
                # First check the currently focused widget (most reliable)
                focus_widget = QApplication.focusWidget()
                if focus_widget and hasattr(focus_widget, 'textCursor'):
                    cursor = focus_widget.textCursor()
                    if cursor.hasSelection():
                        initial_query = cursor.selectedText()
                        self.log(f"[Concordance] Got selection from focus widget: '{initial_query[:30]}...'")
                
                # Fallback: check source/target cells directly
                if not initial_query and hasattr(self, 'table') and self.table:
                    current_row = self.table.currentRow()
                    if current_row >= 0:
                        # Try source column first (column 2)
                        source_widget = self.table.cellWidget(current_row, 2)
                        if source_widget and hasattr(source_widget, 'textCursor'):
                            cursor = source_widget.textCursor()
                            if cursor.hasSelection():
                                initial_query = cursor.selectedText()
                        
                        # If no selection in source, try target column (column 3)
                        if not initial_query:
                            target_widget = self.table.cellWidget(current_row, 3)
                            if target_widget and hasattr(target_widget, 'textCursor'):
                                cursor = target_widget.textCursor()
                                if cursor.hasSelection():
                                    initial_query = cursor.selectedText()
            
            # Get project language pair
            source_lang = getattr(self, 'source_language', None)
            target_lang = getattr(self, 'target_language', None)
            
            # Log for debugging
            if initial_query:
                self.log(f"[Concordance] Opening Superlookup with query: '{initial_query[:50]}...' (lang: {source_lang} → {target_lang})")
            else:
                self.log(f"[Concordance] Opening Superlookup (no selection)")
            
            # Navigate to Superlookup tab
            self._go_to_superlookup()
            
            # Trigger search if we have a query
            if hasattr(self, 'lookup_tab') and self.lookup_tab:
                if initial_query:
                    # Use vertical view for traditional concordance layout
                    # Pass language pair from project
                    self.lookup_tab.search_with_query(
                        initial_query, 
                        switch_to_vertical=True,
                        source_lang=source_lang,
                        target_lang=target_lang
                    )
                else:
                    # Just focus the source text input
                    self.lookup_tab.source_text.setFocus()
                    # Reset language dropdowns to "Any" (index 0) for unrestricted search
                    if hasattr(self.lookup_tab, 'lang_from_combo'):
                        self.lookup_tab.lang_from_combo.setCurrentIndex(0)  # "Any"
                    if hasattr(self.lookup_tab, 'lang_to_combo'):
                        self.lookup_tab.lang_to_combo.setCurrentIndex(0)  # "Any"
                    # Switch to vertical view for consistency
                    if hasattr(self.lookup_tab, 'tm_view_vertical_radio'):
                        self.lookup_tab.tm_view_vertical_radio.setChecked(True)
                        
        except Exception as e:
            self.log(f"Error opening concordance search: {e}")
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "Error", f"Failed to open concordance search:\n{str(e)}")
    
    def show_tm_manager_tab(self, tab_index: int = 0):
        """Show TM Manager dialog opened to specific tab"""
        from modules.tm_manager_qt import TMManagerDialog
        
        try:
            dialog = TMManagerDialog(self, self.db_manager, self.log)
            dialog.tabs.setCurrentIndex(tab_index)
            dialog.exec()
        except Exception as e:
            self.log(f"Error opening TM Manager: {e}")
            QMessageBox.critical(self, "Error", f"Failed to open TM Manager:\n{str(e)}")

    def create_log_tab(self) -> QWidget:
        """Create the Log tab - Session Log"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(10)

        # Top toolbar with detach button
        toolbar = QWidget()
        toolbar_layout = QHBoxLayout(toolbar)
        toolbar_layout.setContentsMargins(0, 0, 0, 0)

        detach_btn = QPushButton("🪟 Detach Log Window")
        detach_btn.setToolTip("Open log in a separate window that can be moved to another screen")
        detach_btn.clicked.connect(self.detach_log_window)
        toolbar_layout.addWidget(detach_btn)

        clear_btn = QPushButton("🗑️ Clear Log")
        clear_btn.setToolTip("Clear all log messages")
        clear_btn.clicked.connect(self.clear_log)
        toolbar_layout.addWidget(clear_btn)

        toolbar_layout.addStretch()
        layout.addWidget(toolbar)

        # Log display area
        self.session_log = QPlainTextEdit()
        self.session_log.setReadOnly(True)
        self.session_log.setStyleSheet("""
            QPlainTextEdit {
                font-family: 'Consolas', 'Courier New', monospace;
                font-size: 9px;
                border: 1px solid #ccc;
            }
        """)
        layout.addWidget(self.session_log)

        # Initialize with welcome message
        self.session_log.setPlainText("Session Log - Ready\n" + "="*50 + "\n")

        # List to track detached log windows
        self.detached_log_windows = []

        return tab

    def detach_log_window(self):
        """Create a detached log window"""
        detached_window = DetachedLogWindow(self)
        detached_window.show()
        self.detached_log_windows.append(detached_window)
        self.log("🪟 Log window detached")

    def clear_log(self):
        """Clear the log"""
        reply = QMessageBox.question(
            self,
            "Clear Log",
            "Are you sure you want to clear the log?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            self.session_log.setPlainText("Session Log - Ready\n" + "="*50 + "\n")
            # Also clear detached windows
            for window in self.detached_log_windows:
                if window and not window.isHidden():
                    window.log_display.setPlainText("Session Log - Ready\n" + "="*50 + "\n")

    def _on_main_tab_changed(self, index: int):
        """Handle main tab changes (Grid/Project resources/Tools/Settings)"""
        try:
            if index == 0:  # Grid
                # Grid refreshes automatically when segments change
                pass
            elif index == 1:  # Project resources
                # Refresh AI Assistant LLM client when switching to Project resources tab
                # Prompts tab is inside Project resources, so refresh LLM settings when going there
                if hasattr(self, 'prompt_manager_qt'):
                    llm_settings = self.load_llm_settings()
                    self.current_provider = llm_settings.get('provider', 'openai')
                    provider_key = f"{self.current_provider}_model"
                    self.current_model = llm_settings.get(provider_key)
                    # Reinitialize AI Assistant's LLM client
                    self.prompt_manager_qt._init_llm_client()
        except Exception as e:
            self.log(f"⚠️ Error switching main tabs: {e}")
            import traceback
            traceback.print_exc()
    
    def detach_superlookup(self):
        """Detach Superlookup into a separate window for second screen use"""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QPushButton
        
        try:
            # If already detached, just show the window
            if hasattr(self, 'lookup_detached_window') and self.lookup_detached_window and self.lookup_detached_window.isVisible():
                self.lookup_detached_window.raise_()
                self.lookup_detached_window.activateWindow()
                return
            
            # Create detached window
            self.lookup_detached_window = QDialog(self)
            self.lookup_detached_window.setWindowTitle("🔍 Superlookup - Supervertaler")
            self.lookup_detached_window.setMinimumSize(600, 700)
            self.lookup_detached_window.resize(700, 800)
            
            # Set window flags to ensure it appears as a proper window
            self.lookup_detached_window.setWindowFlags(
                Qt.WindowType.Window | 
                Qt.WindowType.WindowCloseButtonHint | 
                Qt.WindowType.WindowMinimizeButtonHint |
                Qt.WindowType.WindowMaximizeButtonHint
            )
            
            # Position window to the right of main window, or center on same screen if no space
            main_geometry = self.geometry()
            
            # Get the screen that contains the main window (handles multi-monitor setups)
            app = QApplication.instance()
            screen = app.screenAt(main_geometry.center())
            if not screen:
                # Fallback to primary screen
                screen = app.primaryScreen()
            
            screen_geometry = screen.geometry()
            
            # Try to place to the right of main window
            new_x = main_geometry.right() + 20
            new_y = main_geometry.top()
            
            # If window would go off-screen to the right, center it on the main window's screen
            if new_x + 700 > screen_geometry.right():
                # Center horizontally on screen, place above main window
                new_x = screen_geometry.left() + (screen_geometry.width() - 700) // 2
                new_y = max(screen_geometry.top() + 50, main_geometry.top() - 100)
                
            # Ensure window stays on screen
            new_x = max(screen_geometry.left() + 10, min(new_x, screen_geometry.right() - 710))
            new_y = max(screen_geometry.top() + 10, min(new_y, screen_geometry.bottom() - 810))
            
            self.lookup_detached_window.move(int(new_x), int(new_y))
            
            # Make sure window is raised and activated - use QTimer to ensure it's after window is shown
            QTimer.singleShot(100, lambda: (
                self.lookup_detached_window.raise_(),
                self.lookup_detached_window.activateWindow(),
                self.lookup_detached_window.setFocus()
            ))
            
            layout = QVBoxLayout(self.lookup_detached_window)
            layout.setContentsMargins(10, 10, 10, 10)
            layout.setSpacing(5)
            
            # Header with reattach button
            header_layout = QVBoxLayout()
            
            header_title = QLabel("🔍 Superlookup")
            header_title.setStyleSheet("font-size: 16px; font-weight: bold; color: #333;")
            header_layout.addWidget(header_title)
            
            button_layout = QVBoxLayout()
            reattach_btn = QPushButton("📥 Attach to Main Window")
            reattach_btn.setToolTip("Re-attach Superlookup to the Home tab")
            reattach_btn.setStyleSheet("font-size: 9pt; padding: 4px 12px; max-width: 200px;")
            reattach_btn.clicked.connect(self.reattach_superlookup)
            button_layout.addWidget(reattach_btn, alignment=Qt.AlignmentFlag.AlignRight)
            header_layout.addLayout(button_layout)
            
            layout.addLayout(header_layout)
            
            # Create new Superlookup instance for detached window
            # Or move the existing one - better to create new to avoid widget parenting issues
            detached_lookup = SuperlookupTab(self.lookup_detached_window)
            
            # Explicitly copy theme_manager reference
            detached_lookup.theme_manager = self.theme_manager
            
            # Copy TM database reference if available
            if hasattr(self, 'tm_database') and self.tm_database:
                detached_lookup.tm_database = self.tm_database
                if detached_lookup.engine:
                    detached_lookup.engine.set_tm_database(self.tm_database)
            
            # Copy home lookup state if it exists
            if hasattr(self, 'home_lookup_widget') and self.home_lookup_widget:
                # Copy source text
                source_text = self.home_lookup_widget.source_text.toPlainText()
                detached_lookup.source_text.setPlainText(source_text)
                
                # Copy TM database reference
                if hasattr(self.home_lookup_widget, 'tm_database'):
                    detached_lookup.tm_database = self.home_lookup_widget.tm_database
            
            layout.addWidget(detached_lookup, stretch=1)
            
            # Store reference for cleanup
            self.lookup_detached_widget = detached_lookup
            
            # Handle window close
            def on_close():
                self.lookup_detached_window = None
                self.lookup_detached_widget = None
            
            self.lookup_detached_window.finished.connect(on_close)
            
            # Show window (non-modal)
            self.lookup_detached_window.setWindowModality(Qt.WindowModality.NonModal)
            self.lookup_detached_window.show()
            self.log("Superlookup detached to separate window")
        except Exception as e:
            import traceback
            error_msg = f"Error detaching Superlookup: {str(e)}\n{traceback.format_exc()}"
            self.log(error_msg)
            QMessageBox.warning(self, "Error", f"Could not detach Superlookup:\n{str(e)}")
    
    def reattach_superlookup(self):
        """Re-attach Superlookup to the Home tab"""
        if not self.lookup_detached_window:
            return
        
        # Copy state back to home widget if it exists
        if (hasattr(self, 'home_lookup_widget') and self.home_lookup_widget and 
            hasattr(self, 'lookup_detached_widget') and self.lookup_detached_widget):
            # Copy source text
            source_text = self.lookup_detached_widget.source_text.toPlainText()
            self.home_lookup_widget.source_text.setPlainText(source_text)
        
        # Close detached window
        self.lookup_detached_window.close()
        self.lookup_detached_window = None
        if hasattr(self, 'lookup_detached_widget'):
            self.lookup_detached_widget = None
        
        self.log("Superlookup re-attached to Home tab")
    
    def create_resources_tab(self):
        """Create the Project resources tab with nested sub-tabs"""
        from PyQt6.QtWidgets import QTabWidget
        
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        # Create nested tab widget
        resources_tabs = QTabWidget()
        resources_tabs.tabBar().setFocusPolicy(Qt.FocusPolicy.NoFocus)
        resources_tabs.tabBar().setDrawBase(False)
        resources_tabs.setStyleSheet("QTabBar::tab { outline: 0; } QTabBar::tab:focus { outline: none; } QTabBar::tab:selected { border-bottom: 1px solid #2196F3; background-color: rgba(33, 150, 243, 0.08); }")
        self.resources_tabs = resources_tabs  # Store for navigation
        
        # Add nested tabs
        tm_tab = self.create_translation_memories_tab()
        resources_tabs.addTab(tm_tab, "💾 TMs")
        
        termbase_tab = self.create_termbases_tab()
        resources_tabs.addTab(termbase_tab, "🏷️ Glossaries")
        
        nt_tab = self.create_non_translatables_tab()
        resources_tabs.addTab(nt_tab, "🚫 Non-Translatables")
        
        ref_tab = self.create_reference_images_tab()
        resources_tabs.addTab(ref_tab, "🎯 Image Context")
        
        layout.addWidget(resources_tabs)
        
        return tab
    
    def create_specialised_tools_tab(self):
        """Create the Specialised Tools tab with nested sub-tabs"""
        from PyQt6.QtWidgets import QTabWidget
        
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        # Create nested tab widget
        modules_tabs = QTabWidget()
        modules_tabs.tabBar().setFocusPolicy(Qt.FocusPolicy.NoFocus)
        modules_tabs.tabBar().setDrawBase(False)
        modules_tabs.setStyleSheet("QTabBar::tab { outline: 0; } QTabBar::tab:focus { outline: none; } QTabBar::tab:selected { border-bottom: 1px solid #2196F3; background-color: rgba(33, 150, 243, 0.08); }")
        self.modules_tabs = modules_tabs  # Store for navigation
        
        # Add nested tabs (alphabetical order)
        autofingers_tab = AutoFingersWidget(self)
        modules_tabs.addTab(autofingers_tab, "✋ AutoFingers")
        
        # Superconverter - Format conversion tools
        superconverter_tab = self.create_superconverter_tab()
        modules_tabs.addTab(superconverter_tab, "🔄 Superconverter")
        
        pdf_tab = self.create_pdf_rescue_tab()
        modules_tabs.addTab(pdf_tab, "📄 PDF Rescue")
        
        # Superbench
        leaderboard_tab = self.create_llm_leaderboard_tab()
        modules_tabs.addTab(leaderboard_tab, "📊 Superbench")
        
        # Superbrowser - Multi-Chat AI Browser
        superbrowser_tab = self.create_superbrowser_tab()
        modules_tabs.addTab(superbrowser_tab, "🌐 Superbrowser")

        supercleaner_tab = self.create_supercleaner_tab()
        modules_tabs.addTab(supercleaner_tab, "🧹 Supercleaner")

        # Superdocs removed (online GitBook will be used instead)

        print("[DEBUG] About to create SuperlookupTab...")
        lookup_tab = SuperlookupTab(self)
        print("[DEBUG] SuperlookupTab created successfully")
        self.lookup_tab = lookup_tab  # Store reference for later use
        modules_tabs.addTab(lookup_tab, "🔍 Superlookup")
        print("[DEBUG] Superlookup tab added to modules_tabs")

        # Supervoice - Voice Commands & Dictation
        supervoice_tab = self._create_voice_dictation_settings_tab()
        modules_tabs.addTab(supervoice_tab, "🎤 Supervoice")

        encoding_tab = self.create_encoding_repair_tab()
        modules_tabs.addTab(encoding_tab, "🔧 Text Encoding Repair")
        
        tmx_tab = self.create_tmx_editor_tab()
        modules_tabs.addTab(tmx_tab, "✏️ TMX Editor")
        
        tracked_tab = self.create_tracked_changes_tab()
        modules_tabs.addTab(tracked_tab, "🔄 Tracked Changes")

        layout.addWidget(modules_tabs)

        return tab
    
    def create_translation_memories_tab(self):
        """Create the Translation Memories tab with nested sub-tabs"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        # Check if database is available
        if not (hasattr(self, 'db_manager') and self.db_manager):
            placeholder = QLabel("Translation Memories Manager\n\nDatabase not initialized.")
            placeholder.setAlignment(Qt.AlignmentFlag.AlignCenter)
            placeholder.setStyleSheet("color: #888; font-size: 12px;")
            layout.addWidget(placeholder, stretch=1)
            return tab
        
        # Import TM metadata manager
        from modules.tm_metadata_manager import TMMetadataManager
        tm_metadata_mgr = TMMetadataManager(self.db_manager, self.log)
        self.tm_metadata_mgr = tm_metadata_mgr  # Store for later use
        
        # Create nested tab widget for TM functionality
        tm_tabs = QTabWidget()
        tm_tabs.tabBar().setFocusPolicy(Qt.FocusPolicy.NoFocus)
        tm_tabs.tabBar().setDrawBase(False)
        tm_tabs.setStyleSheet("QTabBar::tab { outline: 0; } QTabBar::tab:focus { outline: none; } QTabBar::tab:selected { border-bottom: 1px solid #2196F3; background-color: rgba(33, 150, 243, 0.08); }")
        
        # Tab 1: TM List (Management) - manage multiple TMs
        tm_list_tab = self._create_tm_list_tab(tm_metadata_mgr)
        tm_tabs.addTab(tm_list_tab, "📋 TM List")
        
        # Tab 2: Browse All - browse ALL active TMs together
        from modules.tm_manager_qt import TMManagerDialog
        temp_manager = TMManagerDialog(self, self.db_manager, self.log)
        tm_tabs.addTab(temp_manager.browser_tab, "📖 Browse All")
        
        # Note: Concordance tab removed - functionality moved to Superlookup (Ctrl+K)
        # Note: Import/Export tab removed - functionality available in TM List tab
        
        # Tab 3: Statistics - aggregate stats for all TMs
        tm_tabs.addTab(temp_manager.stats_tab, "📊 Statistics")
        
        # Tab 4: Maintenance - cleanup and maintenance tools
        tm_tabs.addTab(temp_manager.maintenance_tab, "🧹 Maintenance")
        
        # Store reference to prevent garbage collection
        tab._tm_manager = temp_manager
        
        layout.addWidget(tm_tabs)
        
        return tab
    
    def _create_tm_list_tab(self, tm_metadata_mgr):
        """Create the TM List sub-tab with table and management buttons"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(10, 10, 10, 10)
        
        # Header
        header = QLabel("💾 TMs")
        header.setStyleSheet("font-size: 14px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(header)
        
        # Description
        desc = QLabel("Manage TMs. Activate/deactivate TMs for current project. Import client TMX files as named TMs.")
        desc.setStyleSheet("color: #666; font-size: 11px; margin-bottom: 10px;")
        layout.addWidget(desc)
        
        # Search bar
        search_layout = QHBoxLayout()
        search_box = QLineEdit()
        search_box.setPlaceholderText("Search TMs...")
        search_box.setMaximumWidth(300)
        search_layout.addWidget(search_box)
        search_layout.addStretch()
        layout.addLayout(search_layout)
        
        # Help message for first-time users
        help_msg = QLabel(
            "💡 <b>Translation Memories</b><br>"
            "• <b>Read</b> (green ✓): TM is used for matching segments<br>"
            "• <b>Write</b> (blue ✓): TM is updated with new translations<br>"
            "• <b>Typical Setup</b>: Main TM (Read + Write) + Reference TMs (Read only)"
        )
        help_msg.setWordWrap(True)
        help_msg.setStyleSheet("background-color: #e3f2fd; padding: 8px; border-radius: 4px; color: #1976d2;")
        layout.addWidget(help_msg)
        
        # Bulk action controls
        bulk_layout = QHBoxLayout()
        bulk_layout.addWidget(QLabel("Quick Actions:"))
        
        read_header_checkbox = CheckmarkCheckBox("Select All Read")
        bulk_layout.addWidget(read_header_checkbox)
        
        write_header_checkbox = BlueCheckmarkCheckBox("Select All Write")
        bulk_layout.addWidget(write_header_checkbox)
        
        bulk_layout.addStretch()
        layout.addLayout(bulk_layout)
        
        # TM list with table
        tm_table = QTableWidget()
        tm_table.setColumnCount(7)
        tm_table.setHorizontalHeaderLabels(["TM Name", "Languages", "Entries", "Read", "Write", "Last Modified", "Description"])
        tm_table.horizontalHeader().setStretchLastSection(True)
        tm_table.setColumnWidth(0, 250)
        tm_table.setColumnWidth(1, 120)
        tm_table.setColumnWidth(2, 80)
        tm_table.setColumnWidth(3, 60)
        tm_table.setColumnWidth(4, 60)
        tm_table.setColumnWidth(5, 150)
        
        # Get current project (for lambda closures below)
        current_project = self.current_project if hasattr(self, 'current_project') else None
        project_id = current_project.id if (current_project and hasattr(current_project, 'id')) else None
        
        # Connect header checkboxes to toggle all
        def toggle_all_read(checked):
            for row in range(tm_table.rowCount()):
                checkbox = tm_table.cellWidget(row, 3)
                if checkbox and isinstance(checkbox, CheckmarkCheckBox):
                    checkbox.setChecked(checked)
        
        def toggle_all_write(checked):
            for row in range(tm_table.rowCount()):
                checkbox = tm_table.cellWidget(row, 4)
                if checkbox and isinstance(checkbox, BlueCheckmarkCheckBox):
                    checkbox.setChecked(checked)
        
        read_header_checkbox.toggled.connect(toggle_all_read)
        write_header_checkbox.toggled.connect(toggle_all_write)
        
        # Populate TM list
        def refresh_tm_list():
            # Get current project dynamically (not captured in closure!)
            # Use GLOBAL_PROJECT_ID (0) when no project is loaded for Superlookup support
            current_proj = self.current_project if hasattr(self, 'current_project') else None
            refresh_project_id = current_proj.id if (current_proj and hasattr(current_proj, 'id')) else 0  # 0 = global
            
            tms = tm_metadata_mgr.get_all_tms()
            tm_table.setRowCount(len(tms))
            
            for row, tm in enumerate(tms):
                # Check if active (Read mode) for current project or global (0)
                # Note: is_tm_active now supports project_id=0 for global activations
                is_readable = tm_metadata_mgr.is_tm_active(tm['id'], refresh_project_id)
                # Default: read-only (Write unchecked) - read_only=True means not writable
                # If read_only is not set in database, treat as read-only by default
                is_writable = not tm.get('read_only', True)  # Default to True (read-only) if not set
                
                # TM Name (bold if readable)
                name_item = QTableWidgetItem(tm['name'])
                if is_readable:
                    font = name_item.font()
                    font.setBold(True)
                    name_item.setFont(font)
                tm_table.setItem(row, 0, name_item)
                
                # Languages (normalized format: nl-NL, en-US, etc.)
                from modules.tmx_generator import normalize_lang_variant
                src_lang = normalize_lang_variant(tm['source_lang']) if tm['source_lang'] else '?'
                tgt_lang = normalize_lang_variant(tm['target_lang']) if tm['target_lang'] else '?'
                langs = f"{src_lang} → {tgt_lang}"
                tm_table.setItem(row, 1, QTableWidgetItem(langs))
                
                # Entry count
                tm_table.setItem(row, 2, QTableWidgetItem(str(tm['entry_count'])))
                
                # Read checkbox (green checkmark)
                read_checkbox = CheckmarkCheckBox()
                read_checkbox.setChecked(is_readable)
                read_checkbox.setToolTip("Read: TM is used for matching segments")
                
                def on_read_toggle(checked, tm_id=tm['id'], row_idx=row):
                    # Get current project ID dynamically
                    # Use 0 (global) when no project is loaded - allows Superlookup to work
                    curr_proj = self.current_project if hasattr(self, 'current_project') else None
                    curr_proj_id = curr_proj.id if (curr_proj and hasattr(curr_proj, 'id')) else 0  # 0 = global
                    
                    # Activate/deactivate for reading
                    if checked:
                        success = tm_metadata_mgr.activate_tm(tm_id, curr_proj_id)
                    else:
                        success = tm_metadata_mgr.deactivate_tm(tm_id, curr_proj_id)
                    
                    if success:
                        status = "readable" if checked else "not readable"
                        self.log(f"✅ TM {tm_id} set to {status}")
                        # Update name to be bold if readable
                        name_item = tm_table.item(row_idx, 0)
                        if name_item:
                            font = name_item.font()
                            font.setBold(checked)
                            name_item.setFont(font)
                        
                        # Invalidate translation cache so matches are refreshed with new TM settings
                        self.invalidate_translation_cache(smart_invalidation=False)
                    else:
                        self.log(f"❌ Failed to toggle TM {tm_id}")
                        # Revert checkbox on failure
                        sender_checkbox = tm_table.cellWidget(row_idx, 3)
                        if sender_checkbox:
                            sender_checkbox.blockSignals(True)
                            sender_checkbox.setChecked(not checked)
                            sender_checkbox.blockSignals(False)
                
                read_checkbox.toggled.connect(on_read_toggle)
                tm_table.setCellWidget(row, 3, read_checkbox)
                
                # Write checkbox (blue checkmark)
                write_checkbox = BlueCheckmarkCheckBox()
                write_checkbox.setChecked(is_writable)
                write_checkbox.setToolTip("Write: TM is updated with new translations")
                
                def on_write_toggle(checked, tm_id=tm['id'], row_idx=row):
                    # Invert logic: checked = writable, so set read_only to NOT checked
                    success = tm_metadata_mgr.set_read_only(tm_id, not checked)
                    if success:
                        status = "writable" if checked else "read-only"
                        self.log(f"✅ TM {tm_id} set to {status}")
                    else:
                        # Revert on failure
                        sender = tm_table.cellWidget(row_idx, 4)
                        if sender:
                            sender.blockSignals(True)
                            sender.setChecked(not checked)
                            sender.blockSignals(False)
                
                write_checkbox.toggled.connect(on_write_toggle)
                tm_table.setCellWidget(row, 4, write_checkbox)
                
                # Last modified
                modified = tm['modified_date'] or tm['created_date'] or ''
                if modified:
                    # Format datetime nicely
                    try:
                        from datetime import datetime
                        dt = datetime.fromisoformat(modified)
                        modified = dt.strftime("%Y-%m-%d %H:%M")
                    except:
                        pass
                tm_table.setItem(row, 5, QTableWidgetItem(modified))
                
                # Description
                desc_text = tm['description'] or ''
                tm_table.setItem(row, 6, QTableWidgetItem(desc_text))
            
            # Update header checkbox states based on current selection
            read_header_checkbox.blockSignals(True)
            write_header_checkbox.blockSignals(True)
            all_read_checked = all(tm_table.cellWidget(r, 3).isChecked() if tm_table.cellWidget(r, 3) else False for r in range(tm_table.rowCount())) if tm_table.rowCount() > 0 else False
            all_write_checked = all(tm_table.cellWidget(r, 4).isChecked() if tm_table.cellWidget(r, 4) else False for r in range(tm_table.rowCount())) if tm_table.rowCount() > 0 else False
            read_header_checkbox.setChecked(all_read_checked)
            write_header_checkbox.setChecked(all_write_checked)
            read_header_checkbox.blockSignals(False)
            write_header_checkbox.blockSignals(False)
        
        # Store callback as instance attribute so load_project can refresh UI after restoration
        self.tm_tab_refresh_callback = refresh_tm_list
        
        refresh_tm_list()
        layout.addWidget(tm_table, stretch=1)
        
        # Button bar
        button_layout = QHBoxLayout()
        
        create_btn = QPushButton("+ Create New TM")
        create_btn.clicked.connect(lambda: self._show_create_tm_dialog(tm_metadata_mgr, refresh_tm_list, project_id))
        button_layout.addWidget(create_btn)
        
        import_btn = QPushButton("📥 Import TMX")
        import_btn.setToolTip("Import TMX file as a new TM or add to existing TM")
        import_btn.clicked.connect(lambda: self._import_tmx_as_tm(tm_metadata_mgr, tm_table, refresh_tm_list))
        button_layout.addWidget(import_btn)
        
        export_btn = QPushButton("📤 Export TM")
        export_btn.setToolTip("Export selected TM to TMX file")
        export_btn.clicked.connect(lambda: self._export_tm_to_tmx(tm_metadata_mgr, tm_table))
        button_layout.addWidget(export_btn)
        
        edit_btn = QPushButton("✏️ Edit/Maintain TM")
        edit_btn.setToolTip("Open editor for selected TM (Browse, Import/Export, Maintenance)")
        edit_btn.clicked.connect(lambda: self._show_tm_editor_dialog(tm_metadata_mgr, tm_table))
        button_layout.addWidget(edit_btn)
        
        delete_btn = QPushButton("🗑️ Delete TM")
        delete_btn.clicked.connect(lambda: self._delete_tm(tm_metadata_mgr, tm_table, refresh_tm_list))
        button_layout.addWidget(delete_btn)
        
        button_layout.addStretch()
        layout.addLayout(button_layout)
        
        return tab
    
    def _show_tm_context_menu(self, tm_table, tm_metadata_mgr, refresh_callback, pos):
        """Show context menu for TM table"""
        current_row = tm_table.rowAt(pos.y())
        if current_row < 0:
            return
        
        # Get TM info from the row (column 0 = TM Name)
        name_item = tm_table.item(current_row, 0)
        if not name_item:
            return
        
        tm_name = name_item.text()
        
        # Find the TM in the list
        all_tms = tm_metadata_mgr.get_all_tms()
        tm = None
        for t in all_tms:
            if t['name'] == tm_name:
                tm = t
                break
        
        if not tm:
            return
        
        menu = QMenu(tm_table)
        
        # Check if this TM is already the project TM
        is_project_tm = tm.get('is_project_tm', False)
        
        # Get current project
        current_proj = self.current_project if hasattr(self, 'current_project') else None
        project_id = current_proj.id if (current_proj and hasattr(current_proj, 'id')) else None
        
        if project_id:
            if not is_project_tm:
                set_project_action = menu.addAction("📌 Set as Project TM")
                set_project_action.triggered.connect(
                    lambda: self._set_tm_as_project(tm['id'], project_id, tm_metadata_mgr, refresh_callback)
                )
            else:
                unset_project_action = menu.addAction("📌 Unset as Project TM")
                unset_project_action.triggered.connect(
                    lambda: self._unset_tm_as_project(tm['id'], tm_metadata_mgr, refresh_callback)
                )
        else:
            no_project_action = menu.addAction("⚠️ Load a project first")
            no_project_action.setEnabled(False)
        
        menu.addSeparator()
        
        # Read-only toggle
        is_readonly = tm.get('read_only', False)
        readonly_text = "✏️ Set as Writable" if is_readonly else "🔒 Set as Read-Only"
        readonly_action = menu.addAction(readonly_text)
        readonly_action.triggered.connect(
            lambda: self._toggle_tm_readonly(tm['id'], not is_readonly, tm_metadata_mgr, refresh_callback)
        )
        
        menu.exec(tm_table.viewport().mapToGlobal(pos))
    
    def _set_tm_as_project(self, tm_id, project_id, tm_metadata_mgr, refresh_callback):
        """Set a TM as the project TM"""
        success = tm_metadata_mgr.set_as_project_tm(tm_id, project_id)
        if success:
            self.log(f"✅ Set TM as Project TM")
            refresh_callback()
        else:
            QMessageBox.warning(self, "Error", "Failed to set TM as Project TM")
    
    def _unset_tm_as_project(self, tm_id, tm_metadata_mgr, refresh_callback):
        """Unset a TM as project TM"""
        success = tm_metadata_mgr.unset_project_tm(tm_id)
        if success:
            self.log(f"✅ Unset Project TM")
            refresh_callback()
        else:
            QMessageBox.warning(self, "Error", "Failed to unset Project TM")
    
    def _toggle_tm_readonly(self, tm_id, readonly, tm_metadata_mgr, refresh_callback):
        """Toggle read-only status of a TM"""
        success = tm_metadata_mgr.set_read_only(tm_id, readonly)
        if success:
            status = "read-only" if readonly else "writable"
            self.log(f"✅ Set TM as {status}")
            refresh_callback()
        else:
            QMessageBox.warning(self, "Error", f"Failed to set TM as {'read-only' if readonly else 'writable'}")
    
    def import_tmx_file(self):
        """Import TMX file into translation memory with language variant handling"""
        try:
            file_path, _ = fdh.get_open_file_name(
                self, 
                "Import TMX File", 
                "TMX Files (*.tmx);;All Files (*.*)"
            )
            
            if file_path and self.tm_database:
                # Get target languages
                source_lang = getattr(self.current_project, 'source_lang', 'en') if self.current_project else 'en'
                target_lang = getattr(self.current_project, 'target_lang', 'nl') if self.current_project else 'nl'
                
                # Detect TMX languages
                tmx_langs = self.tm_database.detect_tmx_languages(file_path)
                self.log(f"TMX languages detected: {tmx_langs}")
                
                # Check compatibility
                compat = self.tm_database.check_language_compatibility(tmx_langs, source_lang, target_lang)
                
                if not compat['compatible']:
                    QMessageBox.warning(
                        self, "Language Mismatch",
                        f"TMX file languages ({', '.join(tmx_langs)}) don't match project languages ({source_lang}, {target_lang}).\n\n"
                        "Cannot import this TMX file."
                    )
                    return
                
                # Handle variant mismatch
                if compat.get('variant_match'):
                    choice = self._show_language_variant_dialog(compat)
                    
                    if choice == 'cancel':
                        return
                    
                    # Show progress dialog
                    from PyQt6.QtWidgets import QProgressDialog
                    from PyQt6.QtCore import Qt
                    
                    progress = QProgressDialog("Importing TMX file...", None, 0, 0, self)
                    progress.setWindowTitle("TMX Import")
                    progress.setWindowModality(Qt.WindowModality.WindowModal)
                    progress.setMinimumDuration(0)
                    progress.setValue(0)
                    progress.show()
                    QApplication.processEvents()
                    
                    if choice == 'import_strip':
                        # Import with variant stripping
                        tm_id, count = self.tm_database.load_tmx_file(
                            file_path, source_lang, target_lang, 
                            tm_name=None, read_only=False
                        )
                        progress.close()
                        self.log(f"Imported {count} entries (variants stripped: {compat['tmx_source']}, {compat['tmx_target']} → {source_lang}, {target_lang})")
                        QMessageBox.information(
                            self, "Import Complete",
                            f"TMX imported successfully!\n\nEntries: {count}\n"
                            f"Mapped: {compat['tmx_source']}, {compat['tmx_target']} → {source_lang}, {target_lang}"
                        )
                    elif choice == 'create_new':
                        # Create new TM with variant languages
                        tm_id, count = self.tm_database.load_tmx_file(
                            file_path, compat['tmx_source'], compat['tmx_target'],
                            tm_name=None, read_only=False
                        )
                        progress.close()
                        self.log(f"Created new TM with variants: {compat['tmx_source']}, {compat['tmx_target']}")
                        QMessageBox.information(
                            self, "Import Complete", 
                            f"New TM created!\n\nTM ID: {tm_id}\nEntries: {count}\n"
                            f"Languages: {compat['tmx_source']}, {compat['tmx_target']}"
                        )
                else:
                    # Exact match - proceed normally with progress
                    from PyQt6.QtWidgets import QProgressDialog
                    from PyQt6.QtCore import Qt
                    
                    progress = QProgressDialog("Importing TMX file...", None, 0, 0, self)
                    progress.setWindowTitle("TMX Import")
                    progress.setWindowModality(Qt.WindowModality.WindowModal)
                    progress.setMinimumDuration(0)
                    progress.setValue(0)
                    progress.show()
                    QApplication.processEvents()
                    
                    tm_id, count = self.tm_database.load_tmx_file(file_path, source_lang, target_lang)
                    progress.close()
                    self.log(f"Successfully imported {count} entries from TMX file")
                    QMessageBox.information(self, "Import Complete", f"TMX imported!\n\nEntries: {count}\nTM ID: {tm_id}")
            
        except Exception as e:
            self.log(f"Error importing TMX file: {e}")
            QMessageBox.critical(self, "Import Error", f"Failed to import TMX:\n\n{e}")
    
    def _show_language_variant_dialog(self, compat_info: dict) -> str:
        """Show dialog for handling language variant mismatch. Returns 'import_strip', 'create_new', or 'cancel'"""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QPushButton, QRadioButton, QButtonGroup
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Language Variant Detected")
        dialog.setMinimumWidth(500)
        
        layout = QVBoxLayout(dialog)
        
        # Explanation
        info_text = (
            f"<h3>⚠️ Language Variant Mismatch</h3>"
            f"<p>The TMX file uses language variants that don't exactly match your TM:</p>"
            f"<table style='margin: 10px 0;'>"
            f"<tr><td><b>TMX languages:</b></td><td>{compat_info['tmx_source']} → {compat_info['tmx_target']}</td></tr>"
            f"<tr><td><b>Your TM languages:</b></td><td>{compat_info['target_source']} → {compat_info['target_target']}</td></tr>"
            f"</table>"
            f"<p>How would you like to proceed?</p>"
        )
        
        info_label = QLabel(info_text)
        info_label.setWordWrap(True)
        info_label.setTextFormat(Qt.TextFormat.RichText)
        layout.addWidget(info_label)
        
        # Radio button group
        button_group = QButtonGroup(dialog)
        
        option1 = CheckmarkRadioButton(f"Import into existing TM (strip variants: {compat_info['tmx_source']},{compat_info['tmx_target']} → {compat_info['target_source']},{compat_info['target_target']})")
        option1.setChecked(True)
        option1.setToolTip("Import translations by matching base languages, ignoring regional variants")
        button_group.addButton(option1, 1)
        layout.addWidget(option1)
        
        option2 = CheckmarkRadioButton(f"Create new TM with variant languages ({compat_info['tmx_source']}, {compat_info['tmx_target']})")
        option2.setToolTip("Create a separate TM preserving the exact language variants from the TMX")
        button_group.addButton(option2, 2)
        layout.addWidget(option2)
        
        layout.addSpacing(20)
        
        # Buttons
        from PyQt6.QtWidgets import QHBoxLayout
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        btn_layout.addWidget(cancel_btn)
        
        ok_btn = QPushButton("Continue")
        ok_btn.setDefault(True)
        ok_btn.clicked.connect(dialog.accept)
        btn_layout.addWidget(ok_btn)
        
        layout.addLayout(btn_layout)
        
        if dialog.exec() == QDialog.DialogCode.Accepted:
            if option1.isChecked():
                return 'import_strip'
            elif option2.isChecked():
                return 'create_new'
        
        return 'cancel'
    
    def export_target_only_docx(self):
        """Export target text only as a monolingual DOCX document, preserving original formatting"""
        try:
            if not self.current_project or not self.current_project.segments:
                QMessageBox.warning(self, "No Project", "Please open a project with segments first")
                return
            
            # Check if there are any translations
            segments = list(self.current_project.segments)
            translated_count = sum(1 for seg in segments if seg.target and seg.target.strip())
            
            if translated_count == 0:
                QMessageBox.warning(
                    self, "No Translations", 
                    "No translated segments found.\n\nPlease translate some segments before exporting."
                )
                return
            
            # Warn if not all segments are translated
            total_count = len(segments)
            if translated_count < total_count:
                reply = QMessageBox.question(
                    self, "Incomplete Translation",
                    f"Only {translated_count} of {total_count} segments are translated.\n\n"
                    f"Untranslated segments will use source text.\n\n"
                    f"Continue with export?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                )
                if reply != QMessageBox.StandardButton.Yes:
                    return
            
            # Get save path
            default_name = ""
            if self.current_project.name:
                default_name = self.current_project.name.replace(" ", "_") + "_translated.docx"
            
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Export Target Only DOCX",
                default_name,
                "Word Documents (*.docx);;All Files (*.*)"
            )
            
            if not file_path:
                return
            
            # Ensure .docx extension
            if not file_path.lower().endswith('.docx'):
                file_path += '.docx'
            
            from docx import Document
            import shutil
            
            # Check if we have the original document to use as template
            original_path = getattr(self, 'original_docx', None) or getattr(self, 'current_document_path', None)
            
            if original_path and os.path.exists(original_path):
                # Copy original document and replace text - preserves all formatting
                self.log(f"Using original document as template: {os.path.basename(original_path)}")
                shutil.copy2(original_path, file_path)
                
                doc = Document(file_path)
                
                # Helper function to strip all formatting tags (for matching only)
                import re
                def strip_all_tags(text):
                    """Remove all formatting and list tags from text"""
                    text = re.sub(r'</?li-[bo]>', '', text)  # <li-b>, </li-b>, <li-o>, </li-o>
                    text = re.sub(r'</?li>', '', text)       # <li>, </li>
                    text = re.sub(r'</?[biu]>', '', text)    # <b>, </b>, <i>, </i>, <u>, </u>
                    text = re.sub(r'</?bi>', '', text)       # <bi>, </bi>
                    return text.strip()
                
                def clean_special_chars(text):
                    """Remove problematic Unicode characters like object replacement char"""
                    # Remove Unicode Object Replacement Character (U+FFFC) and similar
                    text = text.replace('\ufffc', '')  # Object Replacement Character
                    text = text.replace('\ufffe', '')  # Noncharacter
                    text = text.replace('\uffff', '')  # Noncharacter
                    text = text.replace('\u0000', '')  # Null character
                    return text
                
                def apply_formatted_text_to_paragraph(para, tagged_text):
                    """
                    Replace paragraph text with tagged text, applying bold/italic/underline formatting.
                    Parses tags like <b>, <i>, <u>, <bi> and creates appropriate runs.
                    """
                    # Clean special characters first
                    text = clean_special_chars(tagged_text)
                    
                    # Strip list tags - they don't affect formatting
                    text = re.sub(r'</?li-[bo]>', '', text)
                    text = re.sub(r'</?li>', '', text)
                    
                    # Clear existing runs
                    for run in para.runs:
                        run.clear()
                    # Remove the cleared runs
                    for run in list(para.runs):
                        run._element.getparent().remove(run._element)
                    
                    # Parse tags and create runs with formatting
                    # Pattern matches tags or text between tags
                    tag_pattern = re.compile(r'(</?(?:b|i|u|bi)>)')
                    parts = tag_pattern.split(text)
                    
                    is_bold = False
                    is_italic = False
                    is_underline = False
                    
                    for part in parts:
                        if not part:
                            continue
                        
                        # Check if this is a tag
                        if part == '<b>':
                            is_bold = True
                        elif part == '</b>':
                            is_bold = False
                        elif part == '<i>':
                            is_italic = True
                        elif part == '</i>':
                            is_italic = False
                        elif part == '<u>':
                            is_underline = True
                        elif part == '</u>':
                            is_underline = False
                        elif part == '<bi>':
                            is_bold = True
                            is_italic = True
                        elif part == '</bi>':
                            is_bold = False
                            is_italic = False
                        else:
                            # This is text content - create a run with current formatting
                            if part.strip() or part:  # Include whitespace
                                run = para.add_run(part)
                                run.bold = is_bold
                                run.italic = is_italic
                                run.underline = is_underline
                
                # Build a mapping of source text (without tags) to raw target text (with tags)
                text_map = {}
                for seg in segments:
                    # Strip tags from source for matching against original DOCX
                    source_clean = strip_all_tags(seg.source) if seg.source else ""
                    source_clean = clean_special_chars(source_clean)
                    # Keep raw target text WITH tags for formatting
                    target_raw = seg.target.strip() if seg.target and seg.target.strip() else seg.source
                    if source_clean and target_raw:
                        text_map[source_clean] = target_raw
                
                def replace_segments_in_text(original_text, text_map):
                    """Replace all matching segments in text, handling partial matches."""
                    result = original_text
                    # Sort by length (longest first) to avoid partial replacement issues
                    for source_clean, target_raw in sorted(text_map.items(), key=lambda x: len(x[0]), reverse=True):
                        if source_clean in result:
                            # Get clean target (no tags) for text replacement
                            target_clean = strip_all_tags(target_raw)
                            target_clean = clean_special_chars(target_clean)
                            result = result.replace(source_clean, target_clean)
                    return result
                
                replaced_count = 0
                
                # Replace text in paragraphs (outside tables)
                for para in doc.paragraphs:
                    para_text = para.text.strip()
                    
                    # First try exact match (single segment = whole paragraph)
                    if para_text in text_map:
                        target_text = text_map[para_text]
                        apply_formatted_text_to_paragraph(para, target_text)
                        replaced_count += 1
                    else:
                        # Try partial replacement (paragraph contains multiple segments)
                        new_text = replace_segments_in_text(para_text, text_map)
                        if new_text != para_text:
                            # Text was changed - update paragraph
                            # For partial replacements, we lose formatting tags but at least translate
                            for run in para.runs:
                                run.clear()
                            for run in list(para.runs):
                                run._element.getparent().remove(run._element)
                            para.add_run(new_text)
                            replaced_count += 1
                
                # Replace text in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                para_text = para.text.strip()
                                
                                # First try exact match
                                if para_text in text_map:
                                    target_text = text_map[para_text]
                                    apply_formatted_text_to_paragraph(para, target_text)
                                    replaced_count += 1
                                else:
                                    # Try partial replacement
                                    new_text = replace_segments_in_text(para_text, text_map)
                                    if new_text != para_text:
                                        for run in para.runs:
                                            run.clear()
                                        for run in list(para.runs):
                                            run._element.getparent().remove(run._element)
                                        para.add_run(new_text)
                                        replaced_count += 1
                
                doc.save(file_path)
                self.log(f"✓ Replaced {replaced_count} text segments in original document structure")
                
            else:
                # No original document - create simple paragraph-based export
                self.log("No original document found - creating new document (formatting may differ)")
                
                # Helper function to strip all formatting tags (if not already defined)
                import re
                def strip_all_tags(text):
                    """Remove all formatting and list tags from text"""
                    text = re.sub(r'</?li-[bo]>', '', text)  # <li-b>, </li-b>, <li-o>, </li-o>
                    text = re.sub(r'</?li>', '', text)       # <li>, </li>
                    text = re.sub(r'</?[biu]>', '', text)    # <b>, </b>, <i>, </i>, <u>, </u>
                    text = re.sub(r'</?bi>', '', text)       # <bi>, </bi>
                    return text.strip()
                
                def clean_special_chars(text):
                    """Remove problematic Unicode characters"""
                    text = text.replace('\ufffc', '')  # Object Replacement Character
                    text = text.replace('\ufffe', '')  # Noncharacter
                    text = text.replace('\uffff', '')  # Noncharacter
                    text = text.replace('\u0000', '')  # Null character
                    return text
                
                def add_formatted_text_to_paragraph(para, tagged_text):
                    """Add text with formatting tags to a paragraph."""
                    # Clean special characters first
                    text = clean_special_chars(tagged_text)
                    
                    # Strip list tags
                    text = re.sub(r'</?li-[bo]>', '', text)
                    text = re.sub(r'</?li>', '', text)
                    
                    # Parse and apply formatting
                    tag_pattern = re.compile(r'(</?(?:b|i|u|bi)>)')
                    parts = tag_pattern.split(text)
                    
                    is_bold = False
                    is_italic = False
                    is_underline = False
                    
                    for part in parts:
                        if not part:
                            continue
                        if part == '<b>':
                            is_bold = True
                        elif part == '</b>':
                            is_bold = False
                        elif part == '<i>':
                            is_italic = True
                        elif part == '</i>':
                            is_italic = False
                        elif part == '<u>':
                            is_underline = True
                        elif part == '</u>':
                            is_underline = False
                        elif part == '<bi>':
                            is_bold = True
                            is_italic = True
                        elif part == '</bi>':
                            is_bold = False
                            is_italic = False
                        else:
                            if part:
                                run = para.add_run(part)
                                run.bold = is_bold
                                run.italic = is_italic
                                run.underline = is_underline
                
                doc = Document()
                
                for seg in segments:
                    raw_text = seg.target.strip() if seg.target and seg.target.strip() else seg.source
                    para = doc.add_paragraph()
                    
                    # Try to apply heading style if segment has heading style
                    if hasattr(seg, 'style') and seg.style:
                        style_lower = seg.style.lower()
                        if 'heading 1' in style_lower:
                            para.style = 'Heading 1'
                        elif 'heading 2' in style_lower:
                            para.style = 'Heading 2'
                        elif 'heading 3' in style_lower:
                            para.style = 'Heading 3'
                        elif 'title' in style_lower:
                            para.style = 'Title'
                    
                    # Add text with formatting
                    if raw_text:
                        add_formatted_text_to_paragraph(para, raw_text)
                
                doc.save(file_path)
            
            self.log(f"✓ Exported {len(segments)} segments to: {os.path.basename(file_path)}")
            
            QMessageBox.information(
                self, "Export Complete",
                f"Successfully exported {translated_count} translated segments to:\n\n{os.path.basename(file_path)}"
            )
            
        except ImportError:
            QMessageBox.critical(
                self, "Missing Dependency",
                "The 'python-docx' library is required for DOCX export.\n\n"
                "Install it with: pip install python-docx"
            )
        except Exception as e:
            self.log(f"✗ Export failed: {str(e)}")
            QMessageBox.critical(self, "Export Error", f"Failed to export DOCX:\n\n{str(e)}")
            import traceback
            traceback.print_exc()
    
    def export_review_table_with_tags(self):
        """Export Supervertaler Bilingual Table with formatting tags visible.
        
        This format is intended for proofreaders who will edit and return the file
        for re-import into Supervertaler.
        """
        self._export_review_table(apply_formatting=False)
    
    def export_review_table_formatted(self):
        """Export Supervertaler Bilingual Table with formatting applied (bold, italic, underline).
        
        This format is intended for end clients who want to see the actual formatting
        rather than the tags.
        """
        # Warn user that this format cannot be re-imported with tags
        reply = QMessageBox.warning(
            self, "Formatted Export - Tags Will Be Applied",
            "This export applies formatting (bold, italic, underline) to the text.\n\n"
            "⚠️ Formatting tags will be converted to actual Word formatting.\n"
            "This version CANNOT be re-imported to restore tagged formatting.\n\n"
            "Use 'Bilingual Table - With Tags' if you need to re-import after review.\n\n"
            "Continue with formatted export?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply != QMessageBox.StandardButton.Yes:
            return
        
        self._export_review_table(apply_formatting=True)
    
    def _add_hyperlink_to_paragraph(self, paragraph, url, text):
        """Add a hyperlink to a paragraph in a Word document.
        
        Args:
            paragraph: The paragraph to add the hyperlink to
            url: The URL for the hyperlink
            text: The display text for the hyperlink
            
        Returns:
            The run containing the hyperlink text (for further formatting)
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # Get the document part
        part = paragraph.part
        
        # Create the relationship
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        # Create the hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # Create a new run for the hyperlink text
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Add the text
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(rPr)
        new_run.append(text_elem)
        hyperlink.append(new_run)
        
        # Add hyperlink to paragraph
        paragraph._p.append(hyperlink)
        
        # Return a reference to the run for further formatting
        # We need to return a proper Run object
        from docx.text.run import Run
        return Run(new_run, paragraph)
    
    def _export_review_table(self, apply_formatting=False):
        """Internal method to export Supervertaler Bilingual Table.
        
        Args:
            apply_formatting: If True, apply bold/italic/underline formatting.
                            If False, show raw Supervertaler tags.
        """
        import re
        
        # Check if we have segments
        if not self.current_project or not self.current_project.segments:
            QMessageBox.warning(self, "No Data", "No segments to export")
            return
        
        segments = list(self.current_project.segments)
        
        if not segments:
            QMessageBox.warning(self, "No Data", "No segments to export")
            return
        
        # Determine default filename
        project_name = getattr(self.current_project, 'name', 'project')
        if hasattr(self, 'current_project_path') and self.current_project_path:
            project_name = Path(self.current_project_path).stem
        
        format_suffix = "_bilingual_formatted" if apply_formatting else "_bilingual"
        default_name = f"{project_name}{format_suffix}.docx"
        
        # Get save path
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export Bilingual Table" if not apply_formatting else "Export Bilingual Table (Formatted)",
            default_name,
            "Word Documents (*.docx);;All Files (*.*)"
        )
        
        if not file_path:
            return
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            doc = Document()

            # Set up document margins and landscape orientation
            from docx.enum.section import WD_ORIENT
            sections = doc.sections
            for section in sections:
                # Set landscape orientation for better visualization of long segments
                section.orientation = WD_ORIENT.LANDSCAPE
                # Swap page dimensions for landscape
                section.page_width, section.page_height = section.page_height, section.page_width
                # Set margins
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
            
            # Language code to display name mapping
            LANG_DISPLAY_NAMES = {
                'en': 'English', 'en-US': 'English (US)', 'en-GB': 'English (UK)', 'en-AU': 'English (AU)',
                'nl': 'Dutch', 'nl-NL': 'Dutch (NL)', 'nl-BE': 'Dutch (Belgium)',
                'de': 'German', 'de-DE': 'German (DE)', 'de-AT': 'German (AT)', 'de-CH': 'German (CH)',
                'fr': 'French', 'fr-FR': 'French (FR)', 'fr-CA': 'French (CA)', 'fr-BE': 'French (BE)',
                'es': 'Spanish', 'es-ES': 'Spanish (ES)', 'es-MX': 'Spanish (MX)', 'es-AR': 'Spanish (AR)',
                'it': 'Italian', 'it-IT': 'Italian (IT)',
                'pt': 'Portuguese', 'pt-PT': 'Portuguese (PT)', 'pt-BR': 'Portuguese (BR)',
                'ru': 'Russian', 'zh': 'Chinese', 'zh-CN': 'Chinese (Simplified)', 'zh-TW': 'Chinese (Traditional)',
                'ja': 'Japanese', 'ko': 'Korean', 'ar': 'Arabic', 'pl': 'Polish', 'sv': 'Swedish',
                'da': 'Danish', 'no': 'Norwegian', 'fi': 'Finnish', 'cs': 'Czech', 'hu': 'Hungarian',
                'tr': 'Turkish', 'el': 'Greek', 'he': 'Hebrew', 'th': 'Thai', 'vi': 'Vietnamese',
            }
            
            def get_language_display_name(lang_value):
                """Convert language code or name to display name."""
                if not lang_value:
                    return "Unknown"
                # If it's already a nice display name (contains space or parentheses), use it
                if ' ' in lang_value or '(' in lang_value:
                    return lang_value
                # Look up in our mapping
                if lang_value in LANG_DISPLAY_NAMES:
                    return LANG_DISPLAY_NAMES[lang_value]
                # Try lowercase
                if lang_value.lower() in LANG_DISPLAY_NAMES:
                    return LANG_DISPLAY_NAMES[lang_value.lower()]
                # Capitalize first letter as fallback
                return lang_value.capitalize() if lang_value else "Unknown"
            
            # Get language names for column headers
            source_lang = "Source"
            target_lang = "Target"
            if self.current_project:
                source_lang = get_language_display_name(self.current_project.source_lang) or "Source"
                target_lang = get_language_display_name(self.current_project.target_lang) or "Target"
            
            # === HEADER SECTION ===
            # Add decorative line above title
            header_line = doc.add_paragraph()
            header_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
            line_run = header_line.add_run("━" * 50)
            line_run.font.color.rgb = RGBColor(0, 102, 204)
            line_run.font.size = Pt(10)
            header_line.paragraph_format.space_after = Pt(6)
            
            # Add title with link to Supervertaler website
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.paragraph_format.space_before = Pt(0)
            title.paragraph_format.space_after = Pt(6)
            
            # Add globe emoji and "Supervertaler Bilingual Table" in blue (no link on title)
            globe_run = title.add_run("🌐 ")
            globe_run.font.size = Pt(18)
            
            title_run = title.add_run("Supervertaler Bilingual Table")
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0, 102, 204)
            
            # Add subtitle with clickable website URL
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.paragraph_format.space_before = Pt(0)
            subtitle.paragraph_format.space_after = Pt(6)
            
            # Make Supervertaler.com a clickable link
            url_link = self._add_hyperlink_to_paragraph(subtitle, "https://supervertaler.com/", "Supervertaler.com")
            url_link.font.size = Pt(10)
            url_link.font.color.rgb = RGBColor(0, 102, 204)
            url_link.font.underline = True
            
            # Add decorative line below title
            footer_line = doc.add_paragraph()
            footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
            line_run2 = footer_line.add_run("━" * 50)
            line_run2.font.color.rgb = RGBColor(0, 102, 204)
            line_run2.font.size = Pt(10)
            footer_line.paragraph_format.space_after = Pt(12)
            
            # Add project info in a nice format
            info = doc.add_paragraph()
            info.add_run(f"Project: ").bold = True
            info.add_run(f"{project_name}\n")
            info.add_run(f"Languages: ").bold = True
            info.add_run(f"{source_lang} → {target_lang}\n")
            info.add_run(f"Segments: ").bold = True
            info.add_run(f"{len(segments)}\n")
            info.add_run(f"Exported: ").bold = True
            info.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
            
            # Add important notice
            notice = doc.add_paragraph()
            notice.paragraph_format.space_before = Pt(6)
            if not apply_formatting:
                warning_run = notice.add_run("⚠️ Important: ")
                warning_run.bold = True
                warning_run.font.color.rgb = RGBColor(180, 100, 0)
                notice.add_run("Do not change segment numbers (#) or source text. ").italic = True
                notice.add_run("This file can be re-imported into Supervertaler after proofreading.").italic = True
            else:
                notice.add_run("Note: This version shows applied formatting and is for client delivery or archiving. It cannot be re-imported.").italic = True
            
            # Create table with 5 columns: #, Source Language, Target Language, Status, Notes
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set column widths
            widths = [Inches(0.5), Inches(3.0), Inches(3.0), Inches(0.8), Inches(1.2)]
            for i, width in enumerate(widths):
                for cell in table.columns[i].cells:
                    cell.width = width
            
            # Add header row with actual language names
            header_cells = table.rows[0].cells
            headers = ['#', source_lang, target_lang, 'Status', 'Notes']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # Make header bold and shaded
                for para in header_cells[i].paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.size = Pt(10)
                # Add shading to header
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'DDDDDD')
                header_cells[i]._tc.get_or_add_tcPr().append(shading)
            
            # Helper function to add formatted text to a cell
            def add_formatted_text_to_cell(cell, text, apply_fmt=False):
                """Add text to cell, optionally applying formatting."""
                from docx.shared import RGBColor
                
                # Clear existing content
                for para in cell.paragraphs:
                    para.clear()
                
                if not text:
                    return
                
                para = cell.paragraphs[0]
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                
                if not apply_fmt:
                    # Show raw text with tags in pink text (like in Supervertaler)
                    # Parse to find tags and color them - includes Supervertaler tags AND CAT tool tags
                    # Supervertaler tags: <b>, <i>, <u>, <bi>, <li>, <li-o>, <li-b>
                    # memoQ tags: {1}, [2}, {3], [uicontrol id="..."], {tagname}
                    # Trados tags: <1>, </1>
                    tag_pattern = re.compile(
                        r'('
                        r'</?(?:b|i|u|bi|li|li-[bo])>'  # Supervertaler tags
                        r'|</?[0-9]+>'  # Trados numeric tags: <1>, </1>
                        r'|\[[^}\]]+\}'  # memoQ mixed: [anything}
                        r'|\{[^\[\]]+\]'  # memoQ mixed: {anything]
                        r'|\[[a-zA-Z][^}\]]*\s[^}\]]*\]'  # memoQ content: [tag attr...]
                        r'|\{[a-zA-Z][a-zA-Z0-9_-]*\}'  # memoQ closing: {tagname}
                        r'|\[[0-9]+\]'  # memoQ numeric: [1]
                        r'|\{[0-9]+\}'  # memoQ/Phrase numeric: {1}
                        r')'
                    )
                    parts = tag_pattern.split(text)
                    
                    for part in parts:
                        if not part:
                            continue
                        run = para.add_run(part)
                        run.font.size = Pt(9)
                        
                        # Check if this is a tag - make text memoQ dark red
                        if tag_pattern.match(part):
                            run.font.color.rgb = RGBColor(127, 0, 1)  # memoQ dark red (#7f0001)
                else:
                    # Parse tags and apply formatting
                    # First handle list tags - convert to visible markers
                    text = re.sub(r'<li-b>\s*', '• ', text)
                    text = re.sub(r'<li-o>\s*', '◦ ', text)  # Open circle for nested
                    text = re.sub(r'<li>\s*', '– ', text)
                    text = re.sub(r'</li-[bo]>', '', text)
                    text = re.sub(r'</li>', '', text)
                    
                    # Parse formatting tags
                    tag_pattern = re.compile(r'(</?(?:b|i|u|bi)>)')
                    parts = tag_pattern.split(text)
                    
                    is_bold = False
                    is_italic = False
                    is_underline = False
                    
                    for part in parts:
                        if not part:
                            continue
                        if part == '<b>':
                            is_bold = True
                        elif part == '</b>':
                            is_bold = False
                        elif part == '<i>':
                            is_italic = True
                        elif part == '</i>':
                            is_italic = False
                        elif part == '<u>':
                            is_underline = True
                        elif part == '</u>':
                            is_underline = False
                        elif part == '<bi>':
                            is_bold = True
                            is_italic = True
                        elif part == '</bi>':
                            is_bold = False
                            is_italic = False
                        else:
                            # Regular text
                            run = para.add_run(part)
                            run.font.size = Pt(9)
                            run.bold = is_bold
                            run.italic = is_italic
                            run.underline = is_underline
            
            # Helper to get status display text
            def get_status_display(status):
                """Convert status to user-friendly display text."""
                status_map = {
                    'translated': 'Translated',
                    'tr_confirmed': 'Confirmed',
                    'confirmed': 'Confirmed',
                    'draft': 'Draft',
                    'not_translated': 'Not Translated',
                    'proofread': 'Proofread',
                    'approved': 'Approved',
                    'rejected': 'Rejected',
                    'needs_review': 'Needs Review',
                    'edited': 'Edited',
                    'mt': 'MT',
                    'tm_match': 'TM Match',
                    'fuzzy_match': 'Fuzzy'
                }
                return status_map.get(status, status.replace('_', ' ').title() if status else '')
            
            # Add segment rows
            for i, seg in enumerate(segments):
                row = table.add_row()
                cells = row.cells
                
                # Segment number
                cells[0].text = str(i + 1)
                for para in cells[0].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(9)
                
                # Source text
                source_text = seg.source if hasattr(seg, 'source') else ''
                add_formatted_text_to_cell(cells[1], source_text, apply_formatting)
                
                # Target text
                target_text = seg.target if hasattr(seg, 'target') else ''
                add_formatted_text_to_cell(cells[2], target_text, apply_formatting)
                
                # Status
                status = seg.status if hasattr(seg, 'status') else ''
                cells[3].text = get_status_display(status)
                for para in cells[3].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(8)
                        # Color-code status
                        if status in ('confirmed', 'tr_confirmed', 'proofread', 'approved'):
                            run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                        elif status in ('not_translated', 'rejected'):
                            run.font.color.rgb = RGBColor(200, 0, 0)  # Red
                        elif status in ('draft', 'needs_review'):
                            run.font.color.rgb = RGBColor(200, 100, 0)  # Orange
                
                # Notes column - populate with segment notes if available
                notes_text = seg.notes if hasattr(seg, 'notes') else ''
                cells[4].text = notes_text
                for para in cells[4].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
            
            # Add footer with branding
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.paragraph_format.space_before = Pt(12)
            
            # Add decorative line
            footer_line_run = footer_para.add_run("\n" + "━" * 50 + "\n")
            footer_line_run.font.size = Pt(10)
            footer_line_run.font.color.rgb = RGBColor(0, 102, 204)
            
            # Add footer text: "Supervertaler Bilingual Table | Supervertaler.com"
            footer_text = doc.add_paragraph()
            footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            brand_run = footer_text.add_run("Supervertaler Bilingual Table")
            brand_run.font.size = Pt(9)
            brand_run.font.color.rgb = RGBColor(100, 100, 100)
            
            separator_run = footer_text.add_run(" | ")
            separator_run.font.size = Pt(9)
            separator_run.font.color.rgb = RGBColor(150, 150, 150)
            
            # Add clickable website link in footer
            footer_link = self._add_hyperlink_to_paragraph(footer_text, "https://supervertaler.com/", "Supervertaler.com")
            footer_link.font.size = Pt(9)
            footer_link.font.color.rgb = RGBColor(0, 102, 204)
            footer_link.font.underline = True
            
            # Save the document
            doc.save(file_path)
            
            format_type = "formatted" if apply_formatting else "with tags"
            self.log(f"✓ Exported bilingual table ({format_type}) with {len(segments)} segments to: {Path(file_path).name}")
            
            QMessageBox.information(
                self, "Export Complete",
                f"Successfully exported {len(segments)} segments to bilingual table:\n\n{os.path.basename(file_path)}\n\n"
                + ("This formatted version is suitable for client review." if apply_formatting 
                   else "This version with tags can be re-imported after proofreading.")
            )
            
        except ImportError:
            QMessageBox.critical(
                self, "Missing Dependency",
                "The 'python-docx' library is required for DOCX export.\n\n"
                "Install it with: pip install python-docx"
            )
        except Exception as e:
            self.log(f"✗ Export failed: {str(e)}")
            QMessageBox.critical(self, "Export Error", f"Failed to export bilingual table:\n\n{str(e)}")
            import traceback
            traceback.print_exc()
    
    def export_tmx_from_grid(self):
        """Export all segments from current project grid as TMX file"""
        try:
            if not self.current_project or not self.current_project.segments:
                QMessageBox.warning(self, "No Project", "Please open a project first")
                return
            
            file_path, _ = fdh.get_save_file_name(
                self, 
                "Export Grid as TMX", 
                "TMX Files (*.tmx);;All Files (*.*)"
            )
            
            if not file_path:
                return
            
            source_segments = []
            target_segments = []
            entry_count = 0
            
            # Export all segments from grid (including empty translations)
            self.log("Exporting all segments from grid...")
            for segment in self.current_project.segments:
                source_text = segment.source
                target_text = segment.target if segment.target else ""
                
                if source_text:
                    source_segments.append(source_text)
                    target_segments.append(target_text)
                    entry_count += 1
            
            self.log(f"Found {entry_count} segments in grid")
            
            if entry_count == 0:
                QMessageBox.warning(self, "No Data", "No segments found in grid")
                return
            
            # Generate and save TMX
            from modules.tmx_generator import TMXGenerator
            tmx_generator = TMXGenerator(log_callback=self.log)
            
            source_lang = self.current_project.source_lang or "en"
            target_lang = self.current_project.target_lang or "nl"
            
            tmx_tree = tmx_generator.generate_tmx(
                source_segments=source_segments,
                target_segments=target_segments,
                source_lang=source_lang,
                target_lang=target_lang
            )
            
            if tmx_generator.save_tmx(tmx_tree, file_path):
                self.log(f"✓ Exported TMX from grid: {file_path} ({entry_count} segments)")
                QMessageBox.information(
                    self, 
                    "Export Complete", 
                    f"Grid exported to TMX successfully!\n\n"
                    f"File: {file_path}\n"
                    f"Segments: {entry_count}\n"
                    f"Language pair: {source_lang} → {target_lang}"
                )
            else:
                QMessageBox.warning(self, "Export Error", "Failed to save TMX file")
            
        except Exception as e:
            self.log(f"✗ Error exporting grid as TMX: {e}")
            import traceback
            self.log(traceback.format_exc())
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n\n{str(e)}")
    
    def export_tmx_from_selected(self):
        """Export only selected segments from grid as TMX file"""
        try:
            if not self.current_project or not self.current_project.segments:
                QMessageBox.warning(self, "No Project", "Please open a project first")
                return
            
            # Get selected rows from table
            selected_rows = set()
            for index in self.table.selectedIndexes():
                selected_rows.add(index.row())
            
            if not selected_rows:
                QMessageBox.warning(
                    self, 
                    "No Selection", 
                    "Please select one or more rows in the grid first.\n\n"
                    "Click on row numbers or use Ctrl+Click to select multiple rows."
                )
                return
            
            file_path, _ = QFileDialog.getSaveFileName(
                self, 
                "Export Selected Segments as TMX", 
                "supervertaler_selected.tmx", 
                "TMX Files (*.tmx);;All Files (*.*)"
            )
            
            if not file_path:
                return
            
            source_segments = []
            target_segments = []
            entry_count = 0
            
            # Export only selected segments
            self.log(f"Exporting {len(selected_rows)} selected segments...")
            for row in sorted(selected_rows):
                if row < len(self.current_project.segments):
                    segment = self.current_project.segments[row]
                    source_text = segment.source
                    target_text = segment.target if segment.target else ""
                    
                    if source_text:
                        source_segments.append(source_text)
                        target_segments.append(target_text)
                        entry_count += 1
            
            self.log(f"Selected {entry_count} segments to export")
            
            if entry_count == 0:
                QMessageBox.warning(self, "No Data", "No valid segments selected")
                return
            
            # Generate and save TMX
            from modules.tmx_generator import TMXGenerator
            tmx_generator = TMXGenerator(log_callback=self.log)
            
            source_lang = self.current_project.source_lang or "en"
            target_lang = self.current_project.target_lang or "nl"
            
            tmx_tree = tmx_generator.generate_tmx(
                source_segments=source_segments,
                target_segments=target_segments,
                source_lang=source_lang,
                target_lang=target_lang
            )
            
            if tmx_generator.save_tmx(tmx_tree, file_path):
                self.log(f"✓ Exported TMX from selected: {file_path} ({entry_count} segments)")
                QMessageBox.information(
                    self, 
                    "Export Complete", 
                    f"Selected segments exported to TMX successfully!\n\n"
                    f"File: {file_path}\n"
                    f"Segments: {entry_count}\n"
                    f"Language pair: {source_lang} → {target_lang}"
                )
            else:
                QMessageBox.warning(self, "Export Error", "Failed to save TMX file")
            
        except Exception as e:
            self.log(f"✗ Error exporting selected segments as TMX: {e}")
            import traceback
            self.log(traceback.format_exc())
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n\n{str(e)}")
    
    def export_tmx_from_tm_database(self):
        """Export translation memory entries as TMX file"""
        try:
            if not self.tm_database:
                QMessageBox.warning(self, "Error", "Translation memory database not available")
                return
            
            # Get TM entries from database
            tm_entries = self.tm_database.get_tm_entries(tm_id=None, limit=None)
            if not tm_entries:
                QMessageBox.warning(
                    self, 
                    "No Data", 
                    "No translation units in Translation Memory database.\n\n"
                    "Please import a TM or save translations to the TM first."
                )
                return
            
            file_path, _ = QFileDialog.getSaveFileName(
                self, 
                "Export TM Database as TMX", 
                "supervertaler_tm_database.tmx", 
                "TMX Files (*.tmx);;All Files (*.*)"
            )
            
            if not file_path:
                return
            
            source_segments = []
            target_segments = []
            
            # Extract segments from TM entries
            self.log(f"Exporting {len(tm_entries)} entries from TM database...")
            for entry in tm_entries:
                source_segments.append(entry.get('source_text', entry.get('source', '')))
                target_segments.append(entry.get('target_text', entry.get('target', '')))
            
            # Generate and save TMX
            from modules.tmx_generator import TMXGenerator
            tmx_generator = TMXGenerator(log_callback=self.log)
            
            # Use project languages or defaults
            source_lang = "en"
            target_lang = "nl"
            if self.current_project:
                source_lang = self.current_project.source_lang or "en"
                target_lang = self.current_project.target_lang or "nl"
            
            tmx_tree = tmx_generator.generate_tmx(
                source_segments=source_segments,
                target_segments=target_segments,
                source_lang=source_lang,
                target_lang=target_lang
            )
            
            if tmx_generator.save_tmx(tmx_tree, file_path):
                self.log(f"✓ Exported TMX from TM database: {file_path} ({len(tm_entries)} entries)")
                QMessageBox.information(
                    self, 
                    "Export Complete", 
                    f"Translation Memory exported to TMX successfully!\n\n"
                    f"File: {file_path}\n"
                    f"Entries: {len(tm_entries)}\n"
                    f"Language pair: {source_lang} → {target_lang}"
                )
            else:
                QMessageBox.warning(self, "Export Error", "Failed to save TMX file")
            
        except Exception as e:
            self.log(f"✗ Error exporting TM database as TMX: {e}")
            import traceback
            self.log(traceback.format_exc())
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n\n{str(e)}")
    
    def export_tm_as_tmx(self):
        """Legacy function - exports grid segments (for backward compatibility)"""
        self.export_tmx_from_grid()
    
    
    
    
    def clear_tm_entries(self):
        """Clear all translation memory entries (with confirmation)"""
        try:
            reply = QMessageBox.question(
                self, 
                "Clear Translation Memory", 
                "⚠️ WARNING: This will permanently delete ALL translation memory entries!\n\nThis action cannot be undone. Are you sure?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                # Double confirmation
                reply2 = QMessageBox.question(
                    self, 
                    "Final Confirmation", 
                    "Last chance! Are you absolutely sure you want to delete all TM entries?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                    QMessageBox.StandardButton.No
                )
                
                if reply2 == QMessageBox.StandardButton.Yes and self.tm_database:
                    # Clear TM entries
                    import sqlite3
                    db_path = self.user_data_path / "Translation_Resources" / "supervertaler.db"
                    conn = sqlite3.connect(str(db_path))
                    cursor = conn.cursor()
                    cursor.execute("DELETE FROM translation_units")
                    conn.commit()
                    conn.close()
                    
                    self.log("All translation memory entries cleared")
                    QMessageBox.information(self, "TM Cleared", "All translation memory entries have been deleted.")
            
        except Exception as e:
            self.log(f"Error clearing TM entries: {e}")
            QMessageBox.critical(self, "Clear Error", f"Failed to clear TM entries:\n\n{e}")
    
    def create_segmentation_rules_tab(self):
        """Create Segmentation Rules management tab"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(10, 10, 10, 10)
        
        # Header
        header = QLabel("📏 Segmentation Rules")
        header.setStyleSheet("font-size: 14px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(header)
        
        # Description
        desc = QLabel("Manage language-specific segmentation rules for accurate sentence/segment boundaries.")
        desc.setStyleSheet("color: #666; font-size: 11px; margin-bottom: 10px;")
        layout.addWidget(desc)
        
        # Language Selection
        lang_group = QGroupBox("Select Language")
        lang_layout = QHBoxLayout()
        
        lang_combo = QComboBox()
        languages = ["English (en)", "Dutch (nl)", "German (de)", "French (fr)", "Spanish (es)", "Italian (it)", "Portuguese (pt)", "Chinese (zh)", "Japanese (ja)", "Arabic (ar)"]
        lang_combo.addItems(languages)
        lang_layout.addWidget(QLabel("Language:"))
        lang_layout.addWidget(lang_combo, 1)
        
        lang_group.setLayout(lang_layout)
        layout.addWidget(lang_group)
        
        # Segmentation Rules
        rules_group = QGroupBox("Segmentation Rules")
        rules_layout = QVBoxLayout()
        
        # Current implementation info
        current_info = QLabel(
            "Current Implementation: SimpleSegmenter (language-agnostic)\n\n"
            "• Segments on: . ! ? (followed by space/newline)\n"
            "• Handles basic abbreviations: Mr. Dr. etc.\n"
            "• Preserves paragraph breaks\n"
            "• Treats each table cell as separate segment"
        )
        current_info.setStyleSheet("padding: 10px; border-radius: 4px; font-family: monospace;")
        rules_layout.addWidget(current_info)
        
        # Future implementation section
        future_label = QLabel("🚧 Language-Specific Rules (Planned):")
        future_label.setStyleSheet("font-weight: bold; margin-top: 15px;")
        rules_layout.addWidget(future_label)
        
        future_info = QLabel(
            "• German: Handle compound sentences, different abbreviations\n"
            "• Chinese/Japanese: Word boundary detection, different punctuation\n"
            "• Arabic: Right-to-left text handling\n"
            "• French: Quotation mark handling, spacing rules\n"
            "• Custom: User-defined regex patterns and exceptions"
        )
        future_info.setStyleSheet("color: #666; margin-left: 20px;")
        rules_layout.addWidget(future_info)
        
        rules_group.setLayout(rules_layout)
        layout.addWidget(rules_group)
        
        # Management buttons (for future implementation)
        buttons_group = QGroupBox("Rule Management")
        buttons_layout = QVBoxLayout()
        
        # Test segmentation button
        test_btn = QPushButton("🧪 Test Segmentation")
        test_btn.setToolTip("Test current segmentation rules on sample text")
        test_btn.clicked.connect(self.test_segmentation_rules)
        buttons_layout.addWidget(test_btn)
        
        # Import/Export buttons (disabled for now)
        import_btn = QPushButton("📥 Import Rules")
        import_btn.setToolTip("Import segmentation rules from file (Coming Soon)")
        import_btn.setEnabled(False)
        buttons_layout.addWidget(import_btn)
        
        export_btn = QPushButton("📤 Export Rules")
        export_btn.setToolTip("Export current segmentation rules (Coming Soon)")
        export_btn.setEnabled(False)
        buttons_layout.addWidget(export_btn)
        
        buttons_group.setLayout(buttons_layout)
        layout.addWidget(buttons_group)
        
        layout.addStretch()
        
        return tab
    
    def test_segmentation_rules(self):
        """Test current segmentation rules with sample text"""
        from PyQt6.QtWidgets import QInputDialog
        
        sample_text = (
            "Hello world. This is a test! How are you? "
            "Mr. Smith went to Dr. Jones. "
            "The U.S.A. is great. "
            "What about this... and that? "
            "End of test."
        )
        
        text, ok = QInputDialog.getMultiLineText(
            self, 
            "Test Segmentation", 
            "Enter text to test segmentation:\n(Default sample text provided)",
            sample_text
        )
        
        if ok and text:
            try:
                # Test with current segmenter
                from modules.simple_segmenter import SimpleSegmenter
                segmenter = SimpleSegmenter()
                
                # Create fake paragraph list for testing
                paragraphs = [(0, text)]
                segments = segmenter.segment_paragraphs(paragraphs)
                
                # Show results
                result_text = f"Input text:\n{text}\n\n"
                result_text += f"Segmentation results ({len(segments)} segments):\n"
                result_text += "=" * 50 + "\n"
                
                for i, (para_id, segment_text) in enumerate(segments, 1):
                    result_text += f"{i}. {segment_text}\n"
                
                QMessageBox.information(
                    self,
                    "Segmentation Test Results",
                    result_text
                )
                
            except Exception as e:
                QMessageBox.critical(
                    self,
                    "Segmentation Test Error",
                    f"Error testing segmentation:\n\n{e}"
                )
    
    def 