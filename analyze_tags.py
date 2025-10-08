from docx import Document
from docx.shared import RGBColor

doc = Document('projects/memoQ bilingual.docx')
table = doc.tables[0]

print("=== ANALYZING MEMOQ TAGS ===\n")

# Look at row 10 (index 11) which has tags
row = table.rows[11]
source_cell = row.cells[1]

print(f"Full text: {source_cell.text}")
print("\n=== Individual Runs ===")

for para_idx, para in enumerate(source_cell.paragraphs):
    for run_idx, run in enumerate(para.runs):
        print(f"\nRun {para_idx}.{run_idx}:")
        print(f"  Text: {repr(run.text)}")
        print(f"  Font name: {run.font.name}")
        print(f"  Bold: {run.bold}")
        print(f"  Italic: {run.italic}")
        print(f"  Underline: {run.underline}")
        
        # Check font color
        if run.font.color and hasattr(run.font.color, 'rgb') and run.font.color.rgb:
            rgb = run.font.color.rgb
            print(f"  Color: RGB({rgb[0]}, {rgb[1]}, {rgb[2]})")
        else:
            print(f"  Color: None/Default")
        
        # Check if it's a tag
        is_tag = any(c in run.text for c in ['[', ']', '{', '}'])
        if is_tag:
            print(f"  >>> THIS IS A TAG <<<")

# Also check row 15 which has more tags
print("\n\n=== ROW 15 (More complex tags) ===")
row15 = table.rows[16]
source_cell15 = row15.cells[1]
print(f"Full text: {source_cell15.text}")

print("\n=== Runs ===")
for para in source_cell15.paragraphs:
    for run in para.runs:
        if any(c in run.text for c in ['[', ']', '{', '}', '<', '>']):
            print(f"TAG: {repr(run.text)} | Color: {run.font.color.rgb if run.font.color and hasattr(run.font.color, 'rgb') and run.font.color.rgb else 'None'}")
