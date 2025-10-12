"""Deep dive into tag formatting - check styles and character formatting"""
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

filepath = 'C:/Dev/Supervertaler/user data_private/Trados Studio 2024 bilingual review format test file (with formatting tags).docx'
doc = Document(filepath)
table = doc.tables[0]

# Find first row with tags (row 24)
row = table.rows[24]
source_cell = row.cells[2]

print("=== Deep formatting analysis of tags ===\n")

for para in source_cell.paragraphs:
    for run_idx, run in enumerate(para.runs):
        text = run.text
        if '<' in text or '>' in text:  # This is a tag
            print(f"Tag Run {run_idx}: {repr(text)}")
            print(f"  Font properties:")
            print(f"    name: {run.font.name}")
            print(f"    size: {run.font.size}")
            print(f"    bold: {run.font.bold}")
            print(f"    italic: {run.font.italic}")
            print(f"    underline: {run.font.underline}")
            print(f"    color.rgb: {run.font.color.rgb if hasattr(run.font.color, 'rgb') else 'N/A'}")
            print(f"    color.theme_color: {run.font.color.theme_color if hasattr(run.font.color, 'theme_color') else 'N/A'}")
            print(f"    highlight_color: {run.font.highlight_color}")
            
            # Check style
            print(f"  Style: {run.style}")
            
            # Check XML for hidden formatting
            print(f"  XML properties:")
            rPr = run._element.rPr
            if rPr is not None:
                print(f"    rPr exists: {rPr.xml}")
                # Check for color in XML
                color_elem = rPr.find(qn('w:color'))
                if color_elem is not None:
                    print(f"    color val: {color_elem.get(qn('w:val'))}")
                    print(f"    color themeColor: {color_elem.get(qn('w:themeColor'))}")
                # Check for highlight
                highlight_elem = rPr.find(qn('w:highlight'))
                if highlight_elem is not None:
                    print(f"    highlight val: {highlight_elem.get(qn('w:val'))}")
                # Check for shading
                shd_elem = rPr.find(qn('w:shd'))
                if shd_elem is not None:
                    print(f"    shading fill: {shd_elem.get(qn('w:fill'))}")
            else:
                print(f"    No rPr element")
            print()
