"""Deep dive into actual target cell content - byte-level comparison"""
from docx import Document
from docx.oxml import parse_xml
import zipfile

trados_file = r'C:\Dev\Supervertaler\user data_private\Projects\trados test\trados-test\External Review\en-US\trados-test.docx.review.docx'
supervertaler_file = r'C:\Dev\Supervertaler\user data_private\Projects\trados test\trados-test\External Review\en-US\trados_bilingual_export.review.docx'

def analyze_target_cells(filepath, label):
    """Analyze target cells in detail"""
    print(f"\n{'='*80}")
    print(f"{label}")
    print(f"{'='*80}")
    
    doc = Document(filepath)
    table = doc.tables[0]
    
    # Check first few rows
    for row_idx in range(1, min(4, len(table.rows))):
        row = table.rows[row_idx]
        target_cell = row.cells[3]
        
        print(f"\n📍 Row {row_idx} - Segment ID: {row.cells[0].text[:40]}...")
        print(f"   Target text: {target_cell.text[:100]}...")
        print(f"   Number of paragraphs: {len(target_cell.paragraphs)}")
        
        for para_idx, para in enumerate(target_cell.paragraphs):
            print(f"   Paragraph {para_idx}:")
            print(f"      Text: {para.text[:80]}...")
            print(f"      Number of runs: {len(para.runs)}")
            print(f"      Style: {para.style.name if para.style else 'None'}")
            
            # Check runs
            for run_idx, run in enumerate(para.runs[:3]):  # First 3 runs
                print(f"      Run {run_idx}: {repr(run.text[:30])}... ")
                print(f"         Style: {run.style.name if run.style else 'None'}")

# Analyze both
analyze_target_cells(trados_file, "TRADOS FILE - TARGET CELLS")
analyze_target_cells(supervertaler_file, "SUPERVERTALER FILE - TARGET CELLS")

# Now check for actual XML differences in document.xml
print(f"\n\n{'='*80}")
print("RAW DOCUMENT.XML COMPARISON (First 5000 bytes)")
print(f"{'='*80}")

with zipfile.ZipFile(trados_file, 'r') as z1:
    trados_doc_xml = z1.read('word/document.xml')

with zipfile.ZipFile(supervertaler_file, 'r') as z2:
    super_doc_xml = z2.read('word/document.xml')

print(f"\n📊 File sizes:")
print(f"   Trados: {len(trados_doc_xml)} bytes")
print(f"   Supervertaler: {len(super_doc_xml)} bytes")

if trados_doc_xml == super_doc_xml:
    print(f"\n✅ document.xml files are BYTE-FOR-BYTE IDENTICAL!")
else:
    print(f"\n❌ document.xml files are DIFFERENT")
    print(f"\nFinding first difference...")
    
    for i, (b1, b2) in enumerate(zip(trados_doc_xml, super_doc_xml)):
        if b1 != b2:
            print(f"First difference at byte {i}:")
            print(f"   Trados:        byte={b1} char={chr(b1) if 32 <= b1 < 127 else '?'}")
            print(f"   Supervertaler: byte={b2} char={chr(b2) if 32 <= b2 < 127 else '?'}")
            print(f"\n   Context (Trados): ...{trados_doc_xml[max(0,i-50):i+50].decode('utf-8', errors='ignore')}...")
            print(f"\n   Context (Super):  ...{super_doc_xml[max(0,i-50):i+50].decode('utf-8', errors='ignore')}...")
            break
    
    # Check length difference
    if len(trados_doc_xml) != len(super_doc_xml):
        print(f"\nLength difference: {len(super_doc_xml) - len(trados_doc_xml)} bytes")
