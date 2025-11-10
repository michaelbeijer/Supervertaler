"""
Verify that pipe symbols in CafeTran exported DOCX are formatted as bold and red.
"""

from docx import Document
from docx.shared import RGBColor

# Load the test output file
doc = Document(r'projects\test_document (CafeTran bilingual docx)_TEST_OUTPUT.docx')
table = doc.tables[0]

print("Checking pipe symbol formatting in target cells:\n")
print("="*70)

# Check a few rows with pipe symbols
test_rows = [
    (10, "Naples|, Italy"),
    (14, "|Atalanta|"),
    (15, "|Juventus FC| ... |Lumezzane| ... |2002-03 Serie C1|")
]

for row_idx, expected_content in test_rows:
    print(f"\nRow {row_idx} - Expected: {expected_content}")
    print("-"*70)
    
    cell = table.rows[row_idx].cells[2]  # Target column
    
    if not cell.paragraphs:
        print("  ⚠ No paragraphs in cell")
        continue
    
    para = cell.paragraphs[0]
    
    if not para.runs:
        print("  ⚠ No runs in paragraph")
        continue
    
    print(f"  Found {len(para.runs)} runs:")
    
    pipe_count = 0
    for i, run in enumerate(para.runs):
        text = run.text
        is_pipe = text == '|'
        
        if is_pipe:
            pipe_count += 1
            is_bold = run.bold == True
            color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
            is_red = color == RGBColor(255, 0, 0) if color else False
            
            status = "✓" if (is_bold and is_red) else "✗"
            print(f"    {status} Run {i}: \"{text}\" | Bold: {is_bold} | Color: {color} | Red: {is_red}")
        else:
            # Show first few regular text runs
            if i < 3 or '|' in ''.join([r.text for r in para.runs[max(0, i-1):i+2]]):
                print(f"      Run {i}: \"{text[:30]}...\" (regular text)")
    
    print(f"  Total pipes found: {pipe_count}")

print("\n" + "="*70)
print("Verification complete!")
