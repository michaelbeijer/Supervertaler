"""Verify what was actually written to the export file"""
from docx import Document

output_file = 'C:/Dev/Supervertaler/user data_private/test_tag_styling_export.docx'

print("=== Checking exported file ===\n")
doc = Document(output_file)
table = doc.tables[0]

# Check rows around 24-26 (should have tags)
for row_idx in [24, 25, 26, 208, 209]:
    if row_idx >= len(table.rows):
        continue
    row = table.rows[row_idx]
    seg_id = row.cells[0].text.strip()
    status = row.cells[1].text.strip()
    target = row.cells[3].text.strip()
    
    print(f"Row {row_idx}:")
    print(f"  ID: {seg_id}")
    print(f"  Status: {status}")
    print(f"  Target: {target[:80]}...")
    
    # Check if target has tags
    if '<' in target and '>' in target:
        print(f"  ✅ Has tags!")
        target_cell = row.cells[3]
        for para in target_cell.paragraphs:
            for run in para.runs:
                if '<' in run.text or '>' in run.text:
                    print(f"    Tag run: {repr(run.text)}")
                    print(f"      Style: {run.style.name if run.style else 'None'}")
                    print(f"      Italic: {run.italic}")
                    try:
                        print(f"      Color: {run.font.color.rgb}")
                    except:
                        print(f"      Color: (error reading)")
    print()
