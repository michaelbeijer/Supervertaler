"""Check the exact Word formatting of tags in Trados file"""
from docx import Document
from docx.shared import RGBColor

filepath = 'C:/Dev/Supervertaler/user data_private/Trados Studio 2024 bilingual review format test file (with formatting tags).docx'
doc = Document(filepath)
table = doc.tables[0]

# Find rows with tags
print("=== Searching for rows with tags ===")
for row_idx, row in enumerate(table.rows):
    source_text = row.cells[2].text
    if '<' in source_text and '>' in source_text:
        print(f"\nRow {row_idx} has tags:")
        print(f"  Source: {source_text[:100]}...")
        print(f"  Target: {row.cells[3].text[:100]}...")

# Now examine the first row with tags in detail
print("\n\n=== Detailed analysis of first row with tags ===")
for row_idx, row in enumerate(table.rows):
    source_text = row.cells[2].text
    if '<' in source_text and '>' in source_text:
        print(f"\nAnalyzing Row {row_idx}:")
        source_cell = row.cells[2]
        
        for para_idx, para in enumerate(source_cell.paragraphs):
            print(f"\n  Source Paragraph {para_idx}:")
            for run_idx, run in enumerate(para.runs):
                text = run.text
                font = run.font
                
                # Get color
                color_info = "None"
                try:
                    if font.color.type == 1:  # RGB color
                        rgb = font.color.rgb
                        if rgb:
                            color_info = f"RGB({rgb[0]}, {rgb[1]}, {rgb[2]})"
                    elif font.color.type == 2:  # Theme color
                        color_info = f"Theme({font.color.theme_color})"
                except:
                    color_info = "Error reading color"
                
                print(f"    Run {run_idx}:")
                print(f"      Text: {repr(text)}")
                print(f"      Font Name: {font.name}")
                print(f"      Font Size: {font.size}")
                print(f"      Color: {color_info}")
                print(f"      Bold: {font.bold}")
                print(f"      Italic: {font.italic}")
        
        # Also check target
        print(f"\n  Target Paragraph:")
        target_cell = row.cells[3]
        for para_idx, para in enumerate(target_cell.paragraphs):
            for run_idx, run in enumerate(para.runs):
                text = run.text
                font = run.font
                
                color_info = "None"
                try:
                    if font.color.type == 1:
                        rgb = font.color.rgb
                        if rgb:
                            color_info = f"RGB({rgb[0]}, {rgb[1]}, {rgb[2]})"
                    elif font.color.type == 2:
                        color_info = f"Theme({font.color.theme_color})"
                except:
                    color_info = "Error"
                
                print(f"    Run {run_idx}:")
                print(f"      Text: {repr(text)}")
                print(f"      Color: {color_info}")
                print(f"      Bold: {font.bold}")
        
        break  # Only analyze first row with tags

