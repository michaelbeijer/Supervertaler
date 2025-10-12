"""Compare tag styling between original and exported files"""
from docx import Document
from docx.oxml.ns import qn

def analyze_tag_runs(filepath, label):
    """Analyze tag runs in a file"""
    print(f"\n=== {label} ===")
    doc = Document(filepath)
    table = doc.tables[0]
    
    # Check row 24 (has <13> tags)
    row = table.rows[24]
    target_cell = row.cells[3]
    
    print(f"Target text: {target_cell.text[:80]}...")
    print(f"\nTag runs:")
    
    for para in target_cell.paragraphs:
        for run in para.runs:
            if '<' in run.text or '>' in run.text:
                print(f"\n  Text: {repr(run.text)}")
                
                # Check style
                print(f"  Style name: {run.style.name if run.style else 'None'}")
                
                # Check XML for rStyle
                rPr = run._element.rPr
                if rPr is not None:
                    rStyle = rPr.find(qn('w:rStyle'))
                    if rStyle is not None:
                        print(f"  XML rStyle: {rStyle.get(qn('w:val'))}")
                    
                    # Check color in XML
                    color_elem = rPr.find(qn('w:color'))
                    if color_elem is not None:
                        print(f"  XML color: {color_elem.get(qn('w:val'))}")
                    
                    # Check italic
                    italic_elem = rPr.find(qn('w:i'))
                    if italic_elem is not None:
                        print(f"  XML italic: Yes")
                
                # Check Python API values
                print(f"  font.italic: {run.font.italic}")
                try:
                    rgb = run.font.color.rgb
                    if rgb:
                        print(f"  font.color.rgb: {rgb}")
                except:
                    pass

# Compare files
original_file = 'C:/Dev/Supervertaler/user data_private/Trados Studio 2024 bilingual review format test file (with formatting tags).docx'
exported_file = 'C:/Dev/Supervertaler/user data_private/test_tag_styling_export.docx'

analyze_tag_runs(original_file, "ORIGINAL FILE")
analyze_tag_runs(exported_file, "EXPORTED FILE")

print("\n\n" + "="*70)
print("SUMMARY:")
print("="*70)
print("✅ Both files should have:")
print("   - Style name: Tag")
print("   - XML rStyle: Tag")
print("   - XML color: FF0066")
print("   - XML italic: Yes")
print("\nIf both match, the tag styling is PERFECT!")
