"""Test the fixed tag handling with real Trados file"""
import sys
sys.path.insert(0, 'C:/Dev/Supervertaler')

from modules.trados_docx_handler import TradosDOCXHandler
from docx import Document

# The test file from user
test_file = r'C:\Dev\Supervertaler\user data_private\Projects\trados test\Trados test project (test_document.docx)\External Review\nl-BE\test_document.docx.review.docx'
output_file = r'C:\Dev\Supervertaler\user data_private\test_tag_fix_verification.docx'

print("Testing tag fix with real Trados file...")

# Create test segment with tags
test_segments = [
    {
        'id': '1c0c7aafe-a056-4f9b-ad8e-7c8b96da134c',  # Segment ID from the file
        'status': 'Translated (100%)',
        'target': '<14>Biagio Pagano</14> (geboren op 29 januari 1983) is een Italiaanse voetballer die momenteel als middenvelder speelt voor Ghivizzano Borgoamozzano.'
    }
]

print(f"\nInput text: {test_segments[0]['target']}")

# Export
try:
    TradosDOCXHandler.update_bilingual_docx(test_file, test_segments, output_file)
    print("\n✅ Export successful!")
    
    # Verify the output
    doc = Document(output_file)
    table = doc.tables[0]
    
    # Find the row with this segment
    for row in table.rows[1:]:
        if '1c0c7aafe' in row.cells[0].text:
            target_text = row.cells[3].text
            print(f"\nOutput text: {target_text}")
            
            # Check for the bug
            if '<14>14>' in target_text or '</14>14>' in target_text:
                print("\n❌ BUG STILL PRESENT! Tags are duplicated!")
            elif '<14>Biagio Pagano</14>' in target_text:
                print("\n✅ BUG FIXED! Tags are correct!")
            else:
                print(f"\n⚠️  Unexpected output")
            
            # Check runs to see tag styling
            print(f"\nTarget cell runs:")
            for para in row.cells[3].paragraphs:
                for run in para.runs:
                    if '<' in run.text or '>' in run.text:
                        print(f"  Tag run: {repr(run.text)} - Style: {run.style.name if run.style else 'None'}")
            break
    
except Exception as e:
    print(f"\n❌ Error: {e}")
    import traceback
    traceback.print_exc()
