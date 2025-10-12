"""Test the tag styling in exported document"""
import sys
sys.path.insert(0, 'C:/Dev/Supervertaler')

from modules.trados_docx_handler import TradosDOCXHandler
from docx import Document

# Test data with Trados tags
test_segments = [
    {
        'id': '171a00032-da93-4bc2-ac51-65ef1f6ae67d',
        'status': 'Translated (100%)',
        'source': 'Uitklapbaar geraamte voor opklapmeubel',
        'target': 'Folding framework for folding furniture'
    },
    {
        'id': '171a00032-da93-4bc2-ac51-65ef1f6ae67e',
        'status': 'Not Translated (0%)',
        'source': 'TECHNISCH DOMEIN',
        'target': 'TECHNICAL FIELD'
    },
    {
        'id': '4b3ad998-d989-453e-a7a8-a4d3e8e99d3c',
        'status': 'Translated (100%)',
        'source': '<13>Figuur 1</13> toont een zijdelingse weergave',
        'target': '<13>Figure 1</13> shows a lateral view'
    },
    {
        'id': '5550936e3bfe-0a3e-40d3-8000-55599f643fea',
        'status': 'Translated (100%)',
        'source': 'een <233>eerste hoofdsteunkoppeling (3) <231/></233>met een eerste uiteinde',
        'target': 'a <233>first head support coupling (3) <231/></233>with a first end'
    }
]

# Paths
original_file = 'C:/Dev/Supervertaler/user data_private/Trados Studio 2024 bilingual review format test file (with formatting tags).docx'
output_file = 'C:/Dev/Supervertaler/user data_private/test_tag_styling_export.docx'

print("Testing tag style preservation...")
print(f"Original file: {original_file}")
print(f"Output file: {output_file}")

try:
    # Update the bilingual document
    TradosDOCXHandler.update_bilingual_docx(original_file, test_segments, output_file)
    print("\n✅ Export successful!")
    
    # Verify tag styling in exported file
    print("\n=== Verifying tag styles in exported file ===")
    doc = Document(output_file)
    table = doc.tables[0]
    
    # Check row with UUID '4b3ad998-d989-453e-a7a8-a4d3e8e99d3c' (has <13> tags)
    for row in table.rows[1:]:
        if '4b3ad998-d989-453e-a7a8-a4d3e8e99d3c' in row.cells[0].text:
            print("\nFound row with <13> tags:")
            target_cell = row.cells[3]
            for para in target_cell.paragraphs:
                for run in para.runs:
                    text = run.text
                    if '<' in text or '>' in text:
                        print(f"  Tag run: {repr(text)}")
                        print(f"    Style: {run.style}")
                        print(f"    Italic: {run.italic}")
                        print(f"    Color RGB: {run.font.color.rgb}")
            break
    
    print("\n✅ Tag styling verification complete!")
    print(f"\nNow open '{output_file}' and compare with original to verify visual match.")
    
except Exception as e:
    print(f"\n❌ Error: {e}")
    import traceback
    traceback.print_exc()
