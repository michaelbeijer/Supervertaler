"""Deep comparison of Trados-exported vs Supervertaler-exported files"""
from docx import Document
from docx.oxml.ns import qn
import sys

trados_file = r'C:\Dev\Supervertaler\user data_private\Projects\trados test\trados-test\External Review\en-US\trados-test.docx.review.docx'
supervertaler_file = r'C:\Dev\Supervertaler\user data_private\Projects\trados test\trados-test\External Review\en-US\trados_bilingual_export.review.docx'

def analyze_document_structure(filepath, label):
    """Analyze complete document structure"""
    print(f"\n{'='*80}")
    print(f"{label}")
    print(f"{'='*80}")
    
    doc = Document(filepath)
    
    # Document-level analysis
    print(f"\n📄 DOCUMENT STRUCTURE:")
    print(f"   Number of tables: {len(doc.tables)}")
    print(f"   Number of paragraphs: {len(doc.paragraphs)}")
    print(f"   Number of sections: {len(doc.sections)}")
    
    # Check title paragraph
    if doc.paragraphs:
        title = doc.paragraphs[0]
        print(f"\n📝 TITLE PARAGRAPH:")
        print(f"   Text: {title.text}")
        print(f"   Runs: {len(title.runs)}")
        if title.runs:
            run = title.runs[0]
            print(f"   First run bold: {run.bold}")
            print(f"   First run font size: {run.font.size}")
            print(f"   First run font name: {run.font.name}")
    
    # Table structure
    if doc.tables:
        table = doc.tables[0]
        print(f"\n📊 TABLE STRUCTURE:")
        print(f"   Rows: {len(table.rows)}")
        print(f"   Columns: {len(table.columns)}")
        print(f"   Table style: {table.style.name if table.style else 'None'}")
        
        # Header row
        print(f"\n   HEADER ROW:")
        header_row = table.rows[0]
        for col_idx, cell in enumerate(header_row.cells):
            print(f"      Col {col_idx}: {cell.text}")
            if cell.paragraphs and cell.paragraphs[0].runs:
                run = cell.paragraphs[0].runs[0]
                print(f"         Bold: {run.bold}")
        
        # Check first data row
        print(f"\n   FIRST DATA ROW (Row 1):")
        if len(table.rows) > 1:
            row = table.rows[1]
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                print(f"      Col {col_idx}: {text[:50]}...")
        
        # Check a row with tags (if exists)
        print(f"\n   CHECKING FOR TAG-CONTAINING ROWS:")
        found_tags = False
        for row_idx, row in enumerate(table.rows[1:21], 1):  # Check first 20 rows
            source_text = row.cells[2].text if len(row.cells) > 2 else ""
            target_text = row.cells[3].text if len(row.cells) > 3 else ""
            if '<' in source_text or '<' in target_text:
                print(f"      Row {row_idx} has tags")
                if not found_tags:
                    # Detailed analysis of first row with tags
                    print(f"         Target: {target_text[:80]}...")
                    target_cell = row.cells[3]
                    for para in target_cell.paragraphs:
                        for run in para.runs:
                            if '<' in run.text or '>' in run.text:
                                print(f"         Tag run: {repr(run.text)}")
                                print(f"            Style: {run.style.name if run.style else 'None'}")
                                
                                # Check XML
                                rPr = run._element.rPr
                                if rPr is not None:
                                    rStyle = rPr.find(qn('w:rStyle'))
                                    if rStyle is not None:
                                        print(f"            XML rStyle: {rStyle.get(qn('w:val'))}")
                    found_tags = True
                if row_idx >= 3:
                    break
    
    # Check available styles
    print(f"\n🎨 AVAILABLE STYLES:")
    tag_style_found = False
    for style in doc.styles:
        if 'tag' in style.name.lower():
            print(f"   ✅ Found: {style.name} (ID: {style.style_id})")
            tag_style_found = True
            
            # Get style properties
            if hasattr(style, 'font'):
                print(f"      Font properties:")
                print(f"         Italic: {style.font.italic}")
                try:
                    rgb = style.font.color.rgb
                    if rgb:
                        print(f"         Color RGB: {rgb}")
                except:
                    pass
    
    if not tag_style_found:
        print(f"   ❌ No 'Tag' style found!")
    
    # Check XML package structure
    print(f"\n📦 XML PACKAGE STRUCTURE:")
    try:
        # Check for custom XML parts
        package = doc.part.package
        print(f"   Package parts count: {len(package.parts)}")
        
        # Look for Trados-specific parts
        for rel in doc.part.rels.values():
            if 'sdl' in rel.target_ref.lower() or 'trados' in rel.target_ref.lower():
                print(f"   ⚠️ Trados-specific relationship: {rel.target_ref}")
    except Exception as e:
        print(f"   Error checking package: {e}")
    
    return doc

# Analyze both files
print("\n" + "🔍 ANALYZING TRADOS-EXPORTED FILE (REFERENCE)")
trados_doc = analyze_document_structure(trados_file, "TRADOS-EXPORTED FILE")

print("\n\n" + "🔍 ANALYZING SUPERVERTALER-EXPORTED FILE")
super_doc = analyze_document_structure(supervertaler_file, "SUPERVERTALER-EXPORTED FILE")

# Direct comparison
print(f"\n\n{'='*80}")
print("🔍 DIRECT COMPARISON")
print(f"{'='*80}")

trados_table = trados_doc.tables[0]
super_table = super_doc.tables[0]

print(f"\nTable rows: Trados={len(trados_table.rows)}, Supervertaler={len(super_table.rows)}")
print(f"Table style: Trados={trados_table.style.name if trados_table.style else 'None'}, Supervertaler={super_table.style.name if super_table.style else 'None'}")

# Compare first few segment IDs
print(f"\nFirst 5 Segment IDs comparison:")
for i in range(1, min(6, len(trados_table.rows), len(super_table.rows))):
    trados_id = trados_table.rows[i].cells[0].text.strip()
    super_id = super_table.rows[i].cells[0].text.strip()
    match = "✅" if trados_id == super_id else "❌"
    print(f"  Row {i}: {match}")
    if trados_id != super_id:
        print(f"     Trados: {trados_id}")
        print(f"     Super:  {super_id}")

# Compare header
print(f"\nHeader comparison:")
for i in range(4):
    trados_header = trados_table.rows[0].cells[i].text.strip()
    super_header = super_table.rows[0].cells[i].text.strip()
    match = "✅" if trados_header == super_header else "❌"
    print(f"  Col {i}: {match} Trados='{trados_header}' vs Super='{super_header}'")

print("\n" + "="*80)
print("ANALYSIS COMPLETE")
print("="*80)
