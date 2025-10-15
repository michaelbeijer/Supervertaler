"""
PDF Rescue Module
Embeddable version of the AI-powered OCR tool for extracting text from poorly formatted PDFs
Uses OpenAI's GPT-4 Vision API

This module can be embedded in the main Supervertaler application as a tab.
"""

import os
import base64
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
from openai import OpenAI
from docx import Document
from docx.shared import Pt


class PDFRescue:
    """
    PDF Rescue feature - extract text from images using AI OCR
    Can be embedded in any tkinter application as a tab or panel
    """
    
    def __init__(self, parent_app):
        """
        Initialize PDF Rescue module
        
        Args:
            parent_app: Reference to the main application (needs .api_keys attribute)
        """
        self.parent_app = parent_app
        self.client = None
        self.image_files = []
        self.extracted_texts = {}
        
        # Initialize OpenAI client
        api_key = None
        if hasattr(parent_app, 'api_keys'):
            api_key = parent_app.api_keys.get('openai')
        elif hasattr(parent_app, 'api_key'):
            api_key = parent_app.api_key
            
        if api_key:
            try:
                self.client = OpenAI(api_key=api_key)
            except Exception as e:
                print(f"Failed to initialize OpenAI client: {e}")
    
    def create_tab(self, parent):
        """
        Create the PDF Rescue tab UI
        
        Args:
            parent: The parent widget (notebook tab or frame)
        """
        # Header
        header_frame = tk.Frame(parent, bg='#e3f2fd', relief='solid', borderwidth=1)
        header_frame.pack(fill='x', padx=5, pady=5)
        
        tk.Label(header_frame, text="🔍 PDF Rescue - AI-Powered OCR", 
                font=('Segoe UI', 10, 'bold'), bg='#e3f2fd').pack(side='left', padx=10, pady=5)
        
        tk.Label(header_frame, text="Extract text from poorly formatted PDF screenshots",
                font=('Segoe UI', 9), bg='#e3f2fd', fg='#666').pack(side='left', padx=(0, 10), pady=5)
        
        # Split view: Files on left, Preview on right
        paned = ttk.PanedWindow(parent, orient='horizontal')
        paned.pack(fill='both', expand=True, padx=5, pady=5)
        
        # LEFT: File list
        left_frame = tk.Frame(paned)
        paned.add(left_frame, weight=1)
        
        tk.Label(left_frame, text="Images to Process", 
                font=('Segoe UI', 9, 'bold')).pack(anchor='w', pady=(0, 5))
        
        # File list with scrollbar
        list_container = tk.Frame(left_frame)
        list_container.pack(fill='both', expand=True)
        
        scroll = tk.Scrollbar(list_container, orient='vertical')
        scroll.pack(side='right', fill='y')
        
        self.file_listbox = tk.Listbox(list_container, yscrollcommand=scroll.set,
                                       font=('Consolas', 9))
        self.file_listbox.pack(fill='both', expand=True)
        scroll.config(command=self.file_listbox.yview)
        self.file_listbox.bind('<<ListboxSelect>>', self._on_file_select)
        
        # Buttons
        btn_frame = tk.Frame(left_frame)
        btn_frame.pack(fill='x', pady=(10, 0))
        
        tk.Button(btn_frame, text="📁 Add Files", command=self._add_files,
                 bg='#2196F3', fg='white', font=('Segoe UI', 8, 'bold'),
                 padx=8, pady=4).pack(side='left', padx=(0, 3))
        
        tk.Button(btn_frame, text="📂 Folder", command=self._add_folder,
                 bg='#2196F3', fg='white', font=('Segoe UI', 8, 'bold'),
                 padx=8, pady=4).pack(side='left', padx=3)
        
        tk.Button(btn_frame, text="Clear", command=self._clear_list,
                 bg='#9E9E9E', fg='white', font=('Segoe UI', 8),
                 padx=8, pady=4).pack(side='left', padx=3)
        
        # RIGHT: Text preview
        right_frame = tk.Frame(paned)
        paned.add(right_frame, weight=2)
        
        tk.Label(right_frame, text="Extracted Text Preview", 
                font=('Segoe UI', 9, 'bold')).pack(anchor='w', pady=(0, 5))
        
        self.preview_text = scrolledtext.ScrolledText(right_frame, wrap='word',
                                                      font=('Segoe UI', 9),
                                                      height=15)
        self.preview_text.pack(fill='both', expand=True)
        
        # Processing options
        options_frame = tk.LabelFrame(parent, text="Processing Options", 
                                     padx=10, pady=10)
        options_frame.pack(fill='x', padx=5, pady=(0, 10))
        
        # Model selection
        model_frame = tk.Frame(options_frame)
        model_frame.pack(fill='x', pady=(0, 5))
        
        tk.Label(model_frame, text="Model:", font=('Segoe UI', 9)).pack(side='left', padx=(0, 5))
        self.model_var = tk.StringVar(value="gpt-4o")
        models = ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo"]
        ttk.Combobox(model_frame, textvariable=self.model_var, values=models,
                    width=20, state='readonly').pack(side='left')
        
        # Custom instructions
        tk.Label(options_frame, text="Extraction Instructions:", 
                font=('Segoe UI', 9)).pack(anchor='w', pady=(5, 2))
        
        self.instructions_text = scrolledtext.ScrolledText(options_frame, wrap='word',
                                                          font=('Segoe UI', 9),
                                                          height=3)
        self.instructions_text.pack(fill='x')
        
        default_instructions = """Extract all text from this image. The image is a screenshot from a poorly formatted PDF.
Please:
- Extract all visible text accurately
- Fix any obvious OCR errors or formatting issues
- Remove extraneous line breaks within paragraphs
- Preserve intentional paragraph breaks
- Maintain the logical flow and structure of the content
- Output clean, readable text only (no commentary)"""
        
        self.instructions_text.insert('1.0', default_instructions)
        
        # Action buttons
        action_frame = tk.Frame(parent, bg='white')
        action_frame.pack(fill='x', padx=5, pady=(0, 10))
        
        tk.Button(action_frame, text="🔍 Process Selected", 
                 command=self._process_selected,
                 bg='#FF9800', fg='white', font=('Segoe UI', 9, 'bold'),
                 padx=15, pady=6).pack(side='left', padx=(0, 5))
        
        tk.Button(action_frame, text="⚡ Process ALL", 
                 command=self._process_all,
                 bg='#4CAF50', fg='white', font=('Segoe UI', 9, 'bold'),
                 padx=15, pady=6).pack(side='left', padx=5)
        
        tk.Button(action_frame, text="💾 Save DOCX", 
                 command=self._save_to_docx,
                 bg='#2196F3', fg='white', font=('Segoe UI', 9, 'bold'),
                 padx=15, pady=6).pack(side='left', padx=5)
        
        tk.Button(action_frame, text="📋 Copy All", 
                 command=self._copy_all_text,
                 bg='#607D8B', fg='white', font=('Segoe UI', 9, 'bold'),
                 padx=15, pady=6).pack(side='left', padx=5)
        
        # Status
        self.status_label = tk.Label(parent, text="Ready - Add images to begin", 
                                     font=('Segoe UI', 9), fg='#666', bg='white')
        self.status_label.pack(pady=(0, 5))
        
        # Progress bar
        self.progress = ttk.Progressbar(parent, mode='determinate')
        self.progress.pack(fill='x', padx=5, pady=(0, 5))
    
    # === File Management Methods ===
    
    def _add_files(self):
        """Add individual image files"""
        files = filedialog.askopenfilenames(
            title="Select Image Files",
            filetypes=[
                ("Image files", "*.jpg *.jpeg *.png *.bmp *.gif *.tiff"),
                ("All files", "*.*")
            ]
        )
        
        if files:
            for file in files:
                if file not in self.image_files:
                    self.image_files.append(file)
            self._update_listbox()
            self.status_label.config(text=f"Added {len(files)} file(s)")
    
    def _add_folder(self):
        """Add all images from a folder"""
        folder = filedialog.askdirectory(title="Select Folder with Images")
        
        if folder:
            image_extensions = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff'}
            files = []
            
            for file in sorted(os.listdir(folder)):
                file_path = os.path.join(folder, file)
                if os.path.isfile(file_path):
                    ext = os.path.splitext(file)[1].lower()
                    if ext in image_extensions and file_path not in self.image_files:
                        files.append(file_path)
            
            self.image_files.extend(files)
            self._update_listbox()
            self.status_label.config(text=f"Added {len(files)} file(s) from folder")
    
    def _clear_list(self):
        """Clear all files"""
        if self.image_files and messagebox.askyesno("Clear", "Remove all files?"):
            self.image_files = []
            self.extracted_texts = {}
            self._update_listbox()
            self.preview_text.delete('1.0', tk.END)
            self.status_label.config(text="List cleared")
    
    def _update_listbox(self):
        """Update file listbox"""
        self.file_listbox.delete(0, tk.END)
        for i, file in enumerate(self.image_files, 1):
            filename = os.path.basename(file)
            status = "✓ " if file in self.extracted_texts else ""
            self.file_listbox.insert(tk.END, f"{status}{i:2d}. {filename}")
    
    def _on_file_select(self, event):
        """Show extracted text when file is selected"""
        selection = self.file_listbox.curselection()
        if not selection:
            return
        
        idx = selection[0]
        if idx < len(self.image_files):
            file = self.image_files[idx]
            if file in self.extracted_texts:
                self.preview_text.delete('1.0', tk.END)
                self.preview_text.insert('1.0', self.extracted_texts[file])
    
    # === OCR Processing Methods ===
    
    def _encode_image(self, image_path):
        """Encode image to base64"""
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
    
    def _extract_text_from_image(self, image_path):
        """Use GPT-4 Vision to extract text from image"""
        if not self.client:
            return "[ERROR: OpenAI client not initialized. Check API key.]"
        
        try:
            base64_image = self._encode_image(image_path)
            instructions = self.instructions_text.get('1.0', 'end-1c').strip()
            
            response = self.client.chat.completions.create(
                model=self.model_var.get(),
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": instructions
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=4000
            )
            
            return response.choices[0].message.content
        
        except Exception as e:
            return f"[ERROR extracting text: {str(e)}]"
    
    def _process_selected(self):
        """Process currently selected image"""
        selection = self.file_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select an image to process")
            return
        
        idx = selection[0]
        if idx >= len(self.image_files):
            return
        
        file = self.image_files[idx]
        filename = os.path.basename(file)
        
        self.status_label.config(text=f"Processing {filename}...")
        if hasattr(self.parent_app, 'root'):
            self.parent_app.root.update()
        
        text = self._extract_text_from_image(file)
        self.extracted_texts[file] = text
        
        self.preview_text.delete('1.0', tk.END)
        self.preview_text.insert('1.0', text)
        
        self._update_listbox()
        self.status_label.config(text=f"✓ Processed {filename}")
    
    def _process_all(self):
        """Process all images in the list"""
        if not self.image_files:
            messagebox.showwarning("No Files", "Please add images first")
            return
        
        if not messagebox.askyesno("Process All", 
                                   f"Process all {len(self.image_files)} images?\n\n"
                                   "This will use API credits and may take several minutes."):
            return
        
        self.progress['maximum'] = len(self.image_files)
        self.progress['value'] = 0
        
        for i, file in enumerate(self.image_files, 1):
            filename = os.path.basename(file)
            self.status_label.config(text=f"Processing {i}/{len(self.image_files)}: {filename}...")
            if hasattr(self.parent_app, 'root'):
                self.parent_app.root.update()
            
            if file not in self.extracted_texts:
                text = self._extract_text_from_image(file)
                self.extracted_texts[file] = text
            
            self.progress['value'] = i
            self._update_listbox()
        
        self.status_label.config(text=f"✓ Processed all {len(self.image_files)} images!")
        messagebox.showinfo("Complete", 
                          f"Successfully processed {len(self.image_files)} images!\n\n"
                          "Click 'Save DOCX' to export the text.")
    
    # === Export Methods ===
    
    def _save_to_docx(self):
        """Save all extracted text to a Word document"""
        if not self.extracted_texts:
            messagebox.showwarning("No Text", "No extracted text to save.\n\n"
                                 "Process images first.")
            return
        
        output_file = filedialog.asksaveasfilename(
            title="Save Extracted Text",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx"), ("All files", "*.*")],
            initialfile="extracted_text.docx"
        )
        
        if not output_file:
            return
        
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading('Extracted Text from Images', 0)
            title.runs[0].font.size = Pt(16)
            
            # Add extracted text in order
            for i, file in enumerate(self.image_files, 1):
                if file in self.extracted_texts:
                    # Page header
                    heading = doc.add_heading(f'Page {i}: {os.path.basename(file)}', level=2)
                    heading.runs[0].font.size = Pt(12)
                    
                    # Text content
                    text = self.extracted_texts[file]
                    para = doc.add_paragraph(text)
                    para.paragraph_format.line_spacing = 1.15
                    para.paragraph_format.space_after = Pt(12)
                    
                    # Page break except for last
                    if i < len(self.image_files):
                        doc.add_page_break()
            
            doc.save(output_file)
            
            self.status_label.config(text=f"✓ Saved to {os.path.basename(output_file)}")
            
            if messagebox.askyesno("Success", 
                                  f"Document saved successfully!\n\n"
                                  f"{len(self.extracted_texts)} pages of text extracted\n\n"
                                  "Open the document now?"):
                os.startfile(output_file)
        
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save document:\n\n{str(e)}")
    
    def _copy_all_text(self):
        """Copy all extracted text to clipboard"""
        if not self.extracted_texts:
            messagebox.showwarning("No Text", "No extracted text to copy")
            return
        
        all_text = []
        for i, file in enumerate(self.image_files, 1):
            if file in self.extracted_texts:
                all_text.append(f"=== Page {i}: {os.path.basename(file)} ===\n")
                all_text.append(self.extracted_texts[file])
                all_text.append("\n\n")
        
        combined = "".join(all_text)
        
        # Get root window from parent app or use clipboard differently
        if hasattr(self.parent_app, 'root'):
            self.parent_app.root.clipboard_clear()
            self.parent_app.root.clipboard_append(combined)
        
        self.status_label.config(text=f"✓ Copied {len(self.extracted_texts)} pages to clipboard")
        messagebox.showinfo("Copied", f"Copied text from {len(self.extracted_texts)} pages to clipboard!")
