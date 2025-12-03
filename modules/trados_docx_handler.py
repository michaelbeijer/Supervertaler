"""
Trados Bilingual DOCX Handler (Review Files)

This module handles the import and export of Trados Studio bilingual review DOCX files.
Trados uses a table-based format with numbered inline tags.

Format Structure:
- Table with columns: Segment ID | Segment status | Source segment | Target segment
- Tags use character style "Tag" and format: <N>text</N>
- Segment IDs are GUIDs with numeric prefixes
- Statuses: "Not Translated", "Draft", "Translated", etc.

Critical for re-import:
- Tags MUST preserve the "Tag" character style
- Tag numbers must match between source and target
- Segment IDs must remain unchanged
"""

import os
import re
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from lxml import etree
from typing import List, Dict, Tuple, Optional
from copy import deepcopy


class TradosSegment:
    """
    Represents a Trados segment with tag information.
    """
    def __init__(self, segment_id: str, status: str, source_text: str, target_text: str = "",
                 source_runs: List[Dict] = None, row_index: int = 0):
        self.segment_id = segment_id
        self.status = status
        self.source_text = source_text  # Plain text with tags as text
        self.target_text = target_text
        self.source_runs = source_runs or []  # List of {text, is_tag, style_xml} dicts
        self.row_index = row_index
        
        # Extract tags from source for validation
        self.source_tags = self._extract_tags(source_text)
    
    def _extract_tags(self, text: str) -> List[str]:
        """Extract all tags from text."""
        pattern = r'</?(\d+)>'
        return re.findall(pattern, text)
    
    @property
    def plain_source(self) -> str:
        """Get source text without tags for translation."""
        return re.sub(r'</?(\d+)>', '', self.source_text)
    
    def __repr__(self):
        return f"TradosSegment(id={self.segment_id[:20]}..., status={self.status}, source={self.source_text[:40]}...)"


class TradosDOCXHandler:
    """
    Handler for Trados Studio bilingual review DOCX files.
    
    This class provides methods to:
    - Load and parse Trados bilingual review DOCX files
    - Extract source segments with tag markers
    - Update target segments with translations (preserving tag style)
    - Save modified files ready for re-import to Trados
    """
    
    # Trados tag pattern: <N> or </N> where N is a number
    TAG_PATTERN = re.compile(r'(</?(\d+)>)')
    
    def __init__(self):
        self.doc = None
        self.table = None
        self.segments: List[TradosSegment] = []
        self.file_path = None
        self.header_row = None
        self.tag_style_xml = None  # Store the Tag style XML for reuse
        
    def load(self, file_path: str) -> bool:
        """
        Load a Trados bilingual review DOCX file.
        
        Args:
            file_path: Path to the Trados bilingual DOCX file
            
        Returns:
            bool: True if loaded successfully, False otherwise
        """
        try:
            self.file_path = file_path
            self.doc = Document(file_path)
            
            # Trados bilingual files should have exactly one table
            if len(self.doc.tables) == 0:
                print(f"ERROR: No table found in {file_path}")
                return False
                
            self.table = self.doc.tables[0]
            
            # Verify the header row
            if len(self.table.rows) < 2:
                print(f"ERROR: Table has insufficient rows")
                return False
                
            self.header_row = [cell.text.strip() for cell in self.table.rows[0].cells]
            
            # Check if this looks like a Trados bilingual DOCX
            expected_headers = ['Segment ID', 'Segment status', 'Source segment', 'Target segment']
            if self.header_row != expected_headers:
                print(f"WARNING: Headers don't match expected Trados format")
                print(f"  Expected: {expected_headers}")
                print(f"  Found: {self.header_row}")
                # Continue anyway if it's close enough
                if 'Segment' not in self.header_row[0]:
                    return False
            
            # Find and store the Tag style XML for later use
            self._capture_tag_style()
            
            print(f"Successfully loaded Trados bilingual DOCX: {file_path}")
            print(f"Header: {self.header_row}")
            print(f"Total rows (including header): {len(self.table.rows)}")
            
            return True
            
        except Exception as e:
            print(f"ERROR loading Trados DOCX: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _capture_tag_style(self):
        """Find and capture the Tag style XML from the document."""
        try:
            # Look through the document for a run with Tag style
            for row in self.table.rows[1:]:
                source_cell = row.cells[2]
                for para in source_cell.paragraphs:
                    for run in para.runs:
                        rPr = run._r.find(qn('w:rPr'))
                        if rPr is not None:
                            style_elem = rPr.find(qn('w:rStyle'))
                            if style_elem is not None and style_elem.get(qn('w:val')) == 'Tag':
                                # Found a Tag style - save the entire rPr as template
                                self.tag_style_xml = deepcopy(rPr)
                                print("Captured Tag style from document")
                                return
        except Exception as e:
            print(f"Warning: Could not capture Tag style: {e}")
    
    def extract_source_segments(self) -> List[TradosSegment]:
        """
        Extract all source segments from the Trados bilingual DOCX.
        
        Returns:
            list: List of TradosSegment objects
        """
        self.segments = []
        
        if not self.table:
            print("ERROR: No table loaded")
            return []
        
        # Skip header row (index 0), process data rows
        for i, row in enumerate(self.table.rows[1:], start=1):
            try:
                cells = row.cells
                
                # Extract data from columns
                segment_id = cells[0].text.strip()
                status = cells[1].text.strip()
                source_cell = cells[2]
                target_cell = cells[3] if len(cells) > 3 else None
                
                # Get source text
                source_text = source_cell.text.strip()
                target_text = target_cell.text.strip() if target_cell else ""
                
                # Extract run information for preserving tag styles
                source_runs = self._extract_runs_with_styles(source_cell)
                
                # Create TradosSegment
                segment = TradosSegment(
                    segment_id=segment_id,
                    status=status,
                    source_text=source_text,
                    target_text=target_text,
                    source_runs=source_runs,
                    row_index=i
                )
                
                self.segments.append(segment)
                
            except Exception as e:
                print(f"WARNING: Error processing row {i}: {e}")
                continue
        
        print(f"Extracted {len(self.segments)} segments from Trados DOCX")
        return self.segments
    
    def _extract_runs_with_styles(self, cell) -> List[Dict]:
        """
        Extract runs from a cell, noting which are tags.
        
        Returns:
            List of dicts with: {text, is_tag, style_xml}
        """
        runs = []
        for para in cell.paragraphs:
            for run in para.runs:
                is_tag = False
                style_xml = None
                
                # Check if this run has Tag style
                rPr = run._r.find(qn('w:rPr'))
                if rPr is not None:
                    style_elem = rPr.find(qn('w:rStyle'))
                    if style_elem is not None and style_elem.get(qn('w:val')) == 'Tag':
                        is_tag = True
                        style_xml = deepcopy(rPr)
                
                runs.append({
                    'text': run.text,
                    'is_tag': is_tag,
                    'style_xml': style_xml
                })
        
        return runs
    
    def update_target_segments(self, translations: Dict[int, str]) -> int:
        """
        Update target segments with translations.
        
        Args:
            translations: Dict mapping row index to translated text
            
        Returns:
            int: Number of segments updated
        """
        updated_count = 0
        
        for idx, translation in translations.items():
            if 0 < idx < len(self.table.rows):
                row = self.table.rows[idx]
                target_cell = row.cells[3]
                
                # Get the source segment for tag info
                source_cell = row.cells[2]
                
                # Clear existing target content
                for para in target_cell.paragraphs:
                    for run in list(para.runs):
                        run._r.getparent().remove(run._r)
                
                # Write target with proper tag styling
                self._write_text_with_tags(target_cell, translation, source_cell)
                
                # Update status to indicate translation
                status_cell = row.cells[1]
                if status_cell.text.strip() == "Not Translated":
                    self._set_cell_text(status_cell, "Translated")
                
                updated_count += 1
                
        print(f"Updated {updated_count} target segments")
        return updated_count
    
    def _write_text_with_tags(self, target_cell, text: str, source_cell):
        """
        Write text to target cell, applying Tag style to tag patterns.
        
        This ensures tags in the target have the same style as in the source,
        which is critical for re-import into Trados.
        """
        if not target_cell.paragraphs:
            return
            
        para = target_cell.paragraphs[0]
        
        # Split text by tags
        parts = self.TAG_PATTERN.split(text)
        
        # parts will be: [text, full_tag, tag_num, text, full_tag, tag_num, ...]
        i = 0
        while i < len(parts):
            part = parts[i]
            
            if self.TAG_PATTERN.match(part):
                # This is a tag - apply Tag style
                run = para.add_run(part)
                self._apply_tag_style(run)
                i += 2  # Skip the captured group (tag number)
            elif part:
                # Regular text
                run = para.add_run(part)
                # Copy language settings from source if available
                self._apply_default_style(run)
            else:
                i += 1
                continue
            
            i += 1
    
    def _apply_tag_style(self, run):
        """Apply the Tag character style to a run."""
        # Create rPr element if needed
        rPr = run._r.find(qn('w:rPr'))
        if rPr is None:
            rPr = etree.SubElement(run._r, qn('w:rPr'))
            run._r.insert(0, rPr)
        
        # Add rStyle element with Tag value
        style_elem = rPr.find(qn('w:rStyle'))
        if style_elem is None:
            style_elem = etree.SubElement(rPr, qn('w:rStyle'))
        style_elem.set(qn('w:val'), 'Tag')
    
    def _apply_default_style(self, run):
        """Apply default style (language settings) to a run."""
        # Copy language from tag_style_xml if available
        if self.tag_style_xml is not None:
            lang_elem = self.tag_style_xml.find(qn('w:lang'))
            if lang_elem is not None:
                rPr = run._r.find(qn('w:rPr'))
                if rPr is None:
                    rPr = etree.SubElement(run._r, qn('w:rPr'))
                    run._r.insert(0, rPr)
                
                new_lang = deepcopy(lang_elem)
                rPr.append(new_lang)
    
    def _set_cell_text(self, cell, text: str):
        """Set cell text, preserving formatting."""
        if cell.paragraphs:
            para = cell.paragraphs[0]
            # Clear existing runs
            for run in list(para.runs):
                run._r.getparent().remove(run._r)
            para.add_run(text)
    
    def save(self, output_path: str = None) -> bool:
        """
        Save the modified document.
        
        Args:
            output_path: Path to save to (defaults to original path)
            
        Returns:
            bool: True if saved successfully
        """
        try:
            save_path = output_path or self.file_path
            self.doc.save(save_path)
            print(f"Saved Trados bilingual DOCX: {save_path}")
            return True
        except Exception as e:
            print(f"ERROR saving Trados DOCX: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def get_segments_for_translation(self) -> List[Tuple[int, str, str]]:
        """
        Get segments that need translation.
        
        Returns:
            List of (row_index, source_text, plain_source) tuples
        """
        result = []
        for seg in self.segments:
            if seg.status == "Not Translated" or not seg.target_text:
                result.append((seg.row_index, seg.source_text, seg.plain_source))
        return result


def detect_bilingual_docx_type(file_path: str) -> str:
    """
    Detect the type of bilingual DOCX file.
    
    Returns:
        str: "trados", "cafetran", "memoq", or "unknown"
    """
    try:
        doc = Document(file_path)
        
        if len(doc.tables) == 0:
            return "unknown"
        
        table = doc.tables[0]
        if len(table.rows) < 1:
            return "unknown"
        
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        
        # Trados: Segment ID | Segment status | Source segment | Target segment
        if headers and headers[0] == "Segment ID" and "Segment status" in headers:
            return "trados"
        
        # CafeTran: ID | filename | filename | Notes | *
        if headers and headers[0] == "ID":
            return "cafetran"
        
        # memoQ: Usually has different structure
        # TODO: Add memoQ detection
        
        return "unknown"
        
    except Exception as e:
        print(f"Error detecting bilingual DOCX type: {e}")
        return "unknown"
