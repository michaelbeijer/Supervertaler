"""Analyze Trados bilingual review file structure"""
from docx import Document
from docx.oxml.ns import qn

doc = Document('user_data_private/Projects/DAT FRANCE/SPV agreement NL_page3.docx.review.docx')
table = doc.tables[0]

print("=== TABLE STRUCTURE ===")
print(f"Rows: {len(table.rows)}, Columns: {len(table.columns)}")
print()

# Look at header
print("=== HEADER ROW ===")
for j, cell in enumerate(table.rows[0].cells):
    print(f"  Col {j}: {cell.text}")
print()

# Look at rows with content
print("=== SAMPLE ROWS WITH TAGS ===")
for i in range(1, min(8, len(table.rows))):
    row = table.rows[i]
    print(f"--- Row {i} ---")
    print(f"  Segment ID: {row.cells[0].text[:40]}...")
    print(f"  Status: {row.cells[1].text}")
    
    # Analyze source cell formatting
    source_cell = row.cells[2]
    print(f"  Source text: {source_cell.text[:60]}...")
    
    # Check runs for formatting
    for para in source_cell.paragraphs:
        for run in para.runs:
            text = run.text
            if '<' in text or '>' in text or run.font.color.rgb:
                color = str(run.font.color.rgb) if run.font.color.rgb else "default"
                bold = run.font.bold
                italic = run.font.italic
                print(f"    TAG RUN: '{text[:40]}' color={color} bold={bold} italic={italic}")
    
    # Check target cell
    target_cell = row.cells[3]
    if target_cell.text.strip():
        print(f"  Target text: {target_cell.text[:60]}...")
    print()

# Look for any existing translated segments
print("=== SEGMENTS WITH TRANSLATIONS ===")
for i, row in enumerate(table.rows[1:], 1):
    target = row.cells[3].text.strip()
    if target:
        print(f"Row {i}: Status='{row.cells[1].text}' Target='{target[:50]}...'")
