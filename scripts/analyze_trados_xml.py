"""Deeper analysis of Trados tag formatting"""
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

doc = Document('user_data_private/Projects/DAT FRANCE/SPV agreement NL_page3.docx.review.docx')
table = doc.tables[0]

# Row 5 has tags - let's look at the XML
row = table.rows[5]
source_cell = row.cells[2]

print("=== RAW XML OF SOURCE CELL (Row 5) ===")
print(etree.tostring(source_cell._tc, pretty_print=True, encoding='unicode')[:3000])
print()

# Check each run's XML
print("=== RUN DETAILS ===")
for para in source_cell.paragraphs:
    for run in para.runs:
        if '<' in run.text or '>' in run.text:
            print(f"Tag text: '{run.text}'")
            # Check the run's XML for color
            rPr = run._r.find(qn('w:rPr'))
            if rPr is not None:
                print("  Run properties XML:")
                print("  " + etree.tostring(rPr, pretty_print=True, encoding='unicode').replace('\n', '\n  '))
            else:
                print("  No run properties (default formatting)")
            print()
